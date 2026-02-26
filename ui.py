import tkinter as tk
from banka_tab import BankaTab
from giris_tab import GirisTab
from tkinter import ttk, messagebox, filedialog
from datetime import date
from typing import Optional, Dict, List
import threading
import time
import os
import sqlite3
from pathlib import Path

DB_PATH = Path(__file__).with_name("erp.db")

# .env dosyasından API key yükle
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

# .env veya env dosyasından API key yükle
if not GEMINI_API_KEY:
    for env_file in [".env", "env"]:
        env_path = Path(__file__).parent / env_file
        if env_path.exists():
            try:
                content = env_path.read_text(encoding="utf-8")
                for line in content.splitlines():
                    line = line.strip()
                    if line.startswith("GEMINI_API_KEY="):
                        GEMINI_API_KEY = line.split("=", 1)[1].strip()
                        break
            except Exception:
                pass
        if GEMINI_API_KEY:
            break

# Gemini
try:
    from google import genai
    if GEMINI_API_KEY:
        _gemini_client = genai.Client(api_key=GEMINI_API_KEY)
    else:
        _gemini_client = None
    GEMINI_AVAILABLE = bool(GEMINI_API_KEY)
except ImportError:
    _gemini_client = None
    GEMINI_AVAILABLE = False

# ttkbootstrap (varsa kullan, yoksa normal ttk)
try:
    import ttkbootstrap as ttkb
    from ttkbootstrap.constants import *
    BOOTSTRAP = True
except ImportError:
    BOOTSTRAP = False

from database import (
    get_all_customers_with_rent_progression,
    get_all_products,
    get_all_invoices,
    get_tufe_for_year,
    save_tufe_for_year,
    import_tufe_from_excel,
    import_customers_from_excel,
    insert_customer,
    update_customer,
    delete_customer,
    insert_product,
    update_product,
    delete_product,
    insert_invoice,
    delete_invoice,
    calculate_rent_progression,
    get_months_list,
    initialize_database,
    get_connection,
    MONTHS_TR,
    save_rent_payment,
    get_rent_payments_for_customer,
    get_rent_payments_for_year,
    get_yearly_totals_for_customer,
    get_all_offices,
    get_office_summary,
    save_office,
    delete_office,
    assign_office_to_customer,
    get_next_office_code,
    get_offices_for_customer,
    initialize_offices,
    insert_tahsilat,
    delete_tahsilat,
    get_tahsilatlar,
    get_musteri_toplam_borc,
    get_musteri_bu_ay_borc,
    get_tahsilat_toplam,
)


# ============================================================================
# ANA PENCERE
# ============================================================================

class BaseWindow(tk.Tk):
    """ERP ana penceresi."""

    def __init__(self) -> None:
        super().__init__()
        self.title("BestOffice ERP — Kira Takip Sistemi")
        self.geometry("1600x900")
        self.minsize(1200, 700)

        # ttkbootstrap tema
        if BOOTSTRAP:
            self.style = ttkb.Style(theme="superhero")  # mavi/cyan tema
        else:
            self.style = ttk.Style(self)
            self._setup_theme()

        self._create_menu()

        # Üst başlık
        header = tk.Frame(self, bg="#0f2537", pady=10)
        header.pack(side=tk.TOP, fill=tk.X)
        tk.Label(
            header,
            text="BestOffice ERP — Kira Takip Sistemi",
            font=("Segoe UI", 16, "bold"),
            bg="#0f2537", fg="#4fc3f7"
        ).pack(side=tk.LEFT, padx=16)

        # Gemini AI butonu (sağ üst)
        ai_btn_text = "✨ Gemini AI Analiz" if GEMINI_AVAILABLE else "✨ Gemini AI (Key Gerekli)"
        ai_color = "#f59e0b" if GEMINI_AVAILABLE else "#6b7280"
        self.ai_btn = tk.Button(
            header, text=ai_btn_text,
            bg=ai_color, fg="white",
            font=("Segoe UI", 10, "bold"),
            relief="flat", padx=12, pady=4, cursor="hand2",
            command=self._open_gemini_panel
        )
        self.ai_btn.pack(side=tk.RIGHT, padx=(0, 10))

        # Sekmeler
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))

        self.giris_tab     = GirisTab(self.notebook)
        self.customer_tab  = CustomerTab(self.notebook)
        self.office_tab    = OfficeTab(self.notebook)
        self.tahsilat_tab  = TahsilatTab(self.notebook)
        self.personel_tab  = PersonelTab(self.notebook)
        self.product_tab   = ProductTab(self.notebook)
        self.invoice_tab   = InvoiceTab(self.notebook)
        self.kargolar_tab  = KargolarTab(self.notebook)
        self.banka_tab     = BankaTab(self.notebook)
        self.tufe_tab      = TufeTab(self.notebook)

        self.notebook.add(self.customer_tab,  text="  Müşteriler  ")
        self.notebook.add(self.office_tab,    text="  🏢 Ofisler  ")
        self.notebook.add(self.tahsilat_tab,  text="  💰 Tahsilat  ")
        self.notebook.add(self.personel_tab,  text="  👤 Personel  ")
        self.notebook.add(self.product_tab,   text="  Ürünler  ")
        self.notebook.add(self.invoice_tab,   text="  Faturalar  ")
        self.notebook.add(self.kargolar_tab,  text="  Kargolar  ")
        self.notebook.add(self.banka_tab,     text="  🏦 Bankalar  ")
        self.notebook.add(self.tufe_tab,      text="  TÜFE  ")
        self.notebook.add(self.giris_tab,     text="  Giriş  ")

        # Global mousewheel scroll — tüm canvas'lara bağla
        self.bind_all("<MouseWheel>", self._global_scroll)

    def _global_scroll(self, event):
        """Fare altındaki canvas'ı scroll et."""
        widget = event.widget
        # Üst canvas'ı bul
        while widget:
            if isinstance(widget, tk.Canvas):
                try:
                    widget.yview_scroll(int(-event.delta / 120), "units")
                except:
                    pass
                return
            try:
                widget = widget.master
            except:
                break


    def _open_gemini_panel(self) -> None:
        """Gemini AI chat/analiz panelini aç."""
        if not GEMINI_AVAILABLE:
            messagebox.showwarning(
                "API Key Gerekli",
                "Gemini API key bulunamadı!\n\n"
                "Proje klasöründeki .env dosyasına şunu ekleyin:\n"
                "GEMINI_API_KEY=your_api_key_here\n\n"
                "Key almak için: aistudio.google.com"
            )
            return
        GeminiPanel(self)

    def _setup_theme(self) -> None:
        bg_dark   = "#1e2a3a"   # Ana arka plan
        bg_mid    = "#243447"   # Panel arka planı
        bg_light  = "#2d4060"   # Açık panel
        accent    = "#00bcd4"   # Cyan vurgu
        accent2   = "#0097a7"   # Koyu cyan
        fg_main   = "#e0f7fa"   # Ana yazı
        fg_sub    = "#b0bec5"   # İkincil yazı
        header_bg = "#162230"   # Başlık arka planı

        self.configure(bg=bg_dark)

        style = self.style
        style.theme_use("clam")

        style.configure("TFrame",            background=bg_dark)
        style.configure("TLabelframe",       background=bg_mid,   relief="flat")
        style.configure("TLabelframe.Label", background=bg_mid,   foreground=accent,
                                             font=("Segoe UI", 9, "bold"))
        style.configure("TLabel",            background=bg_dark,  foreground=fg_main,
                                             font=("Segoe UI", 9))
        style.configure("TButton",           background=accent2,  foreground="white",
                                             font=("Segoe UI", 9, "bold"), padding=(8, 4),
                                             relief="flat")
        style.map("TButton",
                  background=[("active", accent), ("pressed", "#006064")])
        style.configure("TEntry",            fieldbackground=bg_light, foreground=fg_main,
                                             insertcolor=fg_main, padding=4)
        style.configure("TCombobox",         fieldbackground=bg_light, foreground=fg_main,
                                             background=bg_light, selectbackground=accent2)
        style.map("TCombobox",
                  fieldbackground=[("readonly", bg_light)],
                  foreground=[("readonly", fg_main)])
        style.configure("TNotebook",         background=header_bg, tabmargins=[2, 4, 2, 0])
        style.configure("TNotebook.Tab",     background=bg_mid,   foreground=fg_sub,
                                             font=("Segoe UI", 9), padding=(10, 5))
        style.map("TNotebook.Tab",
                  background=[("selected", accent2)],
                  foreground=[("selected", "white")])
        style.configure("Treeview",          background=bg_mid,   foreground=fg_main,
                                             fieldbackground=bg_mid, rowheight=26,
                                             font=("Segoe UI", 9))
        style.configure("Treeview.Heading",  background=accent2,  foreground="white",
                                             font=("Segoe UI", 9, "bold"), relief="flat")
        style.map("Treeview",
                  background=[("selected", accent)],
                  foreground=[("selected", "white")])
        style.configure("TScrollbar",        background=bg_light, troughcolor=bg_dark,
                                             arrowcolor=accent)
        style.configure("TCheckbutton",      background=bg_dark,  foreground=fg_main,
                                             font=("Segoe UI", 9))
        style.map("TCheckbutton",
                  background=[("active", bg_dark)])
        style.configure("TSeparator",        background=accent2)

        # Canvas ve toplam satırı rengi
        self._theme_colors = {
            "bg_dark": bg_dark, "bg_mid": bg_mid, "bg_light": bg_light,
            "accent": accent, "accent2": accent2, "fg_main": fg_main,
            "header_bg": header_bg
        }

    def _create_menu(self) -> None:
        menubar = tk.Menu(self)

        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Tüm Verileri Yenile", command=self._reload_data)
        file_menu.add_separator()
        file_menu.add_command(label="Çıkış", command=self.destroy)
        menubar.add_cascade(label="Dosya", menu=file_menu)

        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(
            label="Hakkında",
            command=lambda: messagebox.showinfo(
                "Hakkında", "BestOffice ERP - Kira Takip Sistemi\nv1.1"
            ),
        )
        menubar.add_cascade(label="Yardım", menu=help_menu)
        self.config(menu=menubar)

    def _reload_data(self) -> None:
        self.customer_tab.refresh()
        self.product_tab.refresh()
        self.invoice_tab.refresh()
        self.tufe_tab.refresh()          # FIX: tufe_tab da yenileniyordu eksikti
        messagebox.showinfo("Bilgi", "Veriler yenilendi.")


# ============================================================================
# YARDIMCI BILEŞEN: FormField
# ============================================================================

class FormField(ttk.Frame):
    """Etiket + Giriş alanı bileşeni."""

    def __init__(
        self,
        master: tk.Widget,
        label: str,
        width: int = 28,
        field_type: str = "entry",
        options: List[str] = None,
        **kwargs,
    ) -> None:
        super().__init__(master, **kwargs)

        ttk.Label(self, text=label, width=20, anchor="w").pack(side=tk.LEFT, padx=(0, 6))
        self.var = tk.StringVar()

        if field_type == "combobox" and options:
            self.entry = ttk.Combobox(
                self, textvariable=self.var, values=options, width=width, state="readonly"
            )
        else:
            self.entry = ttk.Entry(self, textvariable=self.var, width=width)

        self.entry.pack(side=tk.LEFT, fill=tk.X, expand=True)

    def get(self) -> str:
        return self.var.get().strip()

    def set(self, value) -> None:
        self.var.set(str(value) if value is not None else "")

    def clear(self) -> None:
        self.var.set("")



# ============================================================================
# GEMİNİ AI PANELİ
# ============================================================================

class GeminiPanel(tk.Toplevel):
    """Gemini AI Chat ve Analiz Penceresi."""

    def __init__(self, master) -> None:
        super().__init__(master)
        self.title("✨ Gemini AI Asistan")
        self.geometry("700x600")
        self.resizable(True, True)
        self.configure(bg="#1e1e2e")

        # Başlık
        hdr = tk.Frame(self, bg="#2a2a3e", pady=10)
        hdr.pack(fill=tk.X)
        tk.Label(hdr, text="✨ Gemini AI Kira Asistanı",
                 font=("Segoe UI", 14, "bold"),
                 bg="#2a2a3e", fg="#f59e0b").pack(side=tk.LEFT, padx=16)

        # Hızlı analiz butonları
        btn_frame = tk.Frame(self, bg="#1e1e2e", pady=8)
        btn_frame.pack(fill=tk.X, padx=12)
        tk.Label(btn_frame, text="Hızlı Analiz:", bg="#1e1e2e",
                 fg="#a0a0c0", font=("Segoe UI", 9)).pack(side=tk.LEFT)

        quick_btns = [
            ("📊 Kira Özeti",      "Tüm müşterilerin kira durumunu analiz et ve özet çıkar"),
            ("📈 TÜFE Analizi",    "TÜFE oranlarına göre kira artışlarını değerlendir"),
            ("⚠️ Riskli Müşteri", "Kira artışı çok yüksek olan riskli müşterileri belirle"),
        ]
        for label, prompt in quick_btns:
            tk.Button(btn_frame, text=label, bg="#3b3b5c", fg="white",
                      font=("Segoe UI", 8), relief="flat", padx=8, pady=3,
                      cursor="hand2",
                      command=lambda p=prompt: self._send(p)
                      ).pack(side=tk.LEFT, padx=(6, 0))

        # Chat alanı
        chat_frame = tk.Frame(self, bg="#1e1e2e")
        chat_frame.pack(fill=tk.BOTH, expand=True, padx=12, pady=(0, 8))

        self.chat_text = tk.Text(
            chat_frame, wrap=tk.WORD, state="disabled",
            bg="#12121e", fg="#e0e0f0", font=("Segoe UI", 10),
            relief="flat", padx=12, pady=8,
            insertbackground="white"
        )
        chat_scroll = ttk.Scrollbar(chat_frame, command=self.chat_text.yview)
        self.chat_text.configure(yscrollcommand=chat_scroll.set)
        chat_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.chat_text.pack(fill=tk.BOTH, expand=True)

        # Tag renkleri
        self.chat_text.tag_configure("user",  foreground="#60a5fa", font=("Segoe UI", 10, "bold"))
        self.chat_text.tag_configure("ai",    foreground="#34d399", font=("Segoe UI", 10))
        self.chat_text.tag_configure("error", foreground="#f87171", font=("Segoe UI", 10))
        self.chat_text.tag_configure("info",  foreground="#a0a0c0", font=("Segoe UI", 9, "italic"))

        # Giriş alanı
        input_frame = tk.Frame(self, bg="#1e1e2e", pady=8)
        input_frame.pack(fill=tk.X, padx=12)

        self.input_var = tk.StringVar()
        self.input_entry = tk.Entry(
            input_frame, textvariable=self.input_var,
            bg="#2a2a3e", fg="white", insertbackground="white",
            font=("Segoe UI", 10), relief="flat"
        )
        self.input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=8, padx=(0, 8))
        self.input_entry.bind("<Return>", lambda e: self._send())

        self.send_btn = tk.Button(
            input_frame, text="Gönder ➤",
            bg="#f59e0b", fg="white", font=("Segoe UI", 10, "bold"),
            relief="flat", padx=16, pady=6, cursor="hand2",
            command=self._send
        )
        self.send_btn.pack(side=tk.RIGHT)

        # Hoşgeldin mesajı
        self._append("info", "🤖 Gemini AI hazır! Soru sorun veya hızlı analiz butonlarını kullanın.\n\n")

    def _append(self, tag: str, text: str) -> None:
        self.chat_text.configure(state="normal")
        self.chat_text.insert(tk.END, text, tag)
        self.chat_text.see(tk.END)
        self.chat_text.configure(state="disabled")

    def _send(self, prompt: str = "") -> None:
        if not prompt:
            prompt = self.input_var.get().strip()
        if not prompt:
            return

        self.input_var.set("")
        self._append("user", f"👤 Sen: {prompt}\n\n")
        self._append("info", "⏳ Gemini düşünüyor...\n")
        self.send_btn.configure(state="disabled")

        # DB'den müşteri verilerini çek — AI'ye bağlam ver
        from database import fetch_all
        customers = fetch_all("SELECT name, ilk_kira_bedeli, current_rent, rent_start_date FROM customers LIMIT 20")
        context = "Mevcut müşteri verileri (ilk 20):\n"
        for c in customers:
            context += f"- {c['name']}: İlk kira={c['ilk_kira_bedeli']}₺, Güncel={c['current_rent']}₺, Başlangıç={c['rent_start_date']}\n"

        full_prompt = f"""Sen bir Türk ERP sisteminin kira takip asistanısın. 
Aşağıdaki müşteri verilerine göre soruyu yanıtla.

{context}

Soru: {prompt}

Yanıtını Türkçe ver, kısa ve net ol."""

        def call_gemini():
            try:
                from google import genai
                client = genai.Client(api_key=GEMINI_API_KEY)
                response = client.models.generate_content(
                    model="gemini-2.0-flash-lite",
                    contents=full_prompt
                )
                answer = response.text
                self.after(0, lambda: self._append("ai", f"🤖 Gemini: {answer}\n\n"))
            except Exception as e:
                err_msg = str(e)
                if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg:
                    short = "⏳ Kota doldu, lütfen 1-2 dakika bekleyip tekrar deneyin."
                elif "404" in err_msg or "NOT_FOUND" in err_msg:
                    short = "❌ Model bulunamadı. API key'i kontrol edin."
                elif "403" in err_msg or "API_KEY" in err_msg:
                    short = "🔑 API key geçersiz. .env dosyasını kontrol edin."
                else:
                    short = f"❌ Hata: {err_msg[:120]}..."
                self.after(0, lambda m=short: self._append("error", f"{m}\n\n"))
            finally:
                self.after(0, lambda: self.send_btn.configure(state="normal"))

        threading.Thread(target=call_gemini, daemon=True).start()


# ============================================================================
# MÜŞTERİ SEKMESİ
# ============================================================================

class CustomerTab(ttk.Frame):
    """Müşteri ve Kira Takip Sekmesi — Dinamik Yıl Sütunları."""

    def __init__(self, master: tk.Widget) -> None:
        super().__init__(master, padding=10)
        self.selected_id: Optional[int] = None
        self._sort_col: Optional[str] = None
        self._sort_reverse: bool = False

        # Grid layout: row 0 = sabit üst bar, row 1 = scroll'lu içerik
        self.rowconfigure(0, weight=0)
        self.rowconfigure(1, weight=1)
        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=0)

        # Scroll canvas (row 1)
        self._vsb = ttk.Scrollbar(self, orient="vertical")
        self._canvas = tk.Canvas(self, highlightthickness=0,
                                  yscrollcommand=self._vsb.set)
        self._vsb.configure(command=self._canvas.yview)
        self._canvas.grid(row=1, column=0, sticky="nsew")
        self._vsb.grid(row=1, column=1, sticky="ns")

        self._inner = ttk.Frame(self._canvas)
        self._inner_id = self._canvas.create_window((0, 0), window=self._inner, anchor="nw")

        self._inner.bind("<Configure>", self._on_inner_configure)
        self._canvas.bind("<Configure>", self._on_canvas_configure)
        self._canvas.bind("<MouseWheel>", self._on_mousewheel)
        self._inner.bind("<MouseWheel>", self._on_mousewheel)
        self.bind("<MouseWheel>", self._on_mousewheel)

        self._build_ui_top()
        self._build_ui_body()
        self._init_panels()
        self.refresh()

    def _on_inner_configure(self, e=None):
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    def _on_canvas_configure(self, e=None):
        self._canvas.itemconfig(self._inner_id, width=e.width)

    def _on_mousewheel(self, e):
        self._canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")

    def _build_ui_top(self) -> None:
        # ── Üst buton çubuğu — row 0, sabit ──
        top = ttk.Frame(self)
        top.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 6))

        ttk.Label(top, text="Müşteriler", font=("Segoe UI", 13, "bold")).pack(side=tk.LEFT)
        ttk.Button(top, text="Excel'den İçeri Aktar", command=self._on_import_excel).pack(side=tk.LEFT, padx=(10, 0))
        ttk.Button(top, text="↺ Yenile", command=self.refresh).pack(side=tk.LEFT, padx=(4, 0))

        # Arama kutusu
        ttk.Label(top, text="  🔍").pack(side=tk.LEFT, padx=(16, 2))
        self.filter_arama_var = tk.StringVar()
        self.filter_arama_var.trace_add("write", lambda *a: self._apply_filter())
        arama_entry = ttk.Entry(top, textvariable=self.filter_arama_var, width=18)
        arama_entry.pack(side=tk.LEFT)

        ttk.Label(top, text="  Ay:").pack(side=tk.LEFT, padx=(16, 2))
        self.filter_ay_var = tk.StringVar(value="Tümü")
        self.filter_ay = ttk.Combobox(top, textvariable=self.filter_ay_var,
                                       values=["Tümü"] + MONTHS_TR, width=10, state="readonly")
        self.filter_ay.pack(side=tk.LEFT)
        self.filter_ay.bind("<<ComboboxSelected>>", lambda e: self._apply_filter())

        ttk.Label(top, text="  Yıl:").pack(side=tk.LEFT, padx=(8, 2))
        self.filter_yil_var = tk.StringVar(value="Tümü")
        self.filter_yil = ttk.Combobox(top, textvariable=self.filter_yil_var,
                                        values=["Tümü"], width=8, state="readonly")
        self.filter_yil.pack(side=tk.LEFT)
        self.filter_yil.bind("<<ComboboxSelected>>", lambda e: self._apply_filter())

        ttk.Button(top, text="✖ Filtre Temizle", command=self._clear_filter).pack(side=tk.LEFT, padx=(6, 0))

        # Görünüm: Yıl bazlı veya Ay bazlı
        ttk.Label(top, text="  📅 Görünüm:").pack(side=tk.LEFT, padx=(16, 2))
        self.gorunum_yil_var = tk.StringVar(value="2026")
        yillar = ["Tüm Yıllar"] + [str(y) for y in range(2021, 2028)]
        self.gorunum_yil = ttk.Combobox(top, textvariable=self.gorunum_yil_var,
                                         values=yillar, width=10, state="readonly")
        self.gorunum_yil.pack(side=tk.LEFT)
        self.gorunum_yil.bind("<<ComboboxSelected>>", lambda e: self._on_gorunum_degis())

        self.gorunum_ay_var = tk.StringVar(value="Tüm Aylar")
        aylar = ["Tüm Aylar"] + MONTHS_TR
        self.gorunum_ay = ttk.Combobox(top, textvariable=self.gorunum_ay_var,
                                        values=aylar, width=10, state="readonly")
        self.gorunum_ay.pack(side=tk.LEFT, padx=(4, 0))
        self.gorunum_ay.bind("<<ComboboxSelected>>", lambda e: self._on_gorunum_degis())

        # Ödenmeyen ay filtresi
        ttk.Label(top, text="  ⚠️ Ödenmeyen:").pack(side=tk.LEFT, padx=(16, 2))
        self.filter_odenmemis_var = tk.StringVar(value="Tümü")
        self.filter_odenmemis = ttk.Combobox(
            top, textvariable=self.filter_odenmemis_var,
            values=["Tümü", "1+ ay", "2+ ay", "3+ ay", "6+ ay", "12+ ay"],
            width=8, state="readonly"
        )
        self.filter_odenmemis.pack(side=tk.LEFT)
        self.filter_odenmemis.bind("<<ComboboxSelected>>", lambda e: self._apply_filter())

        self._show_form = tk.BooleanVar(value=True)
        self._show_rent = tk.BooleanVar(value=True)

        ttk.Checkbutton(top, text="👤 Müşteri Detayları",
                        variable=self._show_form,
                        command=self._toggle_panels).pack(side=tk.RIGHT, padx=(4, 0))
        ttk.Checkbutton(top, text="💰 Aylık Kira Girişi",
                        variable=self._show_rent,
                        command=self._toggle_panels).pack(side=tk.RIGHT, padx=(4, 0))

    def _build_ui_body(self) -> None:
        # ── Ana alan ──
        self.main_frame = ttk.Frame(self._inner)
        self.main_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # ── Treeview ──
        tree_frame = ttk.Frame(self.main_frame)
        tree_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 0))

        self.tree = ttk.Treeview(tree_frame, show="headings", selectmode="browse", height=15)
        vsb = ttk.Scrollbar(tree_frame, orient="vertical",   command=self.tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        self.tree.bind("<<TreeviewSelect>>", self._on_row_select)
        self.tree.bind("<Double-1>", self._on_musteri_detay)

        # Toplam satırı
        self.totals_frame = ttk.Frame(self.main_frame)
        self.totals_frame.pack(side=tk.TOP, fill=tk.X, pady=(2, 0))
        self.totals_canvas = tk.Canvas(self.totals_frame, height=26, bg="#0097a7", highlightthickness=0)
        self.totals_canvas.pack(fill=tk.X)

        self.tree.bind("<Configure>",       lambda e: self.after(50, self._update_totals))
        self.tree.bind("<ButtonRelease-1>", lambda e: self.after(50, self._update_totals))
        orig_hsb_set = hsb.set
        def hsb_set_and_update(*args):
            orig_hsb_set(*args)
            self.after(10, self._update_totals)
        self.tree.configure(xscrollcommand=hsb_set_and_update)

        # ── Müşteri Detayları Formu (başta gizli) ──
        self.form_frame = ttk.LabelFrame(self.main_frame, text="👤 Müşteri Detayları", padding=8)
        # pack edilmeyecek — toggle ile açılacak

        form_inner = ttk.Frame(self.form_frame)
        form_inner.pack(fill=tk.X)

        left_f  = ttk.Frame(form_inner)
        right_f = ttk.Frame(form_inner)
        left_f.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        right_f.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.form: Dict[str, FormField] = {}

        left_fields = [
            ("name",            "Ad / Ünvan:",      "entry"),
            ("email",           "E-posta:",          "entry"),
            ("phone",           "Telefon:",          "entry"),
            ("address",         "Adres:",            "entry"),
            ("tax_number",      "Vergi No:",         "entry"),
        ]
        right_fields = [
            ("rent_start_year",  "Başlangıç Tarihi:", "entry"),
            ("rent_start_month", "Başlangıç Ayı:",    "combobox"),
            ("ilk_kira_bedeli",  "İlk Kira (₺):",    "entry"),
            ("current_rent",     "Gerçek Kira (₺):", "entry"),
            ("office_code",      "Ofis Kodu:",        "entry"),
        ]
        for key, label, ftype in left_fields:
            f = FormField(left_f, label=label, field_type=ftype,
                          options=get_months_list() if ftype == "combobox" else None)
            f.pack(fill=tk.X, pady=2)
            self.form[key] = f
        for key, label, ftype in right_fields:
            f = FormField(right_f, label=label, field_type=ftype,
                          options=get_months_list() if ftype == "combobox" else None)
            f.pack(fill=tk.X, pady=2)
            self.form[key] = f

        btn_f = ttk.Frame(self.form_frame)
        btn_f.pack(fill=tk.X, pady=(8, 0))
        ttk.Button(btn_f, text="💾 Ekle/Güncelle", command=self._on_save).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(btn_f, text="🗑 Sil",           command=self._on_delete).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(btn_f, text="✖ Temizle",        command=self._clear_form).pack(side=tk.LEFT)

        # ── Aylık Kira Girişi (başta gizli) ──
        self.rent_frame = ttk.LabelFrame(self.main_frame, text="💰 Aylık Kira Girişi", padding=8)
        # pack edilmeyecek — toggle ile açılacak

        rent_top = ttk.Frame(self.rent_frame)
        rent_top.pack(fill=tk.X, pady=(0, 6))
        ttk.Label(rent_top, text="Yıl:").pack(side=tk.LEFT)
        self.rent_year_var = tk.StringVar(value=str(date.today().year))
        # Yıl seçici — Tüm Yıllar seçeneği ile
        ttk.Label(rent_top, text="Yıl:").pack(side=tk.LEFT)
        self.rent_year_var = tk.IntVar(value=date.today().year)
        self.rent_year_spin = tk.Spinbox(
            rent_top, from_=2000, to=2100, width=6,
            textvariable=self.rent_year_var, font=("Segoe UI", 9),
            command=self._on_rent_year_change
        )
        self.rent_year_spin.pack(side=tk.LEFT, padx=(4, 4))
        self.rent_year_spin.bind("<Return>", lambda e: self._on_rent_year_change())

        # Tüm yıllar butonu
        ttk.Button(rent_top, text="📅 Tüm Yıllar Ödenmiş",
                   command=self._mark_all_years_paid).pack(side=tk.LEFT, padx=(4, 12))

        ttk.Label(rent_top, text="TÜFE Kira:").pack(side=tk.LEFT)
        self.tufe_rent_label = ttk.Label(rent_top, text="—", foreground="#1a3a6b",
                                          font=("Segoe UI", 9, "bold"))
        self.tufe_rent_label.pack(side=tk.LEFT, padx=(4, 20))
        ttk.Label(rent_top, text="Ödenen Toplam:").pack(side=tk.LEFT)
        self.paid_total_label = ttk.Label(rent_top, text="—", foreground="#1a6b2a",
                                           font=("Segoe UI", 9, "bold"))
        self.paid_total_label.pack(side=tk.LEFT, padx=(4, 0))

        # Yıllık toplam giriş
        ttk.Label(rent_top, text="   Yıllık Toplam:").pack(side=tk.LEFT, padx=(16, 2))
        self.yearly_total_var = tk.StringVar()
        yearly_entry = ttk.Entry(rent_top, textvariable=self.yearly_total_var, width=12, justify="right")
        yearly_entry.pack(side=tk.LEFT)
        ttk.Button(rent_top, text="Dağıt", command=self._distribute_yearly_total).pack(side=tk.LEFT, padx=(4, 0))

        ttk.Button(rent_top, text="💾 Kaydet", command=self._on_save_rent_payments).pack(side=tk.RIGHT)

        months_grid = ttk.Frame(self.rent_frame)
        months_grid.pack(fill=tk.X, padx=4, pady=4)
        self.rent_month_vars: Dict[str, tk.StringVar] = {}

        for i, month_name in enumerate(MONTHS_TR):
            col_base = (i % 6) * 3   # 6 ay yan yana
            row_idx  = i // 6         # 2 satır

            ttk.Label(months_grid, text=month_name, width=7, anchor="e").grid(
                row=row_idx, column=col_base, padx=(6, 2), pady=3, sticky="e")
            var = tk.StringVar()
            self.rent_month_vars[month_name] = var
            entry = ttk.Entry(months_grid, textvariable=var, width=10, justify="right")
            entry.grid(row=row_idx, column=col_base + 1, padx=(0, 2), pady=3)
            ttk.Button(months_grid, text="✖", width=2,
                       command=lambda m=month_name: self._clear_month(m)
                       ).grid(row=row_idx, column=col_base + 2, padx=(0, 4), pady=3)
            var.trace_add("write", lambda *args: self._update_paid_total())

    def _toggle_panels(self) -> None:
        """Panelleri aç/kapat — sırası her zaman: form → rent."""
        self.form_frame.pack_forget()
        self.rent_frame.pack_forget()

        if self._show_form.get():
            self.form_frame.pack(side=tk.TOP, fill=tk.X, pady=(4, 0))
        if self._show_rent.get():
            self.rent_frame.pack(side=tk.TOP, fill=tk.X, pady=(4, 0))

    def _init_panels(self) -> None:
        """Başlangıçta panelleri göster."""
        self._toggle_panels()

    # ── Sütunlar ──

    def _get_year_range(self, customers) -> List[int]:
        years = set()
        for c in customers:
            years.update(c.get("rent_years_dict", {}).keys())
        return sorted(years)

    def _sort_by_col(self, col) -> None:
        """Sütun başlığına tıklanınca sırala."""
        if self._sort_col == col:
            self._sort_reverse = not self._sort_reverse
        else:
            self._sort_col = col
            self._sort_reverse = False

        items = [(self.tree.set(k, col), k) for k in self.tree.get_children("")]

        # Sayısal mı metin mi?
        def parse_val(v):
            try:
                # Para formatı: '1,234.56' → float
                return float(v.replace(",", "").replace(".", "").replace(" ", "") or "0")
            except:
                # Tarih: 'GG.AA.YYYY' → sıralama için YYYY.AA.GG
                parts = v.split(".")
                if len(parts) == 3 and all(p.isdigit() for p in parts):
                    return f"{parts[2]}.{parts[1]}.{parts[0]}"
                return v.lower()

        items.sort(key=lambda x: parse_val(x[0]), reverse=self._sort_reverse)

        for index, (_, k) in enumerate(items):
            self.tree.move(k, "", index)

        # Başlık okunu güncelle
        cols = self.tree["columns"]
        headers = {
            "name": "Ad / Ünvan", "tax_number": "Vergi No",
            "baslangic_tarihi": "Başlangıç Tarihi",
            "ilk_kira_bedeli": "İlk Kira (₺)", "current_rent": "Gerçek Kira (₺)",
        }
        for c in cols:
            label = headers.get(c, c)
            arrow = (" ▲" if not self._sort_reverse else " ▼") if c == col else ""
            self.tree.heading(c, text=label + arrow)

    def _on_gorunum_degis(self) -> None:
        """Yıl/Ay görünümü değişince sütunları yeniden oluştur."""
        self.refresh()

    def _refresh_columns(self, year_range: List[int]) -> None:
        gorunum_yil = self.gorunum_yil_var.get()
        gorunum_ay  = self.gorunum_ay_var.get()

        base = ["name", "tax_number", "baslangic_tarihi", "baslangic_ayi", "baslangic_yili", "ilk_kira_bedeli"]

        # Ay bazlı görünüm: tek yıl + 12 ay sütunu
        if gorunum_ay != "Tüm Aylar" or (gorunum_yil != "Tüm Yıllar" and gorunum_ay == "Tüm Aylar"):
            if gorunum_yil == "Tüm Yıllar":
                # Ay seçilmiş ama yıl tüm yıllar — tüm yıllar için o ayı göster
                y_cols = [f"{y}_{gorunum_ay}" for y in year_range] if gorunum_ay != "Tüm Aylar" else [str(y) for y in year_range]
            else:
                yil = int(gorunum_yil)
                if gorunum_ay != "Tüm Aylar":
                    # Tek yıl, tek ay — sadece o ayı göster
                    y_cols = [f"{yil}_{gorunum_ay}"]
                else:
                    # Tek yıl, tüm aylar — 12 ay sütunu
                    y_cols = [f"{yil}_{ay}" for ay in MONTHS_TR]
        else:
            # Tüm yıllar, tüm aylar — yıl bazlı özet
            y_cols = [str(y) for y in year_range]

        self._current_y_cols = y_cols
        cols = base + y_cols + ["current_rent", "son_tahsilat"]
        self.tree["columns"] = cols
        self.tree.column("#0", width=0, stretch=tk.NO)

        headers = {
            "name": "Ad / Ünvan", "tax_number": "Vergi No",
            "baslangic_tarihi": "Başlangıç Tarihi", "baslangic_ayi": "Ay",
            "baslangic_yili": "Yıl", "ilk_kira_bedeli": "İlk Kira (₺)",
            "current_rent": "Gerçek Kira (₺)", "son_tahsilat": "Son Tahsilat",
        }

        for col in cols:
            if col in headers:
                label = headers[col]
            elif "_" in col:
                # "2026_Ocak" → "Ocak"
                parts = col.split("_", 1)
                label = f"{parts[1]}" if gorunum_yil != "Tüm Yıllar" else f"{parts[1]} {parts[0]}"
            else:
                label = col

            self.tree.heading(col, text=label, command=lambda c=col: self._sort_by_col(c))
            if col == "son_tahsilat":
                w = 150
            elif col == "name":
                w = 160
            elif col == "baslangic_tarihi":
                w = 110
            elif col in ("baslangic_ayi",):
                w = 70
            elif col == "baslangic_yili":
                w = 50
            else:
                w = 90
            self.tree.column(col, width=w, anchor="w" if col == "name" else "e", minwidth=50)

    # ── Yenile ──

    def refresh(self) -> None:
        for item in self.tree.get_children():
            self.tree.delete(item)

        self._all_customers = get_all_customers_with_rent_progression()
        year_range = self._get_year_range(self._all_customers)
        self._refresh_columns(year_range)

        # Yıl filtresini güncelle
        yil_values = ["Tümü"] + [str(y) for y in sorted(
            set(c.get("rent_start_date", "").split(".")[-1]
                for c in self._all_customers
                if c.get("rent_start_date", ""))
        )]
        self.filter_yil["values"] = yil_values

        self._apply_filter()

    def _apply_filter(self) -> None:
        """Filtre uygula: ay, yıl, arama, ödenmeyen ay."""
        for item in self.tree.get_children():
            self.tree.delete(item)

        filtre_ay        = self.filter_ay_var.get()
        filtre_yil       = self.filter_yil_var.get()
        filtre_arama     = self.filter_arama_var.get().strip().upper()
        filtre_odenmemis = self.filter_odenmemis_var.get()

        esik = 0
        if filtre_odenmemis != "Tümü":
            esik = int(filtre_odenmemis.replace("+ ay", "").strip())

        year_range = self._get_year_range(self._all_customers)

        from database import get_connection
        import datetime
        bugun_yil = datetime.date.today().year
        bugun_ay  = datetime.date.today().month

        for c in self._all_customers:
            rent_date = c.get("rent_start_date", "") or ""
            baslangic_ayi  = ""
            baslangic_yili = ""

            if rent_date:
                try:
                    parts = rent_date.split(".")
                    baslangic_ayi  = MONTHS_TR[int(parts[1]) - 1]
                    baslangic_yili = parts[2]
                except (IndexError, ValueError):
                    pass

            # Ay/Yıl filtresi
            if filtre_ay  != "Tümü" and baslangic_ayi  != filtre_ay:
                continue
            if filtre_yil != "Tümü" and baslangic_yili != filtre_yil:
                continue

            # Arama filtresi
            if filtre_arama and filtre_arama not in c["name"].upper():
                continue

            # Ödenmeyen ay filtresi
            if esik > 0:
                conn = get_connection()
                try:
                    # Başlangıç tarihinden bugüne kadar beklenen ay sayısı
                    try:
                        p = rent_date.split(".")
                        s_ay = int(p[1])
                        s_yil = int(p[2])
                    except Exception:
                        s_ay, s_yil = 1, 2021

                    # Beklenen tüm ayları hesapla
                    beklenen = []
                    y, m = s_yil, s_ay
                    while (y < bugun_yil) or (y == bugun_yil and m <= bugun_ay):
                        beklenen.append((y, MONTHS_TR[m - 1]))
                        m += 1
                        if m > 12:
                            m = 1
                            y += 1

                    # Ödenmiş ayları çek (amount > 0)
                    odenmis_rows = conn.execute(
                        "SELECT year, month FROM rent_payments WHERE customer_id=? AND amount > 0",
                        (c["id"],)
                    ).fetchall()
                    odenmis_set = {(r["year"], r["month"]) for r in odenmis_rows}

                    # Ödenmemiş = beklenen - ödenmiş
                    odenmemis_sayi = sum(1 for (y2, m2) in beklenen if (y2, m2) not in odenmis_set)
                finally:
                    conn.close()

                if odenmemis_sayi < esik:
                    continue

            years_dict = c.get("rent_years_dict", {})
            values = [
                c["name"],
                c.get("tax_number", ""),
                rent_date,
                baslangic_ayi,
                baslangic_yili,
                f"{float(c.get('ilk_kira_bedeli', 0)):,.2f}",
            ]

            y_cols = getattr(self, "_current_y_cols", [str(y) for y in year_range])

            for col in y_cols:
                if "_" in col:
                    # Ay bazlı: "2026_Ocak"
                    yil_str, ay_str = col.split("_", 1)
                    row = get_connection().execute(
                        "SELECT amount FROM rent_payments WHERE customer_id=? AND year=? AND month=?",
                        (c["id"], int(yil_str), ay_str)
                    ).fetchone()
                    amt = float(row["amount"]) if row and row["amount"] else 0.0
                    values.append(f"{amt:,.2f}" if amt > 0 else "—")
                else:
                    # Yıl bazlı özet
                    rent = years_dict.get(int(col), years_dict.get(col, 0))
                    values.append(f"{float(rent):,.2f}")

            current = c.get("current_rent")
            if current:
                values.append(f"{float(current):,.2f}")
            elif year_range:
                last = years_dict.get(year_range[-1], years_dict.get(str(year_range[-1]), 0))
                values.append(f"{float(last):,.2f}")
            else:
                values.append("0,00")

            # Son tahsilat
            from database import fetch_all as _fa
            son = _fa(
                "SELECT tahsilat_tarihi, tutar, odeme_turu FROM tahsilatlar WHERE customer_id=? ORDER BY tahsilat_tarihi DESC, id DESC LIMIT 1",
                (c["id"],)
            )
            if son:
                try:
                    p = son[0]["tahsilat_tarihi"].split("-")
                    t_str = f"{p[2]}.{p[1]}.{p[0]}"
                except Exception:
                    t_str = son[0]["tahsilat_tarihi"]
                tur = "B" if son[0]["odeme_turu"] == "B" else "N"
                values.append(f"{t_str} {float(son[0]['tutar']):,.0f}₺({tur})")
            else:
                values.append("")

            self.tree.insert("", tk.END, iid=str(c["id"]), values=values)

        self._update_totals()

    def _update_totals(self) -> None:
        """Görünen satırların toplamlarını sütun altına hizalı göster."""
        self.totals_canvas.delete("all")
        self.totals_canvas.config(bg="#dde3ed")

        items = self.tree.get_children()
        if not items:
            self.totals_canvas.create_text(6, 14, text="📊 0 müşteri", anchor="w",
                                           font=("Segoe UI", 9, "bold"), fill="#1a3a6b")
            return

        cols = self.tree["columns"]
        totals = {}
        count = len(items)

        for iid in items:
            vals = self.tree.item(iid, "values")
            for i, col in enumerate(cols):
                if col in ("name", "tax_number", "baslangic_tarihi", "baslangic_ayi", "baslangic_yili"):
                    continue
                try:
                    v = float(str(vals[i]).replace(",", ""))
                    totals[col] = totals.get(col, 0.0) + v
                except (ValueError, IndexError):
                    pass

        # Müşteri sayısı sol başa
        self.totals_canvas.create_text(6, 14, text=f"📊 {count} müşteri",
                                       anchor="w", font=("Segoe UI", 9, "bold"), fill="#1a3a6b")

        # Treeview toplam genişliği
        total_width = sum(self.tree.column(c, "width") for c in cols)
        # Scroll oranı (0.0 - 1.0)
        try:
            scroll_left = self.tree.xview()[0]
        except Exception:
            scroll_left = 0.0
        scroll_offset = scroll_left * total_width

        # Canvas genişliği
        canvas_width = self.totals_canvas.winfo_width()

        x_offset = 0
        for col in cols:
            col_width = self.tree.column(col, "width")
            if col in totals:
                # Ekranda görünen x pozisyonu
                screen_x = x_offset - scroll_offset + col_width - 4
                if 0 < screen_x < canvas_width:
                    self.totals_canvas.create_text(
                        screen_x, 14,
                        text=f"{totals[col]:,.2f}",
                        anchor="e",
                        font=("Segoe UI", 9, "bold"),
                        fill="#000000"
                    )
            x_offset += col_width

    def _clear_filter(self) -> None:
        """Filtreleri temizle."""
        self.filter_ay_var.set("Tümü")
        self.filter_yil_var.set("Tümü")
        self.filter_odenmemis_var.set("Tümü")
        self.filter_arama_var.set("")
        self.gorunum_yil_var.set("2026")
        self.gorunum_ay_var.set("Tüm Aylar")
        self._apply_filter()

    # ── Satır Seçimi ──

    def _on_row_select(self, event=None) -> None:
        sel = self.tree.selection()
        if not sel:
            return

        self.selected_id = int(sel[0])

        from database import fetch_one
        row = fetch_one(
            """SELECT id, name, email, phone, address, tax_number,
                      rent_start_date, rent_start_year, rent_start_month,
                      ilk_kira_bedeli, current_rent, office_code
               FROM customers WHERE id = ?""",
            (self.selected_id,)
        )
        if not row:
            return

        self.form["name"].set(row["name"])
        self.form["email"].set(row["email"] or "")
        self.form["phone"].set(row["phone"] or "")
        self.form["address"].set(row["address"] or "")
        self.form["tax_number"].set(row["tax_number"] or "")
        self.form["rent_start_year"].set(row["rent_start_date"] or "")
        self.form["rent_start_month"].set(row["rent_start_month"] or "Ocak")
        self.form["ilk_kira_bedeli"].set(row["ilk_kira_bedeli"] or "")
        self.form["current_rent"].set(row["current_rent"] or "")
        self.form["office_code"].set(row["office_code"] or "")

        # Aylık kira panelini yükle
        self._load_rent_panel()


    def _on_musteri_detay(self, event=None) -> None:
        """Cift tiklama -> Musteri Detay Popup (Cari Ekstre + Kargo)."""
        sel = self.tree.selection()
        if not sel:
            return
        musteri_id = int(sel[0])
        from database import fetch_one
        row = fetch_one("SELECT id, name, phone FROM customers WHERE id=?", (musteri_id,))
        if not row:
            return
        MusteriDetayPopup(self, musteri_id, row["name"], row["phone"] or "")

    def _load_rent_panel(self) -> None:
        """Seçili müşteri ve yıl için aylık kira verilerini panele yükle."""
        if not self.selected_id:
            return
        year = int(self.rent_year_var.get())
        payments = get_rent_payments_for_year(self.selected_id, year)

        for month_name in MONTHS_TR:
            val = payments.get(month_name, "")
            self.rent_month_vars[month_name].set(f"{val:.2f}" if val else "")

        self._update_tufe_label(year)
        self._update_paid_total()

    def _on_rent_year_change(self) -> None:
        self._load_rent_panel()

    def _clear_month(self, month_name: str) -> None:
        """Tek ayın ödemesini sil — ödenmemiş olarak işaretle."""
        if not self.selected_id:
            return
        from database import delete_rent_payment
        year = self.rent_year_var.get()
        delete_rent_payment(self.selected_id, year, month_name)
        self.rent_month_vars[month_name].set("")
        self._update_paid_total()

    def _mark_all_years_paid(self) -> None:
        """Müşterinin başlangıç tarihinden bugüne tüm aylarını ödenmiş yap."""
        if not self.selected_id:
            messagebox.showwarning("Uyarı", "Önce bir müşteri seçin!")
            return
        from database import fetch_one, save_rent_payment
        import datetime
        row = fetch_one(
            "SELECT name, rent_start_year, rent_start_month, ilk_kira_bedeli FROM customers WHERE id=?",
            (self.selected_id,)
        )
        if not row:
            return
        onay = messagebox.askyesno(
            "Toplu Ödeme",
            f"{row['name']} için tüm yıllar/aylar ödenmiş olarak işaretlensin mi?\n"
            f"(Zaten girilmiş veriler değişmez)"
        )
        if not onay:
            return

        start_year = row["rent_start_year"] or 2021
        start_month_name = row["rent_start_month"] or "Ocak"
        try:
            start_month_idx = MONTHS_TR.index(start_month_name) + 1
        except ValueError:
            start_month_idx = 1
        ilk_kira = float(row["ilk_kira_bedeli"] or 0)

        bugun = datetime.date.today()
        count = 0
        y, m = start_year, start_month_idx
        while (y < bugun.year) or (y == bugun.year and m <= bugun.month):
            month_name = MONTHS_TR[m - 1]
            save_rent_payment(self.selected_id, y, month_name, ilk_kira)
            count += 1
            m += 1
            if m > 12:
                m = 1
                y += 1

        self._load_rent_panel()
        messagebox.showinfo("Tamamlandı", f"{count} ay ödenmiş olarak işaretlendi!")

    def _update_tufe_label(self, year: int) -> None:
        if not self.selected_id:
            self.tufe_rent_label.config(text="—")
            return
        from database import fetch_one
        row = fetch_one(
            "SELECT rent_start_year, rent_start_month, ilk_kira_bedeli FROM customers WHERE id=?",
            (self.selected_id,)
        )
        if not row or not row["ilk_kira_bedeli"]:
            self.tufe_rent_label.config(text="—")
            return
        prog = calculate_rent_progression(
            start_year=row["rent_start_year"],
            start_month=row["rent_start_month"] or "Ocak",
            initial_rent=float(row["ilk_kira_bedeli"])
        )
        tufe_val = prog.get("years", {}).get(year, 0)
        self.tufe_rent_label.config(text=f"{tufe_val:,.2f} ₺")

    def _update_paid_total(self) -> None:
        total = 0.0
        for var in self.rent_month_vars.values():
            raw = var.get().replace(",", ".").replace(" ", "")
            try:
                total += float(raw) if raw else 0.0
            except ValueError:
                pass
        self.paid_total_label.config(text=f"{total:,.2f} ₺")

    def _distribute_yearly_total(self) -> None:
        """Yıllık toplam tutarı aylara dağıt."""
        raw = self.yearly_total_var.get().replace(",", ".").replace(" ", "")
        try:
            yearly = float(raw) if raw else 0.0
        except ValueError:
            messagebox.showerror("Hata", "Geçerli bir tutar girin!")
            return

        if yearly <= 0:
            return

        # Müşterinin aylık kira tutarını bul
        if self.selected_id:
            from database import fetch_one
            row = fetch_one("SELECT ilk_kira_bedeli FROM customers WHERE id=?", (self.selected_id,))
            aylik = float(row["ilk_kira_bedeli"]) if row else 0.0
        else:
            aylik = 0.0

        # Aylık tutar bilinemiyorsa eşit böl
        if aylik <= 0:
            aylik = yearly / 12

        kalan = yearly
        for i, month_name in enumerate(MONTHS_TR):
            if kalan <= 0:
                self.rent_month_vars[month_name].set("")
            elif kalan >= aylik:
                self.rent_month_vars[month_name].set(f"{aylik:.2f}")
                kalan -= aylik
            else:
                # Kalan yarım ay
                self.rent_month_vars[month_name].set(f"{kalan:.2f}")
                kalan = 0

        self._update_paid_total()

    def _on_save_rent_payments(self) -> None:
        if not self.selected_id:
            messagebox.showwarning("Uyarı", "Önce bir müşteri seçin!")
            return
        try:
            year = int(self.rent_year_var.get())
        except ValueError:
            messagebox.showerror("Hata", "Geçerli bir yıl girin!")
            return

        for month_name in MONTHS_TR:
            raw = self.rent_month_vars[month_name].get().replace(",", ".").replace(" ", "")
            try:
                amount = float(raw) if raw else 0.0
                save_rent_payment(self.selected_id, year, month_name, amount)
            except ValueError:
                pass

        messagebox.showinfo("Başarılı", f"{year} yılı aylık kira kaydedildi.")
        self.refresh()

    def _clear_form(self) -> None:
        for f in self.form.values():
            f.clear()
        for var in self.rent_month_vars.values():
            var.set("")
        self.tufe_rent_label.config(text="—")
        self.paid_total_label.config(text="—")
        if self.tree.selection():
            self.tree.selection_remove(self.tree.selection())
        self.selected_id = None

    # ── Kaydet ──

    def _on_save(self) -> None:
        try:
            name = self.form["name"].get()
            if not name:
                messagebox.showwarning("Uyarı", "Ad / Ünvan alanı zorunludur!")
                return

            email      = self.form["email"].get()
            phone      = self.form["phone"].get()
            address    = self.form["address"].get()
            tax_number = self.form["tax_number"].get()
            month      = self.form["rent_start_month"].get() or "Ocak"

            year_str = self.form["rent_start_year"].get()  # artık tarih alanı
            year: Optional[int] = None
            month: str = "Ocak"
            rent_date: str = ""

            if year_str:
                # GG.AA.YYYY formatını parse et
                parts = year_str.strip().split('.')
                if len(parts) == 3:
                    try:
                        gun = int(parts[0])
                        ay  = int(parts[1])
                        yil = int(parts[2])
                        year = yil
                        month = get_months_list()[ay - 1]
                        rent_date = f"{gun:02d}.{ay:02d}.{yil}"
                    except (ValueError, IndexError):
                        messagebox.showerror("Hata", "Tarih formatı GG.AA.YYYY olmalıdır! Örn: 01.01.2024")
                        return
                else:
                    # Sadece yıl girilmişse
                    try:
                        year = int(year_str)
                        rent_date = f"01.01.{year}"
                    except ValueError:
                        messagebox.showerror("Hata", "Tarih formatı GG.AA.YYYY olmalıdır! Örn: 01.01.2024")
                        return

            def to_float(field_key: str) -> float:
                raw = self.form[field_key].get().replace(".", "").replace(",", ".")
                return float(raw) if raw else 0.0

            ilk_kira = to_float("ilk_kira_bedeli")
            current  = to_float("current_rent")
            office_code = self.form["office_code"].get().strip()

            if self.selected_id is None:
                insert_customer(
                    name=name, email=email, phone=phone, address=address,
                    tax_number=tax_number, rent_start_date=rent_date,
                    rent_start_year=year, rent_start_month=month,
                    ilk_kira_bedeli=ilk_kira, current_rent=current,
                )
                # Yeni eklenen müşterinin ID'sini bul ve office_code kaydet
                from database import fetch_one as _fo
                new_c = _fo("SELECT id FROM customers WHERE tax_number=? OR name=? ORDER BY id DESC LIMIT 1",
                            (tax_number, name))
                if new_c and office_code:
                    from database import save_customer_office_code
                    save_customer_office_code(new_c["id"], office_code)
                messagebox.showinfo("Başarılı", "Müşteri eklendi.")
            else:
                update_customer(
                    customer_id=self.selected_id,
                    name=name, email=email, phone=phone, address=address,
                    tax_number=tax_number, rent_start_date=rent_date,
                    rent_start_year=year, rent_start_month=month,
                    ilk_kira_bedeli=ilk_kira, current_rent=current,
                )
                if office_code is not None:
                    from database import save_customer_office_code
                    save_customer_office_code(self.selected_id, office_code)
                messagebox.showinfo("Başarılı", "Müşteri güncellendi.")

            self.refresh()
            self._clear_form()

        except ValueError as e:
            messagebox.showerror("Hata", f"Geçersiz veri: {e}")
        except Exception as e:
            messagebox.showerror("Hata", f"Kaydetme hatası: {e}")

    # ── Sil ──

    def _on_delete(self) -> None:
        if not self.selected_id:
            messagebox.showwarning("Uyarı", "Silmek için bir müşteri seçin!")
            return
        if not messagebox.askyesno("Onay", "Seçili müşteri silinecek. Emin misiniz?"):
            return
        try:
            delete_customer(self.selected_id)
            messagebox.showinfo("Başarılı", "Müşteri silindi.")
            self.refresh()
            self._clear_form()
        except Exception as e:
            messagebox.showerror("Hata", f"Silme hatası: {e}")

    # ── Excel ──

    def _on_import_excel(self) -> None:
        path = filedialog.askopenfilename(
            title="Müşteri Excel Dosyası Seç",
            filetypes=[("Excel Dosyaları", "*.xlsx *.xls"), ("Tüm Dosyalar", "*.*")],
        )
        if not path:
            return
        try:
            count = import_customers_from_excel(path)
            messagebox.showinfo("Başarılı", f"{count} müşteri içe aktarıldı.")
            self.refresh()
        except Exception as e:
            messagebox.showerror("Hata", f"Excel aktarma hatası: {e}")


# ============================================================================
# ÜRÜN SEKMESİ
# ============================================================================




# ============================================================================
# ============================================================================
# PERSONEL SEKMESİ
# ============================================================================

class PersonelTab(ttk.Frame):
    """Personel Devam Takip Sekmesi."""

    def __init__(self, master):
        super().__init__(master, padding=10)
        self._build_ui()
        self.refresh()

    def _get_conn(self):
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        return conn

    def _build_ui(self):
        from datetime import date

        # ── Üst bar ──
        top = ttk.Frame(self)
        top.pack(fill=tk.X, pady=(0, 8))

        ttk.Label(top, text="👤 Personel Devam Takip", font=("Segoe UI", 13, "bold")).pack(side=tk.LEFT)

        # Tarih seçici
        ttk.Label(top, text="  Tarih:").pack(side=tk.LEFT, padx=(16, 2))
        self.tarih_var = tk.StringVar(value=date.today().strftime("%d.%m.%Y"))
        tarih_e = ttk.Entry(top, textvariable=self.tarih_var, width=12)
        tarih_e.pack(side=tk.LEFT)
        tarih_e.bind("<Return>", lambda e: self.refresh())
        ttk.Button(top, text="🔍", command=self.refresh).pack(side=tk.LEFT, padx=(4, 0))
        ttk.Button(top, text="↺ Bugün", command=self._bugun).pack(side=tk.LEFT, padx=(4, 0))

        # Sunucu durumu
        self.wifi_lbl = ttk.Label(top, text="🔴 WiFi Takip Kapalı",
                                   font=("Segoe UI", 9, "bold"), foreground="#e74c3c")
        self.wifi_lbl.pack(side=tk.RIGHT, padx=(0, 8))
        self.wifi_ip_lbl = ttk.Label(top, text="", foreground="#4fc3f7", font=("Segoe UI", 9))
        self.wifi_ip_lbl.pack(side=tk.RIGHT, padx=(0, 8))
        self.btn_wifi = ttk.Button(top, text="📡 WiFi Takibi Başlat", command=self._toggle_wifi)
        self.btn_wifi.pack(side=tk.RIGHT)

        # ── Ana içerik: Sekmeli yapı ──
        self.nb = ttk.Notebook(self)
        self.nb.pack(fill=tk.BOTH, expand=True)

        # Devam sekmesi
        devam_frame = ttk.Frame(self.nb, padding=4)
        self.nb.add(devam_frame, text="  📋 Devam  ")

        # İzin sekmesi
        izin_frame = ttk.Frame(self.nb, padding=4)
        self.nb.add(izin_frame, text="  🏖️ İzin Yönetimi  ")

        self._build_devam_ui(devam_frame)
        self._build_izin_ui(izin_frame)

    def _build_devam_ui(self, parent):
        body = ttk.Frame(parent)
        body.pack(fill=tk.BOTH, expand=True)

        # Sol: Günlük rapor
        left = ttk.LabelFrame(body, text="📋 Günlük Devam Listesi", padding=8)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 8))

        cols = ("ad", "giris", "cikis", "durum", "gec_dakika")
        self.tree = ttk.Treeview(left, columns=cols, show="headings", height=12)
        hdrs = {"ad": "Personel", "giris": "Giriş", "cikis": "Çıkış",
                "durum": "Durum", "gec_dakika": "Geç (dk)"}
        wdts = {"ad": 160, "giris": 80, "cikis": 80, "durum": 120, "gec_dakika": 70}
        for c in cols:
            self.tree.heading(c, text=hdrs[c])
            self.tree.column(c, width=wdts[c], anchor="center" if c != "ad" else "w")

        vsb = ttk.Scrollbar(left, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.LEFT, fill=tk.Y)

        self.tree.tag_configure("gec",      foreground="#ff8a65")
        self.tree.tag_configure("zamaninda", foreground="#69f0ae")
        self.tree.tag_configure("gelmedi",  foreground="#90a4ae")

        # Canlı bildirim kutusu
        bildirim_f = ttk.LabelFrame(left, text="📡 Canlı Bildirimler", padding=6)
        bildirim_f.pack(fill=tk.X, pady=(6, 0))
        self.bildirim_text = tk.Text(bildirim_f, height=4, state="disabled",
                                      bg="#0d1f2d", fg="#4fc3f7",
                                      font=("Segoe UI", 9), relief="flat", wrap="word")
        self.bildirim_text.pack(fill=tk.X)

        # Manuel giriş/çıkış
        manuel = ttk.LabelFrame(left, text="✏️ Manuel Düzeltme", padding=6)
        manuel.pack(fill=tk.X, pady=(8, 0))

        mr = ttk.Frame(manuel); mr.pack(fill=tk.X)
        ttk.Label(mr, text="Personel:", width=9, anchor="e").pack(side=tk.LEFT)
        self.m_personel_var = tk.StringVar()
        self.m_personel_cb = ttk.Combobox(mr, textvariable=self.m_personel_var, width=18, state="readonly")
        self.m_personel_cb.pack(side=tk.LEFT, padx=(2, 8))

        ttk.Label(mr, text="Giriş:", anchor="e").pack(side=tk.LEFT)
        self.m_giris_var = tk.StringVar()
        ttk.Entry(mr, textvariable=self.m_giris_var, width=7).pack(side=tk.LEFT, padx=(2, 8))

        ttk.Label(mr, text="Çıkış:", anchor="e").pack(side=tk.LEFT)
        self.m_cikis_var = tk.StringVar()
        ttk.Entry(mr, textvariable=self.m_cikis_var, width=7).pack(side=tk.LEFT, padx=(2, 8))

        ttk.Button(mr, text="💾 Kaydet", command=self._manuel_kaydet).pack(side=tk.LEFT)

        # Sağ: Aylık özet + Personel yönetim
        right = ttk.Frame(body)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Aylık özet
        aylik = ttk.LabelFrame(right, text="📊 Aylık Özet", padding=8)
        aylik.pack(fill=tk.X, pady=(0, 8))

        ay_top = ttk.Frame(aylik); ay_top.pack(fill=tk.X, pady=(0, 6))
        ttk.Label(ay_top, text="Yıl:").pack(side=tk.LEFT)
        self.aylik_yil_var = tk.StringVar(value=str(date.today().year))
        ttk.Entry(ay_top, textvariable=self.aylik_yil_var, width=6).pack(side=tk.LEFT, padx=(2, 8))
        ttk.Label(ay_top, text="Ay:").pack(side=tk.LEFT)
        self.aylik_ay_var = tk.StringVar(value=str(date.today().month).zfill(2))
        ttk.Combobox(ay_top, textvariable=self.aylik_ay_var, width=4, state="readonly",
                     values=[str(i).zfill(2) for i in range(1, 13)]).pack(side=tk.LEFT, padx=(2, 8))
        ttk.Button(ay_top, text="📊 Göster", command=self._aylik_rapor).pack(side=tk.LEFT)

        ay_cols = ("ad", "toplam_gun", "gec_sayisi", "ort_gec")
        self.ay_tree = ttk.Treeview(aylik, columns=ay_cols, show="headings", height=5)
        ay_hdrs = {"ad": "Personel", "toplam_gun": "Toplam Gün",
                   "gec_sayisi": "Geç Sayısı", "ort_gec": "Ort. Geç (dk)"}
        for c in ay_cols:
            self.ay_tree.heading(c, text=ay_hdrs[c])
            self.ay_tree.column(c, width=120, anchor="center" if c != "ad" else "w")
        self.ay_tree.pack(fill=tk.X)
        self.ay_tree.tag_configure("uyari", foreground="#ff8a65")

        # Personel yönetimi
        pers = ttk.LabelFrame(right, text="👥 Personel Yönetimi", padding=8)
        pers.pack(fill=tk.BOTH, expand=True)

        pr = ttk.Frame(pers); pr.pack(fill=tk.X, pady=2)
        ttk.Label(pr, text="Ad:", width=12, anchor="e").pack(side=tk.LEFT)
        self.p_ad_var = tk.StringVar()
        ttk.Entry(pr, textvariable=self.p_ad_var, width=18).pack(side=tk.LEFT, padx=(2, 8))
        ttk.Label(pr, text="Mesai:", anchor="e").pack(side=tk.LEFT)
        self.p_mesai_var = tk.StringVar(value="09:00")
        ttk.Entry(pr, textvariable=self.p_mesai_var, width=7).pack(side=tk.LEFT, padx=(2, 8))
        ttk.Label(pr, text="Durum:", anchor="e").pack(side=tk.LEFT)
        self.p_aktif_var = tk.StringVar(value="aktif")
        ttk.Combobox(pr, textvariable=self.p_aktif_var, width=8, state="readonly",
                     values=["aktif", "pasif"]).pack(side=tk.LEFT, padx=(2, 0))

        pr2 = ttk.Frame(pers); pr2.pack(fill=tk.X, pady=2)
        ttk.Label(pr2, text="📱 MAC Adresi:", width=12, anchor="e").pack(side=tk.LEFT)
        self.p_mac_var = tk.StringVar()
        ttk.Entry(pr2, textvariable=self.p_mac_var, width=20).pack(side=tk.LEFT, padx=(2, 8))
        ttk.Button(pr2, text="📡 Ağı Tara", command=self._mac_tara).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Label(pr2, text="(AA:BB:CC:DD:EE:FF)", foreground="#90a4ae",
                  font=("Segoe UI", 8)).pack(side=tk.LEFT)

        pr3 = ttk.Frame(pers); pr3.pack(fill=tk.X, pady=2)
        ttk.Label(pr3, text="📝 Not:", width=12, anchor="e").pack(side=tk.LEFT)
        self.p_not_var = tk.StringVar()
        ttk.Entry(pr3, textvariable=self.p_not_var, width=42).pack(side=tk.LEFT, padx=(2, 0))

        pb = ttk.Frame(pers); pb.pack(fill=tk.X, pady=4)
        ttk.Button(pb, text="➕ Ekle / Güncelle", command=self._personel_ekle).pack(side=tk.LEFT)
        ttk.Button(pb, text="🟢 Aktif Yap",  command=lambda: self._aktif_degistir(1)).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(pb, text="🔴 Pasife Al",  command=lambda: self._aktif_degistir(0)).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(pb, text="✖ Temizle",     command=self._p_temizle).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(pb, text="🗑 Kalıcı Sil", command=self._personel_sil).pack(side=tk.RIGHT)

        # Filtre
        pf2 = ttk.Frame(pers); pf2.pack(fill=tk.X, pady=(0, 2))
        ttk.Label(pf2, text="Göster:", font=("Segoe UI", 8)).pack(side=tk.LEFT)
        self.p_filtre_var = tk.StringVar(value="tümü")
        for val, txt in [("tümü","Tümü"),("aktif","🟢 Aktif"),("pasif","🔴 Pasif")]:
            ttk.Radiobutton(pf2, text=txt, variable=self.p_filtre_var,
                            value=val, command=self._personel_listesi).pack(side=tk.LEFT, padx=(4,0))

        p_cols = ("ad", "mesai", "durum", "not_", "mac")
        self.p_tree = ttk.Treeview(pers, columns=p_cols, show="headings", height=6)
        p_hdrs = {"ad": "Personel Adı", "mesai": "Mesai", "durum": "Durum",
                  "not_": "Not / Değerlendirme", "mac": "MAC"}
        p_wdts = {"ad": 130, "mesai": 60, "durum": 70, "not_": 160, "mac": 130}
        for c in p_cols:
            self.p_tree.heading(c, text=p_hdrs[c])
            self.p_tree.column(c, width=p_wdts[c], anchor="w" if c in ("ad","not_") else "center")
        p_vsb = ttk.Scrollbar(pers, orient="vertical", command=self.p_tree.yview)
        self.p_tree.configure(yscrollcommand=p_vsb.set)
        self.p_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        p_vsb.pack(side=tk.LEFT, fill=tk.Y)
        self.p_tree.tag_configure("aktif", foreground="#69f0ae")
        self.p_tree.tag_configure("pasif", foreground="#90a4ae")
        self.p_tree.bind("<<TreeviewSelect>>", self._p_select)

        self._wifi_running = False

    def _build_izin_ui(self, parent):
        from datetime import date
        bugun = date.today()

        body = ttk.Frame(parent)
        body.pack(fill=tk.BOTH, expand=True)

        # ── SOL: Personel seç + raporlar ──
        left = ttk.Frame(body)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 6))

        pf = ttk.LabelFrame(left, text="👤 Personel Seç", padding=6)
        pf.pack(fill=tk.X, pady=(0, 4))
        pf_row = ttk.Frame(pf); pf_row.pack(fill=tk.X)
        ttk.Label(pf_row, text="Personel:").pack(side=tk.LEFT)
        self.iz_personel_var = tk.StringVar()
        self.iz_personel_cb = ttk.Combobox(pf_row, textvariable=self.iz_personel_var, width=22, state="readonly")
        self.iz_personel_cb.pack(side=tk.LEFT, padx=(4, 8))
        ttk.Label(pf_row, text="Yıl:").pack(side=tk.LEFT)
        self.iz_yil_var = tk.StringVar(value=str(bugun.year))
        ttk.Entry(pf_row, textvariable=self.iz_yil_var, width=6).pack(side=tk.LEFT, padx=(2, 8))
        ttk.Button(pf_row, text="🔍 Göster", command=self._izin_yukle).pack(side=tk.LEFT)

        # Bakiye kartları
        kart_f = ttk.Frame(left)
        kart_f.pack(fill=tk.X, pady=(0, 4))
        self.iz_kart_hak    = self._kart_olustur(kart_f, "Toplam Hak",    "—", "#4fc3f7")
        self.iz_kart_kul    = self._kart_olustur(kart_f, "Kullanılan",    "—", "#ff8a65")
        self.iz_kart_kalan  = self._kart_olustur(kart_f, "Kalan",         "—", "#69f0ae")
        self.iz_kart_rapor  = self._kart_olustur(kart_f, "Rapor/Sağlık", "—", "#ce93d8")
        self.iz_kart_mazeret= self._kart_olustur(kart_f, "Mazeret",       "—", "#90caf9")

        # Alt sekmeli raporlar
        gec_nb = ttk.Notebook(left)
        gec_nb.pack(fill=tk.BOTH, expand=True)

        gec_f = ttk.Frame(gec_nb)
        gec_nb.add(gec_f, text="  ⚠️ Geç Kalmalar  ")
        gec_cols = ("tarih", "giris", "gec_dk")
        self.gec_tree = ttk.Treeview(gec_f, columns=gec_cols, show="headings", height=8)
        for c, h2, w in [("tarih","Tarih",100),("giris","Giriş",100),("gec_dk","Geç (dk)",80)]:
            self.gec_tree.heading(c, text=h2)
            self.gec_tree.column(c, width=w, anchor="center")
        gec_vsb = ttk.Scrollbar(gec_f, orient="vertical", command=self.gec_tree.yview)
        self.gec_tree.configure(yscrollcommand=gec_vsb.set)
        self.gec_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        gec_vsb.pack(side=tk.LEFT, fill=tk.Y)
        self.gec_tree.tag_configure("gec", foreground="#ff8a65")
        self.gec_toplam_lbl = ttk.Label(gec_f, text="", font=("Segoe UI", 9, "bold"), foreground="#ff8a65")
        self.gec_toplam_lbl.pack(anchor="w", padx=4, pady=2)

        cikis_f = ttk.Frame(gec_nb)
        gec_nb.add(cikis_f, text="  🚶 Giriş/Çıkış Kayıtları  ")
        c_cols = ("tarih", "giris", "cikis")
        self.cikis_tree = ttk.Treeview(cikis_f, columns=c_cols, show="headings", height=8)
        for c, h2 in [("tarih","Tarih"),("giris","Giriş"),("cikis","Çıkış")]:
            self.cikis_tree.heading(c, text=h2)
            self.cikis_tree.column(c, width=110, anchor="center")
        c_vsb = ttk.Scrollbar(cikis_f, orient="vertical", command=self.cikis_tree.yview)
        self.cikis_tree.configure(yscrollcommand=c_vsb.set)
        self.cikis_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        c_vsb.pack(side=tk.LEFT, fill=tk.Y)

        # ── SAĞ: Bilgi formu + İzin girişi + Liste ──
        right = ttk.Frame(body)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        bilgi_f = ttk.LabelFrame(right, text="📝 Personel Ek Bilgileri", padding=6)
        bilgi_f.pack(fill=tk.X, pady=(0, 4))
        b1 = ttk.Frame(bilgi_f); b1.pack(fill=tk.X, pady=1)
        ttk.Label(b1, text="İşe Başlama:", width=13, anchor="e").pack(side=tk.LEFT)
        self.iz_ise_bas_var = tk.StringVar()
        ttk.Entry(b1, textvariable=self.iz_ise_bas_var, width=12).pack(side=tk.LEFT, padx=(2,8))
        ttk.Label(b1, text="Unvan:").pack(side=tk.LEFT)
        self.iz_unvan_var = tk.StringVar()
        ttk.Entry(b1, textvariable=self.iz_unvan_var, width=14).pack(side=tk.LEFT, padx=(2,0))
        b2 = ttk.Frame(bilgi_f); b2.pack(fill=tk.X, pady=1)
        ttk.Label(b2, text="Departman:", width=13, anchor="e").pack(side=tk.LEFT)
        self.iz_dept_var = tk.StringVar()
        ttk.Entry(b2, textvariable=self.iz_dept_var, width=12).pack(side=tk.LEFT, padx=(2,8))
        ttk.Label(b2, text="T.C. No:").pack(side=tk.LEFT)
        self.iz_tc_var = tk.StringVar()
        ttk.Entry(b2, textvariable=self.iz_tc_var, width=12).pack(side=tk.LEFT, padx=(2,8))
        ttk.Label(b2, text="Ek Gün:").pack(side=tk.LEFT)
        self.iz_manuel_gun_var = tk.StringVar(value="0")
        ttk.Entry(b2, textvariable=self.iz_manuel_gun_var, width=4).pack(side=tk.LEFT, padx=(2,0))
        ttk.Button(bilgi_f, text="💾 Bilgileri Kaydet", command=self._izin_bilgi_kaydet).pack(anchor="e", pady=(2,0))

        form_f = ttk.LabelFrame(right, text="➕ Yeni İzin Girişi", padding=6)
        form_f.pack(fill=tk.X, pady=(0, 4))
        iz1 = ttk.Frame(form_f); iz1.pack(fill=tk.X, pady=2)
        ttk.Label(iz1, text="İzin Türü:", width=13, anchor="e").pack(side=tk.LEFT)
        self.iz_tur_var = tk.StringVar(value="Yıllık Ücretli İzin")
        ttk.Combobox(iz1, textvariable=self.iz_tur_var, width=22, state="readonly",
                     values=["Yıllık Ücretli İzin","Ücretsiz İzin","Sağlık / Rapor","Mazeret İzni","Yarım Gün İzin"]
                     ).pack(side=tk.LEFT, padx=(2,0))
        iz2 = ttk.Frame(form_f); iz2.pack(fill=tk.X, pady=2)
        ttk.Label(iz2, text="Başlangıç:", width=13, anchor="e").pack(side=tk.LEFT)
        self.iz_bas_var = tk.StringVar(value=bugun.strftime("%d.%m.%Y"))
        ttk.Entry(iz2, textvariable=self.iz_bas_var, width=12).pack(side=tk.LEFT, padx=(2,8))
        ttk.Label(iz2, text="Bitiş:").pack(side=tk.LEFT)
        self.iz_bit_var = tk.StringVar(value=bugun.strftime("%d.%m.%Y"))
        ttk.Entry(iz2, textvariable=self.iz_bit_var, width=12).pack(side=tk.LEFT, padx=(2,0))
        iz3 = ttk.Frame(form_f); iz3.pack(fill=tk.X, pady=2)
        ttk.Label(iz3, text="Gün Sayısı:", width=13, anchor="e").pack(side=tk.LEFT)
        self.iz_gun_var = tk.StringVar(value="1")
        ttk.Entry(iz3, textvariable=self.iz_gun_var, width=6).pack(side=tk.LEFT, padx=(2,8))
        self.iz_yari_var = tk.BooleanVar()
        ttk.Checkbutton(iz3, text="Yarım Gün", variable=self.iz_yari_var).pack(side=tk.LEFT)
        iz4 = ttk.Frame(form_f); iz4.pack(fill=tk.X, pady=2)
        ttk.Label(iz4, text="Açıklama:", width=13, anchor="e").pack(side=tk.LEFT)
        self.iz_aciklama_var = tk.StringVar()
        ttk.Entry(iz4, textvariable=self.iz_aciklama_var, width=30).pack(side=tk.LEFT, padx=(2,0))
        btn_f = ttk.Frame(form_f); btn_f.pack(fill=tk.X, pady=(4,0))
        ttk.Button(btn_f, text="💾 İzin Kaydet",  command=self._izin_kaydet).pack(side=tk.LEFT)
        ttk.Button(btn_f, text="🖨️ PDF Formu",    command=self._izin_pdf).pack(side=tk.LEFT, padx=(6,0))
        ttk.Button(btn_f, text="💾 Word Formu",   command=self._izin_word).pack(side=tk.LEFT, padx=(6,0))
        ttk.Button(btn_f, text="✖ Temizle",
                   command=lambda: [self.iz_gun_var.set("1"), self.iz_aciklama_var.set(""), self.iz_yari_var.set(False)]
                   ).pack(side=tk.LEFT, padx=(6,0))

        iz_list_f = ttk.LabelFrame(right, text="📋 İzin Kayıtları", padding=4)
        iz_list_f.pack(fill=tk.BOTH, expand=True)
        iz_cols = ("tur","bas","bit","gun","aciklama")
        self.iz_tree = ttk.Treeview(iz_list_f, columns=iz_cols, show="headings", height=7)
        iz_hdrs = {"tur":"İzin Türü","bas":"Başlangıç","bit":"Bitiş","gun":"Gün","aciklama":"Açıklama"}
        iz_wdts = {"tur":130,"bas":85,"bit":85,"gun":45,"aciklama":120}
        for c in iz_cols:
            self.iz_tree.heading(c, text=iz_hdrs[c])
            self.iz_tree.column(c, width=iz_wdts[c], anchor="center" if c!="tur" else "w")
        iz_vsb = ttk.Scrollbar(iz_list_f, orient="vertical", command=self.iz_tree.yview)
        self.iz_tree.configure(yscrollcommand=iz_vsb.set)
        self.iz_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        iz_vsb.pack(side=tk.LEFT, fill=tk.Y)
        self.iz_tree.tag_configure("yillik",  foreground="#4fc3f7")
        self.iz_tree.tag_configure("rapor",   foreground="#ce93d8")
        self.iz_tree.tag_configure("mazeret", foreground="#90caf9")
        self.iz_tree.tag_configure("ucretsiz",foreground="#ff8a65")
        self.iz_tree.tag_configure("yari",    foreground="#fff176")
        ttk.Button(iz_list_f, text="🗑 Seçili İzni Sil", command=self._izin_sil).pack(pady=(4,0))

    def _kart_olustur(self, parent, baslik, deger, renk):
        f = tk.Frame(parent, bg="#1e3a50", relief="flat")
        f.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        tk.Label(f, text=baslik, font=("Segoe UI", 7), bg="#1e3a50", fg="#90a4ae").pack(pady=(4,0))
        lbl = tk.Label(f, text=deger, font=("Segoe UI", 11, "bold"), bg="#1e3a50", fg=renk)
        lbl.pack(pady=(0,4))
        return lbl

    def _izin_yukle(self):
        from database import get_izinler, get_izin_ozet, get_gec_kalma_rapor, get_cikis_rapor, get_personel_bilgi
        ad = self.iz_personel_var.get()
        if not ad:
            messagebox.showwarning("Uyarı", "Personel seçin!"); return
        pid = self._pid_bul(ad)
        if not pid: return
        yil = int(self.iz_yil_var.get())
        bilgi = get_personel_bilgi(pid)
        if bilgi:
            self.iz_ise_bas_var.set(bilgi.get("ise_baslama_tarihi") or "")
            self.iz_unvan_var.set(bilgi.get("unvan") or "")
            self.iz_dept_var.set(bilgi.get("departman") or "")
            self.iz_tc_var.set(bilgi.get("tc_no") or "")
            self.iz_manuel_gun_var.set(str(bilgi.get("manuel_izin_gun") or 0))
        ozet = get_izin_ozet(pid, yil)
        self.iz_kart_hak.config(text=f"{ozet['toplam_hak']} gün")
        self.iz_kart_kul.config(text=f"{ozet['yillik_kullanilan']} gün")
        kalan = ozet['kalan']
        self.iz_kart_kalan.config(text=f"{kalan} gün", fg="#69f0ae" if kalan > 0 else "#e74c3c")
        self.iz_kart_rapor.config(text=f"{ozet['rapor']} gün")
        self.iz_kart_mazeret.config(text=f"{ozet['mazeret']} gün")
        for item in self.iz_tree.get_children(): self.iz_tree.delete(item)
        tur_tag = {"Yıllık Ücretli İzin":"yillik","Sağlık / Rapor":"rapor",
                   "Mazeret İzni":"mazeret","Ücretsiz İzin":"ucretsiz","Yarım Gün İzin":"yari"}
        for iz in get_izinler(pid, yil):
            gun = f"{iz['gun_sayisi']}" + (" (½)" if iz.get("yari_gun") else "")
            self.iz_tree.insert("", tk.END, iid=str(iz["id"]),
                                tags=(tur_tag.get(iz["izin_turu"],""),),
                                values=(iz["izin_turu"], iz["baslangic_tarihi"],
                                        iz["bitis_tarihi"], gun, iz.get("aciklama") or ""))
        for item in self.gec_tree.get_children(): self.gec_tree.delete(item)
        toplam_dk = 0
        for g in get_gec_kalma_rapor(pid):
            try:
                p2 = g["tarih"].split("-"); tarih_str = f"{p2[2]}.{p2[1]}.{p2[0]}"
            except Exception:
                tarih_str = g["tarih"]
            self.gec_tree.insert("", tk.END, tags=("gec",),
                                  values=(tarih_str, g["giris_saati"] or "—", g["gec_dakika"]))
            toplam_dk += int(g["gec_dakika"] or 0)
        saat, dk = divmod(toplam_dk, 60)
        gec_list = get_gec_kalma_rapor(pid)
        self.gec_toplam_lbl.config(text=f"  Toplam {len(gec_list)} kez geç kalındı — {saat}s {dk}dk")
        for item in self.cikis_tree.get_children(): self.cikis_tree.delete(item)
        for c in get_cikis_rapor(pid):
            try:
                p2 = c["tarih"].split("-"); tarih_str = f"{p2[2]}.{p2[1]}.{p2[0]}"
            except Exception:
                tarih_str = c["tarih"]
            self.cikis_tree.insert("", tk.END, values=(tarih_str, c["giris_saati"] or "—", c["cikis_saati"] or "—"))

    def _izin_bilgi_kaydet(self):
        from database import save_personel_bilgi
        ad = self.iz_personel_var.get()
        if not ad: messagebox.showwarning("Uyarı", "Personel seçin!"); return
        pid = self._pid_bul(ad)
        if not pid: return
        try: manuel = int(self.iz_manuel_gun_var.get() or 0)
        except ValueError: manuel = 0
        save_personel_bilgi(pid, self.iz_ise_bas_var.get().strip(),
                             self.iz_unvan_var.get().strip(), self.iz_dept_var.get().strip(),
                             self.iz_tc_var.get().strip(), manuel)
        messagebox.showinfo("✅ Kaydedildi", "Personel bilgileri güncellendi.")
        self._izin_yukle()

    def _izin_kaydet(self):
        from database import insert_izin
        ad = self.iz_personel_var.get()
        if not ad: messagebox.showwarning("Uyarı", "Personel seçin!"); return
        pid = self._pid_bul(ad)
        if not pid: return
        try:
            bas = self._parse_tarih(self.iz_bas_var.get())
            bit = self._parse_tarih(self.iz_bit_var.get())
            yari = 1 if self.iz_yari_var.get() else 0
            gun = 0.5 if yari else float(self.iz_gun_var.get() or 1)
            insert_izin(pid, self.iz_tur_var.get(), bas, bit, gun, yari, self.iz_aciklama_var.get())
            messagebox.showinfo("✅ Kaydedildi", "İzin kaydedildi.")
            self._izin_yukle()
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _izin_sil(self):
        from database import delete_izin
        sel = self.iz_tree.selection()
        if not sel: return
        if messagebox.askyesno("Sil", "Seçili izin kaydı silinsin mi?"):
            delete_izin(int(sel[0])); self._izin_yukle()

    def _pid_bul(self, ad: str):
        conn = self._get_conn()
        r = conn.execute("SELECT id FROM personeller WHERE ad=?", (ad,)).fetchone()
        conn.close()
        return r["id"] if r else None

    def _izin_formu_data(self) -> dict:
        from datetime import date
        from database import get_izin_ozet
        ad = self.iz_personel_var.get()
        pid = self._pid_bul(ad) if ad else None
        bakiye = get_izin_ozet(pid, int(self.iz_yil_var.get())) if pid else None
        return {
            "firma_adi": "BestOffice ERP", "bugun": date.today().strftime("%d.%m.%Y"),
            "personel_ad": ad, "tc_no": self.iz_tc_var.get(),
            "unvan": self.iz_unvan_var.get(), "departman": self.iz_dept_var.get(),
            "ise_baslama": self.iz_ise_bas_var.get(), "izin_turu": self.iz_tur_var.get(),
            "baslangic": self.iz_bas_var.get(), "bitis": self.iz_bit_var.get(),
            "gun_sayisi": self.iz_gun_var.get(), "yari_gun": self.iz_yari_var.get(),
            "aciklama": self.iz_aciklama_var.get(), "izin_bakiye": bakiye,
        }

    def _izin_pdf(self):
        from izin_form_pdf import izin_formu_olustur
        from tkinter import filedialog
        import os
        ad = self.iz_personel_var.get()
        if not ad: messagebox.showwarning("Uyarı", "Personel seçin!"); return
        yol = filedialog.asksaveasfilename(title="PDF Kaydet", defaultextension=".pdf",
                                            filetypes=[("PDF","*.pdf")],
                                            initialfile=f"izin_{ad.replace(' ','_')}.pdf")
        if not yol: return
        try:
            izin_formu_olustur(self._izin_formu_data(), yol)
            messagebox.showinfo("✅ PDF", f"Kaydedildi:\n{yol}")
            if os.name == "nt": os.startfile(yol)
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _izin_word(self):
        from tkinter import filedialog
        import os
        ad = self.iz_personel_var.get()
        if not ad: messagebox.showwarning("Uyarı", "Personel seçin!"); return
        yol = filedialog.asksaveasfilename(title="Word Kaydet", defaultextension=".docx",
                                            filetypes=[("Word","*.docx")],
                                            initialfile=f"izin_{ad.replace(' ','_')}.docx")
        if not yol: return
        try:
            self._word_olustur(self._izin_formu_data(), yol)
            messagebox.showinfo("✅ Word", f"Kaydedildi:\n{yol}")
            if os.name == "nt": os.startfile(yol)
        except Exception as e:
            messagebox.showerror("Hata", f"Word hatası: {e}\npip install python-docx")

    def _word_olustur(self, data: dict, yol: str):
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for sec in doc.sections:
            sec.left_margin = sec.right_margin = Cm(2.5)
            sec.top_margin = sec.bottom_margin = Cm(2)
        def bold_satir(label, deger):
            p = doc.add_paragraph()
            r = p.add_run(f"{label} "); r.bold = True; r.font.size = Pt(10)
            r2 = p.add_run(deger or "________________________"); r2.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(data.get("firma_adi","BestOffice")); r.font.size=Pt(14); r.bold=True
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("PERSONEL İZİN TALEP FORMU"); r2.font.size=Pt(16); r2.bold=True
        doc.add_paragraph(f"Tarih: {data.get('bugun','')}")
        doc.add_paragraph("─"*55)
        doc.add_heading("Personel Bilgileri", 2)
        bold_satir("Adı Soyadı:", data.get("personel_ad"))
        bold_satir("T.C. Kimlik No:", data.get("tc_no"))
        bold_satir("Bölüm / Departman:", data.get("departman"))
        bold_satir("Unvan:", data.get("unvan"))
        bold_satir("İşe Başlama Tarihi:", data.get("ise_baslama"))
        doc.add_paragraph("─"*55)
        doc.add_heading("İzin Detayları", 2)
        secili = data.get("izin_turu","")
        for tur in ["Yıllık Ücretli İzin","Ücretsiz İzin","Sağlık / Rapor","Mazeret İzni","Yarım Gün İzin"]:
            p3 = doc.add_paragraph(style="List Bullet")
            r3 = p3.add_run(f"{'☑' if tur==secili else '☐'}  {tur}")
            r3.font.size=Pt(10); r3.bold = (tur==secili)
        doc.add_paragraph()
        gun_str = str(data.get("gun_sayisi",""))
        if data.get("yari_gun"): gun_str += " (Yarım Gün)"
        bold_satir("İzin Başlangıç:", data.get("baslangic"))
        bold_satir("İzin Bitiş:", data.get("bitis"))
        bold_satir("Toplam Gün:", gun_str)
        bold_satir("Açıklama:", data.get("aciklama"))
        bakiye = data.get("izin_bakiye")
        if bakiye:
            doc.add_paragraph("─"*55)
            doc.add_heading("İzin Bakiyesi", 2)
            bold_satir("Toplam Hak:", f"{bakiye.get('toplam_hak',14)} gün")
            bold_satir("Kullanılan:", f"{bakiye.get('yillik_kullanilan',0)} gün")
            bold_satir("Kalan:", f"{bakiye.get('kalan',14)} gün")
        doc.add_paragraph("─"*55)
        doc.add_paragraph('"Yukarıda belirttiğim tarihlerde izin kullanmak istediğimi beyan ederim."')
        doc.add_paragraph()
        table = doc.add_table(rows=3, cols=3); table.style="Table Grid"
        for i, b in enumerate(["Personel İmzası","Departman Yöneticisi","İK / Genel Müdür"]):
            c = table.cell(0,i); c.text=b
            c.paragraphs[0].runs[0].bold=True
            c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for i in range(3):
            table.cell(1,i).text="\n\n\n"
            p4=table.cell(2,i).paragraphs[0]; p4.add_run("İmza / Tarih")
            p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.save(yol)

    def _bugun(self):

        from datetime import date
        self.tarih_var.set(date.today().strftime("%d.%m.%Y"))
        self.refresh()

    def _parse_tarih(self, s):
        try:
            p = s.strip().split(".")
            return f"{p[2]}-{p[1]:0>2}-{p[0]:0>2}"
        except Exception:
            from datetime import date
            return date.today().isoformat()

    def refresh(self):
        self._gunluk_rapor()
        self._personel_listesi()
        self._aylik_rapor()
        # İzin sekmesi personel listesini güncelle
        try:
            conn = self._get_conn()
            adlar = [r["ad"] for r in conn.execute("SELECT ad FROM personeller WHERE aktif=1 ORDER BY ad").fetchall()]
            conn.close()
            self.iz_personel_cb["values"] = adlar
            self.m_personel_cb["values"] = adlar
        except Exception:
            pass

    def _toggle_wifi(self):
        if self._wifi_running:
            # Durdur
            try:
                from wifi_takip import servis
                servis.durdur()
            except Exception:
                pass
            self._wifi_running = False
            self.wifi_lbl.config(text="🔴 WiFi Takip Kapalı", foreground="#e74c3c")
            self.wifi_ip_lbl.config(text="")
            self.btn_wifi.config(text="📡 WiFi Takibi Başlat")
        else:
            # Başlat
            try:
                from wifi_takip import servis, set_durum_callback
                set_durum_callback(self._wifi_bildirim)
                servis.baslat()
                self._wifi_running = True
                self.wifi_lbl.config(text="🟢 WiFi Takip Aktif", foreground="#69f0ae")
                # IP bilgisi
                import socket
                try:
                    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
                    s.connect(("8.8.8.8", 80))
                    ip = s.getsockname()[0]
                    s.close()
                except Exception:
                    ip = "yerel ağ"
                self.wifi_ip_lbl.config(text=f"📶 {ip}")
                self.btn_wifi.config(text="⏹ WiFi Takibi Durdur")
                self._bildirim_ekle("✅ WiFi takip başladı — MAC adresi kayıtlı personeller izleniyor")
                # İlk taramayı hemen yap
                threading.Thread(target=self._ilk_tarama, daemon=True).start()
            except ImportError:
                messagebox.showerror("Hata", "wifi_takip.py bulunamadı!")
            except Exception as e:
                messagebox.showerror("Hata", str(e))

    def _ilk_tarama(self):
        """Servis başlayınca 5 saniye sonra ilk tarama."""
        time.sleep(5)
        try:
            from wifi_takip import servis
            servis._tara()
            self.after(0, self.refresh)
        except Exception:
            pass

    def _wifi_bildirim(self, mesaj: str, pid: int = None, tip: str = "bilgi"):
        """WiFi takip servisi bildirim callback — UI thread safe."""
        self.after(0, lambda: self._bildirim_ekle(mesaj))
        self.after(0, self.refresh)

    def _bildirim_ekle(self, mesaj: str):
        from datetime import datetime
        saat = datetime.now().strftime("%H:%M:%S")
        self.bildirim_text.config(state="normal")
        self.bildirim_text.insert("1.0", f"[{saat}] {mesaj}\n")
        # Son 50 satırı tut
        lines = int(self.bildirim_text.index("end-1c").split(".")[0])
        if lines > 50:
            self.bildirim_text.delete("51.0", "end")
        self.bildirim_text.config(state="disabled")

    def _mac_tara(self):
        """İki aşamalı tarama: önce mevcut listeyi al, sonra yeni eklenenleri göster."""
        popup = tk.Toplevel(self)
        popup.title("📡 Akıllı MAC Tarama")
        popup.geometry("520x480")
        popup.grab_set()
        popup.resizable(False, False)

        # Başlık
        tk.Label(popup, text="📡 Akıllı MAC Adresi Bulucu",
                 font=("Segoe UI", 12, "bold"), bg="#0f2537", fg="#4fc3f7"
                 ).pack(fill=tk.X, ipady=8)

        # Adım açıklaması
        self._tara_adim_lbl = ttk.Label(popup,
            text="1. Adım: Telefonun WiFi'sini KAPAT, ardından 'Tara' butonuna bas",
            font=("Segoe UI", 9), wraplength=480, foreground="#f59e0b")
        self._tara_adim_lbl.pack(pady=(8, 4), padx=12, anchor="w")

        # Liste frame
        liste_frame = ttk.Frame(popup)
        liste_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=4)

        # Sol: Önceki tarama
        sol = ttk.LabelFrame(liste_frame, text="📋 Mevcut Cihazlar (WiFi Kapalı)", padding=4)
        sol.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 4))
        self._lb_once = tk.Listbox(sol, font=("Courier", 9), bg="#1a2e3f", fg="#90a4ae",
                                    selectbackground="#0097a7", height=12)
        self._lb_once.pack(fill=tk.BOTH, expand=True)

        # Sağ: Yeni gelenler
        sag = ttk.LabelFrame(liste_frame, text="✨ Yeni Eklenenler (Senin Telefon!)", padding=4)
        sag.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self._lb_yeni = tk.Listbox(sag, font=("Courier", 9, "bold"), bg="#0d2818",
                                    fg="#69f0ae", selectbackground="#00897b", height=12)
        self._lb_yeni.pack(fill=tk.BOTH, expand=True)

        # Durum
        self._tara_durum = ttk.Label(popup, text="", foreground="#4fc3f7", font=("Segoe UI", 8))
        self._tara_durum.pack(pady=2)

        # Butonlar
        btn_frame = ttk.Frame(popup)
        btn_frame.pack(fill=tk.X, padx=10, pady=(4, 8))

        self._btn_tara1 = ttk.Button(btn_frame, text="📡 1. Tara (WiFi Kapalıyken)",
                                      command=lambda: self._tara_adim1(popup))
        self._btn_tara1.pack(side=tk.LEFT)

        self._btn_tara2 = ttk.Button(btn_frame, text="📡 2. Tara (WiFi Açıkken)",
                                      command=lambda: self._tara_adim2(popup), state="disabled")
        self._btn_tara2.pack(side=tk.LEFT, padx=(6, 0))

        self._btn_sec = ttk.Button(btn_frame, text="✅ Seç ve Ata",
                                    command=lambda: self._tara_sec(popup), state="disabled")
        self._btn_sec.pack(side=tk.LEFT, padx=(6, 0))

        ttk.Button(btn_frame, text="✖ Kapat",
                   command=popup.destroy).pack(side=tk.RIGHT)

        # İpucu ve manuel giriş
        ipucu_frame = ttk.LabelFrame(popup, text="📱 Telefonda MAC Adresini Bul", padding=6)
        ipucu_frame.pack(fill=tk.X, padx=10, pady=(0, 4))

        ipucu_sol = ttk.Frame(ipucu_frame)
        ipucu_sol.pack(side=tk.LEFT, fill=tk.X, expand=True)

        ttk.Label(ipucu_sol, text="🍎 iPhone:  Ayarlar → Genel → Hakkında → WiFi Adresi",
                  font=("Segoe UI", 8), foreground="#4fc3f7").pack(anchor="w")
        ttk.Label(ipucu_sol, text="🤖 Android: Ayarlar → Telefon Hakkında → Durum → WiFi MAC Adresi",
                  font=("Segoe UI", 8), foreground="#69f0ae").pack(anchor="w")

        manuel_frame = ttk.Frame(ipucu_frame)
        manuel_frame.pack(side=tk.RIGHT, padx=(10, 0))
        ttk.Label(manuel_frame, text="Manuel Gir:", font=("Segoe UI", 8)).pack(side=tk.LEFT)
        self._manuel_mac_var = tk.StringVar()
        manuel_entry = ttk.Entry(manuel_frame, textvariable=self._manuel_mac_var, width=18,
                                  font=("Courier", 9))
        manuel_entry.pack(side=tk.LEFT, padx=(4, 4))
        ttk.Button(manuel_frame, text="✅ Kullan",
                   command=lambda: self._manuel_mac_kullan(popup)).pack(side=tk.LEFT)

        self._onceki_macler: set = set()

    def _manuel_mac_kullan(self, popup):
        """Manuel girilen MAC adresini forma ata."""
        mac = self._manuel_mac_var.get().strip().upper().replace("-", ":")
        if not mac:
            messagebox.showwarning("Uyarı", "MAC adresi giriniz!", parent=popup)
            return
        # Basit format kontrolü
        import re as _re
        if not _re.match(r"^([0-9A-F]{2}:){5}[0-9A-F]{2}$", mac):
            messagebox.showwarning("Format Hatası",
                                   "MAC adresi formatı yanlış!\nDoğru format: AA:BB:CC:DD:EE:FF",
                                   parent=popup)
            return
        self.p_mac_var.set(mac)
        popup.destroy()
        messagebox.showinfo("✅ Eklendi", f"MAC adresi forma eklendi:\n{mac}\n\nŞimdi 'Ekle/Güncelle' butonuna basın.")

    def _tara_adim1(self, popup):
        """1. Tarama: WiFi kapalıyken mevcut cihazları kaydet."""
        self._btn_tara1.config(state="disabled", text="⏳ Taranıyor...")
        self._tara_durum.config(text="Ağ taranıyor, lütfen bekleyin...")

        def tara():
            from wifi_takip import arp_tara
            macler = arp_tara()
            self.after(0, lambda: self._tara_adim1_sonuc(macler))

        threading.Thread(target=tara, daemon=True).start()

    def _tara_adim1_sonuc(self, macler: set):
        self._onceki_macler = macler
        self._lb_once.delete(0, tk.END)

        # Kayıtlı MAC'leri de al
        conn = self._get_conn()
        kayitli = {r["mac_adresi"] for r in conn.execute(
            "SELECT mac_adresi FROM personeller WHERE mac_adresi IS NOT NULL"
        ).fetchall()}
        adlar = {r["mac_adresi"]: r["ad"] for r in conn.execute(
            "SELECT ad, mac_adresi FROM personeller WHERE mac_adresi IS NOT NULL"
        ).fetchall()}
        conn.close()

        for mac in sorted(macler):
            etiket = f" ← {adlar[mac]}" if mac in adlar else (" ← Kayıtlı" if mac in kayitli else "")
            self._lb_once.insert(tk.END, f"{mac}{etiket}")

        self._tara_adim_lbl.config(
            text="2. Adım: Şimdi telefonun WiFi'sini AÇ, bağlan, ardından '2. Tara' butonuna bas",
            foreground="#69f0ae")
        self._tara_durum.config(text=f"✅ {len(macler)} cihaz bulundu. Şimdi telefonu WiFi'ye bağlayın.")
        self._btn_tara2.config(state="normal")
        self._btn_tara1.config(text="↺ Tekrar Tara (1.Adım)")
        self._btn_tara1.config(state="normal")

    def _tara_adim2(self, popup):
        """2. Tarama: WiFi açıkken yeni eklenenler = senin telefonun."""
        self._btn_tara2.config(state="disabled", text="⏳ Taranıyor...")
        self._tara_durum.config(text="Yeni cihazlar aranıyor...")

        def tara():
            from wifi_takip import arp_tara
            import time
            time.sleep(1)
            macler = arp_tara()
            self.after(0, lambda: self._tara_adim2_sonuc(macler))

        threading.Thread(target=tara, daemon=True).start()

    def _tara_adim2_sonuc(self, yeni_macler: set):
        self._lb_yeni.delete(0, tk.END)
        fark = yeni_macler - self._onceki_macler

        if fark:
            for mac in sorted(fark):
                self._lb_yeni.insert(tk.END, mac)
            self._tara_adim_lbl.config(
                text=f"✅ {len(fark)} yeni cihaz bulundu! Sağ listeden telefonunu seç.",
                foreground="#69f0ae")
            self._tara_durum.config(text=f"🎉 Yeni cihaz(lar) yeşil listede gösteriliyor!")
            self._btn_sec.config(state="normal")
            # İlk yeni MAC'i otomatik seç
            self._lb_yeni.selection_set(0)
        else:
            self._tara_adim_lbl.config(
                text="⚠️ Yeni cihaz bulunamadı. Telefon bağlı mı? Private WiFi Address kapalı mı?",
                foreground="#ff8a65")
            self._tara_durum.config(text="Tekrar deneyin: telefonu WiFi'den koparıp bağlayın.")

        self._btn_tara2.config(state="normal", text="📡 2. Tara (WiFi Açıkken)")

    def _tara_sec(self, popup):
        """Seçili MAC'i personel formuna ata."""
        sel = self._lb_yeni.curselection()
        if not sel:
            messagebox.showwarning("Uyarı", "Yeşil listeden bir MAC seçin!", parent=popup)
            return
        mac = self._lb_yeni.get(sel[0]).strip()
        self.p_mac_var.set(mac)
        popup.destroy()
        messagebox.showinfo("✅ Seçildi",
                            f"MAC adresi forma eklendi:\n{mac}\n\nŞimdi 'Ekle/Güncelle' butonuna basın.")

    def _mac_secim_pencere(self, macler: set):
        pass  # Artık _mac_tara direkt popup açıyor

    def _personel_ekle(self):
        ad    = self.p_ad_var.get().strip()
        mesai = self.p_mesai_var.get().strip() or "09:00"
        mac   = self.p_mac_var.get().strip().upper().replace("-", ":")
        aktif = 1 if self.p_aktif_var.get() == "aktif" else 0
        notlar = self.p_not_var.get().strip()
        if not ad:
            messagebox.showwarning("Uyarı", "Ad giriniz!")
            return
        try:
            conn = self._get_conn()
            sel = self.p_tree.selection()
            if sel:
                conn.execute(
                    "UPDATE personeller SET ad=?, mesai_baslangic=?, mac_adresi=?, aktif=?, notlar=? WHERE id=?",
                    (ad, mesai, mac or None, aktif, notlar, int(sel[0]))
                )
                messagebox.showinfo("✅", f"{ad} güncellendi.")
            else:
                conn.execute(
                    "INSERT OR IGNORE INTO personeller (ad, mesai_baslangic, mac_adresi, aktif, notlar) VALUES (?,?,?,?,?)",
                    (ad, mesai, mac or None, aktif, notlar)
                )
                messagebox.showinfo("✅", f"{ad} eklendi.")
            conn.commit(); conn.close()
            self._p_temizle()
            self._personel_listesi()
            self._gunluk_rapor()
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _gunluk_rapor(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        tarih = self._parse_tarih(self.tarih_var.get())
        try:
            conn = self._get_conn()
            # Tüm personeller
            personeller = conn.execute("SELECT id, ad FROM personeller WHERE aktif=1").fetchall()
            devam = {r["personel_id"]: r for r in conn.execute(
                "SELECT * FROM devam WHERE tarih=?", (tarih,)
            ).fetchall()}
            conn.close()

            for p in personeller:
                d = devam.get(p["id"])
                if d:
                    giris = d["giris_saati"] or "—"
                    cikis = d["cikis_saati"] or "—"
                    if d["gec_kaldi"]:
                        durum = f"⚠️ Geç Geldi"
                        tag = "gec"
                    else:
                        durum = "✅ Zamanında"
                        tag = "zamaninda"
                    gec_dk = str(d["gec_dakika"]) if d["gec_kaldi"] else "—"
                else:
                    giris = "—"; cikis = "—"; durum = "❓ Kayıt Yok"; gec_dk = "—"
                    tag = "gelmedi"
                self.tree.insert("", tk.END, tags=(tag,),
                                 values=(p["ad"], giris, cikis, durum, gec_dk))

            # Manuel combobox güncelle
            self.m_personel_cb["values"] = [p["ad"] for p in personeller]
            self._personel_listesi_data = list(personeller)

        except Exception as e:
            pass

    def _aylik_rapor(self):
        for item in self.ay_tree.get_children():
            self.ay_tree.delete(item)
        yil = self.aylik_yil_var.get()
        ay  = self.aylik_ay_var.get()
        try:
            conn = self._get_conn()
            rows = conn.execute("""
                SELECT p.ad,
                       COUNT(d.id) as toplam_gun,
                       SUM(d.gec_kaldi) as gec_sayisi,
                       COALESCE(AVG(CASE WHEN d.gec_kaldi=1 THEN d.gec_dakika END), 0) as ort_gec
                FROM personeller p
                LEFT JOIN devam d ON d.personel_id=p.id AND d.tarih LIKE ?
                WHERE p.aktif=1
                GROUP BY p.id
            """, (f"{yil}-{ay}%",)).fetchall()
            conn.close()
            for r in rows:
                tag = "uyari" if (r["gec_sayisi"] or 0) >= 3 else ""
                self.ay_tree.insert("", tk.END, tags=(tag,),
                                    values=(r["ad"], r["toplam_gun"] or 0,
                                            r["gec_sayisi"] or 0,
                                            f"{r['ort_gec']:.0f}" if r["ort_gec"] else "—"))
        except Exception:
            pass

    def _personel_listesi(self):
        for item in self.p_tree.get_children():
            self.p_tree.delete(item)
        filtre = getattr(self, "p_filtre_var", None)
        filtre_val = filtre.get() if filtre else "tümü"
        try:
            conn = self._get_conn()
            sql = "SELECT id, ad, mesai_baslangic, mac_adresi, aktif, notlar FROM personeller"
            if filtre_val == "aktif":
                sql += " WHERE aktif=1"
            elif filtre_val == "pasif":
                sql += " WHERE aktif=0"
            sql += " ORDER BY aktif DESC, ad"
            rows = conn.execute(sql).fetchall()
            conn.close()
            for r in rows:
                mac_goster = r["mac_adresi"] or "—"
                durum = "🟢 Aktif" if r["aktif"] else "🔴 Pasif"
                tag = "aktif" if r["aktif"] else "pasif"
                self.p_tree.insert("", tk.END, iid=str(r["id"]), tags=(tag,),
                                   values=(r["ad"], r["mesai_baslangic"] or "09:00",
                                           durum, r["notlar"] or "", mac_goster))
        except Exception:
            pass

    def _personel_sil(self):
        sel = self.p_tree.selection()
        if not sel:
            return
        if messagebox.askyesno("Sil", "Bu personel silinsin mi?"):
            conn = self._get_conn()
            conn.execute("DELETE FROM personeller WHERE id=?", (int(sel[0]),))
            conn.commit(); conn.close()
            self._personel_listesi()

    def _p_select(self, _=None):
        sel = self.p_tree.selection()
        if not sel:
            return
        conn = self._get_conn()
        r = conn.execute("SELECT * FROM personeller WHERE id=?", (int(sel[0]),)).fetchone()
        conn.close()
        if r:
            self.p_ad_var.set(r["ad"])
            self.p_mesai_var.set(r["mesai_baslangic"] or "09:00")
            self.p_mac_var.set(r["mac_adresi"] or "")
            self.p_aktif_var.set("aktif" if r["aktif"] else "pasif")
            try:
                self.p_not_var.set(r["notlar"] or "")
            except Exception:
                pass

    def _aktif_degistir(self, aktif: int):
        sel = self.p_tree.selection()
        if not sel:
            messagebox.showwarning("Uyarı", "Listeden personel seçin!")
            return
        conn = self._get_conn()
        r = conn.execute("SELECT ad FROM personeller WHERE id=?", (int(sel[0]),)).fetchone()
        durum_str = "aktif" if aktif else "pasif"
        if not messagebox.askyesno("Onay", f"{r['ad']} → {durum_str} yapılsın mı?"):
            conn.close(); return
        conn.execute("UPDATE personeller SET aktif=? WHERE id=?", (aktif, int(sel[0])))
        conn.commit(); conn.close()
        self._personel_listesi()
        self._gunluk_rapor()

    def _p_temizle(self):
        self.p_ad_var.set("")
        self.p_mesai_var.set("09:00")
        self.p_mac_var.set("")
        self.p_aktif_var.set("aktif")
        try:
            self.p_not_var.set("")
        except Exception:
            pass
        self.p_tree.selection_remove(self.p_tree.selection())

    def _manuel_kaydet(self):
        ad    = self.m_personel_var.get()
        giris = self.m_giris_var.get().strip()
        cikis = self.m_cikis_var.get().strip()
        tarih = self._parse_tarih(self.tarih_var.get())
        if not ad:
            messagebox.showwarning("Uyarı", "Personel seçin!")
            return
        try:
            conn = self._get_conn()
            pid  = conn.execute(
                "SELECT id, mesai_baslangic FROM personeller WHERE ad=?", (ad,)
            ).fetchone()
            if not pid:
                return
            gec_kaldi = 0; gec_dk = 0
            if giris:
                from datetime import datetime as dt2, date as d2
                sinir_str = pid["mesai_baslangic"] or "09:00"
                sinir   = dt2.strptime(sinir_str, "%H:%M").time()
                giris_t = dt2.strptime(giris, "%H:%M").time()
                if giris_t > sinir:
                    fark = int((dt2.combine(d2.today(), giris_t) -
                                dt2.combine(d2.today(), sinir)).total_seconds() // 60)
                    if fark > 0:
                        gec_kaldi = 1; gec_dk = fark
            conn.execute("""
                INSERT INTO devam (personel_id, tarih, giris_saati, cikis_saati, gec_kaldi, gec_dakika)
                VALUES (?, ?, ?, ?, ?, ?)
                ON CONFLICT(personel_id, tarih) DO UPDATE SET
                    giris_saati=excluded.giris_saati, cikis_saati=excluded.cikis_saati,
                    gec_kaldi=excluded.gec_kaldi, gec_dakika=excluded.gec_dakika
            """, (pid["id"], tarih, giris or None, cikis or None, gec_kaldi, gec_dk))
            conn.commit(); conn.close()
            self._gunluk_rapor()
            messagebox.showinfo("Kaydedildi", f"{ad} — {tarih} güncellendi.")
        except Exception as e:
            messagebox.showerror("Hata", str(e))


# TAHSİLAT SEKMESİ
# ============================================================================

class TahsilatTab(ttk.Frame):
    """Tahsilat Yönetim Sekmesi."""

    def __init__(self, master):
        super().__init__(master, padding=10)
        self._all_customers: List[Dict] = []
        self._build_ui()
        self.refresh()

    def _build_ui(self):
        from datetime import date
        bugun = date.today().strftime("%d.%m.%Y")
        self._selected_customer_id: Optional[int] = None

        # ── Üst bar: Tarih filtresi + Özet ──
        top = ttk.Frame(self)
        top.pack(fill=tk.X, pady=(0, 6))

        ttk.Label(top, text="💰 Tahsilat", font=("Segoe UI", 13, "bold")).pack(side=tk.LEFT)
        ttk.Label(top, text="  Başlangıç:").pack(side=tk.LEFT, padx=(16, 2))
        self.bas_var = tk.StringVar(value=bugun)
        bas_e = ttk.Entry(top, textvariable=self.bas_var, width=12)
        bas_e.pack(side=tk.LEFT)
        bas_e.bind("<Button-1>", lambda e: self._open_datepicker(self.bas_var))

        ttk.Label(top, text="  Bitiş:").pack(side=tk.LEFT, padx=(8, 2))
        self.bit_var = tk.StringVar(value=bugun)
        bit_e = ttk.Entry(top, textvariable=self.bit_var, width=12)
        bit_e.pack(side=tk.LEFT)
        bit_e.bind("<Button-1>", lambda e: self._open_datepicker(self.bit_var))

        ttk.Button(top, text="🔍 Filtrele", command=self.refresh).pack(side=tk.LEFT, padx=(6, 0))
        ttk.Button(top, text="↺ Bugün",    command=self._reset_today).pack(side=tk.LEFT, padx=(4, 0))

        self.lbl_toplam = ttk.Label(top, text="📊 Toplam: 0,00 ₺", font=("Segoe UI", 11, "bold"), foreground="#f59e0b")
        self.lbl_toplam.pack(side=tk.RIGHT, padx=(0, 16))
        self.lbl_banka  = ttk.Label(top, text="🏦 Banka: 0,00 ₺",  font=("Segoe UI", 10, "bold"), foreground="#3498db")
        self.lbl_banka.pack(side=tk.RIGHT, padx=(0, 16))
        self.lbl_nakit  = ttk.Label(top, text="💵 Nakit: 0,00 ₺",  font=("Segoe UI", 10, "bold"), foreground="#2ecc71")
        self.lbl_nakit.pack(side=tk.RIGHT, padx=(0, 16))

        # ── Ana body: Sol (müşteriler) + Sağ (form + tahsilat listesi) ──
        body = ttk.Frame(self)
        body.pack(fill=tk.BOTH, expand=True)

        # ── SOL PANEL ──
        left = ttk.LabelFrame(body, text="⚠️ Borçlu Müşteriler", padding=6)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=False, padx=(0, 6))
        left.configure(width=480)

        # Arama
        arama_f = ttk.Frame(left)
        arama_f.pack(fill=tk.X, pady=(0, 4))
        ttk.Label(arama_f, text="🔍").pack(side=tk.LEFT)
        self.arama_var = tk.StringVar()
        self.arama_var.trace_add("write", lambda *a: self._load_customer_list())
        ttk.Entry(arama_f, textvariable=self.arama_var, width=28).pack(side=tk.LEFT, padx=(4, 0))

        # Müşteri treeview
        cust_cols = ("name", "bu_ay", "toplam_borc", "son_tahsilat")
        self.cust_tree = ttk.Treeview(left, columns=cust_cols, show="headings",
                                       selectmode="browse", height=16)
        ch = {"name": "Müşteri", "bu_ay": "Bu Ay (₺)", "toplam_borc": "Toplam Borç (₺)", "son_tahsilat": "Son Tahsilat"}
        cw = {"name": 160, "bu_ay": 80, "toplam_borc": 100, "son_tahsilat": 120}
        for c in cust_cols:
            self.cust_tree.heading(c, text=ch[c])
            self.cust_tree.column(c, width=cw[c], anchor="w" if c in ("name","son_tahsilat") else "e")
        cvsb = ttk.Scrollbar(left, orient="vertical", command=self.cust_tree.yview)
        self.cust_tree.configure(yscrollcommand=cvsb.set)
        self.cust_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        cvsb.pack(side=tk.LEFT, fill=tk.Y)
        self.cust_tree.tag_configure("borc", foreground="white")
        self.cust_tree.bind("<<TreeviewSelect>>", self._on_cust_select)

        # Müşteri toplam satırı
        self.cust_totals = ttk.Label(left, text="", font=("Segoe UI", 9, "bold"), foreground="#f59e0b")
        self.cust_totals.pack(fill=tk.X, pady=(2, 0))

        # ── SAĞ PANEL ──
        right = ttk.Frame(body)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Tahsilat formu
        form = ttk.LabelFrame(right, text="Yeni Tahsilat Girişi", padding=8)
        form.pack(fill=tk.X, pady=(0, 6))

        f0 = ttk.Frame(form); f0.pack(fill=tk.X, pady=2)
        ttk.Label(f0, text="Müşteri:", width=10, anchor="e").pack(side=tk.LEFT)
        self.sel_cust_label = ttk.Label(f0, text="— Sol listeden seçin —",
                                         font=("Segoe UI", 9, "bold"), foreground="#4fc3f7")
        self.sel_cust_label.pack(side=tk.LEFT, padx=(4, 0))

        f1 = ttk.Frame(form); f1.pack(fill=tk.X, pady=2)
        ttk.Label(f1, text="Tarih:", width=10, anchor="e").pack(side=tk.LEFT)
        self.f_tarih_var = tk.StringVar(value=bugun)
        tarih_e = ttk.Entry(f1, textvariable=self.f_tarih_var, width=14)
        tarih_e.pack(side=tk.LEFT, padx=(2, 4))
        tarih_e.bind("<Button-1>", lambda e: self._open_datepicker(self.f_tarih_var))
        ttk.Button(f1, text="📅", width=2, command=lambda: self._open_datepicker(self.f_tarih_var)).pack(side=tk.LEFT, padx=(0, 12))

        ttk.Label(f1, text="Tutar (₺):", anchor="e").pack(side=tk.LEFT)
        self.f_tutar_var = tk.StringVar()
        ttk.Entry(f1, textvariable=self.f_tutar_var, width=14, justify="right").pack(side=tk.LEFT, padx=(2, 12))

        ttk.Label(f1, text="Ödeme:", anchor="e").pack(side=tk.LEFT)
        self.f_odeme_var = tk.StringVar(value="N")
        ttk.Radiobutton(f1, text="💵 Nakit", variable=self.f_odeme_var, value="N").pack(side=tk.LEFT)
        ttk.Radiobutton(f1, text="🏦 Banka", variable=self.f_odeme_var, value="B").pack(side=tk.LEFT, padx=(8, 0))

        f2 = ttk.Frame(form); f2.pack(fill=tk.X, pady=2)
        ttk.Label(f2, text="Açıklama:", width=10, anchor="e").pack(side=tk.LEFT)
        self.f_aciklama_var = tk.StringVar()
        ttk.Entry(f2, textvariable=self.f_aciklama_var, width=50).pack(side=tk.LEFT, padx=(2, 0))

        btn_row = ttk.Frame(form); btn_row.pack(fill=tk.X, pady=(6, 0))
        ttk.Button(btn_row, text="💾 Tahsilat Kaydet", command=self._save_tahsilat).pack(side=tk.LEFT)
        ttk.Button(btn_row, text="✖ Temizle", command=self._clear_form).pack(side=tk.LEFT, padx=(6, 0))

        # Tahsilat listesi
        tah = ttk.LabelFrame(right, text="Tahsilat Kayıtları", padding=6)
        tah.pack(fill=tk.BOTH, expand=True)

        t_cols = ("tarih", "musteri", "tutar", "odeme", "aciklama")
        self.tah_tree = ttk.Treeview(tah, columns=t_cols, show="headings",
                                      selectmode="browse", height=12)
        th2 = {"tarih": "Tarih", "musteri": "Müşteri", "tutar": "Tutar (₺)", "odeme": "Ödeme", "aciklama": "Açıklama"}
        tw2 = {"tarih": 90, "musteri": 180, "tutar": 100, "odeme": 80, "aciklama": 200}
        for c in t_cols:
            self.tah_tree.heading(c, text=th2[c])
            self.tah_tree.column(c, width=tw2[c], anchor="w" if c in ("musteri","aciklama") else "center")
        tvsb = ttk.Scrollbar(tah, orient="vertical", command=self.tah_tree.yview)
        self.tah_tree.configure(yscrollcommand=tvsb.set)
        self.tah_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tvsb.pack(side=tk.LEFT, fill=tk.Y)
        self.tah_tree.tag_configure("banka", foreground="#3498db")
        self.tah_tree.tag_configure("nakit", foreground="#2ecc71")

        ttk.Button(tah, text="🗑 Seçili Kaydı Sil", command=self._delete_tahsilat).pack(pady=(4, 0))

    def _open_datepicker(self, target_var: tk.StringVar):
        """Tarih seçici popup."""
        import calendar
        from datetime import date

        try:
            p = target_var.get().strip().split(".")
            cur = date(int(p[2]), int(p[1]), int(p[0]))
        except Exception:
            cur = date.today()

        popup = tk.Toplevel(self)
        popup.title("Tarih Seç")
        popup.resizable(False, False)
        popup.grab_set()

        state = {"year": cur.year, "month": cur.month}

        header = ttk.Frame(popup, padding=4)
        header.pack(fill=tk.X)

        lbl_ym = ttk.Label(header, font=("Segoe UI", 10, "bold"), width=16, anchor="center")
        lbl_ym.pack(side=tk.LEFT, expand=True)

        cal_frame = ttk.Frame(popup, padding=6)
        cal_frame.pack()

        def render():
            for w in cal_frame.winfo_children():
                w.destroy()
            y, m = state["year"], state["month"]
            lbl_ym.config(text=f"{y}  {MONTHS_TR[m-1]}")
            days = ["Pt","Sa","Ça","Pe","Cu","Ct","Pa"]
            for i, d in enumerate(days):
                ttk.Label(cal_frame, text=d, width=4, anchor="center",
                          font=("Segoe UI", 8, "bold")).grid(row=0, column=i, padx=1, pady=1)
            first_day = calendar.weekday(y, m, 1)
            _, num_days = calendar.monthrange(y, m)
            row, col = 1, first_day
            for day in range(1, num_days + 1):
                d = day
                btn = tk.Button(cal_frame, text=str(d), width=4, height=1,
                                relief="flat", cursor="hand2",
                                bg="#0097a7" if date(y, m, d) == cur else "#243447",
                                fg="white",
                                command=lambda dd=d: pick(dd))
                btn.grid(row=row, column=col, padx=1, pady=1)
                col += 1
                if col > 6:
                    col = 0
                    row += 1

        def pick(day):
            y, m = state["year"], state["month"]
            target_var.set(f"{day:02d}.{m:02d}.{y}")
            popup.destroy()

        def prev_month():
            if state["month"] == 1:
                state["month"] = 12; state["year"] -= 1
            else:
                state["month"] -= 1
            render()

        def next_month():
            if state["month"] == 12:
                state["month"] = 1; state["year"] += 1
            else:
                state["month"] += 1
            render()

        def prev_year():
            state["year"] -= 1; render()

        def next_year():
            state["year"] += 1; render()

        ttk.Button(header, text="◀◀", width=3, command=prev_year).pack(side=tk.LEFT)
        ttk.Button(header, text="◀",  width=3, command=prev_month).pack(side=tk.LEFT)
        ttk.Button(header, text="▶",  width=3, command=next_month).pack(side=tk.RIGHT)
        ttk.Button(header, text="▶▶", width=3, command=next_year).pack(side=tk.RIGHT)

        render()

    def _reset_today(self):
        from datetime import date
        bugun = date.today().strftime("%d.%m.%Y")
        self.bas_var.set(bugun)
        self.bit_var.set(bugun)
        self.refresh()

    def _parse_date(self, dstr: str) -> str:
        """GG.AA.YYYY → YYYY-MM-DD"""
        try:
            p = dstr.strip().split(".")
            return f"{p[2]}-{p[1]:0>2}-{p[0]:0>2}"
        except Exception:
            from datetime import date
            return date.today().strftime("%Y-%m-%d")

    def refresh(self):
        self._load_customer_list()
        self._load_tahsilat_list()
        self._update_summary()

    def _load_customer_list(self):
        from database import fetch_all
        arama = self.arama_var.get().strip().upper()

        for item in self.cust_tree.get_children():
            self.cust_tree.delete(item)

        customers = fetch_all("SELECT id, name FROM customers ORDER BY name")
        self._all_customers = list(customers)

        toplam_bu_ay = 0.0
        toplam_borc  = 0.0
        count = 0

        for c in customers:
            cid = c["id"]
            if arama and arama not in c["name"].upper():
                continue

            bu_ay  = get_musteri_bu_ay_borc(cid)
            toplam = get_musteri_toplam_borc(cid)

            # Sadece borcu olanları göster
            if toplam <= 0 and bu_ay <= 0:
                continue

            son = fetch_all(
                "SELECT tahsilat_tarihi, tutar, odeme_turu FROM tahsilatlar WHERE customer_id=? ORDER BY tahsilat_tarihi DESC, id DESC LIMIT 1",
                (cid,)
            )
            son_str = ""
            if son:
                try:
                    p = son[0]["tahsilat_tarihi"].split("-")
                    t_str = f"{p[2]}.{p[1]}.{p[0]}"
                except Exception:
                    t_str = son[0]["tahsilat_tarihi"]
                tur = "B" if son[0]["odeme_turu"] == "B" else "N"
                son_str = f"{t_str} {float(son[0]['tutar']):,.0f}₺({tur})"

            self.cust_tree.insert("", tk.END, iid=str(cid), tags=("borc",),
                                   values=(c["name"],
                                           f"{bu_ay:,.2f}" if bu_ay > 0 else "—",
                                           f"{toplam:,.2f}",
                                           son_str))
            toplam_bu_ay += bu_ay
            toplam_borc  += toplam
            count += 1

        self.cust_totals.config(
            text=f"  {count} borçlu müşteri   |   Bu Ay: {toplam_bu_ay:,.2f} ₺   |   Toplam: {toplam_borc:,.2f} ₺"
        )

    def _load_tahsilat_list(self):
        for item in self.tah_tree.get_children():
            self.tah_tree.delete(item)

        bas = self._parse_date(self.bas_var.get())
        bit = self._parse_date(self.bit_var.get())
        kayitlar = get_tahsilatlar(bas, bit)

        for k in kayitlar:
            tag = "banka" if k["odeme_turu"] == "B" else "nakit"
            # Tarihi GG.MM.YYYY formatında göster
            try:
                p = k["tahsilat_tarihi"].split("-")
                tarih_str = f"{p[2]}.{p[1]}.{p[0]}"
            except Exception:
                tarih_str = k["tahsilat_tarihi"]

            self.tah_tree.insert("", tk.END, iid=str(k["id"]), tags=(tag,),
                                  values=(tarih_str,
                                          k["customer_name"],
                                          f"{float(k['tutar']):,.2f}",
                                          "🏦 Banka" if k["odeme_turu"] == "B" else "💵 Nakit",
                                          k["aciklama"] or ""))

    def _update_summary(self):
        bas = self._parse_date(self.bas_var.get())
        bit = self._parse_date(self.bit_var.get())
        t = get_tahsilat_toplam(bas, bit)
        self.lbl_nakit.config(text=f"💵 Nakit: {t['N']:,.2f} ₺")
        self.lbl_banka.config(text=f"🏦 Banka: {t['B']:,.2f} ₺")
        self.lbl_toplam.config(text=f"📊 Toplam: {t['toplam']:,.2f} ₺")

    def _on_cust_select(self, _=None):
        sel = self.cust_tree.selection()
        if not sel:
            return
        self._selected_customer_id = int(sel[0])
        name = self.cust_tree.item(sel[0], "values")[0]
        self.sel_cust_label.config(text=name)

    def _save_tahsilat(self):
        if not self._selected_customer_id:
            messagebox.showwarning("Uyarı", "Önce sol listeden bir müşteri seçin!")
            return
        try:
            tutar_str = self.f_tutar_var.get().replace(",", ".").replace(" ", "")
            tutar = float(tutar_str)
            if tutar <= 0:
                raise ValueError("Tutar 0'dan büyük olmalı")
        except ValueError as e:
            messagebox.showerror("Hata", f"Geçersiz tutar: {e}")
            return

        tarih_iso = self._parse_date(self.f_tarih_var.get())
        odeme = self.f_odeme_var.get()
        aciklama = self.f_aciklama_var.get().strip()

        insert_tahsilat(self._selected_customer_id, tutar, odeme, tarih_iso, aciklama)
        self.refresh()
        self._clear_form()
        messagebox.showinfo("Kaydedildi", f"Tahsilat kaydedildi: {tutar:,.2f} ₺ ({odeme})")

    def _delete_tahsilat(self):
        sel = self.tah_tree.selection()
        if not sel:
            return
        if messagebox.askyesno("Sil", "Bu tahsilat kaydı silinsin mi?"):
            delete_tahsilat(int(sel[0]))
            self.refresh()

    def _clear_form(self):
        from datetime import date
        self.f_tarih_var.set(date.today().strftime("%d.%m.%Y"))
        self.f_tutar_var.set("")
        self.f_odeme_var.set("N")
        self.f_aciklama_var.set("")


# ============================================================================
# OFİS SEKMESİ
# ============================================================================

class OfficeTab(ttk.Frame):
    """Ofis / Stok Yönetim Sekmesi."""

    TYPE_COLORS = {
        "Hazır Ofis":      "#1a6b2a",
        "Paylaşımlı Masa": "#1a3a6b",
        "Sanal Ofis":      "#6b1a5a",
    }

    def __init__(self, master):
        super().__init__(master, padding=10)
        self._all_offices: List[Dict] = []
        self._build_ui()
        self.refresh()

    def _build_ui(self):
        from database import get_all_customers_with_rent_progression

        # ── Üst bar ──
        top = ttk.Frame(self)
        top.pack(fill=tk.X, pady=(0, 8))

        ttk.Label(top, text="🏢 Ofis Yönetimi", font=("Segoe UI", 13, "bold")).pack(side=tk.LEFT)
        ttk.Button(top, text="↺ Yenile", command=self.refresh).pack(side=tk.LEFT, padx=(10, 0))
        ttk.Button(top, text="➕ Yeni Ofis", command=self._new_office).pack(side=tk.LEFT, padx=(6, 0))

        # Filtre
        ttk.Label(top, text="  Tür:").pack(side=tk.LEFT, padx=(16, 2))
        self.filter_type_var = tk.StringVar(value="Tümü")
        ttk.Combobox(top, textvariable=self.filter_type_var, state="readonly", width=16,
                     values=["Tümü", "Hazır Ofis", "Paylaşımlı Masa", "Sanal Ofis"]
                     ).pack(side=tk.LEFT)
        self.filter_type_var.trace_add("write", lambda *a: self._apply_filter())

        ttk.Label(top, text="  Durum:").pack(side=tk.LEFT, padx=(8, 2))
        self.filter_status_var = tk.StringVar(value="Tümü")
        ttk.Combobox(top, textvariable=self.filter_status_var, state="readonly", width=8,
                     values=["Tümü", "bos", "dolu"]
                     ).pack(side=tk.LEFT)
        self.filter_status_var.trace_add("write", lambda *a: self._apply_filter())

        ttk.Label(top, text="  Aktif:").pack(side=tk.LEFT, padx=(8, 2))
        self.filter_aktif_var = tk.StringVar(value="aktif")
        ttk.Combobox(top, textvariable=self.filter_aktif_var, state="readonly", width=8,
                     values=["Tümü", "aktif", "pasif"]
                     ).pack(side=tk.LEFT)
        self.filter_aktif_var.trace_add("write", lambda *a: self._apply_filter())

        # ── Özet kartlar ──
        self.summary_frame = ttk.Frame(self)
        self.summary_frame.pack(fill=tk.X, pady=(0, 8))

        # ── Treeview ──
        tree_frame = ttk.Frame(self)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        cols = ("code", "type", "unit_no", "monthly_price", "status", "aktif_pasif", "customer_name", "notes")
        self.tree = ttk.Treeview(tree_frame, columns=cols, show="headings",
                                 selectmode="browse", height=18)

        headers = {"code": "Kod", "type": "Tür", "unit_no": "No",
                   "monthly_price": "Aylık Kira (₺)", "status": "Durum",
                   "aktif_pasif": "Aktif/Pasif",
                   "customer_name": "Müşteri", "notes": "Not"}
        widths   = {"code": 90, "type": 130, "unit_no": 60,
                    "monthly_price": 110, "status": 70,
                    "aktif_pasif": 80,
                    "customer_name": 200, "notes": 180}

        for col in cols:
            self.tree.heading(col, text=headers[col])
            self.tree.column(col, width=widths[col],
                             anchor="w" if col in ("type","customer_name","notes") else "center")

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        tree_frame.rowconfigure(0, weight=1)
        tree_frame.columnconfigure(0, weight=1)

        # Tag renkleri
        self.tree.tag_configure("dolu", foreground="#e74c3c")
        self.tree.tag_configure("bos",  foreground="#2ecc71")

        self.tree.bind("<<TreeviewSelect>>", self._on_select)
        self.tree.bind("<Double-1>", self._on_double_click)

        # ── Alt form ──
        form = ttk.LabelFrame(self, text="Ofis Detayı / Düzenle", padding=8)
        form.pack(fill=tk.X, pady=(8, 0))

        r0 = ttk.Frame(form); r0.pack(fill=tk.X, pady=2)
        ttk.Label(r0, text="Kod:", width=10, anchor="e").pack(side=tk.LEFT)
        self.f_code = ttk.Entry(r0, width=12); self.f_code.pack(side=tk.LEFT, padx=(2,12))

        ttk.Label(r0, text="Tür:", width=6, anchor="e").pack(side=tk.LEFT)
        self.f_type_var = tk.StringVar()
        ttk.Combobox(r0, textvariable=self.f_type_var, width=16, state="readonly",
                     values=["Hazır Ofis","Paylaşımlı Masa","Sanal Ofis"]
                     ).pack(side=tk.LEFT, padx=(2,12))

        ttk.Label(r0, text="No:", width=4, anchor="e").pack(side=tk.LEFT)
        self.f_unit = ttk.Entry(r0, width=8); self.f_unit.pack(side=tk.LEFT, padx=(2,12))

        ttk.Label(r0, text="Aylık Kira:", anchor="e").pack(side=tk.LEFT)
        self.f_price = ttk.Entry(r0, width=12, justify="right"); self.f_price.pack(side=tk.LEFT, padx=(2,12))

        ttk.Label(r0, text="Durum:", anchor="e").pack(side=tk.LEFT)
        self.f_status_var = tk.StringVar(value="bos")
        ttk.Combobox(r0, textvariable=self.f_status_var, width=8, state="readonly",
                     values=["bos","dolu"]).pack(side=tk.LEFT, padx=(2,12))

        ttk.Label(r0, text="Aktif/Pasif:", anchor="e").pack(side=tk.LEFT)
        self.f_aktif_var = tk.StringVar(value="aktif")
        ttk.Combobox(r0, textvariable=self.f_aktif_var, width=8, state="readonly",
                     values=["aktif","pasif"]).pack(side=tk.LEFT, padx=(2,0))

        r1 = ttk.Frame(form); r1.pack(fill=tk.X, pady=2)
        ttk.Label(r1, text="Müşteri:", width=10, anchor="e").pack(side=tk.LEFT)
        self.f_customer_var = tk.StringVar()
        self.f_customer_cb = ttk.Combobox(r1, textvariable=self.f_customer_var,
                                           width=35, state="readonly")
        self.f_customer_cb.pack(side=tk.LEFT, padx=(2,12))

        ttk.Label(r1, text="Not:", anchor="e").pack(side=tk.LEFT)
        self.f_notes = ttk.Entry(r1, width=40); self.f_notes.pack(side=tk.LEFT, padx=(2,0))

        btn_row = ttk.Frame(form); btn_row.pack(fill=tk.X, pady=(6,0))
        ttk.Button(btn_row, text="💾 Kaydet",   command=self._save).pack(side=tk.LEFT)
        ttk.Button(btn_row, text="🗑 Sil",      command=self._delete).pack(side=tk.LEFT, padx=(6,0))
        ttk.Button(btn_row, text="✖ Temizle",   command=self._clear_form).pack(side=tk.LEFT, padx=(6,0))
        ttk.Button(btn_row, text="🔓 Boşalt",   command=self._vacate).pack(side=tk.LEFT, padx=(6,0))

        self._load_customer_list()

    def _load_customer_list(self):
        from database import fetch_all
        rows = fetch_all("SELECT id, name FROM customers ORDER BY name")
        self._customer_map = {r["name"]: r["id"] for r in rows}
        self._customer_map["— Boş —"] = None
        names = ["— Boş —"] + [r["name"] for r in rows]
        self.f_customer_cb["values"] = names

    def refresh(self):
        from database import get_all_offices, get_office_summary, initialize_offices
        initialize_offices()
        self._all_offices = get_all_offices()
        self._update_summary()
        self._apply_filter()
        self._load_customer_list()

    def _update_summary(self):
        from database import get_office_summary
        for w in self.summary_frame.winfo_children():
            w.destroy()
        summary = get_office_summary()
        types = [("Hazır Ofis","🏢","#1a6b2a"),
                 ("Paylaşımlı Masa","🪑","#1a3a6b"),
                 ("Sanal Ofis","🌐","#6b1a5a")]
        for type_, icon, color in types:
            s = summary.get(type_, {"toplam":0,"dolu":0,"bos":0})
            card = tk.Frame(self.summary_frame, bg=color, padx=12, pady=6)
            card.pack(side=tk.LEFT, padx=(0,8))
            tk.Label(card, text=f"{icon} {type_}", bg=color, fg="white",
                     font=("Segoe UI",9,"bold")).pack()
            tk.Label(card, text=f"Toplam: {s['toplam']}  Dolu: {s['dolu']}  Boş: {s['bos']}",
                     bg=color, fg="#d0f0ff", font=("Segoe UI",9)).pack()

    def _apply_filter(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        ft = self.filter_type_var.get()
        fs = self.filter_status_var.get()
        fa = self.filter_aktif_var.get()
        for o in self._all_offices:
            if ft != "Tümü" and o["type"] != ft:
                continue
            if fs != "Tümü" and o["status"] != fs:
                continue
            is_active = o.get("is_active", 1)
            aktif_str = "aktif" if is_active else "pasif"
            if fa != "Tümü" and aktif_str != fa:
                continue
            tag = o["status"]
            self.tree.insert("", tk.END, iid=str(o["id"]), tags=(tag,),
                             values=(o["code"], o["type"], o["unit_no"] or "",
                                     f"{float(o['monthly_price'] or 0):,.2f}",
                                     "✅ Dolu" if o["status"]=="dolu" else "⬜ Boş",
                                     "🟢 Aktif" if is_active else "🔴 Pasif",
                                     o["customer_name"] or "", o["notes"] or ""))

    def _on_select(self, _=None):
        sel = self.tree.selection()
        if not sel:
            return
        iid = sel[0]
        o = next((x for x in self._all_offices if str(x["id"]) == iid), None)
        if not o:
            return
        self.f_code.delete(0, tk.END); self.f_code.insert(0, o["code"])
        self.f_type_var.set(o["type"])
        self.f_unit.delete(0, tk.END); self.f_unit.insert(0, o["unit_no"] or "")
        self.f_price.delete(0, tk.END); self.f_price.insert(0, str(o["monthly_price"] or ""))
        self.f_status_var.set(o["status"])
        self.f_aktif_var.set("aktif" if o.get("is_active", 1) else "pasif")
        self.f_customer_var.set(o["customer_name"] or "— Boş —")
        self.f_notes.delete(0, tk.END); self.f_notes.insert(0, o["notes"] or "")

    def _on_double_click(self, _=None):
        self._on_select()

    def _new_office(self):
        from database import get_next_office_code
        self._clear_form()
        type_ = self.f_type_var.get() or "Sanal Ofis"
        self.f_type_var.set(type_)
        code = get_next_office_code(type_)
        self.f_code.insert(0, code)

    def _save(self):
        from database import save_office, assign_office_to_customer
        code   = self.f_code.get().strip()
        type_  = self.f_type_var.get()
        unit   = self.f_unit.get().strip()
        status = self.f_status_var.get()
        is_active = 1 if self.f_aktif_var.get() == "aktif" else 0
        notes  = self.f_notes.get().strip()
        cname  = self.f_customer_var.get()
        cid    = self._customer_map.get(cname)

        try:
            price = float(self.f_price.get().replace(",","").replace(" ","")) if self.f_price.get() else 0
        except ValueError:
            messagebox.showerror("Hata", "Aylık kira sayı olmalı!")
            return

        if not code or not type_:
            messagebox.showerror("Hata", "Kod ve tür zorunludur!")
            return

        save_office(code, type_, unit, price, status, is_active, cid, notes)
        self.refresh()
        messagebox.showinfo("Kaydedildi", f"{code} kaydedildi!")

    def _delete(self):
        sel = self.tree.selection()
        if not sel:
            return
        iid = sel[0]
        o = next((x for x in self._all_offices if str(x["id"]) == iid), None)
        if not o:
            return
        if messagebox.askyesno("Sil", f"{o['code']} silinsin mi?"):
            from database import delete_office
            delete_office(o["code"])
            self._clear_form()
            self.refresh()

    def _vacate(self):
        sel = self.tree.selection()
        if not sel:
            return
        iid = sel[0]
        o = next((x for x in self._all_offices if str(x["id"]) == iid), None)
        if not o:
            return
        from database import assign_office_to_customer
        assign_office_to_customer(o["code"], None)
        self.refresh()

    def _clear_form(self):
        self.f_code.delete(0, tk.END)
        self.f_unit.delete(0, tk.END)
        self.f_price.delete(0, tk.END)
        self.f_notes.delete(0, tk.END)
        self.f_type_var.set("Hazır Ofis")
        self.f_status_var.set("bos")
        self.f_aktif_var.set("aktif")
        self.f_customer_var.set("— Boş —")


class ProductTab(ttk.Frame):
    """Ürün Yönetimi Sekmesi — CRUD dahil."""

    def __init__(self, master: tk.Widget) -> None:
        super().__init__(master, padding=10)
        self.selected_id: Optional[int] = None
        self._build_ui()
        self.refresh()

    def _build_ui(self) -> None:
        top = ttk.Frame(self)
        top.pack(side=tk.TOP, fill=tk.X, pady=(0, 8))
        ttk.Label(top, text="Ürünler", font=("Segoe UI", 13, "bold")).pack(side=tk.LEFT)
        ttk.Button(top, text="↺ Yenile", command=self.refresh).pack(side=tk.LEFT, padx=(10, 0))

        # Treeview
        tree_frame = ttk.Frame(self)
        tree_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        cols = ("name", "sku", "unit_price", "stock_quantity")
        self.tree = ttk.Treeview(tree_frame, columns=cols, show="headings", selectmode="browse")
        headers = {"name": "Ürün Adı", "sku": "Stok Kodu", "unit_price": "Birim Fiyat (₺)", "stock_quantity": "Stok"}
        for col in cols:
            self.tree.heading(col, text=headers[col])
            self.tree.column(col, width=160, anchor="w")

        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.bind("<<TreeviewSelect>>", self._on_row_select)

        # Form
        form_frame = ttk.LabelFrame(self, text="Ürün Detayları", padding=10)
        form_frame.pack(side=tk.TOP, fill=tk.X, pady=(8, 0))

        self.form: Dict[str, FormField] = {}
        for key, label in [("name", "Ürün Adı:"), ("sku", "Stok Kodu:"),
                            ("unit_price", "Birim Fiyat (₺):"), ("stock_quantity", "Stok Miktarı:")]:
            f = FormField(form_frame, label=label)
            f.pack(fill=tk.X, pady=2)
            self.form[key] = f

        btn = ttk.Frame(form_frame)
        btn.pack(fill=tk.X, pady=(8, 0))
        ttk.Button(btn, text="💾 Ekle / Güncelle", command=self._on_save).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(btn, text="🗑 Sil",             command=self._on_delete).pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(btn, text="✖ Temizle",          command=self._clear_form).pack(side=tk.LEFT)

    def refresh(self) -> None:
        for item in self.tree.get_children():
            self.tree.delete(item)
        for p in get_all_products():
            self.tree.insert("", tk.END, iid=str(p["id"]), values=(
                p["name"], p["sku"] or "",
                f"{p['unit_price']:,.2f}", f"{p['stock_quantity']:,.2f}",
            ))

    def _on_row_select(self, event=None) -> None:
        sel = self.tree.selection()
        if not sel:
            return
        self.selected_id = int(sel[0])
        vals = self.tree.item(sel[0], "values")
        keys = ["name", "sku", "unit_price", "stock_quantity"]
        for k, v in zip(keys, vals):
            self.form[k].set(v)

    def _clear_form(self) -> None:
        for f in self.form.values():
            f.clear()
        if self.tree.selection():
            self.tree.selection_remove(self.tree.selection())
        self.selected_id = None

    def _on_save(self) -> None:
        try:
            name = self.form["name"].get()
            if not name:
                messagebox.showwarning("Uyarı", "Ürün adı zorunludur!")
                return
            sku   = self.form["sku"].get()
            price = float(self.form["unit_price"].get().replace(",", ".") or "0")
            stock = float(self.form["stock_quantity"].get().replace(",", ".") or "0")

            if self.selected_id is None:
                insert_product(name=name, sku=sku, unit_price=price, stock_quantity=stock)
                messagebox.showinfo("Başarılı", "Ürün eklendi.")
            else:
                update_product(self.selected_id, name=name, sku=sku, unit_price=price, stock_quantity=stock)
                messagebox.showinfo("Başarılı", "Ürün güncellendi.")

            self.refresh()
            self._clear_form()
        except ValueError as e:
            messagebox.showerror("Hata", f"Geçersiz değer: {e}")

    def _on_delete(self) -> None:
        if not self.selected_id:
            messagebox.showwarning("Uyarı", "Silmek için bir ürün seçin!")
            return
        if not messagebox.askyesno("Onay", "Ürün silinecek. Emin misiniz?"):
            return
        delete_product(self.selected_id)
        messagebox.showinfo("Başarılı", "Ürün silindi.")
        self.refresh()
        self._clear_form()


# ============================================================================
# FATURA SEKMESİ
# ============================================================================

# ============================================================================
# FATURA SEKMESİ
# ============================================================================

class InvoiceTab(ttk.Frame):
    """E-Arşiv Fatura Yönetimi — Oluştur, Kes, Cari Döküm, Tahsilat."""

    def __init__(self, master):
        super().__init__(master)
        self._kalemler = []
        self._secili_fatura_id = None
        self._musteri_map = {}
        self._build_ui()
        self.after(100, self._gecikmeli_refresh)

    def _gecikmeli_refresh(self):
        """UI tamamen yüklendikten sonra veri çek."""
        try:
            self.refresh()
        except Exception as e:
            import traceback
            print(f"[Fatura Refresh] {e}")
            traceback.print_exc()

    # ──────────────────────────────────────────────────────────────────────────
    def _build_ui(self):
        from datetime import date

        self.fatura_nb = ttk.Notebook(self)
        self.fatura_nb.pack(fill=tk.BOTH, expand=True)

        f1 = ttk.Frame(self.fatura_nb)
        f2 = ttk.Frame(self.fatura_nb)
        f3 = ttk.Frame(self.fatura_nb)
        f4 = ttk.Frame(self.fatura_nb)
        f5 = ttk.Frame(self.fatura_nb)

        self.fatura_nb.add(f5, text="  Aylik Faturalama  ")
        self.fatura_nb.add(f1, text="  Fatura Olustur  ")
        self.fatura_nb.add(f2, text="  Fatura Listesi  ")
        self.fatura_nb.add(f3, text="  Cari Dokum  ")
        self.fatura_nb.add(f4, text="  Firma Ayarlari  ")

        self._build_aylik(f5)
        self._build_olustur(f1)
        self._build_liste(f2)
        self._build_cari(f3)
        self._build_ayarlar(f4)

    # ── SEKME 1: Fatura Oluştur ───────────────────────────────────────────────
    def _build_olustur(self, parent):
        from datetime import date

        # Üst: sol form + sağ kalem listesi
        body = ttk.Frame(parent)
        body.pack(fill=tk.BOTH, expand=True)

        # SOL: Fatura başlık bilgileri
        left = ttk.LabelFrame(body, text=" Fatura Bilgileri", padding=8)
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 6))

        def row(parent, label, var, width=22, combo_vals=None, readonly=False):
            f = ttk.Frame(parent); f.pack(fill=tk.X, pady=2)
            ttk.Label(f, text=label, width=16, anchor="e").pack(side=tk.LEFT)
            if combo_vals is not None:
                w = ttk.Combobox(f, textvariable=var, width=width,
                                  values=combo_vals, state="readonly")
            elif readonly:
                w = ttk.Entry(f, textvariable=var, width=width, state="readonly")
            else:
                w = ttk.Entry(f, textvariable=var, width=width)
            w.pack(side=tk.LEFT, padx=(4, 0))
            return w

        # Fatura No (otomatik)
        self.v_fatura_no  = tk.StringVar()
        self.v_musteri    = tk.StringVar()
        self.v_musteri_vkn= tk.StringVar()
        self.v_musteri_adr= tk.StringVar()
        self.v_tarih      = tk.StringVar(value=date.today().strftime("%d.%m.%Y"))
        self.v_vade       = tk.StringVar()
        self.v_tur        = tk.StringVar(value="SATIŞ")
        self.v_not        = tk.StringVar()

        fn_f = ttk.Frame(left); fn_f.pack(fill=tk.X, pady=2)
        ttk.Label(fn_f, text="Fatura No:", width=16, anchor="e").pack(side=tk.LEFT)
        ttk.Entry(fn_f, textvariable=self.v_fatura_no, width=20,
                  state="readonly").pack(side=tk.LEFT, padx=(4,4))
        ttk.Button(fn_f, text=" Yeni No", command=self._yeni_fatura_no).pack(side=tk.LEFT)

        # Müşteri seçimi
        mc_f = ttk.Frame(left); mc_f.pack(fill=tk.X, pady=2)
        ttk.Label(mc_f, text="Müşteri:", width=16, anchor="e").pack(side=tk.LEFT)
        self.musteri_cb = ttk.Combobox(mc_f, textvariable=self.v_musteri,
                                        width=28, state="readonly")
        self.musteri_cb.pack(side=tk.LEFT, padx=(4,4))
        self.musteri_cb.bind("<<ComboboxSelected>>", self._musteri_sec)

        row(left, "Müşteri VKN/TC:", self.v_musteri_vkn, 20)
        row(left, "Müşteri Adres:", self.v_musteri_adr, 35)
        row(left, "Fatura Tarihi:", self.v_tarih, 12)
        row(left, "Vade Tarihi:", self.v_vade, 12)
        row(left, "Fatura Türü:", self.v_tur, 14,
            combo_vals=["SATIŞ", "İADE", "TEVKİFATLI"])
        row(left, "Not / Açıklama:", self.v_not, 35)

        # SAĞ: Kalem girişi + listesi
        right = ttk.Frame(body)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        kalem_form = ttk.LabelFrame(right, text="+ Kalem Ekle", padding=6)
        kalem_form.pack(fill=tk.X, pady=(0, 4))

        k1 = ttk.Frame(kalem_form); k1.pack(fill=tk.X, pady=2)
        ttk.Label(k1, text="Açıklama:", width=12, anchor="e").pack(side=tk.LEFT)
        self.v_k_aciklama = tk.StringVar()
        ttk.Entry(k1, textvariable=self.v_k_aciklama, width=32).pack(side=tk.LEFT, padx=(4,0))

        k2 = ttk.Frame(kalem_form); k2.pack(fill=tk.X, pady=2)
        ttk.Label(k2, text="Miktar:", width=12, anchor="e").pack(side=tk.LEFT)
        self.v_k_miktar = tk.StringVar(value="1")
        ttk.Entry(k2, textvariable=self.v_k_miktar, width=7).pack(side=tk.LEFT, padx=(4,8))
        ttk.Label(k2, text="Birim:").pack(side=tk.LEFT)
        self.v_k_birim = tk.StringVar(value="Ay")
        ttk.Combobox(k2, textvariable=self.v_k_birim, width=7,
                     values=["Ay","Adet","Hizmet","m2","Saat","Gün"]).pack(side=tk.LEFT, padx=(4,8))
        ttk.Label(k2, text="Birim Fiyat:").pack(side=tk.LEFT)
        self.v_k_fiyat = tk.StringVar(value="0")
        ttk.Entry(k2, textvariable=self.v_k_fiyat, width=10).pack(side=tk.LEFT, padx=(4,0))

        k3 = ttk.Frame(kalem_form); k3.pack(fill=tk.X, pady=2)
        ttk.Label(k3, text="KDV Oranı:", width=12, anchor="e").pack(side=tk.LEFT)
        self.v_k_kdv = tk.StringVar(value="20")
        ttk.Combobox(k3, textvariable=self.v_k_kdv, width=5,
                     values=["0","10","20"], state="readonly").pack(side=tk.LEFT, padx=(4,8))
        ttk.Label(k3, text="İskonto %:").pack(side=tk.LEFT)
        self.v_k_iskonto = tk.StringVar(value="0")
        ttk.Entry(k3, textvariable=self.v_k_iskonto, width=5).pack(side=tk.LEFT, padx=(4,8))
        ttk.Button(k3, text="+ Kalemi Ekle", command=self._kalem_ekle).pack(side=tk.LEFT)
        ttk.Button(k3, text="x Sil", command=self._kalem_sil).pack(side=tk.LEFT, padx=(4,0))

        # Kalem listesi
        k_cols = ("aciklama","miktar","birim","fiyat","iskonto","kdv","toplam")
        self.kalem_tree = ttk.Treeview(right, columns=k_cols, show="headings", height=7)
        k_hdrs = {"aciklama":"Açıklama","miktar":"Miktar","birim":"Birim",
                  "fiyat":"Birim Fiyat","iskonto":"İskonto%","kdv":"KDV%","toplam":"Toplam"}
        k_wdts = {"aciklama":200,"miktar":55,"birim":55,"fiyat":90,"iskonto":65,"kdv":55,"toplam":90}
        for c in k_cols:
            self.kalem_tree.heading(c, text=k_hdrs[c])
            self.kalem_tree.column(c, width=k_wdts[c],
                                   anchor="w" if c=="aciklama" else "center")
        k_vsb = ttk.Scrollbar(right, orient="vertical", command=self.kalem_tree.yview)
        self.kalem_tree.configure(yscrollcommand=k_vsb.set)
        self.kalem_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        k_vsb.pack(side=tk.LEFT, fill=tk.Y)

        # Alt: Toplam özet + butonlar
        alt = ttk.Frame(parent)
        alt.pack(fill=tk.X, pady=(6,0))

        # Toplamlar
        top_f = ttk.LabelFrame(alt, text=" Fatura Özeti", padding=6)
        top_f.pack(side=tk.LEFT, fill=tk.X, expand=True)
        self.lbl_matrah  = ttk.Label(top_f, text="Matrah: 0,00 ₺", font=("Segoe UI",9))
        self.lbl_matrah.pack(side=tk.LEFT, padx=10)
        self.lbl_iskonto = ttk.Label(top_f, text="İskonto: 0,00 ₺", font=("Segoe UI",9))
        self.lbl_iskonto.pack(side=tk.LEFT, padx=10)
        self.lbl_kdv     = ttk.Label(top_f, text="KDV: 0,00 ₺", font=("Segoe UI",9))
        self.lbl_kdv.pack(side=tk.LEFT, padx=10)
        self.lbl_toplam  = ttk.Label(top_f, text="TOPLAM: 0,00 ₺",
                                      font=("Segoe UI",11,"bold"), foreground="#4fc3f7")
        self.lbl_toplam.pack(side=tk.LEFT, padx=16)

        # Butonlar
        btn_f = ttk.Frame(alt)
        btn_f.pack(side=tk.RIGHT)
        ttk.Button(btn_f, text=" Taslak Kaydet",
                   command=self._taslak_kaydet).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_f, text=" PDF Önizle",
                   command=self._pdf_onizle).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_f, text=" FATURA KES",
                   command=self._fatura_kes).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_f, text=" Temizle",
                   command=self._form_temizle).pack(side=tk.LEFT, padx=4)

        self._yeni_fatura_no()

    # ── SEKME 2: Fatura Listesi ───────────────────────────────────────────────
    def _build_liste(self, parent):
        # Filtreler
        filtre_f = ttk.Frame(parent); filtre_f.pack(fill=tk.X, pady=(0,6))
        ttk.Label(filtre_f, text="Durum:").pack(side=tk.LEFT)
        self.v_filtre = tk.StringVar(value="tümü")
        for val, txt in [("tümü","Tümü"),("taslak","📝 Taslak"),
                         ("kesildi"," Kesildi"),("iptal"," İptal")]:
            ttk.Radiobutton(filtre_f, text=txt, variable=self.v_filtre,
                            value=val, command=self._liste_yukle).pack(side=tk.LEFT, padx=(6,0))
        ttk.Separator(filtre_f, orient="vertical").pack(side=tk.LEFT, fill=tk.Y, padx=8)
        ttk.Label(filtre_f, text="Başlangıç:").pack(side=tk.LEFT)
        self.v_bas_tarih = tk.StringVar()
        ttk.Entry(filtre_f, textvariable=self.v_bas_tarih, width=12).pack(side=tk.LEFT, padx=(2,8))
        ttk.Label(filtre_f, text="Bitiş:").pack(side=tk.LEFT)
        self.v_bit_tarih = tk.StringVar()
        ttk.Entry(filtre_f, textvariable=self.v_bit_tarih, width=12).pack(side=tk.LEFT, padx=(2,8))
        ttk.Button(filtre_f, text=" Filtrele", command=self._liste_yukle).pack(side=tk.LEFT)

        # Liste
        l_cols = ("no","musteri","tarih","toplam","durum","tahsilat","kalan")
        self.liste_tree = ttk.Treeview(parent, columns=l_cols, show="headings", height=14)
        l_hdrs = {"no":"Fatura No","musteri":"Müşteri","tarih":"Tarih",
                  "toplam":"Toplam","durum":"Durum","tahsilat":"Tahsilat","kalan":"Kalan"}
        l_wdts = {"no":130,"musteri":160,"tarih":90,"toplam":100,"durum":80,"tahsilat":100,"kalan":100}
        for c in l_cols:
            self.liste_tree.heading(c, text=l_hdrs[c])
            self.liste_tree.column(c, width=l_wdts[c],
                                   anchor="w" if c in ("no","musteri") else "center")
        l_vsb = ttk.Scrollbar(parent, orient="vertical", command=self.liste_tree.yview)
        self.liste_tree.configure(yscrollcommand=l_vsb.set)
        self.liste_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        l_vsb.pack(side=tk.LEFT, fill=tk.Y)
        self.liste_tree.tag_configure("kesildi", foreground="#69f0ae")
        self.liste_tree.tag_configure("taslak",  foreground="#fff176")
        self.liste_tree.tag_configure("iptal",   foreground="#90a4ae")
        self.liste_tree.tag_configure("odenmis", foreground="#69f0ae")
        self.liste_tree.tag_configure("bekliyor",foreground="#ff8a65")
        self.liste_tree.bind("<<TreeviewSelect>>", self._liste_sec)

        # Alt butonlar + Tahsilat
        alt2 = ttk.Frame(parent); alt2.pack(fill=tk.X, pady=(6,0))
        ttk.Button(alt2, text=" PDF İndir", command=self._liste_pdf).pack(side=tk.LEFT, padx=4)
        ttk.Button(alt2, text=" Fatura Kes", command=self._liste_kes).pack(side=tk.LEFT, padx=4)
        ttk.Button(alt2, text=" İptal Et",  command=self._liste_iptal).pack(side=tk.LEFT, padx=4)
        ttk.Separator(alt2, orient="vertical").pack(side=tk.LEFT, fill=tk.Y, padx=8)
        ttk.Label(alt2, text=" Tahsilat:").pack(side=tk.LEFT)
        self.v_tah_tutar   = tk.StringVar()
        ttk.Entry(alt2, textvariable=self.v_tah_tutar, width=10).pack(side=tk.LEFT, padx=(4,4))
        ttk.Label(alt2, text="₺").pack(side=tk.LEFT)
        self.v_tah_sekil = tk.StringVar(value="Banka")
        ttk.Combobox(alt2, textvariable=self.v_tah_sekil, width=8, state="readonly",
                     values=["Banka","Nakit","Çek","Kredi Kartı"]).pack(side=tk.LEFT, padx=(4,4))
        from datetime import date as _dt
        self.v_tah_tarih = tk.StringVar(value=_dt.today().strftime("%d.%m.%Y"))
        ttk.Entry(alt2, textvariable=self.v_tah_tarih, width=11).pack(side=tk.LEFT, padx=(4,4))
        ttk.Button(alt2, text=" Tahsilat Gir",
                   command=self._tahsilat_gir).pack(side=tk.LEFT, padx=4)

    # ── SEKME 3: Cari Döküm ──────────────────────────────────────────────────
    def _build_cari(self, parent):
        top = ttk.Frame(parent); top.pack(fill=tk.X, pady=(0,6))
        ttk.Label(top, text="Müşteri Filtre:").pack(side=tk.LEFT)
        self.v_cari_musteri = tk.StringVar()
        self.cari_cb = ttk.Combobox(top, textvariable=self.v_cari_musteri, width=28)
        self.cari_cb.pack(side=tk.LEFT, padx=(4,8))
        ttk.Button(top, text=" Göster", command=self._cari_yukle).pack(side=tk.LEFT)
        ttk.Button(top, text=" Tümü",   command=lambda: [self.v_cari_musteri.set(""), self._cari_yukle()]).pack(side=tk.LEFT, padx=4)

        c_cols = ("no","musteri","tarih","toplam","tahsilat","kalan","durum")
        self.cari_tree = ttk.Treeview(parent, columns=c_cols, show="headings", height=14)
        c_hdrs = {"no":"Fatura No","musteri":"Müşteri","tarih":"Tarih",
                  "toplam":"Toplam","tahsilat":"Tahsilat","kalan":"Kalan","durum":"Durum"}
        c_wdts = {"no":130,"musteri":160,"tarih":90,"toplam":110,"tahsilat":110,"kalan":110,"durum":80}
        for c in c_cols:
            self.cari_tree.heading(c, text=c_hdrs[c])
            self.cari_tree.column(c, width=c_wdts[c],
                                   anchor="w" if c in ("no","musteri") else "center")
        c_vsb = ttk.Scrollbar(parent, orient="vertical", command=self.cari_tree.yview)
        self.cari_tree.configure(yscrollcommand=c_vsb.set)
        self.cari_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        c_vsb.pack(side=tk.LEFT, fill=tk.Y)
        self.cari_tree.tag_configure("odenmis", foreground="#69f0ae")
        self.cari_tree.tag_configure("bekliyor", foreground="#ff8a65")

        # Özet
        self.cari_ozet_lbl = ttk.Label(parent, text="", font=("Segoe UI",9,"bold"))
        self.cari_ozet_lbl.pack(anchor="w", pady=(4,0))

    # ── SEKME 4: Firma Ayarları ───────────────────────────────────────────────
    def _build_ayarlar(self, parent):
        f = ttk.LabelFrame(parent, text=" Firma Bilgileri (Faturada Görünecek)", padding=12)
        f.pack(fill=tk.X, pady=8, padx=20)

        self.v_fa = {}
        alanlar = [
            ("firma_adi",           "Firma Adı:"),
            ("firma_vkn",           "VKN / TCKN:"),
            ("firma_vergi_dairesi", "Vergi Dairesi:"),
            ("firma_adres",         "Adres:"),
            ("firma_tel",           "Telefon:"),
            ("fatura_seri",         "Fatura Serisi:"),
            ("baslangic_no",        "Başlangıç No:"),
        ]
        for key, lbl in alanlar:
            r = ttk.Frame(f); r.pack(fill=tk.X, pady=3)
            ttk.Label(r, text=lbl, width=20, anchor="e").pack(side=tk.LEFT)
            self.v_fa[key] = tk.StringVar()
            ttk.Entry(r, textvariable=self.v_fa[key], width=40).pack(side=tk.LEFT, padx=(4,0))

        ttk.Button(f, text=" Kaydet", command=self._ayar_kaydet).pack(anchor="e", pady=(8,0))
        self._ayar_yukle()

    # ── Fonksiyonlar ─────────────────────────────────────────────────────────

    def refresh(self):
        from database import init_fatura_db
        try:
            init_fatura_db()
        except Exception as e:
            print(f"[Fatura DB] {e}")
        try:
            self._musteri_listesi_guncelle()
        except Exception as e:
            print(f"[Fatura] Müşteri: {e}")
        try:
            self._liste_yukle()
        except Exception as e:
            print(f"[Fatura] Liste: {e}")
        try:
            self._cari_yukle()
        except Exception as e:
            print(f"[Fatura] Cari: {e}")
        try:
            self._ayar_yukle()
        except Exception as e:
            print(f"[Fatura] Ayar: {e}")

    def _musteri_listesi_guncelle(self):
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            rows = conn.execute("SELECT id, name FROM customers ORDER BY name").fetchall()
            conn.close()
            self._musteri_map = {r["name"]: r["id"] for r in rows}
            adlar = list(self._musteri_map.keys())
            self.musteri_cb["values"] = adlar
            self.cari_cb["values"] = [""] + adlar
        except Exception as e:
            print(f"[Fatura] Müşteri yükleme hatası: {e}")
            self._musteri_map = {}

    def _musteri_sec(self, _=None):
        ad = self.v_musteri.get()
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            r = conn.execute(
                "SELECT * FROM customers WHERE name=?", (ad,)
            ).fetchone()
            conn.close()
            if r:
                keys = r.keys()
                self.v_musteri_vkn.set(r["tax_number"] if "tax_number" in keys else "")
                self.v_musteri_adr.set(r["address"] if "address" in keys else "")
        except Exception:
            pass

    def _yeni_fatura_no(self):
        from database import yeni_fatura_no
        self.v_fatura_no.set(yeni_fatura_no())

    def _kalem_ekle(self):
        aciklama = self.v_k_aciklama.get().strip()
        if not aciklama:
            messagebox.showwarning("Uyarı", "Açıklama giriniz!")
            return
        try:
            miktar       = float(self.v_k_miktar.get() or 1)
            birim_fiyat  = float(self.v_k_fiyat.get() or 0)
            kdv_oran     = float(self.v_k_kdv.get() or 20)
            iskonto_oran = float(self.v_k_iskonto.get() or 0)
        except ValueError:
            messagebox.showwarning("Uyarı", "Sayısal alanlara sayı girin!")
            return

        brut          = miktar * birim_fiyat
        iskonto_tutar = brut * iskonto_oran / 100
        matrah        = brut - iskonto_tutar
        kdv_tutar     = matrah * kdv_oran / 100
        toplam        = matrah + kdv_tutar

        kalem = {
            "aciklama": aciklama, "miktar": miktar, "birim": self.v_k_birim.get(),
            "birim_fiyat": birim_fiyat, "iskonto_oran": iskonto_oran,
            "kdv_oran": kdv_oran, "matrah": matrah,
            "kdv_tutar": kdv_tutar, "toplam": toplam,
        }
        self._kalemler.append(kalem)
        self._kalem_listesi_guncelle()
        # Formu temizle
        self.v_k_aciklama.set("")
        self.v_k_miktar.set("1")
        self.v_k_fiyat.set("0")
        self.v_k_iskonto.set("0")

    def _kalem_sil(self):
        sel = self.kalem_tree.selection()
        if not sel:
            return
        idx = self.kalem_tree.index(sel[0])
        if 0 <= idx < len(self._kalemler):
            self._kalemler.pop(idx)
        self._kalem_listesi_guncelle()

    def _kalem_listesi_guncelle(self):
        for item in self.kalem_tree.get_children():
            self.kalem_tree.delete(item)
        t_matrah = t_iskonto = t_kdv = t_toplam = 0
        for k in self._kalemler:
            self.kalem_tree.insert("", tk.END, values=(
                k["aciklama"],
                f"{k['miktar']:,.2f}",
                k["birim"],
                f"{k['birim_fiyat']:,.2f} ₺",
                f"%{k['iskonto_oran']:.0f}" if k["iskonto_oran"] else "—",
                f"%{k['kdv_oran']:.0f}",
                f"{k['toplam']:,.2f} ₺",
            ))
            t_matrah  += k["matrah"]
            t_iskonto += k["matrah"] * k["iskonto_oran"] / (100 - k["iskonto_oran"] + 0.001) if k["iskonto_oran"] else 0
            t_kdv     += k["kdv_tutar"]
            t_toplam  += k["toplam"]
        self.lbl_matrah.config( text=f"Matrah: {t_matrah:,.2f} ₺")
        self.lbl_iskonto.config(text=f"İskonto: {t_iskonto:,.2f} ₺")
        self.lbl_kdv.config(    text=f"KDV: {t_kdv:,.2f} ₺")
        self.lbl_toplam.config( text=f"TOPLAM: {t_toplam:,.2f} ₺")

    def _fatura_data(self) -> dict:
        from database import get_firma_ayar
        firma = get_firma_ayar()
        musteri_id = self._musteri_map.get(self.v_musteri.get())
        return {
            "firma_adi":             firma.get("firma_adi",""),
            "firma_vkn":             firma.get("firma_vkn",""),
            "firma_adres":           firma.get("firma_adres",""),
            "firma_tel":             firma.get("firma_tel",""),
            "firma_vergi_dairesi":   firma.get("firma_vergi_dairesi",""),
            "musteri_id":            musteri_id,
            "musteri_adi":           self.v_musteri.get(),
            "musteri_vkn":           self.v_musteri_vkn.get(),
            "musteri_adres":         self.v_musteri_adr.get(),
            "fatura_no":             self.v_fatura_no.get(),
            "fatura_tarihi":         self.v_tarih.get(),
            "vade_tarihi":           self.v_vade.get(),
            "fatura_turu":           self.v_tur.get(),
            "not_aciklama":          self.v_not.get(),
            "kalemler":              self._kalemler,
            "toplam_matrah":         sum(k["matrah"] for k in self._kalemler),
            "toplam_kdv":            sum(k["kdv_tutar"] for k in self._kalemler),
            "toplam_iskonto":        sum(k["birim_fiyat"]*k["miktar"]*k["iskonto_oran"]/100 for k in self._kalemler),
            "genel_toplam":          sum(k["toplam"] for k in self._kalemler),
        }

    def _validate(self) -> bool:
        if not self.v_musteri.get():
            messagebox.showwarning("Uyarı", "Müşteri seçin!"); return False
        if not self._kalemler:
            messagebox.showwarning("Uyarı", "En az bir kalem ekleyin!"); return False
        if not self.v_fatura_no.get():
            messagebox.showwarning("Uyarı", "Fatura no boş!"); return False
        return True

    def _taslak_kaydet(self):
        if not self._validate(): return
        from database import insert_fatura
        data = self._fatura_data()
        data["durum"] = "taslak"
        try:
            insert_fatura(data, self._kalemler)
            messagebox.showinfo("", "Taslak kaydedildi.")
            self._form_temizle()
            self._liste_yukle()
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _pdf_onizle(self):
        if not self._validate(): return
        import tempfile, os
        from fatura_pdf import fatura_pdf_olustur
        data = self._fatura_data()
        yol = str(Path(tempfile.gettempdir()) / f"onizleme_{data['fatura_no']}.pdf")
        try:
            fatura_pdf_olustur(data, yol)
            if os.name == "nt":
                os.startfile(yol)
            else:
                import subprocess
                subprocess.Popen(["xdg-open", yol])
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _fatura_kes(self):
        if not self._validate(): return
        if not messagebox.askyesno("Fatura Kes",
            f"! {self.v_fatura_no.get()} numaralı fatura kesilecek!\n\n"
            "Kesildikten sonra değiştirilemez. Devam?"):
            return
        from database import insert_fatura, fatura_kesildi
        from fatura_pdf import fatura_pdf_olustur
        from tkinter import filedialog
        import os
        data = self._fatura_data()
        data["durum"] = "taslak"
        try:
            fid = insert_fatura(data, self._kalemler)
            # PDF kaydet
            pdf_dir = Path(DB_PATH).parent / "Faturalar"
            pdf_dir.mkdir(exist_ok=True)
            pdf_yolu = str(pdf_dir / f"{data['fatura_no']}.pdf")
            fatura_pdf_olustur(data, pdf_yolu)
            fatura_kesildi(fid, pdf_yolu)
            messagebox.showinfo(" Fatura Kesildi!",
                f"Fatura No: {data['fatura_no']}\n"
                f"Tutar: {data['genel_toplam']:,.2f} ₺\n\n"
                f"PDF: {pdf_yolu}")
            if os.name == "nt":
                os.startfile(pdf_yolu)
            self._form_temizle()
            self._liste_yukle()
            self._yeni_fatura_no()
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _form_temizle(self):
        self.v_musteri.set("")
        self.v_musteri_vkn.set("")
        self.v_musteri_adr.set("")
        self.v_not.set("")
        self._kalemler.clear()
        self._kalem_listesi_guncelle()
        self._yeni_fatura_no()

    def _liste_yukle(self):
        from database import get_faturalar
        for item in self.liste_tree.get_children():
            self.liste_tree.delete(item)
        filtre = self.v_filtre.get() if hasattr(self,"v_filtre") else "tümü"
        try:
            rows = get_faturalar(filtre)
            for r in rows:
                tahsil = r.get("tahsilat",0) or 0
                kalan  = r.get("kalan", r["genel_toplam"]) 
                # Tahsilat hesapla
                try:
                    conn2 = sqlite3.connect(DB_PATH); conn2.row_factory=sqlite3.Row
                    t = conn2.execute("SELECT COALESCE(SUM(tutar),0) as s FROM fatura_tahsilat WHERE fatura_id=?",
                                       (r["id"],)).fetchone()["s"]
                    conn2.close()
                    kalan = r["genel_toplam"] - t
                    tahsil = t
                except Exception:
                    pass
                durum = r["durum"]
                tag = durum
                if durum == "kesildi":
                    tag = "odenmis" if kalan <= 0 else "bekliyor"
                self.liste_tree.insert("", tk.END, iid=str(r["id"]), tags=(tag,),
                    values=(r["fatura_no"], r["musteri_adi"],
                            r["fatura_tarihi"],
                            f"{r['genel_toplam']:,.2f} ₺",
                            durum.upper(),
                            f"{tahsil:,.2f} ₺",
                            f"{kalan:,.2f} ₺"))
        except Exception as e:
            print(f"Liste yükleme hatası: {e}")

    def _liste_sec(self, _=None):
        sel = self.liste_tree.selection()
        if sel:
            self._secili_fatura_id = int(sel[0])

    def _liste_pdf(self):
        from database import get_fatura_detay, get_firma_ayar
        from fatura_pdf import fatura_pdf_olustur
        import os
        if not self._secili_fatura_id:
            messagebox.showwarning("Uyarı","Listeden fatura seçin!"); return
        detay = get_fatura_detay(self._secili_fatura_id)
        if not detay: return
        # Kaydedilmiş PDF varsa direkt aç
        if detay.get("pdf_yolu") and Path(detay["pdf_yolu"]).exists():
            if os.name == "nt": os.startfile(detay["pdf_yolu"])
            return
        # Yoksa yeni oluştur
        from database import get_firma_ayar
        firma = get_firma_ayar()
        data = {**firma, **detay, "kalemler": detay["kalemler"]}
        pdf_dir = Path(DB_PATH).parent / "Faturalar"
        pdf_dir.mkdir(exist_ok=True)
        pdf_yolu = str(pdf_dir / f"{detay['fatura_no']}.pdf")
        try:
            fatura_pdf_olustur(data, pdf_yolu)
            if os.name == "nt": os.startfile(pdf_yolu)
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _liste_kes(self):
        from database import fatura_kesildi, get_fatura_detay, get_firma_ayar
        from fatura_pdf import fatura_pdf_olustur
        import os
        if not self._secili_fatura_id:
            messagebox.showwarning("Uyarı","Listeden fatura seçin!"); return
        detay = get_fatura_detay(self._secili_fatura_id)
        if not detay: return
        if detay["durum"] != "taslak":
            messagebox.showwarning("Uyarı","Sadece taslak faturalar kesilebilir!"); return
        if not messagebox.askyesno("Onayla", f"{detay['fatura_no']} faturası kesilsin mi?"): return
        firma = get_firma_ayar()
        data = {**firma, **detay, "kalemler": detay["kalemler"]}
        pdf_dir = Path(DB_PATH).parent / "Faturalar"
        pdf_dir.mkdir(exist_ok=True)
        pdf_yolu = str(pdf_dir / f"{detay['fatura_no']}.pdf")
        try:
            fatura_pdf_olustur(data, pdf_yolu)
            fatura_kesildi(self._secili_fatura_id, pdf_yolu)
            messagebox.showinfo("","Fatura kesildi ve PDF oluşturuldu.")
            if os.name=="nt": os.startfile(pdf_yolu)
            self._liste_yukle()
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _liste_iptal(self):
        from database import fatura_iptal
        if not self._secili_fatura_id:
            messagebox.showwarning("Uyarı","Listeden fatura seçin!"); return
        if messagebox.askyesno("İptal","Seçili fatura iptal edilsin mi?"):
            fatura_iptal(self._secili_fatura_id)
            self._liste_yukle()

    def _tahsilat_gir(self):
        from database import insert_fatura_tahsilat
        if not self._secili_fatura_id:
            messagebox.showwarning("Uyarı","Listeden fatura seçin!"); return
        try:
            tutar = float(self.v_tah_tutar.get().replace(",",".") or 0)
            if tutar <= 0:
                messagebox.showwarning("Uyarı","Tutar girin!"); return
            tarih = self._parse_tarih(self.v_tah_tarih.get())
            insert_fatura_tahsilat(self._secili_fatura_id, tarih, tutar,
                                    self.v_tah_sekil.get(), "")
            messagebox.showinfo("",f"{tutar:,.2f} ₺ tahsilat kaydedildi.")
            self.v_tah_tutar.set("")
            self._liste_yukle()
            self._cari_yukle()
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _parse_tarih(self, s):
        try:
            p = s.strip().split(".")
            return f"{p[2]}-{p[1]:0>2}-{p[0]:0>2}"
        except Exception:
            from datetime import date as _d2
            return _d2.today().isoformat()

    def _cari_yukle(self):
        from database import get_cari_dokum
        for item in self.cari_tree.get_children():
            self.cari_tree.delete(item)
        musteri = getattr(self, "v_cari_musteri", tk.StringVar()).get()
        try:
            rows = get_cari_dokum(musteri_adi=musteri if musteri else None)
            t_toplam = t_tahsil = t_kalan = 0
            for r in rows:
                tag = "odenmis" if r["kalan"] <= 0 else "bekliyor"
                self.cari_tree.insert("", tk.END, tags=(tag,),
                    values=(r["fatura_no"], r["musteri_adi"],
                            r["fatura_tarihi"],
                            f"{r['genel_toplam']:,.2f} ₺",
                            f"{r['tahsil_edilen']:,.2f} ₺",
                            f"{r['kalan']:,.2f} ₺",
                            r["durum"].upper()))
                t_toplam += r["genel_toplam"]
                t_tahsil += r["tahsil_edilen"]
                t_kalan  += r["kalan"]
            self.cari_ozet_lbl.config(
                text=f"  Toplam: {t_toplam:,.2f} ₺  |  "
                     f"Tahsilat: {t_tahsil:,.2f} ₺  |  "
                     f"Kalan: {t_kalan:,.2f} ₺",
                foreground="#ff8a65" if t_kalan > 0 else "#69f0ae"
            )
        except Exception as e:
            print(f"Cari yükleme hatası: {e}")

    def _ayar_yukle(self):
        from database import get_firma_ayar
        try:
            ayar = get_firma_ayar()
            for key, var in self.v_fa.items():
                var.set(str(ayar.get(key,"") or ""))
        except Exception:
            pass

    def _ayar_kaydet(self):
        from database import save_firma_ayar
        try:
            ayar = {k: v.get() for k,v in self.v_fa.items()}
            try: ayar["baslangic_no"] = int(ayar["baslangic_no"] or 1)
            except: ayar["baslangic_no"] = 1
            save_firma_ayar(ayar)
            messagebox.showinfo("","Firma ayarlari kaydedildi.")
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    # ── SEKME 5: Aylık Faturalama ─────────────────────────────────────────────
    def _build_aylik(self, parent):
        from datetime import date
        AYLAR = ["Ocak","Subat","Mart","Nisan","Mayis","Haziran",
                 "Temmuz","Agustos","Eylul","Ekim","Kasim","Aralik"]
        OFIS_TURLERI = ["Tumu","Oda","Sanal","Paylasimli"]

        # ── Üst kontrol çubuğu ──
        ust = ttk.Frame(parent)
        ust.pack(fill=tk.X, padx=8, pady=6)

        ttk.Label(ust, text="Yil:", font=("Segoe UI",9,"bold")).pack(side=tk.LEFT)
        self.v_ay_yil = tk.StringVar(value=str(date.today().year))
        yil_cb = ttk.Combobox(ust, textvariable=self.v_ay_yil, width=6, state="readonly",
                               values=[str(y) for y in range(2020, date.today().year+3)])
        yil_cb.pack(side=tk.LEFT, padx=(4,12))

        ttk.Label(ust, text="Ay:", font=("Segoe UI",9,"bold")).pack(side=tk.LEFT)
        self.v_ay_ay = tk.StringVar(value=AYLAR[date.today().month - 1])
        ay_cb = ttk.Combobox(ust, textvariable=self.v_ay_ay, width=9, state="readonly",
                              values=AYLAR)
        ay_cb.pack(side=tk.LEFT, padx=(4,12))

        ttk.Label(ust, text="Ofis Turu:", font=("Segoe UI",9,"bold")).pack(side=tk.LEFT)
        self.v_ay_tur = tk.StringVar(value="Oda")
        tur_cb = ttk.Combobox(ust, textvariable=self.v_ay_tur, width=12, state="readonly",
                               values=OFIS_TURLERI)
        tur_cb.pack(side=tk.LEFT, padx=(4,12))

        ttk.Button(ust, text="Listele", command=self._aylik_yukle).pack(side=tk.LEFT, padx=4)
        ttk.Button(ust, text="Tumu Kes", command=self._aylik_tumu_kes).pack(side=tk.LEFT, padx=4)

        # Toplam bilgi etiketi
        self.lbl_ay_toplam = ttk.Label(ust, text="Toplam: -",
                                        font=("Segoe UI",10,"bold"), foreground="#4fc3f7")
        self.lbl_ay_toplam.pack(side=tk.RIGHT, padx=12)
        self.lbl_ay_ozet = ttk.Label(ust, text="", font=("Segoe UI",9))
        self.lbl_ay_ozet.pack(side=tk.RIGHT, padx=8)

        # ── Ana içerik: Sol liste + Sağ detay ──
        body = ttk.Frame(parent)
        body.pack(fill=tk.BOTH, expand=True, padx=8, pady=(0,8))

        # Sol: Firma listesi
        sol = ttk.LabelFrame(body, text="Firmalar", padding=4)
        sol.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Treeview sütunları
        ay_cols = ("firma","ofis","kira","durum","fatura_no","fatura_tutari")
        self.ay_tree = ttk.Treeview(sol, columns=ay_cols, show="headings",
                                     selectmode="browse", height=20)
        ay_hdrs = {"firma":"Firma Adi","ofis":"Ofis","kira":"Kira (TL)",
                   "durum":"Durum","fatura_no":"Fatura No","fatura_tutari":"Fatura Tutari"}
        ay_wdts = {"firma":180,"ofis":80,"kira":90,"durum":80,"fatura_no":120,"fatura_tutari":110}
        for c in ay_cols:
            self.ay_tree.heading(c, text=ay_hdrs[c])
            self.ay_tree.column(c, width=ay_wdts[c],
                                 anchor="w" if c in ("firma","ofis","fatura_no") else "center")
        ay_vsb = ttk.Scrollbar(sol, orient="vertical", command=self.ay_tree.yview)
        self.ay_tree.configure(yscrollcommand=ay_vsb.set)
        self.ay_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        ay_vsb.pack(side=tk.LEFT, fill=tk.Y)

        # Renk etiketleri
        self.ay_tree.tag_configure("kesildi",   foreground="#69f0ae")  # Yesil
        self.ay_tree.tag_configure("bekliyor",  foreground="#fff176")  # Sari
        self.ay_tree.tag_configure("pasif",     foreground="#78909c")  # Gri

        # Çift tıkla firma detayı
        self.ay_tree.bind("<Double-Button-1>", self._aylik_firma_detay)
        self.ay_tree.bind("<<TreeviewSelect>>", self._aylik_sec)

        # Sağ: Fatura hazırlama paneli
        sag = ttk.LabelFrame(body, text="Fatura Hazirla", padding=8)
        sag.pack(side=tk.LEFT, fill=tk.Y, padx=(8,0), ipadx=4)
        sag.configure(width=300)

        # Seçili firma bilgisi
        self.lbl_sec_firma = ttk.Label(sag, text="Firma secilmedi",
                                        font=("Segoe UI",9,"bold"), foreground="#4fc3f7")
        self.lbl_sec_firma.pack(anchor="w", pady=(0,6))
        self.lbl_sec_kira  = ttk.Label(sag, text="Kira: -", font=("Segoe UI",9))
        self.lbl_sec_kira.pack(anchor="w", pady=(0,8))

        ttk.Separator(sag, orient="horizontal").pack(fill=tk.X, pady=4)

        # Kalemler listesi (düzenlenebilir)
        ttk.Label(sag, text="Fatura Kalemleri:", font=("Segoe UI",9,"bold")).pack(anchor="w")

        k_cols2 = ("aciklama","tutar","kdv")
        self.ay_kalem_tree = ttk.Treeview(sag, columns=k_cols2, show="headings", height=5)
        for c, h, w in [("aciklama","Aciklama",180),("tutar","Tutar",80),("kdv","KDV%",50)]:
            self.ay_kalem_tree.heading(c, text=h)
            self.ay_kalem_tree.column(c, width=w, anchor="w" if c=="aciklama" else "center")
        self.ay_kalem_tree.pack(fill=tk.X, pady=(2,4))

        # Kalem ekleme
        k_ekle = ttk.Frame(sag); k_ekle.pack(fill=tk.X, pady=2)
        self.v_ay_k_acik = tk.StringVar()
        self.v_ay_k_tutar = tk.StringVar(value="0")
        self.v_ay_k_kdv   = tk.StringVar(value="20")
        ttk.Label(k_ekle, text="Aciklama:", width=10).pack(side=tk.LEFT)
        ttk.Entry(k_ekle, textvariable=self.v_ay_k_acik, width=18).pack(side=tk.LEFT, padx=(2,4))

        k_ekle2 = ttk.Frame(sag); k_ekle2.pack(fill=tk.X, pady=2)
        ttk.Label(k_ekle2, text="Tutar:", width=10).pack(side=tk.LEFT)
        ttk.Entry(k_ekle2, textvariable=self.v_ay_k_tutar, width=10).pack(side=tk.LEFT, padx=(2,6))
        ttk.Label(k_ekle2, text="KDV%:").pack(side=tk.LEFT)
        ttk.Combobox(k_ekle2, textvariable=self.v_ay_k_kdv, width=4,
                     values=["0","10","20"], state="readonly").pack(side=tk.LEFT, padx=(2,0))

        k_btn = ttk.Frame(sag); k_btn.pack(fill=tk.X, pady=4)
        ttk.Button(k_btn, text="+ Kalem Ekle",  command=self._ay_kalem_ekle).pack(side=tk.LEFT, padx=2)
        ttk.Button(k_btn, text="x Kalemi Sil",  command=self._ay_kalem_sil).pack(side=tk.LEFT, padx=2)

        ttk.Separator(sag, orient="horizontal").pack(fill=tk.X, pady=6)

        # Vade tarihi
        vd = ttk.Frame(sag); vd.pack(fill=tk.X, pady=2)
        ttk.Label(vd, text="Vade:", width=10).pack(side=tk.LEFT)
        self.v_ay_vade = tk.StringVar()
        ttk.Entry(vd, textvariable=self.v_ay_vade, width=12).pack(side=tk.LEFT, padx=(2,0))

        # Not
        nt = ttk.Frame(sag); nt.pack(fill=tk.X, pady=2)
        ttk.Label(nt, text="Not:", width=10).pack(side=tk.LEFT)
        self.v_ay_not = tk.StringVar()
        ttk.Entry(nt, textvariable=self.v_ay_not, width=22).pack(side=tk.LEFT, padx=(2,0))

        ttk.Separator(sag, orient="horizontal").pack(fill=tk.X, pady=6)

        # Özet toplam
        self.lbl_ay_kalem_top = ttk.Label(sag, text="Toplam: 0,00 TL",
                                            font=("Segoe UI",10,"bold"), foreground="#4fc3f7")
        self.lbl_ay_kalem_top.pack(anchor="e", pady=4)

        # Aksiyon butonları
        ttk.Button(sag, text="FATURA KES",
                   command=self._aylik_fatura_kes).pack(fill=tk.X, pady=3, ipady=4)
        ttk.Button(sag, text="PDF Onizle",
                   command=self._aylik_pdf_onizle).pack(fill=tk.X, pady=2)
        ttk.Button(sag, text="Temizle",
                   command=self._ay_form_temizle).pack(fill=tk.X, pady=2)

        # İç state
        self._ay_kalemler = []   # [{"aciklama","tutar","kdv_oran"}]
        self._ay_secili_musteri = None  # {"id","name","kira","ofis_adi","musteri_id"}

    # ── Aylık Faturalama Fonksiyonları ───────────────────────────────────────

    def _aylik_yukle(self):
        """Seçilen yıl/ay/tür için firma listesini yükle."""
        for item in self.ay_tree.get_children():
            self.ay_tree.delete(item)
        # Sağ paneli temizle
        self._ay_secili_musteri = None
        self._ay_kalemler = []
        self._ay_kalem_listele()
        self.lbl_sec_firma.config(text="Firma secilmedi")
        self.lbl_sec_kira.config(text="Kira: -")
        self.v_ay_vade.set("")
        self.v_ay_not.set("")

        yil = self.v_ay_yil.get()
        ay_adi = self.v_ay_ay.get()
        tur  = self.v_ay_tur.get()

        AYLAR = ["Ocak","Subat","Mart","Nisan","Mayis","Haziran",
                 "Temmuz","Agustos","Eylul","Ekim","Kasim","Aralik"]
        ay_no = AYLAR.index(ay_adi) + 1
        ay_str_bas = f"{yil}-{ay_no:02d}-01"
        ay_str_bit = f"{yil}-{ay_no:02d}-31"

        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row

            # Ofis türü filtresi
            # DB'de: offices.type = 'Hazır Ofis' | 'Sanal Ofis' | 'Paylaşımlı Masa'
            # Müşteri-ofis bağlantısı: offices.customer_id = customers.id
            # YA DA customers.office_code ile INNER JOIN
            tur_filtre = ""
            tur_param  = []
            if tur != "Tumu":
                tur_map = {
                    "Oda":        "Hazır Ofis",
                    "Sanal":      "Sanal Ofis",
                    "Paylasimli": "Paylaşımlı Masa",
                }
                tur_filtre = "AND o.type = ?"
                tur_param  = [tur_map.get(tur, tur)]

            sql = f"""
                SELECT c.id as musteri_id, c.name as firma_adi,
                       COALESCE(o.code, c.office_code, '') as ofis_kodu,
                       COALESCE(o.type, '') as ofis_turu,
                       c.current_rent as kira,
                       f.id as fatura_id, f.fatura_no, f.genel_toplam, f.durum
                FROM customers c
                LEFT JOIN offices o ON (o.customer_id = c.id OR o.code = c.office_code)
                LEFT JOIN faturalar f ON f.musteri_id = c.id
                    AND f.fatura_tarihi >= ? AND f.fatura_tarihi <= ?
                    AND f.durum != 'iptal'
                WHERE c.id IS NOT NULL {tur_filtre}
                GROUP BY c.id
                ORDER BY c.name
            """
            params = [ay_str_bas, ay_str_bit] + tur_param
            rows = conn.execute(sql, params).fetchall()
            conn.close()

            toplam_kesildi = 0.0
            toplam_bekliyor = 0.0
            kesildi_sayi = 0
            toplam_sayi = 0

            for r in rows:
                kira  = float(r["kira"] or 0)
                durum_str = ""
                tag   = "bekliyor"
                fatura_no  = ""
                fatura_tut = ""

                if r["fatura_id"]:
                    durum_str  = "✓ Kesildi" if r["durum"] == "kesildi" else r["durum"].upper()
                    fatura_no  = r["fatura_no"] or ""
                    fatura_tut = f"{float(r['genel_toplam'] or 0):,.2f} TL"
                    if r["durum"] == "kesildi":
                        tag = "kesildi"
                        toplam_kesildi += float(r["genel_toplam"] or 0)
                        kesildi_sayi += 1
                    else:
                        toplam_bekliyor += float(r["genel_toplam"] or 0)
                else:
                    durum_str = "Bekliyor"
                    toplam_bekliyor += kira
                toplam_sayi += 1

                self.ay_tree.insert("", tk.END,
                    iid=str(r["musteri_id"]),
                    tags=(tag,),
                    values=(r["firma_adi"], r["ofis_kodu"] or "-",
                            f"{kira:.2f}", durum_str,
                            fatura_no, fatura_tut))

            genel = toplam_kesildi + toplam_bekliyor
            self.lbl_ay_toplam.config(
                text=f"Toplam: {genel:,.2f} TL")
            self.lbl_ay_ozet.config(
                text=f"Kesildi: {kesildi_sayi}/{toplam_sayi}  |  "
                     f"Kesilen: {toplam_kesildi:,.2f} TL")

        except Exception as e:
            import traceback; traceback.print_exc()
            messagebox.showerror("Hata", str(e))

    def _aylik_sec(self, _=None):
        """Listeden firma seç → sağ paneli doldur."""
        sel = self.ay_tree.selection()
        if not sel:
            return
        musteri_id = int(sel[0])
        vals = self.ay_tree.item(sel[0], "values")
        firma_adi  = vals[0]
        ofis_adi   = vals[1]
        kira_str   = vals[2]

        self._ay_secili_musteri = {
            "id": musteri_id,
            "name": firma_adi,
            "ofis_adi": ofis_adi,
            "kira": kira_str,
        }

        self.lbl_sec_firma.config(text=firma_adi)
        self.lbl_sec_kira.config(text=f"Ofis: {ofis_adi}  |  Kira: {kira_str} TL")

        # Otomatik kira kalemini oluştur
        ay_adi = self.v_ay_ay.get()
        yil    = self.v_ay_yil.get()
        try:
            # Format: "900.00" (nokta ondalık) - sadece boşluk/TL temizle
            kira = float(kira_str.replace(" TL","").replace(" ","").strip() or 0)
        except:
            kira = 0.0

        self._ay_kalemler = [{
            "aciklama": f"{ofis_adi} Kira Bedeli - {ay_adi} {yil}",
            "tutar": kira,
            "kdv_oran": 20.0,
        }]
        self._ay_kalem_listele()

        # Vade: ayın son günü
        AYLAR = ["Ocak","Subat","Mart","Nisan","Mayis","Haziran",
                 "Temmuz","Agustos","Eylul","Ekim","Kasim","Aralik"]
        ay_no = AYLAR.index(ay_adi) + 1
        import calendar
        son_gun = calendar.monthrange(int(yil), ay_no)[1]
        self.v_ay_vade.set(f"{son_gun:02d}.{ay_no:02d}.{yil}")
        self.v_ay_not.set(f"{ay_adi} {yil} donemi kira faturasi")

    def _ay_kalem_listele(self):
        for item in self.ay_kalem_tree.get_children():
            self.ay_kalem_tree.delete(item)
        toplam = 0.0
        for k in self._ay_kalemler:
            kdv_tutar = k["tutar"] * k["kdv_oran"] / 100
            brut = k["tutar"] + kdv_tutar
            toplam += brut
            self.ay_kalem_tree.insert("", tk.END, values=(
                k["aciklama"],
                f"{k['tutar']:,.2f}",
                f"%{k['kdv_oran']:.0f}",
            ))
        self.lbl_ay_kalem_top.config(text=f"Toplam: {toplam:,.2f} TL")

    def _ay_kalem_ekle(self):
        acik  = self.v_ay_k_acik.get().strip()
        if not acik:
            messagebox.showwarning("Uyari","Aciklama giriniz!"); return
        try:
            tutar = float(self.v_ay_k_tutar.get().replace(",",".") or 0)
            kdv   = float(self.v_ay_k_kdv.get() or 20)
        except:
            messagebox.showwarning("Uyari","Gecersiz tutar!"); return
        self._ay_kalemler.append({"aciklama":acik,"tutar":tutar,"kdv_oran":kdv})
        self._ay_kalem_listele()
        self.v_ay_k_acik.set("")
        self.v_ay_k_tutar.set("0")

    def _ay_kalem_sil(self):
        sel = self.ay_kalem_tree.selection()
        if not sel: return
        idx = self.ay_kalem_tree.index(sel[0])
        if 0 <= idx < len(self._ay_kalemler):
            self._ay_kalemler.pop(idx)
        self._ay_kalem_listele()

    def _ay_form_temizle(self):
        self._ay_kalemler = []
        self._ay_secili_musteri = None
        self.lbl_sec_firma.config(text="Firma secilmedi")
        self.lbl_sec_kira.config(text="Kira: -")
        self._ay_kalem_listele()
        self.v_ay_vade.set("")
        self.v_ay_not.set("")

    def _ay_fatura_data(self):
        """Aylık sekmeden fatura data dict'i üret."""
        from database import get_firma_ayar, yeni_fatura_no
        firma = get_firma_ayar()
        m = self._ay_secili_musteri

        # Kalem hesapları
        kalemler = []
        t_matrah = t_kdv = t_iskonto = t_toplam = 0.0
        for k in self._ay_kalemler:
            kdv_t = k["tutar"] * k["kdv_oran"] / 100
            top   = k["tutar"] + kdv_t
            t_matrah  += k["tutar"]
            t_kdv     += kdv_t
            t_toplam  += top
            kalemler.append({
                "aciklama": k["aciklama"],
                "miktar": 1, "birim": "Hizmet",
                "birim_fiyat": k["tutar"],
                "iskonto_oran": 0,
                "kdv_oran": k["kdv_oran"],
                "matrah": k["tutar"],
                "kdv_tutar": kdv_t,
                "toplam": top,
            })

        # Müşteri bilgilerini DB'den çek
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            c = conn.execute("SELECT * FROM customers WHERE id=?", (m["id"],)).fetchone()
            conn.close()
            musteri_vkn = c["tax_number"] if c and "tax_number" in c.keys() else ""
            musteri_adr = c["address"]    if c and "address"    in c.keys() else ""
        except:
            musteri_vkn = musteri_adr = ""

        ay_adi = self.v_ay_ay.get()
        yil    = self.v_ay_yil.get()
        from datetime import date
        AYLAR = ["Ocak","Subat","Mart","Nisan","Mayis","Haziran",
                 "Temmuz","Agustos","Eylul","Ekim","Kasim","Aralik"]
        ay_no = AYLAR.index(ay_adi) + 1
        fatura_tarihi = f"01.{ay_no:02d}.{yil}"

        return {
            "firma_adi":           firma.get("firma_adi",""),
            "firma_vkn":           firma.get("firma_vkn",""),
            "firma_adres":         firma.get("firma_adres",""),
            "firma_tel":           firma.get("firma_tel",""),
            "firma_vergi_dairesi": firma.get("firma_vergi_dairesi",""),
            "musteri_id":          m["id"],
            "musteri_adi":         m["name"],
            "musteri_vkn":         musteri_vkn,
            "musteri_adres":       musteri_adr,
            "fatura_no":           yeni_fatura_no(),
            "fatura_tarihi":       fatura_tarihi,
            "vade_tarihi":         self.v_ay_vade.get(),
            "fatura_turu":         "SATIS",
            "not_aciklama":        self.v_ay_not.get(),
            "kalemler":            kalemler,
            "toplam_matrah":       t_matrah,
            "toplam_kdv":          t_kdv,
            "toplam_iskonto":      0.0,
            "genel_toplam":        t_toplam,
        }

    def _aylik_pdf_onizle(self):
        if not self._ay_secili_musteri:
            messagebox.showwarning("Uyari","Firma secin!"); return
        if not self._ay_kalemler:
            messagebox.showwarning("Uyari","Kalem ekleyin!"); return
        import tempfile, os, traceback
        try:
            from fatura_pdf import fatura_pdf_olustur
            data = self._ay_fatura_data()
            yol = str(Path(tempfile.gettempdir()) / f"onizle_{data['fatura_no']}.pdf")
            fatura_pdf_olustur(data, yol)
            if os.name == "nt":
                os.startfile(yol)
            else:
                import subprocess; subprocess.Popen(["xdg-open", yol])
        except Exception as e:
            traceback.print_exc()
            messagebox.showerror("PDF Hatasi", f"{type(e).__name__}: {e}")

    def _aylik_fatura_kes(self):
        if not self._ay_secili_musteri:
            messagebox.showwarning("Uyari","Firma secin!"); return
        if not self._ay_kalemler:
            messagebox.showwarning("Uyari","Kalem ekleyin!"); return
        m = self._ay_secili_musteri
        if not messagebox.askyesno("Onayla",
                f"{m['name']} icin fatura kesilsin mi?"):
            return
        from database import insert_fatura, fatura_kesildi
        from fatura_pdf import fatura_pdf_olustur
        import os
        data = self._ay_fatura_data()
        data["durum"] = "taslak"
        try:
            fid = insert_fatura(data, data["kalemler"])
            pdf_dir = Path(DB_PATH).parent / "Faturalar"
            pdf_dir.mkdir(exist_ok=True)
            pdf_yolu = str(pdf_dir / f"{data['fatura_no']}.pdf")
            fatura_pdf_olustur(data, pdf_yolu)
            fatura_kesildi(fid, pdf_yolu)
            messagebox.showinfo("Fatura Kesildi!",
                f"Fatura No: {data['fatura_no']}\n"
                f"Tutar: {data['genel_toplam']:,.2f} TL\n"
                f"PDF: {pdf_yolu}")
            if os.name == "nt": os.startfile(pdf_yolu)
            self._aylik_yukle()
            self._ay_form_temizle()
        except Exception as e:
            messagebox.showerror("Hata", str(e))

    def _aylik_tumu_kes(self):
        """Tüm bekleyen firmaların faturalarını toplu kes."""
        bekleyenler = []
        for iid in self.ay_tree.get_children():
            vals = self.ay_tree.item(iid, "values")
            durum = vals[3]
            if "Bekliyor" in durum or durum == "":
                bekleyenler.append(iid)
        if not bekleyenler:
            messagebox.showinfo("Bilgi","Bekleyen fatura yok!"); return
        if not messagebox.askyesno("Toplu Kes",
                f"{len(bekleyenler)} firma icin fatura kesilecek. Emin misiniz?"):
            return

        from database import insert_fatura, fatura_kesildi, yeni_fatura_no, get_firma_ayar
        from fatura_pdf import fatura_pdf_olustur
        firma = get_firma_ayar()
        ay_adi = self.v_ay_ay.get()
        yil    = self.v_ay_yil.get()
        AYLAR  = ["Ocak","Subat","Mart","Nisan","Mayis","Haziran",
                  "Temmuz","Agustos","Eylul","Ekim","Kasim","Aralik"]
        ay_no  = AYLAR.index(ay_adi) + 1
        import calendar

        basarili = 0; hatali = 0
        for iid in bekleyenler:
            try:
                vals = self.ay_tree.item(iid, "values")
                musteri_id = int(iid)
                firma_adi  = vals[0]
                ofis_adi   = vals[1]
                kira       = float(vals[2].replace(",","") or 0)

                conn = sqlite3.connect(DB_PATH)
                conn.row_factory = sqlite3.Row
                c = conn.execute("SELECT * FROM customers WHERE id=?",
                                  (musteri_id,)).fetchone()
                conn.close()
                musteri_vkn = c["tax_number"] if c and "tax_number" in c.keys() else ""
                musteri_adr = c["address"]    if c and "address"    in c.keys() else ""

                kdv_t = kira * 0.20
                toplam = kira + kdv_t
                son_gun = calendar.monthrange(int(yil), ay_no)[1]
                kalem = [{
                    "aciklama": f"{ofis_adi} Kira Bedeli - {ay_adi} {yil}",
                    "miktar": 1, "birim": "Hizmet",
                    "birim_fiyat": kira,
                    "iskonto_oran": 0, "kdv_oran": 20.0,
                    "matrah": kira, "kdv_tutar": kdv_t, "toplam": toplam,
                }]
                data = {
                    "firma_adi":           firma.get("firma_adi",""),
                    "firma_vkn":           firma.get("firma_vkn",""),
                    "firma_adres":         firma.get("firma_adres",""),
                    "firma_tel":           firma.get("firma_tel",""),
                    "firma_vergi_dairesi": firma.get("firma_vergi_dairesi",""),
                    "musteri_id":   musteri_id,
                    "musteri_adi":  firma_adi,
                    "musteri_vkn":  musteri_vkn,
                    "musteri_adres":musteri_adr,
                    "fatura_no":    yeni_fatura_no(),
                    "fatura_tarihi":f"01.{ay_no:02d}.{yil}",
                    "vade_tarihi":  f"{son_gun:02d}.{ay_no:02d}.{yil}",
                    "fatura_turu":  "SATIS",
                    "not_aciklama": f"{ay_adi} {yil} donemi kira faturasi",
                    "kalemler":     kalem,
                    "toplam_matrah": kira,
                    "toplam_kdv":    kdv_t,
                    "toplam_iskonto":0.0,
                    "genel_toplam":  toplam,
                    "durum":        "taslak",
                }
                fid = insert_fatura(data, kalem)
                pdf_dir = Path(DB_PATH).parent / "Faturalar"
                pdf_dir.mkdir(exist_ok=True)
                pdf_yolu = str(pdf_dir / f"{data['fatura_no']}.pdf")
                fatura_pdf_olustur(data, pdf_yolu)
                fatura_kesildi(fid, pdf_yolu)
                basarili += 1
            except Exception as e:
                print(f"[Toplu Kes] {iid}: {e}")
                hatali += 1

        messagebox.showinfo("Toplu Kesim Tamamlandi",
            f"Basarili: {basarili} fatura\nHatali: {hatali} fatura")
        self._aylik_yukle()

    def _aylik_firma_detay(self, _=None):
        """Çift tıkla: Firmaya ait tüm faturaları popup'ta göster."""
        sel = self.ay_tree.selection()
        if not sel: return
        musteri_id  = int(sel[0])
        firma_adi   = self.ay_tree.item(sel[0], "values")[0]

        popup = tk.Toplevel(self)
        popup.title(f"Fatura Gecmisi — {firma_adi}")
        popup.geometry("820x480")
        popup.transient(self)
        popup.grab_set()

        ttk.Label(popup, text=f"{firma_adi} — Fatura Gecmisi",
                  font=("Segoe UI",11,"bold")).pack(pady=(12,6))

        # Treeview
        cols = ("fatura_no","tarih","tutar","kdv","toplam","durum","pdf")
        tree = ttk.Treeview(popup, columns=cols, show="headings", height=14)
        hdrs = {"fatura_no":"Fatura No","tarih":"Tarih","tutar":"Matrah",
                "kdv":"KDV","toplam":"Genel Toplam","durum":"Durum","pdf":"PDF"}
        wdts = {"fatura_no":130,"tarih":90,"tutar":100,"kdv":80,
                "toplam":110,"durum":80,"pdf":60}
        for c in cols:
            tree.heading(c, text=hdrs[c])
            tree.column(c, width=wdts[c], anchor="center")
        vsb = ttk.Scrollbar(popup, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=vsb.set)
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(8,0), pady=8)
        vsb.pack(side=tk.LEFT, fill=tk.Y, pady=8)

        tree.tag_configure("kesildi", foreground="#69f0ae")
        tree.tag_configure("taslak",  foreground="#fff176")
        tree.tag_configure("iptal",   foreground="#90a4ae")

        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            rows = conn.execute("""
                SELECT fatura_no, fatura_tarihi, toplam_matrah, toplam_kdv,
                       genel_toplam, durum, pdf_yolu
                FROM faturalar WHERE musteri_id=?
                ORDER BY fatura_tarihi DESC
            """, (musteri_id,)).fetchall()
            conn.close()

            t_toplam = 0.0
            for r in rows:
                tree.insert("", tk.END, tags=(r["durum"],), values=(
                    r["fatura_no"],
                    r["fatura_tarihi"],
                    f"{float(r['toplam_matrah'] or 0):,.2f}",
                    f"{float(r['toplam_kdv'] or 0):,.2f}",
                    f"{float(r['genel_toplam'] or 0):,.2f}",
                    r["durum"].upper(),
                    "Var" if r["pdf_yolu"] else "-",
                ))
                if r["durum"] != "iptal":
                    t_toplam += float(r["genel_toplam"] or 0)

            alt = ttk.Frame(popup)
            alt.pack(fill=tk.X, padx=8, pady=(0,8))
            ttk.Label(alt, text=f"Toplam Ciro: {t_toplam:,.2f} TL",
                      font=("Segoe UI",10,"bold"), foreground="#4fc3f7").pack(side=tk.LEFT)

            # PDF aç butonu
            def pdf_ac():
                sel2 = tree.selection()
                if not sel2: return
                idx = tree.index(sel2[0])
                if idx < len(rows):
                    pdf = rows[idx]["pdf_yolu"]
                    if pdf and Path(pdf).exists():
                        import os
                        if os.name == "nt": os.startfile(pdf)
                    else:
                        messagebox.showinfo("Bilgi","PDF bulunamadi.")
            ttk.Button(alt, text="PDF Ac", command=pdf_ac).pack(side=tk.RIGHT)
            ttk.Button(alt, text="Kapat", command=popup.destroy).pack(side=tk.RIGHT, padx=4)

        except Exception as e:
            messagebox.showerror("Hata", str(e))
            popup.destroy()



# ============================================================================
# KARGOLAR ANA SEKMESİ
# ============================================================================

class KargolarTab(ttk.Frame):
    """Tüm müşterilerin kargolarını gösteren ana sekme."""

    def __init__(self, master):
        super().__init__(master, padding=8)
        self._build_ui()
        self.refresh()

    def _build_ui(self):
        # ── Üst filtre çubuğu ──
        ust = ttk.Frame(self)
        ust.pack(fill=tk.X, pady=(0, 6))

        ttk.Label(ust, text="Kargolar", font=("Segoe UI", 13, "bold")).pack(side=tk.LEFT)

        ttk.Label(ust, text="  Musteri:").pack(side=tk.LEFT, padx=(16, 2))
        self.v_filtre_musteri = tk.StringVar()
        e = ttk.Entry(ust, textvariable=self.v_filtre_musteri, width=18)
        e.pack(side=tk.LEFT)
        e.bind("<Return>", lambda _: self.refresh())

        ttk.Label(ust, text="  Durum:").pack(side=tk.LEFT, padx=(8, 2))
        self.v_filtre_durum = tk.StringVar(value="Tumu")
        ttk.Combobox(ust, textvariable=self.v_filtre_durum,
                     values=["Tumu", "bekliyor", "odenmis", "ucretsiz"],
                     width=10, state="readonly").pack(side=tk.LEFT)

        ttk.Button(ust, text="Listele", command=self.refresh).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(ust, text="Yenile",  command=self.refresh).pack(side=tk.LEFT, padx=4)

        # Özet etiketleri
        self.lbl_ozet = ttk.Label(ust, text="",
                                   font=("Segoe UI", 9), foreground="#4fc3f7")
        self.lbl_ozet.pack(side=tk.RIGHT, padx=8)

        # ── Treeview ──
        cols = ("tarih", "musteri", "kargo_firmasi", "teslim_alan",
                "takip_no", "odeme", "durum", "wp", "resim")
        self.tree = ttk.Treeview(self, columns=cols, show="headings",
                                  selectmode="browse", height=22)

        hdrs  = {"tarih":"Tarih", "musteri":"Musteri",
                 "kargo_firmasi":"Kargo Firmasi", "teslim_alan":"Teslim Alan",
                 "takip_no":"Takip No", "odeme":"Odeme (TL)",
                 "durum":"Odeme Durumu", "wp":"WhatsApp", "resim":"Resim"}
        wdts  = {"tarih":90, "musteri":160, "kargo_firmasi":110, "teslim_alan":120,
                 "takip_no":120, "odeme":90, "durum":100, "wp":80, "resim":60}
        anch  = {"odeme":"e", "resim":"center"}

        for c in cols:
            self.tree.heading(c, text=hdrs[c],
                              command=lambda _c=c: self._sort(_c))
            self.tree.column(c, width=wdts[c],
                             anchor=anch.get(c, "w"))

        vsb = ttk.Scrollbar(self, orient="vertical",   command=self.tree.yview)
        hsb = ttk.Scrollbar(self, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.LEFT, fill=tk.Y)

        self.tree.tag_configure("bekliyor", foreground="#ff8a65")
        self.tree.tag_configure("odendi",   foreground="#69f0ae")
        self.tree.tag_configure("ucretsiz", foreground="#b0bec5")
        self.tree.tag_configure("wp_ok",    foreground="#69f0ae")

        self.tree.bind("<Double-1>", self._detay_ac)
        self._sort_col = None
        self._sort_rev = False

    def refresh(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        from database import get_tum_kargolar
        musteri = self.v_filtre_musteri.get().strip()
        durum_f = self.v_filtre_durum.get()
        if durum_f == "Tumu":
            durum_f = ""

        try:
            rows = get_tum_kargolar(musteri, durum_f)
        except Exception as e:
            self.lbl_ozet.config(text=f"Hata: {e}")
            return

        t_kargo = len(rows)
        t_odeme = sum(float(r.get("odeme_tutari") or 0) for r in rows)
        t_bekliyor = sum(float(r.get("odeme_tutari") or 0)
                         for r in rows if r.get("odeme_durumu") == "bekliyor")
        self.lbl_ozet.config(
            text=f"Toplam: {t_kargo}  |  "
                 f"Tahsil bekleyen: {t_bekliyor:,.2f} TL  |  "
                 f"Toplam odeme: {t_odeme:,.2f} TL")

        for r in rows:
            tutar = float(r.get("odeme_tutari") or 0)
            odeme_str = f"{tutar:,.2f}" if tutar > 0 else "-"
            durum = r.get("odeme_durumu", "")
            if tutar == 0:
                tag = "ucretsiz"; durum_str = "Ucretsiz"
            elif durum == "odendi":
                tag = "odendi"; durum_str = "Odendi"
            else:
                tag = "bekliyor"; durum_str = "Bekliyor"

            wp_str  = "Gonderildi" if r.get("whatsapp_gonderildi") else "-"
            resim_n = r.get("resim_sayisi", 0)

            self.tree.insert("", tk.END, iid=str(r["id"]), tags=(tag,), values=(
                r.get("tarih", ""),
                r.get("musteri_adi", ""),
                r.get("kargo_firmasi", "") or "",
                r.get("teslim_alan",  "") or "",
                r.get("takip_no",     "") or "",
                odeme_str,
                durum_str,
                wp_str,
                f"{resim_n} resim" if resim_n else "-",
            ))

    def _sort(self, col):
        items = [(self.tree.set(k, col), k) for k in self.tree.get_children("")]
        rev = (self._sort_col == col) and not self._sort_rev
        try:
            items.sort(key=lambda x: float(x[0].replace(",","").replace("-","0")),
                       reverse=rev)
        except:
            items.sort(key=lambda x: x[0].lower(), reverse=rev)
        for i, (_, k) in enumerate(items):
            self.tree.move(k, "", i)
        self._sort_col = col; self._sort_rev = rev

    def _detay_ac(self, _=None):
        """Çift tıkla → müşteri detay popup'ını aç."""
        sel = self.tree.selection()
        if not sel: return
        kargo_id = int(sel[0])
        # Kargo kaydından müşteri bilgisini al
        from database import fetch_one
        r = fetch_one("""
            SELECT k.musteri_id, c.name, c.phone
            FROM kargolar k JOIN customers c ON c.id=k.musteri_id
            WHERE k.id=?
        """, (kargo_id,))
        if r:
            MusteriDetayPopup(self, r["musteri_id"], r["name"], r["phone"] or "")


# ============================================================================
# MUSTERI DETAY POPUP — Cari Ekstre + Kargo
# ============================================================================

class MusteriDetayPopup(tk.Toplevel):
    """Musteri cift tiklama popup: sol=cari ekstre, sag=kargo gecmisi."""

    KARGO_FIRMALARI = [
        "",
        # Geleneksel Kargo
        "Aras Kargo", "Yurtici Kargo", "MNG Kargo", "Surat Kargo", "PTT Kargo",
        # E-Ticaret Odakli
        "Trendyol Express", "Hepsijet", "Kolay Gelsin", "Jetizz", "Sendeo",
        "Vigo", "Paket Taxi",
        # Uluslararasi
        "DHL", "UPS", "FedEx", "TNT",
        # Diger
        "Diger",
    ]

    def __init__(self, master, musteri_id: int, musteri_adi: str, telefon: str = ""):
        super().__init__(master)
        self.musteri_id  = musteri_id
        self.musteri_adi = musteri_adi
        self.telefon     = telefon
        self._secili_kargo_id = None
        self._yeni_resimler   = []   # [(yol, ad), ...]

        self.title(f"Musteri Detay — {musteri_adi}")
        self.geometry("1200x640")
        self.transient(master)
        self.grab_set()
        self.resizable(True, True)

        # Kargo resim klasörü
        from pathlib import Path
        from database import DB_PATH
        self._resim_dir = Path(DB_PATH).parent / "KargoResimleri" / str(musteri_id)
        self._resim_dir.mkdir(parents=True, exist_ok=True)

        self._build_ui()
        self._cari_yukle()
        self._kargo_yukle()

    # ── UI ────────────────────────────────────────────────────────────────

    def _build_ui(self):
        # Başlık
        ust = ttk.Frame(self)
        ust.pack(fill=tk.X, padx=10, pady=(8, 4))
        ttk.Label(ust, text=self.musteri_adi,
                  font=("Segoe UI", 12, "bold")).pack(side=tk.LEFT)
        tel_txt = f"  |  Tel: {self.telefon}" if self.telefon else ""
        ttk.Label(ust, text=tel_txt, font=("Segoe UI", 9)).pack(side=tk.LEFT)
        ttk.Button(ust, text="Kapat", command=self.destroy).pack(side=tk.RIGHT)

        ttk.Separator(self, orient="horizontal").pack(fill=tk.X, padx=10)

        # Ana alan: sol + sag
        body = ttk.Frame(self)
        body.pack(fill=tk.BOTH, expand=True, padx=10, pady=8)

        # ── SOL: Cari Ekstre ──────────────────────────────────────
        sol = ttk.LabelFrame(body, text="Cari Ekstre", padding=6)
        sol.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 6))

        cari_cols = ("tarih", "fatura_no", "tutar", "tahsilat", "kalan", "durum")
        self.cari_tree = ttk.Treeview(sol, columns=cari_cols, show="headings", height=16)
        hdrs = {"tarih":"Tarih","fatura_no":"Fatura No","tutar":"Tutar",
                "tahsilat":"Tahsilat","kalan":"Kalan","durum":"Durum"}
        wdts = {"tarih":85,"fatura_no":120,"tutar":90,"tahsilat":90,"kalan":90,"durum":75}
        for c in cari_cols:
            self.cari_tree.heading(c, text=hdrs[c])
            self.cari_tree.column(c, width=wdts[c],
                                   anchor="w" if c in ("fatura_no",) else "center")
        c_vsb = ttk.Scrollbar(sol, orient="vertical", command=self.cari_tree.yview)
        self.cari_tree.configure(yscrollcommand=c_vsb.set)
        self.cari_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        c_vsb.pack(side=tk.LEFT, fill=tk.Y)

        self.cari_tree.tag_configure("kesildi",  foreground="#69f0ae")
        self.cari_tree.tag_configure("taslak",   foreground="#fff176")
        self.cari_tree.tag_configure("iptal",    foreground="#90a4ae")
        self.cari_tree.tag_configure("odenmis",  foreground="#69f0ae")
        self.cari_tree.tag_configure("bekliyor", foreground="#ff8a65")

        # Cari özet
        self.lbl_cari_ozet = ttk.Label(sol, text="",
                                        font=("Segoe UI", 9, "bold"), foreground="#4fc3f7")
        self.lbl_cari_ozet.pack(fill=tk.X, pady=(4, 0))

        # ── SAG: Kargo Sekmeli ────────────────────────────────────
        sag = ttk.Frame(body)
        sag.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        kargo_nb = ttk.Notebook(sag)
        kargo_nb.pack(fill=tk.BOTH, expand=True)

        f_liste = ttk.Frame(kargo_nb)
        f_yeni  = ttk.Frame(kargo_nb)
        kargo_nb.add(f_liste, text="  Kargo Gecmisi  ")
        kargo_nb.add(f_yeni,  text="  Yeni Kargo Ekle  ")

        self._build_kargo_liste(f_liste)
        self._build_kargo_yeni(f_yeni)

    def _build_kargo_liste(self, parent):
        # Üst özet
        ozet_f = ttk.Frame(parent)
        ozet_f.pack(fill=tk.X, pady=(0, 4))
        self.lbl_kargo_ozet = ttk.Label(ozet_f, text="",
                                         font=("Segoe UI", 9), foreground="#4fc3f7")
        self.lbl_kargo_ozet.pack(side=tk.LEFT)

        # Treeview
        k_cols = ("tarih", "kargo_firmasi", "teslim_alan", "takip_no", "wp", "notlar")
        self.kargo_tree = ttk.Treeview(parent, columns=k_cols, show="headings",
                                        selectmode="browse", height=10)
        hdrs = {"tarih":"Tarih","kargo_firmasi":"Kargo Firmasi","teslim_alan":"Teslim Alan",
                "takip_no":"Takip No","wp":"WhatsApp","notlar":"Notlar"}
        wdts = {"tarih":85,"kargo_firmasi":90,"teslim_alan":110,"takip_no":100,"wp":70,"notlar":130}
        for c in k_cols:
            self.kargo_tree.heading(c, text=hdrs[c])
            self.kargo_tree.column(c, width=wdts[c],
                                    anchor="w" if c in ("teslim_alan","notlar") else "center")
        k_vsb = ttk.Scrollbar(parent, orient="vertical", command=self.kargo_tree.yview)
        self.kargo_tree.configure(yscrollcommand=k_vsb.set)
        self.kargo_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        k_vsb.pack(side=tk.LEFT, fill=tk.Y)

        self.kargo_tree.tag_configure("wp_ok",    foreground="#69f0ae")
        self.kargo_tree.tag_configure("wp_hayir", foreground="#fff176")

        self.kargo_tree.bind("<<TreeviewSelect>>", self._kargo_sec)

        # Alt butonlar
        alt = ttk.Frame(parent)
        alt.pack(fill=tk.X, pady=(4, 0))
        ttk.Button(alt, text="WhatsApp Gonder",
                   command=self._whatsapp_gonder).pack(side=tk.LEFT, padx=2)
        ttk.Button(alt, text="Resimleri Goster",
                   command=self._resimleri_goster).pack(side=tk.LEFT, padx=2)
        ttk.Button(alt, text="Kaydi Sil",
                   command=self._kargo_sil).pack(side=tk.LEFT, padx=2)

        # Resim önizleme etiketi
        self.lbl_resim_bilgi = ttk.Label(parent, text="", font=("Segoe UI", 8))
        self.lbl_resim_bilgi.pack(fill=tk.X, pady=(2, 0))

        # ── Borç Paneli (seçili kargo ödemeliyse açılır) ──────────────────
        self.borc_frame = ttk.LabelFrame(parent, text="Kargo Borcu", padding=6)
        self.borc_frame.pack(fill=tk.X, pady=(6, 0))

        b_ust = ttk.Frame(self.borc_frame); b_ust.pack(fill=tk.X)
        self.lbl_borc_tutar = ttk.Label(b_ust,
            text="Odeme tutari: -",
            font=("Segoe UI", 10, "bold"), foreground="#ff8a65")
        self.lbl_borc_tutar.pack(side=tk.LEFT)

        self.lbl_borc_durum = ttk.Label(b_ust,
            text="",
            font=("Segoe UI", 9, "bold"))
        self.lbl_borc_durum.pack(side=tk.LEFT, padx=(12, 0))

        b_btn = ttk.Frame(self.borc_frame); b_btn.pack(fill=tk.X, pady=(4, 0))
        ttk.Button(b_btn, text="Odendi Isaretle",
                   command=lambda: self._kargo_odeme_guncelle("odendi")).pack(side=tk.LEFT, padx=2)
        ttk.Button(b_btn, text="Bekliyor Isaretle",
                   command=lambda: self._kargo_odeme_guncelle("bekliyor")).pack(side=tk.LEFT, padx=2)

        self.lbl_borc_fatura = ttk.Label(self.borc_frame,
            text="", font=("Segoe UI", 8), foreground="#aaaaaa")
        self.lbl_borc_fatura.pack(anchor="w", pady=(2, 0))

    # Barkod → kargo firması eşleştirme tablosu
    BARKOD_FIRMA_MAP = {
        # Prefix bazlı tespitler
        "1": "Aras Kargo",       # Aras: 1 ile başlar (genelde 10-11 hane)
        "6": "Yurtici Kargo",    # Yurtici: 6 ile başlar
        "7": "MNG Kargo",        # MNG: 7 ile başlar
        "5": "Surat Kargo",      # Surat: 5 ile başlar
        "99": "PTT Kargo",       # PTT: 99 ile başlar
        "TY": "Trendyol Express",
        "HP": "Hepsijet",
        "KG": "Kolay Gelsin",
        "JT": "Jetizz",
        "SD": "Sendeo",
        "DHL": "DHL",
        "UPS": "UPS",
        "1Z": "UPS",             # UPS tracking "1Z" ile başlar
        "FX": "FedEx",
        "JD": "Trendyol Express",
    }

    def _barkod_firma_tespit(self, barkod: str) -> str:
        """Barkod numarasından kargo firmasını otomatik tespit et."""
        b = barkod.strip().upper()
        # Harf prefix kontrolü (2-3 karakter)
        for prefix in ["DHL", "UPS", "1Z", "TY", "HP", "KG", "JT", "SD", "FX", "JD", "99"]:
            if b.startswith(prefix):
                return self.BARKOD_FIRMA_MAP.get(prefix, "")
        # Tek harf prefix
        if b and b[0] in self.BARKOD_FIRMA_MAP:
            return self.BARKOD_FIRMA_MAP[b[0]]
        return ""

    def _build_kargo_yeni(self, parent):
        frm = ttk.Frame(parent, padding=10)
        frm.pack(fill=tk.BOTH, expand=True)

        def lbl_row(label, widget):
            f = ttk.Frame(frm); f.pack(fill=tk.X, pady=3)
            ttk.Label(f, text=label, width=14, anchor="e").pack(side=tk.LEFT)
            widget.pack(in_=f, side=tk.LEFT, padx=(6, 0), fill=tk.X, expand=True)
            return widget

        from datetime import date
        self.v_k_tarih   = tk.StringVar(value=date.today().strftime("%d.%m.%Y"))
        self.v_k_firma   = tk.StringVar()
        self.v_k_teslim  = tk.StringVar()
        self.v_k_takip   = tk.StringVar()
        self.v_k_notlar  = tk.StringVar()

        # ── Barkod Okuma alanı ──────────────────────────────────
        barkod_lf = ttk.LabelFrame(frm, text="Barkod Okuyucu", padding=6)
        barkod_lf.pack(fill=tk.X, pady=(0, 8))

        barkod_ust = ttk.Frame(barkod_lf); barkod_ust.pack(fill=tk.X)
        ttk.Label(barkod_ust, text="Barkod:", width=10, anchor="e").pack(side=tk.LEFT)
        self.v_barkod = tk.StringVar()
        self.barkod_entry = ttk.Entry(barkod_ust, textvariable=self.v_barkod,
                                       width=24, font=("Segoe UI", 10))
        self.barkod_entry.pack(side=tk.LEFT, padx=(6, 6))
        ttk.Button(barkod_ust, text="Uygula",
                   command=self._barkod_uygula).pack(side=tk.LEFT, padx=2)
        ttk.Button(barkod_ust, text="Temizle",
                   command=lambda: self.v_barkod.set("")).pack(side=tk.LEFT, padx=2)

        self.lbl_barkod_sonuc = ttk.Label(barkod_lf, text="Barkod okuyucuyu bu alana odaklayip okutun",
                                           font=("Segoe UI", 8), foreground="#aaaaaa")
        self.lbl_barkod_sonuc.pack(anchor="w", pady=(2, 0))

        # Enter tuşuna basınca otomatik uygula (barkod okuyucu Enter gönderir)
        self.barkod_entry.bind("<Return>", lambda e: self._barkod_uygula())
        # Popup açılınca barkod alanına focus ver
        self.after(200, lambda: self.barkod_entry.focus_set())

        # ── Normal Form Alanları ────────────────────────────────
        # Tarih
        f_tarih = ttk.Frame(frm); f_tarih.pack(fill=tk.X, pady=3)
        ttk.Label(f_tarih, text="Tarih:", width=14, anchor="e").pack(side=tk.LEFT)
        ttk.Entry(f_tarih, textvariable=self.v_k_tarih, width=12).pack(side=tk.LEFT, padx=(6,0))

        # Kargo Firması
        f_firma = ttk.Frame(frm); f_firma.pack(fill=tk.X, pady=3)
        ttk.Label(f_firma, text="Kargo Firmasi:", width=14, anchor="e").pack(side=tk.LEFT)
        self.firma_cb = ttk.Combobox(f_firma, textvariable=self.v_k_firma,
                                      values=self.KARGO_FIRMALARI, state="normal", width=18)
        self.firma_cb.pack(side=tk.LEFT, padx=(6, 0))

        # Teslim Alan
        f_teslim = ttk.Frame(frm); f_teslim.pack(fill=tk.X, pady=3)
        ttk.Label(f_teslim, text="Teslim Alan:", width=14, anchor="e").pack(side=tk.LEFT)
        ttk.Entry(f_teslim, textvariable=self.v_k_teslim, width=22).pack(side=tk.LEFT, padx=(6,0))

        # Takip No (barkod ile dolar)
        f_takip = ttk.Frame(frm); f_takip.pack(fill=tk.X, pady=3)
        ttk.Label(f_takip, text="Takip No:", width=14, anchor="e").pack(side=tk.LEFT)
        self.takip_entry = ttk.Entry(f_takip, textvariable=self.v_k_takip, width=22)
        self.takip_entry.pack(side=tk.LEFT, padx=(6, 0))

        # Ödeme Tutarı
        f_odeme = ttk.Frame(frm); f_odeme.pack(fill=tk.X, pady=3)
        ttk.Label(f_odeme, text="Odeme Tutari:", width=14, anchor="e").pack(side=tk.LEFT)
        self.v_k_odeme = tk.StringVar(value="")
        odeme_e = ttk.Entry(f_odeme, textvariable=self.v_k_odeme, width=12)
        odeme_e.pack(side=tk.LEFT, padx=(6, 4))
        ttk.Label(f_odeme, text="TL  (0 = ucretsiz)",
                  font=("Segoe UI", 8), foreground="#aaaaaa").pack(side=tk.LEFT)

        # Notlar
        f_notlar = ttk.Frame(frm); f_notlar.pack(fill=tk.X, pady=3)
        ttk.Label(f_notlar, text="Notlar:", width=14, anchor="e").pack(side=tk.LEFT)
        ttk.Entry(f_notlar, textvariable=self.v_k_notlar, width=30).pack(side=tk.LEFT, padx=(6,0))

        # Resim seçimi
        r_f = ttk.LabelFrame(frm, text="Resimler", padding=6)
        r_f.pack(fill=tk.X, pady=(8, 4))

        self.resim_listbox = tk.Listbox(r_f, height=3, selectmode=tk.SINGLE,
                                         font=("Segoe UI", 8))
        self.resim_listbox.pack(fill=tk.X)

        r_btn = ttk.Frame(r_f); r_btn.pack(fill=tk.X, pady=(4, 0))
        ttk.Button(r_btn, text="+ Resim Sec",
                   command=self._resim_sec).pack(side=tk.LEFT, padx=2)
        ttk.Button(r_btn, text="x Kaldir",
                   command=self._resim_kaldir).pack(side=tk.LEFT, padx=2)

        # Kaydet butonu
        ttk.Separator(frm, orient="horizontal").pack(fill=tk.X, pady=8)
        ttk.Button(frm, text="KARGOYU KAYDET",
                   command=self._kargo_kaydet).pack(fill=tk.X, ipady=6)

    def _barkod_uygula(self):
        """Barkod alanındaki değeri takip no ve firma tespitine uygula."""
        barkod = self.v_barkod.get().strip()
        if not barkod:
            return

        # Takip no alanına aktar
        self.v_k_takip.set(barkod)

        # Firma otomatik tespit
        firma = self._barkod_firma_tespit(barkod)
        if firma:
            self.v_k_firma.set(firma)
            self.lbl_barkod_sonuc.config(
                text=f"Tespit edildi: {firma}  |  Takip No: {barkod}",
                foreground="#69f0ae")
        else:
            self.lbl_barkod_sonuc.config(
                text=f"Firma tespit edilemedi. Manuel seçin.  |  Takip No: {barkod}",
                foreground="#fff176")

        # Barkod alanını temizle, teslim alan alanına geç
        self.v_barkod.set("")
        self.v_k_teslim.focus_set()

    # ── Veri Yükleme ─────────────────────────────────────────────────────

    def _cari_yukle(self):
        for item in self.cari_tree.get_children():
            self.cari_tree.delete(item)
        from database import get_cari_dokum
        try:
            rows = get_cari_dokum(musteri_id=self.musteri_id)
            t_tutar = t_tahsilat = t_kalan = 0.0
            for r in rows:
                tag = r.get("durum", "taslak")
                if r["kalan"] <= 0:
                    tag = "odenmis"
                elif tag == "kesildi":
                    tag = "bekliyor"
                self.cari_tree.insert("", tk.END, tags=(tag,), values=(
                    r.get("fatura_tarihi", ""),
                    r.get("fatura_no", ""),
                    f"{float(r.get('genel_toplam',0)):,.2f}",
                    f"{float(r.get('tahsil_edilen',0)):,.2f}",
                    f"{float(r.get('kalan',0)):,.2f}",
                    r.get("durum", "").upper(),
                ))
                t_tutar    += float(r.get("genel_toplam", 0))
                t_tahsilat += float(r.get("tahsil_edilen", 0))
                t_kalan    += float(r.get("kalan", 0))
            self.lbl_cari_ozet.config(
                text=f"Toplam: {t_tutar:,.2f} TL  |  "
                     f"Tahsilat: {t_tahsilat:,.2f} TL  |  "
                     f"Kalan: {t_kalan:,.2f} TL")
        except Exception as e:
            self.lbl_cari_ozet.config(text=f"Hata: {e}")

    def _kargo_yukle(self):
        for item in self.kargo_tree.get_children():
            self.kargo_tree.delete(item)
        from database import get_kargolar, get_kargo_ozet
        try:
            rows = get_kargolar(self.musteri_id)
            ozet = get_kargo_ozet(self.musteri_id)
            self.lbl_kargo_ozet.config(
                text=f"Toplam kargo: {ozet['toplam']}  |  "
                     f"WhatsApp: {ozet['wp_gonderilen']}  |  "
                     f"Son kargo: {ozet['son_kargo'] or '-'}")
            for r in rows:
                wp_str = "Gonderildi" if r["whatsapp_gonderildi"] else "Gonderilmedi"
                tag    = "wp_ok" if r["whatsapp_gonderildi"] else "wp_hayir"
                resim_sayisi = len(r["resim_listesi"])
                notlar = r.get("notlar") or ""
                if resim_sayisi:
                    notlar = f"[{resim_sayisi} resim] " + notlar
                self.kargo_tree.insert("", tk.END, iid=str(r["id"]), tags=(tag,), values=(
                    r["tarih"],
                    r.get("kargo_firmasi", "") or "",
                    r.get("teslim_alan", "") or "",
                    r.get("takip_no", "") or "",
                    wp_str,
                    notlar,
                ))
        except Exception as e:
            self.lbl_kargo_ozet.config(text=f"Hata: {e}")

    # ── Kargo İşlemleri ───────────────────────────────────────────────────

    def _kargo_sec(self, _=None):
        sel = self.kargo_tree.selection()
        if not sel:
            self._secili_kargo_id = None
            self.lbl_resim_bilgi.config(text="")
            self.lbl_borc_tutar.config(text="Odeme tutari: -")
            self.lbl_borc_durum.config(text="")
            self.lbl_borc_fatura.config(text="")
            return
        self._secili_kargo_id = int(sel[0])
        from database import get_kargolar
        rows = get_kargolar(self.musteri_id)
        for r in rows:
            if r["id"] == self._secili_kargo_id:
                n = len(r["resim_listesi"])
                self.lbl_resim_bilgi.config(text=f"Bu kargoda {n} resim var.")
                # Borç panelini güncelle
                tutar = float(r.get("odeme_tutari") or 0)
                if tutar > 0:
                    kdv = tutar * 0.20
                    self.lbl_borc_tutar.config(
                        text=f"Odeme tutari: {tutar:,.2f} TL  (KDV dahil: {tutar+kdv:,.2f} TL)",
                        foreground="#ff8a65")
                    durum = r.get("odeme_durumu", "bekliyor")
                    if durum == "odendi":
                        self.lbl_borc_durum.config(text="ODENDI", foreground="#69f0ae")
                    else:
                        self.lbl_borc_durum.config(text="BEKLIYOR", foreground="#ff8a65")
                    fno = r.get("kargo_fatura_no") or r.get("fatura_id") or ""
                    self.lbl_borc_fatura.config(
                        text=f"Cari kayit: {fno}  |  Musteri borcuna eklendi")
                else:
                    self.lbl_borc_tutar.config(text="Ucretsiz kargo", foreground="#69f0ae")
                    self.lbl_borc_durum.config(text="")
                    self.lbl_borc_fatura.config(text="")
                break

    def _resim_sec(self):
        from tkinter import filedialog
        dosyalar = filedialog.askopenfilenames(
            title="Kargo Resimlerini Sec",
            filetypes=[("Resimler", "*.jpg *.jpeg *.png *.gif *.bmp *.webp"),
                       ("Tum Dosyalar", "*.*")])
        for d in dosyalar:
            if d not in [item for item in self._yeni_resimler]:
                self._yeni_resimler.append(d)
                from pathlib import Path
                self.resim_listbox.insert(tk.END, Path(d).name)

    def _resim_kaldir(self):
        sel = self.resim_listbox.curselection()
        if not sel:
            return
        idx = sel[0]
        self.resim_listbox.delete(idx)
        if 0 <= idx < len(self._yeni_resimler):
            self._yeni_resimler.pop(idx)

    def _kargo_kaydet(self):
        tarih   = self.v_k_tarih.get().strip()
        firma   = self.v_k_firma.get().strip()
        teslim  = self.v_k_teslim.get().strip()
        takip   = self.v_k_takip.get().strip()
        notlar  = self.v_k_notlar.get().strip()

        if not tarih:
            messagebox.showwarning("Uyari", "Tarih giriniz!", parent=self)
            return

        odeme_str = self.v_k_odeme.get().strip().replace(",", ".")
        try:
            odeme_tutari = float(odeme_str) if odeme_str else 0.0
        except:
            odeme_tutari = 0.0

        from database import kargo_ekle, kargo_resim_ekle
        import shutil
        try:
            # Tarihi ISO formatına çevir
            try:
                from datetime import datetime
                if "." in tarih:
                    dt = datetime.strptime(tarih, "%d.%m.%Y")
                    tarih_iso = dt.strftime("%Y-%m-%d")
                else:
                    tarih_iso = tarih
            except:
                tarih_iso = tarih

            kargo_id = kargo_ekle(self.musteri_id, tarih_iso, teslim, firma, takip, notlar,
                                   odeme_tutari=odeme_tutari)

            # Resimleri kopyala ve kaydet
            for dosya_yolu in self._yeni_resimler:
                from pathlib import Path
                hedef = self._resim_dir / Path(dosya_yolu).name
                # Aynı isimde varsa numara ekle
                sayac = 1
                while hedef.exists():
                    hedef = self._resim_dir / f"{Path(dosya_yolu).stem}_{sayac}{Path(dosya_yolu).suffix}"
                    sayac += 1
                shutil.copy2(dosya_yolu, hedef)
                kargo_resim_ekle(kargo_id, str(hedef), hedef.name)

            messagebox.showinfo("Kaydedildi",
                f"Kargo kaydedildi!\n{len(self._yeni_resimler)} resim eklendi.",
                parent=self)

            # Formu sıfırla
            from datetime import date
            self.v_k_tarih.set(date.today().strftime("%d.%m.%Y"))
            self.v_k_firma.set("")
            self.v_k_teslim.set("")
            self.v_k_takip.set("")
            self.v_k_notlar.set("")
            self.v_k_odeme.set("")
            self._yeni_resimler.clear()
            self.resim_listbox.delete(0, tk.END)

            self._kargo_yukle()

        except Exception as e:
            import traceback; traceback.print_exc()
            messagebox.showerror("Hata", str(e), parent=self)

    def _whatsapp_gonder(self):
        if not self._secili_kargo_id:
            messagebox.showwarning("Uyari", "Önce kargo seçin!", parent=self)
            return
        if not self.telefon:
            messagebox.showwarning("Uyari",
                "Bu müşterinin telefon numarası kayıtlı değil!\n"
                "Müşteri kaydına telefon ekleyin.", parent=self)
            return

        from database import get_kargolar, kargo_whatsapp_guncelle
        rows = get_kargolar(self.musteri_id)
        kargo = next((r for r in rows if r["id"] == self._secili_kargo_id), None)
        if not kargo:
            return

        # Resmi panoya kopyala
        resimler = kargo["resim_listesi"]
        if resimler:
            try:
                import subprocess, os
                ilk_resim = resimler[0]
                if os.path.exists(ilk_resim):
                    # Windows'ta resmi panoya kopyala
                    if os.name == "nt":
                        from PIL import Image
                        import io
                        import win32clipboard
                        img = Image.open(ilk_resim)
                        output = io.BytesIO()
                        img.convert("RGB").save(output, "BMP")
                        data = output.getvalue()[14:]  # BMP header'ı atla
                        output.close()
                        win32clipboard.OpenClipboard()
                        win32clipboard.EmptyClipboard()
                        win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
                        win32clipboard.CloseClipboard()
                        messagebox.showinfo("Resim Kopyalandi",
                            f"{len(resimler)} resimden ilki panoya kopyalandı.\n"
                            "WhatsApp Web açılıyor — Ctrl+V ile yapıştırın.",
                            parent=self)
            except ImportError:
                # win32clipboard yoksa sadece bilgi ver
                messagebox.showinfo("Bilgi",
                    f"Resim konumu:\n{resimler[0]}\n\n"
                    "WhatsApp Web açıldıktan sonra resmi manuel ekleyin.",
                    parent=self)
            except Exception as e:
                messagebox.showwarning("Resim Hatasi", str(e), parent=self)

        # WhatsApp Web URL'si oluştur ve aç
        tel = self.telefon.strip().replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
        if tel.startswith("0"):
            tel = "90" + tel[1:]
        elif not tel.startswith("90"):
            tel = "90" + tel

        tarih_str = kargo.get("tarih", "")
        mesaj = f"Sayın {self.musteri_adi}, {tarih_str} tarihinde kargolorunuz teslim alınmıştır."
        if kargo.get("teslim_alan"):
            mesaj += f" Teslim alan: {kargo['teslim_alan']}."

        import urllib.parse, webbrowser
        url = f"https://wa.me/{tel}?text={urllib.parse.quote(mesaj)}"
        webbrowser.open(url)

        # Gönderildi olarak işaretle
        kargo_whatsapp_guncelle(self._secili_kargo_id, True)
        self._kargo_yukle()

    def _resimleri_goster(self):
        if not self._secili_kargo_id:
            messagebox.showwarning("Uyari", "Önce kargo seçin!", parent=self)
            return
        from database import get_kargolar
        rows = get_kargolar(self.musteri_id)
        kargo = next((r for r in rows if r["id"] == self._secili_kargo_id), None)
        if not kargo or not kargo["resim_listesi"]:
            messagebox.showinfo("Bilgi", "Bu kargoda resim yok.", parent=self)
            return

        # Resim görüntüleme penceresi
        rp = tk.Toplevel(self)
        rp.title("Kargo Resimleri")
        rp.geometry("700x520")
        rp.transient(self)

        nb = ttk.Notebook(rp)
        nb.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)

        import os
        from pathlib import Path
        for dosya in kargo["resim_listesi"]:
            if not os.path.exists(dosya):
                continue
            f = ttk.Frame(nb)
            nb.add(f, text=Path(dosya).name[:15])
            try:
                from PIL import Image, ImageTk
                img = Image.open(dosya)
                img.thumbnail((660, 450))
                photo = ImageTk.PhotoImage(img)
                lbl = ttk.Label(f, image=photo)
                lbl.image = photo  # referans tut
                lbl.pack(expand=True)
            except ImportError:
                ttk.Label(f, text=f"PIL kurulu degil.\nDosya: {dosya}").pack(expand=True)
                ttk.Button(f, text="Dosyayi Ac",
                           command=lambda p=dosya: os.startfile(p) if os.name=="nt" else None
                           ).pack()
            except Exception as e:
                ttk.Label(f, text=f"Resim yuklenemedi: {e}\n{dosya}").pack(expand=True)

        ttk.Button(rp, text="Kapat", command=rp.destroy).pack(pady=4)

    def _kargo_odeme_guncelle(self, durum: str):
        """Seçili kargonun ödeme durumunu güncelle."""
        if not self._secili_kargo_id:
            messagebox.showwarning("Uyari", "Once kargo secin!", parent=self)
            return
        from database import kargo_odeme_guncelle
        kargo_odeme_guncelle(self._secili_kargo_id, durum)
        mesaj = "Odendi olarak isaretlendi!" if durum == "odendi" else "Bekliyor olarak isaretlendi."
        messagebox.showinfo("Guncellendi", mesaj, parent=self)
        self._kargo_yukle()
        # Seçimi koru
        try:
            self.kargo_tree.selection_set(str(self._secili_kargo_id))
            self._kargo_sec()
        except: pass

    def _kargo_sil(self):
        if not self._secili_kargo_id:
            messagebox.showwarning("Uyari", "Önce kargo seçin!", parent=self)
            return
        if not messagebox.askyesno("Sil", "Bu kargo kaydı silinsin mi?", parent=self):
            return
        from database import kargo_sil
        kargo_sil(self._secili_kargo_id)
        self._secili_kargo_id = None
        self._kargo_yukle()


# ============================================================================
# TÜFE SEKMESİ
# ============================================================================

class TufeTab(ttk.Frame):
    """TÜFE Yönetimi Sekmesi — TCMB otomatik çekme + manuel giriş + sil."""

    MONTHS = ["Ocak","Şubat","Mart","Nisan","Mayıs","Haziran",
               "Temmuz","Ağustos","Eylül","Ekim","Kasım","Aralık"]

    def __init__(self, master: tk.Widget) -> None:
        super().__init__(master, padding=10)
        self._build_ui()

    def _build_ui(self) -> None:
        from datetime import date as _date
        import threading

        self._today = _date.today()
        self._thread_running = False

        # ── Üst araç çubuğu ───────────────────────────────────────────────
        top = ttk.Frame(self); top.pack(side=tk.TOP, fill=tk.X, pady=(0, 8))

        ttk.Label(top, text="TÜFE Oranları",
                  font=("Segoe UI", 13, "bold")).pack(side=tk.LEFT)
        ttk.Label(top, text="  Yıl:").pack(side=tk.LEFT, padx=(8, 2))
        self.year_var = tk.IntVar(value=self._today.year)
        tk.Spinbox(top, from_=2003, to=self._today.year, width=7,
                   textvariable=self.year_var, font=("Segoe UI", 9),
                   command=self._on_load).pack(side=tk.LEFT)

        ttk.Button(top, text="📂 Yükle",
                   command=self._on_load).pack(side=tk.LEFT, padx=(10, 0))
        ttk.Button(top, text="💾 Kaydet",
                   command=self._on_save).pack(side=tk.LEFT, padx=(4, 0))
        ttk.Button(top, text="🗑 Sil",
                   command=self._on_delete).pack(side=tk.LEFT, padx=(4, 0))
        ttk.Button(top, text="📊 Excel'den Aktar",
                   command=self._on_import_excel).pack(side=tk.LEFT, padx=(4, 0))

        self.lbl_tcmb = ttk.Label(top, text="", font=("Segoe UI", 8),
                                   foreground="#888888")
        self.lbl_tcmb.pack(side=tk.LEFT, padx=(12, 0))

        ttk.Button(top, text="🌐 TCMB'den Çek",
                   command=self._tcmb_cek).pack(side=tk.RIGHT, padx=(4, 0))

        # ── Ay grid'i ──────────────────────────────────────────────────────
        self.month_vars: dict = {}
        self.month_entries: dict = {}
        grid = ttk.Frame(self)
        grid.pack(side=tk.TOP, fill=tk.X, pady=(0, 6))

        for i, month_name in enumerate(self.MONTHS):
            row, col = divmod(i, 3)
            frame = ttk.LabelFrame(grid, text=month_name + " (%)", padding=8)
            frame.grid(row=row, column=col, sticky="ew", padx=6, pady=4)
            grid.columnconfigure(col, weight=1)
            var = tk.StringVar()
            self.month_vars[i + 1] = var
            entry = ttk.Entry(frame, textvariable=var, width=14, justify="right")
            entry.pack()
            self.month_entries[i + 1] = entry

        # İzin verilen ay aralığını güncelle
        self.year_var.trace_add("write", lambda *_: self._izin_guncelle())
        self._izin_guncelle()

        # ── Alt bilgi ──────────────────────────────────────────────────────
        self.info_label = ttk.Label(
            self, text="", foreground="#69f0ae", font=("Segoe UI", 9))
        self.info_label.pack(side=tk.BOTTOM, anchor="w", pady=(4, 0))

        ttk.Label(self,
            text="ℹ  Yalnızca geçmiş ve mevcut aya ait veriler girilebilir. "
                 "Bir sonraki ayın verisi henüz açıklanmadığı için girilemez. "
                 "TCMB her ayın ilk 3 günü veriyi yayımlar.",
            font=("Segoe UI", 8), foreground="#888888",
            wraplength=900).pack(side=tk.BOTTOM, anchor="w")

        # Açılışta mevcut yılı yükle
        self._on_load()

    def _izin_guncelle(self, *_):
        """Sadece bugüne kadar olan ayları düzenlenebilir yap."""
        try: yil = int(self.year_var.get())
        except: return
        t = self._today
        for ay in range(1, 13):
            entry = self.month_entries[ay]
            # Mevcut yıl → sadece açıklanan aylara izin (bugünün ayı dahil)
            # Geçmiş yıllar → tüm aylar serbest
            # Gelecek yıllar → hiçbir ay serbest değil
            if yil < t.year:
                entry.configure(state="normal")
            elif yil == t.year:
                if ay <= t.month:
                    entry.configure(state="normal")
                else:
                    entry.configure(state="disabled")
                    self.month_vars[ay].set("")
            else:
                entry.configure(state="disabled")
                self.month_vars[ay].set("")

    def refresh(self) -> None:
        self._on_load()

    def _get_year(self):
        try: return int(self.year_var.get())
        except: return None

    def _on_load(self) -> None:
        from database import get_tufe_for_year
        year = self._get_year()
        if not year: return
        data = get_tufe_for_year(year)
        for i, month_name in enumerate(self.MONTHS):
            val = data.get(month_name)
            self.month_vars[i + 1].set(str(val) if val is not None else "")
        self._izin_guncelle()
        dolu = sum(1 for v in self.month_vars.values() if v.get().strip())
        if data:
            self.info_label.config(
                text=f"✓ {year} yılı yüklendi — {dolu} ay kayıtlı.",
                foreground="#69f0ae")
        else:
            self.info_label.config(
                text=f"{year} yılına ait kayıt bulunamadı.",
                foreground="#888888")

    def _on_save(self) -> None:
        from database import save_tufe_for_year
        # kira_senaryo.py'daki AYLIK_YILLIK'i de güncelle
        year = self._get_year()
        if not year:
            messagebox.showwarning("Uyarı", "Geçerli bir yıl seçin!"); return
        rates = {}
        for i, month_name in enumerate(self.MONTHS):
            raw = self.month_vars[i + 1].get().strip().replace(",", ".")
            if not raw: continue
            try:
                rates[month_name] = float(raw)
            except ValueError:
                messagebox.showwarning("Uyarı", f"{month_name} için geçersiz oran!")
                return
        if not rates:
            messagebox.showwarning("Uyarı", "En az bir oran girin."); return
        save_tufe_for_year(year, rates)
        # kira_senaryo modülündeki cache'e de yaz
        try:
            import kira_senaryo as ks
            ay_isimleri = ks.AYLAR_TR
            for ay_idx, month_name in enumerate(self.MONTHS):
                if month_name in rates:
                    ks.AYLIK_YILLIK[(year, ay_idx + 1)] = rates[month_name]
        except: pass
        self.info_label.config(
            text=f"✓ {year} yılı kaydedildi — {len(rates)} ay.",
            foreground="#69f0ae")
        messagebox.showinfo("Başarılı", f"{year} yılı TÜFE oranları kaydedildi.")

    def _on_delete(self) -> None:
        year = self._get_year()
        if not year: return
        # Hangi ayı sil?
        pop = tk.Toplevel(self)
        pop.title("Hangi veriyi silmek istiyorsunuz?")
        pop.geometry("320x280"); pop.transient(self); pop.grab_set()
        ttk.Label(pop, text=f"{year} yılı — Silinecek ayları seçin:",
                  font=("Segoe UI",9,"bold")).pack(pady=(10,4))
        check_vars = {}
        cf = ttk.Frame(pop); cf.pack(fill=tk.X, padx=16)
        for i, mn in enumerate(self.MONTHS):
            var = tk.IntVar(value=0)
            check_vars[i+1] = var
            ttk.Checkbutton(cf, text=mn, variable=var).grid(
                row=i//3, column=i%3, sticky="w", padx=4, pady=2)
        bf = ttk.Frame(pop); bf.pack(pady=8)
        def _hepsini():
            for v in check_vars.values(): v.set(1)
        def _sil():
            from database import get_connection
            secilen = [i for i,v in check_vars.items() if v.get()]
            if not secilen:
                messagebox.showwarning("Uyarı","En az bir ay seçin!",parent=pop); return
            ay_isimleri = self.MONTHS
            try:
                conn = get_connection()
                for ay in secilen:
                    conn.execute(
                        "DELETE FROM tufe_verileri WHERE year=? AND month=?",
                        (year, ay_isimleri[ay-1]))
                conn.commit(); conn.close()
                # kira_senaryo cache temizle
                try:
                    import kira_senaryo as ks
                    for ay in secilen:
                        ks.AYLIK_YILLIK.pop((year,ay), None)
                except: pass
            except Exception as e:
                messagebox.showerror("Hata",str(e),parent=pop); return
            pop.destroy()
            self._on_load()
            messagebox.showinfo("Tamam",f"{len(secilen)} ay verisi silindi.")
        ttk.Button(bf, text="Tümünü Seç", command=_hepsini).pack(side=tk.LEFT, padx=4)
        ttk.Button(bf, text="🗑 Sil", command=_sil).pack(side=tk.LEFT, padx=4)
        ttk.Button(bf, text="İptal", command=pop.destroy).pack(side=tk.LEFT, padx=4)

    def _tcmb_cek(self) -> None:
        """TCMB web sitesinden yıllık TÜFE verilerini çek."""
        import threading
        if self._thread_running: return
        self._thread_running = True
        year = self._get_year()
        self.lbl_tcmb.config(text="⏳ TCMB'den çekiliyor...", foreground="#fff176")
        threading.Thread(target=self._tcmb_thread, args=(year,), daemon=True).start()

    def _tcmb_thread(self, hedef_yil):
        try:
            import urllib.request, re
            url = ("https://www.tcmb.gov.tr/wps/wcm/connect/TR/TCMB+TR/"
                   "Main+Menu/Istatistikler/Enflasyon+Verileri/Tuketici+Fiyatlari")
            req = urllib.request.Request(url, headers={
                "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/120"})
            html = urllib.request.urlopen(req, timeout=12).read().decode("utf-8","ignore")
            # | MM-YYYY | yillik | aylik |
            pattern = r"\|\s*(\d{2})-(\d{4})\s*\|\s*([\d.]+)\s*\|"
            matches = re.findall(pattern, html)
            yil_data = {}  # {yil: {ay: oran}}
            for ay_s, yil_s, oran_s in matches:
                ay = int(ay_s); yil = int(yil_s); oran = float(oran_s)
                yil_data.setdefault(yil, {})[ay] = oran
                # kira_senaryo cache'e yaz
                try:
                    import kira_senaryo as ks
                    ks.AYLIK_YILLIK[(yil, ay)] = oran
                except: pass
            # DB'ye kaydet
            if hedef_yil in yil_data:
                from database import save_tufe_for_year
                ay_isimleri = self.MONTHS
                rates = {ay_isimleri[ay-1]: oran
                         for ay, oran in yil_data[hedef_yil].items()}
                if rates:
                    save_tufe_for_year(hedef_yil, rates)
            self.after(0, lambda: self._tcmb_bitti(yil_data, hedef_yil))
        except Exception as e:
            self.after(0, lambda: self._tcmb_hata(str(e)))

    def _tcmb_bitti(self, yil_data, hedef_yil):
        self._thread_running = False
        if hedef_yil in yil_data:
            say = len(yil_data[hedef_yil])
            self.lbl_tcmb.config(
                text=f"✓ {hedef_yil} için {say} ay TCMB'den yüklendi.",
                foreground="#69f0ae")
            self._on_load()
        else:
            self.lbl_tcmb.config(
                text=f"⚠ {hedef_yil} için TCMB'de veri bulunamadı.",
                foreground="#ff8a65")

    def _tcmb_hata(self, hata):
        self._thread_running = False
        self.lbl_tcmb.config(
            text=f"⚠ TCMB erişim hatası: {hata[:50]}",
            foreground="#ff8a65")

    def _on_import_excel(self) -> None:
        from database import import_tufe_from_excel
        from tkinter import filedialog
        path = filedialog.askopenfilename(
            title="TÜFE Excel Dosyası Seç",
            filetypes=[("Excel Dosyaları", "*.xlsx *.xls"), ("Tüm Dosyalar", "*.*")])
        if not path: return
        try:
            count = import_tufe_from_excel(path)
            messagebox.showinfo("Başarılı", f"{count} TÜFE kaydı içe aktarıldı.")
            self._on_load()
        except Exception as e:
            messagebox.showerror("Hata", f"Excel aktarma hatası: {e}")

