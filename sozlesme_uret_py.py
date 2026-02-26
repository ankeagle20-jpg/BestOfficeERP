"""
Hazır Ofis / Sanal Ofis Hizmet Sözleşmesi Üretici — python-docx
Orijinal sözleşme maddeleri birebir korunmuştur.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def _alt_cizgi(p, renk="555555", kalinlik=4):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(kalinlik))
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), renk)
    pBdr.append(b)
    pPr.append(pBdr)


def _ust_cizgi(p, renk="555555", kalinlik=4):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:top')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(kalinlik))
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), renk)
    pBdr.append(b)
    pPr.append(pBdr)


def _tablo_kenarlari_gizle(tablo):
    for satir in tablo.rows:
        for hucre in satir.cells:
            tc = hucre._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for k in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                el = OxmlElement(f'w:{k}')
                el.set(qn('w:val'), 'nil')
                tcBorders.append(el)
            tcPr.append(tcBorders)


def sozlesme_olustur(data: dict, cikti: str) -> str:
    doc = Document()

    # ── Sayfa ────────────────────────────────────────────────────────────────
    s = doc.sections[0]
    s.page_width   = Cm(21);   s.page_height  = Cm(29.7)
    s.left_margin  = Cm(2.5);  s.right_margin = Cm(2.5)
    s.top_margin   = Cm(2);    s.bottom_margin = Cm(2)

    F  = "Times New Roman"
    SZ = 11   # normal punto

    # ── Yardımcılar ──────────────────────────────────────────────────────────
    def p(metin="", bold=False, sz=SZ, align=WD_ALIGN_PARAGRAPH.LEFT,
          oncesi=0, sonrasi=3):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(oncesi)
        para.paragraph_format.space_after  = Pt(sonrasi)
        if metin:
            r = para.add_run(metin)
            r.bold = bold; r.font.size = Pt(sz); r.font.name = F
        return para

    def p2(parcalar, oncesi=0, sonrasi=3):
        """parcalar = [(metin, bold), ...]"""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(oncesi)
        para.paragraph_format.space_after  = Pt(sonrasi)
        for metin, bold in parcalar:
            r = para.add_run(metin)
            r.bold = bold; r.font.size = Pt(SZ); r.font.name = F
        return para

    def madde(no, baslik):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after  = Pt(3)
        r = para.add_run(f"MADDE {no} – {baslik}")
        r.bold = True; r.font.size = Pt(SZ); r.font.name = F
        return para

    def bolum(metin):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after  = Pt(4)
        r = para.add_run(metin)
        r.bold = True; r.font.size = Pt(SZ); r.font.name = F
        _alt_cizgi(para)
        _ust_cizgi(para)
        return para

    def bilgi(etiket, deger, noktalama=True):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        r1 = para.add_run(f"{etiket}: ")
        r1.bold = True; r1.font.size = Pt(SZ); r1.font.name = F
        val = deger if deger else ("." * 55 if noktalama else "")
        r2 = para.add_run(val)
        r2.font.size = Pt(SZ); r2.font.name = F
        return para

    def imza_satiri(etiket):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        r1 = para.add_run(f"{etiket}: ")
        r1.bold = True; r1.font.size = Pt(SZ); r1.font.name = F
        r2 = para.add_run("." * 55)
        r2.font.size = Pt(SZ); r2.font.name = F
        return para

    def pj(metin, sz=SZ, oncesi=0, sonrasi=3):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(oncesi)
        para.paragraph_format.space_after  = Pt(sonrasi)
        r = para.add_run(metin)
        r.font.size = Pt(sz); r.font.name = F
        return para

    # ── Veriler ──────────────────────────────────────────────────────────────
    soz_no   = data.get("sozlesme_no", "SZL-2026-0001")
    tarih    = data.get("sozlesme_tarihi") or datetime.now().strftime("%d.%m.%Y")
    bitis    = data.get("sozlesme_bitis", "")
    hiz      = data.get("hizmet_turu", "Sanal Ofis")
    ofis_k   = data.get("ofis_kodu", "")
    unvan    = data.get("sirket_unvani", "")
    vno      = data.get("vergi_no", "")
    vdr      = data.get("vergi_dairesi", "")
    mersis   = data.get("mersis_no", "")
    sicil    = data.get("ticaret_sicil_no", "")
    faaliyet = data.get("faaliyet_konusu", "")
    adres    = data.get("yeni_adres", "")
    eski_adr = data.get("eski_adres", "")

    yt_ad    = data.get("yetkili_adsoyad", "")
    yt_tc    = data.get("yetkili_tcno", "")
    yt_dog   = data.get("yetkili_dogum", "")
    yt_ik    = data.get("yetkili_ikametgah", "")
    yt_tel   = data.get("yetkili_tel", "")
    yt_tel2  = data.get("yetkili_tel2", "")
    yt_mail  = data.get("yetkili_email", "")

    o1_ad    = data.get("ortak1_adsoyad", "")
    o1_pay   = data.get("ortak1_pay", "")
    o2_ad    = data.get("ortak2_adsoyad", "")
    o2_pay   = data.get("ortak2_pay", "")
    o3_ad    = data.get("ortak3_adsoyad", "")
    o3_pay   = data.get("ortak3_pay", "")
    yab_ad   = data.get("yabanci_adsoyad", "")
    yab_uy   = data.get("yabanci_uyruk", "")
    yab_pas  = data.get("yabanci_pasaport", "")

    sanal = "sanal" in hiz.lower()

    try:    aylik  = f"{float(data.get('aylik_kira') or 0):,.2f}"
    except: aylik  = str(data.get("aylik_kira", ""))
    try:    yillik = f"{float(data.get('yillik_kira') or 0):,.2f}"
    except: yillik = str(data.get("yillik_kira", ""))

    # ═════════════════════════════════════════════════════════════════════════
    # BELGE İÇERİĞİ
    # ═════════════════════════════════════════════════════════════════════════

    # Başlık
    p("HAZIR OFİS / SANAL OFİS", bold=True, sz=16,
      align=WD_ALIGN_PARAGRAPH.CENTER, oncesi=0, sonrasi=2)
    p("ADRES KULLANIM VE HİZMET SÖZLEŞMESİ", bold=True, sz=14,
      align=WD_ALIGN_PARAGRAPH.CENTER, oncesi=0, sonrasi=2)
    para = p(f"Sözleşme No: {soz_no}  |  Tarih: {tarih}",
             sz=10, align=WD_ALIGN_PARAGRAPH.CENTER, sonrasi=8)
    _alt_cizgi(para, "000000", 8)

    # ── MADDE 1 ──────────────────────────────────────────────────────────────
    madde(1, "TARAFLAR")

    p("Hizmet Veren:", bold=True, oncesi=4, sonrasi=2)
    p("Ofisbir Ofis ve Danışmanlık Hizmetleri A.Ş.", sonrasi=2)
    p("Adres: Kavaklıdere Mah. Esat Caddesi No:12 İç Kapı No:1 Çankaya/Ankara", sonrasi=2)
    p("Vergi No: 6340871926", sonrasi=2)
    p('(İşbu sözleşmede "OFİSBİR" olarak anılacaktır.)', sonrasi=6)

    bolum("HİZMET ALAN (ŞİRKET BİLGİLERİ)")
    bilgi("Unvan", unvan)
    bilgi("Vergi No", vno)
    bilgi("Vergi Dairesi", vdr)
    bilgi("MERSİS No", mersis)
    bilgi("Ticaret Sicil No", sicil)
    bilgi("Faaliyet Konusu", faaliyet)
    bilgi("Merkez Adresi", adres)
    if eski_adr:
        bilgi("Önceki Adres", eski_adr)
    if ofis_k:
        bilgi("Ofis/Masa Kodu", ofis_k)

    bolum("YETKİLİ KİŞİ BİLGİLERİ")
    bilgi("Ad Soyad", yt_ad)
    bilgi("T.C. Kimlik No", yt_tc)
    bilgi("Doğum Tarihi", yt_dog)
    bilgi("İkamet Adresi", yt_ik)
    bilgi("Cep Telefonu", yt_tel)
    if yt_tel2:
        bilgi("Cep Telefonu 2", yt_tel2)
    bilgi("E-Posta", yt_mail)
    imza_satiri("Islak İmza")

    bolum("ORTAKLIK BİLGİLERİ")
    bilgi("Ortak 1 Ad Soyad / Unvan", o1_ad)
    bilgi("Pay Oranı (%)", f"%{o1_pay}" if o1_pay else "")
    p("")
    bilgi("Ortak 2 Ad Soyad / Unvan", o2_ad)
    bilgi("Pay Oranı (%)", f"%{o2_pay}" if o2_pay else "")
    p("")
    bilgi("Ortak 3 Ad Soyad / Unvan", o3_ad)
    bilgi("Pay Oranı (%)", f"%{o3_pay}" if o3_pay else "")

    if yab_ad or yab_uy or yab_pas:
        p("Yabancı Ortak Varsa:", bold=True, oncesi=6, sonrasi=2)
        bilgi("Ad Soyad", yab_ad)
        bilgi("Uyruğu", yab_uy)
        bilgi("Pasaport No", yab_pas)
    else:
        p("Yabancı Ortak Varsa:", bold=True, oncesi=6, sonrasi=2)
        bilgi("Ad Soyad", "")
        bilgi("Uyruğu", "")
        bilgi("Pasaport No", "")

    p("")
    pj('(İşbu sözleşmede "MÜŞTERİ" olarak anılacaktır.)', oncesi=4, sonrasi=2)
    pj("Müşteri adına imza atan yetkili, şirket ile birlikte müştereken ve müteselsilen sorumludur.",
       sonrasi=6)

    # ── MADDE 2 ──────────────────────────────────────────────────────────────
    madde(2, "HİZMET TÜRÜ")
    pj("Taraflar aşağıdaki hizmet türlerinden birini seçmiştir:", sonrasi=4)

    p(f"{'☑' if sanal else '☐'} SANAL OFİS HİZMETİ", bold=True, sonrasi=2)
    p(f"{'☑' if not sanal else '☐'} HAZIR OFİS HİZMETİ", bold=True, sonrasi=4)
    pj("(Seçilen hizmet türü sözleşmenin ayrılmaz parçasıdır.)", sonrasi=6)

    # ── BÖLÜM A – SANAL OFİS ─────────────────────────────────────────────────
    if sanal:
        bolum("BÖLÜM A – SANAL OFİS HİZMETİ")

        madde("3A", "KAPSAM")
        pj("Sanal ofis hizmeti; yasal adres tahsisi, posta/kargo/tebligat teslim alma ve "
           "sekreterya bilgilendirme hizmetlerini kapsar.", sonrasi=4)
        pj("Bu sözleşme kira sözleşmesi değildir. Taşınmaz üzerinde kiracılık hakkı doğurmaz. "
           "Ancak MÜŞTERİ'ye sözleşme süresince yasal adres kullanım hakkı verir.", sonrasi=4)
        pj("İşbu sözleşme kapsamında MÜŞTERİ'ye yasal adres kullanım hakkı verilmiş olup, "
           "bu adres vergi mevzuatı çerçevesinde işyeri adresi olarak bildirilebilir.", sonrasi=6)

        madde("4A", "Fiziki Kullanım")
        pj("Sanal ofis müşterisi sürekli masa veya oda kullanım hakkına sahip değildir. "
           "Ofise eşya bırakamaz ve ticari mal bulunduramaz.", sonrasi=6)

        madde("5A", "Haciz Güvencesi")
        pj("MÜŞTERİ, ofis adresinde kendisine ait mal bulunmadığını, ofisteki tüm demirbaşların "
           "OFİSBİR'e ait olduğunu ve haciz halinde OFİSBİR'in üçüncü kişi olduğunu kabul eder.",
           sonrasi=4)
        pj("Haciz bildirgesi gelmesi halinde sözleşme kendiliğinden feshedilir.", sonrasi=6)

    # ── BÖLÜM B – HAZIR OFİS ─────────────────────────────────────────────────
    else:
        bolum("BÖLÜM B – HAZIR OFİS HİZMETİ")

        madde("3B", "KAPSAM")
        pj("Hazır ofis hizmeti; belirlenen oda/masa kullanımı, sekreterya hizmeti, "
           "ortak alan kullanımı ve yasal adres hizmetlerini kapsar.", sonrasi=4)
        pj("Bu sözleşme kira sözleşmesi değildir; hizmet sözleşmesidir.", sonrasi=6)

        madde("4B", "Demirbaşlar")
        pj("Ofiste bulunan tüm mobilya, masa, sandalye, dolap ve ekipman OFİSBİR mülkiyetindedir.",
           sonrasi=4)
        pj("MÜŞTERİ yalnızca şahsi bilgisayar ve küçük kişisel eşyalarını bulundurabilir.",
           sonrasi=6)

        madde("5B", "Haciz")
        pj("MÜŞTERİ borçları nedeniyle haciz gelmesi halinde OFİSBİR'in üçüncü kişi olduğunu "
           "kabul eder ve doğabilecek zararları tazmin etmeyi taahhüt eder.", sonrasi=4)
        pj("Haciz bildirgesi fesih sebebidir.", sonrasi=6)

    # ── ORTAK HÜKÜMLER ────────────────────────────────────────────────────────
    bolum("ORTAK HÜKÜMLER")

    madde(6, "HİZMET BEDELİ")
    bilgi("Yıllık Hizmet Bedeli", f"{yillik} TL + KDV" if yillik else "")
    bilgi("Aylık Hizmet Bedeli", f"{aylik} TL + KDV" if aylik else "")
    p("")
    pj("Ödemeler aylık olarak OFİSBİR'in bildirdiği banka hesabına yapılacaktır.", sonrasi=2)
    pj("İki aylık ödeme gecikmesi halinde sözleşme tek taraflı feshedilebilir.", sonrasi=6)

    madde(7, "ERKEN FESİH")
    pj("MÜŞTERİ, sözleşme süresi dolmadan ayrılmak isterse yazılı bildirim yapmak "
       "kaydıyla sözleşmesini feshedebilir.", sonrasi=4)
    pj("Erken fesih halinde 2 (iki) aylık hizmet bedeli tutarında erken fesih bedeli "
       "ödemeyi kabul eder.", sonrasi=4)
    pj("Bu bedel makul cezai şart niteliğindedir.", sonrasi=6)

    madde(8, "OTOMATİK YENİLEME")
    pj("Sözleşme bitiminden 15 gün önce yazılı fesih yapılmazsa 1 yıl süreyle aynı "
       "şartlarla yenilenir.", sonrasi=6)

    madde(9, "MÜTESELSİL SORUMLULUK")
    pj("Şirket yetkilisi işbu sözleşmeden doğan borçlardan şirket ile birlikte müştereken "
       "ve müteselsilen sorumludur.", sonrasi=6)

    madde(10, "YETKİLİ MAHKEME")
    pj("İşbu sözleşmeden doğacak uyuşmazlıklarda Ankara Mahkemeleri ve İcra Daireleri "
       "yetkilidir.", sonrasi=6)

    madde(11, "YÜRÜRLÜK")
    pj(f"İşbu sözleşme {tarih} tarihinde iki nüsha olarak düzenlenmiş ve imza altına "
       "alınmıştır.", sonrasi=10)

    # ── İMZA TABLOSU (3 kişi) ─────────────────────────────────────────────────
    tablo = doc.add_table(rows=3, cols=3)
    tablo.style = "Table Grid"
    W = [3000, 3000, 3000]  # DXA yaklaşık

    basliklar = ["OFİSBİR", "MÜŞTERİ", "Yetkili (Müteselsil Sorumlu)"]
    for i, b in enumerate(basliklar):
        hucre = tablo.cell(0, i)
        hucre.text = b
        for para in hucre.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True; run.font.size = Pt(SZ); run.font.name = F

    # Boş imza satırı (orta satır — boşluk)
    for i in range(3):
        hucre = tablo.cell(1, i)
        hucre.text = "\n\n\n"

    # İmza çizgisi
    for i in range(3):
        hucre = tablo.cell(2, i)
        hucre.text = "İmza:"
        for para in hucre.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(SZ); run.font.name = F

    _tablo_kenarlari_gizle(tablo)

    # Alt not
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    r = para.add_run(f"Sözleşme No: {soz_no}  |  {tarih}")
    r.font.size = Pt(9); r.font.name = F
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(cikti)
    return cikti
