from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file
from auth import yetki_gerekli, giris_gerekli
from db import (
    fetch_all,
    fetch_one,
    execute,
    execute_returning,
    ensure_faturalar_amount_columns,
    ensure_customers_durum,
    ensure_customers_quick_edit_columns,
    ensure_customers_kapanis_tarihi,
    ensure_customers_hazir_ofis_oda,
    ensure_customers_grup2_secimleri,
    ensure_grup2_bizim_hesap_into_array,
    ensure_grup2_etiketleri_table,
    ensure_musteri_kyc_hazir_ofis_oda_no,
    ensure_musteri_kyc_kira_banka,
    ensure_musteri_kyc_odeme_duzeni,
    db as get_db,
    clear_all_customers,
    get_conn,
)
from utils.musteri_arama import customers_arama_sql_giris_genis, customers_arama_params_giris_genis
import pandas as pd
import calendar
import json
from io import BytesIO
from datetime import date, datetime, timedelta
from docx import Document
import os
import sys
import re
from services.cari_service import CariService

_HAZIR_OFIS_ODA_MIN, _HAZIR_OFIS_ODA_MAX = 200, 230


def _kyc_optional_money(val):
    """Formdan opsiyonel para; boş → None, negatif → None."""
    if val is None:
        return None
    s = str(val).strip().replace(",", ".")
    if not s:
        return None
    try:
        n = float(s)
        if n < 0:
            return None
        return round(n, 2)
    except (TypeError, ValueError):
        return None


def _hizmet_str_hazir_ofis_mi(ht: str) -> bool:
    s = (ht or "").strip().lower().replace("ı", "i")
    return "hazir" in s and "ofis" in s


def _coerce_hazir_ofis_oda_no(data: dict, hizmet_turu: str, durum_pasif: bool):
    """None = yok; False = geçersiz aralık; int = oda no."""
    if durum_pasif:
        return None
    if not _hizmet_str_hazir_ofis_mi(hizmet_turu):
        return None
    raw = data.get("hazir_ofis_oda_no")
    if raw is None or str(raw).strip() == "":
        return None
    try:
        n = int(float(str(raw).strip()))
    except (TypeError, ValueError):
        return False
    if n < _HAZIR_OFIS_ODA_MIN or n > _HAZIR_OFIS_ODA_MAX:
        return False
    return n


# Web kökü (gemini_helper import için)
_web_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _web_root not in sys.path:
    sys.path.insert(0, _web_root)
try:
    from gemini_helper import analiz_yap as gemini_analiz_yap, GEMINI_AVAILABLE
except ImportError:
    GEMINI_AVAILABLE = False
    def gemini_analiz_yap(*args, **kwargs):
        return False, "Gemini modülü yüklenemedi."


def _vergi_no_normalize_veya_hata_kyc(tax_raw, yetkili_tc_raw):
    """10 hane VKN veya Yetkili T.C. ile birebir aynı 11 hane T.C. (Giriş / kaydet ile aynı kural)."""
    v = "".join(c for c in str(tax_raw or "") if c.isdigit())
    tc = "".join(c for c in str(yetkili_tc_raw or "") if c.isdigit())
    if not v:
        return "Vergi numarası zorunludur.", None
    if len(v) == 10:
        return None, v
    if len(v) == 11 and len(tc) == 11 and v == tc:
        return None, v
    if len(v) == 11:
        return "Vergi no 11 hane yalnızca Yetkili T.C. Kimlik No ile aynı olduğunda kabul edilir.", None
    return "Vergi no 10 haneli VKN veya Yetkili T.C. ile aynı 11 haneli T.C. olmalıdır.", None

# helper month names (Turkish)
MONTHS_TR = ["Ocak","Şubat","Mart","Nisan","Mayıs","Haziran","Temmuz","Ağustos","Eylül","Ekim","Kasım","Aralık"]


def _cur_fetch_one(conn, sql, params=()):
    cur = conn.cursor()
    cur.execute(sql, params)
    row = cur.fetchone()
    return dict(row) if row else None


def _cur_fetch_all(conn, sql, params=()):
    cur = conn.cursor()
    cur.execute(sql, params)
    return [dict(r) for r in cur.fetchall()]


bp = Blueprint("musteriler", __name__)


def _fintech_dashboard_data():
    """Müşteri Fintech Komuta Paneli için KPI, analitik, risk ve liste verileri. Tek bağlantı kullanır."""
    try:
        ensure_faturalar_amount_columns()
    except Exception:
        pass
    bugun = date.today()
    otuz_gun_once = bugun - timedelta(days=30)
    alti_ay_once = bugun - timedelta(days=180)
    yedi_gun = bugun - timedelta(days=7)

    def _row(r):
        return dict(r) if r is not None else {}

    try:
        with get_db() as conn:
            cur = conn.cursor()

            # KPI: Toplam Müşteri
            cur.execute("SELECT COUNT(*) as n FROM customers")
            toplam_musteri = _row(cur.fetchone()).get("n", 0) or 0

            # Aktif: en az bir ödenmemiş fatura veya sözleşmesi olan
            cur.execute("""
                SELECT COUNT(DISTINCT c.id) as n FROM customers c
                WHERE EXISTS (SELECT 1 FROM faturalar f WHERE f.musteri_id = c.id AND COALESCE(f.durum,'') != 'odendi')
                   OR EXISTS (SELECT 1 FROM musteri_kyc k WHERE k.musteri_id = c.id AND (
                     k.sozlesme_bitis IS NULL OR TRIM(COALESCE(k.sozlesme_bitis, '')) = ''
                     OR (NULLIF(TRIM(k.sozlesme_bitis), '')::date) >= %s
                   ))
            """, (bugun,))
            aktif_musteri = _row(cur.fetchone()).get("n", 0) or 0

            # Kritik: 30+ gün gecikmiş
            cur.execute("""
                SELECT COUNT(DISTINCT musteri_id) as n FROM faturalar
                WHERE COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL AND (vade_tarihi::date) <= %s
            """, (otuz_gun_once,))
            kritik_musteri = _row(cur.fetchone()).get("n", 0) or 0

            # Toplam aylık tahakkuk
            cur.execute("""
                SELECT COALESCE(SUM(COALESCE(toplam, tutar)), 0) as t FROM faturalar
                WHERE COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL
                  AND EXTRACT(YEAR FROM (vade_tarihi::date)) = %s AND EXTRACT(MONTH FROM (vade_tarihi::date)) = %s
            """, (bugun.year, bugun.month))
            toplam_aylik_tahakkuk = float(_row(cur.fetchone()).get("t", 0) or 0)

            # Toplam gecikme
            cur.execute("""
                SELECT COALESCE(SUM(COALESCE(toplam, tutar)), 0) as t FROM faturalar
                WHERE COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL AND (vade_tarihi::date) < %s
            """, (bugun,))
            toplam_gecikme = float(_row(cur.fetchone()).get("t", 0) or 0)

            # Tahsilat oranı: tahakkuk ve tahsilat son 6 ay
            cur.execute("""
                SELECT COALESCE(SUM(COALESCE(toplam, tutar)), 0) as t FROM faturalar
                WHERE vade_tarihi IS NOT NULL AND (vade_tarihi::date) >= %s
            """, (alti_ay_once,))
            tahakkuk_6ay = float(_row(cur.fetchone()).get("t", 0) or 0)
            cur.execute("""
                SELECT COALESCE(SUM(tutar), 0) as t FROM tahsilatlar
                WHERE (tahsilat_tarihi::date) >= %s
            """, (alti_ay_once,))
            tahsilat_6ay = float(_row(cur.fetchone()).get("t", 0) or 0)
            tahsilat_orani = round((tahsilat_6ay / tahakkuk_6ay * 100), 0) if tahakkuk_6ay else 100

            # Ortalama kira
            cur.execute("SELECT COALESCE(AVG(aylik_kira), 0) as t FROM musteri_kyc WHERE aylik_kira IS NOT NULL AND aylik_kira > 0")
            ortalama_kira = round(float(_row(cur.fetchone()).get("t", 0) or 0), 0)

            # Tahsilat trendi (son 6 ay)
            tahsilat_trend = []
            for i in range(5, -1, -1):
                d = bugun - timedelta(days=30 * i)
                y, m = d.year, d.month
                cur.execute("""
                    SELECT COALESCE(SUM(tutar), 0) as t FROM tahsilatlar
                    WHERE EXTRACT(YEAR FROM (tahsilat_tarihi::date)) = %s AND EXTRACT(MONTH FROM (tahsilat_tarihi::date)) = %s
                """, (y, m))
                tutar = float(_row(cur.fetchone()).get("t", 0) or 0)
                tahsilat_trend.append({"ay": MONTHS_TR[m - 1], "yil": y, "tutar": tutar})

            # Gecikme dağılımı: 0-7, 7-30, 30+ gün
            cur.execute("""
                SELECT COALESCE(SUM(COALESCE(toplam, tutar)), 0) as t FROM faturalar
                WHERE COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL
                  AND (vade_tarihi::date) < %s AND (vade_tarihi::date) > %s
            """, (bugun, yedi_gun))
            gecikme_0_7 = float(_row(cur.fetchone()).get("t", 0) or 0)
            cur.execute("""
                SELECT COALESCE(SUM(COALESCE(toplam, tutar)), 0) as t FROM faturalar
                WHERE COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL
                  AND (vade_tarihi::date) <= %s AND (vade_tarihi::date) > %s
            """, (yedi_gun, otuz_gun_once))
            gecikme_7_30 = float(_row(cur.fetchone()).get("t", 0) or 0)
            cur.execute("""
                SELECT COALESCE(SUM(COALESCE(toplam, tutar)), 0) as t FROM faturalar
                WHERE COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL AND (vade_tarihi::date) <= %s
            """, (otuz_gun_once,))
            gecikme_30_plus = float(_row(cur.fetchone()).get("t", 0) or 0)
            toplam_gecikme_dagilim = gecikme_0_7 + gecikme_7_30 + gecikme_30_plus
            gecikme_dagilimi = [
                {"label": "0-7 Gün", "tutar": gecikme_0_7, "yuzde": round(gecikme_0_7 / toplam_gecikme_dagilim * 100, 1) if toplam_gecikme_dagilim else 0},
                {"label": "7-30 Gün", "tutar": gecikme_7_30, "yuzde": round(gecikme_7_30 / toplam_gecikme_dagilim * 100, 1) if toplam_gecikme_dagilim else 0},
                {"label": "30+ Gün", "tutar": gecikme_30_plus, "yuzde": round(gecikme_30_plus / toplam_gecikme_dagilim * 100, 1) if toplam_gecikme_dagilim else 0},
            ]

            # En riskli 5 müşteri
            cur.execute("""
                SELECT f.musteri_id, c.name,
                       SUM(COALESCE(f.toplam, f.tutar)) as borc,
                       MIN(f.vade_tarihi) as en_eski_vade
                FROM faturalar f
                JOIN customers c ON c.id = f.musteri_id
                WHERE COALESCE(f.durum,'') != 'odendi' AND f.vade_tarihi IS NOT NULL AND (f.vade_tarihi::date) < %s
                GROUP BY f.musteri_id, c.name
                ORDER BY MIN(f.vade_tarihi)
                LIMIT 5
            """, (bugun,))
            risk_list = [dict(r) for r in cur.fetchall()]
            en_riskli_5 = []
            for row in risk_list:
                vd = row.get("en_eski_vade")
                gun = (bugun - vd).days if hasattr(vd, "year") else 0
                try:
                    if not hasattr(vd, "year") and vd:
                        vd = date(*[int(x) for x in str(vd)[:10].split("-")])
                        gun = (bugun - vd).days
                except Exception:
                    pass
                en_riskli_5.append({"name": row.get("name") or "—", "geciken_gun": gun, "musteri_id": row.get("musteri_id"), "borc": round(float(row.get("borc") or 0), 2)})

            genel_risk_puan = max(0, min(100, 100 - (kritik_musteri * 3) - int(toplam_gecikme / 10000))) if toplam_musteri else 100

            # Sözleşme 30 gün içinde bitecekler
            cur.execute("""
                SELECT k.musteri_id, c.name, c.office_code, k.sozlesme_bitis
                FROM musteri_kyc k
                JOIN customers c ON c.id = k.musteri_id
                WHERE NULLIF(TRIM(COALESCE(k.sozlesme_bitis, '')), '') IS NOT NULL
                  AND (NULLIF(TRIM(k.sozlesme_bitis), '')::date) >= %s
                  AND (NULLIF(TRIM(k.sozlesme_bitis), '')::date) <= %s
                ORDER BY (NULLIF(TRIM(k.sozlesme_bitis), '')::date)
                LIMIT 10
            """, (bugun, bugun + timedelta(days=30)))
            sozlesme_30 = [dict(r) for r in cur.fetchall()]
            sozlesme_30_list = []
            for row in sozlesme_30:
                bitis = row.get("sozlesme_bitis")
                kalan = (bitis - bugun).days if hasattr(bitis, "year") else 0
                if not hasattr(bitis, "year") and bitis:
                    try:
                        bitis = date(*[int(x) for x in str(bitis)[:10].split("-")])
                        kalan = (bitis - bugun).days
                    except Exception:
                        kalan = 0
                sozlesme_30_list.append({"name": row.get("name") or "—", "office_code": row.get("office_code") or "—", "kalan_gun": kalan, "musteri_id": row.get("musteri_id")})

            # Toplu Tahsilat: kritik müşteriler
            cur.execute("""
                SELECT f.musteri_id, c.name, c.phone, c.office_code,
                       SUM(COALESCE(f.toplam, f.tutar)) as toplam_alacak,
                       MIN(f.vade_tarihi) as en_eski_vade
                FROM faturalar f
                JOIN customers c ON c.id = f.musteri_id
                WHERE COALESCE(f.durum,'') != 'odendi' AND f.vade_tarihi IS NOT NULL AND (f.vade_tarihi::date) <= %s
                GROUP BY f.musteri_id, c.name, c.phone, c.office_code
                ORDER BY MIN(f.vade_tarihi)
                LIMIT 50
            """, (otuz_gun_once,))
            tahsilat_kritik = [dict(r) for r in cur.fetchall()]
            tahsilat_kritik_list = []
            for row in tahsilat_kritik:
                vd = row.get("en_eski_vade")
                gun = (bugun - vd).days if hasattr(vd, "year") else 0
                try:
                    if not hasattr(vd, "year") and vd:
                        vd = date(*[int(x) for x in str(vd)[:10].split("-")])
                        gun = (bugun - vd).days
                except Exception:
                    pass
                tahsilat_kritik_list.append({
                    "musteri_id": row.get("musteri_id"),
                    "name": row.get("name") or "—",
                    "phone": (row.get("phone") or "").replace(" ", "").replace("-", "").replace("(", "").replace(")", ""),
                    "office_code": row.get("office_code") or "—",
                    "geciken_gun": gun,
                    "toplam_alacak": round(float(row.get("toplam_alacak") or 0), 2),
                })

            # Müşteri listesi (drawer; tam liste /musteriler/list sayfasında)
            cur.execute("SELECT * FROM customers ORDER BY name LIMIT 5000")
            musteriler = [dict(r) for r in cur.fetchall()]

            # Müşteri bazlı borç/gecikme
            cur.execute("""
                SELECT musteri_id, SUM(COALESCE(toplam, tutar)) as toplam, MIN(vade_tarihi) as min_vade
                FROM faturalar WHERE COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL
                GROUP BY musteri_id
            """)
            fat_borc = [dict(r) for r in cur.fetchall()]
            musteri_borc = {}
            for row in fat_borc:
                mid = row["musteri_id"]
                musteri_borc[mid] = {"borc": float(row.get("toplam") or 0), "min_vade": row.get("min_vade")}
            for m in musteriler:
                mb = musteri_borc.get(m["id"], {})
                m["geciken_gun"] = 0
                if mb.get("min_vade"):
                    vd = mb["min_vade"]
                    if hasattr(vd, "year"):
                        m["geciken_gun"] = (bugun - vd).days
                    else:
                        try:
                            vd = date(*[int(x) for x in str(vd)[:10].split("-")])
                            m["geciken_gun"] = (bugun - vd).days
                        except Exception:
                            pass
                m["toplam_borc"] = round(mb.get("borc", 0), 2)

            toplam_bakiye = sum((m.get("toplam_borc") or 0) for m in musteriler)

            kargo_bugun = []
            kargo_teslim_bekleyen = []

            return {
                "kpi": {
                    "toplam_musteri": toplam_musteri,
                    "aktif_musteri": aktif_musteri,
                    "kritik_musteri": kritik_musteri,
                    "toplam_aylik_tahakkuk": round(toplam_aylik_tahakkuk, 2),
                    "toplam_gecikme": round(toplam_gecikme, 2),
                    "tahsilat_orani": tahsilat_orani,
                    "ortalama_kira": ortalama_kira,
                },
                "tahsilat_trend": tahsilat_trend,
                "gecikme_dagilimi": gecikme_dagilimi,
                "genel_risk_puan": genel_risk_puan,
                "en_riskli_5": en_riskli_5,
                "sozlesme_30_list": sozlesme_30_list,
                "musteriler": musteriler,
                "toplam_bakiye": round(toplam_bakiye, 2),
                "tahsilat_kritik_list": tahsilat_kritik_list,
                "kargo_bugun": kargo_bugun,
                "kargo_teslim_bekleyen": kargo_teslim_bekleyen,
                "gemini_available": GEMINI_AVAILABLE,
            }
    except Exception:
        return _fintech_defaults()


_TUM_YILLAR_ODENMIS_FROM = """
FROM customers c
WHERE NOT EXISTS (
    SELECT 1 FROM faturalar f
    WHERE f.musteri_id = c.id AND (f.durum IS NULL OR f.durum != 'odendi')
)
"""


def _musteri_liste_grup2_slugs_from_request():
    """URL ?grup2=a&grup2=b veya tek parametrede virgül; yalnızca grup2_etiketleri tablosunda kayıtlı slug'lar."""
    parts = request.args.getlist("grup2")
    if len(parts) == 1 and parts[0] and "," in parts[0]:
        parts = [p.strip() for p in parts[0].split(",") if p.strip()]
    slugs = [str(s).strip() for s in parts if str(s).strip()]
    if not slugs:
        return []
    ensure_grup2_etiketleri_table()
    rows = fetch_all(
        "SELECT slug FROM grup2_etiketleri WHERE COALESCE(aktif, TRUE) AND slug = ANY(%s)",
        (list(slugs),),
    )
    allowed = {r["slug"] for r in (rows or [])}
    return [s for s in slugs if s in allowed]


def _musteri_liste_grup2_sql_and_params(table_alias: str):
    """Müşteri kartındaki Grup 2: seçilen tüm etiketler caride işaretli olmalı (@>). Boş liste = filtre yok."""
    slugs = _musteri_liste_grup2_slugs_from_request()
    if not slugs:
        return "", ()
    a = f"{table_alias.strip()}." if table_alias and table_alias.strip() else ""
    # Legacy uyum: grup2_secimleri boş olsa bile bizim_hesap/vergi_dairesi alanlarından türet.
    sql = (
        " AND ("
        f"COALESCE({a}grup2_secimleri, ARRAY[]::text[])"
        f" || CASE WHEN COALESCE({a}bizim_hesap, FALSE) THEN ARRAY['bizim_hesap']::text[] ELSE ARRAY[]::text[] END"
        f" || CASE WHEN COALESCE(NULLIF(TRIM(COALESCE({a}vergi_dairesi, '')), ''), '') != ''"
        " THEN ARRAY['vergi_dairesi']::text[] ELSE ARRAY[]::text[] END"
        ") @> %s::text[]"
    )
    return sql, (list(slugs),)


def _musteri_liste_grup2_etiket_rows():
    ensure_grup2_etiketleri_table()
    return fetch_all(
        """
        SELECT slug, etiket FROM grup2_etiketleri
        WHERE COALESCE(aktif, TRUE)
        ORDER BY
            CASE slug
                WHEN 'bizim_hesap' THEN 0
                WHEN 'vergi_dairesi' THEN 1
                WHEN 'vergi_dairesi_terk' THEN 2
                ELSE 3
            END,
            sira NULLS LAST,
            etiket
        """
    )


def _musteri_liste_attach_grup2_display(musteriler, etiket_map):
    """Satırda gösterilecek etiket metni (müşteri kartı Grup 2 ile aynı sıra anlamı)."""
    from routes.giris_routes import _parse_pg_text_array_grup2

    rows = musteriler or []
    ids = [m.get("id") for m in rows if m.get("id") is not None]
    if ids:
        try:
            g2_rows = fetch_all(
                "SELECT id, grup2_secimleri, bizim_hesap FROM customers WHERE id = ANY(%s)",
                (list(ids),),
            )
            by_id = {int(r["id"]): r for r in (g2_rows or []) if r.get("id") is not None}
            for m in rows:
                mid = m.get("id")
                if mid is None:
                    continue
                r = by_id.get(int(mid))
                if not r:
                    continue
                cur_g2 = _parse_pg_text_array_grup2(m.get("grup2_secimleri"))
                db_g2 = _parse_pg_text_array_grup2(r.get("grup2_secimleri"))
                merged_g2 = list(dict.fromkeys(cur_g2 + db_g2))
                if merged_g2:
                    m["grup2_secimleri"] = merged_g2
                elif r.get("grup2_secimleri") is not None:
                    m["grup2_secimleri"] = r.get("grup2_secimleri")
                m["bizim_hesap"] = bool(m.get("bizim_hesap")) or bool(r.get("bizim_hesap"))
        except Exception:
            pass

    for m in rows:
        g2_list = _parse_pg_text_array_grup2(m.get("grup2_secimleri"))
        if bool(m.get("bizim_hesap")) and "bizim_hesap" not in g2_list:
            g2_list = ["bizim_hesap"] + g2_list
        if (m.get("vergi_dairesi") or "").strip() and "vergi_dairesi" not in g2_list:
            g2_list.append("vergi_dairesi")
        if not g2_list and bool(m.get("bizim_hesap")):
            g2_list = ["bizim_hesap"]
        g2_list = list(dict.fromkeys(g2_list))
        parts = [etiket_map.get(str(x), str(x).replace("_", " ").title()) for x in g2_list if x]
        m["grup2_display"] = ", ".join(parts) if parts else ""


def _musteri_liste_data():
    """Müşteri listesi sayfası için ortak veri (ana sayfa ve /list).

    Sayfalama (isteğe bağlı): ?sayfa=1&limit=50 — yalnızca bu parametrelerden biri
    URL'de varsa LIMIT/OFFSET uygulanır; aksi halde eski davranış (tüm kayıtlar).
    """
    arama = request.args.get("q", "").strip()
    tum_yillar_odenmis = request.args.get("tum_yillar_odenmis") == "1"
    paginate = ("sayfa" in request.args) or ("limit" in request.args)
    sayfa = max(1, request.args.get("sayfa", default=1, type=int) or 1)
    limit = request.args.get("limit", type=int)
    if limit is None:
        limit = 50
    limit = max(1, min(500, int(limit)))

    ensure_grup2_etiketleri_table()
    ensure_customers_grup2_secimleri()
    ensure_grup2_bizim_hesap_into_array()
    grup2_filter_slugs = _musteri_liste_grup2_slugs_from_request()
    grup2_etiketler = _musteri_liste_grup2_etiket_rows() or []
    grup2_etiket_map = {r["slug"]: r["etiket"] for r in grup2_etiketler}

    if tum_yillar_odenmis:
        g2_sql, g2_params = _musteri_liste_grup2_sql_and_params("c")
        sql = f"SELECT c.* {_TUM_YILLAR_ODENMIS_FROM}{g2_sql} ORDER BY c.name"
        count_sql = f"SELECT COUNT(*) AS n {_TUM_YILLAR_ODENMIS_FROM}{g2_sql}"
        params = tuple(g2_params)
    elif arama:
        g2_sql, g2_params = _musteri_liste_grup2_sql_and_params("")
        w3 = customers_arama_sql_giris_genis("")
        sql = f"SELECT * FROM customers WHERE {w3}{g2_sql} ORDER BY name"
        count_sql = f"SELECT COUNT(*) AS n FROM customers WHERE {w3}{g2_sql}"
        params = tuple(customers_arama_params_giris_genis(arama)) + tuple(g2_params)
    else:
        g2_sql, g2_params = _musteri_liste_grup2_sql_and_params("")
        if g2_sql:
            sql = f"SELECT * FROM customers WHERE 1=1{g2_sql} ORDER BY name"
            count_sql = f"SELECT COUNT(*) AS n FROM customers WHERE 1=1{g2_sql}"
            params = tuple(g2_params)
        else:
            sql = "SELECT * FROM customers ORDER BY name"
            count_sql = "SELECT COUNT(*) AS n FROM customers"
            params = ()

    toplam_musteri = None
    toplam_sayfa = None
    if paginate:
        cnt = fetch_one(count_sql, params) or {}
        try:
            toplam_musteri = int(cnt.get("n") or 0)
        except (TypeError, ValueError):
            toplam_musteri = 0
        toplam_sayfa = max(1, (toplam_musteri + limit - 1) // limit) if toplam_musteri else 1
        offset = (sayfa - 1) * limit
        musteriler = fetch_all(sql + " LIMIT %s OFFSET %s", (*params, limit, offset))
    else:
        musteriler = fetch_all(sql, params)

    return {
        "musteriler": musteriler,
        "arama": arama,
        "tum_yillar_odenmis": tum_yillar_odenmis,
        "paginate": paginate,
        "sayfa": sayfa if paginate else 1,
        "limit": limit if paginate else None,
        "toplam_musteri": toplam_musteri if paginate else None,
        "toplam_sayfa": toplam_sayfa if paginate else None,
        "grup2_filter_slugs": grup2_filter_slugs,
        "grup2_etiketler": grup2_etiketler,
        "grup2_etiket_map": grup2_etiket_map,
    }


def _enrich_musteri_list_with_borc_gecikme(musteriler, *, skip_dedupe: bool = False):
    """Müşteri listesine toplam_borc, geciken_gun, son_odeme_tarihi, rent_start_date (kyc'den) ekler."""
    if not musteriler:
        return
    bugun = date.today()
    ids = [m["id"] for m in musteriler]

    with get_db() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT
              c.id,
              COALESCE(f.toplam_borc, 0) as toplam_borc,
              f.min_vade,
              t.son_tahsilat,
              k.sozlesme_tarihi,
              k.aylik_kira
            FROM customers c
            LEFT JOIN (
              SELECT musteri_id,
                SUM(COALESCE(toplam,tutar)) as toplam_borc,
                MIN(vade_tarihi) as min_vade
              FROM faturalar
              WHERE COALESCE(durum,'') != 'odendi'
              GROUP BY musteri_id
            ) f ON f.musteri_id = c.id
            LEFT JOIN (
              SELECT musteri_id, MAX(tahsilat_tarihi) as son_tahsilat
              FROM tahsilatlar GROUP BY musteri_id
            ) t ON t.musteri_id = c.id
            LEFT JOIN LATERAL (
              SELECT sozlesme_tarihi, aylik_kira
              FROM musteri_kyc
              WHERE musteri_id = c.id
              ORDER BY id DESC LIMIT 1
            ) k ON true
            WHERE c.id = ANY(%s)
            """,
            (ids,),
        )
        rows = cur.fetchall()
        fat_borc = {
            r["id"]: {"borc": float(r.get("toplam_borc") or 0), "min_vade": r.get("min_vade")}
            for r in rows
        }
        son_tahsilat = {r["id"]: r.get("son_tahsilat") for r in rows}
        kyc_tarih = {r["id"]: r.get("sozlesme_tarihi") for r in rows}
    for m in musteriler:
        if not m.get("rent_start_date") and kyc_tarih.get(m["id"]):
            m["rent_start_date"] = kyc_tarih[m["id"]]
        m["guncel_kira_bedeli"] = float(m.get("guncel_kira_bedeli") or m.get("current_rent") or 0)
        mb = fat_borc.get(m["id"], {})
        if "manuel_borc" in m and m["manuel_borc"] is not None:
            m["toplam_borc"] = round(float(m["manuel_borc"] or 0), 2)
        else:
            m["toplam_borc"] = round(mb.get("borc", 0), 2)
        m["geciken_gun"] = 0
        if mb.get("min_vade"):
            vd = mb["min_vade"]
            try:
                if hasattr(vd, "year"):
                    m["geciken_gun"] = (bugun - vd).days
                else:
                    vd = date(*[int(x) for x in str(vd)[:10].split("-")])
                    m["geciken_gun"] = (bugun - vd).days
            except Exception:
                pass
        m["son_odeme_tarihi"] = m.get("son_odeme_tarihi") or son_tahsilat.get(m["id"])
    _normalize_musteri_durum_for_liste(musteriler)
    if not skip_dedupe:
        _dedupe_musteri_liste_by_identity(musteriler)


def _normalize_musteri_durum_for_liste(musteriler):
    """Tek kaynak: customers.durum ile kapanış tarihi tutarlı olsun.
    Kapanış tarihi doluysa müşteri pasiftir (eski kayıtta durum=aktif kalmış olsa bile).
    Böylece Aktif / Pasif listeleri birbirini dışlar."""
    if not musteriler:
        return
    for m in musteriler:
        raw = (m.get("durum") or "").strip().lower()
        kap = m.get("kapanis_tarihi")
        has_kap = False
        if kap is not None:
            if hasattr(kap, "year"):
                has_kap = True
            else:
                s = str(kap).strip()
                if len(s) >= 10 and s[4] == "-":
                    has_kap = True
                elif s and s not in ("None", "null", ""):
                    has_kap = True
        if raw == "pasif":
            m["durum"] = "pasif"
            continue
        if has_kap:
            m["durum"] = "pasif"
            continue
        if raw == "aktif":
            m["durum"] = "aktif"
            continue
        m["durum"] = "aktif"


def _dedupe_musteri_liste_by_identity(musteriler):
    """Aynı vergi numarası veya aynı cep telefonu ile birden fazla müşteri satırı varsa (çift kart)
    listede tek kayıt gösterilir. Tercih: pasif + kapanış tarihi, yoksa daha yüksek id (genelde güncel)."""
    if not musteriler:
        return
    from collections import defaultdict

    def bucket_key(m):
        tn = "".join(c for c in str(m.get("tax_number") or "") if c.isdigit())
        if len(tn) >= 10:
            return ("tax", tn)
        ph = "".join(c for c in str(m.get("phone") or "") if c.isdigit())
        if len(ph) >= 10:
            return ("ph", ph[-10:])
        return None

    def has_kapanis(m):
        kap = m.get("kapanis_tarihi")
        if kap is None:
            return False
        if hasattr(kap, "year"):
            return True
        s = str(kap).strip()
        return bool(s) and s not in ("None", "null", "")

    def _g2_parse(val):
        if val is None:
            return []
        if isinstance(val, (list, tuple)):
            return [str(x).strip() for x in val if str(x).strip()]
        s = str(val).strip()
        if not s or s in ("{}", "[]", "None", "null"):
            return []
        if s.startswith("{") and s.endswith("}"):
            inner = s[1:-1].strip()
            return [p.strip().strip('"') for p in inner.split(",") if p.strip()] if inner else []
        return [s]

    def score_row(m):
        s = int(m.get("id") or 0)
        if (m.get("durum") or "").strip().lower() == "pasif":
            s += 1_000_000
        if has_kapanis(m):
            s += 500_000
        return s

    buckets = defaultdict(list)
    no_key = []
    for m in musteriler:
        k = bucket_key(m)
        if k is None:
            no_key.append(m)
        else:
            buckets[k].append(m)

    out = list(no_key)
    for _key, group in buckets.items():
        if len(group) == 1:
            out.append(group[0])
            continue
        best = max(group, key=score_row)
        for o in group:
            if o["id"] == best["id"]:
                continue
            best_g2 = _g2_parse(best.get("grup2_secimleri"))
            other_g2 = _g2_parse(o.get("grup2_secimleri"))
            if bool(o.get("bizim_hesap")) and "bizim_hesap" not in other_g2:
                other_g2 = ["bizim_hesap"] + other_g2
            merged = list(dict.fromkeys(best_g2 + other_g2))
            if merged:
                best["grup2_secimleri"] = merged
                best["bizim_hesap"] = "bizim_hesap" in merged or bool(best.get("bizim_hesap")) or bool(
                    o.get("bizim_hesap")
                )
            try:
                best["toplam_borc"] = max(
                    float(best.get("toplam_borc") or 0),
                    float(o.get("toplam_borc") or 0),
                )
            except (TypeError, ValueError):
                pass
            best["geciken_gun"] = max(
                int(best.get("geciken_gun") or 0),
                int(o.get("geciken_gun") or 0),
            )
            if o.get("manuel_borc") is not None and best.get("manuel_borc") is not None:
                try:
                    best["manuel_borc"] = max(
                        float(best.get("manuel_borc") or 0),
                        float(o.get("manuel_borc") or 0),
                    )
                except (TypeError, ValueError):
                    pass
        try:
            best["toplam_borc"] = round(float(best.get("toplam_borc") or 0), 2)
        except (TypeError, ValueError):
            pass
        out.append(best)

    out.sort(key=lambda m: ((m.get("name") or "").lower(), m.get("id") or 0))
    musteriler[:] = out


@bp.route("/list")
@giris_gerekli
def list_full():
    embed = str(request.args.get("embed") or "").lower() in ("1", "true", "yes", "on")
    try:
        ensure_customers_quick_edit_columns()
        ensure_customers_kapanis_tarihi()
        ensure_grup2_etiketleri_table()
        ensure_customers_grup2_secimleri()
    except Exception:
        pass
    data = _musteri_liste_data()
    musteriler = data["musteriler"]
    if musteriler:
        _enrich_musteri_list_with_borc_gecikme(
            musteriler,
            skip_dedupe=bool(data.get("paginate")),
        )
    _musteri_liste_attach_grup2_display(musteriler, data.get("grup2_etiket_map") or {})
    paginate = bool(data.get("paginate"))
    toplam_musteri = data.get("toplam_musteri") if paginate else len(musteriler)
    toplam_sayfa = data.get("toplam_sayfa") if paginate else 1
    return render_template(
        "musteriler/index.html",
        musteriler=musteriler,
        arama=data["arama"],
        tum_yillar_odenmis=data["tum_yillar_odenmis"],
        paginate=paginate,
        sayfa=data.get("sayfa", 1),
        limit=data.get("limit") or 50,
        toplam_musteri=toplam_musteri,
        toplam_sayfa=toplam_sayfa or 1,
        embed=embed,
        grup2_filter_slugs=data.get("grup2_filter_slugs") or [],
        grup2_etiketler=data.get("grup2_etiketler") or [],
    )


@bp.route("/")
@giris_gerekli
def index():
    """Müşteri Fintech Komuta Paneli — ana sayfa. Hata olursa Dashboard'a yönlendirir, 500 dönmez."""
    import traceback
    try:
        try:
            data = _fintech_dashboard_data()
        except Exception as e:
            traceback.print_exc()
            print(f"Fintech dashboard error (index): {e}")
            data = _fintech_defaults()
        import_sonuc = request.args.get("import_sonuc")
        imported = request.args.get("imported", type=int)
        import_hatalar = request.args.get("import_hatalar", type=int) or 0
        import_percent = request.args.get("import_percent", type=int)
        import_total = request.args.get("import_total", type=int)
        bugun = date.today()
        MONTHS_TR = ["Ocak","Şubat","Mart","Nisan","Mayıs","Haziran","Temmuz","Ağustos","Eylül","Ekim","Kasım","Aralık"]
        try:
            return render_template(
                "musteriler/fintech.html",
                **data,
                import_sonuc=import_sonuc,
                imported=imported or 0,
                import_hatalar=import_hatalar,
                import_percent=import_percent,
                import_total=import_total,
                bugun=bugun,
                now_year=bugun.year,
                now_month=bugun.month,
                MONTHS_TR=MONTHS_TR,
            )
        except Exception as e:
            traceback.print_exc()
            print(f"Fintech template render error (index): {e}")
            data = _fintech_defaults()
            try:
                return render_template(
                    "musteriler/fintech.html",
                    **data,
                    import_sonuc=import_sonuc,
                    imported=imported or 0,
                    import_hatalar=import_hatalar,
                    bugun=bugun,
                    now_year=bugun.year,
                    now_month=bugun.month,
                    MONTHS_TR=MONTHS_TR,
                )
            except Exception:
                flash("Müşteriler sayfası geçici olarak yüklenemedi. Lütfen tekrar deneyin.", "warning")
                return redirect(url_for("dashboard.index"))
    except Exception as e:
        traceback.print_exc()
        print(f"Musteriler index unexpected error: {e}")
        flash("Müşteriler sayfası açılamadı. Lütfen tekrar deneyin.", "warning")
        return redirect(url_for("dashboard.index"))


@bp.route("/ozet")
@giris_gerekli
def ozet():
    """Müşteri özet / Fintech komuta paneli — KPI kartları + sağ panel drawer."""
    import traceback
    try:
        data = _fintech_dashboard_data()
    except Exception as e:
        traceback.print_exc()
        print(f"Fintech dashboard error (ozet): {e}")
        data = _fintech_defaults()
    import_sonuc = request.args.get("import_sonuc")
    imported = request.args.get("imported", type=int)
    import_hatalar = request.args.get("import_hatalar", type=int) or 0
    import_percent = request.args.get("import_percent", type=int)
    import_total = request.args.get("import_total", type=int)
    bugun = date.today()
    return render_template(
        "musteriler/fintech.html",
        **data,
        import_sonuc=import_sonuc,
        imported=imported or 0,
        import_hatalar=import_hatalar,
        import_percent=import_percent,
        import_total=import_total,
        now_year=bugun.year,
        now_month=bugun.month,
        MONTHS_TR=MONTHS_TR,
    )


def _fintech_defaults():
    """Fintech sayfası için varsayılan veri (hata durumunda şablon kırılmasın)."""
    return {
        "kpi": {"toplam_musteri": 0, "aktif_musteri": 0, "kritik_musteri": 0, "toplam_aylik_tahakkuk": 0, "toplam_gecikme": 0, "tahsilat_orani": 100, "ortalama_kira": 0},
        "tahsilat_trend": [],
        "gecikme_dagilimi": [{"label": "0-7 Gün", "tutar": 0}, {"label": "7-30 Gün", "tutar": 0}, {"label": "30+ Gün", "tutar": 0}],
        "genel_risk_puan": 100,
        "en_riskli_5": [],
        "sozlesme_30_list": [],
        "musteriler": [],
        "toplam_bakiye": 0,
        "tahsilat_kritik_list": [],
        "kargo_bugun": [],
        "kargo_teslim_bekleyen": [],
        "fatura_bekleyen_count": 0,
        "gemini_available": GEMINI_AVAILABLE,
    }


def _komuta_merkezi_data(mid):
    """Müşteri komuta merkezi için tüm veriyi topla: hero, risk, 12 ay, kargo, TÜFE, WhatsApp metni."""
    bugun = date.today()
    bu_yil = bugun.year
    alti_ay_once = bugun - timedelta(days=180)

    with get_db() as conn:
        musteri = _cur_fetch_one(conn, "SELECT * FROM customers WHERE id=%s", (mid,))
        if not musteri:
            return None
        kyc = _cur_fetch_one(
            conn,
            "SELECT sozlesme_tarihi, sozlesme_bitis, hizmet_turu, aylik_kira FROM musteri_kyc WHERE musteri_id=%s ORDER BY id DESC LIMIT 1",
            (mid,),
        )
        # Tahsilat: toplam ödenen, son ödeme tarihi (musteri_id veya customer_id)
        tah = _cur_fetch_one(
            conn,
            """SELECT COALESCE(SUM(tutar), 0) as toplam_odenen,
                      MAX(tahsilat_tarihi) as son_odeme_tarihi
               FROM tahsilatlar WHERE musteri_id = %s OR customer_id = %s""",
            (mid, mid),
        )
        # Faturalar: toplam borç, gecikme, bu ayki borç (bu ay vadesi gelen ödenmemiş)
        faturalar_odenmemis = _cur_fetch_all(
            conn,
            """
            SELECT id, COALESCE(toplam, tutar) as toplam, vade_tarihi, fatura_no
            FROM faturalar WHERE musteri_id=%s AND COALESCE(durum,'') != 'odendi'
            """,
            (mid,),
        )
        # Risk: son 6 ay gecikme sayısı, ortalama gecikme, puan
        faturalar_son6 = _cur_fetch_all(
            conn,
            """
            SELECT vade_tarihi, durum FROM faturalar
            WHERE musteri_id=%s AND vade_tarihi IS NOT NULL AND (vade_tarihi::date) >= %s
            ORDER BY vade_tarihi
            """,
            (mid, alti_ay_once),
        )
        # 12 aylık grid: sadece cari yıl (tüm geçmişi çekmeyi önler)
        tum_faturalar = _cur_fetch_all(
            conn,
            """
            SELECT id, COALESCE(toplam, tutar) as toplam, vade_tarihi, durum
            FROM faturalar
            WHERE musteri_id=%s AND vade_tarihi IS NOT NULL
              AND EXTRACT(YEAR FROM vade_tarihi::date) = %s
            ORDER BY vade_tarihi
            """,
            (mid, bu_yil),
        )
        # Fatura başına son tahsilat tarihi (tek sorgu, DB tarafında DISTINCT ON)
        tahsilatlar = _cur_fetch_all(
            conn,
            """
            SELECT DISTINCT ON (fatura_id) fatura_id, tahsilat_tarihi
            FROM tahsilatlar
            WHERE (musteri_id = %s OR customer_id = %s) AND fatura_id IS NOT NULL
            ORDER BY fatura_id, tahsilat_tarihi DESC NULLS LAST
            """,
            (mid, mid),
        )

        # Son 6 ay tahsilat toplamı (grafik için) — aynı bağlantıda
        son_6_ay_chart = []
        for i in range(5, -1, -1):
            d = bugun - timedelta(days=30 * i)
            ay_bas = d.replace(day=1)
            if i == 0:
                ay_son = bugun
            else:
                nxt = ay_bas.month % 12 + 1
                y = ay_bas.year + (1 if nxt == 1 else 0)
                ay_son = date(y, nxt, 1) - timedelta(days=1)
            row = _cur_fetch_one(
                conn,
                """SELECT COALESCE(SUM(tutar), 0) as t FROM tahsilatlar
                   WHERE (musteri_id = %s OR customer_id = %s) AND (tahsilat_tarihi::date) >= %s AND (tahsilat_tarihi::date) <= %s""",
                (mid, mid, ay_bas, ay_son),
            )
            son_6_ay_chart.append({"ay": MONTHS_TR[ay_bas.month - 1][:3], "tutar": float(row.get("t") or 0)})

        # Kargolar timeline
        kargolar_raw = _cur_fetch_all(
            conn,
            "SELECT id, tarih, takip_no, kargo_firmasi, teslim_alan, notlar, created_at FROM kargolar WHERE musteri_id=%s ORDER BY created_at DESC LIMIT 30",
            (mid,),
        )

        tufe_oran = 0
        artis_sonrasi_kira = None
        aylik_kira = float(kyc.get("aylik_kira") or 0) if kyc else 0
        sozlesme_bitis = kyc.get("sozlesme_bitis") if kyc else None
        if sozlesme_bitis and aylik_kira > 0:
            yil = sozlesme_bitis.year if hasattr(sozlesme_bitis, "year") else int(str(sozlesme_bitis)[:4])
            r = _cur_fetch_one(
                conn,
                "SELECT oran FROM tufe_verileri WHERE year = %s ORDER BY month DESC LIMIT 1",
                (yil,),
            )
            tufe_oran = float(r.get("oran") or 0) if r else 0
            artis_sonrasi_kira = round(aylik_kira * (1 + tufe_oran / 100), 2)

    sozlesme_baslangic = kyc.get("sozlesme_tarihi") if kyc else None
    sozlesme_bitis = kyc.get("sozlesme_bitis") if kyc else None
    sozlesme_bas_str = _tarih_str(sozlesme_baslangic)
    sozlesme_bitis_str = _tarih_str(sozlesme_bitis)
    kalan_gun = None
    if sozlesme_bitis:
        d = sozlesme_bitis if hasattr(sozlesme_bitis, "year") else _parse_date_str(str(sozlesme_bitis)[:10])
        if d:
            kalan_gun = (d - bugun).days

    toplam_odenen = float(tah.get("toplam_odenen") or 0)
    son_odeme_tarihi = tah.get("son_odeme_tarihi")
    son_odeme_str = _tarih_str(son_odeme_tarihi) or "—"

    toplam_borc = sum(float(f.get("toplam") or 0) for f in (faturalar_odenmemis or []))
    geciken_gun = 0
    bu_ay = bugun.month
    bu_ayki_borc = 0
    for f in (faturalar_odenmemis or []):
        vd = _parse_vade(f.get("vade_tarihi"))
        if vd and vd < bugun:
            geciken_gun = max(geciken_gun, (bugun - vd).days)
        if vd and vd.month == bu_ay and vd.year == bu_yil:
            bu_ayki_borc += float(f.get("toplam") or 0)

    gecikme_sayisi_6ay = 0
    gecikme_gunleri = []
    for f in (faturalar_son6 or []):
        vd = _parse_vade(f.get("vade_tarihi"))
        if not vd:
            continue
        if (f.get("durum") or "").lower() != "odendi":
            if vd < bugun:
                gecikme_sayisi_6ay += 1
                gecikme_gunleri.append((bugun - vd).days)
    ortalama_gecikme = int(sum(gecikme_gunleri) / len(gecikme_gunleri)) if gecikme_gunleri else 0
    # Risk puanı 0-100: yüksek = iyi (az risk)
    risk_puan = max(0, min(100, 100 - gecikme_sayisi_6ay * 15 - ortalama_gecikme * 2))
    if risk_puan >= 70:
        risk_durum = "stabil"
        risk_aksiyon = "Şu an için bir aksiyona gerek yok."
    elif risk_puan >= 40:
        risk_durum = "dikkat"
        risk_aksiyon = "Hatırlatma mesajı gönderin."
    else:
        risk_durum = "kritik"
        risk_aksiyon = "Hukuki uyarı ve takip önerilir."

    fatura_odeme = {t["fatura_id"]: t["tahsilat_tarihi"] for t in (tahsilatlar or []) if t.get("fatura_id")}
    aylik_detay = []
    cari_yil = bugun.year
    for ay in range(1, 13):
        ay_faturalar = [f for f in (tum_faturalar or []) if _parse_vade(f.get("vade_tarihi")) and _parse_vade(f.get("vade_tarihi")).month == ay and _parse_vade(f.get("vade_tarihi")).year == cari_yil]
        if not ay_faturalar:
            aylik_detay.append({"ay_adi": MONTHS_TR[ay - 1], "tutar": 0, "odeme_tarihi": None, "gecikme_gun": None, "durum": "gelecek", "fatura_id": None})
            continue
        f = ay_faturalar[0]
        vd = _parse_vade(f.get("vade_tarihi"))
        tutar = float(f.get("toplam") or 0)
        odendi = (f.get("durum") or "").lower() == "odendi"
        odeme_tarihi = fatura_odeme.get(f.get("id")) if odendi else None
        gecikme_gun = (bugun - vd).days if vd and not odendi and vd < bugun else None
        if odendi:
            durum = "odendi"
        elif vd and vd < bugun:
            durum = "gecikti"
        elif vd and (vd == bugun or (vd - bugun).days <= 7):
            durum = "bugun_yakin"
        else:
            durum = "gelecek"
        aylik_detay.append({
            "ay_adi": MONTHS_TR[ay - 1],
            "tutar": tutar,
            "odeme_tarihi": odeme_tarihi,
            "odeme_tarihi_str": _tarih_str(odeme_tarihi) if odeme_tarihi else "—",
            "gecikme_gun": gecikme_gun,
            "durum": durum,
            "fatura_id": f.get("id"),
        })

    kargolar = []
    for k in (kargolar_raw or []):
        teslim = (k.get("teslim_alan") or "").strip()
        durum = "Teslim alındı" if teslim else "Bekliyor"
        created = k.get("created_at")
        saat_str = created.strftime("%H:%M") if hasattr(created, "strftime") else (str(created)[11:16] if created else "—")
        kargolar.append({
            "id": k.get("id"),
            "tarih": _tarih_str(k.get("tarih")) or "—",
            "barkod_takip": k.get("takip_no") or "—",
            "notlar": k.get("notlar") or "—",
            "teslim_durum": durum,
            "saat": saat_str,
            "kargo_firmasi": k.get("kargo_firmasi") or "—",
        })

    # WhatsApp metni (gecikme gününe göre)
    if geciken_gun <= 0:
        whatsapp_metin = f"Sayın {musteri.get('name') or 'Müşteri'}, {musteri.get('office_code') or ''} numaralı ofisinizle ilgili herhangi bir kira gecikmesi söz konusu değildir. Teşekkür ederiz."
    elif geciken_gun <= 7:
        whatsapp_metin = f"Merhaba, unutkanlık olmuş olabilir; ödeme hatırlatması yapıyoruz. Lütfen kalan tutarı ({toplam_borc:,.2f} ₺) zamanında ödeyiniz."
    elif geciken_gun <= 30:
        whatsapp_metin = f"Merhaba, hizmet devamı için ödemenizin yapılması gerekmektedir. Lütfen kalan tutarı {toplam_borc:,.2f} ₺ ödeyiniz."
    else:
        whatsapp_metin = f"Ödeme yapılmadığı takdirde hukuki işlem başlatılacaktır. Lütfen derhal {toplam_borc:,.2f} ₺ tutarındaki bakiyeyi ödeyiniz."

    # Ödeme durumu rozeti
    if geciken_gun > 30:
        odeme_rozet = "kritik"
    elif geciken_gun > 0 or toplam_borc > 0:
        odeme_rozet = "dikkat"
    else:
        odeme_rozet = "stabil"

    return {
        "musteri": musteri,
        "kyc": kyc,
        "sozlesme_bas_str": sozlesme_bas_str or "—",
        "sozlesme_bitis_str": sozlesme_bitis_str or "—",
        "kalan_gun": kalan_gun,
        "toplam_odenen": toplam_odenen,
        "toplam_borc": toplam_borc,
        "bu_ayki_borc": bu_ayki_borc,
        "son_odeme_str": son_odeme_str,
        "geciken_gun": geciken_gun,
        "risk_gecikme_sayisi_6ay": gecikme_sayisi_6ay,
        "risk_ortalama_gecikme": ortalama_gecikme,
        "risk_puan": risk_puan,
        "risk_durum": risk_durum,
        "risk_aksiyon": risk_aksiyon,
        "aylik_detay": aylik_detay,
        "son_6_ay_chart": son_6_ay_chart,
        "kargolar": kargolar,
        "tufe_oran": tufe_oran,
        "artis_sonrasi_kira": artis_sonrasi_kira,
        "aylik_kira": aylik_kira,
        "whatsapp_metin": whatsapp_metin,
        "odeme_rozet": odeme_rozet,
        "hizmet_turu": (kyc.get("hizmet_turu") or "—") if kyc else "—",
    }


def _tarih_str(d):
    """date veya string -> dd.mm.yyyy"""
    if not d:
        return None
    if hasattr(d, "strftime"):
        return d.strftime("%d.%m.%Y")
    s = str(d)[:10]
    return s[8:10] + "." + s[5:7] + "." + s[0:4] if len(s) >= 10 else s


def _parse_date_str(s):
    try:
        parts = str(s)[:10].split("-")
        if len(parts) == 3:
            return date(int(parts[0]), int(parts[1]), int(parts[2]))
    except Exception:
        pass
    return None


@bp.route("/<int:mid>")
@giris_gerekli
def detay(mid):
    """Müşteri detay sayfası - Komuta Merkezi (tam ekran)."""
    data = _komuta_merkezi_data(mid)
    if not data:
        flash("Müşteri bulunamadı.", "danger")
        return redirect(url_for("musteriler.index"))
    return render_template("musteriler/detay.html", **data)


def _parse_vade(v):
    """vade_tarihi (date veya text) -> date veya None."""
    if not v:
        return None
    if hasattr(v, "year"):
        return v
    try:
        s = str(v)[:10]
        return date(*[int(x) for x in s.split("-")])
    except Exception:
        return None


@bp.route("/api/ara")
@giris_gerekli
def api_ara():
    """İsme göre müşteri ara (autocomplete). ?q= ile sorgu."""
    q = (request.args.get("q") or "").strip()
    if not q or len(q) < 1:
        return jsonify([])
    w3 = customers_arama_sql_giris_genis("")
    rows = fetch_all(
        f"SELECT id, name, musteri_adi FROM customers WHERE {w3} ORDER BY name LIMIT 25",
        customers_arama_params_giris_genis(q),
    )
    return jsonify(rows or [])


@bp.route("/<int:mid>/api/popup")
@giris_gerekli
def api_popup(mid):
    """Tek müşteri için detay popup verisi: ödeme durumu, 12 ay grid, kargo geçmişi, sözleşme bitiş."""
    musteri = fetch_one("SELECT * FROM customers WHERE id=%s", (mid,))
    if not musteri:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    bugun = date.today()

    # KYC / sözleşme bitiş
    kyc = fetch_one(
        "SELECT sozlesme_bitis, hizmet_turu FROM musteri_kyc WHERE musteri_id=%s ORDER BY id DESC LIMIT 1",
        (mid,),
    )
    sozlesme_bitis = kyc.get("sozlesme_bitis") if kyc else None
    sozlesme_str = ""
    sozlesme_yakin = False  # 1 ay kala kırmızı
    if sozlesme_bitis:
        if hasattr(sozlesme_bitis, "strftime"):
            sozlesme_str = sozlesme_bitis.strftime("%d.%m.%Y")
            d = sozlesme_bitis
        else:
            sozlesme_str = str(sozlesme_bitis)[:10]
            try:
                d = date(*[int(x) for x in str(sozlesme_bitis)[:10].split("-")])
            except Exception:
                d = None
        if d and bugun <= d <= bugun + timedelta(days=31):
            sozlesme_yakin = True

    # Ödenmemiş faturalar
    faturalar_odenmemis = fetch_all("""
        SELECT id, COALESCE(toplam, tutar) as toplam, vade_tarihi, durum
        FROM faturalar WHERE musteri_id=%s AND COALESCE(durum,'') != 'odendi' AND vade_tarihi IS NOT NULL
    """, (mid,))
    toplam_alacak = sum(float(f.get("toplam") or 0) for f in (faturalar_odenmemis or []))
    geciken_gun = 0
    for f in (faturalar_odenmemis or []):
        vd = _parse_vade(f.get("vade_tarihi"))
        if vd and vd < bugun:
            geciken_gun = max(geciken_gun, (bugun - vd).days)
    if geciken_gun > 30:
        odeme_durumu = "kritik"
    elif geciken_gun > 0 or (faturalar_odenmemis and len(faturalar_odenmemis) > 0):
        odeme_durumu = "yakin"
    else:
        odeme_durumu = "tam"

    # Tüm faturalar (12 ay grid)
    tum_faturalar = fetch_all("""
        SELECT COALESCE(toplam, tutar) as toplam, vade_tarihi, durum
        FROM faturalar WHERE musteri_id=%s AND vade_tarihi IS NOT NULL
    """, (mid,))
    fat_ay = {}
    for f in (tum_faturalar or []):
        vd = _parse_vade(f.get("vade_tarihi"))
        if not vd:
            continue
        ay = vd.month
        if ay not in fat_ay:
            fat_ay[ay] = []
        fat_ay[ay].append(f)
    aylik_durum = []
    for ay in range(1, 13):
        durum = "gelecek"
        for f in fat_ay.get(ay, []):
            vd = _parse_vade(f.get("vade_tarihi"))
            if not vd:
                continue
            if (f.get("durum") or "").lower() == "odendi":
                durum = "odendi"
                break
            if vd < bugun:
                durum = "gecikti"
                break
            if vd == bugun or (vd - bugun).days <= 7:
                durum = "bugun_yakin"
                break
            durum = "gelecek"
            break
        aylik_durum.append(durum)

    # Kargolar (geçmiş + son durum)
    kargolar_raw = fetch_all(
        "SELECT id, tarih, takip_no, kargo_firmasi, teslim_alan, notlar, created_at FROM kargolar WHERE musteri_id=%s ORDER BY created_at DESC LIMIT 30",
        (mid,),
    )
    kargolar = []
    for k in (kargolar_raw or []):
        teslim = (k.get("teslim_alan") or "").strip()
        durum = "Teslim" if teslim else "Bekliyor"
        created = k.get("created_at")
        saat_str = created.strftime("%H:%M") if hasattr(created, "strftime") else (str(created)[11:16] if created else "—")
        tarih_str = k.get("tarih")
        if hasattr(tarih_str, "strftime"):
            tarih_str = tarih_str.strftime("%d.%m.%Y")
        else:
            tarih_str = str(tarih_str)[:10] if tarih_str else "—"
        kargolar.append({
            "id": k.get("id"),
            "tarih": tarih_str,
            "takip_no": k.get("takip_no") or "—",
            "teslim_alan": teslim or "—",
            "durum": durum,
            "saat": saat_str,
            "kargo_firmasi": k.get("kargo_firmasi") or "—",
        })
    kargo_bekleyen = any((k.get("durum") or "") == "Bekliyor" for k in kargolar)
    son_kargo_durum = kargolar[0]["durum"] if kargolar else "—"

    # Son faturalar (tahsilat / ekstre için)
    faturalar_list = fetch_all("""
        SELECT id, fatura_no, COALESCE(toplam, tutar) as toplam, fatura_tarihi, vade_tarihi, durum
        FROM faturalar WHERE musteri_id=%s ORDER BY COALESCE(vade_tarihi, fatura_tarihi) DESC NULLS LAST LIMIT 15
    """, (mid,))
    faturalar_json = []
    for f in (faturalar_list or []):
        vt = f.get("vade_tarihi")
        ft = f.get("fatura_tarihi")
        faturalar_json.append({
            "id": f.get("id"),
            "fatura_no": f.get("fatura_no") or "—",
            "toplam": float(f.get("toplam") or 0),
            "fatura_tarihi": ft.strftime("%d.%m.%Y") if hasattr(ft, "strftime") else (str(ft)[:10] if ft else "—"),
            "vade_tarihi": vt.strftime("%d.%m.%Y") if hasattr(vt, "strftime") else (str(vt)[:10] if vt else "—"),
            "durum": (f.get("durum") or "—"),
        })

    return jsonify({
        "ok": True,
        "musteri": {
            "id": musteri["id"],
            "name": musteri.get("name") or "—",
            "phone": musteri.get("phone") or "—",
            "email": musteri.get("email") or "—",
            "tax_number": musteri.get("tax_number") or "—",
            "address": musteri.get("address") or "—",
            "office_code": musteri.get("office_code") or "—",
            "notes": musteri.get("notes") or "—",
        },
        "faturalar": faturalar_json,
        "odeme_durumu": odeme_durumu,
        "toplam_alacak": round(toplam_alacak, 2),
        "geciken_gun": geciken_gun,
        "aylik_durum": aylik_durum,
        "sozlesme_bitis_str": sozlesme_str or "—",
        "sozlesme_yakin": sozlesme_yakin,
        "hizmet_turu": (kyc.get("hizmet_turu") or "—") if kyc else "—",
        "kargo_bekleyen": kargo_bekleyen,
        "son_kargo_durum": son_kargo_durum,
        "kargolar": kargolar,
    })


@bp.route("/giris", methods=["GET"])
@giris_gerekli
def giris():
    """Ayrıntılı müşteri giriş / KYC ekranı (masaüstü gibi)"""
    return render_template("musteriler/giris.html")


@bp.route("/ekle", methods=["GET", "POST"])
@giris_gerekli
def ekle():
    """Yeni müşteri ekle"""
    if request.method == "POST":
        ma = (request.form.get("musteri_adi") or "").strip() or None
        row = execute_returning(
            """INSERT INTO customers (name, musteri_adi, email, phone, address, notes)
               VALUES (%s,%s,%s,%s,%s,%s) RETURNING id""",
            (request.form.get("name"),
             ma,
             request.form.get("email"),
             request.form.get("phone"),
             request.form.get("address"),
             request.form.get("notes")))
        flash("✓ Müşteri eklendi.", "success")
        return redirect(url_for("musteriler.detay", mid=row["id"]))
    return render_template("musteriler/ekle.html")


@bp.route("/import", methods=["GET", "POST"])
@giris_gerekli
def import_excel():
    """Excel'den müşteri içeri aktar"""
    if request.method == "POST":
        try:
            f = request.files.get("file")
            if not f or not f.filename:
                flash("Lütfen bir Excel dosyası seçin.", "warning")
                return redirect(request.url)
            raw = f.read()
            if not raw:
                flash("Dosya boş veya okunamadı.", "danger")
                return redirect(request.url)
            try:
                fn = (f.filename or "").lower()
                if fn.endswith(".xls") and not fn.endswith(".xlsx"):
                    try:
                        df = pd.read_excel(BytesIO(raw), header=0)
                    except Exception:
                        flash("Eski .xls formatı desteklenmiyor. Dosyayı .xlsx olarak kaydedip tekrar deneyin.", "danger")
                        return redirect(request.url)
                else:
                    df = pd.read_excel(BytesIO(raw), engine="openpyxl", header=0)
            except ImportError:
                flash("Excel desteği için openpyxl yüklü değil. Kurulum: pip install openpyxl", "danger")
                return redirect(request.url)
            except Exception as e:
                flash(f"Excel okunamadı: {e}", "danger")
                return redirect(request.url)

            if df.empty or len(df) == 0:
                flash("Excel dosyasında veri satırı yok.", "warning")
                return redirect(request.url)

            cols = []
            for i, c in enumerate(df.columns):
                s = (str(c).strip().lower() if c is not None else "") or f"unnamed_{i}"
                cols.append(s)
            df.columns = cols

            def find_col(keys):
                for k in keys:
                    for col in df.columns:
                        if k in col:
                            return col
                return None

            name_col = find_col([
                "ad", "name", "ünvan", "unvan", "firma", "müşteri", "musteri", "adı", "adi",
                "cari", "baslik", "başlık", "musteri adi", "müşteri adı"
            ])
            email_col = find_col(["email", "e-posta", "eposta", "mail"])
            phone_col = find_col(["telefon", "phone", "tel", "gsm", "cep"])
            addr_col = find_col(["adres", "address"])
            tax_col = find_col(["vergi", "tax", "vkn", "tckn", "vergi no", "vergino"])
            yetkili_col = find_col(["yetkili", "yetkili kişi", "contact", "yetkili kisi"])
            durum_col = find_col(["durum", "status", "durumu"])
            rent_start_col = find_col(["başlangıç tarihi", "baslangic tarihi"])
            ilk_kira_col = find_col(["ilk kira"])
            guncel_kira_col = find_col(["güncel kira", "guncel kira"])

            if not name_col and not yetkili_col:
                flash("Excel'de müşteri adı veya yetkili sütunu bulunamadı. Ad/Unvan veya Yetkili Kişi gerekli.", "danger")
                return redirect(request.url)

            def _cell(row, col):
                if not col:
                    return None
                v = row.get(col)
                if v is None:
                    return None
                try:
                    if pd.isna(v):
                        return None
                except (TypeError, ValueError):
                    pass
                s = str(v).strip()
                if not s or s.lower() == "nan":
                    return None
                return s

            ensure_customers_rent_columns()
            ensure_customers_excel_columns()

            def _norm_tax(t):
                """Vergi No: 10 hane ise baştaki sıfırları kaldır."""
                if not t:
                    return None
                s = str(t).strip()
                if not s:
                    return None
                if s.isdigit() and len(s) == 10:
                    return s.lstrip("0") or "0"
                return s

            def _norm_durum(val):
                """Excel'deki faal→aktif, terk→pasif."""
                if not val:
                    return None
                v = str(val).strip().lower()
                if v == "faal":
                    return "aktif"
                if v == "terk":
                    return "pasif"
                return v if v in ("aktif", "pasif") else None

            def _parse_date_excel(value):
                """Excel tarih hücresi -> date."""
                if value is None:
                    return None
                if isinstance(value, (datetime, date)):
                    return value.date() if isinstance(value, datetime) else value
                s = str(value).strip()
                if not s:
                    return None
                for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                    try:
                        return datetime.strptime(s[:10], fmt).date()
                    except Exception:
                        continue
                try:
                    return pd.to_datetime(s).date()
                except Exception:
                    return None

            def _parse_money_excel(value) -> float:
                """Türkçe para formatı: 1.200 = 1200, 1.234,56 = 1234.56."""
                if value is None:
                    return 0.0
                if isinstance(value, (int, float)):
                    return float(value)
                s = str(value).strip()
                if not s:
                    return 0.0
                s = s.replace(" ", "").replace("\u00a0", "")
                s = re.sub(r"[^0-9,.-]", "", s)
                if not s:
                    return 0.0
                # Virgül varsa Türkçe ondalık: 1.234,56 → 1234.56
                if "," in s:
                    if "." in s:
                        s = s.replace(".", "").replace(",", ".")
                    else:
                        s = s.replace(",", ".")
                else:
                    # Sadece nokta: 1.200 = 1200 (binlik) mı yoksa 1.25 = 1.25 (ondalık) mı?
                    if "." in s:
                        parts = s.split(".")
                        if len(parts) == 2 and len(parts[1]) == 3 and parts[1].isdigit():
                            s = s.replace(".", "")
                try:
                    return float(s)
                except Exception:
                    return 0.0

            try:
                ensure_customers_durum()
            except Exception:
                pass
            inserted = 0
            updated = 0
            errors = []
            total_rows = len(df)
            # Tek bağlantı ile tüm import (çok daha hızlı)
            with get_db() as conn:
                cur = conn.cursor()
                # Mevcut vergi no -> id haritası (bir sorguda)
                cur.execute("SELECT id, TRIM(COALESCE(tax_number,'')) AS tn FROM customers")
                existing_by_tax = {}
                for r in cur.fetchall() or []:
                    tn = str(r.get("tn") or "").strip()
                    if tn:
                        existing_by_tax[tn] = r.get("id")
                    if tn and str(tn).isdigit() and len(str(tn)) == 10:
                        normalized = tn.lstrip("0") or "0"
                        if normalized not in existing_by_tax:
                            existing_by_tax[normalized] = r.get("id")
                for idx, row in df.iterrows():
                    name = _cell(row, name_col) or _cell(row, yetkili_col)
                    if not name:
                        name = f"Müşteri-{idx + 2}"
                    email = _cell(row, email_col)
                    phone = _cell(row, phone_col)
                    address = _cell(row, addr_col)
                    tax_raw = _cell(row, tax_col)
                    tax = _norm_tax(tax_raw) if tax_raw else None
                    durum_val = _norm_durum(_cell(row, durum_col))
                    rent_start = _parse_date_excel(row.get(rent_start_col)) if rent_start_col else None
                    ilk_kira = _parse_money_excel(row.get(ilk_kira_col)) if ilk_kira_col else 0.0
                    guncel_kira = _parse_money_excel(row.get(guncel_kira_col)) if guncel_kira_col else 0.0
                    if guncel_kira <= 0 and ilk_kira > 0:
                        guncel_kira = ilk_kira
                    rent_start_year = rent_start.year if rent_start else None
                    rent_start_month = rent_start.strftime("%B") if rent_start else None
                    try:
                        existing_id = None
                        if tax:
                            existing_id = existing_by_tax.get(tax) or existing_by_tax.get((tax_raw or "").strip())
                        if existing_id:
                            cur.execute(
                                "UPDATE customers SET name=%s,email=%s,phone=%s,address=%s,tax_number=%s,durum=%s,"
                                "rent_start_date=%s,rent_start_year=%s,rent_start_month=%s,"
                                "ilk_kira_bedeli=%s,guncel_kira_bedeli=%s WHERE id=%s",
                                (
                                    name,
                                    email or None,
                                    phone or None,
                                    address or None,
                                    tax or None,
                                    durum_val,
                                    rent_start,
                                    rent_start_year,
                                    rent_start_month,
                                    ilk_kira or 0.0,
                                    guncel_kira or 0.0,
                                    existing_id,
                                ),
                            )
                            updated += 1
                        else:
                            cur.execute(
                                "INSERT INTO customers (name, email, phone, address, tax_number, durum,"
                                "rent_start_date, rent_start_year, rent_start_month, ilk_kira_bedeli, guncel_kira_bedeli) "
                                "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                                (
                                    name,
                                    email or None,
                                    phone or None,
                                    address or None,
                                    tax or None,
                                    durum_val,
                                    rent_start,
                                    rent_start_year,
                                    rent_start_month,
                                    ilk_kira or 0.0,
                                    guncel_kira or 0.0,
                                ),
                            )
                            inserted += 1
                    except Exception as e:
                        errors.append(f"Satır {idx + 2}: {name[:30]} — {e}")

            processed = inserted + updated
            percent = int(round(processed * 100 / total_rows)) if total_rows else 0
            msg_extra = f" ({inserted} yeni, {updated} güncelleme)" if updated else ""

            if errors:
                flash(
                    f"Excel aktarımı tamamlandı: {processed}/{total_rows} (%{percent}) satır işlendi{msg_extra}. "
                    f"{len(errors)} satırda hata. Detay: "
                    + "; ".join(errors[:3]) + ("..." if len(errors) > 3 else ""),
                    "warning",
                )
                return redirect(
                    url_for(
                        "musteriler.index",
                        import_sonuc="uyari",
                        imported=processed,
                        import_hatalar=len(errors),
                        import_percent=percent,
                        import_total=total_rows,
                    )
                )
            flash(f"Excel aktarımı başarılı: {processed}/{total_rows} (%{percent}) satır işlendi{msg_extra}. Toplam benzersiz müşteri: {inserted} yeni + güncellemeler.", "success")
            return redirect(
                url_for(
                    "musteriler.index",
                    import_sonuc="ok",
                    imported=processed,
                    import_hatalar=0,
                    import_percent=percent,
                    import_total=total_rows,
                )
            )
        except Exception as e:
            flash(f"Aktarım sırasında hata oluştu: {e}", "danger")
            return redirect(request.url)

    return render_template("musteriler/import.html")


@bp.route("/tumunu-sil", methods=["GET", "POST"])
@giris_gerekli
def tumunu_sil():
    """Tüm müşteri verilerini siler (tahsilat, fatura, kargo, sözleşme dahil). Geri alınamaz."""
    if request.method == "POST":
        onay = (request.form.get("onay") or request.args.get("onay") or "").strip().lower()
        if onay != "evet":
            flash("İşlem iptal edildi. Onay için 'evet' yazmanız gerekiyor.", "warning")
            return redirect(url_for("musteriler.tumunu_sil"))
        try:
            clear_all_customers()
            flash("Tüm müşteri verileri silindi. Şimdi Excel'den yeniden yükleyebilirsiniz.", "success")
            return redirect(url_for("musteriler.import_excel"))
        except Exception as e:
            flash(f"Müşteriler silinirken hata: {e}", "danger")
            return redirect(url_for("musteriler.tumunu_sil"))
    return render_template("musteriler/tumunu_sil_onay.html")


@bp.route("/export")
@giris_gerekli
def export_excel():
    """Müşteri listesini Excel olarak dışa aktar"""
    tum_yillar_odenmis = request.args.get("tum_yillar_odenmis") == "1"
    if tum_yillar_odenmis:
        rows = fetch_all("""
            SELECT c.* FROM customers c
            WHERE NOT EXISTS (
                SELECT 1 FROM faturalar f
                WHERE f.musteri_id = c.id AND (f.durum IS NULL OR f.durum != 'odendi')
            )
            ORDER BY c.name
        """)
    else:
        rows = fetch_all("SELECT * FROM customers ORDER BY name")
    if not rows:
        flash("Dışa aktarılacak müşteri yok.", "warning")
        return redirect(url_for("musteriler.index"))
    df = pd.DataFrame(rows)
    for col in list(df.columns):
        try:
            df[col] = df[col].apply(
                lambda x: x.isoformat()[:10] if hasattr(x, "isoformat") and x is not None else (x if x is not None else "")
            )
        except Exception:
            pass
    buf = BytesIO()
    df.to_excel(buf, index=False, engine="openpyxl")
    buf.seek(0)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=f"musteriler_{date.today().isoformat()}.xlsx",
    )


@bp.route("/gemini-analiz")
@giris_gerekli
def gemini_analiz():
    """Gemini AI Analiz sayfası"""
    return render_template("musteriler/gemini_analiz.html", gemini_available=GEMINI_AVAILABLE)


def _musteri_ozet_metni():
    """Müşteri verilerinden analiz için kısa özet metni üretir."""
    toplam = fetch_one("SELECT COUNT(*) AS n FROM customers")
    n = (toplam or {}).get("n") or 0
    kira = fetch_one(
        "SELECT COUNT(*) AS n, COALESCE(SUM(ilk_kira_bedeli), 0) AS toplam FROM customers WHERE ilk_kira_bedeli IS NOT NULL"
    )
    kira_n = (kira or {}).get("n") or 0
    kira_toplam = float((kira or {}).get("toplam") or 0)
    odendi = fetch_one(
        "SELECT COUNT(DISTINCT musteri_id) AS n FROM faturalar WHERE durum = %s",
        ("odendi",),
    )
    odenen_musteri = (odendi or {}).get("n") or 0
    bekleyen = fetch_one(
        "SELECT COUNT(DISTINCT musteri_id) AS n FROM faturalar WHERE durum IS NULL OR durum != %s",
        ("odendi",),
    )
    bekleyen_musteri = (bekleyen or {}).get("n") or 0
    son_5 = fetch_all(
        "SELECT name, ilk_kira_bedeli, rent_start_date FROM customers ORDER BY id DESC LIMIT 5"
    )
    satirlar = [
        f"Toplam müşteri sayısı: {n}",
        f"Kira bedeli girilmiş müşteri: {kira_n}, toplam aylık kira (ilk bedel): {kira_toplam:,.2f} TL",
        f"Faturası ödenmiş müşteri sayısı: {odenen_musteri}",
        f"Ödenmemiş fatura olan müşteri sayısı: {bekleyen_musteri}",
        "",
        "Son eklenen 5 müşteri (ad, ilk kira, başlangıç tarihi):",
    ]
    for m in (son_5 or []):
        ad = (m.get("name") or "").strip() or "—"
        k = m.get("ilk_kira_bedeli")
        k_str = f"{float(k):,.2f} TL" if k is not None else "—"
        t = m.get("rent_start_date")
        t_str = str(t)[:10] if t else "—"
        satirlar.append(f"  - {ad} | {k_str} | {t_str}")
    return "\n".join(satirlar)


def _sistem_ozet_metni():
    """Sistem analizi için genişletilmiş özet: müşteri + fatura + tahsilat + bakiye."""
    base = _musteri_ozet_metni()
    extra = []
    try:
        fatura = fetch_one(
            "SELECT COUNT(*) AS n, COALESCE(SUM(COALESCE(toplam, tutar, 0)), 0) AS toplam FROM faturalar"
        )
        if fatura:
            extra.append(
                f"Toplam fatura sayısı: {fatura.get('n') or 0}, toplam tutar: {float(fatura.get('toplam') or 0):,.2f} TL"
            )
        tahsilat = fetch_one(
            "SELECT COUNT(*) AS n, COALESCE(SUM(tutar), 0) AS toplam FROM tahsilatlar"
        )
        if tahsilat:
            extra.append(
                f"Toplam tahsilat sayısı: {tahsilat.get('n') or 0}, toplam tahsilat: {float(tahsilat.get('toplam') or 0):,.2f} TL"
            )
        bakiye = fetch_one(
            "SELECT COALESCE(SUM(current_balance), 0) AS t FROM customers"
        )
        if bakiye is not None:
            extra.append(f"Toplam cari bakiye (fatura - tahsilat): {float(bakiye.get('t') or 0):,.2f} TL")
        geciken = fetch_one(
            """SELECT COUNT(DISTINCT c.id) AS n FROM customers c
               JOIN faturalar f ON f.musteri_id = c.id
               WHERE (COALESCE(f.vade_tarihi::date, f.fatura_tarihi::date + INTERVAL '30 days') < CURRENT_DATE)
               AND COALESCE(f.durum, '') != 'odendi'"""
        )
        if geciken and (geciken.get("n") or 0) > 0:
            extra.append(f"Gecikmiş ödemesi olan müşteri sayısı: {geciken.get('n')}")
    except Exception:
        pass
    if extra:
        base += "\n\n--- Ek sistem verileri ---\n" + "\n".join(extra)
    return base


def _sistem_ozet_metni_kisa():
    """Sistem analizi için kısa özet (token tasarrufu, ücretsiz kota): sadece sayılar, liste yok."""
    satirlar = []
    try:
        toplam = fetch_one("SELECT COUNT(*) AS n FROM customers")
        n = (toplam or {}).get("n") or 0
        kira = fetch_one(
            "SELECT COUNT(*) AS n, COALESCE(SUM(ilk_kira_bedeli), 0) AS toplam FROM customers WHERE ilk_kira_bedeli IS NOT NULL"
        )
        kira_n = (kira or {}).get("n") or 0
        kira_toplam = float((kira or {}).get("toplam") or 0)
        satirlar.append(f"Toplam müşteri: {n}. Kira girilmiş: {kira_n}, aylık toplam kira: {kira_toplam:,.0f} TL.")
        fatura = fetch_one(
            "SELECT COUNT(*) AS n, COALESCE(SUM(COALESCE(toplam, tutar, 0)), 0) AS toplam FROM faturalar"
        )
        if fatura:
            satirlar.append(f"Fatura: {fatura.get('n') or 0} adet, toplam {float(fatura.get('toplam') or 0):,.0f} TL.")
        tahsilat = fetch_one(
            "SELECT COUNT(*) AS n, COALESCE(SUM(tutar), 0) AS toplam FROM tahsilatlar"
        )
        if tahsilat:
            satirlar.append(f"Tahsilat: {tahsilat.get('n') or 0} adet, toplam {float(tahsilat.get('toplam') or 0):,.0f} TL.")
        bakiye = fetch_one("SELECT COALESCE(SUM(current_balance), 0) AS t FROM customers")
        if bakiye is not None:
            satirlar.append(f"Cari bakiye: {float(bakiye.get('t') or 0):,.0f} TL.")
    except Exception:
        pass
    return "\n".join(satirlar) if satirlar else "Veri yok."


@bp.route("/api/gemini-analiz", methods=["POST"])
@giris_gerekli
def api_gemini_analiz():
    """Müşteri verisi özeti + isteğe bağlı kullanıcı sorusu ile Gemini'den analiz alır."""
    data = request.get_json(silent=True) or request.form
    soru = (data.get("soru") or data.get("prompt") or "").strip()
    sistem_analizi = data.get("sistem_analizi") in (True, "true", "1")
    if sistem_analizi:
        # Kısa özet + tek istek: ücretsiz planda kota aşımını önler (1 tıklama = 1 API çağrısı)
        context = _sistem_ozet_metni_kisa()
        soru = (
            "Bu ERP sisteminin mevcut müşteri, fatura ve tahsilat verilerine göre kısa bir sistem analizi yap. "
            "Güçlü yönler, olası riskler ve 3-5 maddelik iyileştirme önerileri ver. Yanıtı Türkçe, net ve kısa tut."
        )
        ok, metin = gemini_analiz_yap(context, soru, tek_istek=True)
    else:
        context = _musteri_ozet_metni()
        ok, metin = gemini_analiz_yap(context, soru)
    if ok:
        return jsonify({"ok": True, "metin": metin})
    return jsonify({"ok": False, "hata": metin}), 400


@bp.route("/<int:mid>/duzenle", methods=["GET", "POST"])
@giris_gerekli
def duzenle(mid):
    """Müşteri düzenle"""
    musteri = fetch_one("SELECT * FROM customers WHERE id=%s", (mid,))
    if not musteri:
        return redirect(url_for("musteriler.index"))
    if request.method == "POST":
        execute(
            """UPDATE customers SET name=%s,email=%s,phone=%s,address=%s,notes=%s
               WHERE id=%s""",
            (request.form.get("name"), request.form.get("email"),
             request.form.get("phone"), request.form.get("address"),
             request.form.get("notes"), mid))
        flash("✓ Müşteri güncellendi.", "success")
        return redirect(url_for("musteriler.detay", mid=mid))
    return render_template("musteriler/duzenle.html", musteri=musteri)


@bp.route("/<int:mid>/sil", methods=["POST"])
@giris_gerekli
def sil(mid):
    """Müşteri sil"""
    execute("DELETE FROM customers WHERE id=%s", (mid,))
    flash("Müşteri silindi.", "info")
    return redirect(url_for("musteriler.index"))


# ── API ENDPOİNTS ────────────────────────────────────────────────────────────

@bp.route("/api/liste")
@giris_gerekli
def api_liste():
    """Müşteri listesi JSON (dropdown için)"""
    musteriler = fetch_all("SELECT id, name, musteri_adi FROM customers ORDER BY name")
    return jsonify(musteriler)


@bp.route("/api/list_full")
@giris_gerekli
def api_list_full():
    """Tam müşteri listesi JSON.

    Eski davranış: parametre yok → düz dizi (geriye dönük uyumluluk).
    ?sayfa=1&limit=50 → { paginate: true, rows, toplam_musteri, toplam_sayfa, sayfa, limit }.
    """
    base_sql = (
        "SELECT id, name, email, phone, address, office_code, notes, created_at FROM customers ORDER BY name"
    )
    paginate = ("sayfa" in request.args) or ("limit" in request.args)
    sayfa = max(1, request.args.get("sayfa", default=1, type=int) or 1)
    limit = request.args.get("limit", type=int)
    if limit is None:
        limit = 50
    limit = max(1, min(500, int(limit)))

    if not paginate:
        rows = fetch_all(base_sql)
        return jsonify(rows)

    cnt = fetch_one("SELECT COUNT(*) AS n FROM customers") or {}
    try:
        toplam_musteri = int(cnt.get("n") or 0)
    except (TypeError, ValueError):
        toplam_musteri = 0
    toplam_sayfa = max(1, (toplam_musteri + limit - 1) // limit) if toplam_musteri else 1
    offset = (sayfa - 1) * limit
    rows = fetch_all(base_sql + " LIMIT %s OFFSET %s", (limit, offset))
    return jsonify(
        {
            "paginate": True,
            "rows": rows,
            "toplam_musteri": toplam_musteri,
            "toplam_sayfa": toplam_sayfa,
            "sayfa": sayfa,
            "limit": limit,
        }
    )


@bp.route("/api/bulk-update", methods=["POST"])
@giris_gerekli
def api_bulk_update():
    """Hızlı bilgi düzenleme: tablodaki değişiklikleri toplu kaydet. Cari kart (customers) güncellenir."""
    try:
        ensure_customers_quick_edit_columns()
        ensure_customers_kapanis_tarihi()
        ensure_grup2_etiketleri_table()
        ensure_customers_grup2_secimleri()
    except Exception:
        pass
    allowed_grup2_slugs = set()
    try:
        rows = fetch_all("SELECT slug FROM grup2_etiketleri WHERE COALESCE(aktif, TRUE)")
        allowed_grup2_slugs = {str(r.get("slug") or "").strip() for r in (rows or []) if str(r.get("slug") or "").strip()}
    except Exception:
        allowed_grup2_slugs = set()
    data = request.get_json(silent=True) or {}
    updates = data.get("updates") or []
    if not isinstance(updates, list):
        return jsonify({"ok": False, "message": "Geçersiz veri."}), 400
    updated = 0
    with get_db() as conn:
        cur = conn.cursor()
        for row in updates:
            mid = row.get("id")
            if not mid:
                continue
            try:
                mid = int(mid)
            except (TypeError, ValueError):
                continue
            sets = []
            params = []
            rent_start_date_val = None
            for key, col in (
                ("yetkili_kisi", "yetkili_kisi"),
                ("hizmet_turu", "hizmet_turu"),
                ("phone", "phone"),
                ("rent_start_date", "rent_start_date"),
                ("ilk_kira_bedeli", "ilk_kira_bedeli"),
                ("guncel_kira_bedeli", "guncel_kira_bedeli"),
                ("reel_kira_bedeli", "reel_kira_bedeli"),
                ("manuel_borc", "manuel_borc"),
                ("son_odeme_tarihi", "son_odeme_tarihi"),
                ("durum", "durum"),
                ("kapanis_tarihi", "kapanis_tarihi"),
            ):
                if key not in row:
                    continue
                val = row[key]
                if key in ("ilk_kira_bedeli", "guncel_kira_bedeli", "reel_kira_bedeli", "manuel_borc") and val is not None:
                    try:
                        val = float(val)
                    except (TypeError, ValueError):
                        val = None
                if key in ("son_odeme_tarihi", "rent_start_date", "kapanis_tarihi"):
                    if val is None or val == "":
                        val = None
                    elif hasattr(val, "year"):
                        pass
                    else:
                        s = str(val).strip()[:10]
                        if s and s not in ("—", ""):
                            try:
                                val = date(*[int(x) for x in s.split("-")])
                            except Exception:
                                val = None
                        else:
                            val = None
                    if key == "rent_start_date":
                        rent_start_date_val = val  # date objesi veya parse edilmiş if hasattr(val, "year") else val
                if key == "durum" and val is not None:
                    val = (str(val).strip().lower() or None)
                    if val and val not in ("aktif", "pasif"):
                        val = None
                sets.append(f"{col} = %s")
                params.append(val)
            if "grup2_secimleri" in row:
                raw = row.get("grup2_secimleri")
                g2_list = []
                if isinstance(raw, list):
                    g2_list = [str(x).strip() for x in raw if str(x).strip()]
                elif isinstance(raw, str):
                    s = raw.strip()
                    if s.startswith("["):
                        try:
                            j = json.loads(s)
                            if isinstance(j, list):
                                g2_list = [str(x).strip() for x in j if str(x).strip()]
                        except Exception:
                            g2_list = []
                    elif s:
                        g2_list = [x.strip() for x in s.split("|") if x.strip()]
                if allowed_grup2_slugs:
                    g2_list = [s for s in g2_list if s in allowed_grup2_slugs]
                g2_list = list(dict.fromkeys(g2_list))
                sets.append("grup2_secimleri = %s")
                params.append(g2_list)
                sets.append("bizim_hesap = %s")
                params.append("bizim_hesap" in g2_list)
            if rent_start_date_val:
                sets.append("rent_start_year = %s")
                params.append(rent_start_date_val.year)
                sets.append("rent_start_month = %s")
                params.append(MONTHS_TR[rent_start_date_val.month - 1] if rent_start_date_val.month else None)
            if not sets:
                continue
            params.append(mid)
            cur.execute(
                "UPDATE customers SET " + ", ".join(sets) + " WHERE id = %s",
                params,
            )
            updated += cur.rowcount or 0
    return jsonify({"ok": True, "updated": updated, "message": "Veriler başarıyla kaydedildi ve cari kartlar güncellendi."})


# ── Giriş / KYC API ──────────────────────────────────────────────────────────

@bp.route("/api/musteri/ara")
@giris_gerekli
def api_musteri_ara():
    """Mevcut müşteriye bağlamak için arama (q=)"""
    q = (request.args.get("q") or "").strip()
    if not q:
        rows = fetch_all("SELECT id, name, musteri_adi FROM customers ORDER BY name LIMIT 50")
    else:
        w3 = customers_arama_sql_giris_genis("")
        rows = fetch_all(
            f"SELECT id, name, musteri_adi FROM customers WHERE {w3} ORDER BY name LIMIT 30",
            customers_arama_params_giris_genis(q),
        )
    return jsonify(rows)


@bp.route("/api/kyc/getir")
@giris_gerekli
def api_kyc_getir():
    """Müşteriye ait KYC kaydını getir"""
    musteri_id = request.args.get("musteri_id")
    if not musteri_id:
        return jsonify(None)
    row = fetch_one(
        "SELECT * FROM musteri_kyc WHERE musteri_id = %s ORDER BY id DESC LIMIT 1",
        (musteri_id,),
    )

    # Eğer henüz KYC kaydı yoksa, en azından customers tablosundaki temel bilgileri döndür
    if not row:
        cust = fetch_one(
            "SELECT id, name, musteri_adi, email, phone, address, tax_number, durum, kapanis_tarihi FROM customers WHERE id = %s",
            (musteri_id,),
        )
        if not cust:
            return jsonify(None)
        # frontend'deki formu doldurmak için alan isimlerini eşleştir
        return jsonify(
            {
                "musteri_id": cust.get("id"),
                "sirket_unvani": cust.get("name"),
                "unvan": cust.get("name"),
                "musteri_adi": cust.get("musteri_adi"),
                "durum": cust.get("durum"),
                "kapanis_tarihi": cust.get("kapanis_tarihi"),
                "email": cust.get("email"),
                "yetkili_email": cust.get("email"),
                "yetkili_tel": cust.get("phone"),
                "yeni_adres": cust.get("address"),
                "vergi_no": cust.get("tax_number"),
            }
        )

    # Tarih ve sayısal alanları JSON uyumlu yap
    out = dict(row)
    for k in list(out.keys()):
        v = out[k]
        if hasattr(v, "isoformat"):
            out[k] = v.isoformat()[:10] if v else None
        elif hasattr(v, "days"):  # timedelta
            out[k] = str(v)
    return jsonify(out)


@bp.route("/api/kyc/kaydet", methods=["POST"])
@giris_gerekli
def api_kyc_kaydet():
    """Giriş formundan KYC kaydı kaydet / güncelle"""
    try:
        data = request.json or request.form
        ensure_customers_hazir_ofis_oda()
        ensure_musteri_kyc_hazir_ofis_oda_no()
        ensure_musteri_kyc_kira_banka()
        ensure_musteri_kyc_odeme_duzeni()
        musteri_id = data.get("musteri_id")
        sirket_unvani = (data.get("sirket_unvani") or data.get("unvan") or "").strip()
        musteri_adi = (data.get("musteri_adi") or "").strip() or None
        vergi_no = (data.get("vergi_no") or "").strip()
        vergi_dairesi = (data.get("vergi_dairesi") or "").strip()
        mersis_no = (data.get("mersis_no") or "").strip()
        ticaret_sicil_no = (data.get("ticaret_sicil_no") or "").strip()
        kurulus_tarihi = data.get("kurulus_tarihi")
        faaliyet_konusu = (data.get("faaliyet_konusu") or "").strip()
        nace_kodu = (data.get("nace_kodu") or "").strip()
        eski_adres = (data.get("eski_adres") or "").strip()
        yeni_adres = (data.get("yeni_adres") or "").strip()
        sube_merkez = (data.get("sube_merkez") or "Merkez").strip()
        yetkili_adsoyad = (data.get("yetkili_adsoyad") or "").strip()
        yetkili_tcno = (data.get("yetkili_tcno") or "").strip()
        vergi_err, vergi_no_norm = _vergi_no_normalize_veya_hata_kyc(vergi_no, yetkili_tcno)
        if vergi_err:
            return jsonify({"ok": False, "mesaj": vergi_err}), 400
        vergi_no = vergi_no_norm
        yetkili_dogum = data.get("yetkili_dogum")
        yetkili_ikametgah = (data.get("yetkili_ikametgah") or "").strip()
        yetkili_tel = (data.get("yetkili_tel") or "").strip()
        yetkili_tel2 = (data.get("yetkili_tel2") or "").strip()
        yetkili_tel_aciklama = (data.get("yetkili_tel_aciklama") or "").strip()
        yetkili_tel2_aciklama = (data.get("yetkili_tel2_aciklama") or "").strip()
        yetkili_email = (data.get("yetkili_email") or "").strip()
        email = (data.get("email") or "").strip()
        # customers.email: arama / liste için şirket e-postası yoksa yetkili e-postası (KYC öncesi kayıt silinmesin)
        cust_row_email = email or yetkili_email or None
        hizmet_turu = (data.get("hizmet_turu") or "Sanal Ofis").strip()
        _df = (data.get("duzenli_fatura") or "duzenle").strip().lower()
        _df = _df.replace(" ", "_").replace("-", "_")
        _df = re.sub(r"[^a-z0-9_]", "", _df)
        if not _df or len(_df) > 120:
            _df = "duzenle"
        duzenli_fatura = _df
        _odemd = (data.get("odeme_duzeni") or "aylik").strip().lower()
        if _odemd not in ("aylik", "yillik", "3_aylik", "6_aylik", "manuel"):
            _odemd = "aylik"
        odeme_duzeni_manuel_txt = (
            (data.get("odeme_duzeni_manuel") or "").strip()[:200] if _odemd == "manuel" else ""
        )
        durum_m = (data.get("durum") or "aktif").strip().lower()
        if durum_m not in ("aktif", "pasif"):
            durum_m = "aktif"
        aylik_kira = float(str(data.get("aylik_kira") or 0).replace(",", ".") or 0)
        yillik_kira = float(str(data.get("yillik_kira") or 0).replace(",", ".") or 0)
        sozlesme_no = (data.get("sozlesme_no") or "").strip()
        sozlesme_tarihi = data.get("sozlesme_tarihi")
        sozlesme_bitis = data.get("sozlesme_bitis")
        kira_artis_tarihi = data.get("kira_artis_tarihi")
        notlar = (data.get("notlar") or "").strip()

        def parse_date(s):
            if not s:
                return None
            s = str(s).strip()[:10]
            for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"):
                try:
                    from datetime import datetime
                    return datetime.strptime(s, fmt).date()
                except ValueError:
                    continue
            return None

        kurulus_tarihi = parse_date(kurulus_tarihi)
        yetkili_dogum = parse_date(yetkili_dogum)
        sozlesme_tarihi = parse_date(sozlesme_tarihi)
        sozlesme_bitis = parse_date(sozlesme_bitis)
        kira_artis_tarihi = parse_date(kira_artis_tarihi)
        kapanis_tarihi = parse_date(data.get("kapanis_tarihi")) if durum_m == "pasif" else None

        durum_pasif = durum_m == "pasif"
        hazir_oda_val = _coerce_hazir_ofis_oda_no(data, hizmet_turu, durum_pasif)
        if hazir_oda_val is False:
            return jsonify(
                {"ok": False, "mesaj": "Hazır Ofis oda numarası 200–230 arasında olmalıdır."}
            ), 400

        try:
            _ks = int(float(str(data.get("kira_suresi_ay") or "").replace(",", ".").strip() or "0"))
            kira_suresi_ay = _ks if 1 <= _ks <= 240 else None
        except (TypeError, ValueError):
            kira_suresi_ay = None

        def _add_months_safe(d, months):
            if not d or months < 1:
                return d
            m = d.month - 1 + int(months)
            y = d.year + m // 12
            mo = m % 12 + 1
            last = calendar.monthrange(y, mo)[1]
            return date(y, mo, min(d.day, last))

        # Sözleşme başlangıcı geçmişte olsa bile, aktif müşteride bitiş daima bugüne göre
        # bir sonraki yıldönümüne taşınır — ancak formdan gelen kira süresi (ay) varsa
        # tam süre kullanılır; aksi halde aylık TÜFE grid'i yalnızca ~12 ay üretirdi.
        if sozlesme_tarihi:
            def _same_day_month_safe(src_date, target_year):
                m = int(src_date.month)
                d = int(src_date.day)
                try:
                    return date(target_year, m, d)
                except ValueError:
                    # 29 Şubat gibi tarihlerde ayın son gününe yuvarla.
                    if m == 12:
                        return date(target_year, 12, 31)
                    return date(target_year, m + 1, 1) - timedelta(days=1)

            today = date.today()
            kira_artis_tarihi = _same_day_month_safe(sozlesme_tarihi, today.year)
            if kira_suresi_ay:
                sozlesme_bitis = _add_months_safe(sozlesme_tarihi, kira_suresi_ay)
            elif durum_m != "pasif":
                next_renewal = _same_day_month_safe(sozlesme_tarihi, today.year)
                if next_renewal <= today:
                    next_renewal = _same_day_month_safe(sozlesme_tarihi, today.year + 1)
                sozlesme_bitis = next_renewal

        evrak_imza_sirkuleri = 1 if data.get("evrak_imza_sirkuleri") in (True, 1, "1", "on") else 0
        evrak_vergi_levhasi = 1 if data.get("evrak_vergi_levhasi") in (True, 1, "1", "on") else 0
        evrak_ticaret_sicil = 1 if data.get("evrak_ticaret_sicil") in (True, 1, "1", "on") else 0
        evrak_faaliyet_belgesi = 1 if data.get("evrak_faaliyet_belgesi") in (True, 1, "1", "on") else 0
        evrak_kimlik_fotokopi = 1 if data.get("evrak_kimlik_fotokopi") in (True, 1, "1", "on") else 0
        evrak_ikametgah = 1 if data.get("evrak_ikametgah") in (True, 1, "1", "on") else 0
        evrak_kase = 1 if data.get("evrak_kase") in (True, 1, "1", "on") else 0
        kira_nakit_tutar = _kyc_optional_money(data.get("kira_nakit_tutar"))
        kira_banka_tutar = _kyc_optional_money(data.get("kira_banka_tutar"))
        kira_nakit = data.get("kira_nakit") in (True, 1, "1", "on", "true", "yes")
        kira_banka = data.get("kira_banka") in (True, 1, "1", "on", "true", "yes")
        karma_kira = (
            kira_nakit_tutar is not None
            and kira_banka_tutar is not None
            and kira_nakit_tutar > 0
            and kira_banka_tutar > 0
        )
        if not karma_kira:
            if kira_nakit:
                kira_banka = False
            elif kira_banka:
                kira_nakit = False

        zorunlu = ["sirket_unvani", "vergi_no", "vergi_dairesi", "yeni_adres", "yetkili_adsoyad", "yetkili_tcno"]
        dolu = sum(1 for k in zorunlu if (data.get(k) or "").strip())
        tamamlanma_yuzdesi = int(dolu / len(zorunlu) * 100) if zorunlu else 0

        mevcut = fetch_one("SELECT id FROM musteri_kyc WHERE musteri_id = %s ORDER BY id DESC LIMIT 1", (musteri_id,)) if musteri_id else None

        def _hazir_oda_cakisma_kontrol(mid):
            if hazir_oda_val is None or not mid:
                return None
            crow = fetch_one(
                """
                SELECT id,
                    COALESCE(NULLIF(TRIM(musteri_adi), ''), NULLIF(TRIM(name), ''), '') AS ad_gor
                FROM customers
                WHERE id <> %s AND hazir_ofis_oda_no = %s
                  AND COALESCE(is_active, TRUE) = TRUE
                  AND LOWER(TRIM(COALESCE(durum, ''))) = 'aktif'
                LIMIT 1
                """,
                (int(mid), hazir_oda_val),
            )
            return crow

        if mevcut:
            ck = _hazir_oda_cakisma_kontrol(musteri_id)
            if ck:
                return jsonify(
                    {
                        "ok": False,
                        "mesaj": "Oda %s başka aktif müşteride (ID: %s, %s)."
                        % (hazir_oda_val, ck["id"], (ck.get("ad_gor") or "—").strip() or "—"),
                    }
                ), 409
            execute(
                """UPDATE musteri_kyc SET
                   sirket_unvani=%s, unvan=%s, musteri_adi=%s, vergi_no=%s, vergi_dairesi=%s, mersis_no=%s, ticaret_sicil_no=%s,
                   kurulus_tarihi=%s, faaliyet_konusu=%s, nace_kodu=%s, eski_adres=%s, yeni_adres=%s, sube_merkez=%s,
                   yetkili_adsoyad=%s, yetkili_tcno=%s, yetkili_dogum=%s, yetkili_ikametgah=%s,
                   yetkili_tel=%s, yetkili_tel2=%s, yetkili_tel_aciklama=%s, yetkili_tel2_aciklama=%s, yetkili_email=%s, email=%s,
                   hizmet_turu=%s, duzenli_fatura=%s, aylik_kira=%s, yillik_kira=%s, sozlesme_no=%s, sozlesme_tarihi=%s, sozlesme_bitis=%s,
                   kira_artis_tarihi=%s, kira_suresi_ay=%s, kira_nakit=%s, kira_banka=%s, kira_nakit_tutar=%s, kira_banka_tutar=%s, hazir_ofis_oda_no=%s,
                   evrak_imza_sirkuleri=%s, evrak_vergi_levhasi=%s, evrak_ticaret_sicil=%s, evrak_faaliyet_belgesi=%s,
                   evrak_kimlik_fotokopi=%s, evrak_ikametgah=%s, evrak_kase=%s, notlar=%s, tamamlanma_yuzdesi=%s,
                   odeme_duzeni=%s, odeme_duzeni_manuel=%s, updated_at=NOW()
                   WHERE id=%s""",
                (sirket_unvani, sirket_unvani, musteri_adi, vergi_no, vergi_dairesi, mersis_no, ticaret_sicil_no,
                 kurulus_tarihi, faaliyet_konusu, nace_kodu, eski_adres, yeni_adres, sube_merkez,
                 yetkili_adsoyad, yetkili_tcno, yetkili_dogum, yetkili_ikametgah,
                 yetkili_tel, yetkili_tel2, yetkili_tel_aciklama, yetkili_tel2_aciklama, yetkili_email, email,
                 hizmet_turu, duzenli_fatura, aylik_kira, yillik_kira, sozlesme_no, sozlesme_tarihi, sozlesme_bitis,
                 kira_artis_tarihi, kira_suresi_ay, kira_nakit, kira_banka, kira_nakit_tutar, kira_banka_tutar, hazir_oda_val,
                 evrak_imza_sirkuleri, evrak_vergi_levhasi, evrak_ticaret_sicil, evrak_faaliyet_belgesi,
                 evrak_kimlik_fotokopi, evrak_ikametgah, evrak_kase, notlar, tamamlanma_yuzdesi,
                 _odemd, odeme_duzeni_manuel_txt or None, mevcut["id"])
            )
            kyc_id = mevcut["id"]
            if musteri_id:
                execute(
                    """UPDATE customers SET name=%s, musteri_adi=%s, email=%s, phone=%s, address=%s, tax_number=%s,
                       durum=%s, kapanis_tarihi=%s, hizmet_turu=%s, hazir_ofis_oda_no=%s WHERE id=%s""",
                    (sirket_unvani or None, musteri_adi, cust_row_email, yetkili_tel or None, yeni_adres or None, vergi_no or None,
                     durum_m, kapanis_tarihi, hizmet_turu, hazir_oda_val, musteri_id)
                )
        else:
            if musteri_id:
                ck = _hazir_oda_cakisma_kontrol(musteri_id)
                if ck:
                    return jsonify(
                        {
                            "ok": False,
                            "mesaj": "Oda %s başka aktif müşteride (ID: %s, %s)."
                            % (hazir_oda_val, ck["id"], (ck.get("ad_gor") or "—").strip() or "—"),
                        }
                    ), 409
                execute(
                    """UPDATE customers SET name=%s, musteri_adi=%s, email=%s, phone=%s, address=%s, tax_number=%s,
                       durum=%s, kapanis_tarihi=%s, hizmet_turu=%s, hazir_ofis_oda_no=%s WHERE id=%s""",
                    (sirket_unvani or None, musteri_adi, cust_row_email, yetkili_tel or None, yeni_adres or None, vergi_no or None,
                     durum_m, kapanis_tarihi, hizmet_turu, hazir_oda_val, musteri_id)
                )
            else:
                # Yeni müşteri oluştur (oda çakışması: kayıt öncesi, başka aktif kartta aynı oda var mı)
                if hazir_oda_val is not None:
                    ex0 = fetch_one(
                        """
                        SELECT id,
                            COALESCE(NULLIF(TRIM(musteri_adi), ''), NULLIF(TRIM(name), ''), '') AS ad_gor
                        FROM customers
                        WHERE hazir_ofis_oda_no = %s
                          AND COALESCE(is_active, TRUE) = TRUE
                          AND LOWER(TRIM(COALESCE(durum, ''))) = 'aktif'
                        LIMIT 1
                        """,
                        (hazir_oda_val,),
                    )
                    if ex0:
                        return jsonify(
                            {
                                "ok": False,
                                "mesaj": "Oda %s başka aktif müşteride (ID: %s, %s)."
                                % (hazir_oda_val, ex0["id"], (ex0.get("ad_gor") or "—").strip() or "—"),
                            }
                        ), 409
                cust = execute_returning(
                    """INSERT INTO customers (name, musteri_adi, email, phone, address, notes, tax_number, durum, kapanis_tarihi, hizmet_turu, hazir_ofis_oda_no)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id""",
                    (sirket_unvani or "Yeni Müşteri", musteri_adi, cust_row_email, yetkili_tel or None, yeni_adres or None, notlar or None, vergi_no or None, durum_m, kapanis_tarihi, hizmet_turu, hazir_oda_val)
                )
                musteri_id = cust["id"] if cust else None
            row = execute_returning(
                """INSERT INTO musteri_kyc (
                   musteri_id, sirket_unvani, unvan, musteri_adi, vergi_no, vergi_dairesi, mersis_no, ticaret_sicil_no,
                   kurulus_tarihi, faaliyet_konusu, nace_kodu, eski_adres, yeni_adres, sube_merkez,
                   yetkili_adsoyad, yetkili_tcno, yetkili_dogum, yetkili_ikametgah,
                   yetkili_tel, yetkili_tel2, yetkili_tel_aciklama, yetkili_tel2_aciklama, yetkili_email, email,
                   hizmet_turu, duzenli_fatura, odeme_duzeni, odeme_duzeni_manuel, aylik_kira, yillik_kira, sozlesme_no, sozlesme_tarihi, sozlesme_bitis,
                   kira_artis_tarihi, kira_suresi_ay, kira_nakit, kira_banka, kira_nakit_tutar, kira_banka_tutar, hazir_ofis_oda_no,
                   evrak_imza_sirkuleri, evrak_vergi_levhasi, evrak_ticaret_sicil, evrak_faaliyet_belgesi,
                   evrak_kimlik_fotokopi, evrak_ikametgah, evrak_kase, notlar, tamamlanma_yuzdesi
                ) VALUES (
                   %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,
                   %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,
                   %s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,
                   %s,%s,%s,%s,%s,%s,%s
                ) RETURNING id""",
                (musteri_id, sirket_unvani, sirket_unvani, musteri_adi, vergi_no, vergi_dairesi, mersis_no, ticaret_sicil_no,
                 kurulus_tarihi, faaliyet_konusu, nace_kodu, eski_adres, yeni_adres, sube_merkez,
                 yetkili_adsoyad, yetkili_tcno, yetkili_dogum, yetkili_ikametgah,
                 yetkili_tel, yetkili_tel2, yetkili_tel_aciklama, yetkili_tel2_aciklama, yetkili_email, email,
                 hizmet_turu, duzenli_fatura, _odemd, odeme_duzeni_manuel_txt or None, aylik_kira, yillik_kira, sozlesme_no, sozlesme_tarihi, sozlesme_bitis,
                 kira_artis_tarihi, kira_suresi_ay, kira_nakit, kira_banka, kira_nakit_tutar, kira_banka_tutar, hazir_oda_val,
                 evrak_imza_sirkuleri, evrak_vergi_levhasi, evrak_ticaret_sicil, evrak_faaliyet_belgesi,
                 evrak_kimlik_fotokopi, evrak_ikametgah, evrak_kase, notlar, tamamlanma_yuzdesi)
            )
            kyc_id = row["id"] if row else None
        if musteri_id:
            try:
                from routes.giris_routes import _ensure_aylik_grid_cache_table, _upsert_aylik_grid_cache

                if _odemd == "manuel":
                    _ensure_aylik_grid_cache_table()
                    execute(
                        "DELETE FROM musteri_aylik_grid_cache WHERE musteri_id = %s",
                        (int(musteri_id),),
                    )
                else:
                    _upsert_aylik_grid_cache(int(musteri_id))
            except Exception:
                pass
        # Opsiyonel: hiyerarşik grup ataması (mevcut update akışına ek alan olarak)
        if musteri_id:
            try:
                raw_parent = data.get("parent_cari_id")
                parent_cari_id = None
                if raw_parent not in (None, "", 0, "0"):
                    parent_cari_id = int(raw_parent)
                CariService.set_parent(int(musteri_id), parent_cari_id)
            except Exception:
                pass
        return jsonify({"ok": True, "kyc_id": kyc_id})
    except Exception as e:
        return jsonify({"ok": False, "mesaj": str(e)}), 400


@bp.route("/api/kyc/sozlesme")
@giris_gerekli
def api_kyc_sozlesme():
    """
    Seçili müşteri için KYC bilgileriyle Word sözleşmesi üret ve indir.
    Aynı müşteride tekrar çağrıldığında aynı sözleşme numarasını (yıl bazlı artan) kullanır.
    """
    musteri_id = request.args.get("musteri_id")
    if not musteri_id:
        flash("Müşteri seçilmedi.", "danger")
        return redirect(url_for("musteriler.giris"))

    kyc = fetch_one(
        "SELECT * FROM musteri_kyc WHERE musteri_id = %s ORDER BY id DESC LIMIT 1",
        (musteri_id,),
    )
    if not kyc:
        flash("Bu müşteri için KYC bilgisi bulunamadı. Önce Kaydet yapın.", "danger")
        return redirect(url_for("musteriler.giris"))

    # Sözleşme tarihi ve numarası
    soz_tarih = kyc.get("sozlesme_tarihi") or date.today()
    if isinstance(soz_tarih, str):
        try:
            soz_tarih = datetime.strptime(soz_tarih[:10], "%Y-%m-%d").date()
        except Exception:
            soz_tarih = date.today()
    yil = soz_tarih.year
    soz_no = (kyc.get("sozlesme_no") or "").strip()

    if not soz_no:
        # O yıla ait son sözleşme numarasını bul → bir artır
        like_pattern = f"SZL-{yil}-%"
        last = fetch_one(
            "SELECT sozlesme_no FROM musteri_kyc WHERE sozlesme_no LIKE %s ORDER BY sozlesme_no DESC LIMIT 1",
            (like_pattern,),
        )
        next_seq = 1
        if last and last.get("sozlesme_no"):
            try:
                parca = str(last["sozlesme_no"]).split("-")[-1]
                next_seq = int(parca) + 1
            except Exception:
                next_seq = 1
        soz_no = f"SZL-{yil}-{next_seq:04d}"
        execute(
            "UPDATE musteri_kyc SET sozlesme_no=%s, sozlesme_tarihi=%s WHERE id=%s",
            (soz_no, soz_tarih, kyc["id"]),
        )

    # Sözleşme metni için kullanılacak bilgiler
    unvan = kyc.get("sirket_unvani") or kyc.get("unvan") or ""
    vergi_no = kyc.get("vergi_no") or ""
    vergi_dairesi = kyc.get("vergi_dairesi") or ""
    mersis_no = kyc.get("mersis_no") or "......................................................."
    ticaret_sicil_no = kyc.get("ticaret_sicil_no") or "......................................................."
    faaliyet_konusu = kyc.get("faaliyet_konusu") or "......................................................."
    merkez_adresi = kyc.get("yeni_adres") or kyc.get("eski_adres") or ""

    yet_ad = kyc.get("yetkili_adsoyad") or ""
    yet_tc = kyc.get("yetkili_tcno") or ""
    yet_dogum = ""
    yd = kyc.get("yetkili_dogum")
    if yd:
        if hasattr(yd, "strftime"):
            yet_dogum = yd.strftime("%d.%m.%Y")
        else:
            try:
                yet_dogum = datetime.strptime(str(yd)[:10], "%Y-%m-%d").strftime("%d.%m.%Y")
            except Exception:
                yet_dogum = str(yd)
    yet_ikamet = kyc.get("yetkili_ikametgah") or "......................................................."
    yet_tel = kyc.get("yetkili_tel") or ""
    yet_tel2 = kyc.get("yetkili_tel2") or ""
    yet_email = kyc.get("yetkili_email") or kyc.get("email") or ""

    hizmet_turu = (kyc.get("hizmet_turu") or "Sanal Ofis").upper()
    aylik = float(kyc.get("aylik_kira") or 0)
    yillik = float(kyc.get("yillik_kira") or 0)
    if not yillik and aylik:
        yillik = aylik * 12

    def tl_fmt(v):
        return f"{v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    doc = Document()

    baslik = doc.add_heading("HAZIR OFİS / SANAL OFİS\nADRES KULLANIM VE HİZMET SÖZLEŞMESİ", level=0)
    baslik.alignment = 1  # center

    doc.add_paragraph(f"Sözleşme No: {soz_no}  |  Tarih: {soz_tarih.strftime('%d.%m.%Y')}")

    doc.add_heading("MADDE 1 – TARAFLAR", level=1)
    doc.add_paragraph(
        "Hizmet Veren:\n"
        "Ofisbir Ofis ve Danışmanlık Hizmetleri A.Ş.\n"
        "Adres: Kavaklıdere Mah. Esat Caddesi No:12 İç Kapı No:1 Çankaya/Ankara\n"
        "Vergi No: 6340871926\n"
        '(İşbu sözleşmede "OFİSBİR" olarak anılacaktır.)\n'
    )
    doc.add_paragraph(
        "HİZMET ALAN (ŞİRKET BİLGİLERİ)\n"
        f"Unvan: {unvan}\n"
        f"Vergi No: {vergi_no}\n"
        f"Vergi Dairesi: {vergi_dairesi}\n"
        f"MERSİS No: {mersis_no}\n"
        f"Ticaret Sicil No: {ticaret_sicil_no}\n"
        f"Faaliyet Konusu: {faaliyet_konusu}\n"
        f"Merkez Adresi: {merkez_adresi}\n\n"
        "YETKİLİ KİŞİ BİLGİLERİ\n"
        f"Ad Soyad: {yet_ad}\n"
        f"T.C. Kimlik No: {yet_tc}\n"
        f"Doğum Tarihi: {yet_dogum or '.......................................................'}\n"
        f"İkamet Adresi: {yet_ikamet}\n"
        f"Cep Telefonu: {yet_tel}\n"
        f"Cep Telefonu 2: {yet_tel2 or '.......................................................'}\n"
        f"E-Posta: {yet_email or '.......................................................'}\n"
        "Islak İmza: .......................................................\n\n"
        "ORTAKLIK BİLGİLERİ\n"
        "Ortak 1 Ad Soyad / Unvan: .......................................................\n"
        "Pay Oranı (%): .......................................................\n\n"
        "Ortak 2 Ad Soyad / Unvan: .......................................................\n"
        "Pay Oranı (%): .......................................................\n\n"
        "Ortak 3 Ad Soyad / Unvan: .......................................................\n"
        "Pay Oranı (%): .......................................................\n\n"
        "Yabancı Ortak Varsa:\n"
        "Ad Soyad: .......................................................\n"
        "Uyruğu: .......................................................\n"
        "Pasaport No: .......................................................\n\n"
        '(İşbu sözleşmede "MÜŞTERİ" olarak anılacaktır.)\n'
        "Müşteri adına imza atan yetkili, şirket ile birlikte müştereken ve müteselsilen sorumludur.\n"
    )

    doc.add_heading("MADDE 2 – HİZMET TÜRÜ", level=1)
    sanal_sec = "☑" if "SANAL" in hizmet_turu else "☐"
    hazir_sec = "☑" if "HAZIR" in hizmet_turu else "☐"
    doc.add_paragraph(
        "Taraflar aşağıdaki hizmet türlerinden birini seçmiştir:\n"
        f"{sanal_sec} SANAL OFİS HİZMETİ\n"
        f"{hazir_sec} HAZIR OFİS HİZMETİ\n"
        "(Seçilen hizmet türü sözleşmenin ayrılmaz parçasıdır.)\n"
    )

    doc.add_heading("BÖLÜM A – SANAL OFİS HİZMETİ", level=1)
    doc.add_heading("MADDE 3A – KAPSAM", level=2)
    doc.add_paragraph(
        "Sanal ofis hizmeti; yasal adres tahsisi, posta/kargo/tebligat teslim alma ve sekreterya "
        "bilgilendirme hizmetlerini kapsar.\n"
        "Bu sözleşme kira sözleşmesi değildir. Taşınmaz üzerinde kiracılık hakkı doğurmaz. Ancak "
        "MÜŞTERİ'ye sözleşme süresince yasal adres kullanım hakkı verir.\n"
        "İşbu sözleşme kapsamında MÜŞTERİ'ye yasal adres kullanım hakkı verilmiş olup, bu adres "
        "vergi mevzuatı çerçevesinde işyeri adresi olarak bildirilebilir.\n"
    )
    doc.add_heading("MADDE 4A – Fiziki Kullanım", level=2)
    doc.add_paragraph(
        "Sanal ofis müşterisi sürekli masa veya oda kullanım hakkına sahip değildir. "
        "Ofise eşya bırakamaz ve ticari mal bulunduramaz.\n"
    )
    doc.add_heading("MADDE 5A – Haciz Güvencesi", level=2)
    doc.add_paragraph(
        "MÜŞTERİ, ofis adresinde kendisine ait mal bulunmadığını, ofisteki tüm demirbaşların "
        "OFİSBİR'e ait olduğunu ve haciz halinde OFİSBİR'in üçüncü kişi olduğunu kabul eder.\n"
        "Haciz bildirgesi gelmesi halinde sözleşme kendiliğinden feshedilir.\n"
    )

    doc.add_heading("ORTAK HÜKÜMLER", level=1)
    doc.add_heading("MADDE 6 – HİZMET BEDELİ", level=2)
    doc.add_paragraph(
        f"Yıllık Hizmet Bedeli: {tl_fmt(yillik)} TL + KDV\n"
        f"Aylık Hizmet Bedeli: {tl_fmt(aylik)} TL + KDV\n\n"
        "Ödemeler aylık olarak OFİSBİR'in bildirdiği banka hesabına yapılacaktır.\n"
        "İki aylık ödeme gecikmesi halinde sözleşme tek taraflı feshedilebilir.\n"
    )

    doc.add_heading("MADDE 7 – ERKEN FESİH", level=2)
    doc.add_paragraph(
        "MÜŞTERİ, sözleşme süresi dolmadan ayrılmak isterse yazılı bildirim yapmak kaydıyla "
        "sözleşmesini feshedebilir.\n"
        "Erken fesih halinde 2 (iki) aylık hizmet bedeli tutarında erken fesih bedeli ödemeyi kabul eder.\n"
        "Bu bedel makul cezai şart niteliğindedir.\n"
    )

    doc.add_heading("MADDE 8 – OTOMATİK YENİLEME", level=2)
    doc.add_paragraph(
        "Sözleşme bitiminden 15 gün önce yazılı fesih yapılmazsa 1 yıl süreyle aynı şartlarla yenilenir.\n"
    )

    doc.add_heading("MADDE 9 – MÜTESELSİL SORUMLULUK", level=2)
    doc.add_paragraph(
        "Şirket yetkilisi işbu sözleşmeden doğan borçlardan şirket ile birlikte müştereken ve "
        "müteselsilen sorumludur.\n"
    )

    doc.add_heading("MADDE 10 – YETKİLİ MAHKEME", level=2)
    doc.add_paragraph(
        "İşbu sözleşmeden doğacak uyuşmazlıklarda Ankara Mahkemeleri ve İcra Daireleri yetkilidir.\n"
    )

    doc.add_heading("MADDE 11 – YÜRÜRLÜK", level=2)
    doc.add_paragraph(
        f"İşbu sözleşme {soz_tarih.strftime('%d.%m.%Y')} tarihinde iki nüsha olarak düzenlenmiş ve "
        "imza altına alınmıştır.\n"
    )

    doc.add_paragraph("\nOFİSBİR\t\tMÜŞTERİ\t\tYetkili (Müteselsil Sorumlu)\n\n\n")
    doc.add_paragraph("İmza:\t\tİmza:\t\tİmza:\n")
    doc.add_paragraph(f"Sözleşme No: {soz_no}  |  {soz_tarih.strftime('%d.%m.%Y')}")

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    out_dir = os.path.join(base_dir, "uploads", "sozlesmeler")
    os.makedirs(out_dir, exist_ok=True)

    safe_unvan = (unvan or "Musteri").replace("/", "-").replace("\\", "-").strip()
    filename = f"Sozlesme_{soz_no}_{safe_unvan[:40].replace(' ', '_')}.docx"
    filepath = os.path.join(out_dir, filename)

    doc.save(filepath)

    return send_file(filepath, as_attachment=True, download_name=filename)


@bp.route("/api/rent_progression")
@giris_gerekli
def api_rent_progression():
    """
    Müşterilerin yıllara göre kira artış verileri
    JSON: {rows: [{id, name, tax_number, ilk_kira_bedeli, current_rent, years: {2023:1234.0}}], years: [2021,2022...]}
    """
    customers = fetch_all(
        "SELECT id, name, tax_number, rent_start_date, rent_start_year, rent_start_month, ilk_kira_bedeli FROM customers ORDER BY name"
    )
    tufe_rows = fetch_all("SELECT year, month, oran FROM tufe_verileri")
    tufe = {}
    for r in tufe_rows:
        try:
            y = int(r.get("year"))
            m = int(r.get("month"))
            tufe[(y,m)] = float(r.get("oran") or 0)
        except Exception:
            continue

    result = []
    years_set = set()
    today = date.today()

    def parse_month(m):
        if m is None:
            return 1
        try:
            return int(m)
        except Exception:
            s = str(m).strip()
            if not s:
                return 1
            for idx, name in enumerate(MONTHS_TR, start=1):
                if name.lower() in s.lower():
                    return idx
            return 1

    def parse_year(y, rent_start_date):
        if y:
            try:
                return int(y)
            except Exception:
                pass
        if rent_start_date:
            s = str(rent_start_date)
            parts = s.split(".")
            if len(parts) >= 3:
                try:
                    return int(parts[2])
                except Exception:
                    pass
        return today.year

    for c in customers:
        initial = float(c.get("ilk_kira_bedeli") or 0)
        start_year = parse_year(c.get("rent_start_year"), c.get("rent_start_date"))
        start_month = parse_month(c.get("rent_start_month"))

        rent = initial
        years_dict = {}
        for y in range(start_year, today.year + 1):
            if y == start_year:
                years_dict[y] = round(rent,2)
                years_set.add(y)
                continue
            rate = tufe.get((y, start_month))
            if rate:
                rent = rent * (1.0 + float(rate)/100.0)
            years_dict[y] = round(rent,2)
            years_set.add(y)

        current_rent = round(rent,2)

        result.append({
            "id": c.get("id"),
            "name": c.get("name"),
            "tax_number": c.get("tax_number"),
            "ilk_kira_bedeli": initial,
            "current_rent": current_rent,
            "years": years_dict
        })

    return jsonify({"rows": result, "years": sorted(list(years_set))})