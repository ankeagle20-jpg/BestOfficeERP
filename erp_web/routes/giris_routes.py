"""
Giriş / Müşteri Kaydı Routes
Desktop'taki gibi tam fonksiyonel + Sözleşme oluşturma
"""
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, Response, current_app
from flask_login import current_user
from auth import giris_gerekli
from collections import defaultdict
from cache_utils import (
    CACHE_KEY_DUZENLI_FATURA,
    CACHE_KEY_HIZMET_TURLERI,
    CACHE_KEY_TUFE_VERILERI,
    CACHE_TTL_SEC,
    simple_cache_get,
    simple_cache_invalidate,
    simple_cache_set,
)
from db import (
    fetch_all,
    fetch_one,
    execute,
    execute_returning,
    ensure_hizmet_turleri_table,
    ensure_duzenli_fatura_secenekleri_table,
    ensure_faturalar_amount_columns,
    ensure_contracts_engine,
    ensure_customer_financial_profile,
    ensure_customers_durum,
    ensure_customers_is_active,
    ensure_customers_kapanis_sonrasi_borc_ay,
    ensure_customers_musteri_no,
    ensure_customers_hazir_ofis_oda,
    ensure_musteri_kyc_columns,
    ensure_musteri_kyc_hazir_ofis_oda_no,
    ensure_musteri_kyc_kira_banka,
    ensure_customers_bizim_hesap,
    ensure_customers_grup2_secimleri,
    ensure_grup2_etiketleri_table,
    ensure_grup2_bizim_hesap_into_array,
    ensure_cari_kart_perf_indexes,
    db as get_db,
    get_conn,
    sql_expr_fatura_not_gib_taslak,
)
from datetime import datetime, date, timedelta
import calendar
import time
import threading
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import io
import base64
import re
import math
import logging
import urllib.parse
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from utils.text_utils import turkish_lower
from utils.musteri_arama import (
    customers_arama_params_giris_genis,
    customers_arama_sql_giris_genis,
    customers_arama_sql_params_giris_genis_tokens,
    musteri_arama_ilike_pattern_email_duz,
)
import json
from pathlib import Path
import psycopg2
import secrets
from decimal import Decimal

# Aylık grid «tam ödendi» / tahsil dağıtım mantığı değişince artırın; musteri_aylik_grid_cache yeniden üretilir.
AYLIK_GRID_COMPUTE_REV = 23
AYLIK_GRID_TAM_ODENDI_TOLERANS = 0.05  # kurus farklarini (dagitim/yuvarlama) tam odendi say


# BTUFRT (|BTUFRT|TÜFE borç+tahsil toplu): bu tarihte ve sonrasına ilişkin aylar grid «ödenen» toplamına
# dahil EDİLMEZ (gerçek tahsil ile karışmasın); öncesi için muhasebe çifti yeşilde kalır — tüm zamanları
# dışarı atınca 2023 vb. hep kırmızı oluşmuştu.
BTUFRT_GRID_ODENME_DAHIL_DEGIL = date(2026, 1, 1)


def _aylik_btufrt_row_skip_grid_odeme(ac: str, marker_isos: list[str], r: dict) -> bool:
    """True ise bu tahsil satırı aylık grid ödenen dağılımına hiç yazılmaz."""
    if "|BTUFRT|" not in (ac or ""):
        return False
    esik = BTUFRT_GRID_ODENME_DAHIL_DEGIL
    if marker_isos:
        found_valid = False
        all_ge_esik = True
        for iso_raw in marker_isos:
            try:
                dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                found_valid = True
                if date(dd.year, dd.month, 1) < esik:
                    all_ge_esik = False
            except ValueError:
                continue
        # Satırı ancak işaretçilerin TAMAMI eşik ve sonrasıysa dışla.
        # Karışık satırda (örn. 2023 + 2026) eski ayların ödemesi gridde sayılmalı.
        if found_valid:
            return all_ge_esik
        return False
    d = r.get("fatura_tarihi") or r.get("tahsilat_tarihi")
    if hasattr(d, "year"):
        return date(int(d.year), int(d.month), 1) >= esik
    if d:
        try:
            dd = datetime.strptime(str(d)[:10], "%Y-%m-%d").date()
            return date(dd.year, dd.month, 1) >= esik
        except (ValueError, TypeError):
            pass
    # İşaretçi/tarih çıkmadıysa satırı silme — fatura/tahsilat ayıyla dağıtıma düşsün (2023 vb. tamamen kırmızı kalmasın).
    return False


def _tahsil_row_ekstre_eslesme_ay_iso(r: dict) -> str | None:
    """Cari ekstre tahsilat sorgusu ile aynı öncelik: fatura_tarihi → |AYLIK_TAH| → tahsilat_tarihi (ayın 1'i ISO)."""
    try:
        d = (r or {}).get("fatura_tarihi")
        if d and hasattr(d, "year"):
            return date(int(d.year), int(d.month), 1).isoformat()
        ac = str((r or {}).get("aciklama") or "")
        mm = re.search(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac)
        if mm:
            dd = datetime.strptime(mm.group(1)[:10], "%Y-%m-%d").date()
            return date(dd.year, dd.month, 1).isoformat()
        d2 = (r or {}).get("tahsilat_tarihi")
        if d2 and hasattr(d2, "year"):
            return date(int(d2.year), int(d2.month), 1).isoformat()
        if d2:
            dd = datetime.strptime(str(d2)[:10], "%Y-%m-%d").date()
            return date(dd.year, dd.month, 1).isoformat()
    except Exception:
        return None
    return None


def _register_arial():
    """Türkçe karakter için Arial veya alternatif font kaydet."""
    if getattr(_register_arial, "_done", False):
        return
    candidates = []
    win = os.environ.get("WINDIR") or os.environ.get("SystemRoot") or "C:\\Windows"
    for f in ("arial.ttf", "Arial.ttf", "ARIAL.TTF"):
        candidates.append(os.path.join(win, "Fonts", f))
    candidates.extend([
        "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ])
    for path in candidates:
        if path and os.path.isfile(path):
            try:
                pdfmetrics.registerFont(TTFont("Arial", path))
                _register_arial._done = True
                return
            except Exception:
                pass
    _register_arial._done = True

bp = Blueprint('giris', __name__)

@bp.route('/api/tahsilat-personeller')
@giris_gerekli
def api_tahsilat_personeller():
    rows = fetch_all(
        """
        SELECT id, ad_soyad
        FROM personel
        WHERE COALESCE(is_active, TRUE) = TRUE
        ORDER BY ad_soyad
        """
    ) or []
    out = []
    for r in rows:
        ad = str((r or {}).get("ad_soyad") or "").strip()
        if not ad:
            continue
        out.append({"id": r.get("id"), "ad_soyad": ad})
    return jsonify(out)


# Dosya yükleme ayarları
UPLOAD_FOLDER = 'uploads/musteri_dosyalari'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'pdf', 'docx'}


# Process içinde tek seferlik DDL koruması: ensure_*_table fonksiyonları her HTTP
# isteğinde tetiklenip ~150-300ms ek gecikme yaratıyordu (Supabase round-trip).
# İlk çağrıda tabloyu garanti altına alıp sonraki çağrılarda no-op'a düşüyoruz.
_AYLIK_GRID_CACHE_TABLE_READY = False
_REEL_DONEM_TUTAR_TABLE_READY = False
_TAHSILAT_PANEL_DETAY_TABLE_READY = False

# Süreç-içi grup cache'i: parent_cari_id çözümlemesi için her tıklamada 500 satır
# çekiyorduk (~190ms). Grup listesi sık değişmiyor; 60 saniye TTL ile cacheliyoruz.
_GIRIS_GRUP_HARITA_CACHE = {"ts": 0.0, "data": {}}
_GIRIS_GRUP_HARITA_TTL = 60.0  # saniye


def _giris_grup_uuid_id_haritasi():
    """parent_id (uuid) → {id, name} haritası. Süreç içinde 60 sn TTL ile tutulur."""
    import time as _time
    now = _time.monotonic()
    if (now - _GIRIS_GRUP_HARITA_CACHE.get("ts", 0.0)) < _GIRIS_GRUP_HARITA_TTL:
        d = _GIRIS_GRUP_HARITA_CACHE.get("data")
        if isinstance(d, dict):
            return d
    try:
        from services.cari_service import CariService
        groups = fetch_all(
            "SELECT id, name FROM customers WHERE COALESCE(is_group, FALSE)=TRUE LIMIT 500"
        ) or []
        m = {}
        for g in groups:
            try:
                gid = int(g.get("id") or 0)
            except (TypeError, ValueError):
                continue
            if gid <= 0:
                continue
            try:
                key = str(CariService.customer_uuid(gid))
            except Exception:
                continue
            m[key] = {"id": gid, "name": (g.get("name") or "").strip()}
        _GIRIS_GRUP_HARITA_CACHE["data"] = m
        _GIRIS_GRUP_HARITA_CACHE["ts"] = now
        return m
    except Exception:
        return _GIRIS_GRUP_HARITA_CACHE.get("data") or {}


def _ensure_aylik_grid_cache_table():
    global _AYLIK_GRID_CACHE_TABLE_READY
    if _AYLIK_GRID_CACHE_TABLE_READY:
        return
    try:
        execute(
            """
            CREATE TABLE IF NOT EXISTS musteri_aylik_grid_cache (
                musteri_id INTEGER PRIMARY KEY REFERENCES customers(id) ON DELETE CASCADE,
                payload TEXT NOT NULL,
                updated_at TIMESTAMP NOT NULL DEFAULT NOW()
            )
            """
        )
    except Exception:
        pass
    _AYLIK_GRID_CACHE_TABLE_READY = True


def _ensure_musteri_reel_donem_tutar_table():
    global _REEL_DONEM_TUTAR_TABLE_READY
    if _REEL_DONEM_TUTAR_TABLE_READY:
        return
    try:
        execute(
            """
            CREATE TABLE IF NOT EXISTS musteri_reel_donem_tutar (
                musteri_id INTEGER NOT NULL REFERENCES customers(id) ON DELETE CASCADE,
                donem_yil INTEGER NOT NULL,
                tutar_kdv_dahil NUMERIC(14, 2) NOT NULL,
                giris_tip TEXT,
                giris_tutar NUMERIC(14, 2),
                hibrit_toplam NUMERIC(14, 2),
                hibrit_net NUMERIC(14, 2),
                hibrit_banka NUMERIC(14, 2),
                updated_at TIMESTAMP NOT NULL DEFAULT NOW(),
                PRIMARY KEY (musteri_id, donem_yil)
            )
            """
        )
        for col_sql in (
            "ALTER TABLE musteri_reel_donem_tutar ADD COLUMN IF NOT EXISTS giris_tip TEXT",
            "ALTER TABLE musteri_reel_donem_tutar ADD COLUMN IF NOT EXISTS giris_tutar NUMERIC(14, 2)",
            "ALTER TABLE musteri_reel_donem_tutar ADD COLUMN IF NOT EXISTS hibrit_toplam NUMERIC(14, 2)",
            "ALTER TABLE musteri_reel_donem_tutar ADD COLUMN IF NOT EXISTS hibrit_net NUMERIC(14, 2)",
            "ALTER TABLE musteri_reel_donem_tutar ADD COLUMN IF NOT EXISTS hibrit_banka NUMERIC(14, 2)",
        ):
            try:
                execute(col_sql)
            except Exception:
                pass
    except Exception:
        pass
    _REEL_DONEM_TUTAR_TABLE_READY = True


def _ensure_tahsilat_panel_detay_table():
    global _TAHSILAT_PANEL_DETAY_TABLE_READY
    if _TAHSILAT_PANEL_DETAY_TABLE_READY:
        return
    try:
        execute(
            """
            CREATE TABLE IF NOT EXISTS musteri_tahsilat_panel_detay (
                musteri_id INTEGER PRIMARY KEY REFERENCES customers(id) ON DELETE CASCADE,
                by_iso TEXT NOT NULL DEFAULT '{}',
                updated_at TIMESTAMP NOT NULL DEFAULT NOW()
            )
            """
        )
    except Exception:
        pass
    _TAHSILAT_PANEL_DETAY_TABLE_READY = True


def _load_musteri_panel_by_iso(musteri_id: int) -> dict:
    """DB panel kaynağı: {YYYY-MM-01: {aylik, tahsil, kalan, tahsil_tarih?}}."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return {}
    if mid <= 0:
        return {}
    _ensure_tahsilat_panel_detay_table()
    row = fetch_one(
        "SELECT by_iso FROM musteri_tahsilat_panel_detay WHERE musteri_id = %s",
        (mid,),
    )
    if not row or not row.get("by_iso"):
        return {}
    try:
        raw = row.get("by_iso")
        obj = json.loads(raw) if isinstance(raw, str) else raw
    except (json.JSONDecodeError, TypeError, ValueError):
        return {}
    if not isinstance(obj, dict):
        return {}
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    out = {}
    for iso_raw, prow in obj.items():
        if not isinstance(prow, dict):
            continue
        try:
            parts = str(iso_raw).strip()[:10].split("-")
            iso_key = date(int(parts[0]), int(parts[1]), 1).isoformat()
        except (ValueError, IndexError, TypeError):
            continue
        try:
            pt = round(float(prow.get("tahsil") or 0), 2)
            pk = round(max(float(prow.get("kalan") or 0), 0), 2)
            pb = round(float(prow.get("aylik") or prow.get("brut") or 0), 2)
        except (TypeError, ValueError):
            continue
        if pb > tol and pt > pb + tol:
            pt = pb
            pk = 0.0
        elif pb > tol and pk <= tol and pt + tol < pb:
            pk = round(max(pb - pt, 0), 2)
        out[iso_key] = {
            "aylik": pb,
            "tahsil": pt,
            "kalan": pk,
            "tahsil_tarih": str(prow.get("tahsil_tarih") or "")[:10] or None,
        }
    return out


def _save_musteri_panel_by_iso(musteri_id: int, by_iso: dict, prune_no_db_tahsil: bool = False) -> None:
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return
    if mid <= 0:
        return
    _ensure_tahsilat_panel_detay_table()
    ser = {}
    for iso_k, prow in (by_iso or {}).items():
        if not isinstance(prow, dict):
            continue
        try:
            parts = str(iso_k).strip()[:10].split("-")
            iso_key = date(int(parts[0]), int(parts[1]), 1).isoformat()
        except (ValueError, IndexError, TypeError):
            continue
        try:
            pt = round(float(prow.get("tahsil") or 0), 2)
            pk = round(max(float(prow.get("kalan") or 0), 0), 2)
            pb = round(float(prow.get("aylik") or prow.get("brut") or 0), 2)
        except (TypeError, ValueError):
            continue
        if pt <= 0.004 and pk <= 0.004:
            continue
        ser[iso_key] = {
            "aylik": pb,
            "tahsil": pt,
            "kalan": pk,
            "tahsil_tarih": str(prow.get("tahsil_tarih") or "")[:10] or None,
        }
    merged = dict(_load_musteri_panel_by_iso(mid))
    if prune_no_db_tahsil:
        tol_pr = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
        tahsil_map_pr = _aylik_tahsil_tutar_map(mid)
        for iso_pr in list(merged.keys()):
            if iso_pr in ser:
                continue
            try:
                db_t = round(float(tahsil_map_pr.get(iso_pr) or 0), 2)
            except (TypeError, ValueError):
                db_t = 0.0
            if db_t <= tol_pr:
                merged.pop(iso_pr, None)
    merged.update(ser)
    execute(
        """
        INSERT INTO musteri_tahsilat_panel_detay (musteri_id, by_iso, updated_at)
        VALUES (%s, %s, NOW())
        ON CONFLICT (musteri_id)
        DO UPDATE SET by_iso = EXCLUDED.by_iso, updated_at = NOW()
        """,
        (mid, json.dumps(merged, ensure_ascii=False)),
    )


def _panel_by_iso_from_tahsil_map(
    musteri_id: int,
    payload: dict | None,
    tahsilat_tarihi: str | None = None,
    trust_grid_odenen: bool = True,
) -> dict:
    """|AYLIK_PAY| tahsil haritası + grid brüt → panel by_iso (popup ile aynı kaynak)."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return {}
    if mid <= 0:
        return {}
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    tahsil_map = _aylik_tahsil_tutar_map(mid)
    existing = _load_musteri_panel_by_iso(mid)
    tarih_s = str(tahsilat_tarihi or "")[:10] if tahsilat_tarihi else ""
    by_iso = {}
    if not isinstance(payload, dict):
        return existing
    dagitim_isos = set()
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict):
            continue
        try:
            iso_m = date(int(a.get("yil")), int(a.get("ay")), 1).isoformat()
        except (TypeError, ValueError):
            continue
        brut = round(float(a.get("brut_tutar_kdv") or a.get("tutar_kdv_dahil") or 0), 2)
        tah = round(float(tahsil_map.get(iso_m) or 0), 2)
        try:
            odenen_g = round(float(a.get("odenen_tutar_kdv") or 0), 2)
        except (TypeError, ValueError):
            odenen_g = 0.0
        if trust_grid_odenen and tah <= tol and odenen_g > tol:
            tah = odenen_g
        if brut <= tol and tah <= tol:
            continue
        try:
            kalan_g = round(max(float(a.get("kalan_tutar_kdv") or 0), 0), 2)
        except (TypeError, ValueError):
            kalan_g = -1.0
        kalan_calc = round(max(brut - tah, 0), 2) if brut > tol else 0.0
        if brut > tol and tah >= brut - tol:
            kalan = 0.0
        else:
            kismi_g = bool(a.get("kismi_tahsilat")) or (tah > tol and kalan_g > tol)
            if kismi_g and kalan_g >= 0:
                kalan = kalan_g
            else:
                kalan = kalan_calc
        prev = existing.get(iso_m) or {}
        th = tarih_s or prev.get("tahsil_tarih") or ""
        by_iso[iso_m] = {
            "aylik": brut,
            "tahsil": tah,
            "kalan": kalan,
            "tahsil_tarih": th,
        }
        dagitim_isos.add(iso_m)
    if tarih_s:
        for iso_m in dagitim_isos:
            if iso_m in by_iso:
                by_iso[iso_m]["tahsil_tarih"] = tarih_s
    return by_iso


def _ekstre_panel_filter_db_tahsil(musteri_id: int, panel_by_iso: dict | None) -> dict:
    """Ekstre: panelde kalan sahte tahsil (DB yok) grid/ekstre satırına yansımasın."""
    if not panel_by_iso:
        return {}
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return dict(panel_by_iso)
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    tahsil_map = _aylik_tahsil_tutar_map(mid)
    out = {}
    for iso_k, prow in panel_by_iso.items():
        if not isinstance(prow, dict):
            continue
        try:
            pt = round(float(prow.get("tahsil") or 0), 2)
            db_t = round(float(tahsil_map.get(iso_k) or 0), 2)
        except (TypeError, ValueError):
            continue
        if db_t > tol or pt <= tol:
            out[iso_k] = prow
    return out


def _iso_from_aylik_tah_marker(text: str) -> str | None:
    m = re.search(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", str(text or ""))
    if not m:
        return None
    try:
        dd = datetime.strptime(m.group(1)[:10], "%Y-%m-%d").date()
        return date(dd.year, dd.month, 1).isoformat()
    except ValueError:
        return None


def _panel_by_iso_from_grid_payload(payload: dict | None) -> dict:
    """Grid önbellek kısmi/tam tahsil → panel by_iso (yenilemede kısmi ay kaybolmasın)."""
    if not isinstance(payload, dict):
        return {}
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    out = {}
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict):
            continue
        try:
            iso_m = date(int(a.get("yil")), int(a.get("ay")), 1).isoformat()
        except (TypeError, ValueError):
            continue
        try:
            brut = round(float(a.get("brut_tutar_kdv") or a.get("tutar_kdv_dahil") or 0), 2)
            odenen = round(float(a.get("odenen_tutar_kdv") or 0), 2)
            kalan = round(max(float(a.get("kalan_tutar_kdv") or 0), 0), 2)
        except (TypeError, ValueError):
            continue
        if brut <= tol:
            continue
        kismi = bool(a.get("kismi_tahsilat")) or (odenen > tol and kalan > tol)
        tam = bool(a.get("tahsil_edildi")) or (kalan <= tol and odenen > tol)
        if not kismi and not tam and odenen <= tol:
            continue
        if tam and odenen <= tol:
            odenen = brut
            kalan = 0.0
        out[iso_m] = {
            "aylik": brut,
            "tahsil": odenen,
            "kalan": kalan,
            "tahsil_tarih": str(a.get("tahsil_tarih") or "")[:10] or None,
        }
    return out


def _panel_by_iso_pick_richer(old: dict | None, new: dict | None, tol: float) -> dict:
    """İki panel satırından tahsil/kalan açısından daha bilgilendirici olanı seç."""
    o = old if isinstance(old, dict) else {}
    n = new if isinstance(new, dict) else {}
    ot = round(float(o.get("tahsil") or 0), 2)
    nt = round(float(n.get("tahsil") or 0), 2)
    ok = round(max(float(o.get("kalan") or 0), 0), 2)
    nk = round(max(float(n.get("kalan") or 0), 0), 2)
    ob = round(float(o.get("aylik") or 0), 2)
    nb = round(float(n.get("aylik") or 0), 2)
    if nt > ot + tol:
        return n
    if nt + tol < ot:
        return o
    if nk + tol < ok and nt >= ot - tol:
        return n
    if ok + tol < nk and ot >= nt - tol:
        return o
    if nb > ob + tol:
        return n
    return n if nt >= ot else o


def sync_musteri_panel_from_tahsil_and_dagitim(
    musteri_id: int,
    dagitim_items: list | None = None,
    tahsilat_tarihi: str | None = None,
    payload: dict | None = None,
) -> dict:
    """
    Tahsilat makbuzu sonrası: |AYLIK_PAY| haritası + popup dağıtım → DB panel + grid önbelleği.
    Dönüş: by_iso sözlüğü.
    """
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return {}
    if mid <= 0:
        return {}
    if payload is None:
        payload = _read_aylik_grid_cache_payload(mid)
    if not isinstance(payload, dict):
        try:
            payload = _build_aylik_grid_cache_payload(mid, tufe_map=_tufe_map_by_year_month_cached())
        except Exception:
            payload = None
    by_iso = _panel_by_iso_from_tahsil_map(mid, payload, tahsilat_tarihi)
    if dagitim_items:
        tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
        for it in dagitim_items:
            if not isinstance(it, dict):
                continue
            iso_raw = str(it.get("iso") or "").strip()[:10]
            if not re.match(r"^\d{4}-\d{2}-\d{2}$", iso_raw):
                continue
            try:
                dd = datetime.strptime(iso_raw, "%Y-%m-%d").date()
                iso_key = date(dd.year, dd.month, 1).isoformat()
                yy_d, mm_d = int(dd.year), int(dd.month)
            except ValueError:
                continue
            try:
                pay = round(float(it.get("tutar") or 0), 2)
            except (TypeError, ValueError):
                pay = 0.0
            if pay <= tol:
                continue
            prow = by_iso.get(iso_key) or {}
            brut = round(float(prow.get("aylik") or 0), 2)
            if brut <= tol and isinstance(payload, dict):
                brut = _month_brut_from_grid_payload(payload, yy_d, mm_d)
            if brut <= tol:
                brut = round(float(prow.get("tahsil") or 0), 2)
            if brut <= tol:
                brut = pay
            tah = round(float(prow.get("tahsil") or pay), 2)
            if brut > tol and tah > brut:
                tah = brut
            kalan = round(max(brut - tah, 0), 2) if brut > tol else 0.0
            by_iso[iso_key] = {
                "aylik": brut,
                "tahsil": tah,
                "kalan": kalan,
                "tahsil_tarih": str(tahsilat_tarihi or "")[:10] or prow.get("tahsil_tarih"),
            }
    _save_musteri_panel_by_iso(mid, by_iso, prune_no_db_tahsil=True)
    if isinstance(payload, dict):
        _apply_panel_by_iso_to_grid_payload(payload, by_iso)
        _ensure_aylik_grid_cache_table()
        execute(
            """
            INSERT INTO musteri_aylik_grid_cache (musteri_id, payload, updated_at)
            VALUES (%s, %s, NOW())
            ON CONFLICT (musteri_id)
            DO UPDATE SET payload = EXCLUDED.payload, updated_at = NOW()
            """,
            (mid, json.dumps(payload, ensure_ascii=False)),
        )
        try:
            _aylik_grid_payload_mem[mid] = (time.time(), payload)
        except (TypeError, ValueError):
            pass
        _invalidate_aylik_grid_payload_mem(mid)
    return by_iso


def apply_makbuz_dagitim_to_panel_db(
    musteri_id: int,
    dagitim_items: list | None,
    tahsilat_tarihi: str | None = None,
    payload: dict | None = None,
) -> dict:
    """Makbuz popup dağıtım tutarlarını panel DB'ye kalıcı yazar (ekstre ile uyumlu)."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return {}
    if mid <= 0 or not dagitim_items:
        return {}
    if payload is None:
        payload = _read_aylik_grid_cache_payload(mid)
    if not isinstance(payload, dict):
        try:
            payload = _build_aylik_grid_cache_payload(
                mid, tufe_map=_tufe_map_by_year_month_cached()
            )
        except Exception:
            payload = None
    by_iso = dict(_load_musteri_panel_by_iso(mid))
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    tarih_s = str(tahsilat_tarihi or "")[:10] if tahsilat_tarihi else ""
    for it in dagitim_items:
        if not isinstance(it, dict):
            continue
        iso_raw = str(it.get("iso") or "").strip()[:10]
        if not re.match(r"^\d{4}-\d{2}-\d{2}$", iso_raw):
            continue
        try:
            dd = datetime.strptime(iso_raw, "%Y-%m-%d").date()
            iso_key = date(dd.year, dd.month, 1).isoformat()
            yy_d, mm_d = int(dd.year), int(dd.month)
        except ValueError:
            continue
        try:
            pay = round(float(it.get("tutar") or 0), 2)
        except (TypeError, ValueError):
            pay = 0.0
        if pay <= tol:
            continue
        brut = round(float((by_iso.get(iso_key) or {}).get("aylik") or 0), 2)
        if brut <= tol and isinstance(payload, dict):
            brut = _month_brut_from_grid_payload(payload, yy_d, mm_d)
        if brut <= tol:
            brut = pay
        tah = round(min(pay, brut), 2) if brut > tol else pay
        kalan = round(max(brut - tah, 0), 2) if brut > tol else 0.0
        prev = by_iso.get(iso_key) or {}
        by_iso[iso_key] = {
            "aylik": brut,
            "tahsil": tah,
            "kalan": kalan,
            "tahsil_tarih": tarih_s or prev.get("tahsil_tarih") or "",
        }
    if by_iso:
        _save_musteri_panel_by_iso(mid, by_iso)
        if isinstance(payload, dict):
            _apply_panel_by_iso_to_grid_payload(payload, by_iso)
            _ensure_aylik_grid_cache_table()
            execute(
                """
                INSERT INTO musteri_aylik_grid_cache (musteri_id, payload, updated_at)
                VALUES (%s, %s, NOW())
                ON CONFLICT (musteri_id)
                DO UPDATE SET payload = EXCLUDED.payload, updated_at = NOW()
                """,
                (mid, json.dumps(payload, ensure_ascii=False)),
            )
            _invalidate_aylik_grid_payload_mem(mid)
    return by_iso


def _apply_panel_by_iso_to_grid_payload(payload: dict, panel_by_iso: dict) -> None:
    """Panel DB kaynağı grid önbellek aylarına yazar (yenilemede turuncu/yeşil kalır)."""
    if not isinstance(payload, dict) or not panel_by_iso:
        return
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict):
            continue
        try:
            iso_m = date(int(a.get("yil")), int(a.get("ay")), 1).isoformat()
        except (TypeError, ValueError):
            continue
        prow = panel_by_iso.get(iso_m)
        if not prow:
            continue
        try:
            brut = round(float(prow.get("aylik") or a.get("brut_tutar_kdv") or 0), 2)
            tah = round(float(prow.get("tahsil") or 0), 2)
            kalan = round(max(float(prow.get("kalan") or 0), 0), 2)
        except (TypeError, ValueError):
            continue
        if brut > tol and tah >= brut - tol:
            kalan = 0.0
        if brut > tol:
            a["brut_tutar_kdv"] = brut
            a["tutar_kdv_dahil"] = brut
        a["odenen_tutar_kdv"] = tah
        a["kalan_tutar_kdv"] = kalan
        a["tahsil_edildi"] = kalan <= tol
        a["kismi_tahsilat"] = tah > tol and kalan > tol
        if kalan > tol and tah <= tol:
            a["acik_aylik_borc_faturasi"] = True
        elif tah >= brut - tol and brut > tol:
            a["acik_aylik_borc_faturasi"] = False


def _month_brut_from_grid_payload(payload: dict | None, yil: int, ay: int) -> float:
    if not isinstance(payload, dict):
        return 0.0
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict):
            continue
        try:
            if int(a.get("yil")) != int(yil) or int(a.get("ay")) != int(ay):
                continue
        except (TypeError, ValueError):
            continue
        try:
            brut = float(a.get("brut_tutar_kdv") or a.get("tutar_kdv_dahil") or 0)
        except (TypeError, ValueError):
            brut = 0.0
        if brut > tol:
            return round(brut, 2)
    return 0.0


def sync_musteri_panel_borclu_from_satirlar(musteri_id: int, satirlar: list | None) -> dict:
    """
    Borçlandır / tahsilden çıkar sonrası panel DB: tahsil=0, kalan=aylık brüt.
    Grid ve ekstre bu kayıttan güncellenir.
    """
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return {}
    if mid <= 0 or not satirlar:
        return _load_musteri_panel_by_iso(mid)
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    by_iso = _load_musteri_panel_by_iso(mid)
    payload_hint = _read_aylik_grid_cache_payload(mid)
    if payload_hint is None:
        try:
            payload_hint = _build_aylik_grid_cache_payload(
                mid, tufe_map=_tufe_map_by_year_month_cached()
            )
        except Exception:
            payload_hint = None
    for raw in satirlar:
        if not isinstance(raw, dict):
            continue
        try:
            yil = int(raw.get("yil"))
            ay = int(raw.get("ay"))
        except (TypeError, ValueError):
            continue
        if ay < 1 or ay > 12:
            continue
        iso_key = date(yil, ay, 1).isoformat()
        try:
            brut = round(float(raw.get("tutar_kdv_dahil")), 2)
        except (TypeError, ValueError):
            brut = 0.0
        if brut <= tol:
            brut = _month_brut_from_grid_payload(payload_hint, yil, ay)
        if brut <= tol:
            continue
        by_iso[iso_key] = {
            "aylik": brut,
            "tahsil": 0.0,
            "kalan": brut,
            "tahsil_tarih": None,
        }
    if by_iso:
        _save_musteri_panel_by_iso(mid, by_iso)
    return by_iso


def _persist_grid_cache_with_panel(musteri_id: int, payload: dict | None = None) -> dict | None:
    """Önbelleği panel DB ile hizala ve diske yaz."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return payload
    if payload is None:
        payload = _read_aylik_grid_cache_payload(mid)
    if payload is None:
        return _upsert_aylik_grid_cache(mid)
    payload = _aylik_grid_cache_payload_tahsil_guncelle(mid, payload)
    _ensure_aylik_grid_cache_table()
    execute(
        """
        INSERT INTO musteri_aylik_grid_cache (musteri_id, payload, updated_at)
        VALUES (%s, %s, NOW())
        ON CONFLICT (musteri_id)
        DO UPDATE SET payload = EXCLUDED.payload, updated_at = NOW()
        """,
        (mid, json.dumps(payload, ensure_ascii=False)),
    )
    try:
        _aylik_grid_payload_mem[mid] = (time.time(), payload)
    except (TypeError, ValueError):
        pass
    _invalidate_aylik_grid_payload_mem(mid)
    return payload


def _kyc_date_iso(v):
    if v is None:
        return None
    if hasattr(v, "isoformat"):
        return str(v.isoformat())[:10]
    s = str(v).strip()
    return s[:10] if len(s) >= 10 else (s or None)


def _aylik_grid_cache_matches_kyc(musteri_id, cache_obj):
    if not isinstance(cache_obj, dict):
        return False
    kyc = fetch_one(
        """
        SELECT mk.sozlesme_tarihi, mk.sozlesme_bitis, mk.kira_suresi_ay, mk.aylik_kira, mk.kira_nakit, mk.kira_artis_tarihi,
               mk.kira_nakit_tutar, mk.kira_banka_tutar,
               c.kapanis_tarihi, c.kapanis_sonrasi_borc_ay, c.durum
        FROM customers c
        LEFT JOIN LATERAL (
            SELECT sozlesme_tarihi, sozlesme_bitis, kira_suresi_ay, aylik_kira, kira_nakit, kira_artis_tarihi,
                   kira_nakit_tutar, kira_banka_tutar
            FROM musteri_kyc
            WHERE musteri_id = c.id
            ORDER BY id DESC
            LIMIT 1
        ) mk ON TRUE
        WHERE c.id = %s
        """,
        (musteri_id,),
    )
    if not kyc:
        return False

    # Eski payload'lar: taban kirası yok → Sözleşmeler güncellense bile rapor 0 kalabiliyordu
    if "taban_aylik_net" not in cache_obj:
        return False

    taban_db = _aylik_grid_coerce_money(kyc.get("aylik_kira"))
    taban_ca = _aylik_grid_coerce_money(cache_obj.get("taban_aylik_net"))
    if taban_db != taban_ca:
        return False

    if bool(kyc.get("kira_nakit")) != bool(cache_obj.get("kira_nakit")):
        return False

    if _aylik_grid_coerce_money(kyc.get("kira_nakit_tutar")) != _aylik_grid_coerce_money(
        cache_obj.get("kira_nakit_tutar")
    ):
        return False
    if _aylik_grid_coerce_money(kyc.get("kira_banka_tutar")) != _aylik_grid_coerce_money(
        cache_obj.get("kira_banka_tutar")
    ):
        return False

    bas_k = _aylik_grid_coerce_date(kyc.get("sozlesme_tarihi"))
    if not bas_k:
        return False
    if cache_obj.get("baslangic") != bas_k.isoformat():
        return False

    d_bit = _aylik_grid_coerce_date(kyc.get("sozlesme_bitis"))
    if d_bit is not None:
        bit_eff = _aylik_grid_effective_bitis(kyc, d_bit) or d_bit
        if cache_obj.get("bitis") != bit_eff.isoformat():
            return False
    # KYC'de bitiş yok: önbellek sentetik bitiş tutar; bitis alanını zorlamayız.

    kap_db = _aylik_grid_coerce_date(kyc.get("kapanis_tarihi"))
    kap_cache = _aylik_grid_coerce_date(cache_obj.get("kapanis_tarihi"))
    if _kyc_date_iso(kap_db) != _kyc_date_iso(kap_cache):
        return False
    try:
        ek_db = int(kyc.get("kapanis_sonrasi_borc_ay") or 0)
    except (TypeError, ValueError):
        ek_db = 0
    try:
        ek_cache = int(cache_obj.get("kapanis_sonrasi_borc_ay") or 0)
    except (TypeError, ValueError):
        ek_cache = 0
    if ek_db != ek_cache:
        return False

    artis_b = _aylik_grid_coerce_date(kyc.get("kira_artis_tarihi")) or bas_k
    try:
        artis_ay_db = int(artis_b.month)
    except Exception:
        artis_ay_db = int(bas_k.month)
    try:
        artis_ay_ca = int(cache_obj.get("artis_ay"))
    except (TypeError, ValueError):
        return False
    if artis_ay_db != artis_ay_ca:
        return False

    def _int_or_none(x):
        if x is None or x == "":
            return None
        try:
            return int(x)
        except (TypeError, ValueError):
            return None

    if _int_or_none(cache_obj.get("kira_suresi_ay")) != _int_or_none(kyc.get("kira_suresi_ay")):
        return False
    try:
        if int(cache_obj.get("compute_rev") or 0) != AYLIK_GRID_COMPUTE_REV:
            return False
    except (TypeError, ValueError):
        return False
    # Tahsilat tarafı değiştiyse (eski kayıtlar/manuel düzeltmeler dahil) cache'i geçersiz say.
    if cache_obj.get("tahsilat_imza") != _aylik_tahsil_cache_imza(musteri_id):
        return False
    return True


_tufe_map_mem = {"ts": 0.0, "data": None}


def _tufe_map_mem_reset():
    """TÜFE tablosu değişince GET cache ve grid/ekstre haritasını temizle."""
    simple_cache_invalidate(CACHE_KEY_TUFE_VERILERI)
    _tufe_map_mem["data"] = None
    _tufe_map_mem["ts"] = 0.0


def _tufe_map_by_year_month():
    rows = fetch_all("SELECT year, month, oran FROM tufe_verileri WHERE year IS NOT NULL AND month IS NOT NULL") or []
    ay_tr_to_num = {ad.lower(): i + 1 for i, ad in enumerate(_AY_ADLARI)}
    out = {}
    for r in rows:
        try:
            yv = int(r.get("year"))
        except Exception:
            continue
        mv_raw = str(r.get("month") or "").strip()
        if not mv_raw:
            continue
        try:
            mv = int(mv_raw)
        except Exception:
            mv = ay_tr_to_num.get(mv_raw.lower())
        if not mv or mv < 1 or mv > 12:
            continue
        try:
            oran = float(r.get("oran") or 0)
        except Exception:
            oran = 0.0
        out.setdefault(yv, {})[mv] = oran
    return out


def _tufe_map_by_year_month_cached(max_age_sec: float = 300.0):
    """TÜFE tablosu nadiren değişir; ekstre/grid her istekte yeniden okumasın."""
    now = time.time()
    if _tufe_map_mem["data"] is not None and (now - float(_tufe_map_mem["ts"] or 0)) < max_age_sec:
        return _tufe_map_mem["data"]
    m = _tufe_map_by_year_month()
    _tufe_map_mem["ts"] = now
    _tufe_map_mem["data"] = m
    return m


def _tufe_latest_positive_oran_in_year_map(year_map) -> float:
    """Takvim yılı haritasında (1..12 → %) pozitif oranı olan en büyük ay numarası; TCMB ileri ayları yayınlamadığında projeksiyon."""
    if not isinstance(year_map, dict) or not year_map:
        return 0.0
    best_m, best_o = 0, 0.0
    for mk, ow in year_map.items():
        try:
            mi = int(mk)
        except (TypeError, ValueError):
            continue
        if mi < 1 or mi > 12:
            continue
        try:
            ovv = float(ow or 0)
        except (TypeError, ValueError):
            ovv = 0.0
        if ovv > 0 and math.isfinite(ovv) and mi > best_m:
            best_m, best_o = mi, ovv
    return best_o


def _aylik_grid_months_inclusive_from(bas_first: date, target_first: date) -> int:
    """bas_first ve target_first ayın 1'i; target dahil kaç ay var."""
    if target_first < bas_first:
        return 0
    return (target_first.year - bas_first.year) * 12 + (target_first.month - bas_first.month) + 1


def _aylik_grid_coerce_date(val):
    """KYC / Excel: YYYY-MM-DD, DD.MM.YYYY, date/datetime — grid için tek tip date."""
    if val is None:
        return None
    if isinstance(val, datetime):
        return val.date()
    if isinstance(val, date):
        return val
    s = str(val).strip()
    if not s:
        return None
    if re.match(r"^\d{4}-\d{2}-\d{2}", s):
        try:
            return datetime.strptime(s[:10], "%Y-%m-%d").date()
        except ValueError:
            return None
    m = re.match(r"^(\d{1,2})[./-](\d{1,2})[./-](\d{4})", s)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            return date(y, mo, d)
        except ValueError:
            return None
    return None


def _kyc_reel_anchor_month_day_for_grid(kyc):
    """Reel dönem yıldönümü: JS sozlesmelerReelDonemAnchor ve _aylik_grid_contract_core ile aynı mantık
    (kira_artis_tarihi yoksa veya geçersizse sözleşme başlangıcı; yalnızca sozlesme_tarihi fallback'i kullanılmaz)."""
    kyc = dict(kyc or {})
    bas = _aylik_grid_coerce_date(kyc.get("sozlesme_tarihi"))
    if not bas:
        bas = _aylik_grid_coerce_date(kyc.get("rent_start_date"))
    if not bas:
        return 1, 1
    artis_src_raw = kyc.get("kira_artis_tarihi") or bas
    artis_d = _aylik_grid_coerce_date(artis_src_raw) or bas
    try:
        am = int(artis_d.month)
    except Exception:
        am = int(bas.month)
    try:
        ad = int(artis_d.day)
    except Exception:
        ad = int(bas.day)
    return max(1, min(12, am)), ad


def _aylik_grid_coerce_money(val):
    """Excel/TR: 1.234,56 veya 1234,56 → float (aylık kira tabanı)."""
    if val is None or val == "":
        return 0.0
    if isinstance(val, bool):
        return 0.0
    if isinstance(val, Decimal):
        try:
            x = float(val)
            return round(x, 2) if math.isfinite(x) else 0.0
        except (TypeError, ValueError, OverflowError):
            return 0.0
    if isinstance(val, (int, float)):
        try:
            x = float(val)
            return round(x, 2) if math.isfinite(x) else 0.0
        except (TypeError, ValueError):
            return 0.0
    s = str(val).strip().replace("\xa0", " ")
    if not s:
        return 0.0
    s = s.replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        x = float(s)
        return round(x, 2) if math.isfinite(x) else 0.0
    except (TypeError, ValueError):
        return 0.0


def _kyc_karma_kira_paylari(kyc, aylik_net: float):
    """Nakit + banka tutarı toplam kiraya (toleransla) denkse karma ödeme; KDV yalnız banka payına."""
    aylik_net = _aylik_grid_coerce_money(aylik_net)
    if aylik_net <= 0:
        return False, 0.0, 0.0, 0.0, 0.0
    kyc = kyc or {}
    n = _aylik_grid_coerce_money(kyc.get("kira_nakit_tutar"))
    b = _aylik_grid_coerce_money(kyc.get("kira_banka_tutar"))
    if n <= 0 or b <= 0:
        return False, n, b, 0.0, 0.0
    tol = max(0.02, abs(aylik_net) * 0.01 + 1e-9)
    if abs((n + b) - aylik_net) > tol:
        return False, n, b, 0.0, 0.0
    r_n = n / aylik_net
    r_b = b / aylik_net
    return True, n, b, r_n, r_b


def _aylik_grid_contract_core(kyc, tufe_map):
    """
    KYC + TÜFE ile sözleşme ufku ve yıllık KDV dahil tutar haritası.
    Firma özeti raporunda her satır için 240 aylık liste üretmek yerine tek ay okumak için kullanılır.
    """
    kyc = dict(kyc or {})
    bas_raw = kyc.get("sozlesme_tarihi")
    bit_raw = kyc.get("sozlesme_bitis")
    bas = _aylik_grid_coerce_date(bas_raw)
    if not bas:
        bas = _aylik_grid_coerce_date(kyc.get("rent_start_date"))
    if not bas:
        return None

    try:
        ks_raw_early = kyc.get("kira_suresi_ay")
        ks_int_early = int(ks_raw_early) if ks_raw_early is not None and str(ks_raw_early).strip() != "" else 0
    except (TypeError, ValueError):
        ks_int_early = 0

    bit = _aylik_grid_coerce_date(bit_raw) if bit_raw is not None and str(bit_raw).strip() != "" else None
    bit = _aylik_grid_effective_bitis(kyc, bit)
    bit_kullanici = bit is not None

    bugun = date.today()
    if not bit:
        if 1 <= ks_int_early <= 240:
            bit = _add_months(bas, ks_int_early)
        else:
            # Bitiş alanı boş çok kayıtta var; firma özeti / güncel ay sütunu 0,00 kalmasın diye ufkı kapat.
            y_roll, m_roll = bugun.year, bugun.month + 11
            while m_roll > 12:
                m_roll -= 12
                y_roll += 1
            roll_son = date(y_roll, m_roll, 1)
            h_ay = max(date(bugun.year, 12, 1), roll_son)
            ly, lm = h_ay.year, h_ay.month
            inclusive_last = date(ly, lm, calendar.monthrange(ly, lm)[1])
            bit = inclusive_last + timedelta(days=1)

    if bit_kullanici:
        ay_sayisi = 0
        while ay_sayisi < 240 and _add_months(bas, ay_sayisi) < bit:
            ay_sayisi += 1
        ay_sayisi = max(0, min(240, ay_sayisi))
    else:
        bit_end = bit - timedelta(days=1)
        ay_sayisi = ((bit_end.year - bas.year) * 12 + (bit_end.month - bas.month) + 1) if bit_end >= bas else 1
        ay_sayisi = max(1, min(240, ay_sayisi))
    ks_int = ks_int_early
    # kira_suresi_ay yalnızca bitiş tarihi yokken / türetilmiş bitişte kullanılsın; aksi halde
    # sözleşme bitişi uzunken süre alanı kısa kaldıysa Nisan 2026 vb. aylar 0,00 kalıyordu.
    if not bit_kullanici and 1 <= ks_int <= 240:
        ay_sayisi = ks_int
    # Görünüm: içinde bulunulan takvim yılının tamamı + bugünkü aydan itibaren 12 ay (yenileme / öngörü)
    bas_first = date(bas.year, bas.month, 1)
    dec_son = date(bugun.year, 12, 1)
    y_roll, m_roll = bugun.year, bugun.month + 11
    while m_roll > 12:
        m_roll -= 12
        y_roll += 1
    roll_son = date(y_roll, m_roll, 1)
    if not bit_kullanici:
        horizon_need = max(
            _aylik_grid_months_inclusive_from(bas_first, dec_son),
            _aylik_grid_months_inclusive_from(bas_first, roll_son),
        )
        ay_sayisi = min(240, max(ay_sayisi, horizon_need))
    aylik_net = _aylik_grid_coerce_money(kyc.get("aylik_kira"))
    try:
        kdv_oran = float(kyc.get("kdv_oran") if kyc.get("kdv_oran") is not None else 20)
    except (TypeError, ValueError):
        kdv_oran = 20.0
    if not math.isfinite(kdv_oran) or kdv_oran < 0:
        kdv_oran = 20.0
    kira_nakit = bool(kyc.get("kira_nakit"))
    split_ok, _nak0, _ban0, r_n, r_b = _kyc_karma_kira_paylari(kyc, aylik_net)
    kdv_mult = 1.0 + kdv_oran / 100.0
    current = aylik_net
    start_year = bas.year
    max_year = start_year + max(0, (ay_sayisi - 1) // 12)
    artis_raw = kyc.get("kira_artis_tarihi") or bas
    artis = _aylik_grid_coerce_date(artis_raw) or bas
    try:
        artis_month = int(artis.month)
    except Exception:
        artis_month = bas.month
    yillik_map = {}
    for yil in range(start_year, max_year + 1):
        if split_ok:
            nak_kisim = round(current * r_n, 2)
            ban_kisim = round(current - nak_kisim, 2)
            yillik_map[yil] = round(nak_kisim + ban_kisim * kdv_mult, 2)
        elif kira_nakit:
            yillik_map[yil] = round(current * 1.0, 2)
        else:
            yillik_map[yil] = round(current * kdv_mult, 2)
        if yil < max_year:
            sonraki = yil + 1
            inner = tufe_map.get(sonraki) or {}
            if not isinstance(inner, dict):
                inner = {}
            raw_o = inner.get(artis_month)
            if raw_o is None and artis_month is not None:
                raw_o = inner.get(str(artis_month))
            try:
                oran = float(raw_o or 0)
            except (TypeError, ValueError):
                oran = 0.0
            # Gelecek yılın TÜFE satırı henüz yoksa (örn. 2026 Eylül), reel zincirdeki gibi
            # bir önceki takvim yılının aynı artış ayı oranıyla projeksiyon yap.
            if (not oran or not math.isfinite(oran)) and sonraki > start_year:
                inner_prev = tufe_map.get(sonraki - 1) or {}
                if isinstance(inner_prev, dict):
                    raw_p = inner_prev.get(artis_month)
                    if raw_p is None and artis_month is not None:
                        raw_p = inner_prev.get(str(artis_month))
                    try:
                        oran2 = float(raw_p or 0)
                    except (TypeError, ValueError):
                        oran2 = 0.0
                    if oran2 > 0 and math.isfinite(oran2):
                        oran = oran2
            # Hedef yılda artış ayı (örn. Eylül) henüz boş; Ocak–Mayıs gibi girilmiş son ayı kullan.
            if (not oran or not math.isfinite(oran)):
                o3 = _tufe_latest_positive_oran_in_year_map(inner)
                if o3 > 0 and math.isfinite(o3):
                    oran = o3
            # Hedef yıl tamamen boşsa bir önceki yılın en son girilmiş ayı.
            if (not oran or not math.isfinite(oran)) and sonraki > start_year:
                inner_prev2 = tufe_map.get(sonraki - 1) or {}
                o4 = _tufe_latest_positive_oran_in_year_map(inner_prev2)
                if o4 > 0 and math.isfinite(o4):
                    oran = o4
            if oran > 0 and math.isfinite(oran):
                current = round(current * (1 + oran / 100.0), 2)
    return {
        "bas": bas,
        "bit": bit,
        "ay_sayisi": ay_sayisi,
        "start_year": start_year,
        "yillik_map": yillik_map,
        "artis_month": artis_month,
        "aylik_net": aylik_net,
        "ks_int": ks_int_early,
        "kira_nakit": kira_nakit,
        "split_kira_odeme": split_ok,
    }


def _aylik_grid_single_month_kdv_from_core(core, ref_y, ref_m) -> float:
    """contract_core + takvim ayı → o ayın KDV dahil taban tutarı (reel katmanı ayrı)."""
    if not core:
        return 0.0
    bas = core["bas"]
    ay_sayisi = core["ay_sayisi"]
    try:
        i = (int(ref_y) - bas.year) * 12 + (int(ref_m) - bas.month)
    except (TypeError, ValueError):
        return 0.0
    if i < 0 or i >= ay_sayisi:
        return 0.0
    start_year = core["start_year"]
    proj_yil = start_year + (i // 12)
    try:
        tutar = float(core["yillik_map"].get(proj_yil) or 0)
    except (TypeError, ValueError):
        return 0.0
    return round(tutar, 2) if math.isfinite(tutar) else 0.0


def _aylik_grid_compute(musteri_id, kyc, tufe_map, tahsil_tutar_map=None):
    """
    KYC satırı + önceden yüklenmiş TÜFE haritası ile aylık grid payload üretir.
    Tam ödendi: yalnızca o aya dağıtılan tahsilatların KDV dahil kiraya göre kalanı ≤ 0,01 TL.
    (Eski: fatura/tahsil ayında herhangi bir ödeme olsa «tam ödendi» sayılıyordu — kısmi/yanlış yeşil.)
    """
    core = _aylik_grid_contract_core(kyc, tufe_map)
    if not core:
        return None
    bas = core["bas"]
    bit = core["bit"]
    ay_sayisi = core["ay_sayisi"]
    start_year = core["start_year"]
    yillik_map = core["yillik_map"]
    ks_int = core["ks_int"]
    kira_nakit = core["kira_nakit"]
    aylik_net = core["aylik_net"]
    split_kira = bool(core.get("split_kira_odeme"))
    artis_month = core["artis_month"]
    aylar = []
    tahsil_tutar_map = tahsil_tutar_map or {}
    for i in range(ay_sayisi):
        ay_toplam = (bas.month - 1) + i
        y = bas.year + (ay_toplam // 12)
        m = (ay_toplam % 12) + 1
        d = date(y, m, 1)
        proj_yil = start_year + (i // 12)
        tutar = float(yillik_map.get(proj_yil) or 0)
        ay_key = f"{d.year}-{d.month}"
        ay_iso = d.isoformat()
        odenen = float(tahsil_tutar_map.get(ay_iso) or 0)
        kalan = max(round(tutar - odenen, 2), 0.0)
        tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
        kismi = odenen > 0 and kalan > tol
        tam_odendi = kalan <= tol
        # Hücre metni: her ay sözleşme taban KDV dahil tutarı (yıl içi tutarlı); kalan ayrı alanda.
        gosterim_tutar = tutar
        aylar.append({
            "yil": d.year,
            "ay": d.month,
            "ay_key": ay_key,
            "tutar_kdv_dahil": round(gosterim_tutar, 2),
            "tahsil_edildi": tam_odendi,
            "kismi_tahsilat": kismi,
            "odenen_tutar_kdv": round(odenen, 2),
            "kalan_tutar_kdv": round(kalan, 2),
            "brut_tutar_kdv": round(tutar, 2),
        })
    return {
        "musteri_id": musteri_id,
        "baslangic": bas.isoformat(),
        "bitis": bit.isoformat(),
        "kapanis_tarihi": _kyc_date_iso(kyc.get("kapanis_tarihi")),
        "kapanis_sonrasi_borc_ay": _normalize_kapanis_sonrasi_borc_ay(kyc, str(kyc.get("durum") or "")),
        "kira_suresi_ay": ks_int if 1 <= ks_int <= 240 else None,
        "kira_nakit": kira_nakit,
        "split_kira_odeme": split_kira,
        "kira_nakit_tutar": round(_aylik_grid_coerce_money(kyc.get("kira_nakit_tutar")), 2),
        "kira_banka_tutar": round(_aylik_grid_coerce_money(kyc.get("kira_banka_tutar")), 2),
        "taban_aylik_net": aylik_net,
        "artis_ay": artis_month,
        "aylar": aylar,
        "compute_rev": AYLIK_GRID_COMPUTE_REV,
        "updated_at": datetime.now().isoformat(timespec="seconds"),
    }


_musteri_kyc_grid_mem: dict[int, dict] = {}


def _musteri_kyc_grup_for_aylik_grid(musteri_id: int):
    """customers + son musteri_kyc — aylık grid / cari ekstre aynı satır."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return {}
    if mid <= 0:
        return {}
    now = time.time()
    hit = _musteri_kyc_grid_mem.get(mid)
    if hit and (now - float(hit.get("ts") or 0)) < 45.0 and hit.get("row") is not None:
        return dict(hit["row"])
    row = fetch_one(
        """
        SELECT mk.sozlesme_tarihi, mk.sozlesme_bitis, mk.aylik_kira, mk.kira_artis_tarihi, mk.kira_suresi_ay, mk.kira_nakit,
               mk.kira_nakit_tutar, mk.kira_banka_tutar, mk.kdv_oran,
               c.kapanis_tarihi, c.kapanis_sonrasi_borc_ay, c.durum, c.rent_start_date
        FROM customers c
        LEFT JOIN LATERAL (
            SELECT sozlesme_tarihi, sozlesme_bitis, aylik_kira, kira_artis_tarihi, kira_suresi_ay, kira_nakit,
                   kira_nakit_tutar, kira_banka_tutar, kdv_oran
            FROM musteri_kyc
            WHERE musteri_id = c.id
            ORDER BY id DESC
            LIMIT 1
        ) mk ON TRUE
        WHERE c.id = %s
        """,
        (mid,),
    ) or {}
    _musteri_kyc_grid_mem[mid] = {"ts": now, "row": dict(row) if row else {}}
    if len(_musteri_kyc_grid_mem) > 200:
        stale = [k for k, v in _musteri_kyc_grid_mem.items() if (now - float(v.get("ts") or 0)) > 90.0]
        for k in stale:
            _musteri_kyc_grid_mem.pop(k, None)
    return row


def _musteri_reel_donem_manual_dict_from_db(musteri_id: int) -> dict[int, float]:
    """musteri_reel_donem_tutar: donem_yil -> tutar_kdv_dahil (boş dict olabilir)."""
    _ensure_musteri_reel_donem_tutar_table()
    rows = fetch_all(
        "SELECT donem_yil, tutar_kdv_dahil FROM musteri_reel_donem_tutar WHERE musteri_id = %s",
        (int(musteri_id),),
    ) or []
    out: dict[int, float] = {}
    for r in rows:
        try:
            dy = int(r.get("donem_yil"))
            out[dy] = round(float(r.get("tutar_kdv_dahil") or 0), 2)
        except (TypeError, ValueError):
            continue
    return out


def _reel_manual_merge_db_and_client(musteri_id: int, client: dict | None) -> dict[int, float]:
    """DB reel satırları + tarayıcıdan gelen önizleme (Uygula) üzerine yazar."""
    out = _musteri_reel_donem_manual_dict_from_db(musteri_id)
    if not client or not isinstance(client, dict):
        return out
    for k, v in client.items():
        try:
            yi = int(k)
            vf = float(v)
            if math.isfinite(vf) and vf >= 0:
                out[yi] = round(vf, 2)
        except (TypeError, ValueError):
            continue
    return out


def _aylik_grid_apply_reel_donem_overlay_to_payload(
    musteri_id: int, kyc: dict, tufe_map: dict, payload: dict, manual_reel_by_year=None
) -> None:
    """musteri_reel_donem_tutar (DB) → 12 ay sabit KDV dahil; grid önbelleğinde TÜFE zinciri kullanılmaz."""
    if not payload or not isinstance(payload.get("aylar"), list) or not kyc:
        return
    bas_soz = _aylik_grid_coerce_date(kyc.get("sozlesme_tarihi"))
    if not bas_soz:
        bas_soz = _aylik_grid_coerce_date(kyc.get("rent_start_date"))
    if not bas_soz:
        return
    artis_d = _aylik_grid_coerce_date(kyc.get("kira_artis_tarihi")) or bas_soz
    artis_month = int(artis_d.month)
    artis_day = int(artis_d.day)
    bit = _aylik_grid_coerce_date(kyc.get("sozlesme_bitis"))
    y_end = date.today().year
    try:
        aylar = payload.get("aylar") or []
        if aylar:
            la = aylar[-1]
            y_end = max(y_end, int(la.get("yil") or 0), bas_soz.year)
    except (TypeError, ValueError, IndexError):
        y_end = max(y_end, bas_soz.year)
    if bit:
        try:
            y_end = max(y_end, int(bit.year))
        except (TypeError, ValueError):
            pass
    manual = manual_reel_by_year
    if manual is None:
        manual = _musteri_reel_donem_manual_dict_from_db(int(musteri_id))
    elif not isinstance(manual, dict):
        manual = {}
    try:
        reel_map = _reel_ay_key_tutar_map_db_flat_only(
            bas_soz, artis_month, artis_day, manual
        )
    except Exception:
        reel_map = {}
    if not reel_map:
        return
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    for a in payload["aylar"]:
        if not isinstance(a, dict):
            continue
        key = str(a.get("ay_key") or "").strip()
        if not key:
            try:
                key = f"{int(a.get('yil'))}-{int(a.get('ay'))}"
            except (TypeError, ValueError):
                continue
        nk = _firma_ozet_normalize_tahsil_ay_key(key)
        if not nk or nk not in reel_map:
            continue
        try:
            new_t = round(float(reel_map[nk]), 2)
        except (TypeError, ValueError):
            continue
        if not math.isfinite(new_t) or new_t <= 0:
            continue
        try:
            odenen = float(a.get("odenen_tutar_kdv") or 0)
        except (TypeError, ValueError):
            odenen = 0.0
        if not math.isfinite(odenen):
            odenen = 0.0
        odenen = round(min(max(odenen, 0.0), new_t), 2)
        kalan = max(round(new_t - odenen, 2), 0.0)
        a["brut_tutar_kdv"] = new_t
        a["tutar_kdv_dahil"] = new_t
        a["odenen_tutar_kdv"] = round(odenen, 2)
        a["kalan_tutar_kdv"] = kalan
        a["tahsil_edildi"] = kalan <= tol
        a["kismi_tahsilat"] = odenen > tol and kalan > tol


def _aylik_grid_payload_reel_overlay_from_db(musteri_id: int, payload: dict | None) -> dict | None:
    """Önbellekten dönen payload'a DB'deki reel dönem tutarlarını uygular."""
    if not payload or not isinstance(payload.get("aylar"), list):
        return payload
    manual = _musteri_reel_donem_manual_dict_from_db(int(musteri_id))
    if not manual:
        return payload
    kyc = _musteri_kyc_grup_for_aylik_grid(int(musteri_id))
    if not kyc:
        return payload
    tm = _tufe_map_by_year_month_cached()
    _aylik_grid_apply_reel_donem_overlay_to_payload(
        int(musteri_id), kyc, tm, payload, manual_reel_by_year=manual
    )
    return payload


def _build_aylik_grid_cache_payload(musteri_id, tufe_map=None, kyc_row=None, manual_reel_by_year=None):
    if kyc_row is not None:
        kyc = dict(kyc_row)
    else:
        kyc = _musteri_kyc_grup_for_aylik_grid(int(musteri_id))
    if not kyc:
        return None
    tm = tufe_map if tufe_map is not None else _tufe_map_by_year_month_cached()
    tahsil_map = _aylik_tahsil_tutar_map(musteri_id)
    payload = _aylik_grid_compute(musteri_id, kyc, tm, tahsil_map)
    if isinstance(payload, dict):
        if manual_reel_by_year is None:
            manual_reel_by_year = _musteri_reel_donem_manual_dict_from_db(int(musteri_id))
        if isinstance(manual_reel_by_year, dict) and manual_reel_by_year:
            _aylik_grid_apply_reel_donem_overlay_to_payload(
                int(musteri_id), kyc, tm, payload, manual_reel_by_year=manual_reel_by_year
            )
        acik_aylik_tutar_aylari = _aylik_grid_acik_tutar_ay_keys_normalized(musteri_id)
        tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
        tahsil_pl_grid = tahsil_map or {}
        batch_grid = _ekstre_tahsil_batch_maps_from_rows(
            _ekstre_tahsil_rows_for_musteri(int(musteri_id))
        )
        # |AYLIK_TAH| marker tutarı FIFO'dan kaçsa bile her ay için birleştir (yalnızca «açık fatura» aylarında değil).
        for a in payload.get("aylar") or []:
            if not isinstance(a, dict):
                continue
            try:
                brut = float(a.get("brut_tutar_kdv") or a.get("tutar_kdv_dahil") or 0)
            except (TypeError, ValueError):
                brut = 0.0
            try:
                odenen = float(a.get("odenen_tutar_kdv") or 0)
            except (TypeError, ValueError):
                odenen = 0.0
            try:
                yy = int(a.get("yil"))
                mm = int(a.get("ay"))
                iso_m1 = date(yy, mm, 1).isoformat()
            except Exception:
                iso_m1 = ""
            mk_t = tm_t = te_t = tt_t = 0.0
            if iso_m1 and brut > tol:
                mk_t = round(float((batch_grid.get("marker") or {}).get(iso_m1) or 0), 2)
                tm_t = round(float((tahsil_pl_grid or {}).get(iso_m1) or 0), 2)
                te_t = round(float((batch_grid.get("eslesme") or {}).get(iso_m1) or 0), 2)
                tt_t = round(float((batch_grid.get("tarih_ay") or {}).get(iso_m1) or 0), 2)
                odenen = _grid_payload_ay_odenen_kdv(brut, odenen, mk_t, tm_t, te_t, tt_t, tol)
            reel_kap_b = 0.0
            if isinstance(manual_reel_by_year, dict):
                try:
                    reel_kap_b = round(float(manual_reel_by_year.get(yy) or 0), 2)
                except (TypeError, ValueError):
                    reel_kap_b = 0.0
            kap_b = reel_kap_b if reel_kap_b > tol else round(brut, 2)
            if _grid_payload_marker_panel_tam_kapandi(mk_t, brut, yy, manual_reel_by_year, tol):
                a["odenen_tutar_kdv"] = round(mk_t, 2)
                a["kalan_tutar_kdv"] = 0.0
                a["tahsil_edildi"] = True
                a["kismi_tahsilat"] = False
            elif mk_t > tol and kap_b > tol:
                kalan_p = round(max(kap_b - mk_t, 0), 2)
                a["odenen_tutar_kdv"] = round(mk_t, 2)
                a["kalan_tutar_kdv"] = kalan_p
                a["tahsil_edildi"] = kalan_p <= tol
                a["kismi_tahsilat"] = kalan_p > tol
            else:
                kalan = round(max(brut - odenen, 0), 2)
                a["odenen_tutar_kdv"] = round(odenen, 2)
                a["kalan_tutar_kdv"] = kalan
                a["tahsil_edildi"] = kalan <= tol
                a["kismi_tahsilat"] = odenen > tol and kalan > tol
            a["tutar_kdv_dahil"] = round(max(brut, 0.01), 2)
        if acik_aylik_tutar_aylari:
            for a in payload.get("aylar") or []:
                if not isinstance(a, dict):
                    continue
                nk = _firma_ozet_normalize_tahsil_ay_key(
                    str(a.get("ay_key") or f"{a.get('yil')}-{a.get('ay')}")
                )
                if not nk or nk not in acik_aylik_tutar_aylari:
                    continue
                try:
                    odenen = float(a.get("odenen_tutar_kdv") or 0)
                except (TypeError, ValueError):
                    odenen = 0.0
                try:
                    kalan = float(a.get("kalan_tutar_kdv") or 0)
                except (TypeError, ValueError):
                    kalan = 0.0
                if odenen > tol:
                    a["acik_aylik_borc_faturasi"] = kalan > tol
                else:
                    a["acik_aylik_borc_faturasi"] = True
                    if not bool(a.get("tahsil_edildi")):
                        a["tahsil_edildi"] = False
                        a["kismi_tahsilat"] = False
                        a["kalan_tutar_kdv"] = round(max(kalan, 0), 2)
        payload["tahsilat_imza"] = _aylik_tahsil_cache_imza(musteri_id)
    return payload


def _aylik_tahsil_cache_imza(musteri_id):
    row = fetch_one(
        """
        SELECT
            COUNT(*)::bigint AS cnt,
            COALESCE(SUM(COALESCE(tutar, 0)), 0)::numeric AS toplam,
            COALESCE(MAX(tahsilat_tarihi::timestamp), TIMESTAMP '1970-01-01') AS mx
        FROM tahsilatlar
        WHERE (musteri_id = %s OR customer_id = %s)
          AND COALESCE(tutar, 0) > 0
        """,
        (musteri_id, musteri_id),
    ) or {}
    try:
        cnt = int(row.get("cnt") or 0)
    except (TypeError, ValueError):
        cnt = 0
    try:
        toplam = round(float(row.get("toplam") or 0), 2)
    except (TypeError, ValueError):
        toplam = 0.0
    mx_raw = row.get("mx")
    if hasattr(mx_raw, "isoformat"):
        mx = str(mx_raw.isoformat())
    else:
        mx = str(mx_raw or "1970-01-01T00:00:00")
    return f"{cnt}|{toplam:.2f}|{mx}"


def _aylik_remaining_brut_by_iso_from_kyc(kyc, tufe_map=None) -> dict[str, float]:
    """FIFO dağıtımı için ay başına sözleşme brüt borç (cache okumadan)."""
    remaining_by_iso: dict[str, float] = {}
    if not kyc:
        return remaining_by_iso
    try:
        tm = tufe_map if tufe_map is not None else _tufe_map_by_year_month_cached()
        core = _aylik_grid_contract_core(kyc, tm)
        if not core:
            return remaining_by_iso
        bas = core["bas"]
        ay_sayisi = int(core["ay_sayisi"] or 0)
        start_year = int(core["start_year"] or bas.year)
        yillik_map = core["yillik_map"] or {}
        for i in range(max(0, ay_sayisi)):
            ay_toplam = (bas.month - 1) + i
            y = bas.year + (ay_toplam // 12)
            m = (ay_toplam % 12) + 1
            proj_yil = start_year + (i // 12)
            tut = round(float(yillik_map.get(proj_yil) or 0), 2)
            if tut <= 0:
                continue
            remaining_by_iso[date(y, m, 1).isoformat()] = tut
    except Exception:
        pass
    return remaining_by_iso


def _aylik_tahsil_tutar_map(musteri_id, tahsil_rows=None, remaining_by_iso=None, kyc_row=None, tufe_map=None):
    """Tahsilat tutarlarını aya dağıtır: |AYLIK_TAH|YYYY-MM-DD| varsa oraya; yoksa fatura ayı, o da yoksa tahsilat tarihi.

    - |BTUFRT|: yalnızca BTUFRT_GRID_ODENME_DAHIL_DEGIL (2026-01-01) ve sonrası ayları hedefliyorsa grid
      ödemesinden sayılmaz; önceki ay muhasebe çiftleri gridde yeşilde kalır.
    - Birden fazla |AYLIK_TAH| varsa tutar aylara bölünür.
    tahsil_rows / remaining_by_iso / kyc_row: ekstre gibi sıcak yollarda tekrarlayan SQL'i keser.
    """
    if tahsil_rows is not None:
        rows = tahsil_rows
    else:
        rows = fetch_all(
            """
            SELECT t.id, COALESCE(t.aciklama, '') AS aciklama, COALESCE(t.tutar, 0) AS tutar,
                   t.tahsilat_tarihi, f.fatura_tarihi
            FROM tahsilatlar t
            LEFT JOIN faturalar f ON f.id = t.fatura_id
            WHERE (t.musteri_id = %s OR t.customer_id = %s)
              AND COALESCE(t.tutar, 0) > 0
            ORDER BY t.tahsilat_tarihi ASC NULLS LAST, t.id ASC
            """,
            (musteri_id, musteri_id),
        ) or []
    if remaining_by_iso is not None:
        remaining_by_iso = dict(remaining_by_iso)
    else:
        remaining_by_iso = {}
        try:
            cr = fetch_one("SELECT payload FROM musteri_aylik_grid_cache WHERE musteri_id = %s", (musteri_id,))
            payload_raw = (cr or {}).get("payload")
            if payload_raw:
                pobj = json.loads(payload_raw) if isinstance(payload_raw, str) else payload_raw
                aylar = pobj if isinstance(pobj, list) else ((pobj or {}).get("aylar") or [])
                if isinstance(aylar, list):
                    for a in aylar:
                        if not isinstance(a, dict):
                            continue
                        try:
                            yy = int(a.get("yil"))
                            mm = int(a.get("ay"))
                            tv = round(float(a.get("tutar_kdv_dahil") or 0), 2)
                        except (TypeError, ValueError):
                            continue
                        if tv <= 0 or mm < 1 or mm > 12:
                            continue
                        remaining_by_iso[date(yy, mm, 1).isoformat()] = tv
        except Exception:
            remaining_by_iso = {}

        if not remaining_by_iso:
            if kyc_row is not None:
                remaining_by_iso = _aylik_remaining_brut_by_iso_from_kyc(kyc_row, tufe_map)
            else:
                try:
                    kyc = fetch_one(
                        """
                        SELECT mk.sozlesme_tarihi, mk.sozlesme_bitis, mk.aylik_kira, mk.kira_artis_tarihi,
                               mk.kira_suresi_ay, mk.kira_nakit, mk.kira_nakit_tutar, mk.kira_banka_tutar
                        FROM customers c
                        LEFT JOIN LATERAL (
                            SELECT *
                            FROM musteri_kyc
                            WHERE musteri_id = c.id
                            ORDER BY id DESC
                            LIMIT 1
                        ) mk ON TRUE
                        WHERE c.id = %s
                        """,
                        (musteri_id,),
                    ) or {}
                    remaining_by_iso = _aylik_remaining_brut_by_iso_from_kyc(kyc, tufe_map)
                except Exception:
                    pass

    def _alloc_oldest(rem_tutar):
        try:
            rem = round(float(rem_tutar or 0), 2)
        except (TypeError, ValueError):
            rem = 0.0
        if rem <= 0:
            return []
        if not remaining_by_iso:
            return []
        out_alloc = []
        for iso in sorted(remaining_by_iso.keys()):
            if rem <= 0.004:
                break
            acik = round(float(remaining_by_iso.get(iso) or 0), 2)
            if acik <= 0:
                continue
            pay = round(min(acik, rem), 2)
            if pay <= 0:
                continue
            out_alloc.append((iso, pay))
            remaining_by_iso[iso] = round(acik - pay, 2)
            rem = round(rem - pay, 2)
        return out_alloc
    out = defaultdict(float)
    for r in rows:
        ac = str(r.get("aciklama") or "")
        tut = float(r.get("tutar") or 0)
        if tut <= 0:
            continue
        marker_isos = re.findall(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac)
        pay_tokens = re.findall(r"\|AYLIK_PAY\|([0-9]{4}-[0-9]{2}-[0-9]{2})=([0-9]+(?:\.[0-9]+)?)\|", ac)
        if _aylik_btufrt_row_skip_grid_odeme(ac, marker_isos, r):
            # Grid ödemesinden çıkarmak ekstre ile çelişiyordu (ekstrede alacak varken tüm aylar kırmızı).
            # BTUFRT satırını yine de ekstre ile aynı aya yaz — yeşil/kısmi durumu DB tahsilatı ile hizalı kalsın.
            iso_e = _tahsil_row_ekstre_eslesme_ay_iso(r)
            if iso_e and tut > 0:
                out[iso_e] += tut
            continue
        if pay_tokens:
            used_any = False
            for iso_raw, tut_raw in pay_tokens:
                try:
                    dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                    iso = date(dd.year, dd.month, 1).isoformat()
                    pv = round(float(tut_raw or 0), 2)
                except Exception:
                    continue
                if pv <= 0:
                    continue
                out[iso] += pv
                if iso in remaining_by_iso:
                    remaining_by_iso[iso] = round(max(float(remaining_by_iso.get(iso) or 0) - pv, 0), 2)
                used_any = True
            if used_any:
                continue
        # Marker yoksa oldest-open dağıt (manuel/legacy markersız kayıtlar).
        # Marker varsa o aya/aylara sadık kal.
        ac_plain = re.sub(r"\|AYLIK_TAH\|[0-9]{4}-[0-9]{2}-[0-9]{2}\|", " ", ac)
        ac_plain = re.sub(r"\|AYLIK_PAY\|[0-9]{4}-[0-9]{2}-[0-9]{2}=[0-9]+(?:\.[0-9]+)?\|", " ", ac_plain)
        ac_plain = " ".join(ac_plain.split()).strip()
        if not marker_isos:
            al = _alloc_oldest(tut)
            if al:
                for iso, pv in al:
                    out[iso] += pv
                continue
        if marker_isos:
            n = len(marker_isos)
            cents_total = int(round(tut * 100))
            if cents_total <= 0 or n <= 0:
                continue
            base = cents_total // n
            rem = cents_total % n
            for i, iso_raw in enumerate(marker_isos):
                share_cents = base + (1 if i < rem else 0)
                if share_cents <= 0:
                    continue
                try:
                    dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                    iso = date(dd.year, dd.month, 1).isoformat()
                except ValueError:
                    continue
                out[iso] += share_cents / 100.0
            continue
        # Eski kayıtlar marker'sız olabilir: önce fatura ayı, yoksa tahsilat ayı.
        d = r.get("fatura_tarihi") or r.get("tahsilat_tarihi")
        if hasattr(d, "year"):
            iso = date(int(d.year), int(d.month), 1).isoformat()
            out[iso] += tut
        elif d:
            try:
                ds = str(d)[:10]
                dd = datetime.strptime(ds, "%Y-%m-%d").date()
                iso = date(dd.year, dd.month, 1).isoformat()
                out[iso] += tut
            except Exception:
                pass
    return dict(out)


def _load_aylik_tahsil_ay_keys_by_musteri():
    """musteri_id -> {'YYYY-MM-DD', ...} AYLIK_TAH marker anahtarları (ayın 1'i)."""
    rows = fetch_all(
        """
        SELECT musteri_id, customer_id, aciklama, tutar
        FROM tahsilatlar
        WHERE COALESCE(aciklama, '') LIKE '%%|AYLIK_TAH|%%'
          AND COALESCE(tutar, 0) > 0
        """
    ) or []
    by_mid = defaultdict(set)
    for r in rows:
        try:
            mid = int(r.get("musteri_id") or r.get("customer_id") or 0)
        except (TypeError, ValueError):
            continue
        if mid <= 0:
            continue
        for k in re.findall(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", str(r.get("aciklama") or "")):
            by_mid[mid].add(k)
    return by_mid


def _load_manual_fatura_ay_by_musteri():
    """AYLIK_TUTAR işaretçisi olmayan faturalar: müşteri + ay başı (manuel / dış kayıt)."""
    rows = fetch_all(
        f"""
        SELECT musteri_id, (DATE_TRUNC('month', fatura_tarihi::date))::date AS m
        FROM faturalar
        WHERE fatura_tarihi IS NOT NULL
          AND COALESCE(notlar, '') NOT LIKE '%%|AYLIK_TUTAR|%%'
          AND {sql_expr_fatura_not_gib_taslak("notlar")}
        """
    ) or []
    by_mid = defaultdict(set)
    for r in rows:
        try:
            mid = int(r.get("musteri_id") or 0)
        except (TypeError, ValueError):
            continue
        if mid <= 0:
            continue
        m = r.get("m")
        if m and hasattr(m, "year"):
            by_mid[mid].add(m)
        elif m:
            try:
                by_mid[mid].add(datetime.strptime(str(m)[:10], "%Y-%m-%d").date())
            except Exception:
                pass
    return by_mid


def _upsert_aylik_grid_cache(musteri_id, tufe_map=None):
    payload = _build_aylik_grid_cache_payload(musteri_id, tufe_map=tufe_map)
    if not payload:
        return None
    return _persist_grid_cache_with_panel(musteri_id, payload)


# Kısa süreli bellek önbelleği: aynı müşteri için grid-cache + tahsil-durum arka arkaya gelince
# _build_aylik_grid_cache_payload tekrar çalışmasın (Supabase round-trip + ağır hesap).
_aylik_grid_payload_mem: dict = {}


def _invalidate_aylik_grid_payload_mem(musteri_id=None) -> None:
    if musteri_id is None:
        _aylik_grid_payload_mem.clear()
        return
    try:
        _aylik_grid_payload_mem.pop(int(musteri_id), None)
    except (TypeError, ValueError):
        pass


_giris_kaydet_schema_ready = False
_giris_kaydet_schema_lock = threading.Lock()


def _giris_kaydet_schema_ensure_once() -> None:
    global _giris_kaydet_schema_ready
    if _giris_kaydet_schema_ready:
        return
    with _giris_kaydet_schema_lock:
        if _giris_kaydet_schema_ready:
            return
        ensure_customers_musteri_no()
        ensure_customers_bizim_hesap()
        ensure_customers_grup2_secimleri()
        ensure_customers_kapanis_sonrasi_borc_ay()
        ensure_grup2_etiketleri_table()
        _giris_kaydet_schema_ready = True


def _defer_aylik_grid_cache_rebuild(musteri_id) -> None:
    """Kaydet yanıtını bekletmeden grid önbelleğini arka planda yenile."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return
    if mid <= 0:
        return
    _invalidate_aylik_grid_payload_mem(mid)
    try:
        app = current_app._get_current_object()
    except Exception:
        return

    def _work():
        with app.app_context():
            try:
                _upsert_aylik_grid_cache(mid)
            except Exception as ex:
                try:
                    current_app.logger.warning("defer grid cache mid=%s: %r", mid, ex)
                except Exception:
                    pass

    threading.Thread(target=_work, daemon=True).start()


def _parse_aylik_grid_cache_payload_raw(raw):
    if raw is None:
        return None
    try:
        obj = json.loads(raw) if isinstance(raw, str) else raw
    except Exception:
        return None
    return obj if isinstance(obj, dict) else None


def _read_aylik_grid_cache_payload(musteri_id, mem_ttl: float = 12.0):
    """musteri_aylik_grid_cache satırından payload; yoksa None (yeniden hesap çağıran karar verir)."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return None
    now = time.time()
    hit = _aylik_grid_payload_mem.get(mid)
    if hit and (now - float(hit[0] or 0)) < mem_ttl:
        return hit[1]
    row = fetch_one(
        "SELECT payload FROM musteri_aylik_grid_cache WHERE musteri_id = %s",
        (mid,),
    )
    payload = _parse_aylik_grid_cache_payload_raw((row or {}).get("payload"))
    if payload:
        _aylik_grid_payload_mem[mid] = (now, payload)
    return payload


def _read_aylik_grid_cache_payload_batch(musteri_ids, mem_ttl: float = 12.0) -> dict:
    """Birden çok müşteri için tek SQL ile grid cache okur."""
    mids: list[int] = []
    seen = set()
    for x in musteri_ids or []:
        try:
            i = int(x)
        except (TypeError, ValueError):
            continue
        if i <= 0 or i in seen:
            continue
        seen.add(i)
        mids.append(i)
    if not mids:
        return {}
    now = time.time()
    out: dict = {}
    need_db: list[int] = []
    for mid in mids:
        hit = _aylik_grid_payload_mem.get(mid)
        if hit and (now - float(hit[0] or 0)) < mem_ttl:
            if hit[1]:
                out[mid] = hit[1]
        else:
            need_db.append(mid)
    if need_db:
        rows = fetch_all(
            "SELECT musteri_id, payload FROM musteri_aylik_grid_cache WHERE musteri_id = ANY(%s::bigint[])",
            (need_db,),
        ) or []
        for r in rows:
            try:
                mid = int(r.get("musteri_id") or 0)
            except (TypeError, ValueError):
                continue
            if mid <= 0:
                continue
            payload = _parse_aylik_grid_cache_payload_raw(r.get("payload"))
            if payload:
                out[mid] = payload
                _aylik_grid_payload_mem[mid] = (now, payload)
    return out


def _aylik_tahsil_edilen_aylar_from_payload(payload) -> set[str]:
    """Grid payload içinden tam ödenmiş ay anahtarları (YYYY-M)."""
    out: set[str] = set()
    if not isinstance(payload, dict):
        return out
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict) or not a.get("tahsil_edildi"):
            continue
        nk = _firma_ozet_normalize_tahsil_ay_key(str(a.get("ay_key") or ""))
        if nk:
            out.add(nk)
    return out


def _aylik_tahsil_durum_finalize_ay_set(
    musteri_id: int,
    ay_set: set[str],
    *,
    payload=None,
    marker_ay_only=None,
    ekstre_ay_only=None,
) -> list[str]:
    """Marker + ekstre birleşimi ve kısmi tahsilat kuralları (grid payload tek kaynak)."""
    mid = int(musteri_id)
    ay_set = set(ay_set or [])
    if marker_ay_only is None:
        marker_ay_only = _aylik_tahsil_marker_aylar_set_normalized(mid)
    if ekstre_ay_only is None:
        ekstre_ay_only = _aylik_tahsil_ekstre_eslesme_aylar_set_normalized(mid)
    ay_set.update(marker_ay_only)
    ay_set.update(ekstre_ay_only)
    if payload is None:
        payload = _read_aylik_grid_cache_payload(mid)
    if payload is None:
        try:
            payload = _build_aylik_grid_cache_payload(mid, tufe_map=_tufe_map_by_year_month_cached())
        except Exception:
            payload = None
    if isinstance(payload, dict):
        tol = AYLIK_GRID_TAM_ODENDI_TOLERANS
        for a in (payload.get("aylar") or []):
            if not isinstance(a, dict):
                continue
            nk = _firma_ozet_normalize_tahsil_ay_key(str(a.get("ay_key") or ""))
            if not nk:
                continue
            tahsil_edildi = bool(a.get("tahsil_edildi"))
            kismi = bool(a.get("kismi_tahsilat"))
            try:
                kalan = float(a.get("kalan_tutar_kdv") or 0)
            except (TypeError, ValueError):
                kalan = 0.0
            try:
                odenen = float(a.get("odenen_tutar_kdv") or 0)
            except (TypeError, ValueError):
                odenen = 0.0
            kismi_like = kismi or (odenen > tol and kalan > tol)
            if not tahsil_edildi:
                ay_set.discard(nk)
            elif kismi_like:
                ay_set.discard(nk)
            elif nk in ekstre_ay_only and tahsil_edildi and not kismi_like:
                ay_set.add(nk)
    try:
        tol_p = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
        panel_by_iso = _load_musteri_panel_by_iso(mid) or {}
        for iso_k, prow in panel_by_iso.items():
            if not isinstance(prow, dict):
                continue
            nk_p = _firma_ozet_normalize_tahsil_ay_key(str(iso_k))
            if not nk_p:
                continue
            try:
                pt = round(float(prow.get("tahsil") or 0), 2)
                pk = round(max(float(prow.get("kalan") or 0), 0), 2)
            except (TypeError, ValueError):
                continue
            if pk > tol_p and pt <= tol_p:
                ay_set.discard(nk_p)
    except Exception:
        pass
    return sorted(
        ay_set,
        key=lambda s: (int(str(s).split("-")[0]), int(str(s).split("-")[1])),
    )


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@bp.route('/')
@giris_gerekli
def index():
    """Giriş / Müşteri Kaydı ana sayfası"""
    embed = str(request.args.get('embed') or '').lower() in ('1', 'true', 'yes', 'on')
    return render_template('giris/index.html', embed=embed)


@bp.route('/api/potansiyel', methods=['GET', 'POST'])
@giris_gerekli
def api_potansiyel():
    """Potansiyel müşteri (CRM lead) listesi + ekleme/güncelleme."""
    if request.method == 'GET':
        arama = (request.args.get('q') or '').strip()
        durum = (request.args.get('durum') or '').strip() or None
        params = []
        where = []
        if arama:
            norm = turkish_lower(arama)
            where.append("("
                         "LOWER(translate(ad_soyad, 'İIıŞşĞğÜüÖöÇç', 'iiissgguuoocc')) LIKE %s "
                         "OR LOWER(translate(COALESCE(firma_adi,''), 'İIıŞşĞğÜüÖöÇç', 'iiissgguuoocc')) LIKE %s "
                         "OR telefon ILIKE %s)")
            q = f"%{norm}%"
            params.extend([q, q, f"%{arama}%"])
        if durum:
            where.append("LOWER(COALESCE(lead_durumu,'')) = %s")
            params.append(durum.lower())
        sql = "SELECT * FROM crm_leads"
        if where:
            sql += " WHERE " + " AND ".join(where)
        sql += " ORDER BY COALESCE(takip_tarihi::timestamp, son_gorusme::timestamp, ilk_gorusme::timestamp, NOW()) DESC, id DESC"
        rows = fetch_all(sql, tuple(params))
        return jsonify(rows or [])

    data = request.get_json() or {}
    pid = data.get('id')
    ad_soyad = (data.get('ad_soyad') or data.get('ad') or '').strip()
    if not ad_soyad:
        return jsonify({'ok': False, 'mesaj': 'Ad Soyad zorunlu.'}), 400
    firma_adi = (data.get('firma_adi') or '').strip()
    telefon = (data.get('telefon') or '').strip()
    email = (data.get('email') or '').strip()
    sektor = (data.get('sektor') or '').strip()
    hizmet_turu = (data.get('hizmet_turu') or data.get('paket') or '').strip()
    lead_durumu = (data.get('lead_durumu') or data.get('durum') or 'Yeni Lead').strip()
    try:
        lead_skoru = int(data.get('lead_skoru') or 0)
    except Exception:
        lead_skoru = 0
    ilk_gorusme = data.get('ilk_gorusme') or None
    son_gorusme = data.get('son_gorusme') or None
    takip_tarihi = data.get('takip_tarihi') or data.get('hatirlatma_tarihi') or None
    sorumlu_satis = (data.get('sorumlu_satis') or '').strip()
    notlar = (data.get('notlar') or data.get('gorusme_notu') or '').strip()
    # Takip tarihi boşsa, varsayılan: bugün + 2 gün
    if not takip_tarihi:
        takip_tarihi = (date.today() + timedelta(days=2)).isoformat()

    if pid:
        execute(
            """UPDATE crm_leads
                   SET ad_soyad=%s, firma_adi=%s, telefon=%s, email=%s, sektor=%s,
                       hizmet_turu=%s, lead_durumu=%s, lead_skoru=%s,
                       ilk_gorusme=%s, son_gorusme=%s, takip_tarihi=%s,
                       sorumlu_satis=%s, notlar=%s
                 WHERE id=%s""",
            (
                ad_soyad, firma_adi, telefon, email, sektor,
                hizmet_turu, lead_durumu, lead_skoru,
                ilk_gorusme, son_gorusme, takip_tarihi,
                sorumlu_satis, notlar, pid,
            ),
        )
        return jsonify({'ok': True, 'mesaj': 'Potansiyel müşteri güncellendi.', 'id': pid})

    row = execute_returning(
        """INSERT INTO crm_leads (
                ad_soyad, firma_adi, telefon, email, sektor,
                hizmet_turu, lead_durumu, lead_skoru,
                ilk_gorusme, son_gorusme, takip_tarihi,
                sorumlu_satis, notlar
            )
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            RETURNING id""",
        (
            ad_soyad, firma_adi, telefon, email, sektor,
            hizmet_turu, lead_durumu, lead_skoru,
            ilk_gorusme, son_gorusme, takip_tarihi,
            sorumlu_satis, notlar,
        ),
    )
    return jsonify({'ok': True, 'mesaj': 'Potansiyel müşteri eklendi.', 'id': row['id']})


@bp.route('/api/potansiyel/<int:pid>/convert', methods=['POST'])
@giris_gerekli
def api_potansiyel_convert(pid):
    """Potansiyel müşteriyi gerçek müşteriye çevir ve Cari Kart'a yönlendirme linki döndür."""
    pot = fetch_one("SELECT * FROM crm_leads WHERE id = %s", (pid,))
    if not pot:
        return jsonify({'ok': False, 'mesaj': 'Potansiyel müşteri bulunamadı.'}), 404

    ad_soyad = (pot.get('ad_soyad') or '').strip() or 'Yeni Müşteri'
    firma_adi = (pot.get('firma_adi') or '').strip()
    musteri_adi = firma_adi or ad_soyad
    telefon = pot.get('telefon')
    notes_lines = []
    if pot.get('hizmet_turu'):
        notes_lines.append(f"İlgilendiği hizmet: {pot['hizmet_turu']}")
    if pot.get('sektor'):
        notes_lines.append(f"Sektör: {pot['sektor']}")
    if pot.get('notlar'):
        notes_lines.append(f"Lead notları: {pot['notlar']}")
    notes_text = "\n".join(notes_lines) if notes_lines else None

    yeni = execute_returning(
        """INSERT INTO customers (name, phone, notes, created_at)
               VALUES (%s,%s,%s,NOW())
               RETURNING id""",
        (musteri_adi, telefon, notes_text),
    )
    mid = yeni['id']

    # Lead durumunu güncelle (kazanıldı)
    execute(
        "UPDATE crm_leads SET lead_durumu = %s WHERE id = %s",
        ('Kazanıldı', pid),
    )

    url = url_for('cari_kart.index', mid=mid)
    return jsonify({'ok': True, 'mesaj': 'Sözleşme süreci için Cari Kart açıldı.', 'mid': mid, 'cari_kart_url': url})


@bp.route('/api/potansiyel/pending')
@giris_gerekli
def api_potansiyel_pending():
    """Dashboard'da gösterilecek 'Geri Dönüş Bekleyenler' listesi.

    - lead_durumu != Kazanıldı/Kaybedildi
    - takip_tarihi <= bugün
    """
    bugun = date.today()
    rows = fetch_all(
        """SELECT id, ad_soyad, telefon, hizmet_turu, lead_durumu, takip_tarihi
               FROM crm_leads
              WHERE takip_tarihi IS NOT NULL
                AND takip_tarihi <= %s
                AND LOWER(COALESCE(lead_durumu,'')) NOT IN ('kazanıldı','kazanildi','kaybedildi')
              ORDER BY takip_tarihi ASC, id DESC""",
        (bugun,),
    )
    out = []
    for r in rows or []:
        ad = (r.get('ad_soyad') or '').strip()
        ilk = ad.split()[0] if ad else 'Merhaba'
        mesaj = f"{ilk} Bey selamlar, BestOffice'deki kahve davetimiz hala geçerli, kampanya bitmeden bir daha görüşelim mi?"
        tel_raw = (r.get('telefon') or '').strip()
        num = ''.join(ch for ch in tel_raw if ch.isdigit())
        if num.startswith('0'):
            num = '90' + num[1:]
        elif num and not num.startswith('90'):
            num = '90' + num
        whatsapp_url = f"https://wa.me/{num}?text=" + urllib.parse.quote(mesaj) if num else ''
        r['whatsapp_url'] = whatsapp_url
        r['mesaj'] = mesaj
        out.append(r)
    return jsonify(out)


def _api_musteriler_row_json(row):
    """Tarih ve boolean alanlarını /giris/api/musteriler JSON çıktısı için düzleştir."""
    if not row:
        return row
    r = dict(row)
    r["rent_start_date"] = _musteri_serialize_val(r.get("rent_start_date"))
    r["kyc_sozlesme_tarihi"] = _musteri_serialize_val(r.get("kyc_sozlesme_tarihi"))
    ia = r.get("is_active")
    r["is_active"] = True if ia is None else bool(ia)
    dr = r.get("durum")
    r["durum"] = str(dr).strip() if dr is not None else ""
    return r


@bp.route('/api/musteriler')
@giris_gerekli
def api_musteriler():
    """Müşteri listesi - AJAX; ünvan, müşteri adı, vergi, adres, yetkili iletişim, KYC alanları."""
    arama = (request.args.get('q') or '').strip()
    base = (
        "SELECT id, name, musteri_adi, musteri_no, tax_number, phone, email, office_code, "
        "NULLIF(TRIM(COALESCE(durum, '')), '') AS durum, "
        "COALESCE(is_active, TRUE) AS is_active, "
        "rent_start_date, "
        "(SELECT mk.sozlesme_tarihi FROM musteri_kyc mk WHERE mk.musteri_id = customers.id "
        "ORDER BY mk.id DESC NULLS LAST LIMIT 1) AS kyc_sozlesme_tarihi "
        "FROM customers "
    )
    # Autocomplete dropdown için 100 yeterli; dar aramada zaten eşleşenler üste gelir.
    limit_n = 100 if arama else 1000
    if not arama:
        rows = fetch_all(base + f"ORDER BY name LIMIT {limit_n}")
    else:
        # Boşlukla ayrılmış her kelime tüm alanlarda aranır, kelimeler AND ile birleştirilir.
        # «Mehmet Erdoğdu» ve «Erdoğdu Mehmet» aynı kartı bulur.
        w, p = customers_arama_sql_params_giris_genis_tokens(arama, "")
        rows = fetch_all(
            base + f"WHERE {w} ORDER BY name LIMIT {limit_n}",
            p,
        )
        # Geniş WHERE bazen e-postayı kaçırır (@, normalizasyon); boşsa sadece e-posta kolonlarında düz ILIKE dene
        if (not rows) and ("@" in arama):
            pat = musteri_arama_ilike_pattern_email_duz(arama)
            fb = (
                "TRIM(COALESCE(email, '')) ILIKE %s ESCAPE '\\' OR EXISTS ("
                " SELECT 1 FROM musteri_kyc mk WHERE mk.musteri_id = customers.id AND ("
                " TRIM(COALESCE(mk.email, '')) ILIKE %s ESCAPE '\\'"
                " OR TRIM(COALESCE(mk.yetkili_email, '')) ILIKE %s ESCAPE '\\'"
                "))"
            )
            rows = fetch_all(
                base + f"WHERE ({fb}) ORDER BY name LIMIT {limit_n}",
                (pat, pat, pat),
            )
    out = [_api_musteriler_row_json(r) for r in (rows or [])]
    return jsonify(out)


def _musteri_serialize_val(v):
    """Tarih/sayı alanlarını JSON uyumlu string yap."""
    if v is None:
        return ""
    if hasattr(v, "isoformat"):
        return v.isoformat()[:10] if v else ""
    return str(v).strip() if v else ""


def _parse_kapanis_tarihi(s):
    """Formdan gelen kapanış tarihi (YYYY-MM-DD veya GG.AA.YYYY)."""
    if not s:
        return None
    s = str(s).strip()[:10]
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def _normalize_musteri_durum_kapanis(data):
    """durum: aktif|pasif; pasif değilse kapanis_tarihi None."""
    dr = (data.get("durum") or "aktif").strip().lower()
    if dr not in ("aktif", "pasif"):
        dr = "aktif"
    kap = _parse_kapanis_tarihi(data.get("kapanis_tarihi")) if dr == "pasif" else None
    return dr, kap


def _normalize_kapanis_sonrasi_borc_ay(data, durum: str | None = None):
    """Pasif müşteri için kapanıştan sonra ek borç ayı: 1-12, aksi halde None (Hepsi)."""
    dr = (durum or data.get("durum") or "aktif").strip().lower()
    if dr != "pasif":
        return None
    raw = data.get("kapanis_sonrasi_borc_ay")
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    try:
        val = int(s)
    except (TypeError, ValueError):
        return None
    return val if 1 <= val <= 12 else None


def _aylik_grid_effective_bitis(kyc: dict, bit: date | None) -> date | None:
    """Pasif müşteride kapanış + ek ay seçimi varsa, sözleşme bitişini buna göre kısaltır.

    Kural: kapanış ayı DAHIL sayılır.
    Örn. kapanış=2026-03-01, ek_ay=3 -> borçlanacak aylar: Mart, Nisan, Mayıs.
    """
    src = dict(kyc or {})
    durum = str(src.get("durum") or src.get("musteri_durum") or "").strip().lower()
    if durum != "pasif":
        return bit
    kap = _aylik_grid_coerce_date(src.get("kapanis_tarihi"))
    if not kap:
        return bit
    try:
        ek_ay = int(src.get("kapanis_sonrasi_borc_ay") or 0)
    except (TypeError, ValueError):
        ek_ay = 0
    if ek_ay <= 0:
        return bit
    # bitiş sınırı dışlayıcı kullanılıyor (donemBas >= bit -> gösterme),
    # bu yüzden kapanış ayı dahil N ay için +ek_ay yeterlidir.
    sinir = _add_months(date(kap.year, kap.month, 1), ek_ay)
    if bit is None:
        return sinir
    return min(bit, sinir)


@bp.route("/api/hizmet-turleri", methods=["GET", "POST"])
@giris_gerekli
def api_hizmet_turleri():
    """Hizmet türü listesi (GET) veya yeni tür ekleme (POST)."""
    ensure_hizmet_turleri_table()
    if request.method == "GET":
        cached = simple_cache_get(CACHE_KEY_HIZMET_TURLERI, CACHE_TTL_SEC)
        if cached is not None:
            return jsonify(cached)
        rows = fetch_all("SELECT id, ad FROM hizmet_turleri ORDER BY sira NULLS LAST, ad")
        payload = {"ok": True, "turler": [{"id": r["id"], "ad": r["ad"]} for r in (rows or [])]}
        simple_cache_set(CACHE_KEY_HIZMET_TURLERI, payload)
        return jsonify(payload)
    data = request.get_json(silent=True) or {}
    act = (data.get("action") or "").strip().lower()
    if act in ("sil_hizmet_turu", "delete_hizmet_turu", "delete"):
        return _api_hizmet_turu_sil_json(data)
    if act in ("guncelle_hizmet_turu", "update_hizmet_turu", "update"):
        return _api_hizmet_turu_guncelle_json(data)
    ad = (data.get("ad") or "").strip()
    if not ad:
        return jsonify({"ok": False, "mesaj": "Hizmet türü adı boş olamaz."}), 400
    if len(ad) > 200:
        return jsonify({"ok": False, "mesaj": "En fazla 200 karakter girebilirsiniz."}), 400
    mx = fetch_one("SELECT COALESCE(MAX(sira), 0) + 1 AS n FROM hizmet_turleri")
    next_sira = int(mx["n"] or 1) if mx else 1
    ins = execute_returning(
        "INSERT INTO hizmet_turleri (ad, sira) VALUES (%s, %s) ON CONFLICT (ad) DO NOTHING RETURNING id, ad",
        (ad, next_sira),
    )
    if not ins:
        ins = fetch_one("SELECT id, ad FROM hizmet_turleri WHERE ad = %s", (ad,))
    rows = fetch_all("SELECT id, ad FROM hizmet_turleri ORDER BY sira NULLS LAST, ad")
    simple_cache_invalidate(CACHE_KEY_HIZMET_TURLERI)
    return jsonify(
        {
            "ok": True,
            "turler": [{"id": r["id"], "ad": r["ad"]} for r in (rows or [])],
            "secilen": {"id": ins.get("id"), "ad": ins.get("ad")} if ins else None,
        }
    )


@bp.route("/api/hizmet-turleri/guncelle", methods=["POST", "PUT"])
@giris_gerekli
def api_hizmet_turu_guncelle():
    """Hizmet türü adını güncelle; müşteri/KYC kayıtlarındaki metin de eşlenir."""
    data = request.get_json(silent=True) or {}
    return _api_hizmet_turu_guncelle_json(data)


def _api_hizmet_turu_guncelle_json(data):
    """Hizmet türü adını güncelle; müşteri/KYC kayıtlarındaki metin de eşlenir."""
    ensure_hizmet_turleri_table()
    try:
        tid = int(data.get("id") or 0)
    except (TypeError, ValueError):
        tid = 0
    yeni_ad = (data.get("ad") or "").strip()
    if tid <= 0:
        return jsonify({"ok": False, "mesaj": "Geçersiz hizmet türü."}), 400
    if not yeni_ad:
        return jsonify({"ok": False, "mesaj": "Hizmet türü adı boş olamaz."}), 400
    if len(yeni_ad) > 200:
        return jsonify({"ok": False, "mesaj": "En fazla 200 karakter girebilirsiniz."}), 400
    row = fetch_one("SELECT id, ad FROM hizmet_turleri WHERE id = %s", (tid,))
    if not row:
        return jsonify({"ok": False, "mesaj": "Hizmet türü bulunamadı."}), 404
    eski_ad = (row.get("ad") or "").strip()
    if turkish_lower(eski_ad) == turkish_lower(yeni_ad):
        return jsonify({"ok": True, "id": tid, "ad": eski_ad})
    dup = fetch_one(
        "SELECT id FROM hizmet_turleri WHERE lower(trim(ad)) = lower(trim(%s)) AND id <> %s LIMIT 1",
        (yeni_ad, tid),
    )
    if dup:
        return jsonify({"ok": False, "mesaj": "Bu hizmet türü adı zaten var."}), 400
    execute("UPDATE hizmet_turleri SET ad = %s WHERE id = %s", (yeni_ad, tid))
    if eski_ad and eski_ad != yeni_ad:
        execute(
            "UPDATE customers SET hizmet_turu = %s WHERE TRIM(COALESCE(hizmet_turu, '')) = %s",
            (yeni_ad, eski_ad),
        )
        execute(
            "UPDATE musteri_kyc SET hizmet_turu = %s WHERE TRIM(COALESCE(hizmet_turu, '')) = %s",
            (yeni_ad, eski_ad),
        )
    rows = fetch_all("SELECT id, ad FROM hizmet_turleri ORDER BY sira NULLS LAST, ad")
    simple_cache_invalidate(CACHE_KEY_HIZMET_TURLERI)
    return jsonify(
        {
            "ok": True,
            "id": tid,
            "ad": yeni_ad,
            "eski_ad": eski_ad,
            "turler": [{"id": r["id"], "ad": r["ad"]} for r in (rows or [])],
        }
    )


@bp.route("/api/hizmet-turleri/sil", methods=["POST", "DELETE"])
@giris_gerekli
def api_hizmet_turu_sil():
    """Hizmet türünü listeden kaldır (müşteri kartındaki kayıtlı değer korunur)."""
    data = request.get_json(silent=True) or {}
    return _api_hizmet_turu_sil_json(data)


def _api_hizmet_turu_sil_json(data):
    """Hizmet türünü listeden kaldır (müşteri kartındaki kayıtlı değer korunur)."""
    ensure_hizmet_turleri_table()
    try:
        tid = int(data.get("id") or 0)
    except (TypeError, ValueError):
        tid = 0
    if tid <= 0:
        ad_ara = (data.get("ad") or "").strip()
        if ad_ara:
            row = fetch_one(
                "SELECT id, ad FROM hizmet_turleri WHERE TRIM(ad) = %s LIMIT 1",
                (ad_ara,),
            )
            if row:
                tid = int(row["id"])
    if tid <= 0:
        return jsonify({"ok": False, "mesaj": "Geçersiz hizmet türü."}), 400
    row = fetch_one("SELECT id, ad FROM hizmet_turleri WHERE id = %s", (tid,))
    if not row:
        return jsonify({"ok": False, "mesaj": "Hizmet türü bulunamadı."}), 404
    execute("DELETE FROM hizmet_turleri WHERE id = %s", (tid,))
    rows = fetch_all("SELECT id, ad FROM hizmet_turleri ORDER BY sira NULLS LAST, ad")
    simple_cache_invalidate(CACHE_KEY_HIZMET_TURLERI)
    return jsonify(
        {
            "ok": True,
            "id": tid,
            "ad": row.get("ad"),
            "turler": [{"id": r["id"], "ad": r["ad"]} for r in (rows or [])],
        }
    )


def _duzenli_fatura_kod_slug(etiket: str) -> str:
    s = turkish_lower((etiket or "").strip())
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return (s or "ozel")[:80]


def _parse_grup2_secimleri_from_request(data) -> list:
    """Form / JSON: grup2_secimleri dizi veya JSON dizisi (string)."""
    raw = data.get("grup2_secimleri")
    if raw is None:
        return []
    if isinstance(raw, list):
        return [str(x).strip() for x in raw if str(x).strip()]
    if isinstance(raw, str):
        s = raw.strip()
        if not s:
            return []
        if s.startswith("["):
            try:
                j = json.loads(s)
                if isinstance(j, list):
                    return [str(x).strip() for x in j if str(x).strip()]
            except Exception:
                return []
        return [s]
    return []


def _parse_pg_text_array_grup2(val):
    """customers.grup2_secimleri: list/tuple; '{a,b}' PG metni; '["a"]' JSON; bazı sürücülerde iterable."""
    if val is None:
        return []
    if isinstance(val, (list, tuple)):
        return [str(x).strip() for x in val if str(x).strip()]
    if not isinstance(val, (str, bytes, bytearray)) and hasattr(val, "__iter__"):
        try:
            return [str(x).strip() for x in val if str(x).strip()]
        except (TypeError, ValueError):
            pass
    s = str(val).strip()
    if not s or s in ("{}", "[]") or s.lower() in ("null", "none"):
        return []
    if s.startswith("[") and s.endswith("]"):
        try:
            j = json.loads(s)
            if isinstance(j, list):
                return [str(x).strip() for x in j if str(x).strip()]
        except Exception:
            pass
    if s.startswith("{") and s.endswith("}"):
        inner = s[1:-1].strip()
        if not inner:
            return []
        return [p.strip().strip('"') for p in inner.split(",") if p.strip()]
    return [s] if s else []


def _filter_grup2_slugs_known(slugs: list) -> list:
    if not slugs:
        return []
    ensure_grup2_etiketleri_table()
    rows = fetch_all(
        "SELECT slug FROM grup2_etiketleri WHERE COALESCE(aktif, TRUE) AND slug = ANY(%s)",
        (list(slugs),),
    )
    allowed = {r["slug"] for r in (rows or [])}
    return [s for s in slugs if s in allowed]


_VD_LIST_CACHE = None


def _tr_vergi_daireleri_list():
    """data/tr_vergi_daireleri.json — form datalist için (tam resmi liste değil)."""
    global _VD_LIST_CACHE
    if _VD_LIST_CACHE is not None:
        return _VD_LIST_CACHE
    p = Path(__file__).resolve().parent.parent / "data" / "tr_vergi_daireleri.json"
    try:
        raw = p.read_text(encoding="utf-8")
        j = json.loads(raw)
        if isinstance(j, list):
            seen = set()
            out = []
            for x in j:
                s = str(x).strip()
                if not s:
                    continue
                k = turkish_lower(s)
                if k in seen:
                    continue
                seen.add(k)
                out.append(s)
            _VD_LIST_CACHE = sorted(out, key=turkish_lower)
        else:
            _VD_LIST_CACHE = ["Kavaklıdere", "Çankaya", "Ostim"]
    except Exception:
        _VD_LIST_CACHE = ["Kavaklıdere", "Çankaya", "Ostim"]
    return _VD_LIST_CACHE


@bp.route("/api/vergi-daireleri", methods=["GET"])
@giris_gerekli
def api_vergi_daireleri():
    """Giriş formu vergi dairesi datalist / arama (aynı origin, DB gerektirmez)."""
    q = (request.args.get("q") or "").strip()
    rows = list(_tr_vergi_daireleri_list())
    if q:
        ql = turkish_lower(q)
        rows = [x for x in rows if ql in turkish_lower(x)]
    try:
        lim = int(request.args.get("limit", 400))
    except (TypeError, ValueError):
        lim = 400
    lim = max(1, min(lim, 600))
    return jsonify({"ok": True, "liste": rows[:lim]})


@bp.route("/api/grup2-etiketleri", methods=["GET", "POST", "PUT", "DELETE"])
@giris_gerekli
def api_grup2_etiketleri():
    """Grup 2 etiket listesi (GET), yeni etiket (POST), güncelleme/silme (PUT/DELETE veya POST+action)."""
    _log = logging.getLogger(__name__)
    try:
        ensure_grup2_etiketleri_table()
        if request.method == "GET":
            ensure_grup2_bizim_hesap_into_array()
            rows = fetch_all(
                """
                SELECT id, slug, etiket
                FROM grup2_etiketleri
                WHERE COALESCE(aktif, TRUE)
                ORDER BY
                    CASE slug
                        WHEN 'bizim_hesap' THEN 0
                        WHEN 'vergi_dairesi' THEN 1
                        WHEN 'vergi_dairesi_terk' THEN 2
                        ELSE 3
                    END,
                    sira NULLS LAST,
                    etiket
                """
            )
            return jsonify(
                {
                    "ok": True,
                    "etiketler": [
                        {"id": r["id"], "slug": r["slug"], "etiket": r["etiket"]} for r in (rows or [])
                    ],
                }
            )
        q = request.args.to_dict(flat=True) or {}
        frm = request.form.to_dict(flat=True) or {}
        # Önce ham gövde: get_json / Content-Type zinciri bazen {} döndürüyor; sil isteği put sanılıp "Etiket adı boş olamaz" oluyordu.
        raw = request.get_data(cache=True) or b""
        j = {}
        if raw.strip():
            try:
                p = json.loads(raw.decode("utf-8", errors="replace"))
                if isinstance(p, dict):
                    j = p
            except Exception:
                j = {}
        if not j and request.method in ("POST", "PUT", "PATCH", "DELETE"):
            gj = request.get_json(silent=True, force=True)
            if isinstance(gj, dict):
                j = gj
        data = {**q, **frm, **j}
        op = None
        act = ""
        if request.method == "PUT":
            op = "put"
        elif request.method == "DELETE":
            op = "delete"
        elif request.method == "POST":
            act = str((j.get("action") if isinstance(j, dict) else None) or data.get("action") or "").strip().lower()
            if act in ("update_etiket", "put"):
                op = "put"
            elif act in ("delete_etiket", "delete"):
                op = "delete"
            elif str(data.get("slug") or "").strip() and not str(data.get("etiket") or "").strip():
                op = "delete"
            elif str(data.get("slug") or "").strip() and str(data.get("etiket") or "").strip():
                op = "put"
            else:
                op = "create"
        _log.info(
            "api_grup2_etiketleri %s op=%s act=%r keys=%s ct=%s raw_len=%s",
            request.method,
            op,
            act if request.method == "POST" else None,
            list(data.keys()),
            (request.content_type or "")[:80],
            len(raw or b""),
        )
        if op == "put":
            slug = (data.get("slug") or "").strip()
            etiket = (data.get("etiket") or "").strip()
            if not slug:
                return jsonify({"ok": False, "mesaj": "Slug zorunludur."}), 400
            if not etiket:
                return jsonify({"ok": False, "mesaj": "Etiket adı boş olamaz."}), 400
            if len(etiket) > 200:
                return jsonify({"ok": False, "mesaj": "En fazla 200 karakter girebilirsiniz."}), 400
            row = fetch_one(
                "SELECT id, slug, etiket FROM grup2_etiketleri WHERE slug = %s AND COALESCE(aktif, TRUE) LIMIT 1",
                (slug,),
            )
            if not row:
                return jsonify({"ok": False, "mesaj": "Etiket bulunamadı."}), 400
            dup = fetch_one(
                """
                SELECT id FROM grup2_etiketleri
                WHERE COALESCE(aktif, TRUE)
                  AND lower(trim(etiket)) = lower(trim(%s))
                  AND slug <> %s
                LIMIT 1
                """,
                (etiket, slug),
            )
            if dup:
                return jsonify({"ok": False, "mesaj": "Bu etiket adı zaten kullanılıyor."}), 400
            execute("UPDATE grup2_etiketleri SET etiket = %s WHERE slug = %s", (etiket, slug))
            return jsonify({"ok": True, "slug": slug, "etiket": etiket})
        if act == "delete_etiket":
            slug = data.get("slug") or ""
            if not slug:
                return jsonify({"ok": False, "mesaj": "Slug boş"}), 400
            execute("UPDATE grup2_etiketleri SET aktif = FALSE WHERE slug = %s", (slug,))
            return jsonify({"ok": True})
        if request.method != "POST" or op != "create":
            return jsonify({"ok": False, "mesaj": "Geçersiz istek."}), 405
        etiket = (data.get("etiket") or "").strip()
        if not etiket:
            err = {"ok": False, "mesaj": "Etiket adı boş olamaz."}
            if current_app and current_app.debug:
                err["debug"] = {
                    "sunulan_op": op,
                    "anahtarlar": list(data.keys()),
                    "ipucu": "Sil: action=delete_etiket + slug; güncelle: action=update_etiket + slug + etiket; yeni: yalnızca etiket.",
                }
            _log.warning("grup2 create reddi: op=%s data_keys=%s", op, list(data.keys()))
            return jsonify(err), 400
        if len(etiket) > 200:
            return jsonify({"ok": False, "mesaj": "En fazla 200 karakter girebilirsiniz."}), 400
        ex = fetch_one(
            "SELECT id, slug, etiket FROM grup2_etiketleri WHERE lower(trim(etiket)) = lower(trim(%s)) LIMIT 1",
            (etiket,),
        )
        if ex:
            return jsonify(
                {
                    "ok": True,
                    "slug": ex["slug"],
                    "etiket": ex["etiket"],
                    "mevcut": True,
                }
            )
        rows_all = fetch_all("SELECT slug FROM grup2_etiketleri")
        slug_set = {r["slug"] for r in (rows_all or [])}
        base = _duzenli_fatura_kod_slug(etiket)
        slug_out = None
        for n in range(0, 200):
            cand = (base if n == 0 else f"{base}_{n}")[:80]
            if cand in slug_set:
                continue
            mx = fetch_one("SELECT COALESCE(MAX(sira), 0) + 1 AS n FROM grup2_etiketleri")
            next_sira = int(mx["n"] or 1) if mx else 1
            try:
                execute(
                    "INSERT INTO grup2_etiketleri (slug, etiket, sira) VALUES (%s, %s, %s)",
                    (cand, etiket, next_sira),
                )
                slug_out = cand
                break
            except Exception as ins_err:
                err_l = str(ins_err).lower()
                if "unique" in err_l or "duplicate" in err_l:
                    slug_set.add(cand)
                    continue
                raise
        if not slug_out:
            return jsonify({"ok": False, "mesaj": "Slug üretilemedi."}), 400
        return jsonify({"ok": True, "slug": slug_out, "etiket": etiket})
    except Exception as e:
        _log.exception("api_grup2_etiketleri %s", request.method)
        return jsonify({"ok": False, "mesaj": str(e)}), 500


@bp.route("/api/grup2-etiket-guncelle", methods=["GET", "POST"])
@giris_gerekli
def api_grup2_etiket_guncelle():
    try:
        ensure_grup2_etiketleri_table()
        data = request.get_json(silent=True) or {}
        slug = data.get("slug") or request.args.get("slug") or ""
        etiket = data.get("etiket") or request.args.get("etiket") or ""
        slug = (slug or "").strip()
        etiket = (etiket or "").strip()
        if not slug:
            return jsonify({"ok": False, "mesaj": "Slug zorunludur."}), 400
        if not etiket:
            return jsonify({"ok": False, "mesaj": "Etiket adı boş olamaz."}), 400
        if len(etiket) > 200:
            return jsonify({"ok": False, "mesaj": "En fazla 200 karakter girebilirsiniz."}), 400
        row = fetch_one(
            "SELECT id, slug, etiket FROM grup2_etiketleri WHERE slug = %s AND COALESCE(aktif, TRUE) LIMIT 1",
            (slug,),
        )
        if not row:
            return jsonify({"ok": False, "mesaj": "Etiket bulunamadı."}), 400
        dup = fetch_one(
            """
            SELECT id FROM grup2_etiketleri
            WHERE COALESCE(aktif, TRUE)
              AND lower(trim(etiket)) = lower(trim(%s))
              AND slug <> %s
            LIMIT 1
            """,
            (etiket, slug),
        )
        if dup:
            return jsonify({"ok": False, "mesaj": "Bu etiket adı zaten kullanılıyor."}), 400
        execute("UPDATE grup2_etiketleri SET etiket = %s WHERE slug = %s", (etiket, slug))
        return jsonify({"ok": True, "slug": slug, "etiket": etiket})
    except Exception as e:
        logging.getLogger(__name__).exception("api_grup2_etiket_guncelle")
        return jsonify({"ok": False, "mesaj": str(e)}), 500


@bp.route("/api/grup2-etiket-sil", methods=["GET", "POST"])
@giris_gerekli
def api_grup2_etiket_sil():
    try:
        ensure_grup2_etiketleri_table()
        data = request.get_json(silent=True) or {}
        slug = (data.get("slug") or request.args.get("slug") or "").strip()
        if not slug:
            return jsonify({"ok": False, "mesaj": "Slug boş"}), 400
        # PostgreSQL: aktif BOOLEAN — 0 değil FALSE kullanılmalı (aksi halde 500 / istemci tarafında kırık yanıt).
        execute("UPDATE grup2_etiketleri SET aktif = FALSE WHERE slug = %s", (slug,))
        execute(
            """
            UPDATE customers
            SET grup2_secimleri = array_remove(COALESCE(grup2_secimleri, ARRAY[]::text[]), %s)
            WHERE %s = ANY(COALESCE(grup2_secimleri, ARRAY[]::text[]))
            """,
            (slug, slug),
        )
        return jsonify({"ok": True, "slug": slug})
    except Exception as e:
        logging.getLogger(__name__).exception("api_grup2_etiket_sil")
        return jsonify({"ok": False, "mesaj": str(e)}), 500


def _duzenli_fatura_next_sira() -> int:
    mx = fetch_one("SELECT COALESCE(MAX(sira), 0) + 1 AS n FROM duzenli_fatura_secenekleri")
    return int(mx["n"] or 1) if mx else 1


def _duzenli_fatura_ekle(etiket: str) -> dict:
    """Aynı etiket (Türkçe büyük/küçük duyarsız) varsa mevcut satırı döndür; yoksa yeni kod ile ekler."""
    etiket_c = (etiket or "").strip()
    el = turkish_lower(etiket_c)
    rows_all = fetch_all("SELECT kod, etiket FROM duzenli_fatura_secenekleri")
    kod_set = {r["kod"] for r in (rows_all or [])}
    for r in rows_all or []:
        if turkish_lower((r.get("etiket") or "").strip()) == el:
            return {"kod": r["kod"], "etiket": r["etiket"]}
    base = _duzenli_fatura_kod_slug(etiket_c)
    for n in range(0, 200):
        kod = (base if n == 0 else f"{base}_{n}")[:80]
        if kod in kod_set:
            continue
        execute(
            "INSERT INTO duzenli_fatura_secenekleri (kod, etiket, sira) VALUES (%s, %s, %s)",
            (kod, etiket_c, _duzenli_fatura_next_sira()),
        )
        return {"kod": kod, "etiket": etiket_c}
    raise RuntimeError("Düzenli fatura kodu üretilemedi (çok fazla çakışma).")


@bp.route("/api/duzenli-fatura-secenekleri", methods=["GET", "POST"])
@giris_gerekli
def api_duzenli_fatura_secenekleri():
    """Düzenli Fatura açılır listesi (GET) veya yeni senaryo (POST)."""
    ensure_duzenli_fatura_secenekleri_table()
    if request.method == "GET":
        cached = simple_cache_get(CACHE_KEY_DUZENLI_FATURA, CACHE_TTL_SEC)
        if cached is not None:
            return jsonify(cached)
        rows = fetch_all(
            "SELECT kod, etiket FROM duzenli_fatura_secenekleri ORDER BY sira NULLS LAST, etiket"
        )
        payload = {
            "ok": True,
            "secenekler": [{"kod": r["kod"], "etiket": r["etiket"]} for r in (rows or [])],
        }
        simple_cache_set(CACHE_KEY_DUZENLI_FATURA, payload)
        return jsonify(payload)
    data = request.get_json(silent=True) or {}
    etiket = (data.get("etiket") or "").strip()
    if not etiket:
        return jsonify({"ok": False, "mesaj": "Senaryo adı boş olamaz."}), 400
    if len(etiket) > 200:
        return jsonify({"ok": False, "mesaj": "En fazla 200 karakter girebilirsiniz."}), 400
    try:
        secilen = _duzenli_fatura_ekle(etiket)
    except Exception as e:
        logging.exception("duzenli_fatura_ekle")
        return jsonify({"ok": False, "mesaj": str(e) or "Eklenemedi"}), 500
    rows = fetch_all(
        "SELECT kod, etiket FROM duzenli_fatura_secenekleri ORDER BY sira NULLS LAST, etiket"
    )
    simple_cache_invalidate(CACHE_KEY_DUZENLI_FATURA)
    return jsonify(
        {
            "ok": True,
            "secenekler": [{"kod": r["kod"], "etiket": r["etiket"]} for r in (rows or [])],
            "secilen": secilen,
        }
    )


@bp.route("/api/musteri-komsu", methods=["GET"])
@giris_gerekli
def api_musteri_komsu():
    """
    Ünvana göre sıralı listede (aktif müşteriler) önceki / sonraki müşteri id.
    Giriş ekranında hızlı gezinme için.
    """
    mid = request.args.get("musteri_id", type=int)
    yon = (request.args.get("yon") or "").strip().lower()
    if not mid or yon not in ("onceki", "sonraki"):
        return jsonify({"ok": False, "mesaj": "musteri_id ve yon=onceki|sonraki gerekli."}), 400

    aktif_only = str(request.args.get("aktif_only", "1")).lower() not in ("0", "false", "no", "hayir")

    def _komsu_row(where_sql):
        return fetch_one(
            f"""
            WITH o AS (
                SELECT id, name,
                    LAG(id) OVER (ORDER BY LOWER(TRIM(COALESCE(name, ''))), id) AS prev_id,
                    LEAD(id) OVER (ORDER BY LOWER(TRIM(COALESCE(name, ''))), id) AS next_id
                FROM customers
                WHERE {where_sql}
            )
            SELECT prev_id, next_id FROM o WHERE id = %s
            """,
            (mid,),
        )

    where_aktif = "LOWER(COALESCE(durum, 'aktif')) <> 'pasif'"
    row = _komsu_row(where_aktif) if aktif_only else _komsu_row("TRUE")
    if not row and aktif_only:
        row = _komsu_row("TRUE")
    if not row:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404

    hedef = row.get("prev_id") if yon == "onceki" else row.get("next_id")
    if not hedef:
        return jsonify(
            {
                "ok": True,
                "bos": True,
                "mesaj": "Bu yönde başka müşteri yok.",
            }
        )

    t = fetch_one("SELECT id, name FROM customers WHERE id = %s", (int(hedef),))
    if not t:
        return jsonify({"ok": False, "mesaj": "Komşu kayıt bulunamadı."}), 404

    return jsonify(
        {
            "ok": True,
            "id": int(t["id"]),
            "name": (t.get("name") or "").strip(),
        }
    )


@bp.route('/api/musteri/<int:mid>')
@giris_gerekli
def api_musteri_detay(mid):
    """Tek müşteri tüm alanları - customers + son musteri_kyc birleşik; forma doldurmak için."""
    _od_arg = str(request.args.get("odemeler", "1") or "").strip().lower()
    _hesapla_tahsilat_ozet = _od_arg not in ("0", "false", "hayir", "no")
    force = str(request.args.get("force") or "").lower() in ("1", "true", "yes", "on")
    cache_key = f"{int(mid)}:{1 if _hesapla_tahsilat_ozet else 0}"
    if not force:
        try:
            now = time.time()
            ttl = 35.0
            cache = getattr(api_musteri_detay, "_cache", None)
            if cache is None:
                cache = {}
                setattr(api_musteri_detay, "_cache", cache)
            hit = cache.get(cache_key)
            if hit and (now - float(hit.get("ts") or 0)) <= ttl and hit.get("payload"):
                return jsonify(hit["payload"])
        except Exception:
            pass

    ensure_customers_bizim_hesap()
    ensure_customers_grup2_secimleri()
    ensure_customers_kapanis_sonrasi_borc_ay()
    ensure_grup2_etiketleri_table()
    ensure_grup2_bizim_hesap_into_array()
    # Tek round-trip: customers + en son musteri_kyc + tahsilat toplamı birlikte.
    combined = fetch_all(
        """
        SELECT 'c' AS _src, to_jsonb(c) AS data
        FROM customers c
        WHERE c.id = %s
        UNION ALL
        SELECT 'k' AS _src, to_jsonb(k) AS data
        FROM (
            SELECT * FROM musteri_kyc
            WHERE musteri_id = %s
            ORDER BY id DESC
            LIMIT 1
        ) k
        UNION ALL
        SELECT 't' AS _src, jsonb_build_object('toplam', COALESCE(SUM(tutar), 0)) AS data
        FROM tahsilatlar
        WHERE %s::boolean AND (musteri_id = %s OR customer_id = %s)
        """,
        (mid, mid, _hesapla_tahsilat_ozet, mid, mid),
    ) or []
    row = None
    kyc = None
    _tahsilat_toplam_pre = 0.0
    for r in combined:
        src = r.get("_src")
        data = r.get("data") or {}
        if isinstance(data, str):
            try:
                import json as _json
                data = _json.loads(data) or {}
            except Exception:
                data = {}
        if src == "c":
            row = data
        elif src == "k":
            kyc = data
        elif src == "t":
            try:
                _tahsilat_toplam_pre = float(data.get("toplam") or 0)
            except Exception:
                _tahsilat_toplam_pre = 0.0
    if not row:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    out = {}
    for k, v in row.items():
        if k == "grup2_secimleri":
            continue
        out[k] = _musteri_serialize_val(v)
    out["is_group"] = bool(row.get("is_group"))
    out["parent_cari_id"] = ""
    if row.get("parent_id"):
        try:
            pstr = str(row.get("parent_id"))
            gmap = _giris_grup_uuid_id_haritasi()
            hit = gmap.get(pstr)
            if hit:
                out["parent_cari_id"] = str(hit.get("id"))
                out["parent_name"] = (hit.get("name") or "").strip()
        except Exception:
            pass
    if not kyc and out.get("yetkili_tcno"):
        out["yetkili_tc"] = out["yetkili_tcno"]
    # KYC alanlarını forma uyumlu anahtarlarla birleştir (KYC öncelikli)
    if kyc:
        out["musteri_adi"] = _musteri_serialize_val(out.get("musteri_adi")) or _musteri_serialize_val(
            kyc.get("musteri_adi")
        )
        out["name"] = out.get("name") or _musteri_serialize_val(kyc.get("sirket_unvani"))
        out["tax_number"] = out.get("tax_number") or _musteri_serialize_val(kyc.get("vergi_no"))
        out["vergi_dairesi"] = _musteri_serialize_val(kyc.get("vergi_dairesi")) or out.get("vergi_dairesi", "")
        out["mersis_no"] = _musteri_serialize_val(kyc.get("mersis_no"))
        out["mersis"] = out["mersis_no"]
        out["nace_kodu"] = _musteri_serialize_val(kyc.get("nace_kodu"))
        out["nace"] = out["nace_kodu"]
        out["yetkili_kisi"] = _musteri_serialize_val(kyc.get("yetkili_adsoyad"))
        out["yetkili_ad"] = out["yetkili_kisi"]
        out["yetkili_tc"] = _musteri_serialize_val(kyc.get("yetkili_tcno"))
        out["phone"] = out.get("phone") or _musteri_serialize_val(kyc.get("yetkili_tel"))
        out["phone2"] = _musteri_serialize_val(kyc.get("yetkili_tel2"))
        out["phone_kime"] = _musteri_serialize_val(kyc.get("yetkili_tel_aciklama"))
        out["phone2_kime"] = _musteri_serialize_val(kyc.get("yetkili_tel2_aciklama"))
        # Formdaki «Yetkili e-posta» alanı: KYC yetkili e-postası öncelikli (customers.email şirket yedeği olabilir)
        out["email"] = _musteri_serialize_val(kyc.get("yetkili_email")) or out.get("email")
        out["email_sirket"] = _musteri_serialize_val(kyc.get("email"))
        out["address"] = out.get("address") or _musteri_serialize_val(kyc.get("yeni_adres"))
        out["ev_adres"] = out.get("ev_adres") or _musteri_serialize_val(kyc.get("yetkili_ikametgah"))
        out["notes"] = out.get("notes") or _musteri_serialize_val(kyc.get("notlar"))
        out["hizmet_turu"] = _musteri_serialize_val(kyc.get("hizmet_turu"))
        ho = kyc.get("hazir_ofis_oda_no")
        if ho is not None and str(ho).strip() != "":
            out["hazir_ofis_oda_no"] = _musteri_serialize_val(ho)
        _df = _musteri_serialize_val(kyc.get("duzenli_fatura"))
        out["duzenli_fatura"] = (_df if _df else "duzenle")
        _od = _musteri_serialize_val(kyc.get("odeme_duzeni"))
        out["odeme_duzeni"] = (_od if _od else "aylik")
        out["odeme_duzeni_manuel"] = _musteri_serialize_val(kyc.get("odeme_duzeni_manuel")) or ""
        out["guncel_kira_bedeli"] = _musteri_serialize_val(kyc.get("aylik_kira"))
        out["ilk_kira_bedeli"] = out["guncel_kira_bedeli"]
        out["rent_start_date"] = _musteri_serialize_val(kyc.get("sozlesme_tarihi"))
        out["sozlesme_baslangic"] = out["rent_start_date"]
        out["sozlesme_bitis"] = _musteri_serialize_val(kyc.get("sozlesme_bitis"))
        out["kira_suresi_ay"] = _musteri_serialize_val(kyc.get("kira_suresi_ay"))
        out["kira_artis_tarihi"] = _musteri_serialize_val(kyc.get("kira_artis_tarihi"))
        out["ticaret_sicil"] = _musteri_serialize_val(kyc.get("ticaret_sicil_no"))
        out["kurulus_tarihi"] = _musteri_serialize_val(kyc.get("kurulus_tarihi"))
        out["faaliyet"] = _musteri_serialize_val(kyc.get("faaliyet_konusu"))
        out["onceki_adres"] = _musteri_serialize_val(kyc.get("eski_adres"))
        out["sube_merkez"] = _musteri_serialize_val(kyc.get("sube_merkez"))
        out["kira_nakit"] = bool(kyc.get("kira_nakit"))
        out["kira_banka"] = bool(kyc.get("kira_banka"))
        out["kira_nakit_tutar"] = _musteri_serialize_val(kyc.get("kira_nakit_tutar"))
        out["kira_banka_tutar"] = _musteri_serialize_val(kyc.get("kira_banka_tutar"))

    # Kaç ay ödeme yapıldı (tahsilat / aylık kira KDV dahil)
    try:
        aylik_kira = 0.0
        if out.get("guncel_kira_bedeli"):
            aylik_kira = float(str(out["guncel_kira_bedeli"]).replace(",", ".")) or 0.0
        elif out.get("ilk_kira_bedeli"):
            aylik_kira = float(str(out["ilk_kira_bedeli"]).replace(",", ".")) or 0.0
    except Exception:
        aylik_kira = 0.0
    try:
        kdv_oran = float(str(kyc.get("kdv_oran") or "20").replace(",", ".")) if kyc else 20.0
    except Exception:
        kdv_oran = 20.0
    kira_nakit_m = bool(kyc.get("kira_nakit")) if kyc else False
    split_ok, npay, bpay, _, _ = _kyc_karma_kira_paylari(kyc, aylik_kira) if kyc else (False, 0.0, 0.0, 0.0, 0.0)
    if kyc and split_ok:
        aylik_kdv_dahil = round(npay + bpay * (1 + kdv_oran / 100), 2)
    elif kira_nakit_m and aylik_kira > 0:
        aylik_kdv_dahil = round(aylik_kira, 2)
    else:
        aylik_kdv_dahil = round(aylik_kira * (1 + kdv_oran / 100), 2) if aylik_kira > 0 else 0.0

    odenen_ay_sayisi = 0
    kismi_odeme_var = False
    kismi_ay_eksik_tutar = 0.0  # Kısmi ödenen ayda kalan borç (kutuda gösterilecek)
    if _hesapla_tahsilat_ozet and aylik_kdv_dahil > 0:
        toplam_tahsilat = _tahsilat_toplam_pre
        if toplam_tahsilat > 0:
            odenen_ay_sayisi = int(toplam_tahsilat // aylik_kdv_dahil)
            kalan = toplam_tahsilat - (odenen_ay_sayisi * aylik_kdv_dahil)
            if 0 < kalan < aylik_kdv_dahil:
                kismi_odeme_var = True
                kismi_ay_eksik_tutar = round(aylik_kdv_dahil - kalan, 2)  # O aydan ne kadar eksik kaldı

    out["odenen_ay_sayisi"] = odenen_ay_sayisi
    out["odenen_tam_ay_sayisi"] = odenen_ay_sayisi
    out["kismi_odeme_var"] = kismi_odeme_var
    out["kismi_ay_eksik_tutar"] = kismi_ay_eksik_tutar
    out["toplam_tahsilat"] = round(_tahsilat_toplam_pre, 2) if _hesapla_tahsilat_ozet else 0.0

    if not out.get("hazir_ofis_oda_no") and row.get("hazir_ofis_oda_no") is not None:
        out["hazir_ofis_oda_no"] = _musteri_serialize_val(row.get("hazir_ofis_oda_no"))

    g2_raw = row.get("grup2_secimleri")
    g2_list = _parse_pg_text_array_grup2(g2_raw)
    if bool(row.get("bizim_hesap")) and "bizim_hesap" not in g2_list:
        g2_list = ["bizim_hesap"] + g2_list
    if not g2_list and bool(row.get("bizim_hesap")):
        g2_list = ["bizim_hesap"]
    g2_list = list(dict.fromkeys(g2_list))
    out["grup2_secimleri"] = g2_list
    out["bizim_hesap"] = "bizim_hesap" in g2_list

    payload = {"ok": True, "musteri": out}
    if not force:
        try:
            cache = getattr(api_musteri_detay, "_cache", None)
            if isinstance(cache, dict):
                cache[cache_key] = {"ts": time.time(), "payload": payload}
        except Exception:
            pass
    return jsonify(payload)


@bp.route("/api/hazir-ofis-durum", methods=["GET"])
@giris_gerekli
def api_hazir_ofis_durum():
    """200–230 odalar: dolu (aktif kart + atanmış oda) ve boş liste."""
    ensure_customers_hazir_ofis_oda()
    rows = fetch_all(
        """
        SELECT c.id,
               c.hazir_ofis_oda_no,
               COALESCE(NULLIF(TRIM(c.musteri_adi), ''), NULLIF(TRIM(c.name), ''), '—') AS firma_adi
        FROM customers c
        WHERE c.hazir_ofis_oda_no IS NOT NULL
          AND c.hazir_ofis_oda_no BETWEEN 200 AND 230
          AND COALESCE(c.is_active, TRUE) = TRUE
          AND LOWER(TRIM(COALESCE(c.durum, ''))) = 'aktif'
        ORDER BY c.hazir_ofis_oda_no
        """
    ) or []
    dolu = []
    for r in rows:
        try:
            oda = int(r.get("hazir_ofis_oda_no"))
        except (TypeError, ValueError):
            continue
        dolu.append(
            {
                "oda": oda,
                "musteri_id": r.get("id"),
                "firma_adi": (r.get("firma_adi") or "—").strip(),
            }
        )
    used = {d["oda"] for d in dolu}
    bos = [n for n in range(200, 231) if n not in used]
    return jsonify({"ok": True, "dolu": dolu, "bos": bos})


@bp.route("/api/musteri/<int:mid>/durum", methods=["POST"])
@giris_gerekli
def api_musteri_durum_guncelle(mid):
    """customers.durum + kapanis_tarihi + kapanış sonrası ek borç ayı + is_active."""
    data = request.get_json(silent=True) or {}
    dr_in = (data.get("durum") or "").strip().lower()
    if dr_in not in ("aktif", "pasif"):
        return jsonify({"ok": False, "mesaj": "durum aktif veya pasif olmalıdır."}), 400
    ensure_customers_durum()
    ensure_customers_is_active()
    ensure_customers_kapanis_sonrasi_borc_ay()
    ensure_customers_hazir_ofis_oda()
    row = fetch_one("SELECT id FROM customers WHERE id = %s", (mid,))
    if not row:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    payload = dict(data)
    if dr_in == "pasif" and not (payload.get("kapanis_tarihi") or "").strip():
        payload["kapanis_tarihi"] = date.today().isoformat()
    dr2, kap = _normalize_musteri_durum_kapanis(payload)
    kap_borc_ay = _normalize_kapanis_sonrasi_borc_ay(payload, dr2)
    # Rapor / liste: COALESCE(is_active, TRUE) ve durum NOT IN (pasif, …) birlikte kullanılıyor.
    kart_aktif = dr2 == "aktif"
    pasif_mi = dr2 == "pasif"
    n = execute(
        """
        UPDATE customers
        SET durum = %s, kapanis_tarihi = %s, kapanis_sonrasi_borc_ay = %s, is_active = %s,
            hazir_ofis_oda_no = CASE WHEN %s THEN NULL ELSE hazir_ofis_oda_no END
        WHERE id = %s
        """,
        (dr2, kap, kap_borc_ay, kart_aktif, pasif_mi, int(mid)),
    )
    if n is None or int(n) < 1:
        logging.warning("api_musteri_durum_guncelle: UPDATE rowcount=%s mid=%s user=%s", n, mid, getattr(current_user, "id", None))
        return jsonify(
            {
                "ok": False,
                "mesaj": "Kayıt güncellenemedi (satır yok veya veritabanı kısıtı). Yöneticiye bildirin.",
            }
        ), 409
    kap_out = kap.isoformat() if kap else None
    _upsert_aylik_grid_cache(int(mid))
    return jsonify({
        "ok": True,
        "durum": dr2,
        "kapanis_tarihi": kap_out,
        "kapanis_sonrasi_borc_ay": kap_borc_ay,
        "is_active": kart_aktif,
    })


@bp.route("/api/tahmini-musteri-no", methods=["GET"])
@giris_gerekli
def api_tahmini_musteri_no():
    """Yeni müşteri formu açılırken id/musteri_no rezerve et (customers_id_seq nextval)."""
    row = fetch_one("SELECT nextval('customers_id_seq') AS n")
    if not row or row.get("n") is None:
        return jsonify({"ok": False, "mesaj": "Müşteri numarası rezerve edilemedi."}), 500
    n = int(row["n"])
    return jsonify({"ok": True, "id": n, "musteri_no": n})


@bp.route("/api/musteri-cogalt", methods=["POST"])
@giris_gerekli
def api_musteri_cogalt():
    """
    Kayıtlı müşteriyi yeni id + yeni müşteri_no ile kopyalar (plan değişimi: eski kart pasif, yeni kart).
    Fatura/tahsilat kopyalanmaz. Son KYC satırı yeni müşteriye kopyalanır; sözleşme no ve hazır ofis odası temizlenir.
    """
    ensure_customers_musteri_no()
    ensure_customers_durum()
    ensure_customers_is_active()
    ensure_customers_hazir_ofis_oda()
    ensure_musteri_kyc_columns()
    ensure_musteri_kyc_hazir_ofis_oda_no()
    ensure_musteri_kyc_kira_banka()
    data = request.get_json(silent=True) or {}
    raw_id = data.get("kaynak_id") if data.get("kaynak_id") is not None else data.get("musteri_id")
    try:
        kaynak_id = int(raw_id)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "Geçerli bir müşteri seçin."}), 400

    src = fetch_one("SELECT * FROM customers WHERE id = %s", (kaynak_id,))
    if not src:
        return jsonify({"ok": False, "mesaj": "Kaynak müşteri bulunamadı."}), 404

    kyc = fetch_one(
        "SELECT * FROM musteri_kyc WHERE musteri_id = %s ORDER BY id DESC LIMIT 1",
        (kaynak_id,),
    )

    skip_cols = {"id", "musteri_no", "created_at"}
    new_row = {k: v for k, v in src.items() if k not in skip_cols}
    new_row.pop("updated_at", None)
    new_row["durum"] = "aktif"
    new_row["kapanis_tarihi"] = None
    new_row["is_active"] = True
    new_row["hazir_ofis_oda_no"] = None
    tag = f"Çoğaltma kaynağı: müşteri ID {kaynak_id}."
    prev_notes = (new_row.get("notes") or "").strip()
    new_row["notes"] = f"{prev_notes} {tag}".strip() if prev_notes else tag

    cols = list(new_row.keys())
    vals = [new_row[c] for c in cols]
    placeholders = ", ".join(["%s"] * len(vals))
    sql_cust = (
        f"INSERT INTO customers ({','.join(cols)}) VALUES ({placeholders}) RETURNING id"
    )

    try:
        with get_db() as conn:
            cur = conn.cursor()
            cur.execute(sql_cust, tuple(vals))
            ins = cur.fetchone()
            if not ins:
                raise RuntimeError("INSERT RETURNING sonuç dönmedi.")
            new_id = int(ins["id"])
            cur.execute(
                "UPDATE customers SET musteri_no = %s WHERE id = %s",
                (new_id, new_id),
            )
            new_mno = new_id
            if kyc:
                k2 = {k: v for k, v in kyc.items() if k not in ("id", "created_at", "updated_at")}
                k2["musteri_id"] = new_id
                k2["sozlesme_no"] = None
                k2["hazir_ofis_oda_no"] = None
                kcols = list(k2.keys())
                kvals = [k2[c] for c in kcols]
                kph = ", ".join(["%s"] * len(kvals))
                cur.execute(
                    f"INSERT INTO musteri_kyc ({','.join(kcols)}) VALUES ({kph})",
                    tuple(kvals),
                )
    except psycopg2.errors.UniqueViolation as e:
        logging.warning("api_musteri_cogalt unique violation kaynak=%s: %s", kaynak_id, e)
        return jsonify(
            {
                "ok": False,
                "mesaj": (
                    "Çoğaltılamadı: vergi numarası, müşteri numarası veya başka benzersiz alan çakışıyor. "
                    "Veritabanında vergi no tekil ise önce kaynak müşterinin vergi numarasını güncelleyin "
                    "veya yöneticiye başvurun."
                ),
            }
        ), 409
    except Exception as e:
        logging.exception("api_musteri_cogalt")
        return jsonify({"ok": False, "mesaj": str(e) or "Çoğaltma başarısız."}), 400

    try:
        _upsert_aylik_grid_cache(new_id)
    except Exception:
        pass

    mesaj = f"Yeni müşteri oluşturuldu (ID: {new_id}"
    if new_mno is not None:
        mesaj += f", müşteri no: {new_mno}"
    mesaj += "). Hizmet türü ve sözleşme tarihlerini kontrol edip Kaydet ile onaylayın."
    return jsonify({"ok": True, "id": new_id, "musteri_no": new_mno, "mesaj": mesaj})


def _digits_only(s):
    return "".join(c for c in str(s or "") if c.isdigit())


def _vergi_no_normalize_veya_hata(tax_raw, yetkili_tc_raw):
    """10 hane VKN veya Yetkili T.C. ile birebir aynı 11 hane T.C. (şahıs vergi no) kabul edilir."""
    v = _digits_only(tax_raw)
    tc = _digits_only(yetkili_tc_raw)
    if not v:
        return "Vergi numarası zorunludur.", None
    if len(v) == 10:
        return None, v
    if len(v) == 11 and len(tc) == 11 and v == tc:
        return None, v
    if len(v) == 11:
        return "Vergi no 11 hane yalnızca Yetkili T.C. Kimlik No ile aynı olduğunda kabul edilir.", None
    return "Vergi no 10 haneli VKN veya Yetkili T.C. ile aynı 11 haneli T.C. olmalıdır.", None


@bp.route('/kaydet', methods=['POST'])
@giris_gerekli
def kaydet():
    """Yeni müşteri kaydı veya güncelleme"""
    try:
        _giris_kaydet_schema_ensure_once()
        data = request.get_json()
        vergi_err, tax_norm = _vergi_no_normalize_veya_hata(data.get("tax_number"), data.get("yetkili_tc"))
        if vergi_err:
            return jsonify({"ok": False, "mesaj": vergi_err}), 400

        musteri_id = data.get('id')
        dr, kap = _normalize_musteri_durum_kapanis(data)
        kap_borc_ay = _normalize_kapanis_sonrasi_borc_ay(data, dr)
        bh_raw = data.get("bizim_hesap")
        if isinstance(bh_raw, str):
            bizim_hesap_legacy = bh_raw.strip().lower() in ("1", "true", "yes", "evet", "on")
        else:
            bizim_hesap_legacy = bool(bh_raw)
        g2_list = _parse_grup2_secimleri_from_request(data)
        if not g2_list and bizim_hesap_legacy:
            g2_list = ["bizim_hesap"]
        g2_list = list(dict.fromkeys(_filter_grup2_slugs_known(g2_list)))
        bizim_hesap = "bizim_hesap" in g2_list

        if musteri_id:
            # Güncelleme
            execute("""
                UPDATE customers SET 
                    name = %s,
                    musteri_adi = %s,
                    tax_number = %s,
                    phone = %s,
                    email = %s,
                    address = %s,
                    ev_adres = %s,
                    notes = %s,
                    durum = %s,
                    kapanis_tarihi = %s,
                    kapanis_sonrasi_borc_ay = %s,
                    bizim_hesap = %s,
                    grup2_secimleri = %s
                WHERE id = %s
            """, (
                data.get('name'),
                (data.get('musteri_adi') or '').strip() or None,
                tax_norm,
                data.get('phone'),
                data.get('email'),
                data.get('address'),
                data.get('ev_adres'),
                data.get('notes'),
                dr,
                kap,
                kap_borc_ay,
                bizim_hesap,
                g2_list,
                musteri_id
            ))
            return jsonify({'ok': True, 'mesaj': '✅ Müşteri güncellendi'})
        else:
            # Yeni kayıt: musteri_no = id (aynı transaction); opsiyonel rezerve_edilen_id
            insert_params = (
                data.get('name'),
                (data.get('musteri_adi') or '').strip() or None,
                tax_norm,
                data.get('phone'),
                data.get('email'),
                data.get('address'),
                data.get('ev_adres'),
                data.get('notes'),
                dr,
                kap,
                kap_borc_ay,
                bizim_hesap,
                g2_list,
            )
            rezerve_id = None
            raw_rez = data.get("rezerve_edilen_id")
            if raw_rez is not None and str(raw_rez).strip() != "":
                try:
                    rezerve_id = int(raw_rez)
                    if rezerve_id <= 0:
                        rezerve_id = None
                except (TypeError, ValueError):
                    rezerve_id = None
            try:
                with get_db() as conn:
                    cur = conn.cursor()
                    if rezerve_id is not None:
                        cur.execute(
                            "SELECT 1 AS ok FROM customers WHERE id = %s",
                            (rezerve_id,),
                        )
                        if cur.fetchone():
                            return jsonify(
                                {
                                    "ok": False,
                                    "mesaj": (
                                        "Rezerve müşteri numarası artık geçerli değil. "
                                        "«Yeni Müşteri» butonuna tekrar basın."
                                    ),
                                }
                            ), 409
                        cur.execute(
                            """
                            INSERT INTO customers (
                                id, name, musteri_adi, tax_number, phone, email, address,
                                ev_adres, notes, durum, kapanis_tarihi, kapanis_sonrasi_borc_ay,
                                bizim_hesap, grup2_secimleri, created_at
                            )
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                            RETURNING id
                            """,
                            (rezerve_id,) + insert_params,
                        )
                    else:
                        cur.execute(
                            """
                            INSERT INTO customers (
                                name, musteri_adi, tax_number, phone, email, address,
                                ev_adres, notes, durum, kapanis_tarihi, kapanis_sonrasi_borc_ay,
                                bizim_hesap, grup2_secimleri, created_at
                            )
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                            RETURNING id
                            """,
                            insert_params,
                        )
                    ins = cur.fetchone()
                    if not ins:
                        raise RuntimeError("INSERT RETURNING sonuç dönmedi.")
                    new_id = int(ins["id"])
                    cur.execute(
                        "UPDATE customers SET musteri_no = %s WHERE id = %s",
                        (new_id, new_id),
                    )
                    result = {"id": new_id, "musteri_no": new_id}
            except psycopg2.errors.UniqueViolation as e:
                logging.warning("kaydet yeni musteri unique violation: %s", e)
                mesaj = (
                    "Kayıt oluşturulamadı: vergi numarası veya müşteri numarası "
                    "benzersizlik kuralıyla çakışıyor. Vergi numarasını kontrol edin; "
                    "sorun sürerse yöneticiye başvurun."
                )
                if rezerve_id is not None:
                    mesaj = (
                        "Rezerve müşteri numarası artık geçerli değil veya başka bir alan çakışıyor. "
                        "«Yeni Müşteri» butonuna tekrar basın."
                    )
                return jsonify({"ok": False, "mesaj": mesaj}), 409
            mid = result["id"] if result else None
            mno = result.get("musteri_no") if result else None
            kaydet_mesaj = f"✅ Müşteri kaydedildi (ID: {mid}"
            if mno is not None:
                kaydet_mesaj += f", müşteri no: {mno}"
            kaydet_mesaj += ")"
            return jsonify({"ok": True, "mesaj": kaydet_mesaj, "id": mid, "musteri_no": mno})

    except Exception as e:
        return jsonify({'ok': False, 'mesaj': f'❌ Hata: {str(e)}'}), 400


@bp.route('/resim-yukle/<int:mid>', methods=['POST'])
@giris_gerekli
def resim_yukle(mid):
    """Müşteri dosyası yükle"""
    try:
        if 'file' not in request.files:
            return jsonify({'ok': False, 'mesaj': 'Dosya seçilmedi'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'ok': False, 'mesaj': 'Dosya seçilmedi'}), 400
        
        if file and allowed_file(file.filename):
            # Klasör oluştur
            os.makedirs(UPLOAD_FOLDER, exist_ok=True)
            
            # Dosya adını güvenli hale getir
            filename = secure_filename(f"musteri_{mid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}")
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            
            # Kaydet
            file.save(filepath)
            
            return jsonify({'ok': True, 'mesaj': '✅ Dosya yüklendi', 'filename': filename})
        else:
            return jsonify({'ok': False, 'mesaj': 'Geçersiz dosya formatı'}), 400
            
    except Exception as e:
        return jsonify({'ok': False, 'mesaj': f'❌ Hata: {str(e)}'}), 400


@bp.route('/sozlesme-olustur/<int:mid>')
@giris_gerekli
def sozlesme_olustur(mid):
    """Müşteri sözleşmesi oluştur (Word). ?indir=1 ile doğrudan Word indirilir; yoksa sayfa + Word İndir / WhatsApp."""
    try:
        musteri = fetch_one("SELECT c.* FROM customers c WHERE c.id = %s", (mid,))
        if not musteri:
            return "Müşteri bulunamadı", 404

        if request.args.get("indir") != "1":
            indir_url = url_for("giris.sozlesme_olustur", mid=mid, indir="1", tur=request.args.get("tur", ""))
            tel = (musteri.get("phone") or "").replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
            num = ("90" + tel[1:]) if (tel and tel.startswith("0")) else ("90" + tel if tel else "")
            metin = "Sayın " + (musteri.get("name") or "Müşteri") + ",\n\nHizmet sözleşmeniz ekte yer almaktadır. İncelemenizi rica ederiz.\n\nİyi günler dileriz.\nBESTOFFICE"
            whatsapp_url = "https://wa.me/" + num + "?text=" + urllib.parse.quote(metin) if num else "https://wa.me/?text=" + urllib.parse.quote(metin)
            return render_template(
                "giris/sozlesme_olustur_sayfa.html",
                musteri=musteri,
                indir_url=indir_url,
                whatsapp_url=whatsapp_url,
            )

        # Sözleşme numarası otomatik oluştur
        # Hizmet türüne göre prefix belirle (SO/HO/PO)
        tur_raw = (request.args.get("tur") or "").lower()
        if "hazır" in tur_raw or "hazir" in tur_raw:
            prefix = "HO"
        elif "paylaşımlı" in tur_raw or "paylasimli" in tur_raw:
            prefix = "PO"
        else:
            prefix = "SO"

        today = datetime.now()
        tarih_kod = today.strftime("%d%m%y")  # Örn: 270225

        # Aynı gün ve aynı türdeki sözleşmeler için 600'den başlayan artan numara
        pattern = f"{prefix}{tarih_kod}-%"
        last = fetch_one(
            "SELECT sozlesme_no FROM sozlesmeler WHERE sozlesme_no LIKE %s ORDER BY sozlesme_no DESC LIMIT 1",
            (pattern,),
        )
        if last and last.get("sozlesme_no"):
            try:
                son_parca = str(last["sozlesme_no"]).split("-")[-1]
                sayac = int(son_parca) + 1
            except Exception:
                sayac = 600
        else:
            sayac = 600

        sozlesme_no = f"{prefix}{tarih_kod}-{sayac}"
        
        # Word belgesi oluştur
        doc = Document()
        
        # Başlık
        heading = doc.add_heading('OFİSBİR HİZMET SÖZLEŞMESİ', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sözleşme No ve Tarih (aynı satırda)
        tarih_str = today.strftime('%d.%m.%Y')
        p_no = doc.add_paragraph()
        run_no = p_no.add_run(f"Sözleşme No: {sozlesme_no}    ")
        run_t = p_no.add_run(f"Tarih: {tarih_str}")
        doc.add_paragraph("")
        
        # MADDE 1 - TARAFLAR
        doc.add_heading('MADDE 1 - TARAFLAR', level=2)
        doc.add_paragraph("""
İşbu sözleşme, bir tarafta:

KİRAYA VEREN / HİZMET SAĞLAYICI:
Unvan: OFİSBİR Sanal Ofis Hizmetleri
Adres: Ankara, Türkiye

(Bundan böyle "KİRAYA VEREN" olarak anılacaktır.)

Diğer tarafta:

KİRACI / HİZMET ALAN:""")
        
        doc.add_paragraph(f"""
Unvan: {musteri['name']}
Vergi No: {musteri.get('tax_number') or '-'}
Vergi Dairesi: {musteri.get('vergi_dairesi') or '-'}
Adres: {musteri.get('address') or '-'}
Telefon: {musteri.get('phone') or '-'}
E-posta: {musteri.get('email') or '-'}

(Bundan böyle "KİRACI" olarak anılacaktır.)

arasında aşağıdaki şartlar dâhilinde akdedilmiştir.
        """)
        
        # MADDE 2 - SÖZLEŞMENİN KONUSU
        doc.add_heading('MADDE 2 - SÖZLEŞMENİN KONUSU', level=2)
        doc.add_paragraph(f"""
İşbu sözleşmenin konusu, KİRAYA VEREN'in mülkiyetinde bulunan adres üzerinde; KİRACI'nın işletme adresi olarak kullanması, posta ve kargo hizmetlerinden faydalanması, telefon santralı, sekreterlik ve diğer ofis hizmetlerinden yararlanması amacıyla tarafların hak ve yükümlülüklerinin belirlenmesidir.

Hizmet Türü: Sanal Ofis
Aylık Hizmet Bedeli: {musteri.get('aylik_kira', 0):.2f} TL + KDV
        """)
        
        # MADDE 3 - SÖZLEŞMENİN SÜRESİ
        doc.add_heading('MADDE 3 - SÖZLEŞMENİN SÜRESİ', level=2)
        doc.add_paragraph("""
İşbu sözleşme {sozlesme_baslangic} tarihinde başlamak üzere 1 (bir) yıl süre ile geçerlidir. 
Sözleşme süresi sonunda, taraflardan herhangi biri 1 ay önceden yazılı bildirimde bulunmadığı 
takdirde aynı şartlarla 1 yıl daha uzamış sayılır.
        """)
        
        # MADDE 4 - ÖDEME ŞARTLARI
        doc.add_heading('MADDE 4 - ÖDEME ŞARTLARI', level=2)
        doc.add_paragraph(f"""
4.1. Aylık hizmet bedeli {musteri.get('aylik_kira', 0):.2f} TL + KDV olup, her ayın 1-5'i 
arasında peşin olarak ödenecektir.

4.2. Ödemeler banka havalesi, EFT veya kredi kartı ile yapılabilir.

4.3. KİRACI'nın ödeme yükümlülüğünü yerine getirmemesi halinde, KİRAYA VEREN hizmetleri 
durdurma ve sözleşmeyi tek taraflı feshetme hakkına sahiptir.

4.4. Kira bedeli TÜFE artış oranına göre yıllık olarak güncellenecektir.
        """)
        
        # MADDE 5 - KİRAYA VEREN'İN YÜKÜMLÜLÜKLERİ
        doc.add_heading('MADDE 5 - KİRAYA VEREN\'İN YÜKÜMLÜLÜKLERİ', level=2)
        doc.add_paragraph("""
KİRAYA VEREN aşağıdaki hizmetleri sağlamayı taahhüt eder:

5.1. Sözleşme konusu adresin KİRACI'ya işletme adresi olarak tahsis edilmesi
5.2. Posta ve kargo kabul hizmeti
5.3. Telefon santralı ve çağrı yönlendirme hizmeti
5.4. Sekreterlik hizmeti (çalışma saatleri içinde)
5.5. Toplantı odası kullanımı (ücret karşılığı, rezervasyon ile)
5.6. Ortak alanların kullanımı
5.7. İnternet erişimi (ortak alanlarda)
        """)
        
        # MADDE 6 - KİRACI'NIN YÜKÜMLÜLÜKLERİ
        doc.add_heading('MADDE 6 - KİRACI\'NIN YÜKÜMLÜLÜKLERİ', level=2)
        doc.add_paragraph("""
6.1. Aylık hizmet bedelini zamanında ödemek
6.2. Verilen adresi yasalara ve ahlaka uygun şekilde kullanmak
6.3. Resmi kurumlardan gelen yazı ve bildirimleri düzenli takip etmek
6.4. Kargo ve posta takibini düzenli olarak yapmak
6.5. Toplantı odası kullanımını önceden rezerve ettirmek
6.6. Diğer müşterilere saygılı davranmak ve ortak alanları temiz kullanmak
6.7. Yasadışı faaliyetlerde bulunmamak
        """)
        
        # MADDE 7 - FESİH ŞARTLARI
        doc.add_heading('MADDE 7 - FESİH ŞARTLARI', level=2)
        doc.add_paragraph("""
7.1. Taraflardan herhangi biri, 1 ay önceden yazılı bildirimde bulunmak kaydıyla 
sözleşmeyi feshedebilir.

7.2. KİRACI'nın 2 ay üst üste ödeme yapmaması durumunda, KİRAYA VEREN sözleşmeyi 
tek taraflı olarak feshedebilir.

7.3. KİRACI'nın yasadışı faaliyetlerde bulunması, yasalara veya sözleşme şartlarına 
aykırı hareket etmesi durumunda KİRAYA VEREN derhal fesih hakkına sahiptir.

7.4. Fesih durumunda KİRACI, kullandığı hizmete ait tüm ödemelerini yapmakla yükümlüdür.
        """)
        
        # MADDE 8 - GİZLİLİK
        doc.add_heading('MADDE 8 - GİZLİLİK', level=2)
        doc.add_paragraph("""
8.1. Taraflar, sözleşme konusu hizmetler dolayısıyla öğrendiği karşı tarafa ait ticari 
sırları ve kişisel verileri gizli tutmayı ve üçüncü kişilerle paylaşmamayı taahhüt eder.

8.2. Bu yükümlülük sözleşmenin sona ermesinden sonra da 2 yıl süreyle devam eder.
        """)
        
        # MADDE 9 - UYUŞMAZLIKLARIN ÇÖZÜMÜ
        doc.add_heading('MADDE 9 - UYUŞMAZLIKLARIN ÇÖZÜMÜ', level=2)
        doc.add_paragraph("""
İşbu sözleşmeden doğabilecek her türlü uyuşmazlığın çözümünde Ankara Mahkemeleri 
ve İcra Daireleri yetkilidir.
        """)
        
        # MADDE 10 - YÜRÜRLÜK
        doc.add_heading('MADDE 10 - YÜRÜRLÜK', level=2)
        doc.add_paragraph(f"""
İşbu sözleşme {datetime.now().strftime('%d.%m.%Y')} tarihinde 2 (iki) nüsha olarak 
düzenlenmiş ve taraflarca okunup imzalanarak yürürlüğe girmiştir.
        """)
        
        # İmza alanları (sadeleştirilmiş)
        doc.add_paragraph("")
        doc.add_paragraph("KİRAYA VEREN / HİZMET SAĞLAYICI" + " " * 15 + "KİRACI / HİZMET ALAN")
        doc.add_paragraph(f"OFİSBİR Sanal Ofis Hizmetleri" + " " * 20 + f"{musteri['name']}")
        doc.add_paragraph("İmza: _______________" + " " * 30 + "İmza: _______________")
        
        # Dosya adı
        filename = f"Sozlesme_{sozlesme_no}_{musteri['name'].replace(' ', '_')}.docx"
        filepath = os.path.join('uploads', 'sozlesmeler', filename)
        
        # Klasör oluştur
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Kaydet
        doc.save(filepath)
        
        # Oluşan sözleşmeyi register et
        execute(
            "INSERT INTO sozlesmeler (musteri_id, sozlesme_no, hizmet_turu) VALUES (%s,%s,%s) ON CONFLICT (sozlesme_no) DO NOTHING",
            (mid, sozlesme_no, tur_raw or None),
        )

        # İndir
        return send_file(filepath, as_attachment=True, download_name=filename)
        
    except Exception as e:
        return f"Hata: {str(e)}", 500
@bp.route('/api/tufe-verileri')
@giris_gerekli
def api_tufe_verileri():
    """TCMB TÜFE verilerini getir"""
    try:
        cached = simple_cache_get(CACHE_KEY_TUFE_VERILERI, CACHE_TTL_SEC)
        if cached is not None:
            return jsonify(cached)
        veriler = fetch_all("""
            SELECT year as yil, month as ay, oran 
            FROM tufe_verileri 
            ORDER BY year DESC, 
            CASE month 
                WHEN 'Ocak' THEN 1 WHEN 'Şubat' THEN 2 WHEN 'Mart' THEN 3
                WHEN 'Nisan' THEN 4 WHEN 'Mayıs' THEN 5 WHEN 'Haziran' THEN 6
                WHEN 'Temmuz' THEN 7 WHEN 'Ağustos' THEN 8 WHEN 'Eylül' THEN 9
                WHEN 'Ekim' THEN 10 WHEN 'Kasım' THEN 11 WHEN 'Aralık' THEN 12
            END DESC
            LIMIT 60
        """)
        veriler = veriler or []
        simple_cache_set(CACHE_KEY_TUFE_VERILERI, veriler)
        return jsonify(veriler)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@bp.route('/api/kira-senaryo-excel', methods=['POST'])
@giris_gerekli
def kira_senaryo_excel():
    """Kira senaryo Excel çıktısı"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        
        data = request.get_json() or {}

        satirlar = data.get('satirlar') or []
        baslangic_kira = float(data.get('net_kira') or data.get('baslangic_kira') or 0)
        baslangic_tarih = data.get('baslangic_tarih') or ''
        yil_sayisi = int(data.get('yil_sayisi') or (len(satirlar) or 0) or 0)
        musteri_ismi = (data.get('musteri_ismi') or '').strip()
        nakit_mod = bool(data.get("kira_nakit_senaryo"))
        hibrit_mod = bool(data.get("kira_hibrit_senaryo"))

        # Workbook oluştur
        wb = Workbook()
        ws = wb.active
        ws.title = "Kira Senaryo"

        # Başlık
        _excel_baslik_ek = ' (Nakit / KDV yok)' if nakit_mod else (' (Hibrit)' if hibrit_mod else '')
        ws['A1'] = 'KİRA SENARYO HESAPLAMA' + _excel_baslik_ek
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:E1' if nakit_mod else 'A1:F1')
        ws['A1'].alignment = Alignment(horizontal='center')

        # Parametreler
        ws['A3'] = 'Müşteri İsmi:'
        ws['B3'] = musteri_ismi or '-'
        ws['A4'] = 'Başlangıç Kira:'
        ws['B4'] = baslangic_kira
        ws['A5'] = 'Başlangıç Tarihi:'
        ws['B5'] = baslangic_tarih
        ws['A6'] = 'Yıl Sayısı:'
        ws['B6'] = yil_sayisi

        # Tablo başlıkları (Nakit: KDV sütunu yok — ekranla uyumlu)
        if nakit_mod:
            headers = ['Yıl', 'Aylık Kira', 'Artış %', 'Yıllık Toplam', 'Artış Tutar']
        else:
            headers = ['Yıl', 'Aylık Kira', 'KDV Dahil', 'Artış %', 'Yıllık Toplam', 'Artış Tutar']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=8, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='0097A7', end_color='0097A7', fill_type='solid')
            cell.alignment = Alignment(horizontal='center')

        toplam_gelir = 0.0

        if satirlar:
            # Frontend'de hesaplanan tabloyu birebir Excel'e yaz
            for i, s in enumerate(satirlar, start=1):
                try:
                    yil = int(s.get('yil'))
                except Exception:
                    yil = None
                try:
                    aylik = float(str(s.get('aylik_kira') or '0').replace('.', '').replace(',', '.'))
                except Exception:
                    aylik = 0.0
                try:
                    yillik = float(str(s.get('yillik_toplam') or '0').replace('.', '').replace(',', '.'))
                except Exception:
                    yillik = aylik * 12
                artis_yuzde_raw = (s.get('artis_yuzde') or '').strip()
                if artis_yuzde_raw.endswith('%'):
                    artis_yuzde_raw = artis_yuzde_raw[:-1]
                try:
                    artis_yuzde = float(artis_yuzde_raw.replace(',', '.')) / 100.0
                except Exception:
                    artis_yuzde = 0.0
                try:
                    artis_tutar = float(str(s.get('artis_tutar') or '0').replace('.', '').replace(',', '.'))
                except Exception:
                    artis_tutar = 0.0

                toplam_gelir += yillik

                row = 8 + i
                ws.cell(row=row, column=1).value = yil
                ws.cell(row=row, column=2).value = aylik
                ws.cell(row=row, column=2).number_format = '#,##0.00'
                if nakit_mod:
                    ws.cell(row=row, column=3).value = artis_yuzde
                    ws.cell(row=row, column=3).number_format = '0.00%'
                    ws.cell(row=row, column=4).value = yillik
                    ws.cell(row=row, column=4).number_format = '#,##0.00'
                    ws.cell(row=row, column=5).value = artis_tutar
                    ws.cell(row=row, column=5).number_format = '#,##0.00'
                else:
                    kd_raw = s.get("kdv_dahil_aylik")
                    kd_row = None
                    if kd_raw is not None and kd_raw not in ("", False):
                        if isinstance(kd_raw, (int, float)) and not isinstance(kd_raw, bool):
                            kd_row = float(kd_raw)
                        else:
                            try:
                                st = str(kd_raw).replace(" ", "").replace("TL", "").strip().replace(".", "").replace(",", ".")
                                kd_row = float(st) if st else None
                            except Exception:
                                kd_row = None
                    if hibrit_mod and kd_row is not None and kd_row > 0:
                        kdv_dahil = kd_row
                    else:
                        kdv_dahil = aylik * 1.20
                    ws.cell(row=row, column=3).value = kdv_dahil
                    ws.cell(row=row, column=3).number_format = '#,##0.00'
                    ws.cell(row=row, column=4).value = artis_yuzde
                    ws.cell(row=row, column=4).number_format = '0.00%'
                    ws.cell(row=row, column=5).value = yillik
                    ws.cell(row=row, column=5).number_format = '#,##0.00'
                    ws.cell(row=row, column=6).value = artis_tutar
                    ws.cell(row=row, column=6).number_format = '#,##0.00'
        else:
            # Eski davranış: sabit TÜFE oranı ile hesapla (geriye dönük uyumluluk için)
            tufe_oran = float(data.get('tufe_oran') or 0) / 100.0
            yil = int((baslangic_tarih or '2000-01-01').split('-')[0])
            mevcut_kira = baslangic_kira
            for i in range(1, yil_sayisi + 1):
                yillik_toplam = mevcut_kira * 12
                toplam_gelir += yillik_toplam
                artis_oran = 0 if i == 1 else tufe_oran
                artis_tutar = 0 if i == 1 else mevcut_kira - (mevcut_kira / (1 + tufe_oran or 1))
                row = 8 + i
                ws.cell(row=row, column=1).value = yil + i - 1
                ws.cell(row=row, column=2).value = mevcut_kira
                ws.cell(row=row, column=2).number_format = '#,##0.00'
                if nakit_mod:
                    ws.cell(row=row, column=3).value = artis_oran
                    ws.cell(row=row, column=3).number_format = '0.00%'
                    ws.cell(row=row, column=4).value = yillik_toplam
                    ws.cell(row=row, column=4).number_format = '#,##0.00'
                    ws.cell(row=row, column=5).value = artis_tutar
                    ws.cell(row=row, column=5).number_format = '#,##0.00'
                else:
                    kdv_dahil = mevcut_kira * 1.20
                    ws.cell(row=row, column=3).value = kdv_dahil
                    ws.cell(row=row, column=3).number_format = '#,##0.00'
                    ws.cell(row=row, column=4).value = artis_oran
                    ws.cell(row=row, column=4).number_format = '0.00%'
                    ws.cell(row=row, column=5).value = yillik_toplam
                    ws.cell(row=row, column=5).number_format = '#,##0.00'
                    ws.cell(row=row, column=6).value = artis_tutar
                    ws.cell(row=row, column=6).number_format = '#,##0.00'
                if i < yil_sayisi:
                    mevcut_kira = mevcut_kira * (1 + tufe_oran)

        # Toplam
        satir_sayisi = len(satirlar) or yil_sayisi
        son_satir = 8 + satir_sayisi + 2
        toplam_col = 4 if nakit_mod else 5
        ws.cell(row=son_satir, column=1).value = f'TOPLAM ({satir_sayisi} Yıl):'
        ws.cell(row=son_satir, column=1).font = Font(bold=True)
        ws.cell(row=son_satir, column=toplam_col).value = toplam_gelir
        ws.cell(row=son_satir, column=toplam_col).number_format = '#,##0.00'
        ws.cell(row=son_satir, column=toplam_col).font = Font(bold=True)
        ws.cell(row=son_satir, column=toplam_col).fill = PatternFill(start_color='4CAF50', end_color='4CAF50', fill_type='solid')

        # Sütun genişlikleri
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 18
        ws.column_dimensions['C'].width = 18
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 18
        if not nakit_mod:
            ws.column_dimensions['F'].width = 18
        
        # Dosya kaydet (bellekten gönder)
        filename = f"Kira_Senaryo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def _tarih_fmt(s):
    """YYYY-MM-DD veya DD.MM.YYYY -> DD.MM.YYYY"""
    if not s:
        return ""
    s = str(s).strip()
    if re.match(r"^\d{4}-\d{2}-\d{2}", s):
        parts = s[:10].split("-")
        return f"{parts[2]}.{parts[1]}.{parts[0]}"
    return s[:10] if len(s) >= 10 else s


def _parse_date_str(s):
    """Basit tarih parse: YYYY-MM-DD veya DD.MM.YYYY / DD/MM/YYYY -> date."""
    if not s:
        return None
    s = str(s).strip()
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(s[:10], fmt).date()
        except Exception:
            continue
    return None


def _add_months(d: date, months: int) -> date:
    """Ay ekle (takvim ayı bazlı, yıl devretmeli)."""
    y = d.year + (d.month - 1 + months) // 12
    m = (d.month - 1 + months) % 12 + 1
    # Gün sapmasını engelle: aynı gün yoksa ayın son günü
    day = min(d.day, [31,
                      29 if y % 4 == 0 and (y % 100 != 0 or y % 400 == 0) else 28,
                      31, 30, 31, 30, 31, 31, 30, 31, 30, 31][m - 1])
    return date(y, m, day)


def _generate_installments(contract_id: int, musteri_id: int, baslangic: date, bitis: date | None,
                            sure_ay: int | None, aylik_kira: float, odeme_gunu: int | None):
    """Verilen sözleşme için taksit planını (contract_installments) yeniden üret."""
    if not baslangic or not aylik_kira:
        return
    # Süre yoksa, başlangıç-bitişten ay farkını hesapla
    if not sure_ay and bitis:
        sure_ay = max(1, (bitis.year - baslangic.year) * 12 + (bitis.month - baslangic.month) + 1)
    if not sure_ay:
        sure_ay = 12
    # Eski planı sil
    execute("DELETE FROM contract_installments WHERE contract_id=%s", (contract_id,))
    for i in range(sure_ay):
        vade = _add_months(baslangic, i)
        if odeme_gunu:
            try:
                # Aynı ay, belirtilen gün
                vade = vade.replace(day=min(odeme_gunu, 28 if vade.month == 2 else 30 if vade.day > 30 else odeme_gunu))
            except Exception:
                pass
        execute(
            """
            INSERT INTO contract_installments
                (contract_id, musteri_id, taksit_no, vade_tarihi, tutar, odeme_durumu, odenen_tutar, kalan_tutar)
            VALUES (%s,%s,%s,%s,%s,'planlandi',0,%s)
            """,
            (contract_id, musteri_id, i + 1, vade, aylik_kira, aylik_kira),
        )


def _kira_bildirgesi_yillik_goster(hizmet_turu) -> bool:
    """Yıllık tutar metni yalnızca hizmet türü açıkça 'Sanal Ofis' ise; diğer tüm türlerde yok."""
    raw = str(hizmet_turu or "").strip()
    if not raw:
        return False
    t = raw.casefold()
    key = re.sub(r"[\s_\-/.]+", "", t).replace("ı", "i")
    return key == "sanalofis"


def _kira_bildirgesi_kdv_oran_float(kdv_raw, default=20.0):
    """KDV %0 (nakit) geçerlidir; `x or 20` / `float(x or 20)` kullanılmaz — 0 yanlışlıkla 20 olur."""
    if kdv_raw is None:
        return float(default)
    try:
        return float(kdv_raw)
    except (TypeError, ValueError):
        return float(default)


def build_kira_bildirgesi_pdf(
    musteri_adi,
    sozlesme_tarihi,
    gecerlilik_tarihi,
    kira_net,
    kdv_oran=20,
    hizmet_turu="",
    hibrit_nakit_pay=None,
    hibrit_banka_net=None,
):
    """Kira bildirgesi mektubu A4 PDF (bestoffice / Ofisbir). Arial ile Türkçe karakter desteği.
    hizmet_turu: yalnızca sanal_ofis -> yıllık kira ibaresi; diğerlerinde aylık net + KDV dahil.
    Hibrit: nakit payı KDV dışı, banka net payı üzerinden KDV — toplam KDV dahil = nakit + banka_net * (1+kdv%).
    """
    _register_arial()
    buf = io.BytesIO()
    w, h = A4
    c = canvas.Canvas(buf, pagesize=A4)
    c.setTitle("Kira Bildirgesi")
    font_name = "Arial" if "Arial" in pdfmetrics.getRegisteredFontNames() else "Helvetica"

    kira_net = float(kira_net or 0)
    kdv_oran = _kira_bildirgesi_kdv_oran_float(kdv_oran, 20.0)
    try:
        hn = float(hibrit_nakit_pay)
    except (TypeError, ValueError):
        hn = 0.0
    try:
        hb = float(hibrit_banka_net)
    except (TypeError, ValueError):
        hb = 0.0
    if hn > 0 and hb > 0 and kdv_oran > 0:
        kira_net = round(hn + hb, 2)
        kdv_dahil = round(hn + hb * (1 + kdv_oran / 100), 2)
    else:
        kdv_dahil = round(kira_net * (1 + kdv_oran / 100), 2)
    yillik = round(kdv_dahil * 12, 2)
    sanal_ofis = _kira_bildirgesi_yillik_goster(hizmet_turu)

    soz_fmt, gec_fmt = _soz_ve_guncel_tarih(sozlesme_tarihi, gecerlilik_tarihi)

    # Sol üst logo (blok solda; alt yazı logo merkezine göre ortalanır)
    logo_x = 15 * mm
    logo_center_x = logo_x
    try:
        here = os.path.dirname(os.path.abspath(__file__))
        logo_cands = []
        for nm in (
            "Ofisbir Logo.jpg", "Ofisbir Logo.jpeg", "Ofisbir Logo.png",
            "ofisbir_logo.png", "ofisbir_logo.jpg", "ofisbir_logo.jpeg",
            "ofisbir.png", "ofisbir.jpg", "ofisbir.jpeg", "logo.png", "logo.jpg", "logo.jpeg",
        ):
            logo_cands.append(os.path.abspath(os.path.join(here, "..", "..", "assets", nm)))
            logo_cands.append(os.path.abspath(os.path.join(here, "..", "static", nm)))
        logo_path = next((p for p in logo_cands if os.path.isfile(p)), None)
        if logo_path:
            img = ImageReader(logo_path)
            iw, ih = img.getSize()
            max_w = 96 * mm
            max_h = 28 * mm
            if iw and ih:
                scale = min(max_w / float(iw), max_h / float(ih))
                draw_w = float(iw) * scale
                draw_h = float(ih) * scale
            else:
                draw_w, draw_h = max_w, max_h
            c.drawImage(
                logo_path,
                logo_x,
                h - 40 * mm,
                width=draw_w,
                height=draw_h,
                preserveAspectRatio=False,
                mask='auto',
            )
            logo_center_x = logo_x + (draw_w / 2.0)
    except Exception:
        pass

    y = 44
    c.setFont(font_name, 9)
    c.setFillColorRGB(0, 0, 0)
    unvan_sol = "Ofisbir Ofis ve Dan. Hiz. A.Ş."
    unvan_w = c.stringWidth(unvan_sol, font_name, 9)
    c.drawString(logo_center_x - (unvan_w / 2.0), h - y * mm, unvan_sol)
    y += 20

    hitap_adi = (musteri_adi or "").strip() or "Değerli Kiracımız"
    c.setFont(font_name, 11)
    c.drawString(15 * mm, h - y * mm, "Sayın " + hitap_adi + ",")
    y += 12

    par1 = f"Tarafınızla {soz_fmt} tarihinde imzalanmış olan kira sözleşmesi gereği, {gec_fmt} tarihi itibarıyla kira bedeli güncellemesi yapılması gerekmektedir."
    c.setFont(font_name, 10)
    for chunk in (par1[i:i+95] for i in range(0, len(par1), 95)):
        c.drawString(15 * mm, h - y * mm, chunk)
        y += 5
    y += 6

    if kdv_oran == 0:
        par2_base = (
            f"Mevcut ekonomik koşullar ve yasal düzenlemeler göz önüne alınarak, adı geçen tarihten itibaren uygulanacak yeni kira bedeli TÜFE Yasal Oranı çerçevesinde güncellenecektir. "
            f"Buna göre, {gec_fmt} itibarıyla bedeliniz {kira_net:,.2f} TL dir."
        )
        if sanal_ofis:
            par2 = par2_base + f" Yıllık toplam {yillik:,.2f} TL dir."
        else:
            par2 = par2_base + " "
    else:
        par2_base = (
            f"Mevcut ekonomik koşullar ve yasal düzenlemeler göz önüne alınarak, adı geçen tarihten itibaren uygulanacak yeni kira bedeli TÜFE Yasal Oranı çerçevesinde güncellenecektir. "
            f"Buna göre, {gec_fmt} itibarıyla aylık kira bedeliniz {kira_net:,.2f} TL + %{int(kdv_oran)} KDV dahil {kdv_dahil:,.2f} TL dir."
        )
        if sanal_ofis:
            par2 = par2_base + f" KDV Dahil yıllık {yillik:,.2f} TL dir."
        else:
            par2 = par2_base + " "
    for chunk in (par2[i:i+95] for i in range(0, len(par2), 95)):
        c.drawString(15 * mm, h - y * mm, chunk)
        y += 5
    y += 10

    par3 = "Anlayışınız ve iş birliğiniz için teşekkür eder, sorularınız veya ek talepleriniz olması durumunda bizimle iletişime geçmekten çekinmemenizi rica ederiz."
    for chunk in (par3[i:i+95] for i in range(0, len(par3), 95)):
        c.drawString(15 * mm, h - y * mm, chunk)
        y += 5
    y += 14

    right_margin = w - 20 * mm
    c.setFont(font_name, 10)
    c.drawRightString(right_margin, h - y * mm, "Saygılarımızla,")
    y += 10
    c.setFont(font_name, 11)
    c.drawRightString(right_margin, h - y * mm, "BESTOFFICE")
    y += 6
    c.setFont(font_name, 9)
    unvan_text = "Ofisbir Ofis ve Danışmanlık Hizmetleri A.Ş."
    w_best = c.stringWidth("BESTOFFICE", font_name, 11)
    w_unvan = c.stringWidth(unvan_text, font_name, 9)
    unvan_x = right_margin - w_best / 2 - w_unvan / 2
    c.drawString(unvan_x, h - y * mm, unvan_text)

    c.save()
    buf.seek(0)
    return buf.getvalue()


def _soz_ve_guncel_tarih(sozlesme_tarihi, gecerlilik_tarihi):
    """Sözleşme başlangıç tarihi (ilk tarih) ve bugünkü yıla göre güncel artış tarihini döndürür.

    - İlk tarih: sözleşme tarihi (orijinal yıl)
    - Güncel tarih: sözleşme tarihinin gün/ayı + bugünün yılı
    """
    sozlesme_date = _parse_date_str(sozlesme_tarihi)
    if sozlesme_date:
        soz_fmt = sozlesme_date.strftime("%d.%m.%Y")
        today = date.today()
        try:
            guncel = date(today.year, sozlesme_date.month, sozlesme_date.day)
        except ValueError:
            guncel = date(today.year, sozlesme_date.month, min(sozlesme_date.day, 28))
        gec_fmt = guncel.strftime("%d.%m.%Y")
    else:
        soz_fmt = _tarih_fmt(sozlesme_tarihi)
        gec_fmt = _tarih_fmt(gecerlilik_tarihi) or date.today().strftime("%d.%m.%Y")
    return soz_fmt, gec_fmt


def _kira_bildirgesi_metinleri(sozlesme_tarihi, gecerlilik_tarihi, kira_net, kdv_oran, hizmet_turu=""):
    """Kira bildirgesi paragraf metinlerini döndürür (HTML şablonu için). Yıllık satır yalnızca sanal ofis."""
    soz_fmt, gec_fmt = _soz_ve_guncel_tarih(sozlesme_tarihi, gecerlilik_tarihi)

    kira_net = float(kira_net or 0)
    kdv_oran = _kira_bildirgesi_kdv_oran_float(kdv_oran, 20.0)
    kdv_dahil = round(kira_net * (1 + kdv_oran / 100), 2)
    yillik = round(kdv_dahil * 12, 2)
    yillik_goster = _kira_bildirgesi_yillik_goster(hizmet_turu)

    # HTML içinde tarih ve tutarların satır ortasından bölünmesini engellemek için nowrap span'leri kullan
    soz_html = f'<span class="nowrap">{soz_fmt}</span>' if soz_fmt else ''
    gec_html = f'<span class="nowrap">{gec_fmt}</span>' if gec_fmt else ''
    kira_net_html = f'<span class="nowrap">{kira_net:,.2f} TL</span>'
    kdv_dahil_html = f'<span class="nowrap">{kdv_dahil:,.2f} TL</span>'
    yillik_html = f'<span class="nowrap">{yillik:,.2f} TL</span>'

    # Yeni metin:
    # Konu: Hizmet Bedeli Güncellemesi Hakkında Bilgilendirme
    par1 = (
        "BestOffice bünyesinde devam eden iş birliğimiz ve bize duyduğunuz güven için teşekkür ederiz.<br><br>"
        f"{soz_html} başlangıç tarihli \"Ofis Kullanım ve Hizmet Sözleşmeniz\" uyarınca, hizmet bedeliniz güncellenmiştir. "
        "Mevcut ekonomik veriler ve yasal TÜFE oranları dikkate alınarak yapılan düzenleme neticesinde; "
        f"{gec_html} itibarıyla geçerli olacak yeni dönem hizmet bedeli bilgilerinizi aşağıda bulabilirsiniz."
    )

    if kdv_oran == 0:
        par2 = (
            f"Aylık Hizmet Bedeli: {kira_net_html}<br><br>"
            f"KDV uygulanmamaktadır."
        )
        if yillik_goster:
            par2 += f"<br><br>Yıllık Toplam: {yillik_html}"
    else:
        par2 = (
            f"Aylık Hizmet Bedeli (KDV Hariç): {kira_net_html}<br><br>"
            f"Aylık Toplam (KDV Dahil %{int(kdv_oran)}): {kdv_dahil_html}"
        )
        if yillik_goster:
            par2 += f"<br><br>Yıllık Toplam (KDV Dahil): {yillik_html}"

    par3 = (
        "Yeni döneme ait ödemelerinizi mevcut sözleşme şartlarında belirtilen hesap numaralarımıza yapmanızı rica ederiz. "
        "Başarılarınızın devamını diler, her türlü sorunuz için bizimle iletişime geçmekten çekinmemenizi önemle rica ederiz."
    )
    return par1, par2, par3


def _resolve_kira_bildirgesi_logo_data_uri():
    """Makbuzda kullanılan logo dosyasını bildirge HTML'ine göm."""
    here = os.path.dirname(os.path.abspath(__file__))
    roots = [
        os.path.abspath(os.path.join(here, "..", "..", "assets")),
        os.path.abspath(os.path.join(here, "..", "static")),
    ]
    cands = []
    for rt in roots:
        if not os.path.isdir(rt):
            continue
        # Önce bilinen dosya adları
        for nm in (
            "Ofisbir Logo.jpg", "Ofisbir Logo.jpeg", "Ofisbir Logo.png", "Ofisbir Logo.webp",
            "ofisbir_logo.png", "ofisbir_logo.jpg", "ofisbir_logo.jpeg", "ofisbir_logo.webp",
            "ofisbir.png", "ofisbir.jpg", "ofisbir.jpeg", "ofisbir.webp",
            "logo.png", "logo.jpg", "logo.jpeg", "logo.webp",
        ):
            cands.append(os.path.join(rt, nm))
        # Sonra klasördeki toleranslı eşleşmeler (uzantı/ad varyasyonları için)
        try:
            files = sorted(os.listdir(rt))
        except Exception:
            files = []
        for fn in files:
            low = str(fn).strip().lower()
            if not low:
                continue
            if not (low.endswith(".png") or low.endswith(".jpg") or low.endswith(".jpeg") or low.endswith(".webp")):
                continue
            if ("ofisbir" in low) or ("bestoffice" in low) or ("logo" in low):
                cands.append(os.path.join(rt, fn))
    for pth in cands:
        if not os.path.isfile(pth):
            continue
        try:
            ext = os.path.splitext(pth)[1].lower()
            if ext == ".png":
                mime = "image/png"
            elif ext == ".webp":
                mime = "image/webp"
            else:
                mime = "image/jpeg"
            with open(pth, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("ascii")
            return f"data:{mime};base64,{b64}"
        except Exception:
            continue
    return ""


@bp.route('/kira-bildirgesi-antet')
@giris_gerekli
def kira_bildirgesi_antet():
    """Antetli kira bildirgesi HTML sayfası (önizleme / yazdır)."""
    musteri_adi = (request.args.get('musteri_adi') or '').strip() or 'Müşteri Adı'
    sozlesme_tarihi = request.args.get('sozlesme_tarihi') or ''
    gecerlilik_tarihi = request.args.get('gecerlilik_tarihi') or ''
    try:
        kira_net = float(request.args.get('kira_net') or 0)
        kdv_oran = _kira_bildirgesi_kdv_oran_float(request.args.get('kdv_oran'), 20.0)
    except (TypeError, ValueError):
        kira_net, kdv_oran = 0, 20
    hizmet_turu = (request.args.get('hizmet_turu') or "").strip()
    if not gecerlilik_tarihi:
        gecerlilik_tarihi = sozlesme_tarihi or datetime.now().strftime("%Y-%m-%d")
    par1, par2, par3 = _kira_bildirgesi_metinleri(sozlesme_tarihi, gecerlilik_tarihi, kira_net, kdv_oran, hizmet_turu)
    hitap_adi = (musteri_adi or "").strip() or "Değerli Kiracımız"
    logo_data_uri = _resolve_kira_bildirgesi_logo_data_uri()
    return render_template(
        'giris/kira_bildirgesi_antet.html',
        musteri_adi=musteri_adi,
        hitap_adi=hitap_adi,
        par1=par1,
        par2=par2,
        par3=par3,
        logo_data_uri=logo_data_uri,
    )


@bp.route('/kira-bildirgesi-pdf', methods=['POST'])
@giris_gerekli
def kira_bildirgesi_pdf():
    """Kira bildirgesi PDF oluştur (önizleme / yazdır)."""
    try:
        data = request.get_json()
        musteri_adi = (data.get('musteri_adi') or '').strip() or 'Değerli Kiracımız'
        sozlesme_tarihi = data.get('sozlesme_tarihi') or ''
        gecerlilik_tarihi = data.get('gecerlilik_tarihi') or ''
        kira_net = float(data.get('kira_net') or 0)
        kdv_oran = _kira_bildirgesi_kdv_oran_float(data.get('kdv_oran'), 20.0)
        try:
            hibrit_n = float(data.get("hibrit_nakit_pay"))
        except (TypeError, ValueError):
            hibrit_n = 0.0
        try:
            hibrit_b = float(data.get("hibrit_banka_net"))
        except (TypeError, ValueError):
            hibrit_b = 0.0
        kh = data.get("kira_hibrit")
        use_hibrit_pdf = kh is True or kh == 1 or (isinstance(kh, str) and kh.strip().lower() in ("1", "true", "on", "yes"))
        if not use_hibrit_pdf or hibrit_n <= 0 or hibrit_b <= 0:
            use_hibrit_pdf = False
            hibrit_n, hibrit_b = 0.0, 0.0
        kn = data.get("kira_nakit")
        if not use_hibrit_pdf and (
            kn is True
            or kn == 1
            or (isinstance(kn, str) and kn.strip().lower() in ("1", "true", "on", "yes"))
        ):
            kdv_oran = 0.0
        if not gecerlilik_tarihi:
            return jsonify({'ok': False, 'mesaj': 'Geçerlilik tarihi giriniz.'}), 400
        if kira_net <= 0:
            return jsonify({'ok': False, 'mesaj': 'Kira tutarı 0\'dan büyük olmalıdır.'}), 400
        ht_raw = data.get("hizmet_turu")
        if ht_raw is None:
            hizmet_turu = ""
        else:
            hizmet_turu = str(ht_raw).strip()
        pdf_bytes = build_kira_bildirgesi_pdf(
            musteri_adi,
            sozlesme_tarihi,
            gecerlilik_tarihi,
            kira_net,
            kdv_oran,
            hizmet_turu=hizmet_turu,
            hibrit_nakit_pay=hibrit_n if use_hibrit_pdf else None,
            hibrit_banka_net=hibrit_b if use_hibrit_pdf else None,
        )
        return Response(pdf_bytes, mimetype="application/pdf", headers={
            "Content-Disposition": "inline; filename=Kira_Bildirgesi.pdf"
        })
    except Exception as e:
        return jsonify({'ok': False, 'mesaj': str(e)}), 500


# ── Cari Kart API ───────────────────────────────────────────────────────────

def _odeme_turu_harf(odeme_turu):
    """Tahsilat açıklaması için harf: EFT/Havale/Banka=B, Çek=C, Kredi Kartı=K, Nakit=N."""
    if not odeme_turu:
        return "N"
    o = str(odeme_turu).strip().lower()
    if o in ("havale", "eft", "banka"):
        return "B"
    if o == "cek":
        return "C"
    if o in ("kredi_karti", "kredi kartı"):
        return "K"
    return "N"


def _marker_ay_tarihi_from_text(text):
    """Açıklamadaki |AYLIK_TAH|YYYY-MM-DD| marker'ından ay tarihini çıkar."""
    if not text:
        return None
    m = re.search(r"\|AYLIK_TAH\|(\d{4}-\d{2}-\d{2})\|", str(text))
    if not m:
        return None
    try:
        return datetime.strptime(m.group(1), "%Y-%m-%d").date()
    except Exception:
        return None


def _tahsilat_aciklama_temizle(text):
    """Marker'ı gizle, görünümde temiz açıklama döndür."""
    s = str(text or "").strip()
    if not s:
        return ""
    s = re.sub(r"\|AYLIK_TAH\|\d{4}-\d{2}-\d{2}\|", "", s)
    return re.sub(r"\s{2,}", " ", s).strip()


def _cari_hareketler(musteri_id, banka_tahsilat_only=False, cari_ekstre_b=False, max_rows=None):
    """Fatura (borç) ve tahsilat (alacak) satırlarını tarih sırasına göre birleştirip bakiye hesaplar.
    banka_tahsilat_only=True ise nakit hariç tahsilatlar (havale, eft, banka, kredi kartı, çek).
    cari_ekstre_b=True ise Cari Ekstre B: borçta yalnızca ETTN’li (GİB kesilmiş) faturalar;
    alacakta yalnızca havale, EFT, banka, çek, kredi kartı (nakit yok)."""
    _nt = sql_expr_fatura_not_gib_taslak("notlar")
    limit_n = None
    if max_rows is not None:
        try:
            n = int(max_rows)
            if n > 0:
                limit_n = n
        except (TypeError, ValueError):
            limit_n = None
    limit_sql = " LIMIT %s" if limit_n else ""

    if cari_ekstre_b:
        faturalar = fetch_all(
            f"""SELECT id, fatura_no AS belge_no, fatura_tarihi AS tarih, COALESCE(toplam, tutar, 0) AS tutar, 'Fatura' AS tur, vade_tarihi
               FROM faturalar
               WHERE musteri_id = %s
                 AND NULLIF(TRIM(COALESCE(ettn::text, '')), '') IS NOT NULL
                 AND {_nt}
               ORDER BY fatura_tarihi DESC, id DESC{limit_sql}""",
            ((musteri_id, limit_n) if limit_n else (musteri_id,)),
        )
        tahsilatlar = fetch_all(
            """SELECT t.id,
                      COALESCE(t.makbuz_no, 'Makbuz-' || t.id) AS belge_no,
                      t.tahsilat_tarihi AS tarih,
                      COALESCE(
                        f.fatura_tarihi::date,
                        NULLIF(substring(COALESCE(t.aciklama, '') from '\\|AYLIK_TAH\\|([0-9]{4}-[0-9]{2}-[0-9]{2})\\|'), '')::date,
                        t.tahsilat_tarihi::date
                      ) AS eslesme_tarihi,
                      t.tutar, t.odeme_turu, t.aciklama AS tahsilat_aciklama, 'Tahsilat' AS tur
               FROM tahsilatlar t
               LEFT JOIN faturalar f ON f.id = t.fatura_id
                WHERE (t.musteri_id = %s OR t.customer_id = %s OR f.musteri_id = %s)
                  AND LOWER(TRIM(COALESCE(t.odeme_turu, 'nakit'))) IN ('havale', 'eft', 'banka', 'cek', 'kredi_karti')
               ORDER BY t.tahsilat_tarihi DESC, t.id DESC""" + limit_sql,
            ((musteri_id, musteri_id, musteri_id, limit_n) if limit_n else (musteri_id, musteri_id, musteri_id)),
        )
    else:
        faturalar = fetch_all(
            f"""SELECT id, fatura_no AS belge_no, fatura_tarihi AS tarih, COALESCE(toplam, tutar, 0) AS tutar, 'Fatura' AS tur, vade_tarihi
               FROM faturalar WHERE musteri_id = %s AND {_nt} ORDER BY fatura_tarihi DESC, id DESC{limit_sql}""",
            ((musteri_id, limit_n) if limit_n else (musteri_id,)),
        )
        if banka_tahsilat_only:
            # Nakit hariç tüm tahsilatlar: havale, eft, banka, kredi kartı, çek. Sadece nakit gösterilmez.
            tahsilatlar = fetch_all(
                """SELECT t.id,
                          COALESCE(t.makbuz_no, 'Makbuz-' || t.id) AS belge_no,
                          t.tahsilat_tarihi AS tarih,
                          COALESCE(
                            f.fatura_tarihi::date,
                            NULLIF(substring(COALESCE(t.aciklama, '') from '\\|AYLIK_TAH\\|([0-9]{4}-[0-9]{2}-[0-9]{2})\\|'), '')::date,
                            t.tahsilat_tarihi::date
                          ) AS eslesme_tarihi,
                          t.tutar, t.odeme_turu, t.aciklama AS tahsilat_aciklama, 'Tahsilat' AS tur
                   FROM tahsilatlar t
                   LEFT JOIN faturalar f ON f.id = t.fatura_id
                   WHERE (t.musteri_id = %s OR t.customer_id = %s OR f.musteri_id = %s)
                     AND LOWER(TRIM(COALESCE(t.odeme_turu, 'nakit'))) IN ('havale', 'eft', 'banka', 'kredi_karti', 'cek')
                   ORDER BY t.tahsilat_tarihi DESC, t.id DESC""" + limit_sql,
                ((musteri_id, musteri_id, musteri_id, limit_n) if limit_n else (musteri_id, musteri_id, musteri_id)),
            )
        else:
            tahsilatlar = fetch_all(
                """SELECT t.id,
                          COALESCE(t.makbuz_no, 'Makbuz-' || t.id) AS belge_no,
                          t.tahsilat_tarihi AS tarih,
                          COALESCE(
                            f.fatura_tarihi::date,
                            NULLIF(substring(COALESCE(t.aciklama, '') from '\\|AYLIK_TAH\\|([0-9]{4}-[0-9]{2}-[0-9]{2})\\|'), '')::date,
                            t.tahsilat_tarihi::date
                          ) AS eslesme_tarihi,
                          t.tutar, t.odeme_turu, t.aciklama AS tahsilat_aciklama, 'Tahsilat' AS tur
                   FROM tahsilatlar t
                   LEFT JOIN faturalar f ON f.id = t.fatura_id
                   WHERE (t.musteri_id = %s OR t.customer_id = %s OR f.musteri_id = %s)
                   ORDER BY t.tahsilat_tarihi DESC, t.id DESC""" + limit_sql,
                ((musteri_id, musteri_id, musteri_id, limit_n) if limit_n else (musteri_id, musteri_id, musteri_id)),
            )
    rows = []
    for r in faturalar:
        rows.append({
            "id": r.get("id"), "belge_no": r.get("belge_no") or "", "tarih": str(r.get("tarih") or "")[:10],
            "tur": "Fatura", "borc": float(r.get("tutar") or 0), "alacak": 0, "vade_tarihi": str(r.get("vade_tarihi") or "")[:10] if r.get("vade_tarihi") else None
        })
    for r in tahsilatlar:
        eslesme_tarihi = str(r.get("eslesme_tarihi") or r.get("tarih") or "")[:10]
        harf = _odeme_turu_harf(r.get("odeme_turu"))
        tah_aciklama_raw = (r.get("tahsilat_aciklama") or "").strip()
        marker_tarih = _marker_ay_tarihi_from_text(tah_aciklama_raw)
        tah_aciklama = _tahsilat_aciklama_temizle(tah_aciklama_raw)
        if marker_tarih:
            tah_aciklama = f"{_AY_ADLARI[marker_tarih.month - 1]} {marker_tarih.year} Tahsilat {harf}"
        if not tah_aciklama:
            try:
                td = datetime.strptime(eslesme_tarihi, "%Y-%m-%d").date()
                tah_aciklama = f"{_AY_ADLARI[td.month - 1]} {td.year} Tahsilat {harf}"
            except Exception:
                tah_aciklama = "Tahsilat " + harf
        rows.append({
            "id": "t-" + str(r.get("id")), "belge_no": r.get("belge_no") or "", "tarih": eslesme_tarihi,
            "tur": "Tahsilat", "borc": 0, "alacak": float(r.get("tutar") or 0), "vade_tarihi": None, "odeme_turu": r.get("odeme_turu"),
            "aciklama": tah_aciklama
        })
    rows.sort(key=lambda x: (x["tarih"], x["tur"] == "Fatura" and 0 or 1))
    bakiye = 0
    for r in rows:
        bakiye = bakiye + r["borc"] - r["alacak"]
        r["bakiye"] = round(bakiye, 2)
    return rows


# Ay adları (ekstre açıklama için)
_AY_ADLARI = ("Ocak", "Şubat", "Mart", "Nisan", "Mayıs", "Haziran",
              "Temmuz", "Ağustos", "Eylül", "Ekim", "Kasım", "Aralık")


def _date_add_months(d: date, n: int) -> date:
    """Sözleşme reel grid ile aynı: başlangıç tarihinden n ay ileri (gün taşması güvenli)."""
    if n == 0:
        return d
    total = d.year * 12 + (d.month - 1) + n
    ny = total // 12
    nm = total % 12 + 1
    last = calendar.monthrange(ny, nm)[1]
    nd = min(d.day, last)
    return date(ny, nm, nd)


def _reel_manual_is_explicit_override(y: int, y_start: int, mv: float, prev_eff) -> bool:
    """
    Kayıtlı reel dönem tutarı grid/ekstrede kullanılsın; yalnızca önceki dönemle birebir aynı
    DB kopyası (434→434) TÜFE zincirine bırakılır. Düşük/ yüksek fark etmez (540, 1008).
    """
    if y <= y_start:
        return False
    if prev_eff is None:
        return True
    try:
        prev_f = float(prev_eff)
        mv_f = float(mv)
    except (TypeError, ValueError):
        return True
    if not math.isfinite(prev_f) or prev_f <= 0 or not math.isfinite(mv_f) or mv_f < 0:
        return math.isfinite(mv_f) and mv_f >= 0
    if abs(mv_f - prev_f) <= 0.02:
        return False
    return True


def _reel_donem_ay_keys_for_period(bas_soz: date, donem_yil: int, artis_month: int, artis_day: int):
    """
    Reel dönem: kira artış (yıldönümü) ay/gününden itibaren 12 ay — JS sozlesmeReelDonemAyKeys ile uyumlu.
    Bu 12 ayda KDV dahil tutar sabit; TÜFE yalnızca sonraki dönem başında uygulanır.
    Örn. artış Şubat ise donem_yil=2026 → 2026-2 … 2027-1 (Ocak 2026 önceki dönemde kalır).
    """
    dm = max(1, min(12, int(artis_month)))
    dd = int(artis_day) if artis_day else bas_soz.day
    try:
        start = date(donem_yil, dm, dd)
    except ValueError:
        last = calendar.monthrange(donem_yil, dm)[1]
        start = date(donem_yil, dm, min(dd, last))
    keys = []
    for i in range(12):
        d = _date_add_months(start, i)
        keys.append(f"{d.year}-{d.month}")
    return keys


def _reel_ay_key_tutar_map_db_flat_only(
    bas_soz: date,
    artis_month: int,
    artis_day: int,
    manual_by_year: dict | None,
) -> dict[str, float]:
    """Aylık grid: yalnızca musteri_reel_donem_tutar (dönem yılı → 12 ay sabit KDV dahil); TÜFE zinciri yok."""
    if not bas_soz or not manual_by_year:
        return {}
    manual: dict[int, float] = {}
    for ky, val in manual_by_year.items():
        try:
            yi = int(ky)
            vf = round(float(val or 0), 2)
            if math.isfinite(vf) and vf > 0:
                manual[yi] = vf
        except (TypeError, ValueError):
            continue
    if not manual:
        return {}
    try:
        ad = int(artis_day) if artis_day is not None else bas_soz.day
    except (TypeError, ValueError):
        ad = bas_soz.day
    out: dict[str, float] = {}
    for d_y, tut in manual.items():
        for ky in _reel_donem_ay_keys_for_period(bas_soz, d_y, artis_month, ad):
            out[ky] = tut
    return out


def _reel_kart_yillik_toplam_haritasi(kyc: dict, tufe_map: dict, y_start: int, y_end: int) -> dict[int, float]:
    """Aylık kira kartı / _aylik_grid_contract_core ile aynı KDV dahil yıllık toplam (TÜFE zinciri)."""
    out: dict[int, float] = {}
    if not kyc:
        return out
    try:
        core = _aylik_grid_contract_core(dict(kyc), tufe_map or {})
    except Exception:
        core = None
    if not core or not isinstance(core.get("yillik_map"), dict):
        return out
    for y in range(int(y_start), int(y_end) + 1):
        try:
            v = float((core["yillik_map"] or {}).get(y) or 0)
        except (TypeError, ValueError):
            continue
        if math.isfinite(v) and v > 0:
            out[y] = round(v, 2)
    return out


def _reel_tufe_oran_yillik_gecis(tufe_map: dict, y: int, artis_month: int, y_start: int) -> float:
    inner = tufe_map.get(y) if isinstance(tufe_map, dict) else {}
    if not isinstance(inner, dict):
        inner = {}
    raw_o = inner.get(artis_month)
    if raw_o is None and artis_month is not None:
        raw_o = inner.get(str(artis_month))
    try:
        oran = float(raw_o or 0)
    except (TypeError, ValueError):
        oran = 0.0
    if (not oran or not math.isfinite(oran)) and y > y_start:
        inner_prev = tufe_map.get(y - 1) if isinstance(tufe_map, dict) else {}
        if isinstance(inner_prev, dict):
            raw_p = inner_prev.get(artis_month)
            if raw_p is None and artis_month is not None:
                raw_p = inner_prev.get(str(artis_month))
            try:
                oran2 = float(raw_p or 0)
            except (TypeError, ValueError):
                oran2 = 0.0
            if oran2 > 0 and math.isfinite(oran2):
                oran = oran2
    if not oran or not math.isfinite(oran):
        o3 = _tufe_latest_positive_oran_in_year_map(inner)
        if o3 > 0 and math.isfinite(o3):
            oran = o3
    if (not oran or not math.isfinite(oran)) and y > y_start:
        inner_prev2 = tufe_map.get(y - 1) or {} if isinstance(tufe_map, dict) else {}
        o4 = _tufe_latest_positive_oran_in_year_map(inner_prev2)
        if o4 > 0 and math.isfinite(o4):
            oran = o4
    return oran if (oran and math.isfinite(oran)) else 0.0


def _reel_donem_effective_yillik_for_ekstre(
    manual_by_year: dict, bas_soz: date, tufe_map: dict, artis_month: int, y_end: int, kyc=None
) -> dict:
    """
    Varsayılan: kira kartı yıllık toplam (KDV dahil + TÜFE, karma kira dahil).
    Elle reel dönem yılı girildiyse o yıl sabit; sonraki yıllar o tutardan TÜFE ile artar.
    """
    out = {}
    y_start = bas_soz.year
    bug = date.today().year
    mx = 0
    man: dict = {}
    if isinstance(manual_by_year, dict):
        for k, v in manual_by_year.items():
            try:
                yi = int(k)
                vf = float(v)
                if math.isfinite(vf) and vf >= 0:
                    man[yi] = round(vf, 2)
                    mx = max(mx, yi)
            except (TypeError, ValueError):
                continue
    y_end = max(int(y_end), bug + 2, mx + 5, y_start + 12)
    y_end = min(y_end, max(bug + 25, y_start + 30))
    kart = _reel_kart_yillik_toplam_haritasi(dict(kyc or {}), tufe_map, y_start, y_end) if kyc else {}
    last_explicit: int | None = None
    for y in range(y_start, y_end + 1):
        mv = man.get(y)
        prev_eff = out.get(y - 1)
        # İlk yıl karttan; sonraki yıllarda reel yalnızca önceki dönemden gerçek artışsa sabitlenir.
        if mv is not None and y > y_start and _reel_manual_is_explicit_override(y, y_start, mv, prev_eff):
            out[y] = mv
            last_explicit = y
            continue
        # DB'de önceki yılla aynı reel (434→434): TÜFE/kart — aylık hücre 623 ile ekstre uyumu.
        if mv is not None and y > y_start and not _reel_manual_is_explicit_override(y, y_start, mv, prev_eff):
            prev_f = float(prev_eff) if prev_eff is not None else 0.0
            chosen = None
            kv_dup = kart.get(y)
            if kv_dup is not None:
                try:
                    kvf = float(kv_dup)
                    if math.isfinite(kvf) and kvf > 0 and (prev_f <= 0 or kvf > prev_f * 1.001 + 0.01):
                        chosen = round(kvf, 2)
                except (TypeError, ValueError):
                    chosen = None
            if chosen is None and prev_f > 0:
                oran_dup = _reel_tufe_oran_yillik_gecis(tufe_map, y, artis_month, y_start)
                if oran_dup > 0:
                    nxt_dup = round(prev_f * (1 + oran_dup / 100.0), 2)
                    if math.isfinite(nxt_dup) and nxt_dup > prev_f + 0.02:
                        chosen = nxt_dup
            if chosen is None and kv_dup is not None:
                try:
                    kvf2 = float(kv_dup)
                    if math.isfinite(kvf2) and kvf2 > 0:
                        chosen = round(kvf2, 2)
                except (TypeError, ValueError):
                    pass
            if chosen is not None:
                out[y] = chosen
                last_explicit = y
                continue
        if last_explicit is not None and y > last_explicit and prev_eff and prev_eff > 0:
            prev_f = float(prev_eff)
            oran = _reel_tufe_oran_yillik_gecis(tufe_map, y, artis_month, y_start)
            if oran > 0:
                nxt = round(prev_f * (1 + oran / 100.0), 2)
                if math.isfinite(nxt) and nxt > 0:
                    out[y] = nxt
                    continue
            out[y] = round(prev_f, 2)
            continue
        kv = kart.get(y)
        if kv is not None and kv > 0:
            out[y] = kv
            continue
        if y > y_start and prev_eff and prev_eff > 0:
            prev_f = float(prev_eff)
            oran = _reel_tufe_oran_yillik_gecis(tufe_map, y, artis_month, y_start)
            if oran > 0:
                nxt = round(prev_f * (1 + oran / 100.0), 2)
                if math.isfinite(nxt) and nxt > 0:
                    out[y] = nxt
            else:
                out[y] = round(prev_f, 2)
    return out


def _reel_donem_effective_yilmap_orani_duzelt(
    out: dict,
    manual_by_year: dict,
    tufe_map: dict,
    y_start: int,
    y_end: int,
    kyc: dict,
) -> None:
    """
    TÜFE map'te hedef yılın artış ayı oranı yoksa reel zincir düz kalırdı.
    Aylık grid ile aynı _aylik_grid_contract_core yillik_map oranı out[yy]'yi günceller.
    """
    if not out or not kyc or not isinstance(out, dict):
        return
    kyc_d = dict(kyc or {})
    aylik_net = _aylik_grid_coerce_money(kyc_d.get("aylik_kira"))
    if aylik_net <= 0:
        return
    try:
        core = _aylik_grid_contract_core(kyc_d, tufe_map)
    except Exception:
        core = None
    if not core or not isinstance(core.get("yillik_map"), dict):
        return
    yillik_map = core["yillik_map"]
    man = manual_by_year if isinstance(manual_by_year, dict) else {}
    try:
        y_lo = int(y_start)
        y_hi = int(y_end)
    except (TypeError, ValueError):
        return
    for yy in range(y_lo + 1, y_hi + 1):
        prev = out.get(yy - 1)
        cur = out.get(yy)
        if prev is None or cur is None:
            continue
        try:
            prev_n = float(prev)
            cur_n = float(cur)
        except (TypeError, ValueError):
            continue
        if not math.isfinite(prev_n) or prev_n <= 0 or not math.isfinite(cur_n):
            continue
        mv = man.get(yy)
        if mv is not None:
            try:
                mvn = round(float(mv), 2)
            except (TypeError, ValueError):
                mvn = None
            if mvn is not None and math.isfinite(mvn) and mvn >= 0:
                if _reel_manual_is_explicit_override(yy, y_lo, mvn, prev_n):
                    continue
        if cur_n > prev_n * 1.001:
            continue
        ym = yillik_map.get(yy)
        ym1 = yillik_map.get(yy - 1)
        if ym is None or ym1 is None:
            continue
        try:
            ymf = float(ym)
            ym1f = float(ym1)
        except (TypeError, ValueError):
            continue
        if not math.isfinite(ymf) or not math.isfinite(ym1f) or ym1f <= 0:
            continue
        ratio = ymf / ym1f
        if not math.isfinite(ratio) or ratio <= 0:
            continue
        out[yy] = round(prev_n * ratio, 2)


def _reel_ay_key_tutar_map_musteri(
    musteri_id, bas_soz: date, artis_month: int, artis_day: int, tufe_map: dict, y_end: int,
    manual_by_year=None,
    kyc_for_yilmap=None,
) -> dict:
    """Ay anahtarı (Y-M) -> KDV dahil reel tutar; kayıt yoksa {} (ekstre eski aylık+TÜFE yoluna düşer).
    manual_by_year: {yil: tutar} önceden yüklenmişse DB sorgusu yapılmaz (toplu rapor için). None = DB'den oku."""
    if not bas_soz:
        return {}
    if isinstance(bas_soz, datetime):
        bas_soz = bas_soz.date()
    try:
        am = int(artis_month)
    except (TypeError, ValueError):
        am = bas_soz.month
    artis_month = max(1, min(12, am))
    try:
        ad_raw = int(artis_day) if artis_day is not None else bas_soz.day
    except (TypeError, ValueError):
        ad_raw = bas_soz.day
    manual = {}
    if manual_by_year is not None:
        for ky, val in (manual_by_year.items() if isinstance(manual_by_year, dict) else []):
            try:
                manual[int(ky)] = float(val or 0)
            except (TypeError, ValueError):
                continue
    else:
        manual = _musteri_reel_donem_manual_dict_from_db(musteri_id)
    effective = _reel_donem_effective_yillik_for_ekstre(
        manual, bas_soz, tufe_map, artis_month, y_end, kyc=kyc_for_yilmap
    )
    if kyc_for_yilmap and effective:
        try:
            ys = bas_soz.year
            y_hi = max(int(k) for k in effective.keys())
            _reel_donem_effective_yilmap_orani_duzelt(
                effective, manual, tufe_map, ys, y_hi, kyc_for_yilmap
            )
        except Exception:
            pass
    out = {}
    y_start = bas_soz.year
    ad = ad_raw
    y_paint_end = y_end
    if effective:
        try:
            y_paint_end = max(int(y_paint_end), max(int(k) for k in effective))
        except (TypeError, ValueError):
            pass
    for d_y in range(y_start, y_paint_end + 1):
        if d_y not in effective:
            continue
        tut = effective[d_y]
        if tut is None or not math.isfinite(float(tut)):
            continue
        tut = round(float(tut), 2)
        if tut <= 0:
            continue
        for ky in _reel_donem_ay_keys_for_period(bas_soz, d_y, artis_month, ad):
            out[ky] = tut
    return out


def _aylik_grid_payload_has_month(payload: dict, ref_y: int, ref_m: int) -> bool:
    """Önbellekteki aylar listesi bu takvim ayını içeriyor mu (firma özeti için)."""
    if not isinstance(payload, dict):
        return False
    try:
        ry, rm = int(ref_y), int(ref_m)
    except (TypeError, ValueError):
        return False
    key = f"{ry}-{rm}"
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict):
            continue
        ak = str(a.get("ay_key") or "").strip()
        if ak == key:
            return True
        try:
            if int(a.get("yil") or 0) == ry and int(a.get("ay") or 0) == rm:
                return True
        except (TypeError, ValueError):
            continue
    return False


def firma_ozet_aylik_grid_hucre_kdv_dahil(
    musteri_id: int,
    ref_y: int,
    ref_m: int,
    tufe_map: dict,
    kyc_row: dict,
    cache_payload,
    manual_reel_by_year: dict,
    *,
    skip_disk_cache: bool = False,
    skip_reel_overlay: bool = False,
) -> float:
    """
    Müşteriler (firma özet) raporu: takvim ayı (ref_y/ref_m) için KDV dahil tutar.
    Sözleşmeler aylık/reel grid ile aynı mantık: önbellek veya canlı TÜFE gridi, üzerine reel dönem haritası.
    manual_reel_by_year: {donem_yil: tutar} — toplu raporda DB tekrarını önlemek için (boş dict geçerli).
    """
    try:
        ref_y = int(ref_y)
        ref_m = int(ref_m)
    except (TypeError, ValueError):
        return 0.0
    if ref_m < 1 or ref_m > 12:
        return 0.0
    key = f"{ref_y}-{ref_m}"
    kyc = dict(kyc_row or {})

    def _pd(v):
        return _aylik_grid_coerce_date(v)

    payload = None
    core_fast = None
    base_one = 0.0
    if skip_disk_cache:
        try:
            core_fast = _aylik_grid_contract_core(kyc, tufe_map)
            base_one = _aylik_grid_single_month_kdv_from_core(core_fast, ref_y, ref_m) if core_fast else 0.0
        except Exception:
            core_fast = None
            base_one = 0.0
    else:
        cache_ok = (
            cache_payload
            and isinstance(cache_payload, dict)
            and _aylik_grid_cache_matches_kyc(musteri_id, cache_payload)
            and _aylik_grid_payload_has_month(cache_payload, ref_y, ref_m)
        )
        if cache_ok:
            payload = cache_payload
        else:
            try:
                payload = _aylik_grid_compute(musteri_id, kyc, tufe_map)
            except Exception:
                payload = None
        if not payload or not isinstance(payload.get("aylar"), list):
            return 0.0

    bas_soz = _pd(kyc.get("sozlesme_tarihi"))
    if not bas_soz:
        if skip_disk_cache and core_fast:
            bas_soz = core_fast["bas"]
        elif not skip_disk_cache:
            try:
                bas_soz = datetime.strptime(str(payload.get("baslangic") or "")[:10], "%Y-%m-%d").date()
            except Exception:
                return 0.0
    if not bas_soz:
        return 0.0

    artis_d = _pd(kyc.get("kira_artis_tarihi")) or bas_soz
    artis_month = int(artis_d.month)
    artis_day = int(artis_d.day)
    bit = _pd(kyc.get("sozlesme_bitis"))
    if not bit:
        if skip_disk_cache and core_fast:
            bit = core_fast["bit"]
        elif not skip_disk_cache:
            try:
                bit = datetime.strptime(str(payload.get("bitis") or "")[:10], "%Y-%m-%d").date()
            except Exception:
                bit = bas_soz
        else:
            bit = bas_soz
    if not skip_reel_overlay:
        bugun_y = date.today()
        y_end = max(ref_y, bugun_y.year, bas_soz.year, bit.year)
        reel_manual = manual_reel_by_year if isinstance(manual_reel_by_year, dict) else {}
        try:
            reel_map = _reel_ay_key_tutar_map_musteri(
                musteri_id,
                bas_soz,
                artis_month,
                artis_day,
                tufe_map,
                y_end,
                manual_by_year=reel_manual,
                kyc_for_yilmap=kyc,
            )
        except Exception:
            reel_map = {}
        if key in reel_map:
            try:
                v = float(reel_map[key])
                if math.isfinite(v) and v > 0:
                    return round(v, 2)
            except (TypeError, ValueError):
                pass
    if skip_disk_cache:
        return round(base_one, 2) if math.isfinite(base_one) else 0.0
    for a in payload["aylar"]:
        ak = str(a.get("ay_key") or "")
        if ak == key or (int(a.get("yil") or 0) == ref_y and int(a.get("ay") or 0) == ref_m):
            try:
                v = float(a.get("tutar_kdv_dahil") or 0)
                return round(v, 2) if math.isfinite(v) else 0.0
            except (TypeError, ValueError):
                return 0.0
    return 0.0


def _firma_ozet_normalize_tahsil_ay_key(k: str) -> str | None:
    """Tahsilat ay anahtarını grid ile uyumlu YYYY-M biçimine indirger."""
    s = str(k or "").strip()
    if not s:
        return None
    parts = s.split("-")
    if len(parts) < 2:
        return None
    try:
        y = int(parts[0])
        m = int(parts[1])
    except (TypeError, ValueError):
        return None
    if m < 1 or m > 12:
        return None
    return f"{y}-{m}"


def _aylik_grid_acik_tutar_ay_keys_normalized(musteri_id: int) -> set[str]:
    """Bu müşteride açık (durum != odendi) |AYLIK_TUTAR| faturası bulunan aylar (YYYY-M)."""
    rows = fetch_all(
        """
        SELECT COALESCE(notlar, '') AS notlar
        FROM faturalar
        WHERE musteri_id = %s
          AND COALESCE(notlar, '') LIKE '%%|AYLIK_TUTAR|%%'
          AND COALESCE(durum, '') != 'odendi'
        """,
        (musteri_id,),
    ) or []
    out: set[str] = set()
    for r in rows:
        for iso in re.findall(r"\|AYLIK_TUTAR\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", str((r or {}).get("notlar") or "")):
            nk = _firma_ozet_normalize_tahsil_ay_key(iso)
            if nk:
                out.add(nk)
    return out


def _aylik_tahsil_edilen_aylar_set_normalized(musteri_id: int, payload=None) -> set[str]:
    """Aylık grid ile aynı: tam ödenmiş ay anahtarları (kalan ≤ 0,01 TL), YYYY-M normalize."""
    if payload is None:
        payload = _read_aylik_grid_cache_payload(musteri_id)
    if payload is None:
        payload = _build_aylik_grid_cache_payload(
            musteri_id, tufe_map=_tufe_map_by_year_month_cached()
        )
    return _aylik_tahsil_edilen_aylar_from_payload(payload)


def _aylik_tutar_fatura_id_map(musteri_id: int) -> dict[str, int]:
    """|AYLIK_TUTAR|YYYY-MM-DD| → fatura id (ekstre kira satırı düzenle/sil)."""
    out: dict[str, int] = {}
    rows = fetch_all(
        """
        SELECT id, COALESCE(notlar, '') AS notlar
        FROM faturalar
        WHERE musteri_id = %s AND COALESCE(notlar, '') LIKE '%%|AYLIK_TUTAR|%%'
        ORDER BY id DESC
        """,
        (int(musteri_id),),
    ) or []
    for r in rows:
        try:
            fid = int(r.get("id") or 0)
        except (TypeError, ValueError):
            fid = 0
        if fid <= 0:
            continue
        ac = str(r.get("notlar") or "")
        for iso in re.findall(r"\|AYLIK_TUTAR\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac):
            if iso not in out:
                out[iso] = fid
    return out


def _ekstre_invalidate_after_change(
    musteri_id: int, affected_isos: list | None = None
) -> None:
    """Ekstre düzenle/sil: önce panel DB (grid tahsilden-cikar ile aynı), sonra grid önbelleği."""
    _cari_ekstre_api_cache.clear()
    _invalidate_aylik_grid_payload_mem(musteri_id)
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return
    try:
        _musteri_kyc_grid_mem.pop(mid, None)
    except (TypeError, ValueError):
        pass
    try:
        payload = _upsert_aylik_grid_cache(mid)
    except Exception:
        payload = None
        logging.getLogger(__name__).exception(
            "ekstre invalidate grid cache mid=%s", musteri_id
        )
    tahsil_map = _aylik_tahsil_tutar_map(mid)
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    iso_check: set[str] = set()
    for raw_iso in affected_isos or []:
        try:
            dd = datetime.strptime(str(raw_iso)[:10], "%Y-%m-%d").date()
            iso_check.add(date(dd.year, dd.month, 1).isoformat())
        except ValueError:
            pass
    for iso_k in (_load_musteri_panel_by_iso(mid) or {}).keys():
        iso_check.add(str(iso_k)[:10])
    borc_satirlar = []
    for iso_k in iso_check:
        try:
            db_t = round(float(tahsil_map.get(iso_k) or 0), 2)
        except (TypeError, ValueError):
            db_t = 0.0
        if db_t > tol:
            continue
        try:
            y_i, m_i, _ = str(iso_k)[:10].split("-")
            borc_satirlar.append({"yil": int(y_i), "ay": int(m_i)})
        except (ValueError, TypeError):
            continue
    try:
        if borc_satirlar:
            sync_musteri_panel_borclu_from_satirlar(mid, borc_satirlar)
        by_iso = _panel_by_iso_from_tahsil_map(
            mid, payload, None, trust_grid_odenen=False
        )
        _save_musteri_panel_by_iso(mid, by_iso, prune_no_db_tahsil=True)
        if isinstance(payload, dict):
            panel_now = _load_musteri_panel_by_iso(mid)
            _apply_panel_by_iso_to_grid_payload(payload, panel_now)
            _ensure_aylik_grid_cache_table()
            execute(
                """
                INSERT INTO musteri_aylik_grid_cache (musteri_id, payload, updated_at)
                VALUES (%s, %s, NOW())
                ON CONFLICT (musteri_id)
                DO UPDATE SET payload = EXCLUDED.payload, updated_at = NOW()
                """,
                (mid, json.dumps(payload, ensure_ascii=False)),
            )
            try:
                _aylik_grid_payload_mem[mid] = (time.time(), payload)
            except (TypeError, ValueError):
                pass
            _invalidate_aylik_grid_payload_mem(mid)
        else:
            _upsert_aylik_grid_cache(mid)
    except Exception:
        logging.getLogger(__name__).exception(
            "ekstre invalidate panel sync mid=%s", musteri_id
        )


def _fatura_tutar_kdv_split(toplam_kdv_dahil: float, kira_nakit: bool, kdv_oran: float = 20.0) -> tuple[float, float, float]:
    toplam = round(float(toplam_kdv_dahil or 0), 2)
    if toplam <= 0:
        return 0.0, 0.0, 0.0
    if kira_nakit:
        return toplam, 0.0, toplam
    ko = float(kdv_oran or 20)
    if not math.isfinite(ko) or ko < 0:
        ko = 20.0
    net = round(toplam / (1 + ko / 100.0), 2)
    kdv = round(toplam - net, 2)
    net = round(toplam - kdv, 2)
    return net, kdv, toplam


def _ekstre_tahsil_rows_for_musteri(musteri_id: int) -> list:
    """Müşterinin tüm tahsilatları (ekstre için tek sorgu)."""
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return []
    if mid <= 0:
        return []
    return fetch_all(
        """
        SELECT t.id, COALESCE(t.tutar, 0) AS tutar,
               COALESCE(t.aciklama, '') AS aciklama,
               t.tahsilat_tarihi, f.fatura_tarihi
        FROM tahsilatlar t
        LEFT JOIN faturalar f ON f.id = t.fatura_id
        WHERE (t.musteri_id = %s OR t.customer_id = %s)
          AND COALESCE(t.tutar, 0) > 0
        """,
        (mid, mid),
    ) or []


def _ekstre_tahsil_batch_maps_from_rows(rows) -> dict:
    """Marker / eslesme / tahsilat-tarihi-ay toplamları — ay başına tekrarlayan SQL yok."""
    marker: dict[str, float] = {}
    eslesme: dict[str, float] = {}
    tarih_ay: dict[str, float] = {}
    for r in rows or []:
        if not isinstance(r, dict):
            continue
        try:
            tut = round(float(r.get("tutar") or 0), 2)
        except (TypeError, ValueError):
            continue
        if tut <= 0:
            continue
        ac = str(r.get("aciklama") or "")
        pay_tokens = re.findall(
            r"\|AYLIK_PAY\|([0-9]{4}-[0-9]{2}-[0-9]{2})=([0-9]+(?:\.[0-9]+)?)\|",
            ac,
        )
        if pay_tokens:
            for iso_raw, tut_raw in pay_tokens:
                try:
                    dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                    iso_m = date(dd.year, dd.month, 1).isoformat()
                    pv = round(float(tut_raw or 0), 2)
                except (TypeError, ValueError):
                    continue
                if pv <= 0:
                    continue
                marker[iso_m] = round(float(marker.get(iso_m, 0)) + pv, 2)
        else:
            marker_isos = re.findall(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac)
            if len(marker_isos) == 1:
                try:
                    dd = datetime.strptime(marker_isos[0][:10], "%Y-%m-%d").date()
                    iso_m = date(dd.year, dd.month, 1).isoformat()
                    marker[iso_m] = round(float(marker.get(iso_m, 0)) + tut, 2)
                except ValueError:
                    pass
            elif len(marker_isos) > 1:
                cents_total = int(round(tut * 100))
                n = len(marker_isos)
                base = cents_total // n if n else 0
                rem = cents_total % n if n else 0
                for i, iso_raw in enumerate(marker_isos):
                    share_cents = base + (1 if i < rem else 0)
                    if share_cents <= 0:
                        continue
                    try:
                        dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                        iso_m = date(dd.year, dd.month, 1).isoformat()
                        marker[iso_m] = round(float(marker.get(iso_m, 0)) + share_cents / 100.0, 2)
                    except ValueError:
                        continue
        iso_e = _tahsil_row_ekstre_eslesme_ay_iso(r)
        if iso_e:
            eslesme[iso_e] = round(float(eslesme.get(iso_e, 0)) + tut, 2)
        d2 = r.get("tahsilat_tarihi")
        iso_t = None
        if d2 and hasattr(d2, "year"):
            iso_t = date(int(d2.year), int(d2.month), 1).isoformat()
        elif d2:
            try:
                dd = datetime.strptime(str(d2)[:10], "%Y-%m-%d").date()
                iso_t = date(dd.year, dd.month, 1).isoformat()
            except Exception:
                iso_t = None
        if iso_t:
            tarih_ay[iso_t] = round(float(tarih_ay.get(iso_t, 0)) + tut, 2)
    return {"marker": marker, "eslesme": eslesme, "tarih_ay": tarih_ay}


def _grid_payload_ay_odenen_kdv(brut, odenen_mevcut, mk_t, tm_t, te_t, tt_t, tol=None):
    """Grid/panel: |AYLIK_PAY| (tm_t) ve |AYLIK_TAH| (mk_t) birleşimi; kısmi ay doğru kalsın."""
    tol = float(tol if tol is not None else AYLIK_GRID_TAM_ODENDI_TOLERANS)
    try:
        brut = round(float(brut or 0), 2)
    except (TypeError, ValueError):
        brut = 0.0
    mk_t = round(float(mk_t or 0), 2)
    tm_t = round(float(tm_t or 0), 2)
    best = round(max(mk_t, tm_t), 2)
    if brut <= tol or best <= tol:
        return 0.0
    return round(min(best, brut), 2)


def _grid_payload_marker_panel_tam_kapandi(
    mk_t: float, brut: float, yil: int, manual_reel_by_year: dict | None, tol: float
) -> bool:
    """|AYLIK_TAH| marker reel dönem tutarını (panel) karşılıyorsa grid tam ödendi say."""
    mk_t = round(float(mk_t or 0), 2)
    if mk_t <= tol:
        return False
    reel_kap = 0.0
    if isinstance(manual_reel_by_year, dict):
        try:
            reel_kap = round(float(manual_reel_by_year.get(int(yil)) or 0), 2)
        except (TypeError, ValueError):
            reel_kap = 0.0
    if reel_kap > tol:
        return mk_t + tol >= reel_kap
    try:
        b = round(float(brut or 0), 2)
    except (TypeError, ValueError):
        b = 0.0
    return b > tol and mk_t + tol >= b


def _ekstre_payload_odenen_zenginlestir(
    payload, tahsil_map, batch_maps, manual_reel_by_year=None
) -> None:
    """Grid payload aylarının ödenen alanını batch tahsil özetleriyle güçlendir."""
    if not isinstance(payload, dict):
        return
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    mk_d = (batch_maps or {}).get("marker") or {}
    te_d = (batch_maps or {}).get("eslesme") or {}
    tt_d = (batch_maps or {}).get("tarih_ay") or {}
    tm_d = tahsil_map or {}
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict):
            continue
        try:
            brut = float(a.get("brut_tutar_kdv") or a.get("tutar_kdv_dahil") or 0)
        except (TypeError, ValueError):
            brut = 0.0
        try:
            odenen = float(a.get("odenen_tutar_kdv") or 0)
        except (TypeError, ValueError):
            odenen = 0.0
        try:
            iso_m1 = date(int(a.get("yil")), int(a.get("ay")), 1).isoformat()
        except (TypeError, ValueError):
            continue
        if not iso_m1 or brut <= tol:
            continue
        mk_t = round(float(mk_d.get(iso_m1) or 0), 2)
        tm_t = round(float(tm_d.get(iso_m1) or 0), 2)
        te_t = round(float(te_d.get(iso_m1) or 0), 2)
        tt_t = round(float(tt_d.get(iso_m1) or 0), 2)
        try:
            yy_en = int(a.get("yil"))
        except (TypeError, ValueError):
            yy_en = 0
        reel_kap = 0.0
        if isinstance(manual_reel_by_year, dict) and yy_en:
            try:
                reel_kap = round(float(manual_reel_by_year.get(yy_en) or 0), 2)
            except (TypeError, ValueError):
                reel_kap = 0.0
        kap = reel_kap if reel_kap > tol else round(brut, 2)
        odenen = _grid_payload_ay_odenen_kdv(brut, odenen, mk_t, tm_t, te_t, tt_t, tol)
        if _grid_payload_marker_panel_tam_kapandi(mk_t, brut, yy_en, manual_reel_by_year, tol):
            a["odenen_tutar_kdv"] = round(mk_t, 2)
            a["kalan_tutar_kdv"] = 0.0
            a["tahsil_edildi"] = True
            a["kismi_tahsilat"] = False
        elif mk_t > tol and kap > tol:
            kalan_p = round(max(kap - mk_t, 0), 2)
            a["odenen_tutar_kdv"] = round(mk_t, 2)
            a["kalan_tutar_kdv"] = kalan_p
            a["tahsil_edildi"] = kalan_p <= tol
            a["kismi_tahsilat"] = kalan_p > tol
        else:
            kalan = round(max(brut - odenen, 0), 2)
            a["odenen_tutar_kdv"] = round(odenen, 2)
            a["kalan_tutar_kdv"] = kalan
            a["tahsil_edildi"] = kalan <= tol
            a["kismi_tahsilat"] = odenen > tol and kalan > tol
        a["tutar_kdv_dahil"] = round(max(brut, 0.01), 2)


def _aylik_grid_cache_payload_tahsil_guncelle(musteri_id, payload):
    """Disk önbelleği dönerken tahsil/kısmi alanlarını DB ile yenile (sayfa yenilemesinde turuncu kalan)."""
    if not isinstance(payload, dict):
        return payload
    try:
        mid = int(musteri_id)
        panel_db = _load_musteri_panel_by_iso(mid)
        tahsil_rows = _ekstre_tahsil_rows_for_musteri(mid)
        tahsil_map = _aylik_tahsil_tutar_map(mid, tahsil_rows=tahsil_rows)
        batch = _ekstre_tahsil_batch_maps_from_rows(tahsil_rows)
        manual_reel = _musteri_reel_donem_manual_dict_from_db(mid)
        _ekstre_payload_odenen_zenginlestir(payload, tahsil_map, batch, manual_reel_by_year=manual_reel)
        if panel_db:
            _apply_panel_by_iso_to_grid_payload(payload, panel_db)
        else:
            by_iso = _panel_by_iso_from_tahsil_map(mid, payload, None)
            if by_iso:
                _save_musteri_panel_by_iso(mid, by_iso)
                _apply_panel_by_iso_to_grid_payload(payload, by_iso)
        payload["tahsilat_imza"] = _aylik_tahsil_cache_imza(mid)
    except Exception:
        logging.getLogger(__name__).exception(
            "aylik_grid_cache_payload_tahsil_guncelle musteri_id=%s", musteri_id
        )
    return payload


def _tahsil_toplam_eslesme_ay_iso(musteri_id: int, iso_first: str) -> float:
    """Ekstre eşleme ayı (fatura → marker → tahsilat) bu iso olan tahsilatların toplamı."""
    iso = str(iso_first or "").strip()[:10]
    if len(iso) < 10:
        return 0.0
    try:
        want = date.fromisoformat(iso)
    except ValueError:
        return 0.0
    total = 0.0
    rows = fetch_all(
        """
        SELECT t.tutar, t.tahsilat_tarihi, t.aciklama, f.fatura_tarihi
        FROM tahsilatlar t
        LEFT JOIN faturalar f ON f.id = t.fatura_id
        WHERE (t.musteri_id = %s OR t.customer_id = %s)
          AND COALESCE(t.tutar, 0) > 0
        """,
        (int(musteri_id), int(musteri_id)),
    ) or []
    for r in rows:
        iso_e = _tahsil_row_ekstre_eslesme_ay_iso(r)
        if iso_e != want.isoformat():
            continue
        try:
            total += float(r.get("tutar") or 0)
        except (TypeError, ValueError):
            continue
    return round(total, 2) if math.isfinite(total) else 0.0


def _tahsil_toplam_tahsilat_tarihi_ay(musteri_id: int, ref_y: int, ref_m: int) -> float:
    """Tahsilat tarihi bu takvim ayında olan ödemeler (fatura ayı Mart olsa bile Nisan tahsilatı sayılır)."""
    try:
        y, m = int(ref_y), int(ref_m)
        if m < 1 or m > 12:
            return 0.0
    except (TypeError, ValueError):
        return 0.0
    row = fetch_one(
        """
        SELECT COALESCE(SUM(COALESCE(t.tutar, 0)), 0)::numeric AS s
        FROM tahsilatlar t
        WHERE (t.musteri_id = %s OR t.customer_id = %s)
          AND COALESCE(t.tutar, 0) > 0
          AND t.tahsilat_tarihi IS NOT NULL
          AND EXTRACT(YEAR FROM t.tahsilat_tarihi::date) = %s
          AND EXTRACT(MONTH FROM t.tahsilat_tarihi::date) = %s
        """,
        (int(musteri_id), int(musteri_id), y, m),
    ) or {}
    try:
        v = float(row.get("s") or 0)
    except (TypeError, ValueError):
        v = 0.0
    return round(v, 2) if math.isfinite(v) else 0.0


def _ekstre_tahsil_alacak_hucre(
    musteri_id: int,
    iso_day: str,
    fifo_alacak: float,
    hedef_borc: float,
    tahsil_map: dict | None = None,
    grid_odenen: float | None = None,
    grid_kalan: float | None = None,
    batch_maps: dict | None = None,
) -> float:
    """Ekstre tahsilat = aylık hücre ödemesi; tam kapandıysa o ayın kira (grid) tutarı."""
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    hedef = round(float(hedef_borc or 0), 2)
    if hedef <= tol:
        return 0.0
    iso = str(iso_day or "").strip()[:10]
    if not iso:
        return 0.0
    try:
        kl = float(grid_kalan) if grid_kalan is not None else None
        if kl is not None and math.isfinite(kl) and kl <= tol:
            return hedef
    except (TypeError, ValueError):
        pass
    try:
        og = float(grid_odenen) if grid_odenen is not None else None
        if og is not None and math.isfinite(og) and og >= hedef - tol:
            return hedef
    except (TypeError, ValueError):
        pass
    out = round(float(fifo_alacak or 0), 2)
    if batch_maps is not None:
        try:
            tm = round(float((tahsil_map or {}).get(iso) or 0), 2)
        except (TypeError, ValueError):
            tm = 0.0
        if tm > tol:
            out = round(max(out, tm), 2)
        mk = round(float((batch_maps.get("marker") or {}).get(iso) or 0), 2)
        if mk > tol:
            out = round(max(out, mk), 2)
        te = round(float((batch_maps.get("eslesme") or {}).get(iso) or 0), 2)
        if te > tol:
            out = round(max(out, te), 2)
        tt = round(float((batch_maps.get("tarih_ay") or {}).get(iso) or 0), 2)
        if tt > tol:
            out = round(max(out, tt), 2)
    else:
        if tahsil_map is None:
            try:
                tahsil_map = _aylik_tahsil_tutar_map(int(musteri_id))
            except Exception:
                tahsil_map = {}
        try:
            tm = round(float((tahsil_map or {}).get(iso) or 0), 2)
        except (TypeError, ValueError):
            tm = 0.0
        if tm > tol:
            out = round(max(out, tm), 2)
        try:
            mk = _aylik_tahsil_marker_odeme_toplam_for_iso(int(musteri_id), iso)
        except Exception:
            mk = 0.0
        if mk > tol:
            out = round(max(out, mk), 2)
        try:
            te = _tahsil_toplam_eslesme_ay_iso(int(musteri_id), iso)
        except Exception:
            te = 0.0
        if te > tol:
            out = round(max(out, te), 2)
        try:
            yy, mm, _ = iso.split("-")
            tt = _tahsil_toplam_tahsilat_tarihi_ay(int(musteri_id), int(yy), int(mm))
        except (TypeError, ValueError):
            tt = 0.0
        if tt > tol:
            out = round(max(out, tt), 2)
    if out >= hedef - tol:
        return hedef
    return round(min(out, hedef), 2)


def _ekstre_hucre_borc(
    iso: str,
    fallback: float,
    grid_tutar_by_iso: dict | None,
    reel_ay_map: dict | None = None,
) -> float:
    """Ekstre borç: kayıtlı reel dönem (DB) → grid → hesaplanan fallback."""
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    iso = str(iso or "").strip()[:10]
    try:
        fb = round(float(fallback or 0), 2)
    except (TypeError, ValueError):
        fb = 0.0
    if reel_ay_map and isinstance(reel_ay_map, dict):
        try:
            dd = datetime.strptime(iso[:10], "%Y-%m-%d").date()
            rk = f"{dd.year}-{dd.month}"
            if rk in reel_ay_map:
                r0 = round(float(reel_ay_map.get(rk) or 0), 2)
                if math.isfinite(r0) and r0 > tol:
                    return r0
        except (TypeError, ValueError):
            pass
    g = 0.0
    if isinstance(grid_tutar_by_iso, dict) and iso in grid_tutar_by_iso:
        try:
            g = round(float(grid_tutar_by_iso.get(iso) or 0), 2)
        except (TypeError, ValueError):
            g = 0.0
    if g > tol:
        return g
    if fb > tol:
        return fb
    return fb


def _ekstre_hucre_alacak(
    iso: str,
    grid_tutar_by_iso: dict | None,
    grid_odenen_by_iso: dict | None,
    grid_kalan_by_iso: dict | None,
    grid_tahsil_edildi_by_iso: dict | None = None,
    grid_kismi_by_iso: dict | None = None,
    ayda_tahsilat_var: bool = False,
    reel_ay_map: dict | None = None,
) -> float | None:
    """
    Ekstre tahsilat = aylık grid hücresinde görünen brüt (kira ile aynı rakam).
    Ayda tahsilat yoksa 0; grid'de ay yoksa None → çağıran yedek yola düşer.
    """
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    iso = str(iso or "").strip()[:10]
    hedef = None
    if reel_ay_map and isinstance(reel_ay_map, dict):
        try:
            dp = datetime.strptime(iso[:10], "%Y-%m-%d").date()
            rk = f"{dp.year}-{dp.month}"
            if rk in reel_ay_map:
                rv = float(reel_ay_map.get(rk) or 0)
                if math.isfinite(rv) and rv > tol:
                    hedef = round(rv, 2)
        except (TypeError, ValueError):
            hedef = None
    if hedef is None:
        if not isinstance(grid_tutar_by_iso, dict) or iso not in grid_tutar_by_iso:
            return None
        try:
            hedef = round(float(grid_tutar_by_iso.get(iso) or 0), 2)
        except (TypeError, ValueError):
            hedef = 0.0
    if hedef <= tol:
        return 0.0
    has_pay = bool(ayda_tahsilat_var)
    if not has_pay and bool((grid_tahsil_edildi_by_iso or {}).get(iso)):
        has_pay = True
    if not has_pay:
        try:
            og = round(max(0.0, float((grid_odenen_by_iso or {}).get(iso) or 0)), 2)
            has_pay = og > tol
        except (TypeError, ValueError):
            has_pay = False
    if not has_pay:
        return 0.0
    return hedef


def _aylik_tahsil_marker_odeme_toplam_for_iso(musteri_id: int, iso_day: str) -> float:
    """|AYLIK_TAH|YYYY-MM-DD| bu ay için tahsilat tutarları toplamı (FIFO kaçırsa da). Tarih 2024-9-1 gibi yazılmışsa da yakala."""
    sday = str(iso_day or "").strip()[:10]
    if not sday or len(sday) < 10:
        return 0.0
    pat1 = f"%|AYLIK_TAH|{sday}|%"
    loose_m = None
    try:
        dp = datetime.strptime(sday, "%Y-%m-%d").date()
        loose_m = f"{dp.year}-{dp.month}-{dp.day}"
    except Exception:
        pass
    pat2 = f"%|AYLIK_TAH|{loose_m}|%" if loose_m else None
    if pat2 and pat2 != pat1:
        row = fetch_one(
            """
            SELECT COALESCE(SUM(COALESCE(tutar, 0)), 0)::numeric AS s
            FROM tahsilatlar t
            WHERE (t.musteri_id = %s OR t.customer_id = %s)
              AND COALESCE(t.tutar, 0) > 0
              AND (
                    COALESCE(t.aciklama, '') LIKE %s
                 OR COALESCE(t.aciklama, '') LIKE %s
              )
            """,
            (musteri_id, musteri_id, pat1, pat2),
        ) or {}
    else:
        row = fetch_one(
            """
            SELECT COALESCE(SUM(COALESCE(tutar, 0)), 0)::numeric AS s
            FROM tahsilatlar t
            WHERE (t.musteri_id = %s OR t.customer_id = %s)
              AND COALESCE(t.tutar, 0) > 0
              AND COALESCE(t.aciklama, '') LIKE %s
            """,
            (musteri_id, musteri_id, pat1),
        ) or {}
    try:
        v = float(row.get("s") or 0)
    except (TypeError, ValueError):
        v = 0.0
    return round(v, 2) if math.isfinite(v) else 0.0


def _aylik_tahsil_marker_kayitlari_for_iso(musteri_id: int, iso_day: str) -> list[dict]:
    """|AYLIK_TAH|YYYY-MM-DD| marker'lı tahsilat satırları (en yeni önce)."""
    sday = str(iso_day or "").strip()[:10]
    if not sday or len(sday) < 10:
        return []
    pat1 = f"%|AYLIK_TAH|{sday}|%"
    loose_m = None
    try:
        dp = datetime.strptime(sday, "%Y-%m-%d").date()
        loose_m = f"{dp.year}-{dp.month}-{dp.day}"
    except Exception:
        pass
    pat2 = f"%|AYLIK_TAH|{loose_m}|%" if loose_m else None
    if pat2 and pat2 != pat1:
        rows = fetch_all(
            """
            SELECT t.id, COALESCE(t.tutar, 0) AS tutar, t.makbuz_no
            FROM tahsilatlar t
            WHERE (t.musteri_id = %s OR t.customer_id = %s)
              AND COALESCE(t.tutar, 0) > 0
              AND (
                    COALESCE(t.aciklama, '') LIKE %s
                 OR COALESCE(t.aciklama, '') LIKE %s
              )
            ORDER BY t.id DESC
            """,
            (musteri_id, musteri_id, pat1, pat2),
        ) or []
    else:
        rows = fetch_all(
            """
            SELECT t.id, COALESCE(t.tutar, 0) AS tutar, t.makbuz_no
            FROM tahsilatlar t
            WHERE (t.musteri_id = %s OR t.customer_id = %s)
              AND COALESCE(t.tutar, 0) > 0
              AND COALESCE(t.aciklama, '') LIKE %s
            ORDER BY t.id DESC
            """,
            (musteri_id, musteri_id, pat1),
        ) or []
    out = []
    for r in rows:
        if not isinstance(r, dict):
            continue
        try:
            tid = int(r.get("id") or 0)
        except (TypeError, ValueError):
            tid = 0
        if tid <= 0:
            continue
        try:
            tut = round(float(r.get("tutar") or 0), 2)
        except (TypeError, ValueError):
            tut = 0.0
        out.append({"id": tid, "tutar": tut, "makbuz_no": r.get("makbuz_no")})
    return out


def _ekstre_tahsil_ids_by_iso_from_rows(rows) -> dict[str, list[int]]:
    """Tek tahsilat sorgusundan ay → tahsilat id listesi (ekstre satır butonları, N+1 SQL yok)."""
    by_iso: dict[str, list[int]] = {}
    seen_per: dict[str, set[int]] = {}
    for r in rows or []:
        if not isinstance(r, dict):
            continue
        try:
            tid = int(r.get("id") or 0)
        except (TypeError, ValueError):
            continue
        if tid <= 0:
            continue
        ac = str(r.get("aciklama") or "")
        iso_keys: list[str] = []
        for iso_raw in re.findall(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac):
            try:
                dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                iso_keys.append(date(dd.year, dd.month, 1).isoformat())
            except ValueError:
                pass
        iso_e = _tahsil_row_ekstre_eslesme_ay_iso(r)
        if iso_e:
            iso_keys.append(iso_e)
        tt = r.get("tahsilat_tarihi")
        if tt and hasattr(tt, "year"):
            iso_keys.append(date(int(tt.year), int(tt.month), 1).isoformat())
        for iso_k in iso_keys:
            if not iso_k:
                continue
            st = seen_per.setdefault(iso_k, set())
            if tid in st:
                continue
            st.add(tid)
            by_iso.setdefault(iso_k, []).append(tid)
    return by_iso


def _ekstre_tahsil_ids_for_iso(
    musteri_id: int, iso_first: str, rows: list | None = None
) -> list[int]:
    """Ekstre ay satırı: FIFO id yoksa marker / eşleme / tahsilat tarihi ile tahsilat id'leri."""
    iso = str(iso_first or "").strip()[:10]
    if len(iso) < 10:
        return []
    if rows is not None:
        return list((_ekstre_tahsil_ids_by_iso_from_rows(rows)).get(iso) or [])
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return []
    return _ekstre_tahsil_ids_by_iso_from_rows(_ekstre_tahsil_rows_for_musteri(mid)).get(iso) or []


def _aylik_tahsil_marker_aylar_set_normalized(musteri_id: int) -> set[str]:
    """
    Bu müşteriye ait |AYLIK_TAH| marker'lı aylar (YYYY-M normalize).
    Not: Grid cache ufku dışında kalan (ama daha önce tahsil edilmiş) aylar için
    frontend yeşil durumunun kaybolmaması amacıyla ayrıca döndürülür.
    """
    out: set[str] = set()
    rows = fetch_all(
        """
        SELECT COALESCE(aciklama, '') AS aciklama
        FROM tahsilatlar
        WHERE (musteri_id = %s OR customer_id = %s)
          AND COALESCE(tutar, 0) > 0
          AND COALESCE(aciklama, '') LIKE '%%|AYLIK_TAH|%%'
        """,
        (musteri_id, musteri_id),
    ) or []
    for r in rows:
        ac = str((r or {}).get("aciklama") or "")
        for iso in re.findall(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac):
            nk = _firma_ozet_normalize_tahsil_ay_key(iso)
            if nk:
                out.add(nk)
    return out


def _aylik_tahsil_ekstre_eslesme_aylar_set_normalized(musteri_id: int) -> set[str]:
    """Ekstrede tahsilat satırı üreten aylar (YYYY-M normalize)."""
    rows = fetch_all(
        """
        SELECT
            COALESCE(
                f.fatura_tarihi::date,
                NULLIF(substring(COALESCE(t.aciklama, '') from '\\|AYLIK_TAH\\|([0-9]{4}-[0-9]{2}-[0-9]{2})\\|'), '')::date,
                t.tahsilat_tarihi::date
            ) AS eslesme_tarihi
        FROM tahsilatlar t
        LEFT JOIN faturalar f ON f.id = t.fatura_id
        WHERE (t.musteri_id = %s OR t.customer_id = %s)
          AND COALESCE(t.tutar, 0) > 0
        """,
        (musteri_id, musteri_id),
    ) or []
    out: set[str] = set()
    for r in rows:
        d = (r or {}).get("eslesme_tarihi")
        if not d:
            continue
        try:
            dd = d if hasattr(d, "year") else datetime.strptime(str(d)[:10], "%Y-%m-%d").date()
            nk = _firma_ozet_normalize_tahsil_ay_key(f"{dd.year}-{dd.month}")
            if nk:
                out.add(nk)
        except Exception:
            continue
    return out


def _aylik_tahsil_edilen_aylar_set_normalized_batch(musteri_ids: list[int]) -> dict[int, set[str]]:
    """Birden çok müşteri için tam ödenmiş ay seti (grid ile uyumlu; müşteri başına payload üretir)."""
    mids: list[int] = []
    seen = set()
    for x in musteri_ids or []:
        try:
            i = int(x)
        except (TypeError, ValueError):
            continue
        if i <= 0 or i in seen:
            continue
        seen.add(i)
        mids.append(i)
    if not mids:
        return {}
    out: dict[int, set[str]] = {m: set() for m in mids}
    cached = _read_aylik_grid_cache_payload_batch(mids)
    tm = _tufe_map_by_year_month_cached()
    for mid in mids:
        payload = cached.get(mid)
        if payload is None:
            payload = _build_aylik_grid_cache_payload(mid, tufe_map=tm)
        out[mid] = _aylik_tahsil_edilen_aylar_from_payload(payload)
    return out


def firma_ozet_toplam_borc_ve_geciken_ay(
    musteri_id: int,
    ref_y: int,
    ref_m: int,
    tufe_map: dict,
    kyc_for_grid: dict,
    manual_reel_by_year: dict,
    tahsil_n_override: set[str] | None = None,
    skip_disk_cache_for_months: bool = True,
) -> tuple[float, int]:
    """
    Sözleşme başlangıcından referans ayına kadar (dahil), tahsilatı görülmeyen her ay için
    aylık grid ile aynı KDV dahil tutarları toplar; geciken ay = bu koşulda pozitif tutarlı
    ay adedi (toplam borç ile aynı döngü, tutarların birebir toplamı).

    Not: «Yalnız son ödemeden sonraki borç» değil; ara aylarda ödeme olsa bile öncesi ödenmemiş
    aylar da dahil edilir — böylece Ocak+Şubat+Mart+Nisan gibi üst üste kırmızı ayların toplamı
    raporla örtüşür.
    """
    try:
        mid = int(musteri_id)
        ry = int(ref_y)
        rm = int(ref_m)
    except (TypeError, ValueError):
        return 0.0, 0
    if mid <= 0 or rm < 1 or rm > 12:
        return 0.0, 0
    kyc = dict(kyc_for_grid or {})
    bas_soz = _aylik_grid_coerce_date(kyc.get("sozlesme_tarihi"))
    if not bas_soz:
        return 0.0, 0
    bas_first = date(bas_soz.year, bas_soz.month, 1)
    ref_first = date(ry, rm, 1)
    if ref_first < bas_first:
        return 0.0, 0
    tahsil_n = tahsil_n_override if isinstance(tahsil_n_override, set) else _aylik_tahsil_edilen_aylar_set_normalized(mid)
    reel = manual_reel_by_year if isinstance(manual_reel_by_year, dict) else {}
    artis_d = _aylik_grid_coerce_date(kyc.get("kira_artis_tarihi")) or bas_soz
    artis_month = int(artis_d.month)
    artis_day = int(artis_d.day)
    bit_soz = _aylik_grid_coerce_date(kyc.get("sozlesme_bitis")) or bas_soz
    bugun_y = date.today().year
    y_end = max(ref_first.year, bugun_y, bas_soz.year, bit_soz.year)
    try:
        reel_map = _reel_ay_key_tutar_map_musteri(
            mid,
            bas_soz,
            artis_month,
            artis_day,
            tufe_map,
            y_end,
            manual_by_year=reel,
            kyc_for_yilmap=kyc,
        )
    except Exception:
        reel_map = {}
    total = 0.0
    n_ay = 0
    cur = bas_first
    while cur <= ref_first:
        ay_key = f"{cur.year}-{cur.month}"
        if ay_key not in tahsil_n:
            v = 0.0
            if ay_key in reel_map:
                try:
                    rv = float(reel_map.get(ay_key) or 0)
                    if math.isfinite(rv) and rv > 0:
                        v = rv
                except (TypeError, ValueError):
                    v = 0.0
            if not v:
                v = firma_ozet_aylik_grid_hucre_kdv_dahil(
                    mid,
                    cur.year,
                    cur.month,
                    tufe_map,
                    kyc,
                    None,
                    reel,
                    skip_disk_cache=bool(skip_disk_cache_for_months),
                    # reel_map yukarıda bir kez üretildi; ay bazında tekrar üretme.
                    skip_reel_overlay=True,
                )
            if isinstance(v, (int, float)) and math.isfinite(float(v)) and float(v) > 0:
                total += float(v)
                n_ay += 1
        cur = _add_months(cur, 1)
    return round(total, 2), int(n_ay)


def firma_ozet_toplam_borc_kdv_dahil(
    musteri_id: int,
    ref_y: int,
    ref_m: int,
    tufe_map: dict,
    kyc_for_grid: dict,
    manual_reel_by_year: dict,
) -> float:
    t, _ = firma_ozet_toplam_borc_ve_geciken_ay(
        musteri_id,
        ref_y,
        ref_m,
        tufe_map,
        kyc_for_grid,
        manual_reel_by_year,
    )
    return t


# Firma özeti / fatura raporu ile aynı: giriş tarihi (KYC söz. / rent_start / created_at)
_FIRMA_OZET_GIRIS_TARIHI_SQL = """
COALESCE(
    CASE
        WHEN mk.sozlesme_tarihi IS NULL THEN NULL
        WHEN BTRIM(mk.sozlesme_tarihi::text) = '' THEN NULL
        WHEN BTRIM(mk.sozlesme_tarihi::text) ~ '^[0-9]{4}-[0-9]{2}-[0-9]{2}'
            THEN (SUBSTRING(BTRIM(mk.sozlesme_tarihi::text) FROM 1 FOR 10))::date
        WHEN BTRIM(mk.sozlesme_tarihi::text) ~ '^[0-9]{1,2}\\.[0-9]{1,2}\\.[0-9]{4}'
            THEN TO_DATE(
                REGEXP_REPLACE(BTRIM(mk.sozlesme_tarihi::text), ' .*$', ''),
                'DD.MM.YYYY'
            )
        WHEN BTRIM(mk.sozlesme_tarihi::text) ~ '^[0-9]{1,2}-[0-9]{1,2}-[0-9]{4}'
            THEN TO_DATE(
                REGEXP_REPLACE(BTRIM(mk.sozlesme_tarihi::text), ' .*$', ''),
                'DD-MM-YYYY'
            )
        ELSE NULL
    END,
    c.rent_start_date::date,
    c.created_at::date
)
""".strip()


def _firma_ozet_kyc_dict_from_grid_sql_row(row) -> dict | None:
    """musteri_aylik_grid_* SELECT satırından firma_ozet_aylik_grid_hucre_kdv_dahil için kyc dict."""
    if not row or row.get("id") is None:
        return None
    raw_soz_bas = row.get("kyc_soz_bas")
    raw_soz_bit = row.get("kyc_soz_bit")
    giris_sql = row.get("giris_raw")
    bas_parsed = _aylik_grid_coerce_date(raw_soz_bas)
    bit_parsed = _aylik_grid_coerce_date(raw_soz_bit)
    soz_bas_eff = bas_parsed or giris_sql or raw_soz_bas
    soz_bit_eff = bit_parsed if bit_parsed is not None else raw_soz_bit
    return {
        "sozlesme_tarihi": soz_bas_eff,
        "sozlesme_bitis": soz_bit_eff,
        "durum": row.get("durum"),
        "kapanis_tarihi": row.get("kapanis_tarihi"),
        "kapanis_sonrasi_borc_ay": row.get("kapanis_sonrasi_borc_ay"),
        "aylik_kira": row.get("firma_grid_aylik_net"),
        "kira_artis_tarihi": row.get("kyc_kira_artis"),
        "kira_suresi_ay": row.get("kyc_kira_suresi_ay"),
        "kira_nakit": row.get("kira_nakit"),
        "kira_banka": row.get("kira_banka"),
        "kira_nakit_tutar": row.get("kira_nakit_tutar"),
        "kira_banka_tutar": row.get("kira_banka_tutar"),
        "kdv_oran": row.get("kdv_oran"),
    }


def _musteri_aylik_grid_customer_kyc_select_sql():
    """Tek müşteri veya ANY(musteri_ids) için ortak FROM; WHERE dışarıda eklenir."""
    gsql = _FIRMA_OZET_GIRIS_TARIHI_SQL
    return f"""
        SELECT c.id,
               ({gsql}) AS giris_raw,
               c.guncel_kira_bedeli,
               c.ilk_kira_bedeli,
               c.durum,
               c.kapanis_tarihi,
               c.kapanis_sonrasi_borc_ay,
               mk.sozlesme_tarihi AS kyc_soz_bas,
               mk.sozlesme_bitis AS kyc_soz_bit,
               mk.kira_artis_tarihi AS kyc_kira_artis,
               mk.kira_suresi_ay AS kyc_kira_suresi_ay,
               mk.aylik_kira,
               mk.kira_nakit,
               mk.kira_banka,
               mk.kira_nakit_tutar,
               mk.kira_banka_tutar,
               mk.kdv_oran,
               CASE
                   WHEN mk.aylik_kira IS NOT NULL AND mk.aylik_kira > 0 THEN mk.aylik_kira
                   ELSE COALESCE(c.guncel_kira_bedeli, c.ilk_kira_bedeli, mk.aylik_kira)
               END AS firma_grid_aylik_net
        FROM customers c
        LEFT JOIN (
            SELECT DISTINCT ON (musteri_id)
                musteri_id,
                sozlesme_tarihi,
                sozlesme_bitis,
                kira_artis_tarihi,
                kira_suresi_ay,
                aylik_kira,
                kira_nakit,
                kira_banka,
                kira_nakit_tutar,
                kira_banka_tutar,
                kdv_oran
            FROM musteri_kyc
            ORDER BY musteri_id, id DESC
        ) mk ON mk.musteri_id = c.id
    """


def musteri_aylik_grid_hucre_kdv_dahil_takvim_ayi_batch(musteri_ids: list, ref: date | None = None) -> dict[int, float]:
    """
    Birden çok müşteri için tek TÜFE + tek KYC/reel sorgusu; Grup konsolide raporu için.
    Dönüş: {{ musteri_id: tutar, ... }} — istekte olmayan id'ler için anahtar yoktur.
    """
    d = ref or date.today()
    ref_y, ref_m = int(d.year), int(d.month)
    if ref_m < 1 or ref_m > 12:
        return {}

    mids: list[int] = []
    seen = set()
    for x in musteri_ids or []:
        try:
            i = int(x)
        except (TypeError, ValueError):
            continue
        if i <= 0 or i in seen:
            continue
        seen.add(i)
        mids.append(i)
    if not mids:
        return {}

    _ensure_musteri_reel_donem_tutar_table()
    tufe_map = _tufe_map_by_year_month()

    base_sql = _musteri_aylik_grid_customer_kyc_select_sql()
    rows = fetch_all(base_sql + " WHERE c.id = ANY(%s)", (mids,)) or []
    row_by_id: dict[int, dict] = {}
    for r in rows:
        try:
            rid = int(r.get("id") or 0)
        except (TypeError, ValueError):
            continue
        if rid > 0:
            row_by_id[rid] = r

    reel_by_mid: dict[int, dict[int, float]] = {}
    rrows = (
        fetch_all(
            "SELECT musteri_id, donem_yil, tutar_kdv_dahil FROM musteri_reel_donem_tutar WHERE musteri_id = ANY(%s)",
            (mids,),
        )
        or []
    )
    for rr in rrows:
        try:
            mid_r = int(rr.get("musteri_id") or 0)
            yil = int(rr.get("donem_yil") or 0)
            tut = float(rr.get("tutar_kdv_dahil") or 0)
        except (TypeError, ValueError):
            continue
        if mid_r <= 0:
            continue
        reel_by_mid.setdefault(mid_r, {})[yil] = tut

    out: dict[int, float] = {}
    for mid in mids:
        row = row_by_id.get(mid)
        kyc_for_grid = _firma_ozet_kyc_dict_from_grid_sql_row(row) if row else None
        if not kyc_for_grid:
            out[mid] = 0.0
            continue
        reel_manual = dict(reel_by_mid.get(mid) or {})
        try:
            v = float(
                firma_ozet_aylik_grid_hucre_kdv_dahil(
                    mid,
                    ref_y,
                    ref_m,
                    tufe_map,
                    kyc_for_grid,
                    None,
                    reel_manual,
                    skip_disk_cache=True,
                    skip_reel_overlay=False,
                )
            )
            out[mid] = round(v, 2) if math.isfinite(v) else 0.0
        except Exception:
            out[mid] = 0.0
    return out


def _firma_ozet_cache_payload_usable(payload, ref_y: int, ref_m: int) -> bool:
    """Firma özet raporu için disk önbelleği okunabilir mi (tam yeniden hesap gerekmez)."""
    if not isinstance(payload, dict):
        return False
    if ref_m < 1 or ref_m > 12:
        return False
    try:
        if int(payload.get("compute_rev") or 0) != AYLIK_GRID_COMPUTE_REV:
            return False
    except (TypeError, ValueError):
        return False
    aylar = payload.get("aylar") or []
    if not isinstance(aylar, list) or not aylar:
        return False
    if payload.get("tahsilat_imza"):
        return True
    return any(
        isinstance(a, dict) and ("tahsil_edildi" in a or "kalan_tutar_kdv" in a or "odenen_tutar_kdv" in a)
        for a in aylar
    )


def prewarm_aylik_grid_cache_for_musteriler(
    musteri_ids,
    ref: date | None = None,
    *,
    max_rebuild: int | None = None,
) -> int:
    """
    Tekil müşteri raporu öncesi: eksik/stale grid önbelleklerini doldurur.
    Ağır firma_ozet_toplam_borc hesabının müşteri başına tekrarlanmasını önler.
    """
    d = ref or date.today()
    ref_y, ref_m = int(d.year), int(d.month)
    mids: list[int] = []
    seen = set()
    for x in musteri_ids or []:
        try:
            i = int(x)
        except (TypeError, ValueError):
            continue
        if i <= 0 or i in seen:
            continue
        seen.add(i)
        mids.append(i)
    if not mids:
        return 0
    _ensure_aylik_grid_cache_table()
    tufe_map = _tufe_map_by_year_month_cached()
    payloads = _read_aylik_grid_cache_payload_batch(mids, mem_ttl=0.0)
    need: list[int] = []
    for mid in mids:
        if not _firma_ozet_cache_payload_usable(payloads.get(mid), ref_y, ref_m):
            need.append(mid)
    if not need:
        return 0
    try:
        cap = int(str(os.getenv("FIRMA_OZET_CACHE_REBUILD_PER_REQ", "150")).strip() or "150")
    except ValueError:
        cap = 150
    if max_rebuild is not None:
        try:
            cap = min(cap, max(1, int(max_rebuild)))
        except (TypeError, ValueError):
            pass
    need = need[: max(1, cap)]
    workers = 1
    try:
        workers = max(1, min(4, int(str(os.getenv("FIRMA_OZET_CACHE_REBUILD_WORKERS", "3")).strip() or "3")))
    except ValueError:
        workers = 3
    rebuilt = 0

    def _one(mid: int):
        pl = _upsert_aylik_grid_cache(mid, tufe_map=tufe_map)
        return mid, pl

    if workers <= 1 or len(need) < 2:
        for mid in need:
            try:
                _, pl = _one(mid)
                if pl and _firma_ozet_cache_payload_usable(pl, ref_y, ref_m):
                    rebuilt += 1
            except Exception:
                logging.getLogger(__name__).exception("prewarm grid cache mid=%s", mid)
        return rebuilt

    from concurrent.futures import ThreadPoolExecutor, as_completed

    with ThreadPoolExecutor(max_workers=workers) as pool:
        futs = [pool.submit(_one, mid) for mid in need]
        for fut in as_completed(futs):
            try:
                _, pl = fut.result()
                if pl and _firma_ozet_cache_payload_usable(pl, ref_y, ref_m):
                    rebuilt += 1
            except Exception:
                logging.getLogger(__name__).exception("prewarm grid cache worker")
    if rebuilt and hasattr(musteri_firma_ozet_grid_ozet_batch, "_cache"):
        bc = musteri_firma_ozet_grid_ozet_batch._cache
        if isinstance(bc, dict):
            for mid in need:
                bc.pop((int(mid), ref_y, ref_m), None)
    return rebuilt


def _firma_ozet_ozet_from_grid_cache_payload(payload, ref_y: int, ref_m: int) -> dict | None:
    """
    musteri_aylik_grid_cache payload'ından rapor özet alanları (tam grid yeniden hesap yok).
    firma_ozet_toplam_borc_ve_geciken_ay ile uyumlu: referans aya kadar tahsil edilmemiş pozitif aylar.
    """
    if not _firma_ozet_cache_payload_usable(payload, ref_y, ref_m):
        return None
    if ref_m < 1 or ref_m > 12:
        return None
    ref_first = date(int(ref_y), int(ref_m), 1)
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    borc_month = 0.0
    total_borc = 0.0
    geciken = 0
    for a in payload.get("aylar") or []:
        if not isinstance(a, dict):
            continue
        try:
            yy = int(a.get("yil"))
            mm = int(a.get("ay"))
            cur = date(yy, mm, 1)
        except (TypeError, ValueError):
            continue
        try:
            tut = float(a.get("tutar_kdv_dahil") or a.get("brut_tutar_kdv") or 0)
        except (TypeError, ValueError):
            tut = 0.0
        if not math.isfinite(tut):
            tut = 0.0
        if cur.year == ref_y and cur.month == ref_m:
            borc_month = round(tut, 2)
        if cur > ref_first:
            continue
        tahsil = bool(a.get("tahsil_edildi"))
        try:
            kalan = float(a.get("kalan_tutar_kdv") or 0)
        except (TypeError, ValueError):
            kalan = 0.0
        odenen = float(a.get("odenen_tutar_kdv") or 0) if a.get("odenen_tutar_kdv") is not None else 0.0
        kismi = bool(a.get("kismi_tahsilat")) or (odenen > tol and kalan > tol)
        if (not tahsil) or kismi:
            if tut > tol:
                total_borc += tut
                geciken += 1
    return {
        "borc_month": borc_month,
        "toplam_borc": round(total_borc, 2),
        "geciken_ay": int(geciken),
        "sozlesme_gun": 0,
    }


def musteri_firma_ozet_grid_ozet_batch(musteri_ids: list, ref: date | None = None) -> dict[int, dict]:
    """
    Grup raporları için tek geçişte:
    - borc_month: referans ay KDV dahil aylık grid hücresi
    - geciken_ay: tahsilsiz pozitif ay adedi
    - sozlesme_gun: sözleşme başlangıç günü (1-31)
    """
    d = ref or date.today()
    ref_y, ref_m = int(d.year), int(d.month)
    if ref_m < 1 or ref_m > 12:
        return {}

    mids: list[int] = []
    seen = set()
    for x in musteri_ids or []:
        try:
            i = int(x)
        except (TypeError, ValueError):
            continue
        if i <= 0 or i in seen:
            continue
        seen.add(i)
        mids.append(i)
    if not mids:
        return {}
    if not hasattr(musteri_firma_ozet_grid_ozet_batch, "_cache"):
        musteri_firma_ozet_grid_ozet_batch._cache = {}
    cache = musteri_firma_ozet_grid_ozet_batch._cache
    cache_ttl = 120.0
    now_ts = time.time()
    cached_out: dict[int, dict] = {}
    need_mids: list[int] = []
    for mid in mids:
        ckey = (int(mid), ref_y, ref_m)
        cval = cache.get(ckey)
        if cval and (now_ts - float(cval[0])) <= cache_ttl:
            cached_out[mid] = dict(cval[1] or {})
        else:
            need_mids.append(mid)
    if not need_mids:
        return cached_out

    try:
        prewarm_aylik_grid_cache_for_musteriler(need_mids, date(ref_y, ref_m, 1))
    except Exception:
        pass

    _ensure_musteri_reel_donem_tutar_table()
    tufe_map = _tufe_map_by_year_month_cached()
    base_sql = _musteri_aylik_grid_customer_kyc_select_sql()
    rows = fetch_all(base_sql + " WHERE c.id = ANY(%s)", (need_mids,)) or []
    row_by_id: dict[int, dict] = {}
    for r in rows:
        try:
            rid = int(r.get("id") or 0)
        except (TypeError, ValueError):
            continue
        if rid > 0:
            row_by_id[rid] = r

    reel_by_mid: dict[int, dict[int, float]] = {}
    rrows = (
        fetch_all(
            "SELECT musteri_id, donem_yil, tutar_kdv_dahil FROM musteri_reel_donem_tutar WHERE musteri_id = ANY(%s)",
            (need_mids,),
        )
        or []
    )
    for rr in rrows:
        try:
            mid_r = int(rr.get("musteri_id") or 0)
            yil = int(rr.get("donem_yil") or 0)
            tut = float(rr.get("tutar_kdv_dahil") or 0)
        except (TypeError, ValueError):
            continue
        if mid_r <= 0:
            continue
        reel_by_mid.setdefault(mid_r, {})[yil] = tut

    tahsil_by_mid = _aylik_tahsil_edilen_aylar_set_normalized_batch(need_mids)
    cache_payloads = _read_aylik_grid_cache_payload_batch(need_mids)
    out: dict[int, dict] = {}

    def _sozlesme_gun_from_grid_row(row: dict | None) -> int:
        if not row:
            return 0
        raw_soz_bas = row.get("kyc_soz_bas")
        giris_sql = row.get("giris_raw")
        bas_parsed = _aylik_grid_coerce_date(raw_soz_bas)
        soz_bas_eff = bas_parsed or giris_sql or raw_soz_bas
        if isinstance(soz_bas_eff, datetime):
            giris_iso = soz_bas_eff.date().isoformat()[:10]
        elif isinstance(soz_bas_eff, date):
            giris_iso = soz_bas_eff.isoformat()[:10]
        else:
            gr = giris_sql
            giris_iso = str(gr)[:10] if gr is not None and str(gr).strip() else ""
        if not giris_iso:
            return 0
        try:
            _gds = str(giris_iso).strip()[:10]
            if len(_gds) == 10:
                gd = int(date.fromisoformat(_gds).day)
                return gd if 1 <= gd <= 31 else 0
        except (ValueError, TypeError):
            pass
        return 0

    for mid in need_mids:
        row = row_by_id.get(mid)
        kyc_for_grid = _firma_ozet_kyc_dict_from_grid_sql_row(row) if row else None
        gun = _sozlesme_gun_from_grid_row(row)
        cached_pl = cache_payloads.get(mid)
        fast = (
            _firma_ozet_ozet_from_grid_cache_payload(cached_pl, ref_y, ref_m)
            if cached_pl
            else None
        )
        if fast is not None:
            fast["sozlesme_gun"] = gun
            out[mid] = fast
            cache[(int(mid), ref_y, ref_m)] = (now_ts, dict(fast))
            continue
        if not kyc_for_grid:
            out[mid] = {"borc_month": 0.0, "toplam_borc": 0.0, "geciken_ay": 0, "sozlesme_gun": gun}
            cache[(int(mid), ref_y, ref_m)] = (now_ts, dict(out[mid]))
            continue
        reel_manual = dict(reel_by_mid.get(mid) or {})
        try:
            borc_month = float(
                firma_ozet_aylik_grid_hucre_kdv_dahil(
                    mid,
                    ref_y,
                    ref_m,
                    tufe_map,
                    kyc_for_grid,
                    None,
                    reel_manual,
                    skip_disk_cache=False,
                    skip_reel_overlay=False,
                )
            )
            tborc, gec = firma_ozet_toplam_borc_ve_geciken_ay(
                mid,
                ref_y,
                ref_m,
                tufe_map,
                kyc_for_grid,
                reel_manual,
                tahsil_n_override=tahsil_by_mid.get(mid, set()),
                skip_disk_cache_for_months=False,
            )
            out[mid] = {
                "borc_month": round(borc_month, 2) if math.isfinite(borc_month) else 0.0,
                "toplam_borc": round(float(tborc or 0.0), 2),
                "geciken_ay": int(gec or 0),
                "sozlesme_gun": gun,
            }
        except Exception:
            out[mid] = {"borc_month": 0.0, "toplam_borc": 0.0, "geciken_ay": 0, "sozlesme_gun": gun}
        cache[(int(mid), ref_y, ref_m)] = (now_ts, dict(out[mid]))
    out.update(cached_out)
    return out


def musteri_firma_ozet_geciken_ay_ve_sozlesme_gun_batch(musteri_ids: list, ref: date | None = None) -> dict[int, dict]:
    """
    Firma özeti / fatura raporu ile aynı kaynak: geciken_ay (grid üzerinden pozitif tahsilsiz ay adedi)
    ve sözleşme başlangıç günü (1–31). Grup konsolide satırında alt cariler arası üst sınır (max) için kullanılır.
    Dönüş: { musteri_id: {"geciken_ay": int, "sozlesme_gun": int}, ... }
    """
    all_map = musteri_firma_ozet_grid_ozet_batch(musteri_ids, ref)
    return {
        int(mid): {
            "geciken_ay": int((v or {}).get("geciken_ay") or 0),
            "sozlesme_gun": int((v or {}).get("sozlesme_gun") or 0),
        }
        for mid, v in (all_map or {}).items()
    }


def musteri_aylik_grid_hucre_kdv_dahil_takvim_ayi(musteri_id: int, ref: date | None = None) -> float:
    """
    Cari kart Grup «Aylık borç» ile Sözleşmeler aylık gridi aynı kaynak:
    TÜFE + sözleşme çekirdeği + reel dönem (faturaların fatura_tarihi değil).
    ref: içindeki takvim ayı; None ise bugün.
    """
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return 0.0
    if mid <= 0:
        return 0.0
    m = musteri_aylik_grid_hucre_kdv_dahil_takvim_ayi_batch([mid], ref)
    return float(m.get(mid, 0.0))


def _cari_ekstre_ay_borc_tutar(
    y,
    m,
    prev_borc,
    tahsilat_ay_tutar_map,
    reel_ay_map,
    artis_month,
    tufe_map,
    aylik,
    grid_by_iso=None,
    contract_ay_tutar=None,
):
    """Aylık kira borç satırı tutarı: grid (canlı/disk) ile reel dönem haritasının üst sınırı; sonra TÜFE zinciri.

    Eski önbellekte yalnızca kalan (15,40 vb.) kaldığında grid düşük kalabiliyordu; reel elle/TÜFE dönem tutarı
    daha yüksekse onu kullan — «Aylık Tutarlar» + cari ekstre uyumu.
    contract_ay_tutar: _aylik_grid_contract_core ile aynı ayın KDV dahil taban tutarı; grid/reel hatalı «yalnızca
    artış» (yeni dönem − önceki ay ≈ satır tutarı) deseninde tam kira tutarına yükseltmek için kullanılır.
    """
    ilk_gun = date(y, m, 1)
    tarih_iso = ilk_gun.isoformat()
    rk = f"{y}-{m}"
    gv = None
    if grid_by_iso and tarih_iso in grid_by_iso:
        try:
            v = float(grid_by_iso.get(tarih_iso) or 0)
            if math.isfinite(v) and v > 0:
                gv = round(v, 2)
        except (TypeError, ValueError):
            gv = None
    rv = None
    if reel_ay_map and isinstance(reel_ay_map, dict) and rk in reel_ay_map:
        try:
            r0 = float(reel_ay_map.get(rk) or 0)
            if math.isfinite(r0) and r0 > 0:
                rv = round(r0, 2)
        except (TypeError, ValueError):
            rv = None
    out = None
    # Kayıtlı reel dönem tutarı (musteri_reel_donem_tutar) grid/TÜFE hesabının üstünde.
    if rv is not None:
        out = rv
    elif gv is not None:
        out = gv
    if out is None and prev_borc is not None:
        borc_tutar = float(prev_borc)
        if m == artis_month:
            inner = tufe_map.get(y)
            if not isinstance(inner, dict):
                inner = {}
            raw_o = inner.get(m) if m in inner else inner.get(str(m))
            try:
                oran = float(raw_o or 0)
            except (TypeError, ValueError):
                oran = 0.0
            if oran > 0 and math.isfinite(oran) and math.isfinite(prev_borc):
                borc_tutar = round(float(prev_borc) * (1 + oran / 100.0), 2)
        out = round(borc_tutar, 2)
    elif out is None:
        out = round(float(aylik or 0), 2)

    # Grid/reel bazen yeni dönemin tam tutarı yerine (sözleşme TÜFE tabanı − önceki ay) farkını yazıyordu (~562).
    if contract_ay_tutar is not None and prev_borc is not None:
        try:
            cv = float(contract_ay_tutar)
            pb = float(prev_borc)
            ou = float(out)
        except (TypeError, ValueError):
            return out
        if not all(math.isfinite(x) for x in (cv, pb, ou)):
            return out
        # 1) Satır tutarı ≈ (sözleşme TÜFE tabanı − önceki ay) → yalnızca artış (~562) yazılmış.
        if cv > ou + 0.01 and cv > pb + 0.01:
            inc = cv - pb
            tol_inc = max(2.0, abs(cv) * 0.002)
            if inc > 1.0 and abs(ou - inc) <= tol_inc:
                return round(cv, 2)
        # 2) Önceki ay ile aynı plato; sözleşme çekirdeği yeni dönem tutarından anlamlı düşük (1715 vs 2277).
        if cv > ou + 0.01 and cv > pb * 1.08 and pb - 0.03 <= ou <= pb + 0.03:
            return round(cv, 2)
    return out


def _ekstre_borc_tutar_for_month(
    y,
    m,
    prev_borc,
    grid_tutar_by_iso,
    core_ekstre,
    tahsilat_ay_tutar_map,
    reel_ay_map,
    artis_month,
    tufe_map,
    aylik,
):
    """Kayıtlı reel dönem tutarı varsa önce onu; yoksa grid / sözleşme çekirdeği / TÜFE zinciri."""
    rk = f"{int(y)}-{int(m)}"
    if reel_ay_map and isinstance(reel_ay_map, dict) and rk in reel_ay_map:
        try:
            r0 = float(reel_ay_map.get(rk) or 0)
            if math.isfinite(r0) and r0 > 0:
                return round(r0, 2)
        except (TypeError, ValueError):
            pass
    iso = date(int(y), int(m), 1).isoformat()
    if grid_tutar_by_iso:
        try:
            gv = float(grid_tutar_by_iso.get(iso) or 0)
            if math.isfinite(gv) and gv > 0:
                return round(gv, 2)
        except (TypeError, ValueError):
            pass
    cv_m = None
    if core_ekstre:
        try:
            c0 = _aylik_grid_single_month_kdv_from_core(core_ekstre, y, m)
            if c0 and math.isfinite(c0) and c0 > 0:
                cv_m = c0
        except Exception:
            cv_m = None
    return _cari_ekstre_ay_borc_tutar(
        y,
        m,
        prev_borc,
        tahsilat_ay_tutar_map,
        reel_ay_map,
        artis_month,
        tufe_map,
        aylik,
        grid_tutar_by_iso,
        contract_ay_tutar=cv_m,
    )


def _ekstre_devreden_toplamlari(
    full_borc_for_fifo,
    bas_iso_cmp,
    need_after_fifo=None,
    grid_odenen_by_iso=None,
    grid_kalan_by_iso=None,
    grid_tutar_by_iso=None,
    grid_tahsil_edildi_by_iso=None,
    grid_kismi_by_iso=None,
    tahsil_map=None,
    tol=0.01,
):
    """Ekstre başlangıcından önceki toplam borç/tahsilat (aylık grid kalan/ödenen ile uyumlu)."""
    dev_borc_r = 0.0
    dev_alacak_r = 0.0
    if not isinstance(full_borc_for_fifo, dict):
        return 0.0, 0.0
    for iso in sorted(full_borc_for_fifo.keys()):
        if str(iso) >= str(bas_iso_cmp):
            continue
        try:
            borc_m = round(float(full_borc_for_fifo.get(iso) or 0), 2)
        except (TypeError, ValueError):
            borc_m = 0.0
        try:
            gv = round(float((grid_tutar_by_iso or {}).get(iso) or 0), 2)
            if gv > tol:
                borc_m = gv
        except (TypeError, ValueError):
            pass
        if borc_m <= tol:
            continue
        odenen_m = 0.0
        if isinstance(grid_tutar_by_iso, dict) and iso in grid_tutar_by_iso:
            try:
                og_h = _ekstre_hucre_alacak(
                    iso,
                    grid_tutar_by_iso,
                    grid_odenen_by_iso,
                    grid_kalan_by_iso,
                    grid_tahsil_edildi_by_iso=grid_tahsil_edildi_by_iso,
                    grid_kismi_by_iso=grid_kismi_by_iso,
                )
                if og_h is not None:
                    odenen_m = round(float(og_h), 2)
            except (TypeError, ValueError):
                pass
        dev_borc_r += borc_m
        dev_alacak_r += odenen_m
    return round(dev_borc_r, 2), round(dev_alacak_r, 2)


def _ekstre_acilis_bakiyesi_fifo(
    full_borc_for_fifo,
    bas_iso_cmp,
    pays_fifo,
    floor_pay,
    tol=0.01,
):
    """Dönem başlangıcı öncesi net bakiye: FIFO ile yalnızca bas öncesi ödemeler mahsup edilir."""
    if not isinstance(full_borc_for_fifo, dict):
        return 0.0
    month_order_all = sorted(full_borc_for_fifo.keys())
    pre_months = [iso for iso in month_order_all if str(iso) < str(bas_iso_cmp)]
    if not pre_months:
        return 0.0
    need = {
        iso: round(float(full_borc_for_fifo.get(iso) or 0), 2)
        for iso in month_order_all
    }
    marker_pays = []
    general_pays = []
    for pr in pays_fifo or []:
        ac0 = str(pr.get("tahsilat_aciklama") or "")
        if re.search(r"\|AYLIK_TAH\|\d{4}-\d{2}-\d{2}\|", ac0):
            marker_pays.append(pr)
        else:
            general_pays.append(pr)

    def _fifo_pay_alloc_open(pr, restrict_marker_months=False):
        es = str(pr.get("eslesme_tarihi") or "")[:10]
        if not es or es < floor_pay or es >= str(bas_iso_cmp):
            return
        try:
            v = round(float(pr.get("tutar") or 0), 2)
        except (TypeError, ValueError):
            v = 0.0
        if v <= tol:
            return
        if restrict_marker_months:
            ac_m = str(pr.get("tahsilat_aciklama") or "")
            raw_markers = re.findall(
                r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac_m
            )
            target_isos = []
            for iso_raw in raw_markers:
                try:
                    dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                    iso_m = date(dd.year, dd.month, 1).isoformat()
                    if iso_m in need:
                        target_isos.append(iso_m)
                except ValueError:
                    continue
            if not target_isos:
                return
            n_m = len(target_isos)
            cents_total = int(round(v * 100))
            if cents_total <= 0:
                return
            base_c = cents_total // n_m
            rem_c = cents_total % n_m
            for mi, iso_m in enumerate(target_isos):
                share = (base_c + (1 if mi < rem_c else 0)) / 100.0
                if share <= tol:
                    continue
                rem = float(need.get(iso_m) or 0)
                if rem <= tol:
                    continue
                take = round(min(share, rem), 2)
                if take <= tol:
                    continue
                need[iso_m] = round(rem - take, 2)
            return
        for iso in month_order_all:
            if v <= tol:
                break
            rem = float(need.get(iso) or 0)
            if rem <= tol:
                continue
            take = round(min(v, rem), 2)
            if take <= tol:
                continue
            need[iso] = round(rem - take, 2)
            v = round(v - take, 2)

    for pr in marker_pays:
        _fifo_pay_alloc_open(pr, restrict_marker_months=True)
    for pr in general_pays:
        _fifo_pay_alloc_open(pr, restrict_marker_months=False)
    opening = 0.0
    for iso in pre_months:
        opening += max(0.0, round(float(need.get(iso) or 0), 2))
    return round(opening, 2)


def _ekstre_devreden_satir_toplamlari(full_borc_for_fifo, bas_iso_cmp, opening_bakiye, tol=0.01):
    """Devreden satırı: borç/alacak toplamları, bakiye = tam ekstre ile aynı açılış."""
    if not isinstance(full_borc_for_fifo, dict):
        return 0.0, 0.0
    dev_borc_r = 0.0
    for iso in sorted(full_borc_for_fifo.keys()):
        if str(iso) >= str(bas_iso_cmp):
            continue
        try:
            borc_m = round(float(full_borc_for_fifo.get(iso) or 0), 2)
        except (TypeError, ValueError):
            borc_m = 0.0
        if borc_m > tol:
            dev_borc_r += borc_m
    dev_borc_r = round(dev_borc_r, 2)
    opening = round(float(opening_bakiye or 0), 2)
    dev_alacak_r = round(dev_borc_r - opening, 2)
    return dev_borc_r, dev_alacak_r


def _panel_tahsil_by_iso_parse(raw):
    """Tahsilat paneli JSON → {YYYY-MM-01: {aylik, tahsil, kalan}} (ekstre = panel tutarı)."""
    if not raw:
        return {}
    try:
        obj = json.loads(raw) if isinstance(raw, str) else raw
    except (json.JSONDecodeError, TypeError, ValueError):
        return {}
    if not isinstance(obj, dict):
        return {}
    out = {}
    for iso_raw, row in obj.items():
        if not isinstance(row, dict):
            continue
        try:
            parts = str(iso_raw).strip()[:10].split("-")
            yy, mm = int(parts[0]), int(parts[1])
            iso_key = date(yy, mm, 1).isoformat()
        except (ValueError, IndexError, TypeError):
            continue
        try:
            pt = round(float(row.get("tahsil") or row.get("odenen") or 0), 2)
            pk = round(float(row.get("kalan") or 0), 2)
            pb = round(float(row.get("aylik") or row.get("brut") or 0), 2)
        except (TypeError, ValueError):
            continue
        out[iso_key] = {"aylik": pb, "tahsil": pt, "kalan": pk}
    return out


def _ekstre_grid_cache_tahsil_fresh(musteri_id: int, payload: dict | None) -> bool:
    """Grid önbelleği tahsil imzası güncel mi (ekstrede ağır yenileme atlanır)."""
    if not isinstance(payload, dict) or not payload.get("aylar"):
        return False
    try:
        mid = int(musteri_id)
    except (TypeError, ValueError):
        return False
    try:
        return str(payload.get("tahsilat_imza") or "") == str(
            _aylik_tahsil_cache_imza(mid) or ""
        )
    except Exception:
        return False


def _ekstre_merge_panel_into_grid_maps(
    panel_by_iso,
    grid_tutar_by_iso,
    grid_odenen_by_iso,
    grid_kalan_by_iso,
    grid_tahsil_edildi_by_iso,
    grid_kismi_by_iso,
):
    if not panel_by_iso:
        return
    tol = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
    for iso_k, prow in panel_by_iso.items():
        pt = round(float(prow.get("tahsil") or 0), 2)
        pk = round(float(prow.get("kalan") or 0), 2)
        pb = round(float(prow.get("aylik") or 0), 2)
        if pb > tol:
            grid_tutar_by_iso[iso_k] = pb
        grid_odenen_by_iso[iso_k] = pt
        grid_kalan_by_iso[iso_k] = pk
        grid_tahsil_edildi_by_iso[iso_k] = pk <= tol
        grid_kismi_by_iso[iso_k] = pt > tol and pk > tol


def _cari_ekstre_hareketler(
    musteri_id,
    baslangic,
    bitis,
    aylik_kira,
    use_reel_cells=True,
    tahsilat_borca_hizala=False,
    kira_nakit_ekstre=None,
    reel_client_override=None,
    panel_tahsil_by_iso=None,
):
    """Cari ekstre: aylık kira borç + tahsilat alacak + devir bakiyesi.
    use_reel_cells: API uyumluluğu için bırakıldı; reel dönem (elle + TÜFE) ay tutarları her zaman yüklenir.
    tahsilat_borca_hizala=True ve aylık borç haritası varsa: tahsilat satırları kronolojik FIFO ile en eski
    açık aya mahsup edilir (tek kısmi ay mantığı); Kira tutarları grid ile aynı sözleşme brütüdür.
    kira_nakit_ekstre: API'deki nakit modu; None ise kyc.kira_nakit kullanılır.
    """
    try:
        bas = baslangic if isinstance(baslangic, date) else datetime.strptime(str(baslangic)[:10], "%Y-%m-%d").date()
        bit = bitis if isinstance(bitis, date) else datetime.strptime(str(bitis)[:10], "%Y-%m-%d").date()
    except Exception:
        return []
    aylik = float(aylik_kira or 0)
    rows = []

    # Aylık grid ile aynı kaynak: customers + son KYC (ekstre öncesi yalnızca mk çekilmesi grid'i bozuyordu).
    kyc = _musteri_kyc_grup_for_aylik_grid(musteri_id)
    kn_e = kira_nakit_ekstre if kira_nakit_ekstre is not None else bool(kyc.get("kira_nakit"))
    if aylik > 0:
        kyc = dict(kyc)
        # api_cari_ekstre: aylik_kira parametresi nakitte taban, bankada KDV dahil; _aylik_grid_contract_core taban NET bekler.
        if kn_e:
            kyc["aylik_kira"] = round(aylik, 2)
        else:
            ko = float(kyc.get("kdv_oran") or 20)
            if not math.isfinite(ko) or ko < 0:
                ko = 20.0
            mult = 1.0 + ko / 100.0
            kyc["aylik_kira"] = round(aylik / mult, 4) if mult > 0 else round(aylik, 2)
    bas_soz = None
    soz_raw = kyc.get("sozlesme_tarihi")
    if soz_raw:
        try:
            bas_soz = soz_raw if isinstance(soz_raw, date) else datetime.strptime(str(soz_raw)[:10], "%Y-%m-%d").date()
        except Exception:
            bas_soz = None
    if not bas_soz:
        rs = kyc.get("rent_start_date")
        if rs:
            try:
                bas_soz = rs if isinstance(rs, date) else datetime.strptime(str(rs)[:10], "%Y-%m-%d").date()
            except Exception:
                bas_soz = None
    soz_floor_iso = bas_soz.isoformat() if bas_soz else ""
    soz_first_month = date(bas_soz.year, bas_soz.month, 1) if bas_soz else None
    bas_month_iso = date(bas.year, bas.month, 1).isoformat()
    dev_cutoff_iso = bas_month_iso
    if soz_first_month and soz_first_month.isoformat() > dev_cutoff_iso:
        dev_cutoff_iso = soz_first_month.isoformat()

    dev_bas = bas
    if bas_soz and bas_soz < bas:
        dev_bas = bas_soz

    artis_month, artis_day = _kyc_reel_anchor_month_day_for_grid(kyc)
    tufe_map = _tufe_map_by_year_month_cached()

    # Tek tahsilat sorgusu → batch özet (ay başına tekrarlayan SQL yok).
    tahsil_rows_ek = _ekstre_tahsil_rows_for_musteri(int(musteri_id))
    aylik_fatura_by_iso = _aylik_tutar_fatura_id_map(int(musteri_id))
    ekstre_batch_maps = _ekstre_tahsil_batch_maps_from_rows(tahsil_rows_ek)
    ekstre_tahsil_ids_by_iso = _ekstre_tahsil_ids_by_iso_from_rows(tahsil_rows_ek)
    remaining_fifo = _aylik_remaining_brut_by_iso_from_kyc(kyc, tufe_map)
    try:
        ekstre_tahsil_map = _aylik_tahsil_tutar_map(
            int(musteri_id),
            tahsil_rows=tahsil_rows_ek,
            remaining_by_iso=dict(remaining_fifo),
            kyc_row=kyc,
            tufe_map=tufe_map,
        )
    except Exception:
        ekstre_tahsil_map = {}

    # Reel dönem: DB kayıtları + isteğe bağlı reel_json (Uygula önizlemesi) birleşimi.
    manual_reel_pass = _reel_manual_merge_db_and_client(
        int(musteri_id),
        reel_client_override if isinstance(reel_client_override, dict) else None,
    )
    reel_preview = bool(
        isinstance(reel_client_override, dict) and len(reel_client_override) > 0
    )

    # Grid tutarları: önce önbellek (ekstre hızlı yol); TÜFE/reel zinciri gerekirse sonra.
    grid_tutar_by_iso = {}
    grid_odenen_by_iso = {}
    grid_kalan_by_iso = {}
    grid_tahsil_edildi_by_iso = {}
    grid_kismi_by_iso = {}

    def _ekstre_grid_iso_aylar_doldur(aylar_src, merge=False):
        if not isinstance(aylar_src, list):
            return
        for a in aylar_src:
            if not isinstance(a, dict):
                continue
            try:
                yy = int(a.get("yil"))
                mm = int(a.get("ay"))
                brutv = float(a.get("brut_tutar_kdv") or 0)
                tv = float(a.get("tutar_kdv_dahil") or 0)
                od = float(a.get("odenen_tutar_kdv") or 0)
                kl = float(a.get("kalan_tutar_kdv") or 0)
                if not math.isfinite(brutv):
                    brutv = 0.0
                if not math.isfinite(tv):
                    tv = 0.0
                if not math.isfinite(od):
                    od = 0.0
                if not math.isfinite(kl):
                    kl = 0.0
                # Eski cache / hatalı satırlarda tutar_kdv_dahil yalnızca «kalan» kalmış olabiliyordu (70, 15,40…).
                # Aylık griddeki sözleşme brütü: brüt alanı veya tutar ile (ödenen+kalan) üst sınırının tutarlı birleşimi.
                recon = round(max(0.0, od + kl), 2)
                cell = round(max(0.0, brutv, tv, recon), 2)
            except (TypeError, ValueError):
                continue
            if 1900 <= yy <= 2200 and 1 <= mm <= 12 and cell > 0:
                iso_key = f"{yy:04d}-{mm:02d}-01"
                c = round(cell, 2)
                if merge:
                    # Disk cache eski «kalan» tutarları (70, 15,40…) taşıyabiliyor; canlı gridde olmayan
                    # bir ISO için max(0, disk) o ayı yanlışlıkla gridmiş gibi işaretleyip parça borç üretiyordu.
                    # Birleştirme yalnızca canlıda zaten var olan aylar için: brütü max ile güçlendir.
                    if iso_key not in grid_tutar_by_iso:
                        continue
                    prev = float(grid_tutar_by_iso.get(iso_key) or 0)
                    grid_tutar_by_iso[iso_key] = round(max(prev, c), 2)
                else:
                    grid_tutar_by_iso[iso_key] = c
                try:
                    od_cell = round(max(float(od or 0), 0), 2)
                except (TypeError, ValueError):
                    od_cell = 0.0
                if merge and iso_key in grid_odenen_by_iso:
                    grid_odenen_by_iso[iso_key] = round(max(float(grid_odenen_by_iso.get(iso_key) or 0), od_cell), 2)
                else:
                    grid_odenen_by_iso[iso_key] = od_cell
                kl_cell = round(max(0.0, float(kl or 0)), 2)
                if merge and iso_key in grid_kalan_by_iso:
                    grid_kalan_by_iso[iso_key] = round(min(
                        float(grid_kalan_by_iso.get(iso_key) or 0),
                        kl_cell,
                    ), 2)
                else:
                    grid_kalan_by_iso[iso_key] = kl_cell
                grid_tahsil_edildi_by_iso[iso_key] = bool(a.get("tahsil_edildi"))
                grid_kismi_by_iso[iso_key] = bool(a.get("kismi_tahsilat"))

    live_pl = None
    grid_cache_fast = False
    try:
        live_pl = _read_aylik_grid_cache_payload(int(musteri_id))
        if live_pl is None:
            try:
                mem_hit = _aylik_grid_payload_mem.get(int(musteri_id))
                if mem_hit and isinstance(mem_hit[1], dict):
                    live_pl = mem_hit[1]
            except (TypeError, ValueError):
                pass
        if isinstance(live_pl, dict):
            grid_cache_fast = _ekstre_grid_cache_tahsil_fresh(
                int(musteri_id), live_pl
            ) and not reel_preview
            if not grid_cache_fast:
                live_pl = _aylik_grid_cache_payload_tahsil_guncelle(int(musteri_id), live_pl)
                _ekstre_payload_odenen_zenginlestir(
                    live_pl, ekstre_tahsil_map, ekstre_batch_maps
                )
            if manual_reel_pass and bas_soz and not reel_preview:
                try:
                    _aylik_grid_apply_reel_donem_overlay_to_payload(
                        int(musteri_id),
                        kyc,
                        tufe_map,
                        live_pl,
                        manual_reel_by_year=manual_reel_pass,
                    )
                except Exception:
                    logging.getLogger(__name__).exception(
                        "ekstre reel overlay musteri_id=%s", musteri_id
                    )
            _ekstre_grid_iso_aylar_doldur((live_pl or {}).get("aylar"), merge=False)
        else:
            live_pl = _aylik_grid_compute(int(musteri_id), kyc, tufe_map, ekstre_tahsil_map)
            if isinstance(live_pl, dict):
                _ekstre_payload_odenen_zenginlestir(
                    live_pl, ekstre_tahsil_map, ekstre_batch_maps
                )
                if manual_reel_pass and bas_soz:
                    try:
                        _aylik_grid_apply_reel_donem_overlay_to_payload(
                            int(musteri_id),
                            kyc,
                            tufe_map,
                            live_pl,
                            manual_reel_by_year=manual_reel_pass,
                        )
                    except Exception:
                        logging.getLogger(__name__).exception(
                            "ekstre reel overlay musteri_id=%s", musteri_id
                        )
            _ekstre_grid_iso_aylar_doldur((live_pl or {}).get("aylar"), merge=False)
    except Exception:
        logging.getLogger(__name__).exception("cari ekstre aylık grid musteri_id=%s", musteri_id)

    core_ekstre = None
    reel_ay_map = {}
    if not grid_cache_fast:
        try:
            core_ekstre = _aylik_grid_contract_core(kyc, tufe_map)
        except Exception:
            core_ekstre = None
    if bas_soz:
        bugun_y = date.today()
        y_end = max(
            bit.year, bas.year, bugun_y.year, bas_soz.year if bas_soz else bit.year
        )
        try:
            reel_ay_map = _reel_ay_key_tutar_map_musteri(
                musteri_id,
                bas_soz,
                artis_month,
                artis_day,
                tufe_map,
                y_end,
                manual_by_year=manual_reel_pass if manual_reel_pass else None,
                kyc_for_yilmap=kyc,
            )
        except Exception:
            logging.getLogger(__name__).exception(
                "reel_ay_key_tutar_map_musteri musteri_id=%s", musteri_id
            )
            reel_ay_map = {}
    if not grid_tutar_by_iso:
        try:
            _ensure_aylik_grid_cache_table()
            cache_row = fetch_one(
                "SELECT payload FROM musteri_aylik_grid_cache WHERE musteri_id = %s",
                (musteri_id,),
            )
            payload_raw = (cache_row or {}).get("payload")
            if payload_raw:
                payload_disk = (
                    json.loads(payload_raw) if isinstance(payload_raw, str) else payload_raw
                )
                aylar_disk = (
                    payload_disk
                    if isinstance(payload_disk, list)
                    else ((payload_disk or {}).get("aylar") or [])
                )
                _ekstre_grid_iso_aylar_doldur(aylar_disk, merge=False)
        except Exception:
            logging.getLogger(__name__).exception(
                "cari ekstre disk aylık grid musteri_id=%s", musteri_id
            )
            grid_tutar_by_iso = {}
    visible_iso_set = set(grid_tutar_by_iso.keys()) if grid_tutar_by_iso else None
    # Güvenlik: cache ayları beklenen sözleşme aralığından bariz azsa (stale/eksik cache),
    # ekstreyi görünür-ay filtresiyle budama; sözleşme zincirinden borçları üret.
    try:
        if isinstance(visible_iso_set, set) and bas_soz and bit and visible_iso_set:
            start_m = date(max(bas_soz.year, bas.year), max(bas_soz.month if bas_soz.year == bas.year else 1, 1), 1)
            if start_m > bit:
                start_m = date(bit.year, bit.month, 1)
            exp = 0
            yy, mm = start_m.year, start_m.month
            while (yy, mm) <= (bit.year, bit.month):
                exp += 1
                mm += 1
                if mm > 12:
                    mm, yy = 1, yy + 1
            # Örn: cache 12 ay, beklenen 40 ay => filtreyi devre dışı bırak.
            if exp > len(visible_iso_set) + 2:
                visible_iso_set = None
    except Exception:
        pass

    if panel_tahsil_by_iso:
        panel_tahsil_by_iso = _ekstre_panel_filter_db_tahsil(
            int(musteri_id), panel_tahsil_by_iso
        )
        _ekstre_merge_panel_into_grid_maps(
            panel_tahsil_by_iso,
            grid_tutar_by_iso,
            grid_odenen_by_iso,
            grid_kalan_by_iso,
            grid_tahsil_edildi_by_iso,
            grid_kismi_by_iso,
        )

    # Tahsilatlardan ay bazlı "hedef kira" çıkarımı:
    # Mükerrer/çift kayıtları şişirmemek için ay içindeki en yüksek tek tutarı baz al.
    tahsilat_ay_tutar_map = {}
    bas_iso = bas.isoformat()
    tahsilatlar_ham = []
    for tr in tahsil_rows_ek:
        ac = str(tr.get("aciklama") or "")
        mm = re.search(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac)
        iso_m = None
        if mm:
            try:
                dd = datetime.strptime(mm.group(1)[:10], "%Y-%m-%d").date()
                iso_m = date(dd.year, dd.month, 1).isoformat()
            except ValueError:
                iso_m = None
        ft = tr.get("fatura_tarihi")
        tt = tr.get("tahsilat_tarihi")
        eslesme = None
        if ft and hasattr(ft, "year"):
            eslesme = date(int(ft.year), int(ft.month), 1)
        elif iso_m:
            try:
                eslesme = datetime.strptime(iso_m[:10], "%Y-%m-%d").date()
            except ValueError:
                eslesme = None
        elif tt and hasattr(tt, "year"):
            eslesme = date(int(tt.year), int(tt.month), int(tt.day))
        elif tt:
            try:
                eslesme = datetime.strptime(str(tt)[:10], "%Y-%m-%d").date()
            except Exception:
                eslesme = None
        tahsilatlar_ham.append({
            "tutar": tr.get("tutar"),
            "aciklama": ac,
            "eslesme_tarihi": eslesme,
        })
    for tr in tahsilatlar_ham:
        ts = str(tr.get("eslesme_tarihi") or "")[:10]
        if not ts:
            continue
        if soz_floor_iso and ts < soz_floor_iso:
            continue
        try:
            v = round(float(tr.get("tutar") or 0), 2)
        except Exception:
            v = 0
        # Marker'lı kayıtlarda tutar = reel dönem (varsa) veya aylık grid hücresi.
        try:
            ac = str(tr.get("aciklama") or "")
            mm = re.search(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac)
            if mm:
                iso_m = str(mm.group(1) or "")
                hv = None
                try:
                    dd_m = datetime.strptime(iso_m[:10], "%Y-%m-%d").date()
                    rk_m = f"{dd_m.year}-{dd_m.month}"
                    if reel_ay_map and rk_m in reel_ay_map:
                        r_m = float(reel_ay_map.get(rk_m) or 0)
                        if math.isfinite(r_m) and r_m > 0:
                            hv = round(r_m, 2)
                except (TypeError, ValueError):
                    hv = None
                if hv is None:
                    gv = float(grid_tutar_by_iso.get(iso_m) or 0)
                    if gv > 0:
                        hv = round(gv, 2)
                if hv is not None and hv > 0:
                    tol_m = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
                    if v + tol_m >= hv - tol_m:
                        v = hv
        except Exception:
            pass
        if v <= 0:
            continue
        prev = float(tahsilat_ay_tutar_map.get(ts) or 0)
        if v > prev:
            tahsilat_ay_tutar_map[ts] = v

    # 1) Aylık kira satırları + FIFO borç haritası (sözleşme başından bitişe; tarih filtresinden bağımsız tutar).
    kira_block = []
    fifo_bas = dev_bas
    if soz_first_month and soz_first_month > fifo_bas:
        fifo_bas = soz_first_month
    y, m = fifo_bas.year, fifo_bas.month
    bit_y, bit_m = bit.year, bit.month
    borc_by_tarih = {}
    full_borc_for_fifo = {}
    prev_borc_tutar = None
    while (y, m) <= (bit_y, bit_m):
        ilk_gun = date(y, m, 1)
        borc_tutar = _ekstre_borc_tutar_for_month(
            y,
            m,
            prev_borc_tutar,
            grid_tutar_by_iso,
            core_ekstre,
            tahsilat_ay_tutar_map,
            reel_ay_map,
            artis_month,
            tufe_map,
            aylik,
        )
        prev_borc_tutar = borc_tutar
        tarih_iso = ilk_gun.isoformat()
        borc_tutar = _ekstre_hucre_borc(tarih_iso, borc_tutar, grid_tutar_by_iso, reel_ay_map)
        rk_fifo = f"{y}-{m}"
        reel_kilit = bool(
            reel_ay_map and isinstance(reel_ay_map, dict) and rk_fifo in reel_ay_map
        )
        if panel_tahsil_by_iso and tarih_iso in panel_tahsil_by_iso and not reel_kilit:
            try:
                pbr = round(float(panel_tahsil_by_iso[tarih_iso].get("aylik") or 0), 2)
            except (TypeError, ValueError, KeyError):
                pbr = 0.0
            if pbr > float(AYLIK_GRID_TAM_ODENDI_TOLERANS):
                borc_tutar = pbr
        full_borc_for_fifo[tarih_iso] = borc_tutar
        visible = not (isinstance(visible_iso_set, set) and tarih_iso not in visible_iso_set)
        if not visible and bas_month_iso <= tarih_iso <= bit.isoformat():
            bugun_ek = date.today()
            if y == bugun_ek.year and m == bugun_ek.month:
                visible = True
            elif reel_ay_map and f"{y}-{m}" in reel_ay_map:
                visible = True
        if bas_month_iso <= tarih_iso <= bit.isoformat() and visible:
            aciklama = f"{_AY_ADLARI[m - 1]} {y} Kira"
            borc_by_tarih[tarih_iso] = borc_tutar
            kira_row = {
                "tarih": tarih_iso,
                "aciklama": aciklama,
                "belge_no": aciklama,
                "tur": "Kira",
                "borc": borc_tutar,
                "alacak": 0,
                "bakiye": None,
            }
            try:
                fid_k = int(aylik_fatura_by_iso.get(tarih_iso) or 0)
                if fid_k > 0:
                    kira_row["fatura_id"] = fid_k
            except (TypeError, ValueError):
                pass
            kira_block.append(kira_row)
        m += 1
        if m > 12:
            m, y = 1, y + 1

    # 3) Tahsilat (alacak): varsayılan FIFO — ödeme en eski açık aydan başlayarak kapatır;
    #    AYLIK_PAY ile çok aya bölünmüş görünüm yerine tek kısmi ay + dolu aylar.
    _eslesme_sql = """COALESCE(
                    f.fatura_tarihi::date,
                    NULLIF(substring(COALESCE(t.aciklama, '') from '\\|AYLIK_TAH\\|([0-9]{4}-[0-9]{2}-[0-9]{2})\\|'), '')::date,
                    t.tahsilat_tarihi::date
                  )"""
    use_fifo_tahsil = bool(full_borc_for_fifo)
    if use_fifo_tahsil and tahsilat_borca_hizala:
        month_order_all = sorted(full_borc_for_fifo.keys())
        need = {iso: round(float(full_borc_for_fifo.get(iso) or 0), 2) for iso in month_order_all}
        fifo_alloc_win = {iso: 0.0 for iso in month_order_all}
        fifo_harf = {iso: "B" for iso in month_order_all}
        fifo_ids = defaultdict(list)
        floor_pay = soz_floor_iso if soz_floor_iso else "1900-01-01"
        bit_iso_cmp = bit.isoformat()
        pays_fifo = fetch_all(
            f"""SELECT t.id,
                  COALESCE(t.makbuz_no, 'Makbuz-' || t.id) AS belge_no,
                  t.fatura_id,
                  t.tahsilat_tarihi AS tarih,
                  {_eslesme_sql} AS eslesme_tarihi,
                  t.tutar, t.odeme_turu, t.aciklama AS tahsilat_aciklama
           FROM tahsilatlar t
           LEFT JOIN faturalar f ON f.id = t.fatura_id
           WHERE (t.musteri_id = %s OR t.customer_id = %s)
             AND COALESCE(t.tutar, 0) > 0
             AND {_eslesme_sql} >= %s::date
             AND {_eslesme_sql} <= %s::date
           ORDER BY {_eslesme_sql}, t.tahsilat_tarihi, t.id""",
            (musteri_id, musteri_id, floor_pay, bit),
        ) or []
        tol_f = 0.01
        marker_pays = []
        general_pays = []
        for pr in pays_fifo:
            ac0 = str(pr.get("tahsilat_aciklama") or "")
            if re.search(r"\|AYLIK_TAH\|\d{4}-\d{2}-\d{2}\|", ac0):
                marker_pays.append(pr)
            else:
                general_pays.append(pr)

        def _fifo_pay_alloc(pr, restrict_marker_months=False):
            es = str(pr.get("eslesme_tarihi") or "")[:10]
            if not es or es < floor_pay or es > bit_iso_cmp:
                return
            try:
                v = round(float(pr.get("tutar") or 0), 2)
            except (TypeError, ValueError):
                v = 0.0
            if v <= tol_f:
                return
            harf_p = _odeme_turu_harf(pr.get("odeme_turu"))
            try:
                tid_p = int(pr.get("id"))
            except (TypeError, ValueError):
                tid_p = None
            target_isos = None
            if restrict_marker_months:
                ac_m = str(pr.get("tahsilat_aciklama") or "")
                raw_markers = re.findall(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac_m)
                target_isos = []
                for iso_raw in raw_markers:
                    try:
                        dd = datetime.strptime(iso_raw[:10], "%Y-%m-%d").date()
                        iso_m = date(dd.year, dd.month, 1).isoformat()
                        if iso_m in need:
                            target_isos.append(iso_m)
                    except ValueError:
                        continue
                if not target_isos:
                    return
                n_m = len(target_isos)
                cents_total = int(round(v * 100))
                if cents_total <= 0:
                    return
                base_c = cents_total // n_m
                rem_c = cents_total % n_m
                for mi, iso_m in enumerate(target_isos):
                    share = (base_c + (1 if mi < rem_c else 0)) / 100.0
                    if share <= tol_f:
                        continue
                    rem = float(need.get(iso_m) or 0)
                    if rem <= tol_f:
                        continue
                    take = round(min(share, rem), 2)
                    if take <= tol_f:
                        continue
                    need[iso_m] = round(rem - take, 2)
                    fifo_alloc_win[iso_m] = round(float(fifo_alloc_win.get(iso_m) or 0) + take, 2)
                    fifo_harf[iso_m] = harf_p
                    if tid_p:
                        fifo_ids[iso_m].append(tid_p)
                return
            for iso in month_order_all:
                if v <= tol_f:
                    break
                rem = float(need.get(iso) or 0)
                if rem <= tol_f:
                    continue
                take = round(min(v, rem), 2)
                if take <= tol_f:
                    continue
                need[iso] = round(rem - take, 2)
                fifo_alloc_win[iso] = round(float(fifo_alloc_win.get(iso) or 0) + take, 2)
                fifo_harf[iso] = harf_p
                if tid_p:
                    fifo_ids[iso].append(tid_p)
                v = round(v - take, 2)

        for pr in marker_pays:
            _fifo_pay_alloc(pr, restrict_marker_months=True)
        for pr in general_pays:
            _fifo_pay_alloc(pr, restrict_marker_months=False)
        acilis_fifo = _ekstre_acilis_bakiyesi_fifo(
            full_borc_for_fifo,
            dev_cutoff_iso,
            pays_fifo,
            floor_pay,
            tol=tol_f,
        )
        dev_borc_r, dev_alacak_r = _ekstre_devreden_satir_toplamlari(
            full_borc_for_fifo,
            dev_cutoff_iso,
            acilis_fifo,
            tol=tol_f,
        )
        if dev_borc_r > tol_f or dev_alacak_r > tol_f:
            rows.append({
                "tarih": bas.isoformat(),
                "aciklama": "Devreden bakiye (dönem öncesi)",
                "belge_no": "DEVIR",
                "tur": "Devir",
                "borc": dev_borc_r if dev_borc_r > tol_f else 0.0,
                "alacak": dev_alacak_r if dev_alacak_r > tol_f else 0.0,
                "bakiye": None,
            })
        rows.extend(kira_block)
        for iso in sorted(borc_by_tarih.keys()):
            hedef_grid = round(float(borc_by_tarih.get(iso) or full_borc_for_fifo.get(iso) or 0), 2)
            fifo_amt = round(float(fifo_alloc_win.get(iso) or 0), 2)
            fifo_pay = fifo_amt > tol_f
            panel_pt = None
            if panel_tahsil_by_iso and iso in panel_tahsil_by_iso:
                try:
                    panel_pt = round(float(panel_tahsil_by_iso[iso].get("tahsil") or 0), 2)
                except (TypeError, ValueError):
                    panel_pt = 0.0
            if panel_pt is not None:
                if panel_pt <= tol_f:
                    continue
                db_tah_iso = 0.0
                if ekstre_tahsil_map:
                    try:
                        db_tah_iso = round(float(ekstre_tahsil_map.get(iso) or 0), 2)
                    except (TypeError, ValueError):
                        db_tah_iso = 0.0
                if db_tah_iso <= tol_f and not (fifo_ids.get(iso) or []):
                    panel_pt = None
            if panel_pt is not None:
                alacak_iso = panel_pt
            elif fifo_pay and hedef_grid > tol_f:
                if fifo_amt + tol_f >= hedef_grid - tol_f:
                    alacak_iso = hedef_grid
                else:
                    alacak_iso = fifo_amt
            else:
                alacak_iso = _ekstre_hucre_alacak(
                    iso,
                    grid_tutar_by_iso,
                    grid_odenen_by_iso,
                    grid_kalan_by_iso,
                    grid_tahsil_edildi_by_iso=grid_tahsil_edildi_by_iso,
                    grid_kismi_by_iso=grid_kismi_by_iso,
                    ayda_tahsilat_var=fifo_pay,
                    reel_ay_map=reel_ay_map,
                )
            if alacak_iso is None:
                try:
                    og_iso = float(grid_odenen_by_iso.get(iso) or 0)
                except (TypeError, ValueError):
                    og_iso = None
                try:
                    kl_iso = float(grid_kalan_by_iso.get(iso)) if iso in grid_kalan_by_iso else None
                except (TypeError, ValueError):
                    kl_iso = None
                alacak_iso = _ekstre_tahsil_alacak_hucre(
                    int(musteri_id),
                    iso,
                    float(fifo_alloc_win.get(iso) or 0),
                    hedef_grid,
                    tahsil_map=ekstre_tahsil_map,
                    grid_odenen=og_iso if og_iso and og_iso > 0 else None,
                    grid_kalan=kl_iso,
                    batch_maps=ekstre_batch_maps,
                )
            if alacak_iso <= tol_f:
                continue
            try:
                db_tah_row = (
                    round(float(ekstre_tahsil_map.get(iso) or 0), 2)
                    if ekstre_tahsil_map
                    else 0.0
                )
            except (TypeError, ValueError):
                db_tah_row = 0.0
            seen_tf_pre = set()
            uniq_tf_pre = []
            for x in fifo_ids.get(iso) or []:
                if x not in seen_tf_pre:
                    seen_tf_pre.add(x)
                    uniq_tf_pre.append(x)
            if (
                db_tah_row <= tol_f
                and not uniq_tf_pre
                and fifo_amt <= tol_f
            ):
                continue
            try:
                y_iso, m_iso, _ = iso.split("-")
                yi, mi = int(y_iso), int(m_iso)
            except (ValueError, AttributeError):
                continue
            harf_o = fifo_harf.get(iso) or "B"
            aciklama_o = f"{_AY_ADLARI[mi - 1]} {yi} Tahsilat {harf_o}"
            seen_tf = set()
            uniq_tf = []
            for x in fifo_ids.get(iso) or []:
                if x not in seen_tf:
                    seen_tf.add(x)
                    uniq_tf.append(x)
            if not uniq_tf and alacak_iso > tol_f:
                for x in ekstre_tahsil_ids_by_iso.get(iso) or []:
                    if x not in seen_tf:
                        seen_tf.add(x)
                        uniq_tf.append(x)
            belge_tf = ""
            if uniq_tf:
                belge_tf = ",".join(str(x) for x in uniq_tf[:6])
                if len(uniq_tf) > 6:
                    belge_tf += ",…"
            rows.append({
                "tarih": iso,
                "aciklama": aciklama_o,
                "belge_no": belge_tf,
                "tur": "Tahsilat",
                "borc": 0,
                "alacak": alacak_iso,
                "bakiye": None,
                "tahsilat_ids": uniq_tf,
            })
    else:
        # Eski davranış: takvim / marker tarihine göre gruplama (tahsilat_borca_hizala=0 veya borç haritası yok).
        dev_borc_r, dev_alacak_r = _ekstre_devreden_toplamlari(
            full_borc_for_fifo,
            dev_cutoff_iso,
            need_after_fifo=None,
            grid_odenen_by_iso=grid_odenen_by_iso,
            grid_kalan_by_iso=grid_kalan_by_iso,
            grid_tutar_by_iso=grid_tutar_by_iso,
            grid_tahsil_edildi_by_iso=grid_tahsil_edildi_by_iso,
            grid_kismi_by_iso=grid_kismi_by_iso,
            tahsil_map=ekstre_tahsil_map,
            tol=0.01,
        )
        if dev_borc_r > 0.01 or dev_alacak_r > 0.01:
            rows.append({
                "tarih": bas.isoformat(),
                "aciklama": "Devreden bakiye (dönem öncesi)",
                "belge_no": "DEVIR",
                "tur": "Devir",
                "borc": dev_borc_r if dev_borc_r > 0.01 else 0.0,
                "alacak": dev_alacak_r if dev_alacak_r > 0.01 else 0.0,
                "bakiye": None,
            })
        rows.extend(kira_block)
        tahsilatlar = fetch_all(
            f"""SELECT t.id,
                  COALESCE(t.makbuz_no, 'Makbuz-' || t.id) AS belge_no,
                  t.fatura_id,
                  t.tahsilat_tarihi AS tarih,
                  {_eslesme_sql} AS eslesme_tarihi,
                  t.tutar, t.odeme_turu, t.aciklama AS tahsilat_aciklama
           FROM tahsilatlar t
           LEFT JOIN faturalar f ON f.id = t.fatura_id
           WHERE (t.musteri_id = %s OR t.customer_id = %s)
             AND (
                (
                    {_eslesme_sql} >= %s
                    AND {_eslesme_sql} <= %s
                )
                OR
                (
                    t.tahsilat_tarihi::date >= %s
                    AND t.tahsilat_tarihi::date <= %s
                )
             )
           ORDER BY t.tahsilat_tarihi, t.id""",
            (musteri_id, musteri_id, bas, bit, bas, bit),
        )
        tahsilat_ay_map = {}
        for r in (tahsilatlar or []):
            tarih_eslesme = str(r.get("eslesme_tarihi") or "")[:10]
            tarih_giris = str(r.get("tarih") or "")[:10]
            tah_aciklama_raw = (r.get("tahsilat_aciklama") or "").strip()
            has_marker = bool(re.search(r"\|AYLIK_TAH\|\d{4}-\d{2}-\d{2}\|", tah_aciklama_raw))
            has_pay_token = bool(re.search(r"\|AYLIK_PAY\|\d{4}-\d{2}-\d{2}=[0-9]+(?:\.[0-9]+)?\|", tah_aciklama_raw))
            try:
                fatura_id = int(r.get("fatura_id") or 0)
            except (TypeError, ValueError):
                fatura_id = 0
            manuel_ekstre_tarih = (fatura_id <= 0) and (has_pay_token or has_marker)
            tarih = (tarih_giris if manuel_ekstre_tarih else (tarih_eslesme or tarih_giris))
            if not tarih:
                continue
            if soz_floor_iso and tarih < soz_floor_iso:
                continue
            if has_marker and (not manuel_ekstre_tarih) and isinstance(visible_iso_set, set) and tarih not in visible_iso_set:
                continue
            harf = _odeme_turu_harf(r.get("odeme_turu"))
            marker_tarih = _marker_ay_tarihi_from_text(tah_aciklama_raw)
            tah_aciklama = _tahsilat_aciklama_temizle(tah_aciklama_raw)
            if marker_tarih and (not manuel_ekstre_tarih):
                aciklama = f"{_AY_ADLARI[marker_tarih.month - 1]} {marker_tarih.year} Tahsilat {harf}"
            elif tah_aciklama:
                aciklama = tah_aciklama
            else:
                try:
                    td = datetime.strptime(tarih, "%Y-%m-%d").date()
                    aciklama = f"{_AY_ADLARI[td.month - 1]} {td.year} Tahsilat {harf}"
                except Exception:
                    aciklama = "Tahsilat " + harf
            if tarih not in tahsilat_ay_map:
                tahsilat_ay_map[tarih] = {
                    "tarih": tarih,
                    "aciklama": aciklama,
                    "belge_no": r.get("belge_no") or "",
                    "tur": "Tahsilat",
                    "borc": 0,
                    "alacak_toplam": 0.0,
                    "manuel_tutar_koru": False,
                    "bakiye": None,
                    "tahsilat_ids": [],
                }
            if manuel_ekstre_tarih:
                tahsilat_ay_map[tarih]["manuel_tutar_koru"] = True
            try:
                tid = int(r.get("id"))
            except (TypeError, ValueError):
                tid = None
            if tid:
                tahsilat_ay_map[tarih]["tahsilat_ids"].append(tid)
            tahsilat_ay_map[tarih]["alacak_toplam"] += round(float(r.get("tutar") or 0), 2)

        for tarih, item in tahsilat_ay_map.items():
            alacak_toplam = round(float(item.get("alacak_toplam") or 0), 2)
            if tahsilat_borca_hizala:
                hedef_borc = round(float(borc_by_tarih.get(tarih) or 0), 2)
                if hedef_borc > 0:
                    alacak_grid = _ekstre_hucre_alacak(
                        tarih,
                        grid_tutar_by_iso,
                        grid_odenen_by_iso,
                        grid_kalan_by_iso,
                        grid_tahsil_edildi_by_iso=grid_tahsil_edildi_by_iso,
                        grid_kismi_by_iso=grid_kismi_by_iso,
                        ayda_tahsilat_var=alacak_toplam > 0.01,
                        reel_ay_map=reel_ay_map,
                    )
                    if alacak_grid is not None:
                        alacak_toplam = alacak_grid
                    else:
                        try:
                            og_nl = float(grid_odenen_by_iso.get(tarih) or 0)
                        except (TypeError, ValueError):
                            og_nl = None
                        alacak_toplam = _ekstre_tahsil_alacak_hucre(
                            int(musteri_id),
                            tarih,
                            alacak_toplam,
                            hedef_borc,
                            tahsil_map=ekstre_tahsil_map,
                            grid_odenen=og_nl if og_nl and og_nl > 0 else None,
                            batch_maps=ekstre_batch_maps,
                        )
            item["alacak"] = alacak_toplam
            item.pop("alacak_toplam", None)
            item.pop("manuel_tutar_koru", None)
            seen_t = set()
            uniq_ids = []
            for x in item.get("tahsilat_ids") or []:
                if x not in seen_t:
                    seen_t.add(x)
                    uniq_ids.append(x)
            item["tahsilat_ids"] = uniq_ids
            rows.append(item)

    _tur_sira = {"Devir": 0, "Kira": 1, "Tahsilat": 2}

    def _ekstre_satir_sira(x):
        return _tur_sira.get(x.get("tur") or "", 4)

    rows.sort(key=lambda x: (x["tarih"], _ekstre_satir_sira(x)))
    bakiye = 0
    for r in rows:
        bakiye = bakiye + r["borc"] - r["alacak"]
        r["bakiye"] = round(bakiye, 2)
    return rows


def _risk_skoru_hesapla(musteri_id, gecikmis_gun, gecikmis_tutar):
    """Gecikme ve tutara göre 1-100 risk skoru. 50 altı kritik."""
    if not gecikmis_gun and (not gecikmis_tutar or gecikmis_tutar <= 0):
        return 85
    if gecikmis_gun and gecikmis_gun > 60:
        return max(1, 40 - (gecikmis_gun // 30) * 5)
    if gecikmis_gun and gecikmis_gun > 30:
        return 55
    return 70


def _row_to_plain_dict(row):
    """psycopg satırını JSON uyumlu sözlüğe çevir (date/datetime → str, Decimal → float)."""
    if row is None:
        return None
    if not isinstance(row, dict):
        try:
            row = dict(row)
        except Exception:
            return {}
    out = {}
    for k, v in row.items():
        if v is None:
            out[k] = None
        elif isinstance(v, Decimal):
            try:
                out[k] = float(v)
            except Exception:
                out[k] = str(v)
        elif hasattr(v, "isoformat") and callable(getattr(v, "isoformat")):
            try:
                out[k] = v.isoformat() if hasattr(v, "hour") else str(v)[:10]
            except Exception:
                out[k] = str(v) if v is not None else None
        else:
            out[k] = v
    return out


@bp.route('/api/cari-kart/<int:mid>')
@giris_gerekli
def api_cari_kart(mid):
    """Cari kart verisi: özet (bakiye, gecikmiş, bu ay tahsilat, risk, aging), hareketler, finansal profil."""
    try:
        ensure_contracts_engine()
        ensure_customer_financial_profile()
        ensure_cari_kart_perf_indexes()
    except Exception:
        pass
    try:
        force = str(request.args.get("force") or "").lower() in ("1", "true", "yes", "on")
        if not force:
            try:
                now = time.time()
                ttl = 45.0
                cache = getattr(api_cari_kart, "_cache", None)
                if cache is None:
                    cache = {}
                    setattr(api_cari_kart, "_cache", cache)
                hit = cache.get(int(mid))
                if hit and (now - float(hit.get("ts") or 0)) <= ttl and hit.get("payload"):
                    return jsonify(hit["payload"])
            except Exception:
                pass

        resp = _api_cari_kart_impl(mid)
        if not force:
            try:
                payload = resp.get_json(silent=True) if hasattr(resp, "get_json") else None
                if isinstance(payload, dict) and payload.get("ok"):
                    cache = getattr(api_cari_kart, "_cache", None)
                    if isinstance(cache, dict):
                        cache[int(mid)] = {"ts": time.time(), "payload": payload}
            except Exception:
                pass
        return resp
    except Exception as e:
        logging.getLogger(__name__).exception("api_cari_kart mid=%s", mid)
        return jsonify({"ok": False, "mesaj": f"Cari kart verisi alınamadı: {e}"}), 500


def _api_cari_kart_impl(mid):
    cust = fetch_one("SELECT * FROM customers WHERE id = %s", (mid,))
    if not cust:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    bugun = date.today()
    nt_expr = sql_expr_fatura_not_gib_taslak("notlar")
    ozet_sql = fetch_one(
        f"""
        SELECT
          COALESCE(SUM(COALESCE(toplam, tutar, 0)), 0) AS toplam_borc,
          COALESCE(MAX(CASE WHEN vade_tarihi::date < %s THEN (%s::date - vade_tarihi::date) ELSE 0 END), 0) AS gecikmis_gun,
          COALESCE(SUM(CASE WHEN vade_tarihi::date >= (%s::date - INTERVAL '30 day') AND vade_tarihi::date <= %s::date THEN COALESCE(toplam, tutar, 0) ELSE 0 END), 0) AS aging_0_30,
          COALESCE(SUM(CASE WHEN vade_tarihi::date >= (%s::date - INTERVAL '60 day') AND vade_tarihi::date < (%s::date - INTERVAL '30 day') THEN COALESCE(toplam, tutar, 0) ELSE 0 END), 0) AS aging_31_60,
          COALESCE(SUM(CASE WHEN vade_tarihi::date >= (%s::date - INTERVAL '90 day') AND vade_tarihi::date < (%s::date - INTERVAL '60 day') THEN COALESCE(toplam, tutar, 0) ELSE 0 END), 0) AS aging_61_90,
          COALESCE(SUM(CASE WHEN vade_tarihi::date < (%s::date - INTERVAL '90 day') THEN COALESCE(toplam, tutar, 0) ELSE 0 END), 0) AS aging_91
        FROM faturalar
        WHERE musteri_id = %s
          AND COALESCE(durum, '') != 'odendi'
          AND {nt_expr}
        """,
        (bugun, bugun, bugun, bugun, bugun, bugun, bugun, bugun, bugun, mid),
    ) or {}
    toplam_borc = float(ozet_sql.get("toplam_borc") or 0)
    gecikmis_gun = int(ozet_sql.get("gecikmis_gun") or 0)
    bu_ay_bas = bugun.replace(day=1)
    bu_ay_tahsilat = fetch_one(
        """SELECT COALESCE(SUM(tutar), 0) AS t FROM tahsilatlar
           WHERE (musteri_id = %s OR customer_id = %s)
             AND tahsilat_tarihi::date >= %s
             AND tahsilat_tarihi::date < %s""",
        (mid, mid, bu_ay_bas, bu_ay_bas + timedelta(days=32))
    )
    bu_ay_tahsilat = float(bu_ay_tahsilat.get("t", 0) or 0) if bu_ay_tahsilat else 0
    aging_0_30 = float(ozet_sql.get("aging_0_30") or 0)
    aging_31_60 = float(ozet_sql.get("aging_31_60") or 0)
    aging_61_90 = float(ozet_sql.get("aging_61_90") or 0)
    aging_91 = float(ozet_sql.get("aging_91") or 0)
    risk_skoru = _risk_skoru_hesapla(mid, gecikmis_gun, toplam_borc)
    hareketler = _cari_hareketler(mid, max_rows=250)
    # Sözleşme / taksit özetleri
    contracts = fetch_all(
        """SELECT id, cari_kodu, sozlesme_no, baslangic_tarihi, bitis_tarihi, sure_ay, aylik_kira,
                  toplam_tutar, para_birimi, odeme_gunu, depozito, gecikme_faizi_orani,
                  yillik_artis_orani, muacceliyet_var, durum
           FROM contracts
           WHERE musteri_id = %s
           ORDER BY id DESC""",
        (mid,),
    )
    plan_rows = fetch_all(
        """SELECT id, contract_id, musteri_id, taksit_no, vade_tarihi, tutar,
                  odeme_durumu, odenen_tutar, kalan_tutar
           FROM contract_installments
           WHERE musteri_id = %s
           ORDER BY vade_tarihi, taksit_no""",
        (mid,),
    )
    soz_ozet = fetch_one(
        """
        SELECT
          COALESCE(SUM(tutar),0)                         AS planlanan,
          COALESCE(SUM(CASE WHEN vade_tarihi <= %s THEN tutar END),0) AS tahakkuk,
          COALESCE(SUM(odenen_tutar),0)                  AS odenen,
          COALESCE(SUM(CASE WHEN odeme_durumu IN ('gecikmis','icrada') THEN kalan_tutar ELSE 0 END),0) AS geciken,
          COALESCE(SUM(CASE WHEN vade_tarihi > %s THEN kalan_tutar ELSE 0 END),0) AS gelecek
        FROM contract_installments
        WHERE musteri_id = %s
        """,
        (bugun, bugun, mid),
    ) or {}
    profil = fetch_one("SELECT * FROM customer_financial_profile WHERE musteri_id = %s", (mid,))
    is_admin = getattr(current_user, "role", None) == "admin"
    payload = {
        "ok": True,
        "musteri": {
            "id": cust.get("id"), "name": cust.get("name"), "tax_number": cust.get("tax_number"),
            "phone": cust.get("phone"), "email": cust.get("email"), "address": cust.get("address"),
            "vergi_dairesi": cust.get("vergi_dairesi"), "mersis_no": cust.get("mersis_no"),
            "nace_kodu": cust.get("nace_kodu"), "ofis_tipi": cust.get("ofis_tipi"),
        },
        "ozet": {
            "guncel_bakiye": round(toplam_borc, 2),
            "gecikmis_tutar": round(toplam_borc, 2),
            "gecikmis_gun": gecikmis_gun,
            "bu_ayki_tahsilat": round(bu_ay_tahsilat, 2),
            "risk_skoru": risk_skoru,
            "aging_0_30": round(aging_0_30, 2),
            "aging_31_60": round(aging_31_60, 2),
            "aging_61_90": round(aging_61_90, 2),
            "aging_91_plus": round(aging_91, 2),
        },
        "hareketler": hareketler,
        "contracts": contracts,
        "installments": plan_rows,
        "contracts_ozet": {
            "planlanan": float(soz_ozet.get("planlanan") or 0),
            "tahakkuk": float(soz_ozet.get("tahakkuk") or 0),
            "odenen": float(soz_ozet.get("odenen") or 0),
            "geciken": float(soz_ozet.get("geciken") or 0),
            "gelecek": float(soz_ozet.get("gelecek") or 0),
        },
        "finansal_profil": None,
    }
    if profil:
        payload["finansal_profil"] = {
            "tahmini_odeme_gunu": profil.get("tahmini_odeme_gunu"),
            "yillik_karlilik_endeksi": float(profil.get("yillik_karlilik_endeksi") or 0),
            "hukuki_esk_puan": profil.get("hukuki_esk_puan"),
            "mutabakat_tarihi": str(profil.get("mutabakat_tarihi"))[:10] if profil.get("mutabakat_tarihi") else None,
            "vade_gunu": profil.get("vade_gunu"),
        }
        if is_admin:
            payload["finansal_profil"]["ic_not"] = profil.get("ic_not")
            payload["finansal_profil"]["hukuki_surec"] = profil.get("hukuki_surec")
    # Sözleşme / taksit satırlarında kalan tip uyuşmazlıklarını önle
    payload["contracts"] = [_row_to_plain_dict(r) for r in (payload.get("contracts") or [])]
    payload["installments"] = [_row_to_plain_dict(r) for r in (payload.get("installments") or [])]
    try:
        return jsonify(payload)
    except TypeError:
        logging.getLogger(__name__).exception("api_cari_kart jsonify TypeError mid=%s", mid)
        return jsonify(
            {
                "ok": False,
                "mesaj": "Cari kart verisi JSON’a çevrilemedi (beklenmeyen alan tipi). Yöneticiye bildirin.",
            }
        ), 500


def _cari_ekstre_varsayilan_son_tam_ay():
    """Parametresiz isteklerde: içinde bulunulan takvim ayının 1. ve son günü (cari ay kira borcu dahil)."""
    bugun = date.today()
    y, m = bugun.year, bugun.month
    _, son = calendar.monthrange(y, m)
    return date(y, m, 1), date(y, m, son)


_cari_ekstre_api_cache: dict = {}


@bp.route('/api/cari-ekstre')
@giris_gerekli
def api_cari_ekstre():
    """
    Sözleşme sayfası cari ekstre: Tarih aralığında aylık kira borç + tahsilat alacak.
    Query: musteri_id, baslangic, bitis (YYYY-MM-DD; boşsa sözleşme bitişi, yoksa cari ay), aylik_kira.
    Opsiyonel reel_json: URL-encoded JSON {\"2024\":612,\"2023\":350} — Uygula önizlemesi (Kaydet öncesi) ekstre ile aynı tutarlar.
    """
    musteri_id = request.args.get("musteri_id", type=int)
    if not musteri_id:
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    cust = fetch_one("SELECT id, name FROM customers WHERE id = %s", (musteri_id,))
    if not cust:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    def_b, def_bit_cari_ay = _cari_ekstre_varsayilan_son_tam_ay()
    kyc = _musteri_kyc_grup_for_aylik_grid(int(musteri_id))
    soz_bit = _aylik_grid_coerce_date((kyc or {}).get("sozlesme_bitis"))
    if soz_bit:
        soz_bit = _aylik_grid_effective_bitis(kyc, soz_bit) or soz_bit
    def_bit = soz_bit or def_bit_cari_ay
    baslangic = request.args.get("baslangic")
    bitis = request.args.get("bitis")
    try:
        bas = datetime.strptime(baslangic[:10], "%Y-%m-%d").date() if baslangic else def_b
        bit = datetime.strptime(bitis[:10], "%Y-%m-%d").date() if bitis else def_bit
    except Exception:
        bas, bit = def_b, def_bit
    if bas > bit:
        bas, bit = bit, bas
    # (genislet kaldırıldı - artık sözleşme bitiş tarihi varsayılan bitiş olarak kullanılıyor)
    aylik_kira = request.args.get("aylik_kira", type=float) or 0
    kdv_oran = request.args.get("kdv_oran", type=float) or 20
    kira_nakit_q = request.args.get("kira_nakit", type=str, default="") or ""
    kira_nakit_ekstre = str(kira_nakit_q).lower() in ("1", "true", "on", "yes")
    # form_bazli_kira=1: sözleşme formundaki aylık kira KYC'ye yazılır; reel dönem DB+TÜFE haritası yine de borçta kullanılır.
    _fb = str(request.args.get("form_bazli_kira") or "").strip().lower()
    use_reel_cells = _fb not in ("1", "true", "yes", "on", "evet")
    _tbh = str(request.args.get("tahsilat_borca_hizala") or "1").strip().lower()
    tahsilat_borca_hizala = _tbh in ("1", "true", "yes", "on", "evet")
    # Borçlar: normalde net + KDV; nakit kiracıda forma girilen tutar doğrudan aylık borç
    if kira_nakit_ekstre and aylik_kira:
        aylik_kira_kdv_dahil = round(aylik_kira, 2)
    else:
        aylik_kira_kdv_dahil = round(aylik_kira * (1 + kdv_oran / 100), 2) if aylik_kira else 0
    db_esas = str(request.args.get("db_esas") or "1").strip().lower() in (
        "1", "true", "yes", "on", "evet",
    )
    reel_client = None
    if not db_esas:
        rj_raw = (request.args.get("reel_json") or "").strip()
        if rj_raw and len(rj_raw) < 8000:
            try:
                rj_obj = json.loads(rj_raw)
                if isinstance(rj_obj, dict):
                    reel_client = rj_obj
            except json.JSONDecodeError:
                reel_client = None
    if db_esas:
        reel_merged_key = _musteri_reel_donem_manual_dict_from_db(int(musteri_id))
    else:
        reel_merged_key = _reel_manual_merge_db_and_client(
            int(musteri_id),
            reel_client if isinstance(reel_client, dict) else None,
        )
    reel_key = ""
    if reel_merged_key:
        try:
            reel_key = json.dumps(
                {str(k): reel_merged_key[k] for k in sorted(reel_merged_key)},
                sort_keys=True,
                separators=(",", ":"),
            )
        except Exception:
            reel_key = str(reel_merged_key)
    panel_by_iso = {}
    if not db_esas:
        ptj_raw = (request.args.get("panel_tahsil_json") or "").strip()
        if ptj_raw and len(ptj_raw) < 32000:
            panel_by_iso = _panel_tahsil_by_iso_parse(ptj_raw)
    if not panel_by_iso:
        panel_by_iso = _load_musteri_panel_by_iso(int(musteri_id))
    if panel_by_iso:
        panel_by_iso = _ekstre_panel_filter_db_tahsil(int(musteri_id), panel_by_iso)
    cache_key = (
        int(musteri_id),
        bas.isoformat(),
        bit.isoformat(),
        round(float(aylik_kira_kdv_dahil or 0), 2),
        bool(tahsilat_borca_hizala),
        bool(kira_nakit_ekstre),
        reel_key,
        bool(db_esas),
        "reel_hucre_v6_db",
    )
    now = time.time()
    hit = None
    if not panel_by_iso:
        hit = _cari_ekstre_api_cache.get(cache_key)
        if hit and (now - float(hit.get("ts") or 0)) < 300.0 and hit.get("payload"):
            return jsonify(hit["payload"])
    try:
        hareketler = _cari_ekstre_hareketler(
            musteri_id,
            bas,
            bit,
            aylik_kira_kdv_dahil,
            use_reel_cells=use_reel_cells,
            tahsilat_borca_hizala=tahsilat_borca_hizala,
            kira_nakit_ekstre=kira_nakit_ekstre,
            reel_client_override=reel_client,
            panel_tahsil_by_iso=panel_by_iso or None,
        )
    except Exception as e:
        logging.getLogger(__name__).exception("api_cari_ekstre musteri_id=%s", musteri_id)
        return jsonify({"ok": False, "mesaj": f"Ekstre hesaplanamadı: {e}"}), 500
    toplam_borc = sum(h.get("borc") or 0 for h in hareketler)
    toplam_alacak = sum(h.get("alacak") or 0 for h in hareketler)
    if hareketler:
        try:
            bakiye = round(float(hareketler[-1].get("bakiye") or 0), 2)
        except (TypeError, ValueError):
            bakiye = round(toplam_borc - toplam_alacak, 2)
    else:
        bakiye = 0.0
    payload = {
        "ok": True,
        "musteri_adi": cust.get("name") or "",
        "hareketler": hareketler,
        "toplam_borc": round(toplam_borc, 2),
        "toplam_alacak": round(toplam_alacak, 2),
        "bakiye": bakiye,
    }
    if not panel_by_iso:
        _cari_ekstre_api_cache[cache_key] = {"ts": now, "payload": payload}
    if len(_cari_ekstre_api_cache) > 120:
        stale = [k for k, v in _cari_ekstre_api_cache.items() if (now - float(v.get("ts") or 0)) > 60.0]
        for k in stale:
            _cari_ekstre_api_cache.pop(k, None)
    return jsonify(payload)


def _next_fatura_no_aylik(prefix=None):
    """Yıla göre artan fatura no; finans ile aynı seri (GIB/INV .env)."""
    try:
        from routes.faturalar_routes import _next_fatura_no

        return _next_fatura_no(prefix)
    except Exception:
        pass
    yil = datetime.now().year
    prefix = (prefix or "INV").strip().upper() or "INV"
    like = f"{prefix}{yil}%"
    row = fetch_one("SELECT fatura_no FROM faturalar WHERE fatura_no LIKE %s ORDER BY id DESC LIMIT 1", (like,))
    if not row or not row.get("fatura_no"):
        return f"{prefix}{yil}000001"
    no = str(row["fatura_no"])
    try:
        tail = int(no[-6:])
        return f"{prefix}{yil}{tail + 1:06d}"
    except Exception:
        return f"{prefix}{yil}000001"


def _next_makbuz_no_aylik():
    """Bir sonraki makbuz numarası: 1000, 1001, 1002 ..."""
    row = fetch_one(
        "SELECT makbuz_no FROM tahsilatlar WHERE makbuz_no IS NOT NULL AND makbuz_no <> '' ORDER BY id DESC LIMIT 1",
    )
    if not row or not row.get("makbuz_no"):
        return "1000"
    s = str(row.get("makbuz_no") or "").strip()
    m = re.search(r"(\d+)$", s)
    if not m:
        return "1000"
    try:
        seq = int(m.group(1)) + 1
    except Exception:
        seq = 1000
    if seq < 1000:
        seq = 1000
    return str(seq)


@bp.route('/api/aylik-tutarlardan-borclandir', methods=['POST'])
@giris_gerekli
def api_aylik_tutarlardan_borclandir():
    """
    Aylık Tutarlar gridinden seçilen aylar için ayrı fatura (borç) kaydı oluşturur.
    Tutarlar KDV dahil kabul edilir; net/kdv satırları faturalar.tutar / kdv_tutar olarak bölünür.
    Cari Ekstre B ve genel cari kartta faturalar görünür.
    """
    ensure_faturalar_amount_columns()
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    musteri_id = data.get("musteri_id")
    try:
        musteri_id = int(musteri_id)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    cust = fetch_one("SELECT id, name FROM customers WHERE id = %s", (musteri_id,))
    if not cust:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    musteri_adi = (cust.get("name") or "").strip() or "—"
    satirlar = data.get("satirlar")
    if not isinstance(satirlar, list) or not satirlar:
        return jsonify({"ok": False, "mesaj": "En az bir ay satırı gerekli."}), 400

    KDV_ORAN = 20.0
    raw_nakit = data.get("kira_nakit")
    if raw_nakit is None:
        kn_row = fetch_one(
            "SELECT kira_nakit FROM musteri_kyc WHERE musteri_id = %s ORDER BY id DESC LIMIT 1",
            (musteri_id,),
        )
        kira_nakit_borc = bool(kn_row and kn_row.get("kira_nakit"))
    else:
        kira_nakit_borc = raw_nakit in (True, 1, "1", "true", "on", "yes")

    olusturulan = []
    atlanan = []
    tahsil_silinen = []
    tahsil_silinen_aylar = set()

    for raw in satirlar:
        if not isinstance(raw, dict):
            continue
        try:
            yil = int(raw.get("yil"))
            ay = int(raw.get("ay"))
        except (TypeError, ValueError):
            atlanan.append({"neden": "geçersiz_yil_ay", "satir": raw})
            continue
        if ay < 1 or ay > 12 or yil < 1990 or yil > 2100:
            atlanan.append({"neden": "tarih_aralik", "satir": raw})
            continue
        try:
            tutar = float(raw.get("tutar_kdv_dahil"))
        except (TypeError, ValueError):
            atlanan.append({"neden": "geçersiz_tutar", "satir": raw})
            continue
        if tutar <= 0:
            atlanan.append({"neden": "tutar_sifir", "satir": raw})
            continue

        toplam = round(tutar, 2)

        ay_bir = date(yil, ay, 1)
        ay_anahtar = ay_bir.isoformat()
        marker = f"|AYLIK_TUTAR|{ay_anahtar}|"
        tah_marker_pat = f"%|AYLIK_TAH|{ay_anahtar}|%"
        pay_marker_pat = f"%|AYLIK_PAY|{ay_anahtar}=%"
        # Borçlandırılan (veya zaten borçlu kabul edilen) ayda aylık tahsilat kaydı kalmasın:
        # ekstrede aynı ayda hem borç hem tahsil görünmesini önler.
        silinen_rows = fetch_all(
            """
            DELETE FROM tahsilatlar
            WHERE (musteri_id = %s OR customer_id = %s)
              AND (
                COALESCE(aciklama, '') LIKE %s
                OR COALESCE(aciklama, '') LIKE %s
              )
            RETURNING id, makbuz_no
            """,
            (musteri_id, musteri_id, tah_marker_pat, pay_marker_pat),
        ) or []
        if silinen_rows:
            tahsil_silinen_aylar.add((yil, ay))
            for tr in silinen_rows:
                tahsil_silinen.append(
                    {
                        "id": tr.get("id"),
                        "makbuz_no": tr.get("makbuz_no"),
                        "yil": yil,
                        "ay": ay,
                    }
                )
        var = fetch_one(
            """
            SELECT
                id,
                fatura_no,
                ettn,
                COALESCE(durum, '') AS durum,
                EXISTS(
                    SELECT 1
                    FROM tahsilatlar t
                    WHERE (t.musteri_id = %s OR t.customer_id = %s)
                      AND COALESCE(t.aciklama, '') LIKE %s
                ) AS tahsil_marker_var
            FROM faturalar
            WHERE musteri_id = %s
              AND COALESCE(notlar, '') LIKE %s
            ORDER BY id DESC
            LIMIT 1
            """,
            (musteri_id, musteri_id, tah_marker_pat, musteri_id, f"%{marker}%"),
        )
        if var:
            # Fatura zaten var: yeni fatura açma; tahsil temizlendi → panel/grid borçlu senkronu.
            olusturulan.append(
                {
                    "id": var.get("id"),
                    "fatura_no": var.get("fatura_no"),
                    "yil": yil,
                    "ay": ay,
                    "toplam": toplam,
                    "guncellendi": True,
                    "zaten_vardi": True,
                }
            )
            continue

        # Ay sonu vade (yaklaşık)
        if ay == 12:
            vade = date(yil, 12, 31)
        else:
            vade = date(yil, ay + 1, 1) - timedelta(days=1)

        if kira_nakit_borc:
            net = toplam
            kdv_tutar = 0.0
        else:
            net = round(toplam / (1 + KDV_ORAN / 100.0), 2)
            kdv_tutar = round(toplam - net, 2)
            net = round(toplam - kdv_tutar, 2)

        ay_adi = _AY_ADLARI[ay - 1]
        if kira_nakit_borc:
            notlar = f"{ay_adi} {yil} kira bedeli (nakit / KDV yok, Aylık Tutarlar){marker}"
        else:
            notlar = f"{ay_adi} {yil} kira bedeli (KDV dahil, Aylık Tutarlar){marker}"

        # Marker bulunmadıysa yeni oluştur.
        fatura_no = _next_fatura_no_aylik()
        execute(
            """
            INSERT INTO faturalar (
                fatura_no, musteri_id, musteri_adi, tutar, kdv_tutar,
                toplam, durum, fatura_tarihi, vade_tarihi, notlar
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (
                fatura_no,
                musteri_id,
                musteri_adi,
                net,
                kdv_tutar,
                toplam,
                "odenmedi",
                ay_bir,
                vade,
                notlar,
            ),
        )
        row = fetch_one("SELECT id FROM faturalar WHERE fatura_no = %s ORDER BY id DESC LIMIT 1", (fatura_no,))
        fid = row.get("id") if row else None
        olusturulan.append({"id": fid, "fatura_no": fatura_no, "yil": yil, "ay": ay, "toplam": toplam, "guncellendi": False})

    sync_musteri_panel_borclu_from_satirlar(musteri_id, satirlar)
    _upsert_aylik_grid_cache(musteri_id)
    yeni_n = sum(
        1 for x in olusturulan if isinstance(x, dict) and not x.get("zaten_vardi")
    )
    gunc_n = sum(1 for x in olusturulan if isinstance(x, dict) and x.get("zaten_vardi"))
    if tahsil_silinen and yeni_n == 0:
        mesaj = (
            f"{gunc_n} ay borçlu duruma alındı; "
            f"{len(tahsil_silinen)} tahsilat kaydı silindi."
        )
    elif gunc_n > 0 and yeni_n == 0 and not tahsil_silinen:
        mesaj = f"{gunc_n} ay zaten borçluydu; değişiklik yapılmadı."
    else:
        mesaj = f"{yeni_n} fatura oluşturuldu, {gunc_n} ay güncellendi, {len(atlanan)} satır atlandı."
        if tahsil_silinen:
            mesaj += f" {len(tahsil_silinen)} tahsilat kaydı silindi."
    if not olusturulan and atlanan and all(
        isinstance(x, dict) and x.get("neden") == "zaten_borclandirildi" for x in atlanan
    ):
        mesaj += (
            " Seçilen aylarda zaten «Aylık Tutarlar» borç faturası (|AYLIK_TUTAR|) var; "
            "yeniden borçlandırmak için önce «Borçtan çıkar» ile o ayların aylık faturasını silin "
            "(tahsilat kayıtları ayrı kalır)."
        )
    return jsonify({
        "ok": True,
        "olusturulan": olusturulan,
        "atlanan": atlanan,
        "tahsil_silinen": tahsil_silinen,
        "tahsil_silinen_aylar": [
            {"yil": int(y), "ay": int(a)} for (y, a) in sorted(tahsil_silinen_aylar)
        ],
        "mesaj": mesaj,
    })


@bp.route('/api/aylik-tutarlardan-tahsil-et', methods=['POST'])
@giris_gerekli
def api_aylik_tutarlardan_tahsil_et():
    """
    Aylık Tutarlar gridinden seçilen aylar için ayrı tahsilat (alacak) kaydı.
    Tutarlar griddeki gibi KDV dahil; cari kart ve ekstrelerde tahsilat olarak görünür.
    Aynı ay için tekrar kayıt engellenir; kira/TÜFE artışı sonrası marker tutarı
    grid tutarından düşükse mevcut kayıt güncellenir veya fark için ek kayıt açılır.
    """
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    musteri_id = data.get("musteri_id")
    try:
        musteri_id = int(musteri_id)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    cust = fetch_one("SELECT id, name FROM customers WHERE id = %s", (musteri_id,))
    if not cust:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404

    satirlar = data.get("satirlar")
    if not isinstance(satirlar, list) or not satirlar:
        return jsonify({"ok": False, "mesaj": "En az bir ay satırı gerekli."}), 400

    varsayilan_tahsilat_tarihi = date.today()
    raw_tarih = (data.get("tahsilat_tarihi") or "").strip()
    if raw_tarih:
        try:
            varsayilan_tahsilat_tarihi = datetime.strptime(raw_tarih[:10], "%Y-%m-%d").date()
        except Exception:
            pass

    odeme = (data.get("odeme_turu") or "havale").strip().lower()
    if odeme not in ("nakit", "havale", "eft", "banka", "kredi_karti", "cek"):
        odeme = "havale"

    olusturulan = []
    atlanan = []
    parsed = []

    for raw in satirlar:
        if not isinstance(raw, dict):
            continue
        try:
            yil = int(raw.get("yil"))
            ay = int(raw.get("ay"))
        except (TypeError, ValueError):
            atlanan.append({"neden": "geçersiz_yil_ay", "satir": raw})
            continue
        if ay < 1 or ay > 12 or yil < 1990 or yil > 2100:
            atlanan.append({"neden": "tarih_aralik", "satir": raw})
            continue
        try:
            tutar = float(raw.get("tutar_kdv_dahil"))
        except (TypeError, ValueError):
            atlanan.append({"neden": "geçersiz_tutar", "satir": raw})
            continue
        if tutar <= 0:
            atlanan.append({"neden": "tutar_sifir", "satir": raw})
            continue
        ay_bir = date(yil, ay, 1)
        tahsil_tarih_row = None
        ts_raw = (raw.get("tahsilat_tarihi") or "").strip()[:10]
        if ts_raw:
            try:
                tahsil_tarih_row = datetime.strptime(ts_raw, "%Y-%m-%d").date()
            except Exception:
                tahsil_tarih_row = None
        parsed.append({
            "yil": yil,
            "ay": ay,
            "ay_bir": ay_bir,
            "tutar": round(tutar, 2),
            "tahsilat_tarihi": tahsil_tarih_row,
        })

    def _norm_month_key(v):
        if isinstance(v, datetime):
            return v.date()
        return v

    try:
        with get_db() as conn:
            cur = conn.cursor()
            month_starts = sorted({p["ay_bir"] for p in parsed})
            fatura_by_month = {}
            if month_starts:
                cur.execute(
                    """
                    SELECT id, (DATE_TRUNC('month', fatura_tarihi::date))::date AS m
                    FROM faturalar
                    WHERE musteri_id = %s
                      AND (DATE_TRUNC('month', fatura_tarihi::date))::date IN %s
                    ORDER BY (DATE_TRUNC('month', fatura_tarihi::date))::date, id DESC
                    """,
                    (musteri_id, tuple(month_starts)),
                )
                for row in cur.fetchall() or []:
                    mk = _norm_month_key(row.get("m"))
                    if mk not in fatura_by_month:
                        fatura_by_month[mk] = row.get("id")

            # Duplicate kontrolü: yalnızca aynı aya ait |AYLIK_TAH| marker'lı kayıtlar.
            # Eski davranışta aynı ayda marker'sız herhangi bir tahsilat da duplicate sayılıyordu;
            # bu yüzden "0 oluşturuldu, N atlandı" hatalı görülüyordu.
            dup_makbuz_by_ab = {}
            if month_starts:
                cur.execute(
                    """
                    SELECT t.id, t.makbuz_no, COALESCE(t.aciklama, '') AS aciklama
                    FROM tahsilatlar t
                    WHERE (t.musteri_id = %s OR t.customer_id = %s)
                      AND COALESCE(t.tutar, 0) > 0
                      AND COALESCE(t.aciklama, '') LIKE '%%|AYLIK_TAH|%%'
                    ORDER BY t.id DESC
                    """,
                    (musteri_id, musteri_id),
                )
                by_iso = {}
                for row in cur.fetchall() or []:
                    ac = str(row.get("aciklama") or "")
                    mb = row.get("makbuz_no")
                    for iso in re.findall(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", ac):
                        if iso not in by_iso:
                            by_iso[iso] = mb
                for p in parsed:
                    iso = p["ay_bir"].isoformat()
                    if iso in by_iso:
                        dup_makbuz_by_ab[p["ay_bir"]] = by_iso[iso]

            cur.execute(
                "SELECT makbuz_no FROM tahsilatlar WHERE makbuz_no IS NOT NULL AND makbuz_no <> '' ORDER BY id DESC LIMIT 1"
            )
            row_m = cur.fetchone()
            makbuz_seq = 999
            if row_m and row_m.get("makbuz_no"):
                m = re.search(r"(\d+)$", str(row_m.get("makbuz_no") or "").strip())
                if m:
                    try:
                        makbuz_seq = int(m.group(1))
                    except Exception:
                        makbuz_seq = 999
            if makbuz_seq < 999:
                makbuz_seq = 999

            harf = _odeme_turu_harf(odeme)
            yerel_tahsil_iso = set()

            def _p_tahsil_tarihi(p_item):
                t = p_item.get("tahsilat_tarihi")
                if t:
                    return t
                if raw_tarih:
                    return varsayilan_tahsilat_tarihi
                return p_item["ay_bir"]

            tol_tah = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
            for p in parsed:
                iso = p["ay_bir"].isoformat()
                if p["ay_bir"] in dup_makbuz_by_ab:
                    mevcut_top = _aylik_tahsil_marker_odeme_toplam_for_iso(musteri_id, iso)
                    hedef_t = round(float(p["tutar"]), 2)
                    kayitlar = _aylik_tahsil_marker_kayitlari_for_iso(musteri_id, iso)
                    if abs(mevcut_top - hedef_t) <= tol_tah:
                        atlanan.append(
                            {
                                "neden": "zaten_tahsil",
                                "yil": p["yil"],
                                "ay": p["ay"],
                                "makbuz_no": dup_makbuz_by_ab[p["ay_bir"]],
                            }
                        )
                        continue
                    if len(kayitlar) == 1:
                        kid = int(kayitlar[0]["id"])
                        cur.execute(
                            "UPDATE tahsilatlar SET tutar = %s WHERE id = %s",
                            (hedef_t, kid),
                        )
                        olusturulan.append(
                            {
                                "id": kid,
                                "makbuz_no": kayitlar[0].get("makbuz_no"),
                                "yil": p["yil"],
                                "ay": p["ay"],
                                "tutar": hedef_t,
                                "guncellendi": True,
                            }
                        )
                        yerel_tahsil_iso.add(iso)
                        continue
                    if len(kayitlar) > 1 and mevcut_top > hedef_t + tol_tah:
                        kid_keep = int(kayitlar[0]["id"])
                        cur.execute(
                            "UPDATE tahsilatlar SET tutar = %s WHERE id = %s",
                            (hedef_t, kid_keep),
                        )
                        for extra in kayitlar[1:]:
                            cur.execute(
                                "DELETE FROM tahsilatlar WHERE id = %s",
                                (int(extra["id"]),),
                            )
                        olusturulan.append(
                            {
                                "id": kid_keep,
                                "makbuz_no": kayitlar[0].get("makbuz_no"),
                                "yil": p["yil"],
                                "ay": p["ay"],
                                "tutar": hedef_t,
                                "guncellendi": True,
                                "birlestirildi": True,
                            }
                        )
                        yerel_tahsil_iso.add(iso)
                        continue
                    fark = round(hedef_t - mevcut_top, 2)
                    if fark <= tol_tah:
                        atlanan.append(
                            {
                                "neden": "zaten_tahsil",
                                "yil": p["yil"],
                                "ay": p["ay"],
                                "makbuz_no": dup_makbuz_by_ab[p["ay_bir"]],
                            }
                        )
                        continue
                    makbuz_seq += 1
                    makbuz_no = str(makbuz_seq)
                    ay_adi = _AY_ADLARI[p["ay"] - 1]
                    marker = f"|AYLIK_TAH|{iso}|"
                    harf_up = _odeme_turu_harf(odeme)
                    aciklama = f"{ay_adi} {p['yil']} Tahsilat {harf_up} (fark){marker}"
                    tahsilat_tarihi = _p_tahsil_tarihi(p)
                    fatura_id = fatura_by_month.get(p["ay_bir"])
                    cur.execute(
                        """
                        INSERT INTO tahsilatlar (musteri_id, customer_id, fatura_id, tutar, odeme_turu, aciklama, tahsilat_tarihi, makbuz_no)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                        """,
                        (
                            musteri_id,
                            musteri_id,
                            fatura_id,
                            fark,
                            odeme,
                            aciklama,
                            tahsilat_tarihi,
                            makbuz_no,
                        ),
                    )
                    row_ins = cur.fetchone()
                    tid = row_ins.get("id") if row_ins else None
                    olusturulan.append(
                        {
                            "id": tid,
                            "makbuz_no": makbuz_no,
                            "yil": p["yil"],
                            "ay": p["ay"],
                            "tutar": fark,
                            "fark_tahsilat": True,
                        }
                    )
                    yerel_tahsil_iso.add(iso)
                    continue
                if iso in yerel_tahsil_iso:
                    atlanan.append(
                        {
                            "neden": "zaten_tahsil",
                            "yil": p["yil"],
                            "ay": p["ay"],
                            "makbuz_no": None,
                        }
                    )
                    continue
                makbuz_seq += 1
                makbuz_no = str(makbuz_seq)
                ay_adi = _AY_ADLARI[p["ay"] - 1]
                marker = f"|AYLIK_TAH|{iso}|"
                aciklama = f"{ay_adi} {p['yil']} Tahsilat {harf}{marker}"
                tahsilat_tarihi = _p_tahsil_tarihi(p)
                fatura_id = fatura_by_month.get(p["ay_bir"])
                cur.execute(
                    """
                    INSERT INTO tahsilatlar (musteri_id, customer_id, fatura_id, tutar, odeme_turu, aciklama, tahsilat_tarihi, makbuz_no)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                    """,
                    (
                        musteri_id,
                        musteri_id,
                        fatura_id,
                        p["tutar"],
                        odeme,
                        aciklama,
                        tahsilat_tarihi,
                        makbuz_no,
                    ),
                )
                row_ins = cur.fetchone()
                tid = row_ins.get("id") if row_ins else None
                olusturulan.append(
                    {
                        "id": tid,
                        "makbuz_no": makbuz_no,
                        "yil": p["yil"],
                        "ay": p["ay"],
                        "tutar": p["tutar"],
                    }
                )
                yerel_tahsil_iso.add(iso)
    except Exception as e:
        return jsonify({"ok": False, "mesaj": f"Tahsilat toplu kayıt hatası: {e}"}), 500

    try:
        sync_musteri_panel_from_tahsil_and_dagitim(musteri_id)
    except Exception:
        try:
            _upsert_aylik_grid_cache(musteri_id, tufe_map=_tufe_map_by_year_month())
        except Exception:
            pass
    _cari_ekstre_api_cache.clear()
    zaten_n = sum(
        1 for x in atlanan if isinstance(x, dict) and x.get("neden") == "zaten_tahsil"
    )
    yeni_n = len(olusturulan)
    if yeni_n == 0 and zaten_n > 0:
        mesaj = (
            f"{zaten_n} ay zaten tahsil edilmişti; panel ve grid güncellendi."
        )
    elif yeni_n > 0:
        mesaj = (
            f"{yeni_n} tahsilat kaydı oluşturuldu, {len(atlanan)} satır atlandı."
        )
    else:
        mesaj = f"{yeni_n} tahsilat kaydı oluşturuldu, {len(atlanan)} satır atlandı."
    return jsonify({
        "ok": True,
        "olusturulan": olusturulan,
        "atlanan": atlanan,
        "mesaj": mesaj,
    })


def _aylik_grid_parse_ay_listesi(satirlar):
    """{yil, ay} satırlarından benzersiz (yil, ay) listesi; geçersizleri atlar."""
    seen = set()
    out = []
    if not isinstance(satirlar, list):
        return out
    for raw in satirlar:
        if not isinstance(raw, dict):
            continue
        try:
            yil = int(raw.get("yil"))
            ay = int(raw.get("ay"))
        except (TypeError, ValueError):
            continue
        if ay < 1 or ay > 12 or yil < 1990 or yil > 2100:
            continue
        k = (yil, ay)
        if k in seen:
            continue
        seen.add(k)
        out.append(k)
    return out


@bp.route("/api/aylik-tutarlardan-tahsilden-cikar", methods=["POST"])
@giris_gerekli
def api_aylik_tutarlardan_tahsilden_cikar():
    """
    Grid ile oluşturulmuş aylık tahsilatları (|AYLIK_TAH| işaretçisi) siler;
    ay hücresi caride yeniden ödenmemiş görünür.
    """
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    musteri_id = data.get("musteri_id")
    try:
        musteri_id = int(musteri_id)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    if not fetch_one("SELECT id FROM customers WHERE id = %s", (musteri_id,)):
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404

    aylar = _aylik_grid_parse_ay_listesi(data.get("satirlar"))
    if not aylar:
        return jsonify({"ok": False, "mesaj": "En az bir ay seçin (geçerli yıl/ay)."}), 400

    silinen = []
    atlanan = []

    try:
        with get_db() as conn:
            cur = conn.cursor()
            for yil, ay in aylar:
                iso = date(yil, ay, 1).isoformat()
                pat = f"%|AYLIK_TAH|{iso}|%"
                cur.execute(
                    """
                    DELETE FROM tahsilatlar
                    WHERE (musteri_id = %s OR customer_id = %s)
                      AND COALESCE(aciklama, '') LIKE %s
                    RETURNING id, makbuz_no
                    """,
                    (musteri_id, musteri_id, pat),
                )
                rows = cur.fetchall() or []
                if rows:
                    for r in rows:
                        silinen.append(
                            {
                                "id": r.get("id"),
                                "makbuz_no": r.get("makbuz_no"),
                                "yil": yil,
                                "ay": ay,
                            }
                        )
                else:
                    atlanan.append({"neden": "kayit_yok", "yil": yil, "ay": ay})
    except Exception as e:
        return jsonify({"ok": False, "mesaj": f"Tahsilat silinemedi: {e}"}), 500

    borc_satirlar = []
    raw_satirlar = data.get("satirlar")
    tutar_by_ay = {}
    if isinstance(raw_satirlar, list):
        for raw in raw_satirlar:
            if not isinstance(raw, dict):
                continue
            try:
                yy = int(raw.get("yil"))
                mm = int(raw.get("ay"))
                tt = round(float(raw.get("tutar_kdv_dahil")), 2)
            except (TypeError, ValueError):
                continue
            if 1 <= mm <= 12:
                tutar_by_ay[(yy, mm)] = tt
    for yil, ay in aylar:
        row = {"yil": yil, "ay": ay}
        if (yil, ay) in tutar_by_ay:
            row["tutar_kdv_dahil"] = tutar_by_ay[(yil, ay)]
        borc_satirlar.append(row)
    sync_musteri_panel_borclu_from_satirlar(musteri_id, borc_satirlar)
    _upsert_aylik_grid_cache(musteri_id)
    return jsonify(
        {
            "ok": True,
            "silinen": silinen,
            "atlanan": atlanan,
            "mesaj": f"{len(silinen)} tahsilat silindi, {len(atlanan)} ayda kayıt yoktu.",
        }
    )


@bp.route("/api/aylik-tutarlardan-borctan-cikar", methods=["POST"])
@giris_gerekli
def api_aylik_tutarlardan_borctan_cikar():
    """
    Grid ile oluşturulmuş aylık faturaları (|AYLIK_TUTAR| işaretçisi) siler.
    Önce bu faturaya bağlı tahsilatlar silinir (FK); cari borç düşer.
    """
    ensure_faturalar_amount_columns()
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    musteri_id = data.get("musteri_id")
    try:
        musteri_id = int(musteri_id)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    if not fetch_one("SELECT id FROM customers WHERE id = %s", (musteri_id,)):
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404

    aylar = _aylik_grid_parse_ay_listesi(data.get("satirlar"))
    if not aylar:
        return jsonify({"ok": False, "mesaj": "En az bir ay seçin (geçerli yıl/ay)."}), 400

    silinen = []
    atlanan = []

    try:
        with get_db() as conn:
            cur = conn.cursor()
            for yil, ay in aylar:
                iso = date(yil, ay, 1).isoformat()
                pat = f"%|AYLIK_TUTAR|{iso}|%"
                cur.execute(
                    """
                    SELECT id, fatura_no FROM faturalar
                    WHERE musteri_id = %s AND COALESCE(notlar, '') LIKE %s
                    LIMIT 1
                    """,
                    (musteri_id, pat),
                )
                fr = cur.fetchone()
                if not fr:
                    atlanan.append({"neden": "fatura_yok", "yil": yil, "ay": ay})
                    continue
                fid = fr.get("id")
                fno = fr.get("fatura_no")
                cur.execute("DELETE FROM tahsilatlar WHERE fatura_id = %s", (fid,))
                cur.execute(
                    """
                    DELETE FROM faturalar
                    WHERE id = %s AND musteri_id = %s AND COALESCE(notlar, '') LIKE %s
                    RETURNING id
                    """,
                    (fid, musteri_id, pat),
                )
                if cur.fetchone():
                    silinen.append({"id": fid, "fatura_no": fno, "yil": yil, "ay": ay})
                else:
                    atlanan.append({"neden": "silinemedi", "yil": yil, "ay": ay})
    except Exception as e:
        return jsonify({"ok": False, "mesaj": f"Fatura silinemedi: {e}"}), 500

    _upsert_aylik_grid_cache(musteri_id)
    return jsonify(
        {
            "ok": True,
            "silinen": silinen,
            "atlanan": atlanan,
            "mesaj": f"{len(silinen)} fatura silindi, {len(atlanan)} ayda kayıt yoktu.",
        }
    )


@bp.route('/api/aylik-grid-cache')
@giris_gerekli
def api_aylik_grid_cache():
    musteri_id = request.args.get("musteri_id", type=int)
    if not musteri_id:
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    force = str(request.args.get("force") or "").lower() in ("1", "true", "yes", "on")
    skip_match = str(request.args.get("skip_match") or "").lower() in ("1", "true", "yes", "on")
    _ensure_aylik_grid_cache_table()
    if not force:
        try:
            mem_hit = _aylik_grid_payload_mem.get(int(musteri_id))
            if mem_hit and (time.time() - float(mem_hit[0])) < 60.0 and isinstance(mem_hit[1], dict):
                mem_payload = _aylik_grid_cache_payload_tahsil_guncelle(musteri_id, mem_hit[1])
                mem_payload = _aylik_grid_payload_reel_overlay_from_db(musteri_id, mem_payload)
                try:
                    _aylik_grid_payload_mem[int(musteri_id)] = (time.time(), mem_payload)
                except (TypeError, ValueError):
                    pass
                return jsonify({"ok": True, "cache": mem_payload, "cached": True, "mem": True})
        except (TypeError, ValueError):
            pass
        row = fetch_one("SELECT payload FROM musteri_aylik_grid_cache WHERE musteri_id = %s", (musteri_id,))
        if row and row.get("payload"):
            try:
                cache_obj = json.loads(row["payload"])
                # Hızlı yol: skip_match ile çağrıldıysa normalde kyc-uyum kontrolünü
                # atlayıp önbelleği döndürürüz. Ancak imzasız (eski şema) payload'larda
                # bu davranış stale veriyi sonsuza kadar tutabilir.
                skip_izinli = bool(
                    skip_match
                    and isinstance(cache_obj, dict)
                    and cache_obj.get("tahsilat_imza")
                    and int(cache_obj.get("compute_rev") or 0) == AYLIK_GRID_COMPUTE_REV
                )
                cache_gecerli = skip_izinli or _aylik_grid_cache_matches_kyc(musteri_id, cache_obj)
                # Kart yenileme: KYC alanları formdan henüz gelmemiş olsa bile tam rebuild yapma.
                if not cache_gecerli and skip_match and isinstance(cache_obj, dict):
                    aylar = cache_obj.get("aylar")
                    cache_gecerli = isinstance(aylar, list) and len(aylar) > 0
                if cache_gecerli:
                    cache_obj = _aylik_grid_cache_payload_tahsil_guncelle(musteri_id, cache_obj)
                    cache_obj = _aylik_grid_payload_reel_overlay_from_db(musteri_id, cache_obj)
                    try:
                        _aylik_grid_payload_mem[int(musteri_id)] = (time.time(), cache_obj)
                    except (TypeError, ValueError):
                        pass
                    return jsonify({"ok": True, "cache": cache_obj, "cached": True})
            except Exception:
                pass
    if not fetch_one("SELECT id FROM customers WHERE id = %s", (musteri_id,)):
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    payload = _upsert_aylik_grid_cache(musteri_id)
    return jsonify({"ok": True, "cache": payload or {}, "cached": False})


@bp.route("/api/reel-donem-tutarlar")
@giris_gerekli
def api_reel_donem_tutarlar():
    musteri_id = request.args.get("musteri_id", type=int)
    if not musteri_id:
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    # Müşteri var mı kontrolünü kaldırdık; doğrudan tablo sorgusu yapıyoruz.
    # Yoksa 0 satır döner, çağıran tarafa zaten ok=True/empty map geliyor.
    _ensure_musteri_reel_donem_tutar_table()
    rows = fetch_all(
        """
        SELECT donem_yil, tutar_kdv_dahil, giris_tip, giris_tutar, hibrit_toplam, hibrit_net, hibrit_banka
        FROM musteri_reel_donem_tutar
        WHERE musteri_id = %s
        ORDER BY donem_yil
        """,
        (musteri_id,),
    ) or []
    m = {}
    detay_map = {}
    for r in rows:
        y = r.get("donem_yil")
        t = r.get("tutar_kdv_dahil")
        if y is None or t is None:
            continue
        yk = None
        try:
            yk = str(int(y))
            m[yk] = float(t)
        except (TypeError, ValueError):
            continue
        detay_map[yk] = {
            "tip": (r.get("giris_tip") or "").strip().lower() or None,
            "giris_tutar": float(r.get("giris_tutar")) if r.get("giris_tutar") is not None else None,
            "hibrit_toplam": float(r.get("hibrit_toplam")) if r.get("hibrit_toplam") is not None else None,
            "hibrit_net": float(r.get("hibrit_net")) if r.get("hibrit_net") is not None else None,
            "hibrit_banka": float(r.get("hibrit_banka")) if r.get("hibrit_banka") is not None else None,
        }
    return jsonify({"ok": True, "map": m, "detay_map": detay_map})


@bp.route("/api/reel-donem-tutar", methods=["POST"])
@giris_gerekli
def api_reel_donem_tutar_upsert():
    data = request.get_json(silent=True) or {}
    try:
        musteri_id = int(data.get("musteri_id"))
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    try:
        donem_yil = int(data.get("donem_yil"))
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "donem_yil gerekli."}), 400
    raw_tutar = data.get("tutar_kdv_dahil")
    if raw_tutar is None or str(raw_tutar).strip() == "":
        return jsonify({"ok": False, "mesaj": "tutar_kdv_dahil gerekli."}), 400
    try:
        tutar = float(str(raw_tutar).replace(",", ".").strip())
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "Geçersiz tutar."}), 400
    if tutar < 0:
        return jsonify({"ok": False, "mesaj": "Tutar negatif olamaz."}), 400
    giris_tip = (data.get("giris_tip") or "").strip().lower()
    if giris_tip not in ("dahil", "haric", "net", "hibrit"):
        giris_tip = "dahil"
    def _optf(v):
        if v is None or str(v).strip() == "":
            return None
        try:
            return float(str(v).replace(",", ".").strip())
        except Exception:
            return None
    giris_tutar = _optf(data.get("giris_tutar"))
    hibrit_toplam = _optf(data.get("hibrit_toplam"))
    hibrit_net = _optf(data.get("hibrit_net"))
    hibrit_banka = _optf(data.get("hibrit_banka"))
    if not fetch_one("SELECT id FROM customers WHERE id = %s", (musteri_id,)):
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    _ensure_musteri_reel_donem_tutar_table()
    execute(
        """
        INSERT INTO musteri_reel_donem_tutar (
            musteri_id, donem_yil, tutar_kdv_dahil, giris_tip, giris_tutar, hibrit_toplam, hibrit_net, hibrit_banka
        )
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        ON CONFLICT (musteri_id, donem_yil) DO UPDATE SET
            tutar_kdv_dahil = EXCLUDED.tutar_kdv_dahil,
            giris_tip = EXCLUDED.giris_tip,
            giris_tutar = EXCLUDED.giris_tutar,
            hibrit_toplam = EXCLUDED.hibrit_toplam,
            hibrit_net = EXCLUDED.hibrit_net,
            hibrit_banka = EXCLUDED.hibrit_banka,
            updated_at = NOW()
        """,
        (musteri_id, donem_yil, tutar, giris_tip, giris_tutar, hibrit_toplam, hibrit_net, hibrit_banka),
    )
    try:
        _cari_ekstre_api_cache.clear()
    except Exception:
        pass
    try:
        _invalidate_aylik_grid_payload_mem(musteri_id)
        _upsert_aylik_grid_cache(musteri_id)
    except Exception:
        logging.getLogger(__name__).exception(
            "reel_donem_tutar sonrasi grid cache musteri_id=%s", musteri_id
        )
    try:
        mid_int = int(musteri_id)
        payload_after = _build_aylik_grid_cache_payload(mid_int)
        existing_panel = _load_musteri_panel_by_iso(mid_int)
        patch = {}
        if payload_after and isinstance(payload_after.get("aylar"), list):
            for a in payload_after["aylar"]:
                if not isinstance(a, dict):
                    continue
                yil_p = a.get("yil")
                ay_p = a.get("ay")
                if not yil_p or not ay_p:
                    continue
                iso_k = date(int(yil_p), int(ay_p), 1).isoformat()
                new_brut = round(float(a.get("brut_tutar_kdv") or 0), 2)
                if new_brut <= 0:
                    continue
                existing_row = existing_panel.get(iso_k) or {}
                eski_brut = round(float(existing_row.get("aylik") or 0), 2)
                eski_tahsil = round(float(existing_row.get("tahsil") or 0), 2)
                tol_p = 0.05
                if eski_tahsil >= eski_brut - tol_p and eski_brut > tol_p:
                    yeni_tahsil = new_brut
                    yeni_kalan = 0.0
                else:
                    yeni_tahsil = eski_tahsil
                    yeni_kalan = max(round(new_brut - eski_tahsil, 2), 0.0)
                patch[iso_k] = {
                    "aylik": new_brut,
                    "tahsil": yeni_tahsil,
                    "kalan": yeni_kalan,
                    "tahsil_tarih": existing_row.get("tahsil_tarih"),
                }
        if patch:
            _save_musteri_panel_by_iso(mid_int, patch, prune_no_db_tahsil=False)
            _invalidate_aylik_grid_payload_mem(mid_int)
            _upsert_aylik_grid_cache(mid_int)
    except Exception:
        logging.getLogger(__name__).exception(
            "reel_donem_tutar sonrasi panel aylik senkron musteri_id=%s", musteri_id
        )
    return jsonify({"ok": True})


@bp.route('/api/aylik-grid-cache/rebuild-all', methods=['POST'])
@giris_gerekli
def api_aylik_grid_cache_rebuild_all():
    rows = fetch_all("SELECT id FROM customers ORDER BY id") or []
    updated = 0
    for r in rows:
        mid = r.get("id")
        if not mid:
            continue
        if _upsert_aylik_grid_cache(int(mid)):
            updated += 1
    return jsonify({"ok": True, "updated": updated, "mesaj": f"{updated} müşteri için aylık grid cache güncellendi."})


def _optional_musteri_id_set_from_post(data):
    """
    JSON'da musteri_ids yoksa None → tüm (pasif olmayan) adaylar.
    Varsa pozitif tam sayı kümesi; boş veya geçersiz liste ValueError.
    """
    if not isinstance(data, dict) or "musteri_ids" not in data:
        return None
    raw = data.get("musteri_ids")
    if not isinstance(raw, list):
        raise ValueError("musteri_ids bir liste olmalıdır.")
    out = set()
    for x in raw:
        try:
            mid = int(x)
            if mid > 0:
                out.add(mid)
        except (TypeError, ValueError):
            continue
    if not out:
        raise ValueError("musteri_ids içinde geçerli müşteri numarası yok.")
    return out


@bp.route('/api/aylik-kira-guncelle-ve-borclandir-all', methods=['POST'])
@giris_gerekli
def api_aylik_kira_guncelle_ve_borclandir_all():
    """
    Aktif müşteriler için:
    1) TÜFE hesaplı güncel kira bedelini customers.guncel_kira_bedeli alanına yazar.
    2) Bugüne kadar eksik aylık borç faturalarını (AYLIK_TUTAR marker) toplu oluşturur.

    POST JSON (isteğe bağlı): musteri_ids: [1,2,3] → yalnız bu müşteri kartları.

    Performans: TÜFE ve KYC tek seferde; tahsilat sorgusu atlanır; fatura INSERT tek bağlantıda toplu.
    """
    ensure_faturalar_amount_columns()
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    try:
        id_filter = _optional_musteri_id_set_from_post(data)
    except ValueError as ve:
        return jsonify({"ok": False, "mesaj": str(ve)}), 400
    today = date.today()
    KDV_ORAN = 20.0
    tufe_map = _tufe_map_by_year_month()

    rows = fetch_all(
        """
        SELECT c.id, c.name
        FROM customers c
        WHERE LOWER(COALESCE(c.durum, 'aktif')) != 'pasif'
        ORDER BY c.id
        """
    ) or []
    if id_filter is not None:
        rows = [r for r in rows if int(r.get("id") or 0) in id_filter]

    kyc_rows = fetch_all(
        """
        SELECT DISTINCT ON (musteri_id) musteri_id, sozlesme_tarihi, sozlesme_bitis, aylik_kira, kira_artis_tarihi, kira_suresi_ay
        FROM musteri_kyc
        ORDER BY musteri_id, id DESC
        """
    ) or []
    kyc_by_mid = {}
    for kr in kyc_rows:
        try:
            mid_k = int(kr.get("musteri_id") or 0)
        except (TypeError, ValueError):
            continue
        if mid_k > 0:
            kyc_by_mid[mid_k] = kr

    marker_rows = fetch_all(
        f"""
        SELECT musteri_id, notlar FROM faturalar
        WHERE COALESCE(notlar, '') LIKE '%%|AYLIK_TUTAR|%%'
          AND {sql_expr_fatura_not_gib_taslak("notlar")}
        """
    ) or []
    markers_by_mid = defaultdict(set)
    for mr in marker_rows:
        try:
            mid_m = int(mr.get("musteri_id") or 0)
        except (TypeError, ValueError):
            continue
        if mid_m <= 0:
            continue
        for k in re.findall(r"\|AYLIK_TUTAR\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", str(mr.get("notlar") or "")):
            markers_by_mid[mid_m].add(k)

    prefix = "INV"
    yil = datetime.now().year
    like = f"{prefix}{yil}%"
    row = fetch_one(
        "SELECT fatura_no FROM faturalar WHERE fatura_no LIKE %s ORDER BY id DESC LIMIT 1",
        (like,),
    )
    fatura_tail = 0
    if row and row.get("fatura_no"):
        try:
            fatura_tail = int(str(row["fatura_no"])[-6:])
        except (ValueError, IndexError):
            fatura_tail = 0

    guncel_updated = 0
    borc_eklenen = 0
    atlanan = 0
    customer_updates = []
    insert_rows = []

    for r in rows:
        mid = int(r.get("id") or 0)
        if mid <= 0:
            continue
        kyc = kyc_by_mid.get(mid)
        if not kyc:
            continue
        payload = _aylik_grid_compute(mid, kyc, tufe_map)
        if not payload:
            continue
        aylar = payload.get("aylar") or []
        if not isinstance(aylar, list) or not aylar:
            continue

        last_tutar_kdv = 0.0
        for a in aylar:
            try:
                y = int(a.get("yil"))
                m = int(a.get("ay"))
                if date(y, m, 1) <= today.replace(day=1):
                    last_tutar_kdv = float(a.get("tutar_kdv_dahil") or last_tutar_kdv or 0)
            except Exception:
                continue
        if last_tutar_kdv > 0:
            net_guncel = round(last_tutar_kdv / (1 + KDV_ORAN / 100.0), 2)
            customer_updates.append((net_guncel, mid))
            guncel_updated += 1

        marker_set = set(markers_by_mid[mid])
        musteri_adi = (r.get("name") or "—").strip() or "—"

        for a in aylar:
            try:
                yil_a = int(a.get("yil"))
                ay = int(a.get("ay"))
                ay_bir = date(yil_a, ay, 1)
                if ay_bir > today.replace(day=1):
                    continue
                toplam = round(float(a.get("tutar_kdv_dahil") or 0), 2)
            except Exception:
                atlanan += 1
                continue
            if toplam <= 0:
                atlanan += 1
                continue
            key = ay_bir.isoformat()
            if key in marker_set:
                atlanan += 1
                continue

            if ay == 12:
                vade = date(yil_a, 12, 31)
            else:
                vade = date(yil_a, ay + 1, 1) - timedelta(days=1)
            net = round(toplam / (1 + KDV_ORAN / 100.0), 2)
            kdv_tutar = round(toplam - net, 2)
            net = round(toplam - kdv_tutar, 2)
            ay_adi = _AY_ADLARI[ay - 1]
            marker = f"|AYLIK_TUTAR|{key}|"
            notlar = f"{ay_adi} {yil_a} kira bedeli (otomatik toplu borçlandırma){marker}"
            fatura_tail += 1
            fatura_no = f"{prefix}{yil}{fatura_tail:06d}"
            insert_rows.append(
                (fatura_no, mid, musteri_adi, net, kdv_tutar, toplam, "odenmedi", ay_bir, vade, notlar)
            )
            marker_set.add(key)
            markers_by_mid[mid].add(key)
            borc_eklenen += 1

    conn = get_conn()
    try:
        cur = conn.cursor()
        if customer_updates:
            cur.executemany(
                "UPDATE customers SET guncel_kira_bedeli = %s WHERE id = %s",
                customer_updates,
            )
        ins_sql = """
            INSERT INTO faturalar (
                fatura_no, musteri_id, musteri_adi, tutar, kdv_tutar,
                toplam, durum, fatura_tarihi, vade_tarihi, notlar
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """
        chunk = 400
        for i in range(0, len(insert_rows), chunk):
            cur.executemany(ins_sql, insert_rows[i : i + chunk])
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()

    return jsonify({
        "ok": True,
        "updated_customers": guncel_updated,
        "borc_eklenen": borc_eklenen,
        "atlanan": atlanan,
        "mesaj": f"Güncel kira güncellendi: {guncel_updated}, borç eklenen ay: {borc_eklenen}, atlanan: {atlanan}",
    })


@bp.route("/api/tufe-borclandir-nakit-tahsil-toplu", methods=["POST"])
@giris_gerekli
def api_tufe_borclandir_nakit_tahsil_toplu():
    """
    Sözleşme başlangıcından hedef aya (varsayılan Nisan 2026) kadar TÜFE zincirli aylık tutarlarla:
    - Eksik aylar için borç (|AYLIK_TUTAR|) faturası,
    - Aynı tutarda nakit (veya seçilen ödeme) tahsilat (|AYLIK_TAH|, açıklamada |BTUFRT|),
    böylece cari ekstrede borç ve tahsil çifti görünür.

    Bu otomatik tahsil satırları gerçek para girişi değildir; aylık gridde ödenen toplamı için
    yalnızca BTUFRT_GRID_ODENME_DAHIL_DEGIL (varsayılan 2026–01–01 sonrası) hedefleri hariç tutulur
    ki geçmiş yıllar yeşilden düşmesin.

    Varsayılan: AYLIK_TUTAR dışı fatura yazılmış müşteriler atlanır; tüm hedef aylarda zaten
    borç+tahsil marker'ı olan müşteriler atlanır.

    POST JSON (isteğe bağlı): musteri_ids: [1,2,3] → yalnız bu müşteriler.
    """
    ensure_faturalar_amount_columns()
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    try:
        id_filter_tufe = _optional_musteri_id_set_from_post(data)
    except ValueError as ve:
        return jsonify({"ok": False, "mesaj": str(ve)}), 400
    try:
        hedef_yil = int(data.get("hedef_yil") or 2026)
        hedef_ay = int(data.get("hedef_ay") or 4)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "hedef_yil / hedef_ay sayı olmalı."}), 400
    if hedef_ay < 1 or hedef_ay > 12 or hedef_yil < 2000 or hedef_yil > 2100:
        return jsonify({"ok": False, "mesaj": "Geçersiz hedef tarih."}), 400

    dry_run = bool(data.get("dry_run"))
    manuel_faturali_musteri_atla = data.get("manuel_faturali_musteri_atla")
    if manuel_faturali_musteri_atla is None:
        manuel_faturali_musteri_atla = True
    tam_otomatik_musteri_atla = data.get("tam_otomatik_musteri_atla")
    if tam_otomatik_musteri_atla is None:
        tam_otomatik_musteri_atla = True
    guncel_kira_guncelle = bool(data.get("guncel_kira_guncelle", True))

    odeme = (data.get("odeme_turu") or "nakit").strip().lower()
    if odeme not in ("nakit", "havale", "eft", "banka", "kredi_karti", "cek"):
        odeme = "nakit"

    run_tag = (data.get("run_tag") or "").strip() or secrets.token_hex(6)
    tag_suffix = f"|BTUFRT|{run_tag}|"
    KDV_ORAN = 20.0
    today = date.today()
    hedef_ay_bir = date(hedef_yil, hedef_ay, 1)

    tufe_map = _tufe_map_by_year_month()
    tah_keys_by_mid = _load_aylik_tahsil_ay_keys_by_musteri()
    manual_ay_by_mid = _load_manual_fatura_ay_by_musteri() if manuel_faturali_musteri_atla else {}

    marker_rows = fetch_all(
        f"""
        SELECT musteri_id, notlar, COALESCE(NULLIF(TRIM(ettn), ''), '') AS ettn_g
        FROM faturalar
        WHERE COALESCE(notlar, '') LIKE '%%|AYLIK_TUTAR|%%'
          AND {sql_expr_fatura_not_gib_taslak("notlar")}
        """
    ) or []
    borc_keys_by_mid = defaultdict(set)
    borc_ettn_by_key = {}
    for mr in marker_rows:
        try:
            mid_m = int(mr.get("musteri_id") or 0)
        except (TypeError, ValueError):
            continue
        if mid_m <= 0:
            continue
        notlar = str(mr.get("notlar") or "")
        for k in re.findall(r"\|AYLIK_TUTAR\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", notlar):
            borc_keys_by_mid[mid_m].add(k)
            if (mr.get("ettn_g") or "").strip():
                borc_ettn_by_key[(mid_m, k)] = True

    rows = fetch_all(
        """
        SELECT c.id, c.name
        FROM customers c
        WHERE LOWER(COALESCE(c.durum, 'aktif')) != 'pasif'
        ORDER BY c.id
        """
    ) or []
    if id_filter_tufe is not None:
        rows = [r for r in rows if int(r.get("id") or 0) in id_filter_tufe]

    kyc_rows = fetch_all(
        """
        SELECT DISTINCT ON (musteri_id) musteri_id, sozlesme_tarihi, sozlesme_bitis, aylik_kira,
               kira_artis_tarihi, kira_suresi_ay, kira_nakit
        FROM musteri_kyc
        ORDER BY musteri_id, id DESC
        """
    ) or []
    kyc_by_mid = {}
    for kr in kyc_rows:
        try:
            mk = int(kr.get("musteri_id") or 0)
        except (TypeError, ValueError):
            continue
        if mk > 0:
            kyc_by_mid[mk] = kr

    prefix = "INV"
    yil_fn = datetime.now().year
    like = f"{prefix}{yil_fn}%"
    row = fetch_one(
        "SELECT fatura_no FROM faturalar WHERE fatura_no LIKE %s ORDER BY id DESC LIMIT 1",
        (like,),
    )
    fatura_tail = 0
    if row and row.get("fatura_no"):
        try:
            fatura_tail = int(str(row["fatura_no"])[-6:])
        except (ValueError, IndexError):
            fatura_tail = 0

    row_m = fetch_one(
        "SELECT makbuz_no FROM tahsilatlar WHERE makbuz_no IS NOT NULL AND makbuz_no <> '' ORDER BY id DESC LIMIT 1"
    )
    makbuz_seq = 999
    if row_m and row_m.get("makbuz_no"):
        m = re.search(r"(\d+)$", str(row_m.get("makbuz_no") or "").strip())
        if m:
            try:
                makbuz_seq = int(m.group(1))
            except Exception:
                makbuz_seq = 999
    if makbuz_seq < 999:
        makbuz_seq = 999

    insert_rows = []
    tahsil_plan = []
    touched_mids = set()
    customer_updates = []
    guncel_updated = 0
    musteri_atlanan = 0
    musteri_islenen = 0
    borc_atlandi = 0
    tahsil_atlandi = 0
    harf = _odeme_turu_harf(odeme)

    for r in rows:
        mid = int(r.get("id") or 0)
        if mid <= 0:
            continue
        kyc = kyc_by_mid.get(mid)
        if not kyc:
            musteri_atlanan += 1
            continue

        bas = kyc.get("sozlesme_tarihi")
        bit = kyc.get("sozlesme_bitis")
        try:
            bas = bas if hasattr(bas, "year") else datetime.strptime(str(bas)[:10], "%Y-%m-%d").date()
            bit = bit if hasattr(bit, "year") else datetime.strptime(str(bit)[:10], "%Y-%m-%d").date()
        except Exception:
            musteri_atlanan += 1
            continue
        if not bas or not bit:
            musteri_atlanan += 1
            continue

        bit_end = bit - timedelta(days=1)
        soz_last_month = date(bit_end.year, bit_end.month, 1)
        bas_first = date(bas.year, bas.month, 1)
        last_month = min(hedef_ay_bir, soz_last_month)
        if last_month < bas_first:
            musteri_atlanan += 1
            continue

        kira_nakit_borc = bool(kyc.get("kira_nakit"))
        payload = _aylik_grid_compute(mid, kyc, tufe_map)
        if not payload:
            musteri_atlanan += 1
            continue
        aylar = payload.get("aylar") or []
        if not isinstance(aylar, list) or not aylar:
            musteri_atlanan += 1
            continue

        ay_plan = []
        for a in aylar:
            try:
                yil_a = int(a.get("yil"))
                ay = int(a.get("ay"))
                ay_bir = date(yil_a, ay, 1)
            except Exception:
                continue
            if ay_bir < bas_first or ay_bir > last_month:
                continue
            brut_raw = a.get("brut_tutar_kdv")
            try:
                if brut_raw is not None and brut_raw != "":
                    toplam = round(float(brut_raw), 2)
                else:
                    toplam = round(float(a.get("tutar_kdv_dahil") or 0), 2)
            except Exception:
                continue
            if toplam <= 0:
                continue
            ay_plan.append((ay_bir, yil_a, ay, toplam))

        if not ay_plan:
            musteri_atlanan += 1
            continue

        if manuel_faturali_musteri_atla:
            man_set = manual_ay_by_mid.get(mid) or set()
            if any(ab[0] in man_set for ab in ay_plan):
                musteri_atlanan += 1
                continue

        if tam_otomatik_musteri_atla:
            all_borc_tah = True
            for ay_bir, _y, _m, _t in ay_plan:
                key = ay_bir.isoformat()
                if key not in borc_keys_by_mid[mid] or key not in tah_keys_by_mid[mid]:
                    all_borc_tah = False
                    break
            if all_borc_tah:
                musteri_atlanan += 1
                continue

        musteri_islenen += 1
        musteri_adi = (r.get("name") or "—").strip() or "—"

        if guncel_kira_guncelle:
            last_tutar_kdv = 0.0
            for a in aylar:
                try:
                    y = int(a.get("yil"))
                    m = int(a.get("ay"))
                    if date(y, m, 1) <= today.replace(day=1):
                        last_tutar_kdv = float(a.get("tutar_kdv_dahil") or last_tutar_kdv or 0)
                except Exception:
                    continue
            if last_tutar_kdv > 0:
                if kira_nakit_borc:
                    net_guncel = round(last_tutar_kdv, 2)
                else:
                    net_guncel = round(last_tutar_kdv / (1 + KDV_ORAN / 100.0), 2)
                customer_updates.append((net_guncel, mid))
                guncel_updated += 1

        marker_local = set(borc_keys_by_mid[mid])
        tah_local = set(tah_keys_by_mid[mid])

        for ay_bir, yil_a, ay, toplam in ay_plan:
            key = ay_bir.isoformat()
            marker = f"|AYLIK_TUTAR|{key}|"
            ay_adi = _AY_ADLARI[ay - 1]

            borc_gerek = key not in marker_local and not borc_ettn_by_key.get((mid, key))

            if borc_gerek:
                if ay == 12:
                    vade = date(yil_a, 12, 31)
                else:
                    vade = date(yil_a, ay + 1, 1) - timedelta(days=1)
                if kira_nakit_borc:
                    net = round(toplam, 2)
                    kdv_tutar = 0.0
                else:
                    net = round(toplam / (1 + KDV_ORAN / 100.0), 2)
                    kdv_tutar = round(toplam - net, 2)
                    net = round(toplam - kdv_tutar, 2)
                if kira_nakit_borc:
                    notlar = (
                        f"{ay_adi} {yil_a} kira bedeli (nakit/KDV yok, TÜFE borç+tahsil toplu){marker}{tag_suffix}"
                    )
                else:
                    notlar = (
                        f"{ay_adi} {yil_a} kira bedeli (KDV dahil, TÜFE borç+tahsil toplu){marker}{tag_suffix}"
                    )
                fatura_tail += 1
                fatura_no = f"{prefix}{yil_fn}{fatura_tail:06d}"
                insert_rows.append(
                    (
                        fatura_no,
                        mid,
                        musteri_adi,
                        net,
                        kdv_tutar,
                        toplam,
                        "odenmedi",
                        ay_bir,
                        vade,
                        notlar,
                    )
                )
                marker_local.add(key)
                borc_keys_by_mid[mid].add(key)
                touched_mids.add(mid)
            else:
                borc_atlandi += 1

            tah_gerek = key not in tah_local
            if tah_gerek:
                tah_marker = f"|AYLIK_TAH|{key}|"
                aciklama = f"{ay_adi} {yil_a} Tahsilat {harf}{tah_marker}{tag_suffix}"
                tahsil_plan.append(
                    {
                        "musteri_id": mid,
                        "ay_bir": ay_bir,
                        "tutar": toplam,
                        "aciklama": aciklama,
                        "tahsilat_tarihi": ay_bir,
                    }
                )
                tah_local.add(key)
                tah_keys_by_mid[mid].add(key)
                touched_mids.add(mid)
            else:
                tahsil_atlandi += 1

    borc_eklenen = len(insert_rows)

    if dry_run:
        return jsonify({
            "ok": True,
            "dry_run": True,
            "run_tag": run_tag,
            "hedef": f"{hedef_yil}-{hedef_ay:02d}",
            "musteri_islenen": musteri_islenen,
            "musteri_atlanan": musteri_atlanan,
            "borc_eklenecek": borc_eklenen,
            "borc_atlanan_satir": borc_atlandi,
            "tahsil_eklenecek": len(tahsil_plan),
            "tahsil_atlanan_satir": tahsil_atlandi,
            "guncel_kira_guncellenen": guncel_updated if guncel_kira_guncelle else 0,
            "mesaj": "Dry-run: kayıt yazılmadı.",
        })

    conn = get_conn()
    try:
        cur = conn.cursor()
        if customer_updates and guncel_kira_guncelle:
            cur.executemany(
                "UPDATE customers SET guncel_kira_bedeli = %s WHERE id = %s",
                customer_updates,
            )
        ins_sql = """
            INSERT INTO faturalar (
                fatura_no, musteri_id, musteri_adi, tutar, kdv_tutar,
                toplam, durum, fatura_tarihi, vade_tarihi, notlar
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """
        chunk = 400
        for i in range(0, len(insert_rows), chunk):
            cur.executemany(ins_sql, insert_rows[i : i + chunk])

        mids_tah = sorted({t["musteri_id"] for t in tahsil_plan})
        fmap = {}
        if mids_tah:
            cur.execute(
                f"""
                SELECT DISTINCT ON (f.musteri_id, (DATE_TRUNC('month', f.fatura_tarihi::date))::date)
                    f.id, f.musteri_id, (DATE_TRUNC('month', f.fatura_tarihi::date))::date AS m
                FROM faturalar f
                WHERE f.musteri_id IN %s
                  AND {sql_expr_fatura_not_gib_taslak("f.notlar")}
                ORDER BY f.musteri_id, (DATE_TRUNC('month', f.fatura_tarihi::date))::date, f.id DESC
                """,
                (tuple(mids_tah),),
            )
            for fid, fmid, fm in cur.fetchall() or []:
                fmap[(int(fmid), fm)] = int(fid)

        ins_tah = """
            INSERT INTO tahsilatlar (musteri_id, customer_id, fatura_id, tutar, odeme_turu, aciklama, tahsilat_tarihi, makbuz_no)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """
        tahsil_rows = []
        for t in tahsil_plan:
            mid = t["musteri_id"]
            ab = t["ay_bir"]
            makbuz_seq += 1
            makbuz_no = str(makbuz_seq)
            fid = fmap.get((mid, ab))
            tahsil_rows.append(
                (
                    mid,
                    mid,
                    fid,
                    t["tutar"],
                    odeme,
                    t["aciklama"],
                    t["tahsilat_tarihi"],
                    makbuz_no,
                )
            )
        for i in range(0, len(tahsil_rows), chunk):
            cur.executemany(ins_tah, tahsil_rows[i : i + chunk])

        conn.commit()
    except Exception as e:
        conn.rollback()
        logging.getLogger(__name__).exception("tufe_borclandir_nakit_tahsil_toplu")
        return jsonify({"ok": False, "mesaj": str(e)}), 500
    finally:
        conn.close()

    for mid_cache in sorted(touched_mids):
        try:
            _upsert_aylik_grid_cache(mid_cache, tufe_map=tufe_map)
        except Exception:
            pass

    tahsil_eklenen = len(tahsil_plan)
    return jsonify({
        "ok": True,
        "run_tag": run_tag,
        "hedef": f"{hedef_yil}-{hedef_ay:02d}",
        "musteri_islenen": musteri_islenen,
        "musteri_atlanan": musteri_atlanan,
        "borc_eklenen": borc_eklenen,
        "borc_atlanan_satir": borc_atlandi,
        "tahsil_eklenen": tahsil_eklenen,
        "tahsil_atlanan_satir": tahsil_atlandi,
        "guncel_kira_guncellenen": guncel_updated if guncel_kira_guncelle else 0,
        "mesaj": (
            f"Tamam: borç fatura {borc_eklenen}, tahsilat {tahsil_eklenen}, "
            f"işlenen müşteri {musteri_islenen}, atlanan müşteri {musteri_atlanan}."
        ),
    })


@bp.route('/api/tahsilat-panel-detay', methods=['GET', 'POST'])
@giris_gerekli
def api_tahsilat_panel_detay():
    """Panel DB: GET okuma, POST kayıt → grid önbelleği + ekstre kaynağı."""
    if request.method == "POST":
        try:
            data = request.get_json(force=True, silent=True) or {}
        except Exception:
            data = {}
        musteri_id = data.get("musteri_id")
        try:
            mid = int(musteri_id)
        except (TypeError, ValueError):
            return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
        if not fetch_one("SELECT id FROM customers WHERE id = %s", (mid,)):
            return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
        by_iso_in = data.get("by_iso")
        if not isinstance(by_iso_in, dict):
            return jsonify({"ok": False, "mesaj": "by_iso gerekli."}), 400
        _save_musteri_panel_by_iso(mid, by_iso_in)
        payload = _read_aylik_grid_cache_payload(mid)
        if payload is None:
            try:
                payload = _build_aylik_grid_cache_payload(
                    mid, tufe_map=_tufe_map_by_year_month_cached()
                )
            except Exception:
                payload = None
        if isinstance(payload, dict):
            _persist_grid_cache_with_panel(mid, payload)
        else:
            _upsert_aylik_grid_cache(mid)
        by_iso = _load_musteri_panel_by_iso(mid)
        return jsonify({"ok": True, "by_iso": by_iso, "musteri_id": mid})

    musteri_id = request.args.get("musteri_id", type=int)
    if not musteri_id:
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    mid = int(musteri_id)
    by_iso = _load_musteri_panel_by_iso(mid)
    if by_iso:
        return jsonify({"ok": True, "by_iso": by_iso, "musteri_id": mid})
    payload = _read_aylik_grid_cache_payload(mid)
    if payload is None:
        try:
            payload = _build_aylik_grid_cache_payload(
                mid, tufe_map=_tufe_map_by_year_month_cached()
            )
        except Exception:
            payload = None
    if isinstance(payload, dict):
        fresh = _panel_by_iso_from_tahsil_map(mid, payload, None)
        grid_src = _panel_by_iso_from_grid_payload(payload)
        tol_p = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
        for iso_g, prow_g in grid_src.items():
            fresh[iso_g] = _panel_by_iso_pick_richer(fresh.get(iso_g), prow_g, tol_p)
        if fresh:
            if not by_iso:
                _save_musteri_panel_by_iso(mid, fresh)
                _persist_grid_cache_with_panel(mid, payload)
                by_iso = _load_musteri_panel_by_iso(mid)
            else:
                merged = dict(by_iso)
                changed = False
                for iso, prow in fresh.items():
                    pt = round(float(prow.get("tahsil") or 0), 2)
                    pk = round(float(prow.get("kalan") or 0), 2)
                    if pt <= tol_p and pk <= tol_p:
                        continue
                    old = merged.get(iso) or {}
                    picked = _panel_by_iso_pick_richer(old, prow, tol_p)
                    ot = round(float(old.get("tahsil") or 0), 2)
                    ok = round(float(old.get("kalan") or 0), 2)
                    pt2 = round(float(picked.get("tahsil") or 0), 2)
                    pk2 = round(float(picked.get("kalan") or 0), 2)
                    if (
                        abs(ot - pt2) > tol_p
                        or abs(ok - pk2) > tol_p
                        or iso not in merged
                    ):
                        merged[iso] = picked
                        changed = True
                if changed:
                    _save_musteri_panel_by_iso(mid, merged)
                    by_iso = merged
        grid_resp = _panel_by_iso_from_grid_payload(payload)
        if grid_resp:
            tol_r = float(AYLIK_GRID_TAM_ODENDI_TOLERANS)
            out_r = dict(by_iso)
            for iso_r, prow_r in grid_resp.items():
                pt_r = round(float(prow_r.get("tahsil") or 0), 2)
                pk_r = round(float(prow_r.get("kalan") or 0), 2)
                if pt_r <= tol_r and pk_r <= tol_r:
                    continue
                out_r[iso_r] = _panel_by_iso_pick_richer(out_r.get(iso_r), prow_r, tol_r)
            by_iso = out_r
    return jsonify({"ok": True, "by_iso": by_iso, "musteri_id": mid})


@bp.route('/api/aylik-tahsil-durum')
@giris_gerekli
def api_aylik_tahsil_durum():
    """Aylık grid ile uyumlu: tam ödenmiş ay anahtarları (YYYY-M). Ödeme tarihi ≠ tam kapatma."""
    musteri_id = request.args.get("musteri_id", type=int)
    if not musteri_id:
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    mid = int(musteri_id)
    # Önce DB önbelleği (grid-cache ile aynı payload); yoksa tek seferlik hesap.
    payload = _read_aylik_grid_cache_payload(mid)
    if payload is None:
        payload = _build_aylik_grid_cache_payload(mid, tufe_map=_tufe_map_by_year_month_cached())
    else:
        payload = _aylik_grid_cache_payload_tahsil_guncelle(mid, payload)
    marker_ay_only = _aylik_tahsil_marker_aylar_set_normalized(mid)
    ekstre_ay_only = _aylik_tahsil_ekstre_eslesme_aylar_set_normalized(mid)
    ay_set = _aylik_tahsil_edilen_aylar_from_payload(payload) if payload else set()
    aylar = _aylik_tahsil_durum_finalize_ay_set(
        mid,
        ay_set,
        payload=payload,
        marker_ay_only=marker_ay_only,
        ekstre_ay_only=ekstre_ay_only,
    )
    resp = jsonify({"ok": True, "aylar": aylar})
    resp.headers["Cache-Control"] = "private, max-age=8"
    return resp


@bp.route('/api/cari-ekstre-b')
@giris_gerekli
def api_cari_ekstre_b():
    """
    Cari Ekstre B: borçta yalnızca ETTN’li (GİB) faturalar; alacakta yalnızca havale, EFT, çek, kredi kartı.
    İç fatura, nakit ve «banka» ödeme türü dahil değildir.
    """
    musteri_id = request.args.get("musteri_id", type=int)
    if not musteri_id:
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    cust = fetch_one("SELECT id, name FROM customers WHERE id = %s", (musteri_id,))
    if not cust:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    def_b, def_bit = _cari_ekstre_varsayilan_son_tam_ay()
    baslangic = request.args.get("baslangic")
    bitis = request.args.get("bitis")
    try:
        bas = datetime.strptime(baslangic[:10], "%Y-%m-%d").date() if baslangic else def_b
        bit = datetime.strptime(bitis[:10], "%Y-%m-%d").date() if bitis else def_bit
    except Exception:
        bas, bit = def_b, def_bit
    if bas > bit:
        bas, bit = bit, bas
    rows = _cari_hareketler(musteri_id, cari_ekstre_b=True)
    # Açılış bakiyesi: bas tarihinden önceki hareketlerin net tutarı
    acilis = 0.0
    filtered = []
    for r in rows:
        tarih_str = (r.get("tarih") or "")[:10]
        try:
            t = datetime.strptime(tarih_str, "%Y-%m-%d").date() if tarih_str else None
        except Exception:
            t = None
        borc = round(float(r.get("borc") or 0), 2)
        alacak = round(float(r.get("alacak") or 0), 2)
        if t is None:
            continue
        if t < bas:
            acilis += borc - alacak
            continue
        if t > bit:
            continue
        tur = r.get("tur") or ""
        belge_no = (r.get("belge_no") or "").strip()
        if tur == "Tahsilat":
            aciklama = (r.get("aciklama") or "").strip()
            if not aciklama:
                harf = _odeme_turu_harf(r.get("odeme_turu"))
                aciklama = "Tahsilat " + harf + (" " + belge_no if belge_no else "")
        else:
            aciklama = tur + (" " + belge_no if belge_no else "")
        fatura_id = None
        tahsilat_id = None
        if tur == "Fatura":
            rid = r.get("id")
            try:
                fatura_id = int(rid) if rid is not None else None
            except (TypeError, ValueError):
                fatura_id = None
        elif tur == "Tahsilat":
            rid = r.get("id")
            if isinstance(rid, str) and rid.startswith("t-"):
                try:
                    tahsilat_id = int(rid[2:])
                except ValueError:
                    tahsilat_id = None
        filtered.append({
            "tarih": tarih_str,
            "aciklama": aciklama.strip() or tur,
            "belge_no": r.get("belge_no") or "",
            "tur": r.get("tur") or "",
            "borc": borc,
            "alacak": alacak,
            "fatura_id": fatura_id,
            "tahsilat_id": tahsilat_id,
        })
    # Aynı gün/ay birden çok tahsilatı tek satıra indir.
    tahsilat_ids_by_date = {}
    for h in filtered:
        if (h.get("tur") or "") != "Tahsilat":
            continue
        ts = (h.get("tarih") or "")[:10]
        tid = h.get("tahsilat_id")
        if ts and tid:
            tahsilat_ids_by_date.setdefault(ts, []).append(tid)
    for ts in list(tahsilat_ids_by_date.keys()):
        lst = tahsilat_ids_by_date[ts]
        seen = set()
        uniq = []
        for x in lst:
            if x not in seen:
                seen.add(x)
                uniq.append(x)
        tahsilat_ids_by_date[ts] = uniq
    borc_by_tarih = {}
    tahsilat_by_tarih = {}
    for h in filtered:
        ts = (h.get("tarih") or "")[:10]
        if not ts:
            continue
        if (h.get("tur") or "") == "Tahsilat":
            tahsilat_by_tarih[ts] = round(float(tahsilat_by_tarih.get(ts, 0) or 0) + float(h.get("alacak") or 0), 2)
        else:
            borc_by_tarih[ts] = round(float(borc_by_tarih.get(ts, 0) or 0) + float(h.get("borc") or 0), 2)

    filtered_dedup = []
    seen_tah_tarih = set()
    for h in filtered:
        ts = (h.get("tarih") or "")[:10]
        if (h.get("tur") or "") != "Tahsilat":
            filtered_dedup.append(h)
            continue
        if ts in seen_tah_tarih:
            continue
        seen_tah_tarih.add(ts)
        hedef = borc_by_tarih.get(ts, 0)
        if hedef > 0:
            h["alacak"] = round(hedef, 2)
        else:
            h["alacak"] = round(float(tahsilat_by_tarih.get(ts, h.get("alacak") or 0) or 0), 2)
        h["tahsilat_ids"] = tahsilat_ids_by_date.get(ts, [])
        h.pop("tahsilat_id", None)
        filtered_dedup.append(h)

    filtered = filtered_dedup
    filtered.sort(key=lambda x: (x["tarih"], x["tur"] == "Fatura" and 0 or 1))
    bakiye = acilis
    hareketler = []
    for h in filtered:
        bakiye = round(bakiye + (h["borc"] - h["alacak"]), 2)
        h["bakiye"] = bakiye
        hareketler.append(h)
    toplam_borc = sum(h.get("borc") or 0 for h in hareketler)
    toplam_alacak = sum(h.get("alacak") or 0 for h in hareketler)
    bakiye = round(acilis + toplam_borc - toplam_alacak, 2)
    return jsonify({
        "ok": True,
        "musteri_adi": cust.get("name") or "",
        "hareketler": hareketler,
        "toplam_borc": round(toplam_borc, 2),
        "toplam_alacak": round(toplam_alacak, 2),
        "bakiye": bakiye,
    })


def _serialize_row_dates(row):
    if not row:
        return row
    d = dict(row)
    for k, v in list(d.items()):
        if v is not None and hasattr(v, "strftime"):
            d[k] = v.strftime("%Y-%m-%d %H:%M:%S") if hasattr(v, "hour") else v.strftime("%Y-%m-%d")
    return d


@bp.route("/api/ekstre-tahsilat-guncelle", methods=["POST"])
@giris_gerekli
def api_ekstre_tahsilat_guncelle():
    """Cari ekstre: tahsilat tutarı / tarihi / ödeme türü düzeltme (|AYLIK_TAH| marker korunur)."""
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    try:
        musteri_id = int(data.get("musteri_id"))
        tid = int(data.get("id"))
        tutar = round(float(data.get("tutar")), 2)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id, id ve tutar gerekli."}), 400
    if tutar <= 0:
        return jsonify({"ok": False, "mesaj": "Tutar 0'dan büyük olmalı."}), 400
    raw_tarih = (data.get("tarih") or data.get("tahsilat_tarihi") or "").strip()[:10]
    try:
        tahsilat_tarihi = datetime.strptime(raw_tarih, "%Y-%m-%d").date() if raw_tarih else None
    except Exception:
        tahsilat_tarihi = None
    odeme = (data.get("odeme_turu") or data.get("odeme") or "havale").strip().lower()
    if odeme not in ("nakit", "havale", "eft", "banka", "kredi_karti", "cek"):
        odeme = "havale"
    row = fetch_one(
        """
        SELECT id, COALESCE(aciklama, '') AS aciklama, tahsilat_tarihi
        FROM tahsilatlar
        WHERE id = %s AND (musteri_id = %s OR customer_id = %s)
        """,
        (tid, musteri_id, musteri_id),
    )
    if not row:
        return jsonify({"ok": False, "mesaj": "Tahsilat bulunamadı."}), 404
    affected_iso = _iso_from_aylik_tah_marker(row.get("aciklama"))
    ac_old = str(row.get("aciklama") or "")
    marker_m = re.search(r"(\|AYLIK_TAH\|[0-9]{4}-[0-9]{2}-[0-9]{2}\|)", ac_old)
    marker = marker_m.group(1) if marker_m else ""
    if not tahsilat_tarihi:
        td = row.get("tahsilat_tarihi")
        if td and hasattr(td, "year"):
            tahsilat_tarihi = td if isinstance(td, date) else datetime.strptime(str(td)[:10], "%Y-%m-%d").date()
        else:
            tahsilat_tarihi = date.today()
    if marker:
        mm_iso = re.search(r"\|AYLIK_TAH\|([0-9]{4}-[0-9]{2}-[0-9]{2})\|", marker)
        ay_iso = mm_iso.group(1) if mm_iso else date(tahsilat_tarihi.year, tahsilat_tarihi.month, 1).isoformat()
        try:
            dd = datetime.strptime(ay_iso[:10], "%Y-%m-%d").date()
            ay_adi = _AY_ADLARI[dd.month - 1]
            yil = dd.year
        except Exception:
            ay_adi = _AY_ADLARI[tahsilat_tarihi.month - 1]
            yil = tahsilat_tarihi.year
            ay_iso = date(yil, tahsilat_tarihi.month, 1).isoformat()
            marker = f"|AYLIK_TAH|{ay_iso}|"
        harf = _odeme_turu_harf(odeme)
        aciklama = f"{ay_adi} {yil} Tahsilat {harf}{marker}"
    else:
        aciklama = (data.get("aciklama") or ac_old or "Tahsilat").strip() or "Tahsilat"
    execute(
        """
        UPDATE tahsilatlar
        SET tutar = %s, tahsilat_tarihi = %s, odeme_turu = %s, aciklama = %s
        WHERE id = %s AND (musteri_id = %s OR customer_id = %s)
        """,
        (tutar, tahsilat_tarihi, odeme, aciklama, tid, musteri_id, musteri_id),
    )
    aff_upd = _iso_from_aylik_tah_marker(aciklama) or affected_iso
    _ekstre_invalidate_after_change(
        musteri_id, [aff_upd] if aff_upd else None
    )
    return jsonify({"ok": True, "id": tid, "mesaj": "Tahsilat güncellendi."})


@bp.route("/api/ekstre-tahsilat-sil", methods=["POST"])
@giris_gerekli
def api_ekstre_tahsilat_sil():
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    try:
        musteri_id = int(data.get("musteri_id"))
        tid = int(data.get("id"))
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id ve id gerekli."}), 400
    row_pre = fetch_one(
        """
        SELECT COALESCE(aciklama, '') AS aciklama
        FROM tahsilatlar
        WHERE id = %s AND (musteri_id = %s OR customer_id = %s)
        """,
        (tid, musteri_id, musteri_id),
    )
    affected_iso = (
        _iso_from_aylik_tah_marker(row_pre.get("aciklama")) if row_pre else None
    )
    row = execute_returning(
        """
        DELETE FROM tahsilatlar
        WHERE id = %s AND (musteri_id = %s OR customer_id = %s)
        RETURNING id
        """,
        (tid, musteri_id, musteri_id),
    )
    if not row:
        return jsonify({"ok": False, "mesaj": "Tahsilat bulunamadı."}), 404
    _ekstre_invalidate_after_change(
        musteri_id, [affected_iso] if affected_iso else None
    )
    return jsonify({"ok": True, "mesaj": "Tahsilat silindi."})


@bp.route("/api/ekstre-fatura-guncelle", methods=["POST"])
@giris_gerekli
def api_ekstre_fatura_guncelle():
    """Cari ekstre: aylık kira / borç faturası tutarı ve tarihi (KDV dahil toplam)."""
    ensure_faturalar_amount_columns()
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    try:
        musteri_id = int(data.get("musteri_id"))
        fid = int(data.get("id"))
        toplam = round(float(data.get("tutar") or data.get("tutar_kdv_dahil")), 2)
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id, id ve tutar gerekli."}), 400
    if toplam <= 0:
        return jsonify({"ok": False, "mesaj": "Tutar 0'dan büyük olmalı."}), 400
    raw_tarih = (data.get("tarih") or data.get("fatura_tarihi") or "").strip()[:10]
    try:
        fatura_tarihi = datetime.strptime(raw_tarih, "%Y-%m-%d").date() if raw_tarih else None
    except Exception:
        fatura_tarihi = None
    fr = fetch_one(
        "SELECT id, fatura_tarihi, COALESCE(notlar, '') AS notlar FROM faturalar WHERE id = %s AND musteri_id = %s",
        (fid, musteri_id),
    )
    if not fr:
        return jsonify({"ok": False, "mesaj": "Fatura bulunamadı."}), 404
    if "|AYLIK_TUTAR|" not in str(fr.get("notlar") or ""):
        return jsonify({"ok": False, "mesaj": "Yalnızca Aylık Tutarlar ile oluşturulan faturalar düzenlenebilir."}), 400
    if not fatura_tarihi:
        ft = fr.get("fatura_tarihi")
        if ft and hasattr(ft, "year"):
            fatura_tarihi = ft if isinstance(ft, date) else datetime.strptime(str(ft)[:10], "%Y-%m-%d").date()
        else:
            fatura_tarihi = date.today()
    kn_row = fetch_one(
        "SELECT kira_nakit, kdv_oran FROM musteri_kyc WHERE musteri_id = %s ORDER BY id DESC LIMIT 1",
        (musteri_id,),
    ) or {}
    kira_nakit = bool(kn_row.get("kira_nakit"))
    try:
        kdv_oran = float(kn_row.get("kdv_oran") or 20)
    except (TypeError, ValueError):
        kdv_oran = 20.0
    net, kdv_tutar, toplam_r = _fatura_tutar_kdv_split(toplam, kira_nakit, kdv_oran)
    if fatura_tarihi.month == 12:
        vade = date(fatura_tarihi.year, 12, 31)
    else:
        vade = date(fatura_tarihi.year, fatura_tarihi.month + 1, 1) - timedelta(days=1)
    execute(
        """
        UPDATE faturalar
        SET tutar = %s, kdv_tutar = %s, toplam = %s, fatura_tarihi = %s, vade_tarihi = %s
        WHERE id = %s AND musteri_id = %s
        """,
        (net, kdv_tutar, toplam_r, fatura_tarihi, vade, fid, musteri_id),
    )
    _ekstre_invalidate_after_change(musteri_id)
    return jsonify({"ok": True, "id": fid, "mesaj": "Borç faturası güncellendi."})


@bp.route("/api/ekstre-fatura-sil", methods=["POST"])
@giris_gerekli
def api_ekstre_fatura_sil():
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}
    try:
        musteri_id = int(data.get("musteri_id"))
        fid = int(data.get("id"))
    except (TypeError, ValueError):
        return jsonify({"ok": False, "mesaj": "musteri_id ve id gerekli."}), 400
    fr = fetch_one(
        "SELECT id FROM faturalar WHERE id = %s AND musteri_id = %s AND COALESCE(notlar, '') LIKE '%%|AYLIK_TUTAR|%%'",
        (fid, musteri_id),
    )
    if not fr:
        return jsonify({"ok": False, "mesaj": "Fatura bulunamadı veya silinemez."}), 404
    execute("DELETE FROM tahsilatlar WHERE fatura_id = %s", (fid,))
    execute("DELETE FROM faturalar WHERE id = %s AND musteri_id = %s", (fid, musteri_id))
    _ekstre_invalidate_after_change(musteri_id)
    return jsonify({"ok": True, "mesaj": "Borç faturası silindi."})


@bp.route("/api/tahsilat-detay")
@giris_gerekli
def api_tahsilat_detay():
    """Cari ekstre tıklaması: müşteriye ait tahsilat kayıtlarının özeti (makbuz PDF ile eşleşir)."""
    musteri_id = request.args.get("musteri_id", type=int)
    ids_raw = (request.args.get("ids") or "").strip()
    if not musteri_id:
        return jsonify({"ok": False, "mesaj": "musteri_id gerekli."}), 400
    ids = []
    for part in ids_raw.split(","):
        part = part.strip()
        if not part:
            continue
        try:
            ids.append(int(part))
        except ValueError:
            continue
    if not ids:
        return jsonify({"ok": False, "mesaj": "Geçerli tahsilat id gerekli."}), 400
    if len(ids) > 40:
        return jsonify({"ok": False, "mesaj": "En fazla 40 kayıt sorgulanabilir."}), 400
    ph = ",".join(["%s"] * len(ids))
    sql = f"""
        SELECT t.id, t.makbuz_no, t.tutar, t.odeme_turu, t.tahsilat_tarihi, t.aciklama,
               t.fatura_id, f.fatura_no, t.cek_detay, t.havale_banka
        FROM tahsilatlar t
        LEFT JOIN faturalar f ON f.id = t.fatura_id
        WHERE t.id IN ({ph})
          AND (t.musteri_id = %s OR t.customer_id = %s)
    """
    rows = fetch_all(sql, tuple(ids) + (musteri_id, musteri_id)) or []
    by_id = {r["id"]: r for r in rows}
    ordered = [by_id[i] for i in ids if i in by_id]
    return jsonify({"ok": True, "kayitlar": [_serialize_row_dates(r) for r in ordered]})


@bp.route('/api/cari-kart-pdf/<int:mid>')
@giris_gerekli
def api_cari_kart_pdf(mid):
    """Cari hareketleri BestOffice antetli PDF ekstre olarak indir."""
    cust = fetch_one("SELECT id, name, tax_number FROM customers WHERE id = %s", (mid,))
    if not cust:
        return jsonify({"ok": False, "mesaj": "Müşteri bulunamadı."}), 404
    hareketler = _cari_hareketler(mid)
    _register_arial()
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4
    y = h - 40
    try:
        c.setFont("Arial", 16)
    except Exception:
        c.setFont("Helvetica", 16)
    c.drawString(40, y, "BestOffice - Cari Ekstre")
    y -= 24
    c.setFont("Helvetica", 10)
    c.drawString(40, y, "Müşteri: " + (cust.get("name") or ""))
    c.drawString(40, y - 14, "Vergi No: " + (cust.get("tax_number") or ""))
    y -= 40
    c.drawString(40, y, "Tarih")
    c.drawString(120, y, "Belge No")
    c.drawString(220, y, "Tür")
    c.drawString(300, y, "Borç")
    c.drawString(380, y, "Alacak")
    c.drawString(460, y, "Bakiye")
    y -= 6
    c.line(40, y, 520, y)
    y -= 14
    for row in hareketler:
        if y < 80:
            c.showPage()
            y = h - 40
        c.drawString(40, y, (row.get("tarih") or "")[:10])
        c.drawString(120, y, (row.get("belge_no") or "")[:18])
        c.drawString(220, y, row.get("tur") or "")
        c.drawString(300, y, "{:,.2f}".format(row.get("borc") or 0))
        c.drawString(380, y, "{:,.2f}".format(row.get("alacak") or 0))
        c.drawString(460, y, "{:,.2f}".format(row.get("bakiye") or 0))
        y -= 14
    c.save()
    buf.seek(0)
    return Response(buf.read(), mimetype="application/pdf", headers={
        "Content-Disposition": "attachment; filename=Cari_Ekstre_%s.pdf" % (cust.get("name") or "musteri").replace(" ", "_")[:30]
    })


# ── Sözleşme / Taksit API ─────────────────────────────────────────────────────


@bp.route('/api/contracts/<int:mid>', methods=['GET', 'POST'])
@giris_gerekli
def api_contracts(mid):
    """Belirli bir müşteri için sözleşme ve taksit özetleri."""
    if request.method == 'GET':
        bugun = date.today()
        contracts = fetch_all(
            """SELECT * FROM contracts WHERE musteri_id=%s ORDER BY id DESC""",
            (mid,),
        )
        plan = fetch_all(
            """SELECT * FROM contract_installments
               WHERE musteri_id=%s
               ORDER BY vade_tarihi, taksit_no""",
            (mid,),
        )
        ozet = fetch_one(
            """
            SELECT
              COALESCE(SUM(tutar),0)                         AS planlanan,
              COALESCE(SUM(CASE WHEN vade_tarihi <= %s THEN tutar END),0) AS tahakkuk,
              COALESCE(SUM(odenen_tutar),0)                  AS odenen,
              COALESCE(SUM(CASE WHEN odeme_durumu IN ('gecikmis','icrada') THEN kalan_tutar ELSE 0 END),0) AS geciken,
              COALESCE(SUM(CASE WHEN vade_tarihi > %s THEN kalan_tutar ELSE 0 END),0) AS gelecek
            FROM contract_installments
            WHERE musteri_id = %s
            """,
            (bugun, bugun, mid),
        ) or {}
        return jsonify({
            "ok": True,
            "contracts": contracts or [],
            "installments": plan or [],
            "ozet": {
                "planlanan": float(ozet.get("planlanan") or 0),
                "tahakkuk": float(ozet.get("tahakkuk") or 0),
                "odenen": float(ozet.get("odenen") or 0),
                "geciken": float(ozet.get("geciken") or 0),
                "gelecek": float(ozet.get("gelecek") or 0),
            },
        })

    # POST: yeni sözleşme oluştur / güncelle
    data = request.get_json() or {}
    cid = data.get("id")
    baslangic = _parse_date_str(data.get("baslangic_tarihi"))
    bitis = _parse_date_str(data.get("bitis_tarihi"))
    if not baslangic:
        return jsonify({"ok": False, "mesaj": "Sözleşme başlangıç tarihi zorunlu."}), 400
    try:
        aylik_kira = float(data.get("aylik_kira") or 0)
    except Exception:
        aylik_kira = 0
    if aylik_kira <= 0:
        return jsonify({"ok": False, "mesaj": "Aylık kira tutarı zorunlu."}), 400
    sure_ay = data.get("sure_ay")
    try:
        sure_ay = int(sure_ay) if sure_ay is not None else None
    except Exception:
        sure_ay = None
    try:
        odeme_gunu = int(data.get("odeme_gunu") or 0) or None
    except Exception:
        odeme_gunu = None
    para_birimi = (data.get("para_birimi") or "TRY").strip().upper()
    depozito = data.get("depozito") or 0
    try:
        depozito = float(depozito or 0)
    except Exception:
        depozito = 0
    try:
        gecikme = float(data.get("gecikme_faizi_orani") or 0)
    except Exception:
        gecikme = 0
    try:
        artis = float(data.get("yillik_artis_orani") or 0)
    except Exception:
        artis = 0
    muacceliyet = bool(data.get("muacceliyet_var")) or str(data.get("muacceliyet_var")).lower() in ("1", "true", "evet", "on")
    durum = (data.get("durum") or "aktif").strip().lower()
    sozlesme_no = (data.get("sozlesme_no") or "").strip() or None
    cari_kodu = (data.get("cari_kodu") or "").strip() or None
    toplam_tutar = data.get("toplam_tutar")
    try:
        toplam_tutar = float(toplam_tutar or 0)
    except Exception:
        toplam_tutar = 0
    if not toplam_tutar and sure_ay:
        toplam_tutar = aylik_kira * sure_ay

    if cid:
        execute(
            """
            UPDATE contracts
               SET cari_kodu=%s, sozlesme_no=%s, baslangic_tarihi=%s, bitis_tarihi=%s,
                   sure_ay=%s, aylik_kira=%s, toplam_tutar=%s, para_birimi=%s,
                   odeme_gunu=%s, depozito=%s, gecikme_faizi_orani=%s,
                   yillik_artis_orani=%s, muacceliyet_var=%s, durum=%s,
                   updated_at=NOW()
             WHERE id=%s AND musteri_id=%s
            """,
            (
                cari_kodu, sozlesme_no, baslangic, bitis,
                sure_ay, aylik_kira, toplam_tutar, para_birimi,
                odeme_gunu, depozito, gecikme, artis,
                muacceliyet, durum, cid, mid,
            ),
        )
        contract_id = int(cid)
    else:
        row = execute_returning(
            """
            INSERT INTO contracts
                (musteri_id, cari_kodu, sozlesme_no, baslangic_tarihi, bitis_tarihi,
                 sure_ay, aylik_kira, toplam_tutar, para_birimi,
                 odeme_gunu, depozito, gecikme_faizi_orani,
                 yillik_artis_orani, muacceliyet_var, durum)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            RETURNING id
            """,
            (
                mid, cari_kodu, sozlesme_no, baslangic, bitis,
                sure_ay, aylik_kira, toplam_tutar, para_birimi,
                odeme_gunu, depozito, gecikme, artis,
                muacceliyet, durum,
            ),
        )
        contract_id = row["id"]

    # Taksit planını üret
    _generate_installments(contract_id, mid, baslangic, bitis, sure_ay, aylik_kira, odeme_gunu)

    return jsonify({"ok": True, "id": contract_id, "mesaj": "Sözleşme ve taksit planı kaydedildi."})